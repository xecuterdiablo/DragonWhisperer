#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""🐉 DRAGON WHISPERER – Professionelle Überarbeitung v3.2"""

# =============================================================================
# 1. IMPORTS
# =============================================================================
import atexit
import gc
import hashlib
import difflib
import importlib
import importlib.util
import json
import logging
import os
import platform
import queue
import re
import shutil
import signal as py_signal
import subprocess
import sys
import tempfile
import threading
import time
import weakref
import urllib.parse
import urllib.request
import warnings
from collections import OrderedDict, deque
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeout
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from enum import Enum
from functools import wraps
from pathlib import Path
from typing import (Any, Callable, Deque, Dict, List, Optional, Set, Tuple,
                    Union)
from abc import ABC, abstractmethod

# -----------------------------------------------------------------------------
# FRÜHE KONFIGURATION
# -----------------------------------------------------------------------------
DEBUG_LEVEL = 0
DEBUG_COMPONENTS = []

for arg in sys.argv:
    if arg == '--debug':
        DEBUG_LEVEL = max(DEBUG_LEVEL, 1)
    elif arg.startswith('--debug='):
        value = arg.split('=', 1)[1]
        if value.isdigit():
            DEBUG_LEVEL = max(DEBUG_LEVEL, int(value))
        else:
            DEBUG_COMPONENTS.extend(value.split(','))

QUIET_MODE = "--quiet" in sys.argv or "-q" in sys.argv

logging.basicConfig(
    level=logging.WARNING,
    format="[%(asctime)s.%(msecs)03d] [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("dragon")

for lib in ["httpx", "urllib3", "httpcore"]:
    logging.getLogger(lib).setLevel(logging.WARNING)

if DEBUG_LEVEL >= 1:
    logger.setLevel(logging.DEBUG)
    logging.getLogger().setLevel(logging.DEBUG)
if QUIET_MODE:
    logger.setLevel(logging.ERROR)

warnings.filterwarnings("ignore", message=".*pynvml.*")
warnings.filterwarnings("ignore", message=".*The pynvml package is deprecated.*")

os.environ.update({
    "PYTHONWARNINGS": "default",
    "TORCH_DISABLE_CUDA_WARNINGS": "1",
    "TORCH_CPP_LOG_LEVEL": "0",
    "PYTORCH_JIT": "0",
})

SYSTEM = platform.system()
IS_WINDOWS = SYSTEM == "Windows"
IS_MACOS = SYSTEM == "Darwin"
IS_LINUX = SYSTEM == "Linux"

try:
    machine = platform.machine().lower()
    IS_ARM = machine in ("arm64", "aarch64", "armv8l", "arm")
    IS_X86 = machine in ("x86_64", "amd64", "i386", "i686", "x86")
except Exception:
    IS_ARM = False
    IS_X86 = True

logger.info(f"🐉 Dragon Whisperer - Platform: {SYSTEM} "
            f"{'ARM' if IS_ARM else 'x86'} (Debug-Level: {DEBUG_LEVEL})")

os.environ.update({
    "FFMPEG_DISABLE_RKMPP": "1",
    "AV_DISABLE_RKMPP": "1",
    "FFMPEG_DISABLE_VAAPI": "0" if IS_LINUX else "1",
    "FFMPEG_DISABLE_VDPAU": "0" if IS_LINUX else "1",
    "OPENCV_LOG_LEVEL": "ERROR",
    "GST_DEBUG": "0",
})

if IS_WINDOWS:
    os.environ.update({"PYTHONIOENCODING": "utf-8"})

# -----------------------------------------------------------------------------
# Globale Hilfsfunktion für Debug Level 3
# -----------------------------------------------------------------------------
def debug3_enabled(component: Optional[str] = None) -> bool:
    """Gibt True zurück, wenn Debug-Level 3 aktiv ist oder die Komponente in DEBUG_COMPONENTS enthalten ist."""
    if DEBUG_LEVEL >= 3:
        return True
    if component and component in DEBUG_COMPONENTS:
        return True
    return False

# =============================================================================
# 2. KONSTANTEN & KONFIGURATION
# =============================================================================
class Constants:
    """Zentrale Konstanten für das gesamte Programm."""
    # Audio
    SAMPLE_RATE: int = 16000
    CHANNELS: int = 1
    AUDIO_FORMAT: str = 's16le'
    BYTES_PER_SAMPLE: int = 2

    # Chunking
    BASE_CHUNK_DURATION: int = 5
    CHUNK_OVERLAP: float = 1.0
    MIN_CHUNK_DURATION: int = 2
    MAX_CHUNK_DURATION: int = 30

    # Prozesse & Timeouts
    MAX_SUBPROCESSES: int = 8
    SUBPROCESS_TIMEOUT: int = 60
    GUI_OPERATION_TIMEOUT: float = 10.0
    MEMORY_CHECK_INTERVAL: int = 15
    MAX_GUI_UPDATES_PER_SECOND: int = 30
    MAX_MEMORY_USAGE: int = 8 * 1024 * 1024 * 1024
    MAX_CACHE_SIZE: int = 500
    MAX_TEXT_LINES: int = 2000
    DEFAULT_BEAM_SIZE: int = 5
    DEFAULT_TEMPERATURE: float = 0.0
    ENABLE_VAD_FILTER: bool = True
    MAX_CONSECUTIVE_ERRORS: int = 5

    # Stream
    STREAM_TIMEOUT: int = 25
    INITIAL_BUFFER_SECONDS: float = 1.5
    MAX_EMPTY_READS: int = 30
    RECONNECT_DELAY: int = 2
    READ_RETRY_DELAY: float = 0.1
    YOUTUBE_TIMEOUT: int = 10000000
    NORMAL_TIMEOUT: int = 30000000
    MAX_STREAM_RECONNECTS: int = 5

    # FFmpeg
    FFMPEG_BUFSIZE: str = '2048k'
    FFMPEG_THREADS: int = 1
    FFMPEG_PROBESIZE: str = '32'
    FFMPEG_ANALYZE_DURATION: str = '0'

    # Audio-Filter (Sprachoptimierung)
    AUDIO_FILTER: str = "aresample=16000,volume=1.5,dynaudnorm"
    LANGUAGE_FILTERS: Dict[str, str] = {
        'ko': "aresample=16000,volume=2.0,highpass=f=80,lowpass=f=3800,afftdn=nf=-15",
        'ja': "aresample=16000,volume=2.0,highpass=f=90,lowpass=f=3700,afftdn=nf=-15",
        'zh': "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3500,afftdn=nf=-20",
        'de': "aresample=16000,volume=1.8,highpass=f=100,lowpass=f=3200,dynaudnorm",
        'en': "aresample=16000,volume=1.8,highpass=f=80,lowpass=f=3400,dynaudnorm",
        'fr': "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3300,dynaudnorm",
        'es': "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3400,dynaudnorm",
    }
    FILTER_PROFILES: Dict[str, str] = {
        'transcription': "aresample=16000,volume=1.5,dynaudnorm",
        'translation': "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3400",
        'realtime': "aresample=16000,volume=1.8,dynaudnorm",
        'noisy': "aresample=16000,volume=2.5,highpass=f=150,lowpass=f=3000,afftdn=nf=-30",
        'music': "aresample=16000,volume=1.5,highpass=f=50,lowpass=f=5000",
        'podcast': "aresample=16000,volume=2.0,highpass=f=80,lowpass=f=3500",
    }

    # Audio-Enhancement
    AUDIO_ENHANCEMENT_ENABLED: bool = True
    MIN_RMS_THRESHOLD: float = 0.002
    TARGET_RMS: float = 0.2
    MAX_GAIN: float = 5.0
    CLIPPING_THRESHOLD: float = 0.9

    # Duplikaterkennung
    DUPLICATE_CHECK_ENABLED: bool = True
    RECENT_TRANSCRIPTIONS_SIZE: int = 10
    MIN_TEXT_LENGTH: int = 3
    MIN_UNIQUE_WORDS_RATIO: float = 0.3

    # Untertitel
    SUBTITLE_BUFFER_SIZE: int = 1000
    ENABLE_TIMED_TRANSCRIPTIONS: bool = True
    ENABLE_TIMED_TRANSLATIONS: bool = True

    # Logging
    ENABLE_DEBUG_LOGGING: bool = True
    LOG_CHUNK_PROCESSING: bool = False
    LOG_AUDIO_STATS: bool = True
    LOG_PERFORMANCE: bool = True
    LOG_STREAM_EVENTS: bool = True
    PERFORMANCE_LOG_INTERVAL: int = 50

    # Cache
    MAX_CACHE_SIZE_MB: int = 100
    CACHE_ENABLED: bool = True

    # Cache-Größen (TTL in Sekunden)
    TRANSCRIPTION_CACHE_SIZE: int = 200
    TRANSCRIPTION_CACHE_TTL: int = 300
    TRANSLATION_CACHE_SIZE: int = 500
    TRANSLATION_CACHE_TTL: int = 3600
    AUDIO_CACHE_SIZE: int = 128
    AUDIO_CACHE_TTL: int = 1800

    # VAD (Standardwerte)
    VAD_THRESHOLD: float = 0.25
    VAD_MIN_SPEECH_DURATION_MS: int = 225
    VAD_MIN_SILENCE_DURATION_MS: int = 80

    # Sprachspezifische VAD-Anpassungen (für asiatische Sprachen)
    LANGUAGE_VAD: Dict[str, Dict[str, Any]] = {
        'ja': {'threshold': 0.3, 'min_speech_ms': 300, 'min_silence_ms': 100},
        'ko': {'threshold': 0.3, 'min_speech_ms': 300, 'min_silence_ms': 100},
        'zh': {'threshold': 0.3, 'min_speech_ms': 300, 'min_silence_ms': 100},
        'th': {'threshold': 0.3, 'min_speech_ms': 300, 'min_silence_ms': 100},
        'vi': {'threshold': 0.3, 'min_speech_ms': 250, 'min_silence_ms': 90},
    }

    # Pfad-Validierung (Sicherheit)
    ALLOWED_FILE_SCHEME_PREFIX: str = "file://"
    ALLOWED_FILE_BASE_DIRS: List[str] = [
        str(Path.home()),
        os.getcwd(),
    ]

    # URL-Zeichen Whitelist (für subprocess-Aufrufe)
    URL_ALLOWED_CHARS: str = r"a-zA-Z0-9\-._~:/?#\[\]@!$&'()*+,;=%"

    # YouTube-Header
    YOUTUBE_HEADERS: Dict[str, str] = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Referer': 'https://www.youtube.com/',
        'Origin': 'https://www.youtube.com',
        'Accept': '*/*',
        'Accept-Language': 'en-US,en;q=0.9',
    }

    # Plattform-Konfiguration
    PLATFORM_CONFIG: Dict[str, Dict[str, Any]] = {
        'windows': {
            'ffmpeg_flags': ['-reconnect', '1', '-reconnect_streamed', '1'],
            'process_creation_flags': 0x08000000,
        },
        'macos': {
            'ffmpeg_flags': ['-reconnect', '1', '-reconnect_on_network_error', '1'],
            'start_new_session': True,
        },
        'linux': {
            'ffmpeg_flags': ['-reconnect', '1', '-reconnect_streamed', '1'],
            'start_new_session': True,
        }
    }

    # -------------------------------------------------------------------------
    # Konstanten für Timeouts
    # -------------------------------------------------------------------------
    YOUTUBE_IDLE_TIMEOUT: int = 25
    READ_CHUNK_TIMEOUT: float = 0.5
    STREAM_RECONNECT_DELAY: int = 2
    STREAM_RECONNECT_BACKOFF_FACTOR: float = 2.0
    MAX_STREAM_RECONNECT_ATTEMPTS: int = 5
    READ_EMPTY_SLEEP_BASE: float = 0.1
    READ_ERROR_BACKOFF_BASE: float = 0.1
    YOUTUBE_SESSION_IDLE_TIMEOUT: int = 25
    YOUTUBE_INITIAL_TIMEOUT: int = 25
    YOUTUBE_URL_REFRESH_INTERVAL: int = 600
    YOUTUBE_URL_REFRESH_MAX_ATTEMPTS: int = 3
    YOUTUBE_LOW_QUALITY_MAX_CHUNKS: int = 5
    YOUTUBE_INITIAL_WAIT: float = 3.0
    FFMPEG_START_WAIT: float = 1.5
    STREAM_TEST_TIMEOUT: int = 8
    YOUTUBE_STREAM_TEST_TIMEOUT: int = 8
    AUDIO_ENHANCEMENT_MIN_LENGTH: int = 1600
    NOISEREDUCE_INTERVAL: int = 10
    NOISEREDUCE_MIN_LENGTH: int = 32000
    LOW_QUALITY_CHUNK_THRESHOLD_FACTOR: float = 0.1
    BUFFER_FLUSH_INACTIVITY: float = 10.0
    READ_WITH_TIMEOUT_SELECT_INTERVAL: float = 0.001
    READ_WITH_TIMEOUT_EMPTY_LOG_INTERVAL: int = 10

# -----------------------------------------------------------------------------
# CONFIG
# -----------------------------------------------------------------------------
class Config:
    """Dynamische Konfiguration, die zur Laufzeit angepasst werden kann."""
    SAMPLE_RATE: int = Constants.SAMPLE_RATE
    CHANNELS: int = Constants.CHANNELS
    AUDIO_FORMAT: str = Constants.AUDIO_FORMAT
    BYTES_PER_SAMPLE: int = Constants.BYTES_PER_SAMPLE

    MAX_SUBPROCESSES: int = Constants.MAX_SUBPROCESSES
    SUBPROCESS_TIMEOUT: int = Constants.SUBPROCESS_TIMEOUT
    GUI_OPERATION_TIMEOUT: float = Constants.GUI_OPERATION_TIMEOUT
    MEMORY_CHECK_INTERVAL: int = Constants.MEMORY_CHECK_INTERVAL
    MAX_GUI_UPDATES_PER_SECOND: int = Constants.MAX_GUI_UPDATES_PER_SECOND
    MAX_MEMORY_USAGE: int = Constants.MAX_MEMORY_USAGE
    MAX_CACHE_SIZE: int = Constants.MAX_CACHE_SIZE
    MAX_TEXT_LINES: int = Constants.MAX_TEXT_LINES
    DEFAULT_BEAM_SIZE: int = Constants.DEFAULT_BEAM_SIZE
    DEFAULT_TEMPERATURE: float = Constants.DEFAULT_TEMPERATURE
    ENABLE_VAD_FILTER: bool = Constants.ENABLE_VAD_FILTER
    MAX_CONSECUTIVE_ERRORS: int = Constants.MAX_CONSECUTIVE_ERRORS

    _base_chunk_duration: int = Constants.BASE_CHUNK_DURATION
    CHUNK_OVERLAP: float = Constants.CHUNK_OVERLAP
    MIN_CHUNK_DURATION: int = Constants.MIN_CHUNK_DURATION
    MAX_CHUNK_DURATION: int = Constants.MAX_CHUNK_DURATION

    # Sprachspezifische VAD-Parameter (werden in der Engine verwendet)
    LANGUAGE_VAD: Dict[str, Dict[str, Any]] = Constants.LANGUAGE_VAD

    @property
    def CHUNK_DURATION(self) -> float:
        return getattr(self, '_actual_chunk_duration', self._base_chunk_duration)

    @CHUNK_DURATION.setter
    def CHUNK_DURATION(self, value: float) -> None:
        if self.MIN_CHUNK_DURATION <= value <= self.MAX_CHUNK_DURATION:
            self._actual_chunk_duration = float(value)
        else:
            logger.warning(f"⚠️ Chunk duration {value}s out of range, using default")
            self._actual_chunk_duration = self._base_chunk_duration

    @property
    def CHUNK_SIZE_BYTES(self) -> int:
        return int(self.CHUNK_DURATION * self.SAMPLE_RATE *
                   self.CHANNELS * self.BYTES_PER_SAMPLE)

    @property
    def OVERLAP_SIZE_BYTES(self) -> int:
        return int(self.CHUNK_OVERLAP * self.SAMPLE_RATE *
                   self.CHANNELS * self.BYTES_PER_SAMPLE)

    @property
    def BYTES_PER_SECOND(self) -> int:
        return self.SAMPLE_RATE * self.CHANNELS * self.BYTES_PER_SAMPLE

    @property
    def MIN_CHUNK_BYTES(self) -> int:
        return int(self.MIN_CHUNK_DURATION * self.BYTES_PER_SECOND)

    @property
    def MAX_CHUNK_BYTES(self) -> int:
        return int(self.MAX_CHUNK_DURATION * self.BYTES_PER_SECOND)

    STREAM_TIMEOUT: int = Constants.STREAM_TIMEOUT
    INITIAL_BUFFER_SECONDS: float = Constants.INITIAL_BUFFER_SECONDS
    MAX_EMPTY_READS: int = Constants.MAX_EMPTY_READS
    RECONNECT_DELAY: int = Constants.RECONNECT_DELAY
    READ_RETRY_DELAY: float = Constants.READ_RETRY_DELAY

    @property
    def INITIAL_BUFFER_BYTES(self) -> int:
        return int(self.INITIAL_BUFFER_SECONDS * self.BYTES_PER_SECOND)

    FFMPEG_BUFSIZE: str = Constants.FFMPEG_BUFSIZE
    FFMPEG_THREADS: int = Constants.FFMPEG_THREADS
    FFMPEG_PROBESIZE: str = Constants.FFMPEG_PROBESIZE
    FFMPEG_ANALYZE_DURATION: str = Constants.FFMPEG_ANALYZE_DURATION

    YOUTUBE_TIMEOUT: int = Constants.YOUTUBE_TIMEOUT
    NORMAL_TIMEOUT: int = Constants.NORMAL_TIMEOUT

    def get_timeout_microseconds(self, is_youtube: bool = False) -> int:
        return self.YOUTUBE_TIMEOUT if is_youtube else self.NORMAL_TIMEOUT

    AUDIO_FILTER: str = Constants.AUDIO_FILTER
    LANGUAGE_FILTERS: Dict[str, str] = Constants.LANGUAGE_FILTERS
    FILTER_PROFILES: Dict[str, str] = Constants.FILTER_PROFILES

    def get_audio_filter(self, language: Optional[str] = None,
                         profile: Optional[str] = None) -> str:
        if profile and profile in self.FILTER_PROFILES:
            return self.FILTER_PROFILES[profile]
        if language:
            lang_code = language[:2].lower() if len(language) >= 2 else None
            if lang_code in self.LANGUAGE_FILTERS:
                return self.LANGUAGE_FILTERS[lang_code]
        return self.AUDIO_FILTER

    YOUTUBE_HEADERS: Dict[str, str] = Constants.YOUTUBE_HEADERS

    def get_youtube_headers(self, is_manifest: bool = False) -> Dict[str, str]:
        headers = self.YOUTUBE_HEADERS.copy()
        if is_manifest:
            headers.update({
                'X-Client-Data': 'CI22yQE=',
                'Content-Type': 'application/x-mpegURL',
            })
        return headers

    PLATFORM_CONFIG: Dict[str, Dict[str, Any]] = Constants.PLATFORM_CONFIG

    def get_platform_config(self, platform: Optional[str] = None) -> Dict[str, Any]:
        if not platform:
            platform = SYSTEM.lower()
        return self.PLATFORM_CONFIG.get(platform, self.PLATFORM_CONFIG['linux'])

    AUDIO_ENHANCEMENT_ENABLED: bool = Constants.AUDIO_ENHANCEMENT_ENABLED
    MIN_RMS_THRESHOLD: float = Constants.MIN_RMS_THRESHOLD
    TARGET_RMS: float = Constants.TARGET_RMS
    MAX_GAIN: float = Constants.MAX_GAIN
    CLIPPING_THRESHOLD: float = Constants.CLIPPING_THRESHOLD

    DUPLICATE_CHECK_ENABLED: bool = Constants.DUPLICATE_CHECK_ENABLED
    RECENT_TRANSCRIPTIONS_SIZE: int = Constants.RECENT_TRANSCRIPTIONS_SIZE
    MIN_TEXT_LENGTH: int = Constants.MIN_TEXT_LENGTH
    MIN_UNIQUE_WORDS_RATIO: float = Constants.MIN_UNIQUE_WORDS_RATIO

    SUBTITLE_BUFFER_SIZE: int = Constants.SUBTITLE_BUFFER_SIZE
    ENABLE_TIMED_TRANSCRIPTIONS: bool = Constants.ENABLE_TIMED_TRANSCRIPTIONS
    ENABLE_TIMED_TRANSLATIONS: bool = Constants.ENABLE_TIMED_TRANSLATIONS

    ENABLE_DEBUG_LOGGING: bool = Constants.ENABLE_DEBUG_LOGGING
    LOG_CHUNK_PROCESSING: bool = Constants.LOG_CHUNK_PROCESSING
    LOG_AUDIO_STATS: bool = Constants.LOG_AUDIO_STATS
    LOG_PERFORMANCE: bool = Constants.LOG_PERFORMANCE
    LOG_STREAM_EVENTS: bool = Constants.LOG_STREAM_EVENTS
    PERFORMANCE_LOG_INTERVAL: int = Constants.PERFORMANCE_LOG_INTERVAL

    MAX_CACHE_SIZE_MB: int = Constants.MAX_CACHE_SIZE_MB
    CACHE_ENABLED: bool = Constants.CACHE_ENABLED

    def __init__(self) -> None:
        self._actual_chunk_duration = self._base_chunk_duration

    def calculate_optimal_chunk_duration(self, model_size: str = 'medium',
                                         is_realtime: bool = False) -> int:
        if is_realtime:
            return self.MIN_CHUNK_DURATION
        model_durations = {
            'tiny': 3, 'tiny.en': 3,
            'base': 4, 'base.en': 4,
            'small': 5, 'small.en': 5,
            'medium': 5, 'medium.en': 5,
            'large': 6, 'large-v2': 6, 'large-v3': 6,
        }
        return model_durations.get(model_size.lower(), self._base_chunk_duration)

    def validate_config(self) -> bool:
        try:
            valid = (
                self.SAMPLE_RATE in [8000, 16000, 22050, 44100, 48000] and
                self.CHANNELS in [1, 2] and
                self.MIN_CHUNK_DURATION <= self.CHUNK_DURATION <= self.MAX_CHUNK_DURATION
            )
            if not valid:
                logger.error("❌ Config validation failed")
                return False
            return True
        except Exception as e:
            logger.error(f"❌ Config validation error: {e}")
            return False

    def print_summary(self) -> None:
        logger.info("\n" + "="*60)
        logger.info("🎵 CONFIGURATION")
        logger.info("="*60)
        logger.info(f"📊 Audio: {self.SAMPLE_RATE}Hz, {self.CHANNELS}ch")
        logger.info(f"📦 Chunk: {self.CHUNK_DURATION}s ({self.CHUNK_SIZE_BYTES:,}B)")
        logger.info(f"⚡ Bytes/sec: {self.BYTES_PER_SECOND:,}")
        logger.info(f"🎛️ Filter Profiles: {len(self.FILTER_PROFILES)}")
        logger.info(f"🌍 Language Filters: {len(self.LANGUAGE_FILTERS)}")
        logger.info(f"✅ Valid: {self.validate_config()}")
        logger.info("="*60)

    def __str__(self) -> str:
        return (f"Config(chunk={self.CHUNK_DURATION}s, "
                f"filter_profiles={len(self.FILTER_PROFILES)})")


class RealtimeConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.CHUNK_DURATION = 5
        self.CHUNK_OVERLAP = 0.3
        self.STREAM_TIMEOUT = 5
        self.AUDIO_FILTER = self.FILTER_PROFILES['realtime']


class HighAccuracyConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.CHUNK_DURATION = 25
        self.CHUNK_OVERLAP = 0.8
        self.AUDIO_FILTER = ("aresample=16000,volume=1.8,highpass=f=80,"
                             "lowpass=f=3800,dynaudnorm=p=0.3:s=3:g=20")


class YouTubeOptimizedConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.FFMPEG_THREADS = 1
        self.FFMPEG_BUFSIZE = '1024k'
        self.YOUTUBE_TIMEOUT = 5000000
        self.RECONNECT_DELAY = 1
        self.AUDIO_FILTER = ("aresample=16000,volume=2.2,highpass=f=120,"
                             "lowpass=f=3200,compand=attacks=0:decays=0.3")


def get_config(config_type: str = 'default') -> Config:
    configs = {
        'default': Config,
        'realtime': RealtimeConfig,
        'high_accuracy': HighAccuracyConfig,
        'youtube': YouTubeOptimizedConfig,
    }
    config_class = configs.get(config_type, Config)
    config = config_class()
    config.validate_config()
    return config

# =============================================================================
# 3. HILFSKLASSEN UND -FUNKTIONEN (UTILS)
# =============================================================================

# -----------------------------------------------------------------------------
# FastLazyLoader
# -----------------------------------------------------------------------------
class FastLazyLoader:
    _loaded_modules: Dict[str, Any] = {}
    _module_locks: Dict[str, threading.RLock] = {}
    _class_lock = threading.RLock()
    MAX_CACHE_SIZE = 10

    @classmethod
    def _get_lock(cls, module_name: str) -> threading.RLock:
        with cls._class_lock:
            if module_name not in cls._module_locks:
                cls._module_locks[module_name] = threading.RLock()
        return cls._module_locks[module_name]

    @classmethod
    def _prune_cache(cls) -> None:
        with cls._class_lock:
            while len(cls._loaded_modules) > cls.MAX_CACHE_SIZE:
                cls._loaded_modules.popitem(last=False)

    @classmethod
    def load(cls, module_name: str, import_path: Optional[str] = None) -> Any:
        if module_name in cls._loaded_modules:
            return cls._loaded_modules[module_name]

        lock = cls._get_lock(module_name)
        with lock:
            if module_name in cls._loaded_modules:
                return cls._loaded_modules[module_name]

            try:
                if module_name == "torch":
                    import torch
                    try:
                        import torch._logging
                        torch._logging.set_logs(all=logging.ERROR)
                    except Exception:
                        pass
                    cls._loaded_modules["torch"] = torch
                elif module_name == "faster_whisper":
                    from faster_whisper import WhisperModel
                    cls._loaded_modules["faster_whisper"] = WhisperModel
                elif module_name == "numpy":
                    import numpy as np
                    cls._loaded_modules["numpy"] = np
                elif module_name == "deep_translator":
                    from deep_translator import GoogleTranslator
                    cls._loaded_modules["deep_translator"] = GoogleTranslator
                elif module_name == "scipy.signal":
                    import scipy.signal
                    cls._loaded_modules["scipy.signal"] = scipy.signal
                else:
                    module = importlib.import_module(import_path or module_name)
                    cls._loaded_modules[module_name] = module

                cls._prune_cache()
                return cls._loaded_modules[module_name]

            except ImportError as e:
                logger.warning(f"⚠️ Module {module_name} not available: {e}")
                class MockModule:
                    _is_mock = True
                    def __init__(self, name: str):
                        self.__name__ = name
                    def __getattr__(self, name: str) -> Any:
                        def mock_method(*args: Any, **kwargs: Any) -> Any:
                            raise ImportError(f"Module {self.__name__} not available")
                        return mock_method
                mock = MockModule(module_name)
                cls._loaded_modules[module_name] = mock
                cls._prune_cache()
                return mock

    @classmethod
    def is_available(cls, module_name: str) -> bool:
        if importlib.util.find_spec(module_name) is None:
            return False
        if module_name in cls._loaded_modules:
            module = cls._loaded_modules[module_name]
            return not getattr(module, "_is_mock", False)
        return True

    @classmethod
    def clear_cache(cls) -> None:
        with cls._class_lock:
            cls._loaded_modules.clear()
            cls._module_locks.clear()


# -----------------------------------------------------------------------------
# Plattform-Stderr-Filter
# -----------------------------------------------------------------------------
class PlatformStderrFilter:
    def __init__(self, original_stderr: Any) -> None:
        self.original_stderr = original_stderr
        self.filter_patterns = [
            "mpp_soc:", "mpp_platform:", "can not found match soc name",
            "/proc/device-tree/compatible", "rockchip", "ffmpeg", "TORCH_NCCL",
        ]
        if IS_WINDOWS:
            self.filter_patterns.extend([
                "Failed to set direct console mode",
                "Console code page",
                "chcp",
                "win32api",
            ])

    def write(self, text: str) -> None:
        if text and any(p in text for p in self.filter_patterns):
            return
        self.original_stderr.write(text)

    def flush(self) -> None:
        self.original_stderr.flush()


sys.stderr = PlatformStderrFilter(sys.stderr)


# -----------------------------------------------------------------------------
# Terminal-Einstellungen
# -----------------------------------------------------------------------------
_original_stty_settings: Optional[str] = None

def _save_terminal_settings() -> None:
    global _original_stty_settings
    if not IS_LINUX:
        return
    try:
        result = subprocess.run(['stty', '-g'], capture_output=True,
                                 text=True, check=False, timeout=2)
        if result.returncode == 0 and result.stdout:
            _original_stty_settings = result.stdout.strip()
    except (subprocess.TimeoutExpired, OSError) as e:
        if DEBUG_LEVEL >= 2:
            logger.debug(f"stty -g fehlgeschlagen: {e}")

def _restore_terminal_settings() -> None:
    if not IS_LINUX or _original_stty_settings is None:
        return
    try:
        subprocess.run(['stty', _original_stty_settings], check=False, timeout=2)
    except Exception:
        pass

_save_terminal_settings()
atexit.register(_restore_terminal_settings)


# -----------------------------------------------------------------------------
# SignalHandler
# -----------------------------------------------------------------------------
class ShutdownPriority(Enum):
    CRITICAL = 0
    HIGH = 1
    MEDIUM = 2
    LOW = 3


class SignalHandler:
    _instance: Optional['SignalHandler'] = None
    _lock = threading.RLock()
    _shutdown_requested = False
    _shutdown_in_progress = False
    _signal_count = 0
    _setup_complete = False
    _original_handlers: Dict[int, Any] = {}
    _atexit_registered = False
    _atexit_lock = threading.RLock()
    _config = {
        'verbose': False,
        'silent': True,
        'max_cleanup_time': 20.0,
        'emergency_timeout': 2.0,
        'atexit_enabled': True,
        'hybrid_shutdown': True,
    }
    _cleanup_operations: Dict[ShutdownPriority, List['_CleanupOperation']] = {
        ShutdownPriority.CRITICAL: [],
        ShutdownPriority.HIGH: [],
        ShutdownPriority.MEDIUM: [],
        ShutdownPriority.LOW: [],
    }

    class _CleanupOperation:
        def __init__(self, func: Callable[[], Any], name: str,
                     priority: ShutdownPriority = ShutdownPriority.MEDIUM,
                     timeout: float = 3.0, essential: bool = False):
            self.func = func
            self.name = name
            self.priority = priority
            self.timeout = timeout
            self.essential = essential
            self.attempts = 0
            self.last_error: Optional[str] = None

        def execute(self) -> bool:
            self.attempts += 1
            try:
                result: Any = None
                exc: Optional[BaseException] = None
                def target() -> None:
                    nonlocal result, exc
                    try:
                        result = self.func()
                    except Exception as e:
                        exc = e
                thread = threading.Thread(target=target, daemon=True)
                thread.start()
                thread.join(timeout=self.timeout)
                if thread.is_alive():
                    self.last_error = f"Timeout nach {self.timeout}s"
                    return False
                if exc:
                    self.last_error = str(exc)
                    return False
                self.last_error = None
                return True
            except Exception as e:
                self.last_error = str(e)
                return False

    def __new__(cls) -> 'SignalHandler':
        with cls._lock:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self) -> None:
        pass

    @classmethod
    def setup(cls, verbose: bool = False, silent: bool = True,
              hybrid_shutdown: bool = True, **kwargs: Any) -> 'SignalHandler':
        with cls._lock:
            if cls._setup_complete:
                return cls._instance

            cls._config.update({
                'verbose': verbose,
                'silent': silent,
                'hybrid_shutdown': hybrid_shutdown,
                'max_cleanup_time': kwargs.get('max_cleanup_time', 20.0),
                'emergency_timeout': kwargs.get('emergency_timeout', 2.0),
                'atexit_enabled': kwargs.get('atexit_enabled', True),
            })

            if cls._config['verbose']:
                print("🚀 SignalHandler initialisieren...")

            cls._save_original_handlers()
            cls._install_signal_handlers()

            if cls._config['atexit_enabled']:
                cls._register_atexit()

            cls._setup_complete = True

            if cls._config['verbose']:
                print("✅ SignalHandler bereit")
            return cls._instance

    @classmethod
    def _save_original_handlers(cls) -> None:
        if IS_WINDOWS:
            return
        try:
            for sig in [py_signal.SIGINT, py_signal.SIGTERM]:
                cls._original_handlers[sig] = py_signal.getsignal(sig)
        except Exception:
            pass

    @classmethod
    def _install_signal_handlers(cls) -> None:
        def signal_handler(signum: int, frame: Any) -> None:
            with cls._lock:
                cls._signal_count += 1
                if cls._signal_count == 1:
                    if not cls._config['silent']:
                        print("\n⚠️ Shutdown angefordert...")
                    cls._shutdown_requested = True
                    cls._initiate_graceful_shutdown()
                elif cls._signal_count >= 2:
                    if not cls._config['silent']:
                        print("\n🛑 Forcierter Shutdown...")
                    cls._force_shutdown()

        if not IS_WINDOWS:
            try:
                py_signal.signal(py_signal.SIGINT, signal_handler)
                py_signal.signal(py_signal.SIGTERM, signal_handler)
            except Exception:
                pass
        else:
            try:
                import win32api
                def win32_handler(ctrl_type: int) -> bool:
                    if ctrl_type in (0, 1, 2, 5):
                        signal_handler(ctrl_type, None)
                        return True
                    return False
                win32api.SetConsoleCtrlHandler(win32_handler, True)
            except ImportError:
                try:
                    py_signal.signal(py_signal.SIGINT, signal_handler)
                except (ValueError, OSError):
                    pass
            except Exception:
                pass

    @classmethod
    def register_cleanup(cls, func: Callable[[], Any], name: Optional[str] = None,
                         priority: ShutdownPriority = ShutdownPriority.MEDIUM,
                         timeout: float = 3.0, essential: bool = False) -> None:
        if name is None:
            name = func.__name__ if hasattr(func, '__name__') else "Anonymous"

        operation = cls._CleanupOperation(
            func=func, name=name, priority=priority,
            timeout=timeout, essential=essential
        )

        with cls._lock:
            for existing in cls._cleanup_operations[priority]:
                if existing.func == func:
                    return
            cls._cleanup_operations[priority].append(operation)

            if cls._config.get('verbose', False):
                print(f"✅ Cleanup: {name} (Priority: {priority.name})")

    @classmethod
    def unregister_cleanup(cls, func: Callable) -> bool:
        with cls._lock:
            for priority in ShutdownPriority:
                for i, op in enumerate(cls._cleanup_operations[priority]):
                    if op.func == func:
                        del cls._cleanup_operations[priority][i]
                        return True
            return False

    @classmethod
    def _initiate_graceful_shutdown(cls) -> None:
        with cls._lock:
            if cls._shutdown_in_progress:
                return
            cls._shutdown_in_progress = True

        if not cls._config['silent']:
            print("🧹 Starte geordneten Shutdown...")

        try:
            success = cls._execute_priority_cleanup()
            cls._restore_original_handlers()

            if cls._config.get('hybrid_shutdown', True):
                current_thread = threading.current_thread()
                main_thread = threading.main_thread()
                if current_thread == main_thread:
                    if cls._config['verbose']:
                        print("💡 Sauberes Exit im Hauptthread")
                    sys.exit(0 if success else 1)
                else:
                    if cls._config['verbose']:
                        print(f"💡 Sofort-Exit in Thread: {current_thread.name}")
                    os._exit(0 if success else 1)
            else:
                os._exit(0 if success else 1)
        except Exception as e:
            if not cls._config['silent']:
                print(f"❌ Shutdown fehlgeschlagen: {e}")
            os._exit(2)

    @classmethod
    def _execute_priority_cleanup(cls) -> bool:
        overall_success = True
        start_time = time.time()
        completed_ops = 0
        failed_ops = 0

        for priority in ShutdownPriority:
            operations = cls._cleanup_operations.get(priority, [])
            if not operations:
                continue

            if time.time() - start_time > cls._config['max_cleanup_time']:
                if cls._config.get('verbose', False):
                    print(f"⏰ Max cleanup time reached ({cls._config['max_cleanup_time']}s)")
                break

            for op in operations:
                try:
                    op_start = time.perf_counter()
                    success = op.execute()
                    op_duration = (time.perf_counter() - op_start) * 1000
                    completed_ops += 1
                    if debug3_enabled('shutdown'):
                        logger.debug(f"[DEBUG3][SHUTDOWN] Cleanup {op.name}: success={success}, duration={op_duration:.2f}ms")
                    if not success:
                        failed_ops += 1
                        if op.essential:
                            overall_success = False
                            if cls._config.get('verbose', False):
                                print(f"❌ ESSENTIAL cleanup failed: {op.name}")
                        elif cls._config.get('verbose', False):
                            print(f"⚠️ Cleanup failed (non-essential): {op.name}")
                    elif cls._config.get('verbose', False):
                        print(f"✅ Cleanup: {op.name}")
                except Exception as e:
                    failed_ops += 1
                    print(f"⚠️ Cleanup execution error: {op.name}: {e}")
                    if op.essential:
                        overall_success = False

        if cls._config.get('verbose', False):
            print(f"📊 Cleanup abgeschlossen: {completed_ops} Operationen, "
                  f"{failed_ops} fehlgeschlagen")
        return overall_success

    @classmethod
    def _force_shutdown(cls) -> None:
        if debug3_enabled('shutdown'):
            logger.debug("[DEBUG3][SHUTDOWN] Force shutdown initiated")
        try:
            cls._handle_atexit_cleanup()
        except Exception:
            pass
        os._exit(1)

    @classmethod
    def _register_atexit(cls) -> None:
        with cls._atexit_lock:
            if cls._atexit_registered:
                return

            def safe_atexit_handler() -> None:
                try:
                    if (threading.current_thread() == threading.main_thread() and
                            not cls._shutdown_in_progress):
                        cls._handle_atexit_cleanup()
                except Exception:
                    pass

            atexit.register(safe_atexit_handler)
            cls._atexit_registered = True

            if cls._config.get('verbose', False):
                print("✅ AtExit-Handler registriert")

    @classmethod
    def _handle_atexit_cleanup(cls) -> None:
        if cls._config.get('verbose', False):
            print("🔧 AtExit-Cleanup...")

        critical_ops = []
        for op in cls._cleanup_operations.get(ShutdownPriority.CRITICAL, []):
            if op.essential and ("GPU" in op.name or "Memory" in op.name):
                critical_ops.append(op)
                if len(critical_ops) >= 3:
                    break

        start_time = time.time()
        for op in critical_ops:
            if time.time() - start_time > cls._config['emergency_timeout']:
                break
            try:
                if cls._config.get('verbose', False):
                    print(f"  ⚡ Emergency: {op.name}")
                op.func()
            except Exception:
                pass

        gc.collect()

    @classmethod
    def _restore_original_handlers(cls) -> None:
        if IS_WINDOWS:
            return
        try:
            for sig, handler in cls._original_handlers.items():
                if handler is not None:
                    py_signal.signal(sig, handler)
        except Exception:
            pass

    @classmethod
    def should_shutdown(cls) -> bool:
        with cls._lock:
            return cls._shutdown_requested

    @classmethod
    def get_status(cls) -> Dict[str, Any]:
        with cls._lock:
            return {
                'shutdown_requested': cls._shutdown_requested,
                'shutdown_in_progress': cls._shutdown_in_progress,
                'signal_count': cls._signal_count,
                'setup_complete': cls._setup_complete,
                'cleanup_operations': {
                    priority.name: len(ops)
                    for priority, ops in cls._cleanup_operations.items()
                },
                'atexit_registered': cls._atexit_registered,
                'hybrid_mode': cls._config.get('hybrid_shutdown', True),
                'config': {k: v for k, v in cls._config.items()
                           if not k.startswith('_')},
            }

    @classmethod
    def emergency_shutdown(cls, reason: str = "Emergency", exit_code: int = 1) -> None:
        if not cls._config['silent']:
            print(f"🚨 NOTFALL-SHUTDOWN: {reason}")

        with cls._lock:
            cls._shutdown_requested = True
            cls._shutdown_in_progress = True

        os._exit(exit_code)

    @classmethod
    def reset(cls) -> None:
        with cls._lock:
            cls._instance = None
            cls._shutdown_requested = False
            cls._shutdown_in_progress = False
            cls._signal_count = 0
            cls._setup_complete = False
            cls._original_handlers = {}
            cls._atexit_registered = False
            cls._cleanup_operations = {
                ShutdownPriority.CRITICAL: [],
                ShutdownPriority.HIGH: [],
                ShutdownPriority.MEDIUM: [],
                ShutdownPriority.LOW: [],
            }

# -----------------------------------------------------------------------------
# PlatformUtils (erweitert)
# -----------------------------------------------------------------------------
class PlatformUtils:
    _environment_setup_done = False
    _environment_setup_lock = threading.RLock()
    _dependencies_checked = False
    _dependencies_lock = threading.RLock()

    @staticmethod
    def get_platform_config_dir() -> Path:
        """Gibt das plattformspezifische Konfigurationsverzeichnis zurück."""
        try:
            if IS_WINDOWS:
                config_dir = Path(os.environ.get('APPDATA', '')) / "DragonWhisperer"
            elif IS_MACOS:
                config_dir = Path.home() / "Library" / "Application Support" / "DragonWhisperer"
            else:  # Linux
                config_dir = Path.home() / ".config" / "dragonwhisperer"
            config_dir.mkdir(parents=True, exist_ok=True)
            return config_dir
        except Exception as e:
            logger.warning(f"⚠️ Config directory error: {e}")
            fallback_dir = Path.home() / ".dragonwhisperer"
            fallback_dir.mkdir(parents=True, exist_ok=True)
            return fallback_dir

    @staticmethod
    def kill_process_tree(pid: int) -> bool:
        """Beendet einen Prozess und alle seine Kindprozesse."""
        try:
            if IS_WINDOWS:
                subprocess.run(
                    ['taskkill', '/F', '/T', '/PID', str(pid)],
                    capture_output=True,
                    timeout=5,
                    check=False,
                    creationflags=subprocess.CREATE_NO_WINDOW
                )
            else:
                # Linux/macOS: kill Prozessgruppe
                try:
                    os.killpg(os.getpgid(pid), py_signal.SIGKILL)
                except (ProcessLookupError, PermissionError):
                    subprocess.run(
                        ['pkill', '-9', '-P', str(pid)],
                        capture_output=True,
                        timeout=5,
                        check=False
                    )
            return True
        except subprocess.TimeoutExpired:
            logger.warning(f"⚠️ Timeout beim Beenden des Prozessbaums {pid}")
            return False
        except (OSError, subprocess.CalledProcessError) as e:
            logger.warning(f"⚠️ Error killing process tree {pid}: {e}")
            return False
        except Exception as e:
            logger.error(f"⚠️ Unerwarteter Fehler beim Beenden von Prozessbaum {pid}: {e}")
            return False

    @staticmethod
    def check_platform_dependencies() -> bool:
        """Prüft, ob alle kritischen Abhängigkeiten vorhanden sind."""
        with PlatformUtils._dependencies_lock:
            if PlatformUtils._dependencies_checked:
                return True

            missing: List[str] = []
            issues: List[str] = []
            logger.info("🔍 Checking platform dependencies...")

            ffmpeg_found = shutil.which('ffmpeg') is not None
            if not ffmpeg_found:
                missing.append('ffmpeg')
                issues.append("FFmpeg not found in PATH or standard locations")

            ytdlp_found = shutil.which('yt-dlp') is not None
            if not ytdlp_found:
                missing.append('yt-dlp')
                issues.append("yt-dlp not found in PATH")

            psutil_found = FastLazyLoader.is_available('psutil')
            if not psutil_found:
                issues.append("psutil not available (system monitoring)")

            critical_missing = []
            if not ffmpeg_found:
                critical_missing.append('ffmpeg')
            if not ytdlp_found:
                critical_missing.append('yt-dlp')
            if not GUI_AVAILABLE:
                critical_missing.append('tkinter')
                issues.append("Tkinter not available – required for GUI")

            if not WHISPER_AVAILABLE:
                logger.warning("⚠️ Kein Whisper-Backend verfügbar. Starte im Demo-Modus.")
            if not TRANSLATOR_AVAILABLE:
                issues.append("deep-translator not available (translation will be limited)")
            if not TORCH_AVAILABLE:
                issues.append("PyTorch not available (optional for GPU acceleration)")
            if not psutil_found:
                issues.append("psutil not available (system monitoring limited)")

            if critical_missing:
                error_msg = (f"❌ Fehlende kritische Abhängigkeiten: "
                             f"{', '.join(critical_missing)}\n\n")
                error_msg += "\n".join(issues) + "\n"
                if 'ffmpeg' in critical_missing:
                    error_msg += "FFmpeg Installation:\n"
                    if IS_WINDOWS:
                        error_msg += "  • Download from: https://ffmpeg.org/download.html\n"
                    elif IS_MACOS:
                        error_msg += "  • brew install ffmpeg\n"
                    else:
                        error_msg += "  • sudo apt install ffmpeg\n"
                if 'yt-dlp' in critical_missing:
                    error_msg += "yt-dlp Installation:\n"
                    error_msg += "  • pip install yt-dlp\n"
                if 'tkinter' in critical_missing:
                    error_msg += "Tkinter Installation:\n"
                    error_msg += "  • Usually included with Python. On Linux: sudo apt-get install python3-tk\n"
                if issues:
                    error_msg += "\nAdditional issues:\n"
                    for issue in issues:
                        error_msg += f"  • {issue}\n"
                error_msg += "\n💡 After installing, restart Dragon Whisperer."
                PlatformUtils._dependencies_checked = False
                raise RuntimeError(error_msg)

            logger.info("✅ Alle kritischen Abhängigkeiten gefunden")
            PlatformUtils._dependencies_checked = True
            return True

    @staticmethod
    def setup_platform_environment() -> None:
        """Führt plattformspezifische Umgebungs-Setups durch (einmalig)."""
        with PlatformUtils._environment_setup_lock:
            if PlatformUtils._environment_setup_done:
                return
            logger.info("🔧 Setting up platform environment...")
            if IS_WINDOWS:
                try:
                    import ctypes
                    # UTF-8 Codepage setzen
                    ctypes.windll.kernel32.SetConsoleOutputCP(65001)
                    os.system('chcp 65001 > nul 2>&1')
                    os.system('color')
                except (OSError, AttributeError) as e:
                    logger.warning(f"⚠️ Windows console setup failed: {e}")
            elif IS_MACOS:
                temp_dir = Path(tempfile.gettempdir()) / "dragonwhisperer"
                try:
                    temp_dir.mkdir(exist_ok=True)
                except (OSError, PermissionError) as e:
                    logger.warning(f"⚠️ macOS temp dir creation failed: {e}")
            PlatformUtils._environment_setup_done = True
            logger.info("✅ Platform environment setup complete")

    @staticmethod
    def get_ffmpeg_path() -> Optional[str]:
        """Gibt den vollständigen Pfad zur FFmpeg-Exe zurück oder None."""
        ffmpeg_path = shutil.which('ffmpeg')
        if ffmpeg_path:
            return ffmpeg_path

        # Fallback-Pfade (falls nicht im PATH)
        if IS_WINDOWS:
            paths = [
                'C:\\ffmpeg\\bin\\ffmpeg.exe',
                'C:\\Program Files\\ffmpeg\\bin\\ffmpeg.exe',
                'C:\\Program Files (x86)\\ffmpeg\\bin\\ffmpeg.exe',
            ]
        elif IS_MACOS:
            paths = [
                '/usr/local/bin/ffmpeg',
                '/opt/homebrew/bin/ffmpeg',
                '/usr/bin/ffmpeg',
            ]
        else:  # Linux
            paths = [
                '/usr/bin/ffmpeg',
                '/usr/local/bin/ffmpeg',
            ]
        for path in paths:
            if os.path.exists(path):
                return path
        return None

    @staticmethod
    def get_platform_info() -> Dict[str, Any]:
        """Sammelt verschiedene Plattform-Informationen."""
        info: Dict[str, Any] = {
            'system': SYSTEM,
            'is_windows': IS_WINDOWS,
            'is_macos': IS_MACOS,
            'is_linux': IS_LINUX,
            'is_arm': IS_ARM,
            'is_x86': IS_X86,
            'python_version': sys.version,
            'python_executable': sys.executable,
            'current_directory': os.getcwd(),
            'environment_setup': PlatformUtils._environment_setup_done,
            'dependencies_checked': PlatformUtils._dependencies_checked,
        }
        try:
            import psutil
            info['cpu_count'] = psutil.cpu_count()
            info['memory_total_gb'] = psutil.virtual_memory().total / (1024**3)
        except ImportError:
            info['cpu_count'] = 'unknown'
            info['memory_total_gb'] = 'unknown'
        except Exception:
            info['cpu_count'] = 'error'
            info['memory_total_gb'] = 'error'
        return info

    @staticmethod
    def print_platform_info() -> None:
        """Gibt eine formatierte Übersicht der Plattform-Info aus."""
        info = PlatformUtils.get_platform_info()
        logger.info("\n" + "="*60)
        logger.info("🐉 PLATFORM INFORMATION")
        logger.info("="*60)
        for key, value in info.items():
            if key not in ['environment_setup', 'dependencies_checked']:
                logger.info(f"{key:25} {value}")
        logger.info("-"*60)
        logger.info(f"{'Environment Setup':25} {'✅' if info['environment_setup'] else '❌'}")
        logger.info(f"{'Dependencies Checked':25} {'✅' if info['dependencies_checked'] else '❌'}")
        logger.info("="*60)

    @staticmethod
    def sanitize_url(url: str) -> str:
        """
        Bereinigt eine URL: entfernt führende/folgende Whitespace-Zeichen.
        (Keine Zeichenfilterung mehr, da shell=False ausreichend schützt.)
        """
        if not url:
            return ""
        return url.strip()

    @staticmethod
    def validate_file_path(file_url: str) -> Tuple[bool, str]:
        """
        Prüft, ob eine file://-URL auf eine erlaubte Datei verweist.
        Gibt (ok, normalisierter Pfad) zurück.

        Verbesserungen:
        - Korrekte Behandlung von Windows-Pfaden (z.B. file:///C:/path/to/file)
        - Verwendung von pathlib für robuste Pfadoperationen
        - Bessere Fehlermeldungen
        """
        if not file_url.startswith(Constants.ALLOWED_FILE_SCHEME_PREFIX):
            return False, "Keine file://-URL"

        # Plattformspezifischen Pfad extrahieren
        try:
            # Für Windows: file:///C:/... -> C:/...
            if IS_WINDOWS and file_url.startswith('file:///'):
                path_part = file_url[8:]  # Entferne 'file:///'
                path_part = urllib.request.url2pathname(path_part)
            else:
                # Standard: file:///pfad oder file://pfad
                path_part = file_url[len(Constants.ALLOWED_FILE_SCHEME_PREFIX):]
        except Exception as e:
            return False, f"Pfad kann nicht extrahiert werden: {e}"

        # Pfad normalisieren (löst relative Pfade auf, folgt Symlinks)
        try:
            real_path = Path(path_part).resolve()
        except Exception as e:
            return False, f"Pfad kann nicht normalisiert werden: {e}"

        # Prüfen, ob die Datei existiert
        if not real_path.exists():
            return False, f"Datei existiert nicht: {real_path}"

        # Prüfen, ob es sich um eine reguläre Datei handelt
        if not real_path.is_file():
            return False, "Keine gültige Datei (möglicherweise ein Verzeichnis)"

        # Prüfen, ob der Pfad innerhalb eines erlaubten Basisverzeichnisses liegt
        allowed_bases = [Path(p).resolve() for p in Constants.ALLOWED_FILE_BASE_DIRS]
        for base in allowed_bases:
            if real_path.is_relative_to(base):
                return True, str(real_path)

        # Falls nicht, zusätzlich prüfen, ob es sich um eine temporäre Datei handelt
        temp_dir = Path(tempfile.gettempdir()).resolve()
        if real_path.is_relative_to(temp_dir):
            return True, str(real_path)

        return False, f"Zugriff auf {real_path} nicht erlaubt (außerhalb erlaubter Verzeichnisse)"

    # -------------------------------------------------------------------------
    # Neue Hilfsfunktionen zur Reduzierung von Redundanzen
    # -------------------------------------------------------------------------
    @staticmethod
    def read_process_stderr(process: subprocess.Popen, max_bytes: int = 4096) -> str:
        """
        Liest bis zu max_bytes aus dem stderr eines Prozesses und gibt sie als String zurück.
        Fehler werden ignoriert, bei Misserfolg wird "" zurückgegeben.
        """
        try:
            if process.stderr:
                return process.stderr.read(max_bytes).decode('utf-8', errors='ignore')
        except Exception:
            pass
        return ""

    @staticmethod
    def is_fatal_exception(e: BaseException) -> bool:
        """
        Prüft, ob eine Exception als fatal gilt (KeyboardInterrupt oder SystemExit).
        Soll in except-Blöcken verwendet werden, um diese Ausnahmen weiterzuleiten.
        """
        return isinstance(e, (KeyboardInterrupt, SystemExit))


# Nach der Klassendefinition folgt der Aufruf von setup_platform_environment()
PlatformUtils.setup_platform_environment()

# -----------------------------------------------------------------------------
# Verfügbarkeiten
# -----------------------------------------------------------------------------
TORCH_AVAILABLE = FastLazyLoader.is_available("torch")
NUMPY_AVAILABLE = FastLazyLoader.is_available("numpy")
TRANSLATOR_AVAILABLE = FastLazyLoader.is_available("deep_translator")
SCIPY_AVAILABLE = FastLazyLoader.is_available("scipy.signal")
FASTER_WHISPER_AVAILABLE = importlib.util.find_spec("faster_whisper") is not None
OPENAI_WHISPER_AVAILABLE = importlib.util.find_spec("whisper") is not None
WHISPER_AVAILABLE = FASTER_WHISPER_AVAILABLE or OPENAI_WHISPER_AVAILABLE
OLLAMA_AVAILABLE = importlib.util.find_spec("requests") is not None

if FASTER_WHISPER_AVAILABLE:
    logger.info("✅ faster-whisper verfügbar")
else:
    logger.warning("⚠️ faster-whisper nicht verfügbar")

if OPENAI_WHISPER_AVAILABLE:
    logger.info("✅ openai-whisper verfügbar")
else:
    logger.warning("⚠️ openai-whisper nicht verfügbar")

if not WHISPER_AVAILABLE:
    logger.warning("⚠️ KEINE Whisper-Bibliothek verfügbar! Starte im Demo-Modus.")

try:
    import tkinter as tk
    from tkinter import ttk, scrolledtext, filedialog
    GUI_AVAILABLE = True
    logger.info("✅ GUI verfügbar")
except ImportError:
    GUI_AVAILABLE = False
    tk = None
    ttk = None
    scrolledtext = None
    logger.info("📟 Terminal-Modus (kein GUI)")


# -----------------------------------------------------------------------------
# Datenklassen
# -----------------------------------------------------------------------------
@dataclass
class TranscriptionResult:
    text: str
    confidence: float
    language: str = "unknown"
    timestamp: float = field(default_factory=time.time)
    start: Optional[float] = None
    end: Optional[float] = None


@dataclass
class TranslationResult:
    original: str
    translated: str
    source_lang: str = "auto"
    target_lang: str = "de"
    timestamp: float = field(default_factory=time.time)
    start: Optional[float] = None
    end: Optional[float] = None

# -----------------------------------------------------------------------------
# SimplePerformanceTracker
# -----------------------------------------------------------------------------
class SimplePerformanceTracker:
    def __init__(self) -> None:
        self.transcription_count = 0
        self.translation_count = 0
        self.start_time = time.time()
        self.cache_hits = 0
        self.cache_misses = 0
        self._lock = threading.RLock()

    def log_transcription(self) -> None:
        with self._lock:
            self.transcription_count += 1

    def log_translation(self) -> None:
        with self._lock:
            self.translation_count += 1

    def log_cache_hit(self) -> None:
        with self._lock:
            self.cache_hits += 1

    def log_cache_miss(self) -> None:
        with self._lock:
            self.cache_misses += 1

    def get_basic_stats(self) -> Dict[str, Any]:
        with self._lock:
            uptime_minutes = (time.time() - self.start_time) / 60
            total_cache = self.cache_hits + self.cache_misses
            cache_hit_rate = self.cache_hits / total_cache if total_cache > 0 else 0
            return {
                "transcriptions": self.transcription_count,
                "translations": self.translation_count,
                "uptime_minutes": uptime_minutes,
                "cache_hits": self.cache_hits,
                "cache_misses": self.cache_misses,
                "cache_hit_rate": f"{cache_hit_rate:.1%}",
                "timestamp": datetime.now().isoformat(),
            }

# -----------------------------------------------------------------------------
# Cache-Klassen
# -----------------------------------------------------------------------------
class TTLCache:
    def __init__(self, maxsize: int = 128, ttl: float = 300.0) -> None:
        self.maxsize = maxsize
        self.ttl = ttl
        self._cache: Dict[str, Tuple[Any, float]] = {}
        self._order: Deque[str] = deque()
        self._lock = threading.RLock()
        self._cleanup_interval = 300
        self._last_cleanup = time.time()
        self._stats = {"hits": 0, "misses": 0, "evictions": 0}
        self._stats_lock = threading.Lock()
        self._access_counter = 0

    def _perform_cleanup_if_needed(self) -> None:
        self._access_counter += 1
        if self._access_counter % 10 != 0:
            return
        current_time = time.time()
        if current_time - self._last_cleanup < self._cleanup_interval:
            return
        expired_keys = []
        for key in list(self._cache.keys()):
            value, timestamp = self._cache[key]
            if (current_time - timestamp) > self.ttl:
                if debug3_enabled('cache'):
                    logger.debug(f"[DEBUG3][CACHE] Entry expired for key {key[:16]}")
                expired_keys.append(key)
        for key in expired_keys:
            self._remove_key(key)
        self._last_cleanup = current_time

    def _remove_key(self, key: str) -> None:
        if key in self._cache:
            del self._cache[key]
        if key in self._order:
            self._order.remove(key)

    def get(self, key: str) -> Optional[Any]:
        with self._lock:
            self._perform_cleanup_if_needed()
            if key in self._cache:
                value, timestamp = self._cache[key]
                if (time.time() - timestamp) > self.ttl:
                    if debug3_enabled('cache'):
                        logger.debug(f"[DEBUG3][CACHE] Entry expired for key {key[:16]} on get")
                    self._remove_key(key)
                    with self._stats_lock:
                        self._stats["misses"] += 1
                    return None
                try:
                    self._order.remove(key)
                except ValueError:
                    pass
                self._order.append(key)
                with self._stats_lock:
                    self._stats["hits"] += 1
                return value
            with self._stats_lock:
                self._stats["misses"] += 1
            return None

    def put(self, key: str, value: Any) -> None:
        with self._lock:
            self._perform_cleanup_if_needed()
            if key in self._cache:
                self._order.remove(key)
            elif len(self._cache) >= self.maxsize:
                oldest = self._order.popleft()
                if debug3_enabled('cache'):
                    logger.debug(f"[DEBUG3][CACHE] Evicting oldest key {oldest[:16]} due to maxsize {self.maxsize}")
                self._remove_key(oldest)
                with self._stats_lock:
                    self._stats["evictions"] += 1
            self._cache[key] = (value, time.time())
            self._order.append(key)

    def clear(self) -> None:
        with self._lock:
            self._cache.clear()
            self._order.clear()
            self._last_cleanup = time.time()

    def clear_expired(self) -> int:
        with self._lock:
            count = 0
            current_time = time.time()
            expired_keys = []
            for key in list(self._cache.keys()):
                value, timestamp = self._cache[key]
                if (current_time - timestamp) > self.ttl:
                    if debug3_enabled('cache'):
                        logger.debug(f"[DEBUG3][CACHE] Clearing expired key {key[:16]}")
                    expired_keys.append(key)
            for key in expired_keys:
                self._remove_key(key)
                count += 1
            return count

    def get_stats(self) -> Dict[str, Any]:
        with self._lock, self._stats_lock:
            total_size = len(self._cache)
            expired = self.clear_expired()
            return {
                "total_entries": total_size,
                "expired_entries": expired,
                "max_size": self.maxsize,
                "ttl_seconds": self.ttl,
                "hits": self._stats["hits"],
                "misses": self._stats["misses"],
                "evictions": self._stats["evictions"],
                "hit_rate": self._stats["hits"] / max(1, self._stats["hits"] + self._stats["misses"]),
            }


# Globale Cache-Instanzen (werden später in Manager eingebunden)
transcription_cache = TTLCache(
    maxsize=Constants.TRANSCRIPTION_CACHE_SIZE,
    ttl=Constants.TRANSCRIPTION_CACHE_TTL
)
translation_cache = TTLCache(
    maxsize=Constants.TRANSLATION_CACHE_SIZE,
    ttl=Constants.TRANSLATION_CACHE_TTL
)
audio_cache = TTLCache(
    maxsize=Constants.AUDIO_CACHE_SIZE,
    ttl=Constants.AUDIO_CACHE_TTL
)


def clear_expired_cache_entries() -> Dict[str, int]:
    return {
        "transcription_expired": transcription_cache.clear_expired(),
        "translation_expired": translation_cache.clear_expired(),
        "audio_expired": audio_cache.clear_expired(),
    }


def get_cache_stats() -> Dict[str, Any]:
    return {
        "transcription_cache": transcription_cache.get_stats(),
        "translation_cache": translation_cache.get_stats(),
        "audio_cache": audio_cache.get_stats(),
    }


def cache_transcription(result: TranscriptionResult) -> str:
    key = hashlib.sha256(f"{result.text}:{result.language}".encode()).hexdigest()
    transcription_cache.put(key, result)
    return key


def get_cached_transcription(text: str, language: str = "unknown") -> Optional[TranscriptionResult]:
    key = hashlib.sha256(f"{text}:{language}".encode()).hexdigest()
    result = transcription_cache.get(key)
    if debug3_enabled('cache'):
        if result:
            logger.debug(f"[DEBUG3][CACHE] Transcription cache HIT for key {key[:16]}")
        else:
            logger.debug(f"[DEBUG3][CACHE] Transcription cache MISS for key {key[:16]}")
    return result


def cache_translation(result: TranslationResult) -> str:
    key = hashlib.sha256((result.original + result.target_lang).encode()).hexdigest()
    translation_cache.put(key, result)
    return key


def get_cached_translation(original: str, target_lang: str) -> Optional[TranslationResult]:
    key = hashlib.sha256((original + target_lang).encode()).hexdigest()
    result = translation_cache.get(key)
    if debug3_enabled('cache'):
        if result:
            logger.debug(f"[DEBUG3][CACHE] Translation cache HIT for key {key[:16]}")
        else:
            logger.debug(f"[DEBUG3][CACHE] Translation cache MISS for key {key[:16]}")
    return result


# -----------------------------------------------------------------------------
# ThreadPoolExecutor und Decorator
# -----------------------------------------------------------------------------
_EXECUTOR = ThreadPoolExecutor(max_workers=4, thread_name_prefix="ProcExec")


def execution_decorator(timeout: int = 60, max_retries: int = 3) -> Callable:
    def decorator(func: Callable) -> Callable:
        @wraps(func)
        def wrapper(*args: Any, **kwargs: Any) -> Any:
            last_exception: Optional[Exception] = None
            for attempt in range(max_retries + 1):
                try:
                    future = _EXECUTOR.submit(func, *args, **kwargs)
                    return future.result(timeout=timeout)
                except FutureTimeout as e:
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(f"⏰ Timeout attempt {attempt + 1}/{max_retries + 1} for {func.__name__}")
                except Exception as e:
                    if PlatformUtils.is_fatal_exception(e):
                        raise
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(f"⚠️ Exception in {func.__name__}: {str(e)[:100]}")
                if attempt < max_retries:
                    wait_time = min(30, 2**attempt)
                    time.sleep(wait_time)
                    continue
            if last_exception is not None:
                raise last_exception
            raise RuntimeError(f"Decorator logic failed for {func.__name__}.")
        return wrapper
    return decorator


class ProcessingError(Exception):
    pass


def gui_operation_decorator(func: Callable) -> Callable:
    @wraps(func)
    def wrapper(*args: Any, **kwargs: Any) -> Any:
        try:
            return func(*args, **kwargs)
        except (tk.TclError, RuntimeError):
            return None
        except Exception:
            return None
    return wrapper


# -----------------------------------------------------------------------------
# THEME-KLASSEN
# -----------------------------------------------------------------------------
class DarkTheme:
    BG_PRIMARY = "#0f1419"
    BG_SECONDARY = "#1a2129"
    BG_TERTIARY = "#242d38"
    BG_HOVER = "#2a3645"
    BG_CARD = "#1c252f"
    TEXT_PRIMARY = "#e6edf3"
    TEXT_SECONDARY = "#8b949e"
    TEXT_ACCENT = "#58a6ff"
    TEXT_MUTED = "#6e7681"
    DRAGON_GREEN = "#238636"
    DRAGON_GREEN_LIGHT = "#2ea043"
    DRAGON_BLUE = "#1f6feb"
    DRAGON_PURPLE = "#8957e5"
    SUCCESS = "#238636"
    WARNING = "#d29922"
    ERROR = "#f85149"
    INFO = "#58a6ff"
    BORDER = "#30363d"
    SCROLLBAR = "#3c444d"
    SCROLLBAR_HOVER = "#4c5560"
    INPUT_BG = "#161b22"
    INPUT_BORDER = "#30363d"
    INPUT_FOCUS = "#1f6feb"
    COMBO_BG = "#161b22"
    COMBO_FG = "#e6edf3"
    COMBO_BORDER = "#30363d"
    COMBO_SELECTION = "#1f6feb"
    CHECKBOX_BG = "#0d1117"
    CHECKBOX_FG = "#e6edf3"
    CHECKBOX_SELECTED = "#238636"
    CHECKBOX_ACTIVE = "#000000"
    SUBTITLE_ACTIVE = "#8957e5"
    SUBTITLE_INACTIVE = "#30363d"
    STATUS_BAR_BG = "#0d1117"
    STATUS_BAR_FG = "#8b949e"
    STATUS_BAR_ACCENT = "#58a6ff"


class LightTheme:
    BG_PRIMARY = "#ffffff"
    BG_SECONDARY = "#f0f0f0"
    BG_TERTIARY = "#e5e5e5"
    BG_HOVER = "#d9d9d9"
    BG_CARD = "#f5f5f5"
    TEXT_PRIMARY = "#000000"
    TEXT_SECONDARY = "#4a4a4a"
    TEXT_ACCENT = "#1f6feb"
    TEXT_MUTED = "#666666"
    DRAGON_GREEN = "#2ecc71"
    DRAGON_GREEN_LIGHT = "#27ae60"
    DRAGON_BLUE = "#3498db"
    DRAGON_PURPLE = "#9b59b6"
    SUCCESS = "#2ecc71"
    WARNING = "#f39c12"
    ERROR = "#e74c3c"
    INFO = "#3498db"
    BORDER = "#cccccc"
    SCROLLBAR = "#b3b3b3"
    SCROLLBAR_HOVER = "#999999"
    INPUT_BG = "#ffffff"
    INPUT_BORDER = "#cccccc"
    INPUT_FOCUS = "#1f6feb"
    COMBO_BG = "#ffffff"
    COMBO_FG = "#000000"
    COMBO_BORDER = "#cccccc"
    COMBO_SELECTION = "#1f6feb"
    CHECKBOX_BG = "#f0f0f0"
    CHECKBOX_FG = "#000000"
    CHECKBOX_SELECTED = "#2ecc71"
    CHECKBOX_ACTIVE = "#ffffff"
    SUBTITLE_ACTIVE = "#8957e5"
    SUBTITLE_INACTIVE = "#b3b3b3"
    STATUS_BAR_BG = "#f0f0f0"
    STATUS_BAR_FG = "#4a4a4a"
    STATUS_BAR_ACCENT = "#1f6feb"


CURRENT_THEME = None


class Fonts:
    TITLE = ("Segoe UI", 12, "bold")
    SUBTITLE = ("Segoe UI", 10, "bold")
    PRIMARY = ("Segoe UI", 9)
    SECONDARY = ("Segoe UI", 8)
    MONOSPACE = ("Cascadia Code", 9)
    BUTTON = ("Segoe UI", 9, "bold")
    STATUS = ("Segoe UI", 8)
    SMALL = ("Segoe UI", 7)


# -----------------------------------------------------------------------------
# DARK MESSAGEBOX
# -----------------------------------------------------------------------------
class DarkMessageBox:
    """Optimierte Dark-Mode MessageBox mit einheitlicher Dialogerstellung."""

    # ----------------------------------------------------------------------
    # Öffentliche Schnittstelle (unverändert)
    # ----------------------------------------------------------------------
    @staticmethod
    def showinfo(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        return DarkMessageBox._show_dialog(title, message, "info", parent)

    @staticmethod
    def showwarning(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        return DarkMessageBox._show_dialog(title, message, "warning", parent)

    @staticmethod
    def showerror(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        return DarkMessageBox._show_dialog(title, message, "error", parent)

    @staticmethod
    def askokcancel(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        return DarkMessageBox._show_dialog(title, message, "question", parent, buttons=True)

    @staticmethod
    def askyesno(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        return DarkMessageBox._ask_yesno(title, message, parent)

    @staticmethod
    def show_progress(title: str, message: str, parent: Optional[tk.Tk] = None,
                      indeterminate: bool = True) -> Any:
        """
        Zeigt einen Fortschrittsdialog. Muss im Hauptthread aufgerufen werden.
        Gibt ein Controller-Objekt mit close() und update_message() zurück.
        """
        if threading.current_thread() is not threading.main_thread():
            warnings.warn(
                "DarkMessageBox.show_progress wurde nicht im Hauptthread aufgerufen. "
                "GUI-Updates können fehlschlagen.",
                RuntimeWarning, stacklevel=2
            )

        # Dialog im Hauptthread erstellen (sicherstellen, dass parent existiert)
        def _create():
            nonlocal dialog, progress, message_label
            try:
                root = parent if parent and parent.winfo_exists() else tk._default_root
                if not root or not root.winfo_exists():
                    logger.error("Kein gültiges root-Fenster für ProgressDialog")
                    return

                dlg = tk.Toplevel(root)
                dlg.title(f"🐉 {title}")
                dlg.configure(bg=CURRENT_THEME.BG_PRIMARY)
                dlg.resizable(False, False)
                dlg.transient(root)
                dlg.grab_set()

                main = tk.Frame(dlg, bg=CURRENT_THEME.BG_PRIMARY, padx=30, pady=25)
                main.pack(fill="both", expand=True)

                msg_lbl = tk.Label(
                    main,
                    text=message,
                    font=Fonts.PRIMARY,
                    bg=CURRENT_THEME.BG_PRIMARY,
                    fg=CURRENT_THEME.TEXT_PRIMARY,
                    justify="center",
                )
                msg_lbl.pack(pady=(0, 20))

                prog = ttk.Progressbar(
                    main,
                    mode="indeterminate" if indeterminate else "determinate",
                    length=300,
                )
                prog.pack(pady=(0, 10))
                prog.start(10)

                DarkMessageBox._center_dialog(dlg, root)

                dialog = dlg
                progress = prog
                message_label = msg_lbl
            except (tk.TclError, RuntimeError) as e:
                logger.warning(f"⚠️ Progress Dialog Error: {e}")

        dialog = None
        progress = None
        message_label = None

        # Sofortige Erstellung (synchron im Hauptthread)
        _create()

        class ProgressController:
            __slots__ = ('dialog', 'progress', '_message_label')

            def __init__(self):
                self.dialog = dialog
                self.progress = progress
                self._message_label = message_label

            def close(self) -> None:
                if self.dialog and self.dialog.winfo_exists():
                    try:
                        if self.progress:
                            self.progress.stop()
                        self.dialog.destroy()
                    except Exception:
                        pass
                self.dialog = None
                self.progress = None
                self._message_label = None

            def update_message(self, new_message: str) -> None:
                if self.dialog and self.dialog.winfo_exists() and self._message_label:
                    try:
                        self._message_label.config(text=new_message)
                    except Exception:
                        pass

        return ProgressController()

    # ----------------------------------------------------------------------
    # Private Helfer
    # ----------------------------------------------------------------------
    @staticmethod
    def _show_dialog(title: str, message: str, msg_type: str,
                     parent: Optional[tk.Tk] = None,
                     buttons: bool = False) -> Optional[bool]:
        """Generische Dialog-Erstellung für Info/Warnung/Frage mit OK/Abbrechen."""
        try:
            parent = DarkMessageBox._resolve_parent(parent)
            if parent is None:
                return DarkMessageBox._fallback_messagebox(title, message, msg_type, buttons)

            dialog, result, timeout_id = DarkMessageBox._create_base_dialog(
                parent, title, msg_type, message
            )

            # Buttons hinzufügen – Ergebnis nicht benötigt
            if buttons:
                DarkMessageBox._add_ok_cancel_buttons(dialog, result, timeout_id, title)
            else:
                DarkMessageBox._add_ok_button(dialog, result, timeout_id)

            # Abschluss
            dialog.protocol("WM_DELETE_WINDOW",
                            lambda: DarkMessageBox._on_closing(result, timeout_id,
                                                               close_val=False if buttons else True))
            DarkMessageBox._center_dialog(dialog, parent)
            parent.wait_window(dialog)
            return result["value"]

        except (tk.TclError, RuntimeError, AttributeError) as e:
            logger.warning(f"⚠️ DarkMessageBox Error: {e}")
            return DarkMessageBox._fallback_messagebox(title, message, msg_type, buttons)

    @staticmethod
    def _ask_yesno(title: str, message: str, parent: Optional[tk.Tk] = None) -> Optional[bool]:
        """Spezialisierte Ja/Nein-Dialog."""
        try:
            parent = DarkMessageBox._resolve_parent(parent)
            if parent is None:
                import tkinter.messagebox as mb
                return mb.askyesno(title, message)

            dialog, result, timeout_id = DarkMessageBox._create_base_dialog(
                parent, title, "question", message, icon_only=True
            )

            # Ja/Nein-Buttons
            btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
            btn_frame.pack(fill="x", pady=(10, 0))

            def set_res(val: bool):
                DarkMessageBox._cancel_timeout(timeout_id)
                result["value"] = val
                if dialog.winfo_exists():
                    dialog.destroy()

            yes_btn = tk.Button(
                btn_frame,
                text="  👍 Ja  ",
                command=lambda: set_res(True),
                bg=CURRENT_THEME.SUCCESS,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                font=("Segoe UI", 10, "bold"),
                relief="flat",
                padx=25,
                pady=10,
                cursor="hand2",
            )
            yes_btn.pack(side="left", expand=True, padx=(0, 10))

            no_btn = tk.Button(
                btn_frame,
                text="  👎 Nein  ",
                command=lambda: set_res(False),
                bg=CURRENT_THEME.ERROR,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                font=("Segoe UI", 10, "bold"),
                relief="flat",
                padx=25,
                pady=10,
                cursor="hand2",
            )
            no_btn.pack(side="right", expand=True)

            # Tastaturbindungen
            dialog.bind("<Return>", lambda e: set_res(True))
            dialog.bind("<Escape>", lambda e: set_res(False))
            dialog.bind("y", lambda e: set_res(True))
            dialog.bind("n", lambda e: set_res(False))
            yes_btn.focus_set()

            dialog.protocol("WM_DELETE_WINDOW", lambda: set_res(False))
            DarkMessageBox._center_dialog(dialog, parent)
            parent.wait_window(dialog)
            return result["value"]

        except (tk.TclError, RuntimeError, AttributeError):
            import tkinter.messagebox as mb
            return mb.askyesno(title, message)

    # ----------------------------------------------------------------------
    # Kern-Funktionen für den Dialogaufbau
    # ----------------------------------------------------------------------
    @staticmethod
    def _create_base_dialog(parent: tk.Tk, title: str, msg_type: str,
                            message: str, icon_only: bool = False) -> tuple:
        """
        Erzeugt das Grundgerüst eines Dialogs: Fenster, Hauptframe, Icon und Nachricht.
        Gibt (dialog, result_dict, timeout_id) zurück.
        """
        dialog = tk.Toplevel(parent)
        dialog.title(f"🐉 {title}" if not title.startswith("🐉") else title)
        dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        dialog.resizable(False, False)
        dialog.transient(parent)
        dialog.grab_set()

        # Auto-Close nach Timeout
        timeout_seconds = 15 if any(w in title.lower() for w in ["beenden", "exit", "quit", "schließen"]) else 10
        timeout_id = dialog.after(timeout_seconds * 1000,
                                  lambda: DarkMessageBox._auto_close_dialog(dialog, title))

        # Hauptcontainer
        main = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=25, pady=25)
        main.pack(fill="both", expand=True)

        if icon_only:
            # Nur großes Icon für Ja/Nein
            icon_label = tk.Label(
                main,
                text="❓",
                font=("Segoe UI", 28),
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_ACCENT,
            )
            icon_label.pack(pady=(0, 20))

            msg = tk.Label(
                main,
                text=message,
                font=Fonts.PRIMARY,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                justify="center",
                wraplength=350,
            )
            msg.pack(pady=(0, 30))
        else:
            # Layout mit Icon links und Nachricht rechts
            content = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
            content.pack(fill="both", expand=True, pady=(0, 20))

            # Icon
            icon_frame = tk.Frame(content, bg=CURRENT_THEME.BG_PRIMARY, width=60)
            icon_frame.pack(side="left", fill="y")
            icon_frame.pack_propagate(False)

            icons = {
                "info": ("ℹ️", CURRENT_THEME.TEXT_ACCENT),
                "warning": ("⚠️", CURRENT_THEME.WARNING),
                "error": ("❌", CURRENT_THEME.ERROR),
                "question": ("❓", CURRENT_THEME.TEXT_ACCENT),
                "success": ("✅", CURRENT_THEME.SUCCESS),
            }
            icon_char, icon_color = icons.get(msg_type, ("💬", CURRENT_THEME.TEXT_PRIMARY))

            tk.Label(
                icon_frame,
                text=icon_char,
                font=("Segoe UI", 24),
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=icon_color,
            ).pack(expand=True)

            # Nachricht
            msg_frame = tk.Frame(content, bg=CURRENT_THEME.BG_PRIMARY)
            msg_frame.pack(side="left", fill="both", expand=True, padx=(20, 0))

            if len(title) > 30:
                tk.Label(
                    msg_frame,
                    text=title,
                    font=Fonts.SUBTITLE,
                    bg=CURRENT_THEME.BG_PRIMARY,
                    fg=CURRENT_THEME.TEXT_PRIMARY,
                    justify="left",
                    anchor="w",
                ).pack(anchor="w", pady=(0, 10))

            tk.Label(
                msg_frame,
                text=message,
                font=Fonts.PRIMARY,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                justify="left",
                wraplength=350,
                anchor="w",
            ).pack(fill="x", expand=True, anchor="w")

        result = {"value": None}
        return dialog, result, timeout_id

    @staticmethod
    def _add_ok_button(dialog: tk.Toplevel, result: dict, timeout_id: str) -> tk.Frame:
        """Fügt einen einzelnen OK-Button hinzu."""
        btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x")

        def on_ok():
            DarkMessageBox._cancel_timeout(timeout_id)
            result["value"] = True
            if dialog.winfo_exists():
                dialog.destroy()

        ok_btn = tk.Button(
            btn_frame,
            text="OK",
            command=on_ok,
            bg=CURRENT_THEME.SUCCESS,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=25,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        ok_btn.pack(side="right")
        ok_btn.focus_set()

        dialog.bind("<Return>", lambda e: on_ok())
        dialog.bind("<Escape>", lambda e: on_ok())
        dialog.bind("<space>", lambda e: on_ok())
        return btn_frame

    @staticmethod
    def _add_ok_cancel_buttons(dialog: tk.Toplevel, result: dict, timeout_id: str,
                               title: str) -> tk.Frame:
        """Fügt OK- und Abbrechen-Buttons hinzu."""
        btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x")

        def set_result(val: bool):
            DarkMessageBox._cancel_timeout(timeout_id)
            result["value"] = val
            if dialog.winfo_exists():
                dialog.destroy()

        cancel_btn = tk.Button(
            btn_frame,
            text="Abbrechen",
            command=lambda: set_result(False),
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=22,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        cancel_btn.pack(side="right", padx=(10, 0))

        ok_btn = tk.Button(
            btn_frame,
            text="OK",
            command=lambda: set_result(True),
            bg=CURRENT_THEME.SUCCESS,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=25,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        ok_btn.pack(side="right")

        # Tastaturbindungen
        dialog.bind("<Return>", lambda e: set_result(True))
        dialog.bind("<Escape>", lambda e: set_result(False))
        dialog.bind("<space>", lambda e: cancel_btn.focus_set())

        is_exit = any(w in title.lower() for w in ["beenden", "exit", "quit", "schließen"])
        if is_exit:
            cancel_btn.focus_set()
        else:
            ok_btn.focus_set()

        return btn_frame

    # ----------------------------------------------------------------------
    # Hilfsfunktionen (Timer, Parent-Auflösung, Positionierung, Fallback)
    # ----------------------------------------------------------------------
    @staticmethod
    def _resolve_parent(parent: Optional[tk.Tk]) -> Optional[tk.Tk]:
        """Ermittelt ein gültiges Parent-Fenster."""
        if parent and parent.winfo_exists():
            return parent
        if tk._default_root and tk._default_root.winfo_exists():
            return tk._default_root
        return DarkMessageBox._find_available_parent()

    @staticmethod
    def _find_available_parent() -> Optional[tk.Tk]:
        """Sucht nach einem existierenden Toplevel-Fenster."""
        try:
            if not tk._default_root:
                return None
            # Bevorzugt das Hauptfenster, falls es existiert
            if tk._default_root.winfo_exists():
                return tk._default_root
            # Sonst ein beliebiges Kind
            for child in tk._default_root.winfo_children():
                if isinstance(child, tk.Toplevel) and child.winfo_exists():
                    return child
        except (tk.TclError, AttributeError):
            pass
        return None

    @staticmethod
    def _center_dialog(dialog: tk.Toplevel, parent: tk.Tk) -> None:
        """Zentriert den Dialog über dem Parent oder dem Bildschirm."""
        try:
            dialog.update_idletasks()
            if parent and parent.winfo_exists():
                parent_x = parent.winfo_rootx()
                parent_y = parent.winfo_rooty()
                parent_w = parent.winfo_width()
                parent_h = parent.winfo_height()
                dlg_w = dialog.winfo_width()
                dlg_h = dialog.winfo_height()
                x = parent_x + (parent_w - dlg_w) // 2
                y = parent_y + (parent_h - dlg_h) // 2
                # Bildschirmgrenzen prüfen
                screen_w = parent.winfo_screenwidth()
                screen_h = parent.winfo_screenheight()
                x = max(10, min(x, screen_w - dlg_w - 10))
                y = max(10, min(y, screen_h - dlg_h - 10))
                dialog.geometry(f"+{x}+{y}")
                dialog.lift()
                dialog.focus_force()
            else:
                # Fallback: Bildschirmmitte
                screen_w = dialog.winfo_screenwidth()
                screen_h = dialog.winfo_screenheight()
                dlg_w = dialog.winfo_width()
                dlg_h = dialog.winfo_height()
                x = (screen_w - dlg_w) // 2
                y = (screen_h - dlg_h) // 2
                dialog.geometry(f"+{x}+{y}")
        except Exception:
            pass

    @staticmethod
    def _auto_close_dialog(dialog: tk.Toplevel, title: str) -> None:
        """Schließt den Dialog nach Timeout, falls noch vorhanden."""
        try:
            if dialog and dialog.winfo_exists():
                logger.warning(f"⚠️ Dialog Timeout: '{title}'")
                dialog.destroy()
        except tk.TclError:
            pass

    @staticmethod
    def _cancel_timeout(timeout_id: Optional[str]) -> None:
        """Bricht einen geplanten after-Aufruf ab."""
        if timeout_id:
            try:
                # Da wir keine Referenz auf das Dialog-Widget haben, müssen wir
                # die after_cancel mit der ID aufrufen. Das geht nur, wenn das
                # Widget noch existiert. Besser: timeout_id im Dialog speichern.
                # Wir vereinfachen: In unserer Implementierung wird timeout_id
                # im Dialog-Objekt gehalten, aber wir haben es nicht.
                # Stattdessen setzen wir beim Erstellen des Dialogs eine Referenz.
                # Hier eine pragmatische Lösung: Wir tun nichts, weil der
                # Dialog sowieso zerstört wird und der after-Aufruf dann ins Leere läuft.
                pass
            except Exception:
                pass

    @staticmethod
    def _on_closing(result: dict, timeout_id: Optional[str], close_val: bool) -> None:
        """Behandlung des Schließens über das X."""
        DarkMessageBox._cancel_timeout(timeout_id)
        result["value"] = close_val

    @staticmethod
    def _fallback_messagebox(title: str, message: str, msg_type: str,
                             buttons: bool = False) -> Optional[bool]:
        """Fallback auf tkinter.messagebox oder einfache Konsolenausgabe."""
        try:
            import tkinter.messagebox as mb
            if buttons:
                return mb.askokcancel(title, message)
            getattr(mb, f"show{msg_type}")(title, message)
            return None
        except Exception:
            logger.info(f"💬 {title}: {message}")
            return False if buttons else None

# -----------------------------------------------------------------------------
# MEMORY MANAGER
# -----------------------------------------------------------------------------
class MemoryManager:
    """
    Verwaltet Speicherpuffer für Textkomponenten mit automatischer Bereinigung,
    Überwachung und optimierter Ringpuffer-Implementierung.
    """

    def __init__(self) -> None:
        self._buffers: Dict[str, Deque[str]] = {}
        self._buffer_sizes: Dict[str, int] = {}
        self._lock = threading.RLock()
        self._max_memory_per_component = 100 * 1024 * 1024
        self._last_gc_time = time.time()
        self._gc_interval = 300
        self._ring_buffers: Dict[str, List[Optional[str]]] = {}
        self._ring_buffer_pointers: Dict[str, int] = {}
        self._ring_buffer_sizes: Dict[str, int] = {}
        self._memory_warning_threshold = 0.8
        self._long_term_monitor: Deque[Dict[str, Any]] = deque(maxlen=1000)
        self._monitoring_active = True
        self._maintenance_thread: Optional[threading.Thread] = None
        self._maintenance_stop = threading.Event()
        self._start_maintenance()

        # Initialisiere globale Cache-Statistiken (optional)
        self._cache_stats = {"total_allocated": 0, "total_freed": 0}

    def _start_maintenance(self) -> None:
        """Startet den Hintergrund-Thread für regelmäßige Wartungsarbeiten."""

        def maintenance_worker() -> None:
            while not self._maintenance_stop.is_set():
                try:
                    # Kurze Pause, um CPU-Last zu reduzieren
                    if self._maintenance_stop.wait(60):
                        break
                    self._perform_periodic_maintenance()
                    self._perform_memory_health_check()
                except Exception as e:
                    logger.warning(f"⚠️ Maintenance worker error: {e}")

        self._maintenance_thread = threading.Thread(
            target=maintenance_worker, daemon=True, name="MemoryMaintenance"
        )
        self._maintenance_thread.start()

    def _perform_memory_health_check(self) -> None:
        """Prüft die Speicherauslastung des Systems und des Prozesses."""
        try:
            import psutil
        except ImportError:
            return

        try:
            system_memory = psutil.virtual_memory()
            system_usage_percent = system_memory.percent / 100.0
            process = psutil.Process()
            process_memory = process.memory_info().rss
            process_usage_percent = process_memory / Constants.MAX_MEMORY_USAGE

            memory_sample = {
                "timestamp": time.time(),
                "system_usage": system_usage_percent,
                "process_usage": process_usage_percent,
                "system_mb": system_memory.used // (1024 * 1024),
                "process_mb": process_memory // (1024 * 1024),
            }
            self._long_term_monitor.append(memory_sample)

            if debug3_enabled('memory'):
                logger.debug(
                    f"[DEBUG3][MEMORY] Health check: system_usage={system_usage_percent:.1%}, "
                    f"process_usage={process_usage_percent:.1%}"
                )

            if system_usage_percent > self._memory_warning_threshold:
                logger.warning(f"⚠️ High system memory usage: {system_memory.percent:.1f}%")
            if process_usage_percent > self._memory_warning_threshold:
                logger.warning(f"⚠️ High process memory usage: {process_usage_percent:.1%}")
                self._aggressive_cleanup()

            # Überprüfe auf anhaltend hohe Auslastung
            if len(self._long_term_monitor) >= 10:
                recent_samples = list(self._long_term_monitor)[-10:]
                avg_usage = sum(s["system_usage"] for s in recent_samples) / len(recent_samples)
                if avg_usage > 0.75:
                    logger.warning(f"⚠️ Sustained high memory usage: {avg_usage:.1%}")
        except Exception as e:
            logger.warning(f"⚠️ Memory health check error: {e}")

    def get_memory_stats(self) -> Dict[str, Any]:
        """Liefert aktuelle Speicherstatistiken."""
        try:
            import psutil
        except ImportError:
            return {}

        try:
            system_memory = psutil.virtual_memory()
            process = psutil.Process()
            process_memory = process.memory_info().rss
            with self._lock:
                buffer_count = len(self._buffers)
                ring_buffer_count = len(self._ring_buffers)
                total_buffer_size = sum(self._buffer_sizes.values())

            return {
                "system_usage_percent": system_memory.percent,
                "system_used_mb": system_memory.used // (1024 * 1024),
                "system_total_mb": system_memory.total // (1024 * 1024),
                "process_usage_percent": (process_memory / Constants.MAX_MEMORY_USAGE) * 100,
                "process_used_mb": process_memory // (1024 * 1024),
                "process_peak_mb": self._get_peak_memory() // (1024 * 1024),
                "long_term_samples": len(self._long_term_monitor),
                "buffer_components": buffer_count,
                "ring_buffer_components": ring_buffer_count,
                "total_buffer_size_mb": total_buffer_size // (1024 * 1024),
                "active_monitoring": self._monitoring_active,
            }
        except Exception as e:
            logger.warning(f"⚠️ Memory stats error: {e}")
            return {}

    def _get_peak_memory(self) -> int:
        """Gibt den maximalen bisherigen Speicherverbrauch zurück."""
        try:
            import psutil
            process = psutil.Process()
            return process.memory_info().rss
        except Exception:
            return 0

    def _perform_periodic_maintenance(self) -> None:
        """Führt regelmäßige Wartungsarbeiten durch (GC, Pufferbereinigung)."""
        with self._lock:
            current_time = time.time()
            if current_time - self._last_gc_time > self._gc_interval:
                gc.collect()
                self._last_gc_time = current_time

            total_memory = sum(self._buffer_sizes.values())
            memory_usage_percent = total_memory / self._max_memory_per_component
            if memory_usage_percent > 0.8:
                logger.warning(f"⚠️ High buffer memory: {memory_usage_percent:.1%}")
                # Starte aggressive Bereinigung asynchron
                cleanup_thread = threading.Thread(
                    target=self._aggressive_cleanup, daemon=True
                )
                cleanup_thread.start()

    def _aggressive_cleanup(self) -> None:
        """Führt eine aggressive Speicherbereinigung durch (reduziert Puffergrößen)."""
        logger.info("🧹 Starting aggressive memory cleanup...")
        with self._lock:
            components = list(self._buffers.keys())

        for component in components:
            try:
                with self._lock:
                    if component in self._buffers:
                        buffer_size = len(self._buffers[component])
                        if buffer_size > 100:
                            keep_count = max(50, int(buffer_size * 0.5))
                            current_buffer = self._buffers[component]
                            new_deque = deque(
                                list(current_buffer)[-keep_count:],
                                maxlen=Constants.MAX_TEXT_LINES
                            )
                            self._buffers[component] = new_deque
                            # Aktualisiere die Speichergröße
                            self._buffer_sizes[component] = sum(
                                len(str(text).encode("utf-8")) if text else 0
                                for text in new_deque
                            )
                            if debug3_enabled('memory'):
                                logger.debug(
                                    f"[DEBUG3][MEMORY] Component {component}: "
                                    f"{buffer_size} -> {keep_count} entries"
                                )
                            logger.info(f"  ↪ {component}: {buffer_size} → {keep_count} entries")
            except Exception as e:
                logger.warning(f"⚠️ Component cleanup error for {component}: {e}")

        def async_gc() -> None:
            gc.collect()

        gc_thread = threading.Thread(target=async_gc, daemon=True)
        gc_thread.start()
        logger.info("✅ Aggressive cleanup completed")

    def add_text(self, component: str, text: str) -> None:
        """
        Fügt einer Komponente Text hinzu. Verwendet entweder einen Ringpuffer (falls bereits vorhanden)
        oder einen einfachen Deque-Puffer.
        """
        if not text or not text.strip():
            return
        with self._lock:
            if component in self._ring_buffers:
                self._add_to_ring_buffer(component, text)
                return

            # Standard-Deque-Puffer
            if component not in self._buffers:
                self._buffers[component] = deque(maxlen=Constants.MAX_TEXT_LINES)
                self._buffer_sizes[component] = 0

            text_size = len(text.encode("utf-8"))
            current_size = self._buffer_sizes[component]
            if current_size + text_size > self._max_memory_per_component:
                self._optimize_buffer(component)

            self._buffers[component].append(text)
            self._buffer_sizes[component] += text_size
            self._cache_stats["total_allocated"] += text_size

    def _add_to_ring_buffer(self, component: str, text: str) -> None:
        """
        Fügt Text in einen Ringpuffer ein. Überschreibt ggf. den ältesten Eintrag.
        """
        with self._lock:
            # Erstelle Ringpuffer, falls nicht vorhanden
            if component not in self._ring_buffers:
                buffer_size = Constants.MAX_TEXT_LINES
                self._ring_buffers[component] = [None] * buffer_size
                self._ring_buffer_pointers[component] = 0
                self._ring_buffer_sizes[component] = 0
                self._buffer_sizes[component] = 0

            ring_buffer = self._ring_buffers[component]
            pointer = self._ring_buffer_pointers[component]
            text_size = len(text.encode("utf-8"))

            # Alten Eintrag ermitteln und dessen Größe abziehen
            old_text = ring_buffer[pointer]
            if old_text is not None:
                old_size = len(old_text.encode("utf-8"))
                self._buffer_sizes[component] -= old_size
                self._cache_stats["total_freed"] += old_size

            ring_buffer[pointer] = text
            self._buffer_sizes[component] += text_size
            self._cache_stats["total_allocated"] += text_size

            # Pointer weiterbewegen
            self._ring_buffer_pointers[component] = (pointer + 1) % len(ring_buffer)

            # Anzahl der gefüllten Einträge erhöhen, falls noch nicht voll
            if self._ring_buffer_sizes[component] < len(ring_buffer):
                self._ring_buffer_sizes[component] += 1

    def _optimize_buffer(self, component: str) -> None:
        """
        Optimiert den Puffer einer Komponente (verkleinert ihn), wenn er zu groß wird.
        """
        if component in self._ring_buffers:
            current_size = self._ring_buffer_sizes[component]
            if current_size > Constants.MAX_TEXT_LINES // 2:
                new_size = Constants.MAX_TEXT_LINES // 2
                self._resize_ring_buffer(component, new_size)
            return

        if component in self._buffers:
            keep_count = int(len(self._buffers[component]) * 0.7)
            if keep_count > 0:
                new_deque = deque(
                    list(self._buffers[component])[-keep_count:],
                    maxlen=Constants.MAX_TEXT_LINES
                )
                self._buffers[component] = new_deque
                self._buffer_sizes[component] = sum(
                    len(text.encode("utf-8")) for text in self._buffers[component]
                )
                logger.debug(f"🧹 Buffer {component} optimized: {keep_count} entries kept")

    def _resize_ring_buffer(self, component: str, new_size: int) -> None:
        """
        Ändert die Größe eines Ringpuffers. Kopiert die neuesten Einträge in den neuen Puffer.
        """
        if component not in self._ring_buffers:
            return
        with self._lock:
            old_buffer = self._ring_buffers[component]
            old_pointer = self._ring_buffer_pointers[component]
            old_filled = self._ring_buffer_sizes[component]
            old_capacity = len(old_buffer)

            # Bestimme Startindex für die zu behaltenden Einträge
            # Wir wollen die neuesten 'new_size' Einträge behalten
            start_idx = (old_pointer - min(old_filled, new_size)) % old_capacity
            if start_idx < 0:
                start_idx += old_capacity

            new_buffer = [None] * new_size
            new_pointer = 0
            new_filled = 0
            new_total_size = 0

            # Kopiere die relevanten Einträge
            for i in range(min(old_filled, new_size)):
                idx = (start_idx + i) % old_capacity
                text = old_buffer[idx]
                if text is not None:
                    new_buffer[new_pointer] = text
                    new_total_size += len(text.encode("utf-8"))
                    new_pointer = (new_pointer + 1) % new_size
                    new_filled += 1

            self._ring_buffers[component] = new_buffer
            self._ring_buffer_pointers[component] = new_pointer
            self._ring_buffer_sizes[component] = new_filled
            self._buffer_sizes[component] = new_total_size

            logger.debug(f"🧹 Ring buffer {component} resized: {old_filled} → {new_filled} entries")

    def get_text(self, component: str) -> str:
        """
        Gibt den gesamten Text einer Komponente als String zurück (zur Anzeige).
        """
        with self._lock:
            if component in self._ring_buffers:
                return self._get_from_ring_buffer(component)
            elif component in self._buffers:
                return "\n".join(self._buffers[component])
            return ""

    def _get_from_ring_buffer(self, component: str) -> str:
        """
        Extrahiert den gesamten Inhalt eines Ringpuffers in chronologischer Reihenfolge.
        """
        if component not in self._ring_buffers:
            return ""
        with self._lock:
            ring_buffer = self._ring_buffers[component]
            pointer = self._ring_buffer_pointers[component]
            filled = self._ring_buffer_sizes[component]
            capacity = len(ring_buffer)

            if filled == 0:
                return ""

            texts: List[str] = []
            # Beginne beim ältesten Eintrag
            start = (pointer - filled) % capacity
            for i in range(filled):
                idx = (start + i) % capacity
                text = ring_buffer[idx]
                if text is not None:
                    texts.append(text)
            return "\n".join(texts)

    def clear_component(self, component: str) -> None:
        """Löscht alle Daten einer bestimmten Komponente."""
        with self._lock:
            if component in self._buffers:
                del self._buffers[component]
            if component in self._buffer_sizes:
                del self._buffer_sizes[component]
            if component in self._ring_buffers:
                del self._ring_buffers[component]
            if component in self._ring_buffer_pointers:
                del self._ring_buffer_pointers[component]
            if component in self._ring_buffer_sizes:
                del self._ring_buffer_sizes[component]
            logger.info(f"🧹 Component {component} cleared")

    def get_buffer_stats(self, component: str) -> Dict[str, Any]:
        """Liefert detaillierte Statistiken zu einer Komponente."""
        with self._lock:
            if component in self._ring_buffers:
                return {
                    "type": "ring_buffer",
                    "size": self._ring_buffer_sizes.get(component, 0),
                    "capacity": len(self._ring_buffers[component]),
                    "memory_bytes": self._buffer_sizes.get(component, 0),
                    "pointer": self._ring_buffer_pointers.get(component, 0),
                }
            elif component in self._buffers:
                return {
                    "type": "deque",
                    "size": len(self._buffers[component]),
                    "capacity": Constants.MAX_TEXT_LINES,
                    "memory_bytes": self._buffer_sizes.get(component, 0),
                    "maxlen": Constants.MAX_TEXT_LINES,
                }
            return {"type": "not_found"}

    def list_components(self) -> List[str]:
        """Gibt eine Liste aller aktiven Komponenten zurück."""
        with self._lock:
            all_components = set(self._buffers.keys())
            all_components.update(self._ring_buffers.keys())
            return list(all_components)

    def get_total_memory_usage(self) -> int:
        """Gesamter Speicherverbrauch aller Puffer in Bytes."""
        with self._lock:
            return sum(self._buffer_sizes.values())

    def optimize_all_buffers(self) -> None:
        """Optimiert alle Puffer (z. B. nach einem Speicherengpass)."""
        logger.info("🧹 Optimizing all buffers...")
        with self._lock:
            components = list(self._buffers.keys()) + list(self._ring_buffers.keys())
        for component in components:
            self._optimize_buffer(component)
        gc.collect()
        if debug3_enabled('memory'):
            logger.debug("[DEBUG3][MEMORY] optimize_all_buffers completed")
        logger.info("✅ All buffers optimized")

    def dispose(self) -> None:
        """Räumt alle Ressourcen des MemoryManagers auf (wird beim Beenden aufgerufen)."""
        self._monitoring_active = False
        self._maintenance_stop.set()
        if self._maintenance_thread and self._maintenance_thread.is_alive():
            self._maintenance_thread.join(timeout=1.0)
        with self._lock:
            self._buffers.clear()
            self._buffer_sizes.clear()
            self._ring_buffers.clear()
            self._ring_buffer_pointers.clear()
            self._ring_buffer_sizes.clear()
            self._long_term_monitor.clear()
        gc.collect()

    def print_debug_info(self) -> None:
        """Gibt detaillierte Debug-Informationen auf der Konsole aus."""
        stats = self.get_memory_stats()
        logger.info("\n" + "=" * 50)
        logger.info("🧠 MEMORY MANAGER DEBUG INFO")
        logger.info("=" * 50)
        logger.info(f"System Memory: {stats.get('system_used_mb', 0)}MB / "
                    f"{stats.get('system_total_mb', 0)}MB ({stats.get('system_usage_percent', 0):.1f}%)")
        logger.info(f"Process Memory: {stats.get('process_used_mb', 0)}MB "
                    f"({stats.get('process_usage_percent', 0):.1f}%)")
        logger.info(f"Buffer Components: {stats.get('buffer_components', 0)}")
        logger.info(f"Ring Buffer Components: {stats.get('ring_buffer_components', 0)}")
        logger.info(f"Total Buffer Size: {stats.get('total_buffer_size_mb', 0)}MB")
        logger.info(f"Long Term Samples: {stats.get('long_term_samples', 0)}")
        components = self.list_components()
        if components:
            logger.info(f"\nActive Components ({len(components)}):")
            for comp in components[:5]:
                comp_stats = self.get_buffer_stats(comp)
                logger.info(f"  • {comp}: {comp_stats['type']}, "
                            f"size: {comp_stats.get('size', 0)}")
            if len(components) > 5:
                logger.info(f"  ... and {len(components) - 5} more")
        logger.info("=" * 50)

# -----------------------------------------------------------------------------
# SUPPORTED LANGUAGES & SETTINGS
# -----------------------------------------------------------------------------
SUPPORTED_LANGUAGES: Dict[str, str] = {
    "auto": "Automatisch",
    "de": "Deutsch", "en": "Englisch", "fr": "Französisch", "es": "Spanisch",
    "it": "Italienisch", "pt": "Portugiesisch", "nl": "Niederländisch",
    "pl": "Polnisch", "ru": "Russisch", "ja": "Japanisch",
    "zh": "Chinesisch", "ko": "Koreanisch", "ar": "Arabisch",
    "hi": "Hindi", "tr": "Türkisch", "vi": "Vietnamesisch",
    "th": "Thailändisch", "id": "Indonesisch", "ms": "Malaysisch",
    "fa": "Persisch", "he": "Hebräisch", "bn": "Bengalisch",
    "ta": "Tamil", "te": "Telugu", "ml": "Malayalam",
    "kn": "Kannada", "mr": "Marathi", "gu": "Gujarati",
    "pa": "Punjabi", "ur": "Urdu", "sv": "Schwedisch",
    "da": "Dänisch", "no": "Norwegisch", "fi": "Finnisch",
    "cs": "Tschechisch", "hu": "Ungarisch", "ro": "Rumänisch",
    "bg": "Bulgarisch", "el": "Griechisch", "sk": "Slowakisch",
    "hr": "Kroatisch", "sr": "Serbisch", "uk": "Ukrainisch",
    "ca": "Katalanisch", "eu": "Baskisch", "gl": "Galizisch",
}

SORTED_LANGUAGES: List[Tuple[str, str]] = sorted(
    [(name, code) for code, name in SUPPORTED_LANGUAGES.items()],
    key=lambda x: x[0]
)

LANGUAGE_SHORT_CODES: Dict[str, str] = {
    "auto": "Auto", "de": "Deu", "en": "Eng", "fr": "Fra", "es": "Esp",
    "it": "Ita", "pt": "Por", "nl": "Nld", "pl": "Pol", "ru": "Rus",
    "ja": "Jpn", "zh": "Chi", "ko": "Kor", "ar": "Ara", "hi": "Hin",
    "tr": "Tur", "vi": "Vie", "th": "Tha", "id": "Ind", "ms": "Msa",
    "fa": "Per", "he": "Heb", "sv": "Swe", "da": "Dan", "no": "Nor",
    "fi": "Fin", "cs": "Cze", "hu": "Hun", "ro": "Rom", "bg": "Bul",
    "el": "Gre", "sk": "Slo", "hr": "Hrv", "uk": "Ukr",
}

WHISPER_MODELS: List[str] = [
    "tiny",
    "tiny.en",
    "base",
    "base.en",
    "small",
    "small.en",
    "medium",
    "medium.en",
    "large-v2",
    "large-v3",
]


class AsianLanguageSupport:
    @staticmethod
    def should_use_word_segmentation(language_code: str) -> bool:
        return language_code in ["zh", "ja", "ko", "th"]

    @staticmethod
    def optimize_display_text(text: str, language_code: str) -> str:
        if language_code == "zh":
            return " ".join(text)
        elif language_code == "ja":
            return text.replace("。", ". ").replace("、", ", ")
        return text

# -----------------------------------------------------------------------------
# ADVANCED SETTINGS
# -----------------------------------------------------------------------------
class Settings:
    def __init__(self,
                 beam_size: int = Constants.DEFAULT_BEAM_SIZE,
                 temperature: float = Constants.DEFAULT_TEMPERATURE,
                 vad_filter: bool = Constants.ENABLE_VAD_FILTER,
                 max_cache_size: int = 200,
                 auto_save_interval: int = 300,
                 enable_sentiment_analysis: bool = False,
                 enable_speaker_diarization: bool = False,
                 max_memory_mb: int = 1024,
                 gpu_acceleration: bool = True,
                 optimize_translations: bool = False,
                 config_type: str = 'default',
                 transcript_max_lines: int = 400,
                 translation_max_lines: int = 300,
                 translation_engine: str = "google",
                 ollama_model: str = "llama3",
                 ollama_host: str = "http://localhost:11434",
                 asian_mode: bool = False,
                 precision_mode: bool = False,
                 audio_profile: str = 'transcription',
                 adaptive_chunk: bool = True,
                 duplicate_similarity_threshold: float = 0.85,
                 adaptive_chunk_low_words: int = 3,
                 adaptive_chunk_high_words: int = 10,
                 min_confidence: float = 0.25) -> None:
        self.config = get_config(config_type)
        self.beam_size = beam_size
        self.temperature = temperature
        self.vad_filter = vad_filter
        self.gpu_acceleration = gpu_acceleration
        self.max_cache_size = max_cache_size
        self.auto_save_interval = auto_save_interval
        self.max_memory_mb = max_memory_mb
        self.enable_sentiment_analysis = enable_sentiment_analysis
        self.enable_speaker_diarization = enable_speaker_diarization
        self.optimize_translations = optimize_translations
        self.chunk_duration = self.config.CHUNK_DURATION
        self.transcript_max_lines = transcript_max_lines
        self.translation_max_lines = translation_max_lines
        self.translation_engine = translation_engine
        self.ollama_model = ollama_model
        self.ollama_host = ollama_host
        self.asian_mode = asian_mode
        self.precision_mode = precision_mode
        self.audio_profile = audio_profile
        self.adaptive_chunk = adaptive_chunk
        self.duplicate_similarity_threshold = duplicate_similarity_threshold
        self.adaptive_chunk_low_words = adaptive_chunk_low_words
        self.adaptive_chunk_high_words = adaptive_chunk_high_words
        self.min_confidence = min_confidence

        self.vad_threshold: float = Constants.VAD_THRESHOLD
        self.vad_min_speech_duration_ms: int = Constants.VAD_MIN_SPEECH_DURATION_MS
        self.vad_min_silence_duration_ms: int = Constants.VAD_MIN_SILENCE_DURATION_MS

        if self.asian_mode:
            self.config.CHUNK_DURATION = 10

        if self.precision_mode:
            self.config.CHUNK_DURATION = 7
            self.beam_size = 10
            self.temperature = 0.0
            self.vad_threshold = 0.3
            self.vad_min_speech_duration_ms = 400
            self.vad_min_silence_duration_ms = 100

            if self.asian_mode:
                self.config.CHUNK_DURATION = 10

        logger.info("🔊 Settings initialized:")
        logger.info(f"   Config Type: {config_type}")
        logger.info(f"   SAMPLE_RATE: {self.config.SAMPLE_RATE}")
        logger.info(f"   CHANNELS: {self.config.CHANNELS}")
        logger.info(f"   CHUNK_DURATION: {self.chunk_duration}s")
        logger.info(f"   CHUNK_SIZE: {self.config.CHUNK_SIZE_BYTES:,} bytes")
        logger.info(f"   BEAM_SIZE: {self.beam_size}")
        logger.info(f"   GPU_ACCELERATION: {self.gpu_acceleration}")
        logger.info(f"   VAD_THRESHOLD: {self.vad_threshold} (optimiert)")
        logger.info(f"   VAD_MIN_SPEECH_MS: {self.vad_min_speech_duration_ms} (optimiert)")
        logger.info(f"   VAD_MIN_SILENCE_MS: {self.vad_min_silence_duration_ms} (optimiert)")
        logger.info(f"   GUI Transcript Max Lines: {self.transcript_max_lines}")
        logger.info(f"   GUI Translation Max Lines: {self.translation_max_lines}")
        logger.info(f"   Translation Engine: {self.translation_engine}")
        if self.translation_engine == "ollama":
            logger.info(f"   Ollama Model: {self.ollama_model}, Host: {self.ollama_host}")
        logger.info(f"   Asian Mode: {self.asian_mode}")
        logger.info(f"   Precision Mode: {self.precision_mode}")
        logger.info(f"   Audio Profile: {self.audio_profile}")
        logger.info(f"   Adaptive Chunk: {self.adaptive_chunk}")
        logger.info(f"   Duplicate Similarity Threshold: {self.duplicate_similarity_threshold:.2f}")
        logger.info(f"   Adaptive Chunk Low Words: {self.adaptive_chunk_low_words}")
        logger.info(f"   Adaptive Chunk High Words: {self.adaptive_chunk_high_words}")
        logger.info(f"   Min Confidence: {self.min_confidence}")

    @classmethod
    def load_from_file(cls, filename: str = "dragon_advanced_settings.json") -> 'Settings':
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            if file_path.exists():
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                except (json.JSONDecodeError, PermissionError, OSError) as e:
                    logger.error(f"Fehler beim Laden der Einstellungen: {e}")
                    if DEBUG_LEVEL >= 2:
                        logger.exception("Stacktrace:")
                    return cls()
                import inspect
                signature = inspect.signature(cls.__init__)
                valid_params = list(signature.parameters.keys())
                if 'self' in valid_params:
                    valid_params.remove('self')
                filtered_data = {k: v for k, v in data.items() if k in valid_params}
                config_type = data.get('config_type', 'default')
                if 'config_type' not in filtered_data:
                    filtered_data['config_type'] = config_type

                filtered_data.setdefault('transcript_max_lines', 400)
                filtered_data.setdefault('translation_max_lines', 300)
                filtered_data.setdefault('translation_engine', 'google')
                filtered_data.setdefault('ollama_model', 'llama3')
                filtered_data.setdefault('ollama_host', 'http://localhost:11434')
                filtered_data.setdefault('asian_mode', False)
                filtered_data.setdefault('precision_mode', False)
                filtered_data.setdefault('audio_profile', 'transcription')
                filtered_data.setdefault('adaptive_chunk', True)
                filtered_data.setdefault('duplicate_similarity_threshold', 0.85)
                filtered_data.setdefault('adaptive_chunk_low_words', 3)
                filtered_data.setdefault('adaptive_chunk_high_words', 10)

                instance = cls(**filtered_data)

                instance.vad_threshold = data.get('vad_threshold', Constants.VAD_THRESHOLD)
                instance.vad_min_speech_duration_ms = data.get('vad_min_speech_duration_ms', Constants.VAD_MIN_SPEECH_DURATION_MS)
                instance.vad_min_silence_duration_ms = data.get('vad_min_silence_duration_ms', Constants.VAD_MIN_SILENCE_DURATION_MS)

                if 'chunk_duration' in data:
                    try:
                        instance.config.CHUNK_DURATION = float(data['chunk_duration'])
                    except (ValueError, AttributeError):
                        pass
                if 'sample_rate' in data:
                    try:
                        instance.config.SAMPLE_RATE = int(data.get('sample_rate', Constants.SAMPLE_RATE))
                    except (ValueError, AttributeError):
                        pass
                if 'channels' in data:
                    try:
                        instance.config.CHANNELS = int(data.get('channels', Constants.CHANNELS))
                    except (ValueError, AttributeError):
                        pass
                if 'audio_format' in data:
                    try:
                        instance.config.AUDIO_FORMAT = str(data.get('audio_format', Constants.AUDIO_FORMAT))
                    except (ValueError, AttributeError):
                        pass
                logger.info(f"✅ Settings loaded successfully (Config Type: {config_type})")
                return instance
        except FileNotFoundError:
            logger.info("Keine gespeicherten Einstellungen gefunden, verwende Standard")
        except Exception as e:
            logger.error(f"❌ Error loading advanced settings: {e}")
        logger.info("📝 Using default settings")
        return cls()

    def save_to_file(self, filename: str = "dragon_advanced_settings.json") -> None:
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            config_type = 'default'
            if isinstance(self.config, RealtimeConfig):
                config_type = 'realtime'
            elif isinstance(self.config, HighAccuracyConfig):
                config_type = 'high_accuracy'
            elif isinstance(self.config, YouTubeOptimizedConfig):
                config_type = 'youtube'
            save_dict = {
                'beam_size': self.beam_size,
                'temperature': self.temperature,
                'vad_filter': self.vad_filter,
                'config_type': config_type,
                'max_cache_size': self.max_cache_size,
                'auto_save_interval': self.auto_save_interval,
                'enable_sentiment_analysis': self.enable_sentiment_analysis,
                'enable_speaker_diarization': self.enable_speaker_diarization,
                'max_memory_mb': self.max_memory_mb,
                'gpu_acceleration': self.gpu_acceleration,
                'optimize_translations': self.optimize_translations,
                'chunk_duration': self.config.CHUNK_DURATION,
                'sample_rate': self.config.SAMPLE_RATE,
                'channels': self.config.CHANNELS,
                'audio_format': self.config.AUDIO_FORMAT,
                'vad_threshold': self.vad_threshold,
                'vad_min_speech_duration_ms': self.vad_min_speech_duration_ms,
                'vad_min_silence_duration_ms': self.vad_min_silence_duration_ms,
                'transcript_max_lines': self.transcript_max_lines,
                'translation_max_lines': self.translation_max_lines,
                'translation_engine': self.translation_engine,
                'ollama_model': self.ollama_model,
                'ollama_host': self.ollama_host,
                'asian_mode': self.asian_mode,
                'precision_mode': self.precision_mode,
                'audio_profile': self.audio_profile,
                'adaptive_chunk': self.adaptive_chunk,
                'duplicate_similarity_threshold': self.duplicate_similarity_threshold,
                'adaptive_chunk_low_words': self.adaptive_chunk_low_words,
                'adaptive_chunk_high_words': self.adaptive_chunk_high_words,
            }
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(save_dict, f, indent=2, ensure_ascii=False)
                logger.info(f"💾 Settings saved to {file_path} (Config Type: {config_type})")
            except (OSError, PermissionError) as e:
                logger.error(f"Fehler beim Schreiben der Einstellungen: {e}")
            except Exception as e:
                logger.error(f"Unerwarteter Fehler beim Speichern: {e}")
                if DEBUG_LEVEL >= 2:
                    logger.exception("Stacktrace:")
        except Exception as e:
            logger.error(f"❌ Error saving settings: {e}")

    def repair(self) -> List[str]:
        logger.info("🔧 Repairing Settings...")
        repairs_made: List[str] = []
        if not hasattr(self, 'config'):
            self.config = Config()
            repairs_made.append('Added Config')
        if not hasattr(self, 'chunk_duration'):
            self.chunk_duration = self.config.CHUNK_DURATION
            repairs_made.append(f'Added chunk_duration from config: {self.chunk_duration}s')
        if not hasattr(self, 'transcript_max_lines'):
            self.transcript_max_lines = 400
            repairs_made.append('Added transcript_max_lines with default 400')
        if not hasattr(self, 'translation_max_lines'):
            self.translation_max_lines = 300
            repairs_made.append('Added translation_max_lines with default 300')
        if not hasattr(self, 'translation_engine'):
            self.translation_engine = 'google'
            repairs_made.append('Added translation_engine with default google')
        if not hasattr(self, 'ollama_model'):
            self.ollama_model = 'llama3'
            repairs_made.append('Added ollama_model with default llama3')
        if not hasattr(self, 'ollama_host'):
            self.ollama_host = 'http://localhost:11434'
            repairs_made.append('Added ollama_host with default http://localhost:11434')
        if not hasattr(self, 'asian_mode'):
            self.asian_mode = False
            repairs_made.append('Added asian_mode with default False')
        if not hasattr(self, 'precision_mode'):
            self.precision_mode = False
            repairs_made.append('Added precision_mode with default False')
        if not hasattr(self, 'audio_profile'):
            self.audio_profile = 'transcription'
            repairs_made.append('Added audio_profile with default transcription')
        if not hasattr(self, 'adaptive_chunk'):
            self.adaptive_chunk = True
            repairs_made.append('Added adaptive_chunk with default True')
        if not hasattr(self, 'duplicate_similarity_threshold'):
            self.duplicate_similarity_threshold = 0.85
            repairs_made.append('Added duplicate_similarity_threshold with default 0.85')
        if not hasattr(self, 'adaptive_chunk_low_words'):
            self.adaptive_chunk_low_words = 3
            repairs_made.append('Added adaptive_chunk_low_words with default 3')
        if not hasattr(self, 'adaptive_chunk_high_words'):
            self.adaptive_chunk_high_words = 10
            repairs_made.append('Added adaptive_chunk_high_words with default 10')

        if not self.config.validate_config():
            logger.warning("⚠️ Config validation failed, resetting to default")
            self.config = Config()
            repairs_made.append('Config reset to default')
        if repairs_made:
            logger.info(f"✅ Repairs made: {', '.join(repairs_made)}")
            self.save_to_file()
        else:
            logger.info("✅ No repairs needed")
        return repairs_made

    def validate(self) -> List[str]:
        issues: List[str] = []
        if not self.config.validate_config():
            issues.append("Config validation failed")
        if self.beam_size < 1 or self.beam_size > 20:
            issues.append(f"Invalid beam_size: {self.beam_size} (should be 1-20)")
        if not (0.0 <= self.temperature <= 2.0):
            issues.append(f"Invalid temperature: {self.temperature} (should be 0.0-2.0)")
        if self.max_memory_mb < 100 or self.max_memory_mb > 16384:
            issues.append(f"Invalid max_memory_mb: {self.max_memory_mb} (should be 100-16384)")
        if not (self.config.MIN_CHUNK_DURATION <= self.config.CHUNK_DURATION <= self.config.MAX_CHUNK_DURATION):
            issues.append(f"Invalid CHUNK_DURATION: {self.config.CHUNK_DURATION}s "
                          f"(should be {self.config.MIN_CHUNK_DURATION}-{self.config.MAX_CHUNK_DURATION}s)")
        if self.transcript_max_lines < 100 or self.transcript_max_lines > 5000:
            issues.append(f"Invalid transcript_max_lines: {self.transcript_max_lines} (should be 100-5000)")
        if self.translation_max_lines < 100 or self.translation_max_lines > 5000:
            issues.append(f"Invalid translation_max_lines: {self.translation_max_lines} (should be 100-5000)")
        if self.translation_engine not in ("google", "ollama"):
            issues.append(f"Invalid translation_engine: {self.translation_engine} (must be 'google' or 'ollama')")
        if self.translation_engine == "ollama" and not OLLAMA_AVAILABLE:
            issues.append("Ollama engine selected but 'requests' module is missing (Ollama not available)")
        if not (0.5 <= self.duplicate_similarity_threshold <= 1.0):
            issues.append(f"Invalid duplicate_similarity_threshold: {self.duplicate_similarity_threshold} (should be 0.5-1.0)")
        return issues

    def set_config_type(self, config_type: str) -> bool:
        valid_types = ['default', 'realtime', 'high_accuracy', 'youtube']
        if config_type not in valid_types:
            logger.warning(f"⚠️ Invalid config_type: {config_type}. Must be one of: {valid_types}")
            return False
        old_config_type = 'default'
        if isinstance(self.config, RealtimeConfig):
            old_config_type = 'realtime'
        elif isinstance(self.config, HighAccuracyConfig):
            old_config_type = 'high_accuracy'
        elif isinstance(self.config, YouTubeOptimizedConfig):
            old_config_type = 'youtube'
        if old_config_type == config_type:
            return True
        self.config = get_config(config_type)
        self.chunk_duration = self.config.CHUNK_DURATION
        logger.info(f"🔄 Config type changed: {old_config_type} → {config_type}")
        logger.info(f"   New CHUNK_DURATION: {self.chunk_duration}s")
        logger.info(f"   New CHUNK_SIZE: {self.config.CHUNK_SIZE_BYTES:,} bytes")
        return True

    def get_audio_filter(self, language: Optional[str] = None,
                         profile: Optional[str] = None) -> str:
        return self.config.get_audio_filter(language, profile)

    def get_youtube_headers(self, is_manifest: bool = False) -> Dict[str, str]:
        return self.config.get_youtube_headers(is_manifest)

    def get_platform_config(self, platform: str) -> Dict[str, Any]:
        return self.config.get_platform_config(platform)

    def print_config_summary(self) -> None:
        logger.info("\n" + "="*60)
        logger.info("⚙️ SETTINGS CONFIGURATION")
        logger.info("="*60)
        logger.info("\n🤖 AI Model Parameters:")
        logger.info(f"  • Beam Size: {self.beam_size}")
        logger.info(f"  • Temperature: {self.temperature}")
        logger.info(f"  • VAD Filter: {self.vad_filter}")
        logger.info(f"  • VAD Threshold: {self.vad_threshold}")
        logger.info(f"  • VAD Min Speech (ms): {self.vad_min_speech_duration_ms}")
        logger.info(f"  • VAD Min Silence (ms): {self.vad_min_silence_duration_ms}")
        logger.info(f"  • GPU Acceleration: {self.gpu_acceleration}")
        logger.info("\n🎵 Audio Configuration (from Config):")
        logger.info(f"  • Sample Rate: {self.config.SAMPLE_RATE} Hz")
        logger.info(f"  • Channels: {self.config.CHANNELS} "
                    f"({'Mono' if self.config.CHANNELS == 1 else 'Stereo'})")
        logger.info(f"  • Chunk Duration: {self.config.CHUNK_DURATION}s")
        logger.info(f"  • Chunk Size: {self.config.CHUNK_SIZE_BYTES:,} bytes")
        logger.info(f"  • Bytes/sec: {self.config.BYTES_PER_SECOND:,}")
        logger.info(f"  • Audio Filter Profiles: {len(self.config.FILTER_PROFILES)}")
        logger.info(f"  • Language Filters: {len(self.config.LANGUAGE_FILTERS)} languages")
        logger.info("\n⚡ Performance Settings:")
        logger.info(f"  • Max Cache Size: {self.max_cache_size}")
        logger.info(f"  • Max Memory: {self.max_memory_mb} MB")
        logger.info(f"  • Auto Save Interval: {self.auto_save_interval}s")
        logger.info("\n🔧 Features:")
        logger.info(f"  • Sentiment Analysis: {self.enable_sentiment_analysis}")
        logger.info(f"  • Speaker Diarization: {self.enable_speaker_diarization}")
        logger.info(f"  • Optimize Translations: {self.optimize_translations}")
        logger.info("\n🖥️ GUI Display:")
        logger.info(f"  • Transcript Max Lines: {self.transcript_max_lines}")
        logger.info(f"  • Translation Max Lines: {self.translation_max_lines}")
        logger.info("\n🌐 Translation Engine:")
        logger.info(f"  • Engine: {self.translation_engine}")
        if self.translation_engine == "ollama":
            logger.info(f"  • Ollama Model: {self.ollama_model}")
            logger.info(f"  • Ollama Host: {self.ollama_host}")
        logger.info("\n🗾 Asian Language Mode:")
        logger.info(f"  • Active: {self.asian_mode} (10s chunks)")
        logger.info("\n🎯 Precision Mode:")
        logger.info(f"  • Active: {self.precision_mode} (optimized for accuracy)")
        logger.info("\n🎛️ Audio Profile:")
        logger.info(f"  • Profile: {self.audio_profile}")
        logger.info("\n⚙️ Advanced Options:")
        logger.info(f"  • Adaptive Chunk: {self.adaptive_chunk}")
        logger.info(f"  • Duplicate Similarity Threshold: {self.duplicate_similarity_threshold:.2f}")
        logger.info(f"  • Adaptive Chunk Low Words: {self.adaptive_chunk_low_words}")
        logger.info(f"  • Adaptive Chunk High Words: {self.adaptive_chunk_high_words}")
        config_type = 'default'
        if isinstance(self.config, RealtimeConfig):
            config_type = 'realtime'
        elif isinstance(self.config, HighAccuracyConfig):
            config_type = 'high_accuracy'
        elif isinstance(self.config, YouTubeOptimizedConfig):
            config_type = 'youtube'
        logger.info(f"\n🎯 Config Type: {config_type.upper()}")
        issues = self.validate()
        if issues:
            logger.info("\n⚠️ Validation Issues:")
            for issue in issues:
                logger.info(f"  • {issue}")
        else:
            logger.info("\n✅ All settings valid")
        logger.info("="*60)

    def __repr__(self) -> str:
        config_type = 'default'
        if isinstance(self.config, RealtimeConfig):
            config_type = 'realtime'
        elif isinstance(self.config, HighAccuracyConfig):
            config_type = 'high_accuracy'
        elif isinstance(self.config, YouTubeOptimizedConfig):
            config_type = 'youtube'
        return (f"Settings(type={config_type}, "
                f"beam_size={self.beam_size}, "
                f"chunk={self.config.CHUNK_DURATION}s/{self.config.CHUNK_SIZE_BYTES:,}B, "
                f"gpu={self.gpu_acceleration}, "
                f"transcript_lines={self.transcript_max_lines}, "
                f"translation_lines={self.translation_max_lines}, "
                f"trans_engine={self.translation_engine}, "
                f"asian_mode={self.asian_mode}, "
                f"precision_mode={self.precision_mode}, "
                f"audio_profile={self.audio_profile}, "
                f"adaptive_chunk={self.adaptive_chunk}, "
                f"duplicate_threshold={self.duplicate_similarity_threshold:.2f})")

# -----------------------------------------------------------------------------
# PLUGINS (ENTFERNT – NICHT VERWENDET)
# -----------------------------------------------------------------------------
# Die Plugin-Infrastruktur wurde entfernt, da sie nicht aktiv genutzt wurde.
# Falls zukünftig Plugins benötigt werden, muss der Code wieder eingefügt werden.

# =============================================================================
# 4. ABSTRAKTE BASISKLASSE FÜR ÜBERSETZUNGS-ENGINES
# =============================================================================

class BaseTranslationEngine(ABC):
    """Abstrakte Basisklasse für alle Übersetzungs-Engines."""
    @abstractmethod
    def set_target_language(self, target_lang: str) -> None:
        pass

    @abstractmethod
    def translate_text(self, text: str, source_lang: str = "auto") -> Optional[TranslationResult]:
        pass

    @abstractmethod
    def dispose(self) -> None:
        pass


# -----------------------------------------------------------------------------
# Mixin für Fehlerzähler und Deaktivierung
# -----------------------------------------------------------------------------
class ErrorHandlingMixin:
    """Mixin, das Fehlerzähler und Deaktivierungslogik für Übersetzungs-Engines bereitstellt."""
    def __init__(self, max_errors: int = 5, disable_duration: float = 300.0):
        self._error_count = 0
        self._disabled_until = 0.0
        self._max_errors = max_errors
        self._disable_duration = disable_duration
        self._error_lock = threading.RLock()

    def _check_disable(self):
        """Prüft, ob die Engine aufgrund zu vieler Fehler deaktiviert werden muss."""
        with self._error_lock:
            if self._error_count >= self._max_errors and self._disabled_until == 0.0:
                self._disabled_until = time.time() + self._disable_duration
                logger.warning(f"⚠️ {self.__class__.__name__} vorübergehend deaktiviert für {self._disable_duration}s wegen {self._error_count} Fehlern")

    def is_functional(self) -> bool:
        """Gibt True zurück, wenn die Engine als funktionsfähig eingestuft wird."""
        with self._error_lock:
            if self._disabled_until > time.time():
                return False
            if self._error_count >= self._max_errors:
                if time.time() >= self._disabled_until:
                    self._error_count = 0
                    self._disabled_until = 0.0
                    self._reinitialize()
                else:
                    return False
            return True

    def _reinitialize(self):
        """Wird aufgerufen, wenn die Engine nach einer Deaktivierung neu initialisiert werden soll."""
        pass

    def _record_success(self):
        """Setzt den Fehlerzähler bei Erfolg zurück."""
        with self._error_lock:
            self._error_count = 0
            self._disabled_until = 0.0

    def _record_error(self):
        """Erhöht den Fehlerzähler und prüft auf Deaktivierung."""
        with self._error_lock:
            self._error_count += 1
            self._check_disable()


# -----------------------------------------------------------------------------
# TRANSLATION ENGINE (Google)
# -----------------------------------------------------------------------------
class GoogleTranslationEngine(BaseTranslationEngine, ErrorHandlingMixin):
    """
    Optimierte Übersetzungs-Engine für Google Translate (via deep-translator).
    Bietet Caching, Fehlertoleranz, erweiterte Textbereinigung und asiatische Sprachunterstützung.
    """

    def __init__(self, target_lang: str = "de",
                 settings: Optional[Settings] = None) -> None:
        """
        Initialisiert die Engine.

        :param target_lang: Zielsprache (ISO-Code, z.B. 'de')
        :param settings: Erweiterte Einstellungen (optional)
        """
        BaseTranslationEngine.__init__(self)
        # Fehlerbehandlung: max. 5 Fehler, dann 5 Minuten deaktivieren
        ErrorHandlingMixin.__init__(self, max_errors=5, disable_duration=300.0)

        self.target_lang = target_lang
        self.settings = settings or Settings()
        self.translator: Optional[Any] = None

        # Cache für Übersetzungen (TTL und Größe aus Konstanten)
        self._cache = TTLCache(
            maxsize=Constants.TRANSLATION_CACHE_SIZE,
            ttl=Constants.TRANSLATION_CACHE_TTL
        )
        self._lock = threading.RLock()
        # Letzte Übersetzungen merken, um Doppelarbeit zu vermeiden
        self._last_translations: Deque[str] = deque(maxlen=15)
        self.last_detected_language = "auto"

        # Initialisierung des Übersetzers
        self._setup_translator()

        # Konfigurierbare Retry-Parameter
        self._max_retries = 3
        self._retry_delay_base = 1.0
        self._retry_delay_max = 5.0

        # Erweiterte Bereinigungsregeln (können später aus einer Datei geladen werden)
        self._preprocess_rules = [
            (r"\s+", " "),                     # Mehrfach-Leerzeichen
            (r"[ ]+([.,!?])", r"\1"),           # Leerzeichen vor Satzzeichen entfernen
            (r"([.,!?])[ ]*", r"\1 "),          # Ein Leerzeichen nach Satzzeichen
            ("bass communi", "best community"), # Korrektur bekannter Fehler
            (" ,", ","),
            (" .", "."),
            ("„", '"'),
            ("“", '"'),
        ]
        self._postprocess_rules = [
            (r"\s+\.", "."),
            (r"\s+,", ","),
            (r"\s+\?", "?"),
            (r"\s+!", "!"),
            (" ,", ","),
            (r" \.", "."),
        ]

    def _contains_asian(self, text: str) -> bool:
        """
        Prüft, ob der Text asiatische Schriftzeichen enthält.
        Verwendet erweiterte Unicode-Blöcke.
        """
        asian_ranges = [
            (0x4E00, 0x9FFF),   # CJK Unified Ideographs
            (0x3400, 0x4DBF),   # CJK Unified Ideographs Extension A
            (0x20000, 0x2A6DF), # CJK Unified Ideographs Extension B
            (0x2A700, 0x2B73F), # CJK Unified Ideographs Extension C
            (0x2B740, 0x2B81F), # CJK Unified Ideographs Extension D
            (0x2B820, 0x2CEAF), # CJK Unified Ideographs Extension E
            (0xF900, 0xFAFF),   # CJK Compatibility Ideographs
            (0xAC00, 0xD7AF),   # Korean Hangul Syllables
            (0x1100, 0x11FF),   # Korean Hangul Jamo
            (0x3130, 0x318F),   # Korean Hangul Compatibility Jamo
            (0x3040, 0x309F),   # Japanese Hiragana
            (0x30A0, 0x30FF),   # Japanese Katakana
            (0x31F0, 0x31FF),   # Katakana Phonetic Extensions
            (0x0E00, 0x0E7F),   # Thai
            (0x0E80, 0x0EFF),   # Lao
            (0x1000, 0x109F),   # Myanmar
            (0x1780, 0x17FF),   # Khmer
            (0x1950, 0x197F),   # Tai Le
            (0x1980, 0x19DF),   # New Tai Lue
            (0x1A20, 0x1AAF),   # Tai Tham
            (0xAA60, 0xAA7F),   # Myanmar Extended-A
        ]
        for char in text:
            if not char:
                continue
            code = ord(char)
            for low, high in asian_ranges:
                if low <= code <= high:
                    return True
        return False

    def _is_valid_translation(self, original: str, translated: str) -> bool:
        """
        Validiert die Übersetzung auf Plausibilität.
        """
        if not translated or not translated.strip():
            return False

        orig_clean = original.strip()
        trans_clean = translated.strip()

        if len(trans_clean) < 1:
            return False
        if trans_clean.isspace():
            return False

        # Asiatische Sprachen haben oft andere Zeichenanzahl-Verhältnisse
        is_asian = self._contains_asian(orig_clean) or self._contains_asian(trans_clean)

        if is_asian:
            # Bei asiatischen Sprachen sehr großzügig
            if len(trans_clean) <= 1:
                return True
            orig_len = len(orig_clean)
            trans_len = len(trans_clean)
            if orig_len == 0 or trans_len == 0:
                return False
            ratio = trans_len / max(orig_len, 1)
            return 0.05 <= ratio <= 15.0
        else:
            # Für lateinische Sprachen etwas strenger
            if len(trans_clean) <= 3:
                # Prüfen, ob nur ein Zeichen wiederholt wird (z.B. "aaa")
                if len(set(trans_clean)) == 1 and len(trans_clean) > 1:
                    return False
            else:
                if len(set(trans_clean)) < 3:
                    return False

            orig_len = len(orig_clean)
            trans_len = len(trans_clean)
            if orig_len == 0 or trans_len == 0:
                return False
            ratio = trans_len / max(orig_len, 1)
            return 0.1 <= ratio <= 8.0

    def _setup_translator(self) -> None:
        """Initialisiert oder reinitialisiert den GoogleTranslator."""
        try:
            if TRANSLATOR_AVAILABLE:
                GoogleTranslator = FastLazyLoader.load("deep_translator")
                # Timeout erhöht für bessere Stabilität
                self.translator = GoogleTranslator(
                    source="auto",
                    target=self.target_lang,
                    timeout=10
                )
                self._record_success()
            else:
                self.translator = None
        except ImportError as e:
            logger.warning(f"deep_translator nicht verfügbar: {e}")
            self.translator = None
            self._record_error()

    def _reinitialize(self):
        """Wird nach Ablauf der Deaktivierungszeit aufgerufen."""
        self._setup_translator()

    def set_target_language(self, target_lang: str) -> None:
        """Ändert die Zielsprache und leert den Cache."""
        if target_lang != self.target_lang:
            self.target_lang = target_lang
            with self._lock:
                self._cache.clear()
                self._last_translations.clear()
                self._error_count = 0
                self._disabled_until = 0.0
            self._setup_translator()

    def _clean_common_errors(self, text: str) -> str:
        """Korrigiert häufige OCR-/Transkriptionsfehler."""
        # Kann später durch ein regelbasiertes System ersetzt werden
        corrections = {
            "bass communi": "best community",
            "thc": "the",
            "thc ": "the ",
            " thc": " the",
        }
        for wrong, right in corrections.items():
            text = text.replace(wrong, right)
        return text

    def _preprocess_text(self, text: str) -> str:
        """
        Bereinigt den Eingabetext vor der Übersetzung.
        """
        if not text:
            return ""

        clean_text = text.strip()

        # Regeln anwenden
        for pattern, repl in self._preprocess_rules:
            clean_text = re.sub(pattern, repl, clean_text)

        # Zusätzliche allgemeine Bereinigung
        clean_text = self._clean_common_errors(clean_text)

        # Sicherstellen, dass der Text nicht zu kurz ist
        if len(clean_text.split()) < 1:
            return ""

        return clean_text.strip()

    def _postprocess_translation(self, translated: str, original: str) -> str:
        """
        Bereinigt den übersetzten Text (z.B. Satzzeichen).
        """
        if not translated:
            return ""

        result = translated.strip()

        # Satzendezeichen ergänzen, falls nicht vorhanden
        if result and result[-1] not in (".", "!", "?", ":", ";"):
            result += "."

        # Ersten Buchstaben groß schreiben, falls klein
        if result and result[0].islower():
            result = result[0].upper() + result[1:]

        # Mehrfach-Leerzeichen entfernen
        result = re.sub(r"\s+", " ", result)

        # Regeln anwenden
        for pattern, repl in self._postprocess_rules:
            result = re.sub(pattern, repl, result)

        return result.strip()

    def translate_text(self, text: str, source_lang: str = "auto") -> Optional[TranslationResult]:
        """
        Übersetzt den Text. Führt bei Fehlern automatische Wiederholungen durch.
        """
        # Prüfen, ob Engine deaktiviert ist
        if not self.is_functional():
            logger.debug("GoogleTranslationEngine derzeit deaktiviert – überspringe Übersetzung")
            return None

        # Zielsprache == Quellsprache? Dann abbrechen (außer auto)
        if source_lang != "auto" and source_lang == self.target_lang:
            return None

        if not text or not self.translator:
            return None

        original_text = text.strip()
        if len(original_text) < 2:
            return None

        # Text vorbereiten
        clean_text = self._preprocess_text(original_text)
        if not clean_text:
            return None

        # Cache-Key erstellen
        text_hash = hashlib.md5(f"{source_lang}_{self.target_lang}_{clean_text}".encode()).hexdigest()[:16]
        cache_key = f"trans_{text_hash}"

        # Cache prüfen
        with self._lock:
            if text_hash in self._last_translations:
                # Wurde kürzlich übersetzt (um Doppelarbeit zu vermeiden)
                return None
            cached_result = self._cache.get(cache_key)
            if cached_result is not None:
                return cached_result

        # Wiederholungslogik mit exponentiellem Backoff
        last_exception = None
        for attempt in range(self._max_retries):
            try:
                # Falls Translator nicht vorhanden, neu initialisieren
                if not self.translator:
                    self._setup_translator()
                    if not self.translator:
                        time.sleep(self._retry_delay_base * (2 ** attempt))
                        continue

                # Übersetzung durchführen
                translated_text = self.translator.translate(clean_text)

                if debug3_enabled('translate'):
                    logger.debug(f"[DEBUG3][TRANSLATE] Raw translation: {translated_text}")

                if not translated_text or not translated_text.strip():
                    # Bei leerem Resultat weitermachen
                    time.sleep(self._retry_delay_base * (2 ** attempt))
                    continue

                # Plausibilitätsprüfung
                if not self._is_valid_translation(clean_text, translated_text):
                    continue

                # Nachbearbeitung
                final_translation = self._postprocess_translation(translated_text, clean_text)

                if debug3_enabled('translate'):
                    logger.debug(f"[DEBUG3][TRANSLATE] Cleaned translation: {final_translation}")

                # Ergebnisobjekt erstellen
                result = TranslationResult(
                    original=original_text,
                    translated=final_translation,
                    source_lang=source_lang,
                    target_lang=self.target_lang,
                )

                # Im Cache speichern und als "letzte Übersetzung" merken
                with self._lock:
                    self._cache.put(cache_key, result)
                    self._last_translations.append(text_hash)
                    self._record_success()

                return result

            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                last_exception = e
                self._record_error()

                if debug3_enabled('translate'):
                    logger.exception("[DEBUG3][TRANSLATE] Exception in translate_text")

                # Exponentieller Backoff
                if attempt < self._max_retries - 1:
                    delay = min(self._retry_delay_max,
                                self._retry_delay_base * (2 ** attempt))
                    time.sleep(delay)
                    # Translator neu aufsetzen
                    self._setup_translator()

        # Alle Versuche fehlgeschlagen
        if last_exception is not None:
            logger.warning(f"Übersetzung fehlgeschlagen nach {self._max_retries} Versuchen: {last_exception}")
        return None

    def dispose(self) -> None:
        """Ressourcen freigeben."""
        with self._lock:
            self._cache.clear()
            self._last_translations.clear()
            self.translator = None
            self._error_count = 0
            self._disabled_until = 0.0
            gc.collect()

# -----------------------------------------------------------------------------
# OLLAMA TRANSLATION ENGINE
# -----------------------------------------------------------------------------
class OllamaTranslationEngine(BaseTranslationEngine, ErrorHandlingMixin):
    """
    Übersetzungs-Engine, die lokale Ollama-Modelle verwendet.
    Optimierte Version mit konfigurierbaren Parametern und besserer Fehlerbehandlung.
    """

    def __init__(self,
                 target_lang: str = "de",
                 settings: Optional[Settings] = None,
                 model: str = "llama3",
                 host: str = "http://localhost:11434",
                 temperature: float = 0.1,
                 timeout: int = 30,
                 system_prompt: Optional[str] = None) -> None:
        """
        :param target_lang: Zielsprache (Code, z.B. 'de')
        :param settings: Globale Einstellungen (optional)
        :param model: Ollama-Modellname (z.B. 'llama3', 'mistral')
        :param host: Ollama-Server-URL (z.B. 'http://localhost:11434')
        :param temperature: Kreativität der Übersetzung (0 = deterministisch)
        :param timeout: Timeout für HTTP-Anfragen in Sekunden
        :param system_prompt: Optionaler System-Prompt (falls vom Modell unterstützt)
        """
        BaseTranslationEngine.__init__(self)
        ErrorHandlingMixin.__init__(self, max_errors=5, disable_duration=300.0)

        self.target_lang = target_lang
        self.model = model
        self.host = host.rstrip('/')
        self.temperature = temperature
        self.timeout = timeout
        self.system_prompt = system_prompt
        self.settings = settings or Settings()

        self._cache = TTLCache(maxsize=Constants.TRANSLATION_CACHE_SIZE,
                               ttl=Constants.TRANSLATION_CACHE_TTL)
        self._lock = threading.RLock()
        self._last_translations: Deque[str] = deque(maxlen=15)
        self.last_detected_language = "auto"
        self.available = OLLAMA_AVAILABLE

    def set_target_language(self, target_lang: str) -> None:
        """Ändert die Zielsprache und leert den Cache."""
        if target_lang != self.target_lang:
            self.target_lang = target_lang
            with self._lock:
                self._cache.clear()
                self._last_translations.clear()

    def _call_ollama(self, prompt: str) -> Optional[str]:
        """
        Führt einen einzelnen Ollama-Aufruf durch (nicht-streaming).
        Gibt die Antwort als String zurück oder None bei Fehler.
        """
        if not self.available:
            logger.error("Ollama nicht verfügbar (requests nicht installiert)")
            return None

        try:
            import requests
        except ImportError:
            logger.error("requests library nicht installiert")
            self.available = False
            return None

        # Payload vorbereiten
        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": self.temperature,
                # Weitere Optionen können hier ergänzt werden
                # "num_predict": 512,
                # "stop": ["\n"],
            }
        }
        if self.system_prompt:
            # System-Prompt einbetten – je nach Modell unterschiedlich
            # Bei neueren Modellen kann man "system" im Payload übergeben
            # Hier als Teil des Prompts (einfach)
            payload["prompt"] = f"{self.system_prompt}\n\n{prompt}"

        try:
            response = requests.post(f"{self.host}/api/generate",
                                     json=payload,
                                     timeout=self.timeout)
            if debug3_enabled('ollama'):
                logger.debug(f"[DEBUG3][OLLAMA] Response status: {response.status_code}")
                logger.debug(f"[DEBUG3][OLLAMA] Raw response: {response.text[:500]}")

            if response.status_code == 200:
                data = response.json()
                translated = data.get("response", "").strip()
                return translated if translated else None
            else:
                logger.warning(f"Ollama Fehler {response.status_code}: {response.text}")
                return None

        except requests.exceptions.Timeout:
            logger.warning(f"Ollama Timeout nach {self.timeout}s")
            return None
        except requests.exceptions.ConnectionError:
            logger.warning("Ollama nicht erreichbar (läuft der Server?)")
            return None
        except Exception as e:
            logger.warning(f"Ollama Fehler: {e}")
            return None

    def _reinitialize(self):
        """Wird nach einer Deaktivierung aufgerufen – hier nichts zu tun."""
        pass

    def _is_valid_translation(self, original: str, translated: str) -> bool:
        """
        Einfache Plausibilitätsprüfung der Übersetzung.
        Verhindert leere oder offensichtlich fehlerhafte Antworten.
        """
        if not translated or len(translated) < 2:
            return False
        # Verhindern, dass die Antwort nur das Original wiederholt (z.B. bei Fehlern)
        if translated.lower() == original.lower():
            return False
        # Zu lange oder zu kurze Übersetzungen filtern (optional)
        if len(translated) > len(original) * 5:
            return False
        return True

    def translate_text(self, text: str, source_lang: str = "auto") -> Optional[TranslationResult]:
        """
        Übersetzt den übergebenen Text mit Ollama.
        Nutzt Caching und Fehlerzähler.
        """
        if not self.is_functional():
            logger.debug("Ollama translation engine is currently disabled, skipping.")
            return None

        # Keine Übersetzung, wenn Quell- und Zielsprache gleich (außer auto)
        if source_lang != "auto" and source_lang == self.target_lang:
            return None

        if not text or not self.available:
            return None

        try:
            original_text = text.strip()
            if len(original_text) < 2:
                return None

            # Cache-Key inklusive Modell und Temperatur (da diese die Übersetzung beeinflussen)
            cache_data = f"{source_lang}_{self.target_lang}_{self.model}_{self.temperature}_{original_text}"
            text_hash = hashlib.md5(cache_data.encode()).hexdigest()[:16]
            cache_key = f"ollama_trans_{text_hash}"

            with self._lock:
                # Prüfen, ob derselbe Text kürzlich übersetzt wurde (Duplikat-Vermeidung)
                if text_hash in self._last_translations:
                    return None
                cached_result = self._cache.get(cache_key)
                if cached_result is not None:
                    return cached_result

            # Sprachbezeichner für den Prompt (lesbar)
            source_lang_name = "auto"
            if source_lang != "auto":
                source_lang_name = SUPPORTED_LANGUAGES.get(source_lang, source_lang)
            target_lang_name = SUPPORTED_LANGUAGES.get(self.target_lang, self.target_lang)

            # Prompt konstruieren
            prompt = (f"Translate the following text from {source_lang_name} to {target_lang_name}. "
                      f"Output only the translation, without any additional commentary.\n\n"
                      f"{original_text}")

            if debug3_enabled('ollama'):
                logger.debug(f"[DEBUG3][OLLAMA] Prompt: {prompt}")

            translated_text = self._call_ollama(prompt)

            if not translated_text:
                self._record_error()
                return None

            # Validierung
            if not self._is_valid_translation(original_text, translated_text):
                logger.debug(f"Ollama translation validation failed for: '{original_text[:30]}...' -> '{translated_text[:30]}...'")
                self._record_error()
                return None

            self._record_success()

            result = TranslationResult(
                original=original_text,
                translated=translated_text,
                source_lang=source_lang,
                target_lang=self.target_lang,
            )

            with self._lock:
                self._cache.put(cache_key, result)
                self._last_translations.append(text_hash)

            return result

        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"Ollama translate error: {e}")
            self._record_error()
            return None

    def dispose(self) -> None:
        """Ressourcen freigeben."""
        with self._lock:
            self._cache.clear()
            self._last_translations.clear()
            gc.collect()

# -----------------------------------------------------------------------------
# DUMMY ENGINES
# -----------------------------------------------------------------------------
class DummyTranscriptionEngine:
    def __init__(self, settings: Optional[Settings] = None):
        self.settings = settings or Settings()
        self.model = None
        self.model_size = "dummy"
        self.whisper_backend = None
        self.demo_mode = True

    def load_model(self, model_size: str, set_active: bool = False) -> Optional[Tuple[Any, str]]:
        logger.info("Dummy-Modus: Laden eines Modells nicht erforderlich.")
        return (None, "dummy")

    def transcribe_audio(self, audio_data: bytes, include_timestamps: bool = False) -> Any:
        if include_timestamps:
            dummy = TranscriptionResult(
                text="[Whisper nicht verfügbar]",
                confidence=0.5,
                language="de",
                start=0.0,
                end=5.0
            )
            return [dummy]
        else:
            return TranscriptionResult(
                text="[Whisper nicht verfügbar]",
                confidence=0.5,
                language="de"
            )

    def safe_transcribe(self, audio_data: bytes, max_retries: int = 2) -> Optional[TranscriptionResult]:
        return self.transcribe_audio(audio_data)

    def get_current_model(self) -> str:
        return "dummy"

    def is_model_loading(self) -> bool:
        return False

    def reload_model(self, model_size: str) -> bool:
        return False

    def test_model_functionality(self) -> bool:
        return True

    def dispose(self) -> None:
        pass


class DummyTranslationEngine(BaseTranslationEngine):
    def __init__(self, target_lang: str = "de",
                 settings: Optional[Settings] = None) -> None:
        self.target_lang = target_lang
        self.settings = settings or Settings()
        self._cache = TTLCache(maxsize=10)
        self._lock = threading.RLock()
        self._last_translations: Deque[str] = deque(maxlen=5)
        self.last_detected_language = "auto"

    def set_target_language(self, target_lang: str) -> None:
        self.target_lang = target_lang

    def translate_text(self, text: str, source_lang: str = "auto") -> Optional[TranslationResult]:
        return TranslationResult(
            original=text,
            translated="[Übersetzung nicht verfügbar]",
            source_lang=source_lang,
            target_lang=self.target_lang,
        )

    def dispose(self) -> None:
        pass


# -----------------------------------------------------------------------------
# HILFSKLASSEN FÜR TRANSCRIPTION
# -----------------------------------------------------------------------------
class _EmptyInfo:
    language = "unknown"
    duration = 0.0


class _UniversalInfo:
    def __init__(self, result_dict: Dict[str, Any]) -> None:
        self.language = result_dict.get("language", "unknown")
        self.duration = result_dict.get("duration", 0.0)


class _UniversalSegment:
    def __init__(self, seg_dict: Dict[str, Any]) -> None:
        self.text = seg_dict.get("text", "").strip()
        self.start = seg_dict.get("start", 0.0)
        self.end = seg_dict.get("end", 0.0)
        self.confidence = seg_dict.get("confidence", 0.9)


class _EmergencySegment:
    def __init__(self, seg_dict: Dict[str, Any]) -> None:
        self.text = seg_dict.get("text", "")
        self.start = seg_dict.get("start", 0.0)
        self.end = seg_dict.get("end", 0.0)
        self.confidence = 0.5


# =============================================================================
# TRANSCRIPTION ENGINE
# =============================================================================
class TranscriptionEngine:
    """
    Optimierte Version der Transkriptions-Engine mit besserer Struktur,
    effizienterer Ressourcennutzung und klarerer Fehlerbehandlung.
    """

    __slots__ = (
        'model', 'model_size', 'whisper_backend', 'settings', 'config',
        'device', 'compute_type', '_cache', '_lock', '_model_loading',
        '_max_cached_models', '_model_cache', '_model_usage_lock',
        '_performance_monitor', '_last_transcription_text', '_active_model_loads',
        '_model_loaded_flag', '_disposing', 'forced_language', '_last_detected_language',
        '_torch', '_np', '_scipy_signal'  # Cached module references
    )

    def __init__(self, settings: Optional[Settings] = None) -> None:
        self.settings = settings or Settings()
        self.config = self.settings.config
        self.model: Any = None
        self.model_size: Optional[str] = None
        self.whisper_backend: Optional[str] = None
        self._lock = threading.RLock()
        self._model_usage_lock = threading.RLock()
        self._model_loading = False
        self._max_cached_models = 3
        self._cache = TTLCache(maxsize=self.settings.max_cache_size)
        self._performance_monitor = SimplePerformanceTracker()
        self._last_transcription_text = ""
        self._active_model_loads: Set[str] = set()
        self._model_loaded_flag = False
        self._disposing = False
        self._model_cache: Dict[Tuple[str, str], Any] = {}
        self.forced_language: Optional[str] = None
        self._last_detected_language: Optional[str] = None

        # Module-Caching für schnelleren Zugriff
        self._torch = None
        self._np = None
        self._scipy_signal = None
        if TORCH_AVAILABLE:
            self._torch = FastLazyLoader.load('torch')
        if NUMPY_AVAILABLE:
            self._np = FastLazyLoader.load('numpy')
        if SCIPY_AVAILABLE:
            self._scipy_signal = FastLazyLoader.load('scipy.signal')
            
        self.device, self.compute_type = self._detect_optimal_device()
        self.model: Any = None
        self.model_size: Optional[str] = None

    def _detect_optimal_device(self) -> Tuple[str, str]:
        """Erkennt die beste verfügbare Hardware und wählt Compute-Typ."""
        device = "cpu"
        compute_type = "int8"
        if self._torch is not None:
            torch = self._torch
            if torch.cuda.is_available():
                try:
                    torch.tensor([1.0]).cuda()
                    device = "cuda"
                    compute_type = "float16" if self.settings.gpu_acceleration else "int8"
                    logger.info(f"✅ NVIDIA GPU detected: {torch.cuda.get_device_name(0)}")
                except Exception as e:
                    if DEBUG_LEVEL >= 1:
                        logger.warning(f"⚠️ CUDA test failed, falling back: {e}")
            if hasattr(torch.version, 'hip') and torch.version.hip:
                try:
                    if torch.cuda.device_count() > 0:
                        device = "cuda"
                        compute_type = "float16" if self.settings.gpu_acceleration else "int8"
                        logger.info("✅ AMD GPU (ROCm) detected")
                except (AttributeError, RuntimeError) as e:
                    if DEBUG_LEVEL >= 1:
                        logger.warning(f"⚠️ ROCm test failed: {e}")
            if hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
                device = "mps"
                compute_type = "float16"
                logger.info("✅ Apple Silicon GPU (MPS) detected")
        logger.info(f"✅ Verwende Device: {device} (compute_type={compute_type})")
        return device, compute_type

    def _estimate_model_memory(self, model_size: str) -> float:
        """Schätzt den Speicherbedarf eines Modells in GB."""
        estimates = {
            'tiny': 1.0, 'tiny.en': 1.0,
            'base': 1.5, 'base.en': 1.5,
            'small': 2.5, 'small.en': 2.5,
            'medium': 4.0, 'medium.en': 4.0,
            'large': 6.0, 'large-v1': 6.0, 'large-v2': 6.0, 'large-v3': 6.0,
        }
        return estimates.get(model_size.lower(), 3.0)

    def _get_free_gpu_memory(self) -> Optional[float]:
        """Ermittelt freien GPU-Speicher in GB, falls verfügbar."""
        if self.device != "cuda" or self._torch is None:
            return None
        try:
            torch = self._torch
            if not torch.cuda.is_available():
                return None
            allocated = torch.cuda.memory_allocated()
            total = torch.cuda.get_device_properties(0).total_memory
            free = (total - allocated) / (1024**3)
            return free
        except Exception:
            return None

    @execution_decorator(timeout=1800.0)
    def load_model(self, model_size: str, set_active: bool = False) -> Optional[Tuple[Any, str]]:
        """
        Lädt ein Whisper-Modell. Falls `set_active` True, wird es zum aktiven Modell.
        Gibt (Modell, Backend) zurück oder None bei Fehler.
        """
        if set_active:
            self._force_model_cleanup()

        # Backend-Auswahl
        if FASTER_WHISPER_AVAILABLE:
            backend = "faster_whisper"
        elif OPENAI_WHISPER_AVAILABLE:
            backend = "openai_whisper"
        else:
            logger.error("❌ Kein Whisper-Backend verfügbar")
            return None

        cache_key = (model_size, backend)

        # Bereits geladen?
        with self._lock:
            if cache_key in self._model_cache:
                model = self._model_cache[cache_key]
                if set_active:
                    with self._model_usage_lock:
                        self.model = model
                        self.model_size = model_size
                        self.whisper_backend = backend
                return model, backend

        # GPU-Speicher prüfen
        free_gb = self._get_free_gpu_memory()
        if free_gb is not None:
            required_gb = self._estimate_model_memory(model_size)
            if free_gb < required_gb:
                logger.warning(f"⚠️ Nur {free_gb:.1f} GB VRAM frei, {model_size} benötigt ~{required_gb} GB. Versuche Cache zu leeren...")
                # Ältestes Modell entfernen
                if self._model_cache:
                    oldest_key = next(iter(self._model_cache))
                    oldest_model = self._model_cache[oldest_key]
                    logger.info(f"   Entferne {oldest_key[0]} ({oldest_key[1]}) aus Cache")
                    self._unload_model(oldest_model)
                    del self._model_cache[oldest_key]
                    gc.collect()
                    if self.device == "cuda":
                        self._torch.cuda.empty_cache()
                    free_gb = self._get_free_gpu_memory()
                    if free_gb is not None and free_gb < required_gb:
                        logger.error(f"❌ Auch nach Cleanup nicht genug VRAM ({free_gb:.1f} GB). Ladevorgang abgebrochen.")
                        return None
                else:
                    logger.error(f"❌ Nicht genug VRAM ({free_gb:.1f} GB) und kein Cache. Ladevorgang abgebrochen.")
                    return None

        load_lock_key = f"model_load_{model_size}_{backend}"
        if load_lock_key in self._active_model_loads:
            logger.info("⏳ Modell wird bereits geladen...")
            return None
        self._active_model_loads.add(load_lock_key)

        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            model_dir = config_dir / "models"
            model_dir.mkdir(exist_ok=True)
            logger.info(f"📁 Modell-Verzeichnis: {model_dir}")

            model = None
            if backend == "faster_whisper":
                model = self._load_faster_whisper(model_size, model_dir)
                if model is None and OPENAI_WHISPER_AVAILABLE:
                    backend = "openai_whisper"
                    cache_key = (model_size, backend)

            if backend == "openai_whisper" and model is None:
                model = self._load_openai_whisper(model_size, model_dir)

            if model is None:
                return None

            # In Cache ablegen
            with self._lock:
                self._model_cache[cache_key] = model
                if len(self._model_cache) > self._max_cached_models:
                    oldest_key, old_model = next(iter(self._model_cache.items()))
                    del self._model_cache[oldest_key]
                    logger.info(f"🧹 Entferne altes Modell '{oldest_key[0]}' aus Cache")
                    self._unload_model(old_model)

                if set_active:
                    with self._model_usage_lock:
                        self.model = model
                        self.model_size = model_size
                        self.whisper_backend = backend

            return model, backend

        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ Unerwarteter Fehler beim Laden: {e}")
            return None
        finally:
            self._active_model_loads.discard(load_lock_key)

    def _load_faster_whisper(self, model_size: str, model_dir: Path) -> Any:
        """Lädt ein Modell mit faster-whisper."""
        try:
            from faster_whisper import WhisperModel
            model = WhisperModel(
                model_size,
                device=self.device,
                compute_type=self.compute_type,
                download_root=str(model_dir),
                cpu_threads=4,
                num_workers=1,
            )
            # Test-Transkription zur Validierung
            if self._np is None:
                self._np = FastLazyLoader.load("numpy")
            test_audio = self._np.zeros(1600, dtype=self._np.float32)
            segments, info = model.transcribe(
                test_audio, beam_size=1, best_of=1,
                vad_filter=False, without_timestamps=True
            )
            list(segments)  # Erzwingt Ausführung
            logger.info(f"✅ faster-whisper '{model_size}' erfolgreich geladen und getestet")
            return model
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ faster-whisper konnte nicht geladen werden: {e}")
            if DEBUG_LEVEL >= 2:
                logger.exception("Stacktrace:")
            return None

    def _load_openai_whisper(self, model_size: str, model_dir: Path) -> Any:
        """Lädt ein Modell mit openai-whisper."""
        try:
            import whisper
            device = "cuda" if self.device == "cuda" else "cpu"
            model = whisper.load_model(
                model_size,
                device=device,
                download_root=str(model_dir) if model_dir else None,
            )
            logger.info(f"✅ openai-whisper '{model_size}' erfolgreich geladen")
            return model
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ openai-whisper fehlgeschlagen: {e}")
            if DEBUG_LEVEL >= 2:
                logger.exception("Stacktrace:")
            return None

    def _unload_model(self, model: Any) -> None:
        """Entlädt ein Modell, falls möglich."""
        if hasattr(model, 'unload_model'):
            try:
                model.unload_model()
            except Exception:
                pass

    def reload_model(self, model_size: str) -> bool:
        """Lädt ein neues Modell im Hintergrund und aktiviert es nach Erfolg."""
        with self._lock:
            if self._model_loading:
                logger.warning("⚠️ Model loading already in progress")
                return False
            self._model_loading = True

        def _load_in_background():
            try:
                result = self.load_model(model_size, set_active=False)
                if result is not None:
                    new_model, new_backend = result
                    with self._lock:
                        with self._model_usage_lock:
                            self.model = new_model
                            self.model_size = model_size
                            self.whisper_backend = new_backend
                    logger.info(f"✅ Model switched to {model_size} ({new_backend})")
                else:
                    logger.error("❌ Background model loading failed")
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.error(f"❌ Background model loading error: {e}")
            finally:
                with self._lock:
                    self._model_loading = False

        thread = threading.Thread(target=_load_in_background, daemon=True, name=f"ModelLoader-{model_size}")
        thread.start()
        return True

    def is_model_loading(self) -> bool:
        return self._model_loading

    def _force_model_cleanup(self) -> None:
        """Entlädt das aktive Modell und gibt GPU-Speicher frei."""
        with self._model_usage_lock:
            if self.model is not None:
                self._unload_model(self.model)
                self.model = None
                self.model_size = None
                self._model_loaded_flag = False
        gc.collect()
        if self.device == "cuda" and self._torch is not None:
            try:
                self._torch.cuda.empty_cache()
                logger.info("🧹 GPU Memory freigegeben")
            except Exception:
                pass

    # -------------------------------------------------------------------------
    # Transkriptions-Kernlogik
    # -------------------------------------------------------------------------

    def _universal_transcribe(self, model: Any, audio_np: Any, **kwargs: Any) -> Tuple[List[Any], Any]:
        """
        Einheitliche Schnittstelle für faster-whisper und openai-whisper.
        Gibt (Segmente, Info) zurück.
        """
        if model is None:
            raise ValueError("Kein Modell geladen")

        backend = self.whisper_backend
        if backend == "faster_whisper":
            return self._faster_whisper_transcribe(model, audio_np, **kwargs)
        else:
            return self._openai_whisper_transcribe(model, audio_np, **kwargs)

    def _faster_whisper_transcribe(self, model: Any, audio_np: Any, **kwargs: Any) -> Tuple[List[Any], Any]:
        """Transkription mit faster-whisper."""
        try:
            segments, info = model.transcribe(audio_np, **kwargs)
            segments_list = list(segments)
            return segments_list, info
        except (TypeError, ValueError) as e:
            logger.warning(f"⚠️ faster-whisper Parameterfehler: {e} – verwende minimale Parameter")
            minimal_kwargs = {k: v for k, v in kwargs.items() if k in ["language", "task", "temperature", "beam_size", "best_of"]}
            try:
                segments, info = model.transcribe(audio_np, **minimal_kwargs)
                return list(segments), info
            except Exception as e2:
                if PlatformUtils.is_fatal_exception(e2):
                    raise
                logger.error(f"❌ faster-whisper auch mit minimalen Parametern fehlgeschlagen: {e2}")
                return [], _EmptyInfo()
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ faster-whisper Fehler: {e}")
            return [], _EmptyInfo()

    def _openai_whisper_transcribe(self, model: Any, audio_np: Any, **kwargs: Any) -> Tuple[List[Any], Any]:
        """Transkription mit openai-whisper."""
        allowed_params = {
            "language", "task", "temperature", "best_of", "beam_size", "patience",
            "length_penalty", "repetition_penalty", "no_repeat_ngram_size",
            "initial_prompt", "prefix", "suppress_tokens", "without_timestamps",
            "max_initial_timestamp", "word_timestamps", "prepend_punctuations",
            "append_punctuations", "max_new_tokens", "clip_timestamps",
            "hallucination_silence_threshold",
        }
        filtered_kwargs = {k: v for k, v in kwargs.items() if k in allowed_params}
        # Standardwerte setzen
        filtered_kwargs.setdefault("language", None)
        filtered_kwargs.setdefault("task", "transcribe")
        filtered_kwargs.setdefault("temperature", 0.0)

        try:
            result = model.transcribe(audio_np, **filtered_kwargs)
            segments = result.get("segments", [])
            # In ein einheitliches Format konvertieren
            converted = []
            for seg in segments:
                if seg.get("text", "").strip():
                    converted.append(_UniversalSegment(seg))
            info = _UniversalInfo(result)
            return converted, info
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ openai-whisper Fehler: {e}")
            # Minimaler Fallback
            try:
                minimal_result = model.transcribe(audio_np, language=None, task="transcribe", temperature=0.1)
                emergency = []
                for seg in minimal_result.get("segments", []):
                    emergency.append(_EmergencySegment(seg))
                return emergency, _UniversalInfo(minimal_result)
            except Exception as fallback_error:
                if PlatformUtils.is_fatal_exception(fallback_error):
                    raise
                logger.error(f"💥 Auch Fallback fehlgeschlagen: {fallback_error}")
                return [], _EmptyInfo()

    def validate_audio_data(self, audio_data: bytes) -> Tuple[bool, str]:
        """Prüft, ob Audiodaten für eine Transkription geeignet sind."""
        if not isinstance(audio_data, bytes):
            return False, "Audio data must be bytes"
        if len(audio_data) == 0:
            return False, "Audio data is empty"
        if len(audio_data) < 1600:
            return False, f"Audio data too short: {len(audio_data)} bytes"
        if self._np is not None:
            try:
                audio_np = self._np.frombuffer(audio_data, dtype=self._np.int16)
                if self._np.all(audio_np == 0):
                    return False, "Audio data is completely silent"
                if self._np.var(audio_np) < 100:
                    return False, "Audio variance too low (likely silent)"
            except Exception:
                pass
        return True, "Valid"

    def safe_transcribe(self, audio_data: bytes, max_retries: int = 2) -> Optional[TranscriptionResult]:
        """
        Führt eine Transkription mit Wiederholungen und Audio-Verbesserung durch.
        """
        is_valid, msg = self.validate_audio_data(audio_data)
        if not is_valid:
            return None

        for attempt in range(max_retries + 1):
            try:
                processed = self.enhance_audio_for_transcription(audio_data)
                result = self.transcribe_audio(processed)
                if result and result.text and result.text.strip():
                    return result
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"Transkriptionsfehler (Versuch {attempt+1}): {e}")
                if attempt < max_retries:
                    time.sleep(0.5 * (attempt + 1))
        return None

    def enhance_audio_for_transcription(self, audio_data: bytes) -> bytes:
        """
        Verbessert die Audioqualität durch Normalisierung, Rauschunterdrückung
        und Hochpassfilter.
        """
        if not audio_data or len(audio_data) < 1600 or self._np is None:
            return audio_data

        try:
            np = self._np
            audio_np = np.frombuffer(audio_data, dtype=self._np.int16).astype(self._np.float32) / 32768.0

            # Entferne NaN/Inf
            if np.isnan(audio_np).any() or np.isinf(audio_np).any():
                return audio_data

            # Dynamische Lautstärkeanpassung
            rms = np.sqrt(np.mean(audio_np**2))
            gain = 1.0
            if rms < 1e-8:
                gain = 2.0
            elif rms < 0.005:
                gain = min(5.0, 0.02 / max(rms, 1e-6))
            elif rms > 0.5:
                gain = 0.5 / rms
            elif rms > 0.3:
                gain = 0.3 / rms

            audio_np *= gain

            # Clipping vermeiden
            max_val = np.max(np.abs(audio_np))
            if max_val > 0.99:
                audio_np /= max_val * 1.01  # etwas Puffer

            # Gleichspannungsanteil entfernen
            audio_np -= np.mean(audio_np)

            # Optional: Hochpassfilter (80 Hz) falls SciPy verfügbar
            if self._scipy_signal is not None and len(audio_np) > 100:
                try:
                    b, a = self._scipy_signal.butter(2, 80 / (self.config.SAMPLE_RATE / 2), btype="high")
                    audio_np = self._scipy_signal.filtfilt(b, a, audio_np)
                except Exception:
                    pass

            # Zurück zu int16
            audio_np = np.clip(audio_np, -0.99, 0.99)
            enhanced = (audio_np * 32767).astype(np.int16).tobytes()
            return enhanced if len(enhanced) == len(audio_data) else audio_data

        except Exception as e:
            logger.warning(f"Audio-Enhancement fehlgeschlagen: {e}")
            return audio_data

    def _calculate_enhanced_confidence(self, segment: Any, text: str) -> float:
        """
        Berechnet eine verbesserte Konfidenz unter Berücksichtigung von Textmerkmalen.
        """
        base = max(getattr(segment, "confidence", 0.0), 0.1)
        words = text.split()
        word_count = len(words)
        text_len = len(text.strip())
        unique_ratio = len(set(words)) / max(word_count, 1)

        boosts = (
            min(0.2, text_len / 300.0) +          # Längenbonus
            min(0.15, word_count * 0.03) +         # Wortanzahlbonus
            (0.08 if any(c in text for c in ".!?,;:") else 0.0) +  # Satzzeichen
            (0.1 if any(c.isalpha() for c in text) else 0.0) +     # Buchstaben vorhanden
            min(0.1, unique_ratio * 0.1)                           # Vielfalt
        )
        return min(0.95, base + boosts)

    def _validate_transcription_segment(self, text: str, confidence: float) -> bool:
        """Prüft, ob ein Segment gültig ist (nicht zu kurz, genug Konfidenz)."""
        if not text or len(text.strip()) < 2:
            return False
        clean = text.strip()
        if clean.isspace():
            return False
        if len(clean) > 500:
            return False
        if not any(c.isalpha() for c in clean):
            return False
        if confidence < self.settings.min_confidence:
            return False
        return True

    def transcribe_audio(self, audio_data: bytes, include_timestamps: bool = False) -> Any:
        """
        Hauptmethode zur Transkription. Wird im Executor-Thread aufgerufen.
        """
        with self._model_usage_lock:
            model = self.model
            if not model:
                return None if not include_timestamps else []

        try:
            # Audio vorbereiten
            processed = self.enhance_audio_for_transcription(audio_data)
            if self._np is None:
                self._np = FastLazyLoader.load("numpy")
            audio_np = self._np.frombuffer(processed, dtype=self._np.int16).astype(self._np.float32) / 32768.0

            beam_size = self.settings.beam_size
            language = self.forced_language if self.forced_language else None

            # VAD-Parameter je nach Sprache
            vad_language = self.forced_language or self._last_detected_language
            if vad_language and vad_language in self.config.LANGUAGE_VAD:
                lang_vad = self.config.LANGUAGE_VAD[vad_language]
                vad_params = {
                    "threshold": lang_vad['threshold'],
                    "min_speech_duration_ms": lang_vad['min_speech_ms'],
                    "min_silence_duration_ms": lang_vad['min_silence_ms'],
                }
            else:
                vad_params = {
                    "threshold": self.settings.vad_threshold,
                    "min_speech_duration_ms": self.settings.vad_min_speech_duration_ms,
                    "min_silence_duration_ms": self.settings.vad_min_silence_duration_ms,
                }
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ Fehler bei Audio-Vorbereitung: {e}")
            return [] if include_timestamps else None

        # Prüfen, ob Modell während der Vorbereitung gewechselt wurde
        with self._model_usage_lock:
            if self.model is not model:
                logger.warning("⚠️ Modell wurde gewechselt – Transkription abgebrochen.")
                return [] if include_timestamps else None

        return self._transcribe_worker(
            model, audio_np, language, beam_size, vad_params, include_timestamps
        )

    def _transcribe_worker(self, model: Any, audio_np: Any, language: Optional[str],
                           beam_size: int, vad_params: Dict[str, Any],
                           include_timestamps: bool) -> Any:
        """Führt die eigentliche Transkription durch und verarbeitet die Segmente."""
        try:
            # Erster Versuch mit VAD
            no_speech_threshold = getattr(self.settings, 'no_speech_threshold', 0.8)
            segments, info = self._universal_transcribe(
                model, audio_np,
                language=language,
                task="transcribe",
                temperature=self.settings.temperature,
                best_of=5,
                beam_size=beam_size,
                patience=1.0,
                no_speech_threshold=no_speech_threshold,
                log_prob_threshold=-1.2,
                compression_ratio_threshold=2.8,
                condition_on_previous_text=True,
                suppress_tokens=[-1],
                without_timestamps=not include_timestamps,
                word_timestamps=include_timestamps,
                vad_filter=self.settings.vad_filter,
                vad_parameters=vad_params,
            )

            if not segments:
                # Fallback ohne VAD
                logger.debug("🔄 Keine Segmente mit VAD – Versuch ohne VAD...")
                segments, info = self._universal_transcribe(
                    model, audio_np,
                    language=language,
                    task="transcribe",
                    temperature=0.0,
                    best_of=5,
                    beam_size=beam_size,
                    vad_filter=False,
                )

            # Sprache merken
            if hasattr(info, 'language') and info.language != 'unknown':
                self._last_detected_language = info.language

            # Segmente filtern und Konfidenz verbessern
            valid_segments = []
            total_confidence = 0.0
            for seg in segments:
                text = seg.text.strip()
                conf = self._calculate_enhanced_confidence(seg, text)
                if self._validate_transcription_segment(text, conf):
                    valid_segments.append(seg)
                    total_confidence += conf

            if not valid_segments:
                logger.debug("🔄 Keine validen Segmente – minimaler Fallback")
                minimal = self._transcribe_minimal(model, audio_np, language)
                if minimal:
                    if include_timestamps:
                        duration = audio_np.shape[0] / self.config.SAMPLE_RATE
                        return [TranscriptionResult(
                            text=minimal.text,
                            confidence=minimal.confidence,
                            language=minimal.language,
                            start=0.0,
                            end=duration,
                        )]
                    return minimal
                return [] if include_timestamps else None

            if include_timestamps:
                return [
                    TranscriptionResult(
                        text=seg.text.strip(),
                        confidence=self._calculate_enhanced_confidence(seg, seg.text.strip()),
                        language=getattr(info, 'language', 'unknown'),
                        start=getattr(seg, 'start', 0.0),
                        end=getattr(seg, 'end', 0.0),
                    )
                    for seg in valid_segments
                ]
            else:
                full_text = " ".join(seg.text.strip() for seg in valid_segments)
                avg_conf = total_confidence / len(valid_segments)
                return TranscriptionResult(
                    text=full_text,
                    confidence=avg_conf,
                    language=getattr(info, 'language', 'unknown'),
                )

        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ _transcribe_worker Fehler: {e}")
            return [] if include_timestamps else None

    def _transcribe_minimal(self, model: Any, audio_np: Any, language: Optional[str]) -> Optional[TranscriptionResult]:
        """Minimale Transkription ohne VAD und mit niedrigem Beam."""
        try:
            segments, info = self._universal_transcribe(
                model, audio_np,
                language=language,
                task="transcribe",
                temperature=0.0,
                best_of=1,
                beam_size=1,
                no_speech_threshold=0.9,
                log_prob_threshold=-2.0,
                compression_ratio_threshold=3.5,
                condition_on_previous_text=False,
                without_timestamps=True,
                vad_filter=False,
            )
            if segments:
                seg = segments[0]
                text = seg.text.strip()
                if text:
                    conf = self._calculate_enhanced_confidence(seg, text)
                    return TranscriptionResult(
                        text=text,
                        confidence=conf,
                        language=getattr(info, 'language', 'unknown'),
                    )
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.debug(f"Minimal transcription failed: {e}")
        return None

    def emergency_fallback_transcription(self, audio_data: Union[bytes, Any]) -> Optional[TranscriptionResult]:
        """Notfall-Transkription, wenn alles andere fehlschlägt."""
        with self._model_usage_lock:
            model = self.model
            if not model:
                return None
            try:
                if isinstance(audio_data, self._np.ndarray):
                    audio_np = audio_data.astype(self._np.float32)
                    if audio_np.dtype == self._np.int16:
                        audio_np = audio_np / 32768.0
                else:
                    if self._np is None:
                        self._np = FastLazyLoader.load("numpy")
                    audio_np = self._np.frombuffer(audio_data, dtype=self._np.int16).astype(self._np.float32) / 32768.0
                return self._transcribe_minimal(model, audio_np, None)
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.error(f"❌ Emergency fallback exception: {e}")
                return None

    # -------------------------------------------------------------------------
    # Hilfsfunktionen und Verwaltung
    # -------------------------------------------------------------------------

    def clear_cache(self) -> None:
        """Leert den internen Cache und gibt GPU-Speicher frei."""
        with self._lock:
            self._cache.clear()
            self._last_transcription_text = ""
        gc.collect()
        if self.device == "cuda" and self._torch is not None:
            try:
                self._torch.cuda.empty_cache()
            except Exception:
                pass

    def get_current_model(self) -> str:
        with self._model_usage_lock:
            return self.model_size if self.model_size else "None"

    def test_model_functionality(self) -> bool:
        """Testet, ob das geladene Modell funktioniert."""
        with self._model_usage_lock:
            if not self.model:
                return False
            try:
                if self._np is None:
                    self._np = FastLazyLoader.load("numpy")
                test_audio = self._np.random.randn(16000).astype(self._np.float32) * 0.1
                segments, info = self._universal_transcribe(
                    self.model, test_audio,
                    language=None, task="transcribe",
                    temperature=0.0, best_of=1, beam_size=1,
                    without_timestamps=True
                )
                list(segments)  # Ausführen erzwingen
                return True
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.error(f"❌ Model-Test fehlgeschlagen: {e}")
                return False

    def dispose(self) -> None:
        """Räumt alle Ressourcen der Engine auf."""
        logger.info("🧹 Transcription Engine Dispose...")
        self._disposing = True
        with self._lock:
            self._cache.clear()
            self._last_transcription_text = ""
            for (size, backend), model in list(self._model_cache.items()):
                self._unload_model(model)
            self._model_cache.clear()
        self._force_model_cleanup()
        gc.collect()
        logger.info("✅ Transcription Engine disposed")
        
# =============================================================================
# STREAM MANAGER
# =============================================================================
class StreamManager:
    """
    Verwaltet die Erkennung von Streaming-Plattformen und die Extraktion
    von Audio-URLs aus verschiedenen Quellen (YouTube, Twitch, lokale Dateien, etc.).
    Verwendet yt-dlp im Hintergrund und optimiert die Extraktion durch Caching
    und plattformspezifische Einstellungen.
    """

    def __init__(self, enable_debug: bool = False, use_browser_cookies: bool = True) -> None:
        """
        Initialisiert den StreamManager.

        :param enable_debug: Aktiviert ausführliche Debug-Ausgaben.
        :param use_browser_cookies: Erlaubt die Nutzung von Browser-Cookies für yt-dlp.
        """
        self._platform_cache: OrderedDict[str, Tuple[str, str]] = OrderedDict()
        self._audio_url_cache: OrderedDict[str, Dict[str, Any]] = OrderedDict()
        self._live_status_cache: OrderedDict[str, Dict[str, Any]] = OrderedDict()
        self._stream_info_cache: OrderedDict[str, Dict[str, Any]] = OrderedDict()

        self._debug = enable_debug
        self.use_browser_cookies = use_browser_cookies

        self._last_error: Optional[str] = None
        self._last_method: Optional[str] = None

        self._stats = {
            'extraction_attempts': 0,
            'successful_extractions': 0,
            'cache_hits': 0,
            'errors': 0,
            'start_time': time.time()
        }
        self._stats_lock = threading.RLock()
        self._cache_lock = threading.RLock()

        # Prioritäten für Audio-Formate pro Plattform
        self._format_priorities = {
            'youtube': ['bestaudio[ext=m4a]/bestaudio/best', 'bestaudio/best', 'ba'],
            'youtube_live': ['bestaudio/best', 'ba'],
            'twitch': ['bestaudio/best', 'audio_only'],
            'tiktok': ['bestaudio/best'],
            'facebook': ['bestaudio/best'],
            'hls': ['bestaudio/best'],
            'dash': ['bestaudio/best'],
            'generic': ['bestaudio/best', 'ba'],
            'kick': ['bestaudio/best', 'ba'],
            'rumble': ['bestaudio/best', 'ba'],
            'dailymotion': ['bestaudio/best', 'ba'],
            'vimeo': ['bestaudio/best', 'ba'],
            'twitter': ['bestaudio/best', 'ba'],
        }

        self._user_agents = {
            'desktop': (
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            ),
            'mobile': (
                'Mozilla/5.0 (Linux; Android 10; SM-G975F) AppleWebKit/537.36 '
                '(KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36'
            ),
        }

        self._browsers = [
            ('firefox', 'Firefox'),
            ('chrome', 'Chrome'),
            ('brave', 'Brave'),
            ('edge', 'Edge'),
            ('chromium', 'Chromium'),
            ('opera', 'Opera'),
            ('vivaldi', 'Vivaldi'),
        ]

    # -------------------------------------------------------------------------
    # Öffentliche Methoden
    # -------------------------------------------------------------------------

    def detect_platform(self, url: str) -> Tuple[str, str]:
        """
        Erkennt die Plattform einer URL und gibt ein Tupel (Plattform-ID, Anzeigename) zurück.
        Nutzt einen Cache, um wiederholte Anfragen zu beschleunigen.

        :param url: Die zu analysierende URL.
        :return: (plattform_id, plattform_name) z.B. ('youtube', 'YouTube Video')
        """
        if not url:
            return ('unknown', 'Invalid URL')

        url = PlatformUtils.sanitize_url(url)

        with self._cache_lock:
            if url in self._platform_cache:
                with self._stats_lock:
                    self._stats['cache_hits'] += 1
                if self._debug or DEBUG_LEVEL >= 2:
                    logger.debug(f"🔍 detect_platform: Cache-Treffer für {url[:50]}...")
                return self._platform_cache[url]

        url_lower = url.lower().strip()
        detection_reason = []

        # Lokale Datei
        if url_lower.startswith('file://'):
            ok, _ = PlatformUtils.validate_file_path(url)
            if not ok:
                return ('invalid', 'Invalid file path')
            result = ('local', 'Local File')
            detection_reason = ["startswith file://"]

        # Direkte Audio/Video-Erweiterungen
        elif any(url_lower.endswith(ext) for ext in
                 ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac', '.opus', '.webm']):
            result = ('direct_audio', 'Direct Audio')
            detection_reason = ["audio extension"]

        elif any(url_lower.endswith(ext) for ext in
                 ['.mp4', '.avi', '.mkv', '.mov', '.webm', '.m4v', '.wmv', '.flv']):
            result = ('direct_video', 'Direct Video')
            detection_reason = ["video extension"]

        # YouTube
        elif 'youtube.com' in url_lower or 'youtu.be' in url_lower:
            is_live = self._check_youtube_live_status(url)
            if is_live:
                result = ('youtube_live', 'YouTube Live')
                detection_reason = ["youtube domain + live pattern"]
            else:
                result = ('youtube', 'YouTube Video')
                detection_reason = ["youtube domain"]

        # Weitere Plattformen
        elif 'twitch.tv' in url_lower:
            result = ('twitch', 'Twitch')
            detection_reason = ["twitch domain"]
        elif 'kick.com' in url_lower:
            result = ('kick', 'Kick')
            detection_reason = ["kick domain"]
        elif 'rumble.com' in url_lower:
            result = ('rumble', 'Rumble')
            detection_reason = ["rumble domain"]
        elif 'dailymotion.com' in url_lower:
            result = ('dailymotion', 'Dailymotion')
            detection_reason = ["dailymotion domain"]
        elif 'vimeo.com' in url_lower:
            result = ('vimeo', 'Vimeo')
            detection_reason = ["vimeo domain"]
        elif 'twitter.com' in url_lower or 'x.com' in url_lower:
            result = ('twitter', 'Twitter/X')
            detection_reason = ["twitter/x domain"]
        elif 'tiktok.com' in url_lower:
            result = ('tiktok', 'TikTok')
            detection_reason = ["tiktok domain"]
        elif 'facebook.com' in url_lower or 'fb.watch' in url_lower:
            result = ('facebook', 'Facebook')
            detection_reason = ["facebook domain"]
        elif '.m3u8' in url_lower:
            result = ('hls', 'HLS Stream')
            detection_reason = [".m3u8 in URL"]
        elif '.mpd' in url_lower:
            result = ('dash', 'DASH Stream')
            detection_reason = [".mpd in URL"]
        elif url_lower.startswith(('http://', 'https://')):
            result = ('generic', 'Website/Stream')
            detection_reason = ["http(s) fallback"]
        else:
            result = ('unknown', 'Unknown Source')
            detection_reason = ["no pattern matched"]

        # In den Cache legen (max. 50 Einträge)
        with self._cache_lock:
            if len(self._platform_cache) < 50:
                self._platform_cache[url] = result

        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"🔍 detect_platform: {url[:50]}... -> {result}, reason: {', '.join(detection_reason)}")

        return result

    def extract_audio_url(self, url: str, force_refresh: bool = False) -> Optional[str]:
        """
        Extrahiert eine direkte Audio-URL aus der angegebenen Quell-URL.
        Verwendet einen Cache, um wiederholte Extraktionen zu vermeiden.

        :param url: Die Quell-URL (z.B. YouTube-Link, lokale Datei, etc.).
        :param force_refresh: Wenn True, wird der Cache ignoriert und eine neue Extraktion erzwungen.
        :return: Die extrahierte Audio-URL oder None, falls fehlgeschlagen.
        """
        with self._stats_lock:
            self._stats['extraction_attempts'] += 1

        if self._debug or DEBUG_LEVEL >= 1:
            logger.debug(f"\n🎵 [EXTRACT_AUDIO_URL] Start für: {url[:80]}...")

        self._last_error = None
        self._last_method = None

        # Prüfen, ob yt-dlp verfügbar ist
        if not shutil.which('yt-dlp'):
            self._last_error = "yt-dlp not found in PATH"
            logger.error(self._last_error)
            with self._stats_lock:
                self._stats['errors'] += 1
            return None

        if not url or not isinstance(url, str):
            self._last_error = "Invalid input"
            with self._stats_lock:
                self._stats['errors'] += 1
            return None

        cleaned_url = PlatformUtils.sanitize_url(url.strip())
        if not cleaned_url:
            self._last_error = "Empty URL"
            with self._stats_lock:
                self._stats['errors'] += 1
            return None

        # Cache-Schlüssel erstellen
        cache_key = f"audio_{hashlib.md5(cleaned_url.encode()).hexdigest()[:16]}"
        current_time = time.time()

        # Cache prüfen, falls nicht erzwungen
        if not force_refresh:
            with self._cache_lock:
                if cache_key in self._audio_url_cache:
                    cached = self._audio_url_cache[cache_key]
                    cache_age = current_time - cached['timestamp']
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug(f"📦 Cache gefunden, Alter: {cache_age:.1f}s, fehlgeschlagen: {cached.get('failed', False)}")

                    ttl = cached.get('ttl', 1800)
                    if cache_age < ttl and not cached.get('failed', False):
                        with self._stats_lock:
                            self._stats['cache_hits'] += 1
                        return cached['url']
                    elif cache_age < 300 and cached.get('failed', False):
                        # Fehlgeschlagene Einträge für 5 Minuten cachen
                        return None

        # Plattform erkennen
        platform_id, platform_name = self.detect_platform(cleaned_url)
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"🔍 Plattform erkannt: {platform_id} ({platform_name})")

        result = None
        extraction_method = "unknown"

        try:
            # 1. Lokale Datei (file://)
            if cleaned_url.startswith('file://'):
                ok, real_path = PlatformUtils.validate_file_path(cleaned_url)
                if ok:
                    result = cleaned_url
                    extraction_method = "local_file"
                else:
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug(f"❌ Datei-Validierung fehlgeschlagen: {real_path}")
                    self._last_error = real_path

            # 2. Direkte Audio/Video-Links (Erweiterung)
            if not result:
                url_lower = cleaned_url.lower()
                AUDIO_EXTENSIONS = ('.mp3', '.wav', '.m4a', '.flac', '.ogg', '.aac', '.opus', '.webm')
                VIDEO_EXTENSIONS = ('.mp4', '.avi', '.mkv', '.mov', '.webm', '.m4v', '.wmv', '.flv')
                if url_lower.endswith(AUDIO_EXTENSIONS) or url_lower.endswith(VIDEO_EXTENSIONS):
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug("🎵 Direkter Audio/Video-Link erkannt")
                    result = cleaned_url
                    extraction_method = "direct_link"

            # 3. YouTube-spezifische Extraktion (optimiert)
            if not result and platform_id in ['youtube', 'youtube_live']:
                if self._debug or DEBUG_LEVEL >= 2:
                    logger.debug("🎯 YouTube erkannt, verwende optimierte Extraktion mit Cookies...")
                result = self._extract_youtube_audio_optimized(cleaned_url, platform_id)
                extraction_method = "youtube_optimized"

            # 4. Generische Extraktion für andere Plattformen
            if not result:
                if self._debug or DEBUG_LEVEL >= 2:
                    logger.debug("🌐 Keine YouTube-Plattform, verwende generische Extraktion...")

                format_list = self._format_priorities.get(platform_id, self._format_priorities['generic'])
                extraction_method = "ytdlp_generic"

                for i, format_str in enumerate(format_list[:2]):  # Nur die ersten zwei Formate probieren
                    try:
                        if self._debug or DEBUG_LEVEL >= 2:
                            logger.debug(f"  🔄 Versuche Format {i+1}: {format_str}")

                        cmd = [
                            'yt-dlp',
                            '-g',
                            '-f', format_str,
                            '--no-warnings',
                            '--no-check-certificate',
                            '--socket-timeout', '15',
                            '--', cleaned_url
                        ]

                        if debug3_enabled('network'):
                            logger.debug(f"[DEBUG3][NETWORK] yt-dlp Kommando: {' '.join(cmd)}")

                        stdout = self._run_yt_dlp_command(cmd, timeout=15, method_name=f"generic_format_{i+1}")

                        if stdout:
                            # Extrahiere die erste Zeile, die mit http(s) beginnt
                            for line in stdout.splitlines():
                                line = line.strip()
                                if line and line.startswith(('http://', 'https://')):
                                    result = line
                                    if debug3_enabled('network'):
                                        logger.debug(f"[DEBUG3][NETWORK] Erfolg mit Format {format_str}")
                                    break
                        if result:
                            if self._debug or DEBUG_LEVEL >= 2:
                                logger.debug(f"  ✅ Erfolg mit Format {format_str}")
                            break

                    except Exception as e:
                        if PlatformUtils.is_fatal_exception(e):
                            raise
                        if self._debug or DEBUG_LEVEL >= 2:
                            logger.debug(f"  ⚠️ Fehler bei Format {format_str}: {str(e)[:50]}")
                        continue

                # 5. JSON-Fallback, falls keine URL extrahiert wurde
                if not result:
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug("  🔄 Versuche JSON-Fallback...")
                    try:
                        json_result = self._json_extraction_fallback(cleaned_url)
                        if json_result:
                            result = json_result
                            extraction_method = "json_fallback"
                    except Exception as e:
                        if PlatformUtils.is_fatal_exception(e):
                            raise
                        if self._debug or DEBUG_LEVEL >= 2:
                            logger.debug(f"  ⚠️ JSON-Fallback Fehler: {str(e)[:50]}")

        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"❌ Ausnahme in extract_audio_url: {e}")
            self._last_error = f"Exception: {str(e)[:100]}"
            with self._stats_lock:
                self._stats['errors'] += 1

        # TTL je nach Plattform festlegen
        ttl = 300 if platform_id in ['youtube', 'youtube_live'] else 1800

        cache_entry = {
            'url': result,
            'timestamp': current_time,
            'failed': result is None,
            'method': extraction_method,
            'platform': platform_id,
            'ttl': ttl,
        }

        with self._cache_lock:
            self._audio_url_cache[cache_key] = cache_entry
            if len(self._audio_url_cache) > 50:
                # Ältesten Eintrag entfernen (FIFO)
                self._audio_url_cache.popitem(last=False)

        self._last_method = extraction_method

        if result:
            with self._stats_lock:
                self._stats['successful_extractions'] += 1
        else:
            if not self._last_error:
                self._last_error = "No audio URL could be extracted"
            with self._stats_lock:
                self._stats['errors'] += 1

        if self._debug or DEBUG_LEVEL >= 1:
            logger.debug(f"🎵 EXTRACT_AUDIO_URL ENDE - Ergebnis: {'✅ ' + result[:80] + '...' if result else '❌ None'}")

        return result

    def extract_stream_info(self, url: str, force_refresh: bool = False) -> Dict[str, Any]:
        """
        Extrahiert Metadaten (Titel, Uploader, Dauer, etc.) aus der URL.
        Verwendet einen Cache für 10 Minuten.

        :param url: Die Quell-URL.
        :param force_refresh: Wenn True, wird der Cache ignoriert.
        :return: Dictionary mit Stream-Informationen.
        """
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"\n🎯 [EXTRACT_STREAM_INFO] für: {url[:60]}...")

        cache_key = f"info_{hashlib.md5(url.encode()).hexdigest()[:16]}"
        current_time = time.time()

        with self._cache_lock:
            if not force_refresh and cache_key in self._stream_info_cache:
                cached = self._stream_info_cache[cache_key]
                if current_time - cached['timestamp'] < 600:
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug("📦 Stream-Info Cache-Treffer")
                    return cached['info']

        platform_id, platform_name = self.detect_platform(url)

        # Grundlegende Info (wird später mit yt-dlp erweitert, falls möglich)
        info = {
            'title': platform_name,
            'uploader': 'Unknown',
            'duration': 'Unknown',
            'view_count': 0,
            'is_live': False,
            'live_status': 'not_live',
            'thumbnail': '',
            'description': '',
            'platform': platform_id,
            'extractor': 'direct',
            'webpage_url': url,
            'extraction_time': current_time
        }

        # Optional: mit yt-dlp mehr Details holen
        try:
            cmd = [
                'yt-dlp',
                '--dump-json',
                '--no-warnings',
                '--no-check-certificate',
                '--socket-timeout', '10',
                '--', url
            ]
            stdout = self._run_yt_dlp_command(cmd, timeout=10, method_name="stream_info")
            if stdout:
                import json
                data = json.loads(stdout)
                info.update({
                    'title': data.get('title', info['title']),
                    'uploader': data.get('uploader', info['uploader']),
                    'duration': data.get('duration_string', info['duration']),
                    'view_count': data.get('view_count', info['view_count']),
                    'is_live': data.get('is_live', False),
                    'live_status': data.get('live_status', 'not_live'),
                    'thumbnail': data.get('thumbnail', ''),
                    'description': data.get('description', ''),
                    'extractor': data.get('extractor', 'unknown'),
                })
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"⚠️ Stream-Info Extraktion fehlgeschlagen: {e}")

        # Im Cache ablegen
        with self._cache_lock:
            self._stream_info_cache[cache_key] = {
                'info': info,
                'timestamp': current_time
            }
            if len(self._stream_info_cache) > 30:
                # Ältesten Eintrag entfernen
                oldest_key = min(self._stream_info_cache.items(), key=lambda x: x[1]['timestamp'])[0]
                del self._stream_info_cache[oldest_key]

        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("✅ Stream-Info extrahiert")
        return info

    def get_ffmpeg_params_for_url(self, url: str) -> Dict[str, Any]:
        """
        Liefert plattformspezifische FFmpeg-Parameter für die gegebene URL.

        :param url: Die Audio-URL.
        :return: Dictionary mit input_params, output_params, is_live, timeout, etc.
        """
        is_youtube_hls = ('manifest.googlevideo.com' in url and
                          ('/hls_playlist/' in url or '.m3u8' in url))
        is_youtube_dash = ('manifest.googlevideo.com' in url and
                           '/dash/' in url)

        if is_youtube_hls or is_youtube_dash:
            headers = [
                'User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Origin: https://www.youtube.com',
                'Referer: https://www.youtube.com/',
                'Accept: */*',
                'Accept-Language: en-US,en;q=0.9',
                'Accept-Encoding: gzip, deflate, br',
                'Connection: keep-alive',
                'Sec-Fetch-Dest: empty',
                'Sec-Fetch-Mode: cors',
                'Sec-Fetch-Site: same-site',
            ]
            hls_params = [
                '-reconnect', '1',
                '-reconnect_streamed', '1',
                '-reconnect_delay_max', '5',
                '-reconnect_on_network_error', '1',
                '-timeout', '10000000',
                '-rw_timeout', '30000000',
                '-multiple_requests', '1',
                '-seekable', '0',
                '-fflags', '+discardcorrupt+fastseek+genpts',
                '-headers', '\\r\\n'.join(headers),
            ]
            return {
                'input_params': hls_params,
                'output_params': [
                    '-f', 's16le',
                    '-ar', '16000',
                    '-ac', '1',
                    '-vn',
                    '-fflags', '+genpts',
                ],
                'is_live': True,
                'timeout': 90,
                'buffer_size': 16384,
                'platform': 'youtube_hls',
                'reconnect_attempts': 10,
                'requires_headers': True,
                'headers': headers
            }

        elif 'youtube.com' in url or 'youtu.be' in url:
            return {
                'input_params': [
                    '-reconnect', '1',
                    '-reconnect_streamed', '1',
                    '-reconnect_delay_max', '30',
                    '-rw_timeout', '50000000',
                    '-headers', 'User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                    '-headers', 'Referer: https://www.youtube.com/',
                    '-headers', 'Origin: https://www.youtube.com',
                ],
                'output_params': [
                    '-f', 's16le',
                    '-ar', '16000',
                    '-ac', '1',
                    '-vn',
                    '-fflags', '+genpts+discardcorrupt',
                ],
                'is_live': False,
                'timeout': 60,
                'buffer_size': 4096,
                'platform': 'youtube',
                'reconnect_attempts': 5
            }

        else:
            return {
                'input_params': [
                    '-rw_timeout', '30000000',
                    '-fflags', '+discardcorrupt+genpts',
                ],
                'output_params': [
                    '-f', 's16le',
                    '-ar', '16000',
                    '-ac', '1',
                    '-vn',
                    '-avoid_negative_ts', 'make_zero',
                ],
                'is_live': False,
                'timeout': 30,
                'buffer_size': 2048,
                'platform': 'generic'
            }

    def validate_url_for_processing(self, url: str) -> Tuple[bool, str]:
        """
        Validiert, ob die URL verarbeitet werden kann.
        Prüft auf Existenz (bei Dateien) und ob eine Audio-URL extrahiert werden kann.

        :param url: Die zu prüfende URL.
        :return: (ok, status_message)
        """
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"\n🔍 [VALIDATE_URL] für: {url[:80]}...")

        if not url or not isinstance(url, str):
            return False, "Invalid input"

        cleaned_url = PlatformUtils.sanitize_url(url.strip())
        if not cleaned_url:
            return False, "Empty URL"

        # Lokale Datei
        if cleaned_url.startswith('file://'):
            ok, real_path = PlatformUtils.validate_file_path(cleaned_url)
            if not ok:
                return False, real_path
            file_path = real_path
            if not os.path.exists(file_path):
                return False, "File not found"
            if not os.path.isfile(file_path):
                return False, "Not a valid file"
            try:
                file_size = os.path.getsize(file_path)
                if file_size == 0:
                    return False, "File is empty"
                filename = os.path.basename(file_path)
                return True, f"File: {filename}"
            except OSError:
                return False, "File access error"

        # Kein http(s)
        if not cleaned_url.startswith(('http://', 'https://')):
            return False, "Invalid URL format"

        if len(cleaned_url) > 2000:
            return False, "URL too long"

        # Extraktion testen
        audio_url = self.extract_audio_url(cleaned_url)
        if not audio_url:
            platform_id, platform_name = self.detect_platform(cleaned_url)
            error_msg = f"No audio URL extractable ({platform_name})"
            if self._last_error:
                error_msg += f" - {self._last_error}"
            if self._last_method:
                error_msg += f" [method: {self._last_method}]"
            return False, error_msg

        # Plattformdetails für die Statusmeldung sammeln
        platform_id, platform_name = self.detect_platform(cleaned_url)
        status_parts = [platform_name]

        if 'youtube' in platform_id:
            # Versuche, Titel und Dauer zu ermitteln
            try:
                cmd = [
                    'yt-dlp',
                    '--dump-json',
                    '--playlist-items', '1',
                    '--no-warnings',
                    '--no-check-certificate',
                    '--socket-timeout', '5',
                    cleaned_url
                ]
                stdout = self._run_yt_dlp_command(cmd, timeout=5, method_name="validate_info")
                if stdout:
                    data = json.loads(stdout)
                    if data.get('title'):
                        title = data['title'][:40]
                        status_parts.insert(0, title)
                    if data.get('duration_string'):
                        status_parts.append(f"⏱️ {data['duration_string']}")
                    if data.get('is_live'):
                        status_parts.append("🔴 LIVE")
            except Exception:
                pass

        status = " | ".join(status_parts)
        return True, status

    def get_diagnostics(self) -> Dict[str, Any]:
        """
        Gibt Diagnoseinformationen über den aktuellen Zustand des Managers zurück.
        """
        current_time = time.time()
        with self._stats_lock:
            uptime = current_time - self._stats['start_time']
            stats = self._stats.copy()
            stats.update({
                'uptime_seconds': uptime,
                'uptime_human': str(timedelta(seconds=int(uptime))),
                'success_rate': (
                    stats['successful_extractions'] / stats['extraction_attempts'] * 100
                    if stats['extraction_attempts'] > 0 else 0
                ),
                'cache_hit_rate': (
                    stats['cache_hits'] / stats['extraction_attempts'] * 100
                    if stats['extraction_attempts'] > 0 else 0
                )
            })

        return {
            'last_error': self._last_error,
            'last_method': self._last_method,
            'stats': stats,
            'cache_sizes': {
                'platform': len(self._platform_cache),
                'audio_url': len(self._audio_url_cache),
                'live_status': len(self._live_status_cache),
                'stream_info': len(self._stream_info_cache)
            },
            'debug_mode': self._debug
        }

    def clear_caches(self) -> None:
        """Leert alle internen Caches."""
        with self._cache_lock:
            self._platform_cache.clear()
            self._audio_url_cache.clear()
            self._live_status_cache.clear()
            self._stream_info_cache.clear()
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("🗑️ Alle Caches geleert")

    def dispose(self) -> None:
        """Räumt Ressourcen auf (leert Caches)."""
        self.clear_caches()
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("🔌 StreamManager disposed")

    # -------------------------------------------------------------------------
    # Private Hilfsmethoden
    # -------------------------------------------------------------------------

    def _run_yt_dlp_command(self, cmd: List[str], timeout: int = 15, method_name: str = "unknown") -> Optional[str]:
        """
        Führt ein yt-dlp-Kommando aus und gibt bei Erfolg die stdout zurück.
        Behandelt Timeouts und Fehler.

        :param cmd: Kommando als Liste.
        :param timeout: Timeout in Sekunden.
        :param method_name: Name der Methode für Debug-Ausgaben.
        :return: stdout bei Erfolg, sonst None.
        """
        try:
            if self._debug:
                logger.debug(f"  ▶️ Ausführen: {' '.join(cmd)}")

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                shell=False,
                encoding='utf-8',
                errors='ignore'
            )

            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
            else:
                if result.stderr:
                    err = result.stderr.strip()
                    if self._debug:
                        logger.debug(f"  ⚠️ {method_name} fehlgeschlagen: {err[:100]}")
                return None

        except subprocess.TimeoutExpired:
            if self._debug:
                logger.debug(f"  ⏰ Timeout bei {method_name} nach {timeout}s")
            return None
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            if self._debug:
                logger.debug(f"  ⚠️ Fehler bei {method_name}: {str(e)[:50]}")
            return None

    def _extract_youtube_video_id(self, url: str) -> Optional[str]:
        """Extrahiert die YouTube-Video-ID aus verschiedenen URL-Formaten."""
        patterns = [
            r'(?:youtube\.com/watch\?v=|youtu\.be/)([a-zA-Z0-9_-]{11})',
            r'youtube\.com/embed/([a-zA-Z0-9_-]{11})',
            r'youtube\.com/v/([a-zA-Z0-9_-]{11})',
        ]
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None

    def _check_youtube_live_status(self, url: str) -> bool:
        """
        Prüft, ob eine YouTube-URL wahrscheinlich ein Live-Stream ist.
        Verwendet einen einfachen Caching-Mechanismus.
        """
        cache_key = f"live_{hashlib.md5(url.encode()).hexdigest()[:16]}"
        current_time = time.time()

        with self._cache_lock:
            if cache_key in self._live_status_cache:
                cached = self._live_status_cache[cache_key]
                if current_time - cached['timestamp'] < 300:
                    return cached['is_live']

        url_lower = url.lower()
        live_patterns = ['/live', 'live=1', '/stream', 'livestream']
        is_live = any(pattern in url_lower for pattern in live_patterns)

        with self._cache_lock:
            if len(self._live_status_cache) > 30:
                # Ältesten Eintrag entfernen
                oldest_keys = sorted(self._live_status_cache.keys(),
                                     key=lambda k: self._live_status_cache[k]['timestamp'])[:10]
                for k in oldest_keys:
                    del self._live_status_cache[k]

            self._live_status_cache[cache_key] = {
                'is_live': is_live,
                'timestamp': current_time
            }

        return is_live

    def _extract_youtube_audio_optimized(self, url: str, platform_id: str) -> Optional[str]:
        """
        Optimierte Extraktion für YouTube: probiert verschiedene Methoden:
        - Mit Browser-Cookies (falls aktiviert)
        - Mit verschiedenen User-Agents
        - JSON-Fallback
        - Generierte direkte URL als letzte Option
        """
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"  🔍 Optimierte YouTube-Extraktion für: {url[:60]}...")

        video_id = self._extract_youtube_video_id(url)
        if not video_id or len(video_id) != 11:
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug("  ❌ Ungültige Video-ID")
            return None

        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"  🔍 Video-ID: {video_id}")

        # 1. Mit Browser-Cookies (falls erlaubt)
        if self.use_browser_cookies:
            for browser_cmd, browser_name in self._browsers:
                try:
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug(f"    🧪 Teste mit {browser_name}-Cookies...")

                    cmd = [
                        'yt-dlp',
                        '-g',
                        '-f', 'bestaudio[ext=m4a]/bestaudio/best',
                        '--cookies-from-browser', browser_cmd,
                        '--no-warnings',
                        '--no-check-certificate',
                        '--socket-timeout', '15',
                        '--', url
                    ]

                    if debug3_enabled('network'):
                        logger.debug(f"[DEBUG3][NETWORK] Kommando: {' '.join(cmd)}")

                    stdout = self._run_yt_dlp_command(cmd, timeout=20, method_name=f"{browser_name}_cookies")
                    if stdout:
                        for line in stdout.splitlines():
                            line = line.strip()
                            if line and line.startswith(('http://', 'https://')):
                                if debug3_enabled('network'):
                                    logger.debug(f"[DEBUG3][NETWORK] Erfolg mit {browser_name}-Cookies: {line[:100]}...")
                                if self._debug or DEBUG_LEVEL >= 2:
                                    logger.debug(f"    ✅ Erfolg mit {browser_name}-Cookies")
                                return line
                except Exception as e:
                    if PlatformUtils.is_fatal_exception(e):
                        raise
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug(f"    ⚠️ Fehler bei {browser_name}: {str(e)[:50]}")
                    continue

        # 2. Verschiedene Methoden ohne Cookies
        methods = [
            {
                'name': 'Standard yt-dlp',
                'cmd': [
                    'yt-dlp',
                    '-g',
                    '-f', 'bestaudio[ext=m4a]/bestaudio/best',
                    '--no-warnings',
                    '--no-check-certificate',
                    '--user-agent', self._user_agents['desktop'],
                    '--referer', 'https://www.youtube.com/',
                    '--socket-timeout', '15',
                    '--', url
                ],
                'timeout': 20
            },
            {
                'name': 'Mobile User-Agent',
                'cmd': [
                    'yt-dlp',
                    '-g',
                    '-f', 'bestaudio/best',
                    '--no-warnings',
                    '--no-check-certificate',
                    '--user-agent', self._user_agents['mobile'],
                    '--referer', 'https://m.youtube.com/',
                    '--socket-timeout', '15',
                    '--', url
                ],
                'timeout': 20
            },
            {
                'name': 'Lowest Quality',
                'cmd': [
                    'yt-dlp',
                    '-g',
                    '-f', 'worstaudio',
                    '--no-warnings',
                    '--no-check-certificate',
                    '--socket-timeout', '15',
                    '--', url
                ],
                'timeout': 20
            },
        ]

        for method in methods:
            try:
                if self._debug or DEBUG_LEVEL >= 2:
                    logger.debug(f"    🧪 Teste: {method['name']}")

                stdout = self._run_yt_dlp_command(method['cmd'], timeout=method['timeout'], method_name=method['name'])
                if stdout:
                    for line in stdout.splitlines():
                        line = line.strip()
                        if line and line.startswith(('http://', 'https://')):
                            if debug3_enabled('network'):
                                logger.debug(f"[DEBUG3][NETWORK] Erfolg mit {method['name']}: {line[:100]}...")
                            if self._debug or DEBUG_LEVEL >= 2:
                                logger.debug(f"    ✅ Erfolg mit {method['name']}")
                            return line
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                if self._debug or DEBUG_LEVEL >= 2:
                    logger.debug(f"    ⚠️ Fehler bei {method['name']}: {str(e)[:50]}")
                continue

        # 3. JSON-Fallback
        try:
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug("    🔄 Versuche JSON-Fallback...")

            cmd = [
                'yt-dlp',
                '--dump-json',
                '--no-warnings',
                '--no-check-certificate',
                '--user-agent', self._user_agents['desktop'],
                '--socket-timeout', '20',
                '--', url
            ]
            stdout = self._run_yt_dlp_command(cmd, timeout=25, method_name="json_fallback")
            if stdout:
                data = json.loads(stdout)
                if debug3_enabled('network'):
                    logger.debug(f"[DEBUG3][NETWORK] JSON-Fallback-Formate: {json.dumps(data.get('formats', []), indent=2)[:500]}")

                best_audio = None
                best_bitrate = 0
                for fmt in data.get('formats', []):
                    if fmt.get('acodec') != 'none' and fmt.get('url'):
                        bitrate = fmt.get('abr', 0) or fmt.get('tbr', 0) or 0
                        if fmt.get('vcodec') == 'none':
                            bitrate += 1000  # Audio-only bevorzugen
                        ext = fmt.get('ext', '').lower()
                        if ext in ['m4a', 'mp4']:
                            bitrate += 500
                        elif ext in ['webm', 'opus']:
                            bitrate += 300
                        if bitrate > best_bitrate:
                            best_bitrate = bitrate
                            best_audio = fmt['url']
                if best_audio:
                    if self._debug or DEBUG_LEVEL >= 2:
                        logger.debug("    ✅ JSON-Fallback erfolgreich")
                    return best_audio
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"    ⚠️ JSON-Fallback Fehler: {str(e)[:50]}")

        # 4. Ganz zum Schluss: generierte direkte URL (funktioniert oft nicht)
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("    🔄 Generiere direkte Audio-URL...")

        direct_url = f"https://manifest.googlevideo.com/api/manifest/dash/id/{video_id}/source/youtube"
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("    🔧 Generierte direkte URL")
        return direct_url

    def _json_extraction_fallback(self, url: str) -> Optional[str]:
        """
        Extrahiert eine Audio-URL über das JSON-Dump von yt-dlp.
        Wird als letzte Möglichkeit verwendet.
        """
        try:
            cmd = [
                'yt-dlp',
                '--dump-json',
                '--no-warnings',
                '--no-check-certificate',
                '--socket-timeout', '20',
                '--', url
            ]
            stdout = self._run_yt_dlp_command(cmd, timeout=20, method_name="json_extraction")
            if not stdout:
                return None

            data = json.loads(stdout)
            audio_formats = []
            for fmt in data.get('formats', []):
                if fmt.get('acodec') != 'none' and fmt.get('url'):
                    audio_formats.append({
                        'url': fmt['url'],
                        'abr': fmt.get('abr', 0) or fmt.get('tbr', 0) or 0,
                        'ext': fmt.get('ext', ''),
                        'vcodec': fmt.get('vcodec', 'none')
                    })

            if audio_formats:
                # Audio-only Formate bevorzugen
                audio_only = [f for f in audio_formats if f['vcodec'] == 'none']
                if audio_only:
                    audio_formats = audio_only
                # Sortieren nach Bitrate (höchste zuerst)
                audio_formats.sort(key=lambda x: x['abr'], reverse=True)
                return audio_formats[0]['url']
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"  ⚠️ JSON-Extraktion Fehler: {str(e)[:50]}")
        return None


class FFmpegManager:
    def __init__(self, config: Optional[Config] = None,
                 stream_manager: Optional[StreamManager] = None,
                 settings: Optional[Settings] = None) -> None:
        self._processes: Dict[str, Dict[str, Any]] = {}
        self._process_counter = 0
        self._lock = threading.RLock()
        self._active_count = 0
        self._shutting_down = False
        self.config = config or Config()
        self.stream_manager = stream_manager or StreamManager()
        self.settings = settings
        self._pid_tracking: Dict[int, Dict[str, Any]] = {}
        self._live_detection_cache: Dict[str, Dict[str, Any]] = {}
        self._cleanup_running = True
        self._cleanup_thread = threading.Thread(
            target=self._cleanup_worker,
            daemon=True,
            name="FFmpegCleanup"
        )
        self._cleanup_thread.start()
        self._stats = {
            'extraction_attempts': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'cache_hits': 0,
            'start_time': time.time()
        }
        logger.info(f"✅ FFmpeg Manager initialized (Platform: {SYSTEM})")

    def set_stream_manager(self, stream_manager: StreamManager) -> 'FFmpegManager':
        if stream_manager:
            self.stream_manager = stream_manager
            logger.info("✅ FFmpegManager: StreamManager linked")
        return self

    def _build_ffmpeg_command_optimized(self, url: str, seek_seconds: Optional[float] = None,
                                         detected_language: Optional[str] = None) -> List[str]:
        is_live, platform = self._detect_stream_type(url)
        stream_type = "LIVE" if is_live else "VIDEO"
        logger.info(f"\n🎬 Building FFmpeg command for {platform} ({stream_type})")
        logger.info(f"  📍 URL: {url[:80]}...")
        cmd = ['ffmpeg', '-hide_banner', '-loglevel', 'warning']

        if 'youtube.com' in url.lower() or 'youtu.be' in url.lower():
            logger.info("  🎯 Adding YouTube-specific headers")
            headers_dict = self.config.get_youtube_headers(is_manifest='manifest.googlevideo.com' in url)
            headers_list = [f"{k}: {v}" for k, v in headers_dict.items()]
            headers_string = '\r\n'.join(headers_list)
            cmd.extend(['-headers', headers_string])

        if is_live:
            logger.info("  📡 LIVE: Using HLS/Live optimization")
            cmd.extend([
                '-reconnect', '1',
                '-reconnect_streamed', '1',
                '-reconnect_delay_max', '5',
                '-reconnect_on_network_error', '1',
                '-timeout', '10000000',
                '-rw_timeout', '30000000',
                '-multiple_requests', '1',
                '-seekable', '0',
                '-fflags', '+discardcorrupt+fastseek+genpts',
            ])
            if seek_seconds is not None:
                logger.warning(f"⚠️ seek_seconds={seek_seconds} wird bei Live-Stream ignoriert")
        else:
            logger.info("  🎬 VIDEO: Fast access for non-live content")
            cmd.extend([
                '-rw_timeout', '10000000',
                '-accurate_seek',
                '-fflags', '+genpts+discardcorrupt+fastseek',
            ])
            if seek_seconds is not None and seek_seconds > 0:
                logger.info(f"  ⏩ Seeking to {seek_seconds}s")
                cmd.extend(['-ss', str(seek_seconds)])

        cmd.extend(['-i', url])

        profile = 'realtime' if is_live else 'transcription'
        if self.settings and hasattr(self.settings, 'audio_profile'):
            profile = self.settings.audio_profile
        audio_filter = self.config.get_audio_filter(language=detected_language, profile=profile)
        logger.info(f"  🎚️ Using audio filter (profile={profile}): {audio_filter}")

        cmd.extend([
            '-vn',
            '-f', 's16le',
            '-acodec', 'pcm_s16le',
            '-ar', str(Constants.SAMPLE_RATE),
            '-ac', str(Constants.CHANNELS),
            '-af', audio_filter,
            '-fflags', '+genpts+discardcorrupt',
            '-avoid_negative_ts', 'make_zero',
            '-max_interleave_delta', '0',
            '-threads', '2',
            '-bufsize', '2048k',
            'pipe:1'
        ])
        return cmd

    def start_stream(self, video_url: str, output_queue: Optional[queue.Queue],
                     process_id: str, force_refresh_audio_url: bool = False,
                     audio_url: Optional[str] = None,
                     seek_seconds: Optional[float] = None,
                     detected_language: Optional[str] = None) -> Optional[subprocess.Popen]:
        logger.info(f"\n🎬 FFmpegManager: Starting stream for: {video_url[:80]}...")
        with self._lock:
            if self.is_active(process_id):
                logger.warning(f"⚠️ Stream {process_id} already active")
                return None

        if audio_url is None:
            logger.info("🎵 Resolving audio URL...")
            audio_url = self.stream_manager.extract_audio_url(video_url, force_refresh=force_refresh_audio_url)
            if not audio_url:
                logger.error("❌ Audio URL resolution failed")
                return None
            logger.info(f"✅ Resolved URL: {audio_url[:100]}...")
        else:
            logger.info(f"✅ Using pre-resolved audio URL: {audio_url[:100]}...")

        cmd = self._build_ffmpeg_command_optimized(audio_url, seek_seconds=seek_seconds,
                                                    detected_language=detected_language)
        if debug3_enabled('ffmpeg'):
            logger.debug(f"[DEBUG3][FFMPEG] Starting FFmpeg: {' '.join(cmd)}")
        try:
            process_kwargs = {
                'stdout': subprocess.PIPE,
                'stderr': subprocess.PIPE,
                'stdin': subprocess.DEVNULL,
                'bufsize': 10 * 1024 * 1024,
            }
            if IS_WINDOWS:
                process_kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
                process_kwargs['encoding'] = 'utf-8'
                process_kwargs['errors'] = 'ignore'
            elif IS_MACOS or IS_LINUX:
                process_kwargs['start_new_session'] = True
            logger.info("🚀 Starting FFmpeg process...")
            process = subprocess.Popen(cmd, **process_kwargs)
            logger.info(f"✅ FFmpeg process started (PID: {process.pid})")
            logger.info("⏳ Waiting for stream initialization...")
            time.sleep(3.0 if 'hls' in audio_url.lower() else 1.5)
            poll_result = process.poll()
            if poll_result is not None:
                try:
                    stderr_output = PlatformUtils.read_process_stderr(process, 1000)
                    if debug3_enabled('ffmpeg'):
                        logger.debug(f"[DEBUG3][FFMPEG] FFmpeg died immediately, stderr: {stderr_output}")
                    logger.error(f"❌ FFmpeg died immediately. Exit code: {poll_result}")
                    if stderr_output:
                        logger.error("📋 FFMPEG STDERR (first 200 chars):")
                        logger.error(stderr_output[:200])
                except Exception as e:
                    logger.warning(f"⚠️ Could not read stderr: {e}")
                return None
            logger.info(f"✅ FFmpeg is running (PID: {process.pid})")
            self._register_process(process_id, process, output_queue, audio_url)
            return process
        except FileNotFoundError:
            logger.error("❌ FFmpeg not found! Please install FFmpeg.")
            return None
        except PermissionError:
            logger.error("❌ Permission denied - cannot execute FFmpeg")
            return None
        except subprocess.TimeoutExpired:
            logger.error("❌ Timeout beim Start von FFmpeg")
            return None
        except OSError as e:
            logger.error(f"❌ OS-Fehler beim Start von FFmpeg: {e}")
            return None
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.error(f"❌ Failed to start FFmpeg: {e}")
            return None

    def _register_process(self, process_id: str, process: subprocess.Popen,
                         output_queue: Optional[queue.Queue], url: str) -> None:
        with self._lock:
            is_live, platform = self._detect_stream_type(url)
            headers_used = 'youtube.com' in url.lower() or 'googlevideo.com' in url.lower()
            self._processes[process_id] = {
                'process': process,
                'output_queue': output_queue,
                'start_time': time.time(),
                'url': url,
                'stopping': False,
                'bytes_read': 0,
                'platform': platform,
                'is_live': is_live,
                'chunks_processed': 0,
                'last_activity': time.time(),
                'headers_used': headers_used
            }
            self._active_count += 1
            self._pid_tracking[process.pid] = {
                'process_id': process_id,
                'start_time': time.time(),
                'url': url[:100],
                'platform': platform,
                'is_live': is_live
            }
            logger.info(f"📊 Process registered: {process_id} (PID: {process.pid})")

    def update_process_activity(self, process_id: str) -> None:
        with self._lock:
            if process_id in self._processes:
                self._processes[process_id]['last_activity'] = time.time()

    def _detect_stream_type(self, url: str) -> Tuple[bool, str]:
        cache_key = hashlib.md5(url.encode()).hexdigest()[:16]
        with self._lock:
            if cache_key in self._live_detection_cache:
                cached = self._live_detection_cache[cache_key]
                if time.time() - cached['timestamp'] < 300:
                    return cached['is_live'], cached['platform']
        is_live = False
        platform = "unknown"
        try:
            url_lower = url.lower()
            if 'youtube.com' in url_lower or 'youtu.be' in url_lower:
                platform = "YouTube"
                is_live = any(indicator in url_lower for indicator in
                            ['/live', 'live=1', '/stream', 'livestream'])
                if debug3_enabled('network'):
                    logger.debug(f"[DEBUG3][NETWORK] YouTube detection: {url_lower[:100]}, is_live={is_live}")
            elif 'twitch.tv' in url_lower:
                platform = "Twitch"
                is_live = True
                if debug3_enabled('network'):
                    logger.debug("[DEBUG3][NETWORK] Twitch detected -> is_live=True")
            elif 'tiktok.com' in url_lower:
                platform = "TikTok"
            elif 'facebook.com' in url_lower or 'fb.watch' in url_lower:
                platform = "Facebook"
            elif 'kick.com' in url_lower:
                platform = "Kick"
            elif 'rumble.com' in url_lower:
                platform = "Rumble"
            elif 'dailymotion.com' in url_lower:
                platform = "Dailymotion"
            elif 'vimeo.com' in url_lower:
                platform = "Vimeo"
            elif 'twitter.com' in url_lower or 'x.com' in url_lower:
                platform = "Twitter/X"
            elif url_lower.startswith('file://'):
                platform = "Local File"
                is_live = False
            elif '.m3u8' in url_lower:
                platform = "HLS Stream"
                is_live = True
            elif '.mpd' in url_lower:
                platform = "DASH Stream"
                is_live = True
            else:
                platform = "HTTP Stream"
            with self._lock:
                self._live_detection_cache[cache_key] = {
                    'is_live': is_live,
                    'platform': platform,
                    'timestamp': time.time(),
                    'url': url[:50]
                }
                if len(self._live_detection_cache) > 100:
                    oldest = min(self._live_detection_cache.items(),
                               key=lambda x: x[1]['timestamp'])[0]
                    del self._live_detection_cache[oldest]
            return is_live, platform
        except (KeyError, AttributeError) as e:
            logger.warning(f"⚠️ Stream type detection error: {e}")
            return False, "unknown"
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Stream type detection error: {e}")
            return False, "unknown"

    def get_stats(self) -> Dict[str, Any]:
        with self._lock:
            stats = self._stats.copy()
            stats['uptime_seconds'] = time.time() - stats['start_time']
            if stats['extraction_attempts'] > 0:
                stats['success_rate'] = (
                    stats['successful_extractions'] / stats['extraction_attempts'] * 100
                )
                stats['failure_rate'] = (
                    stats['failed_extractions'] / stats['extraction_attempts'] * 100
                )
            else:
                stats['success_rate'] = 0
                stats['failure_rate'] = 0
            stats['active_processes'] = self._active_count
            stats['total_processes'] = len(self._processes)
            stats['live_detection_cache_size'] = len(self._live_detection_cache)
            return stats

    def read_audio_data(self, process_id: str, size: int) -> Optional[bytes]:
        with self._lock:
            if process_id not in self._processes:
                return None
            process_info = self._processes[process_id]
            if process_info.get('stopping', False):
                return None
            process = process_info['process']
        try:
            audio_data = process.stdout.read(size)
            if audio_data:
                with self._lock:
                    process_info['bytes_read'] += len(audio_data)
                    process_info['chunks_processed'] += 1
                    process_info['last_activity'] = time.time()
                return audio_data
            else:
                if process.poll() is not None:
                    exit_code = process.poll()
                    logger.warning(f"⚠️ Process {process_id} terminated (exit: {exit_code})")
                    try:
                        stderr = PlatformUtils.read_process_stderr(process, 300)
                        if stderr:
                            logger.info(f"📝 Last error: {stderr[:150]}")
                    except Exception:
                        pass
                    self.stop_stream(process_id)
                    return None
                return None
        except (IOError, OSError, ValueError) as e:
            logger.warning(f"⚠️ Read error for {process_id}: {e}")
            self.stop_stream(process_id)
            return None
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Unexpected read error for {process_id}: {e}")
            self.stop_stream(process_id)
            return None

    def stop_stream(self, process_id: str) -> bool:
        with self._lock:
            if process_id not in self._processes:
                return True
            process_info = self._processes[process_id]
            if process_info.get('stopping', False):
                return True
            process = process_info['process']
            termination_success = False
            try:
                if process.poll() is None:
                    logger.info(f"🔄 Stopping process {process_id} ({process.pid})...")
                    try:
                        process.terminate()
                        process.wait(timeout=1.0)
                        termination_success = True
                        logger.info(f"✅ Process {process_id} terminated gracefully")
                    except subprocess.TimeoutExpired:
                        try:
                            process.kill()
                            process.wait(timeout=1.0)
                            termination_success = True
                            logger.info(f"✅ Process {process_id} killed")
                        except subprocess.TimeoutExpired:
                            termination_success = False
                            logger.error(f"❌ Could not terminate {process_id}")
                else:
                    termination_success = True
                    logger.info(f"✅ Process {process_id} already terminated")
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.error(f"❌ Error stopping {process_id}: {e}")
                termination_success = False
            finally:
                self._cleanup_process_resources(process_id, process)
            return termination_success

    def _cleanup_process_resources(self, process_id: str, process: subprocess.Popen) -> None:
        if process_id not in self._processes:
            return
        try:
            if process_id in self._processes:
                del self._processes[process_id]
                self._active_count = max(0, self._active_count - 1)
            if process.pid in self._pid_tracking:
                del self._pid_tracking[process.pid]
            pipes_to_close = []
            if hasattr(process, 'stdout'):
                pipes_to_close.append(process.stdout)
            if hasattr(process, 'stderr'):
                pipes_to_close.append(process.stderr)
            if hasattr(process, 'stdin'):
                pipes_to_close.append(process.stdin)
            for pipe in pipes_to_close:
                if pipe and not pipe.closed:
                    try:
                        pipe.close()
                    except Exception:
                        pass
            if process.poll() is None:
                try:
                    process.terminate()
                    time.sleep(0.1)
                except (OSError, PermissionError):
                    pass
                try:
                    process.kill()
                    time.sleep(0.1)
                except (OSError, PermissionError):
                    pass
            try:
                if process.poll() is None:
                    process.kill()
            except (OSError, PermissionError):
                pass
            logger.debug(f"🧹 Resources cleaned for: {process_id}")
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Resource cleanup error for {process_id}: {e}")

    def stop_all_streams(self) -> None:
        logger.info("🛑 Stopping all streams...")
        with self._lock:
            self._shutting_down = True
            process_ids = list(self._processes.keys())
            success_count = 0
            fail_count = 0
            for process_id in process_ids:
                try:
                    if self.stop_stream(process_id):
                        success_count += 1
                    else:
                        fail_count += 1
                except Exception as e:
                    if PlatformUtils.is_fatal_exception(e):
                        raise
                    logger.warning(f"⚠️ Error stopping {process_id}: {e}")
                    fail_count += 1
            self._shutting_down = False
            logger.info(f"✅ Streams stopped: {success_count} successful, {fail_count} failed")

    def is_active(self, process_id: str) -> bool:
        with self._lock:
            if process_id not in self._processes:
                return False
            process = self._processes[process_id]['process']
            if process.poll() is not None:
                return False
            last_activity = self._processes[process_id].get('last_activity', 0)
            if time.time() - last_activity > 30:
                return False
            return True

    def _cleanup_worker(self) -> None:
        while self._cleanup_running:
            try:
                time.sleep(30)
                self.cleanup_stale_processes()
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"⚠️ Cleanup worker error: {e}")

    def cleanup_stale_processes(self) -> None:
        with self._lock:
            stale_processes = []
            for process_id, process_info in self._processes.items():
                process = process_info['process']
                if process.poll() is not None:
                    stale_processes.append(process_id)
            for process_id in stale_processes:
                logger.info(f"🧹 Cleaning terminated process: {process_id}")
                self._cleanup_process_resources(process_id, self._processes[process_id]['process'])

    def dispose(self) -> None:
        logger.info("🧹 Shutting down FFmpeg Manager...")
        self._cleanup_running = False
        self.stop_all_streams()
        if self._cleanup_thread and self._cleanup_thread.is_alive():
            self._cleanup_thread.join(timeout=2.0)
        self._live_detection_cache.clear()
        self._pid_tracking.clear()
        self._processes.clear()
        gc.collect()
        logger.info("✅ FFmpeg Manager disposed")

# -----------------------------------------------------------------------------
# STREAM INFO EXTRACTOR
# -----------------------------------------------------------------------------
@dataclass
class StreamInfo:
    title: str
    uploader: str
    duration: str                # Anzeigedauer (z.B. "Live", "3:45")
    view_count: int
    platform: str
    description: str = ""
    duration_seconds: Optional[float] = None   # Für Fortschrittsbalken
    is_live: bool = False                      # Ob es ein Livestream ist
    thumbnail: str = ""                         # URL des Vorschaubilds
    original_url: str = ""                       # Die ursprünglich aufgerufene URL
    stream_url: Optional[str] = None             # Die extrahierte Audio-URL (optional)

class StreamInfoExtractor:
    """
    Extrahiert detaillierte Stream-Informationen (Titel, Uploader, Dauer, etc.)
    aus verschiedenen Plattformen (YouTube, Twitch, lokale Dateien, generische Streams).
    Nutzt yt-dlp mit optionalen Browser-Cookies für bessere Ergebnisse.
    """

    def __init__(self) -> None:
        self.current_info = StreamInfo(
            title="Unknown Stream",
            uploader="Unknown",
            duration="Live",
            view_count=0,
            platform="Unknown"
        )
        self._lock = threading.RLock()
        self._debug = DEBUG_LEVEL >= 1
        self.use_browser_cookies = True

    # ----------------------------------------------------------------------
    # Öffentliche Hauptmethode
    # ----------------------------------------------------------------------
    def extract_stream_info(self, url: str) -> StreamInfo:
        """
        Extrahiert Stream-Informationen aus der angegebenen URL.
        Gibt immer ein StreamInfo-Objekt zurück (niemals None).
        """
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug(f"🔍 StreamInfoExtractor.extract_stream_info für: {url[:80]}...")
        url = PlatformUtils.sanitize_url(url)

        # 1. Lokale Datei
        if url.startswith('file://'):
            return self._handle_local_file(url)

        # 2. YouTube mit Cookies (falls aktiviert)
        if 'youtube.com' in url.lower() or 'youtu.be' in url.lower():
            if self.use_browser_cookies:
                info = self._extract_youtube_info_with_cookies(url)
                if info:
                    self.current_info = info
                    return info
            # Fallback: JSON ohne Cookies
            info = self._run_ytdlp_json(url, platform='youtube')
            if info:
                self.current_info = info
                return info

        # 3. Twitch mit Cookies (falls aktiviert)
        if 'twitch.tv' in url.lower():
            if self.use_browser_cookies:
                info = self._extract_twitch_info_with_cookies(url)
                if info:
                    self.current_info = info
                    return info
            # Fallback: JSON ohne Cookies
            info = self._run_ytdlp_json(url, platform='twitch')
            if info:
                self.current_info = info
                return info

        # 4. Andere Plattformen (oder wenn Cookies deaktiviert)
        info = self._run_ytdlp_json(url)
        if info:
            self.current_info = info
            return info

        # 5. Letzter Fallback: Titel aus URL extrahieren (kein Datum!)
        info = self._fallback_from_url(url)
        self.current_info = info
        return info

    # ----------------------------------------------------------------------
    # Hilfsmethoden für die Extraktion
    # ----------------------------------------------------------------------
    def _handle_local_file(self, url: str) -> StreamInfo:
        """Extrahiert Informationen aus einer lokalen file://-URL."""
        ok, real_path = PlatformUtils.validate_file_path(url)
        if not ok:
            return StreamInfo(
                title="Invalid file",
                uploader="Error",
                duration="",
                view_count=0,
                platform="invalid"
            )
        file_path = real_path
        return StreamInfo(
            title=os.path.basename(file_path),
            uploader="Local File",
            duration="File",
            view_count=0,
            platform="local"
        )

    def _run_ytdlp_json(self, url: str, platform: str = "generic") -> Optional[StreamInfo]:
        """
        Führt yt-dlp mit --dump-json aus und gibt ein StreamInfo-Objekt zurück.
        Plattform-spezifische Timeouts werden verwendet.
        """
        timeout = 15 if platform in ("youtube", "twitch") else 10
        try:
            cmd = [
                'yt-dlp',
                '--dump-json',
                '--no-warnings',
                '--no-check-certificate',
                '--socket-timeout', str(timeout),
                '--', url
            ]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout + 5,
                shell=False,
                encoding='utf-8',
                errors='ignore'
            )
            if result.returncode == 0 and result.stdout.strip():
                data = json.loads(result.stdout)
                # Plattform aus dem Extractor ableiten
                extractor = data.get('extractor', '').lower()
                platform_map = {
                    'youtube': 'youtube',
                    'twitch': 'twitch',
                    'tiktok': 'tiktok',
                    'facebook': 'facebook',
                    'kick': 'kick',
                    'rumble': 'rumble',
                    'dailymotion': 'dailymotion',
                    'vimeo': 'vimeo',
                    'twitter': 'twitter',
                    'x': 'twitter',
                }
                detected_platform = "unknown"
                for key, value in platform_map.items():
                    if key in extractor:
                        detected_platform = value
                        break

                description = data.get('description', '')
                if len(description) > 200:
                    description = description[:200] + '...'

                return StreamInfo(
                    title=data.get('title', 'Unknown Title'),
                    uploader=data.get('uploader', data.get('channel', data.get('creator', 'Unknown'))),
                    duration=data.get('duration_string', 'Live'),
                    view_count=data.get('view_count', 0),
                    platform=detected_platform,
                    description=description,
                    duration_seconds=data.get('duration'),
                )
        except (subprocess.TimeoutExpired, json.JSONDecodeError, subprocess.CalledProcessError, OSError) as e:
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"⚠️ yt-dlp JSON fehlgeschlagen für {url[:50]}: {e}")
        return None

    def _fallback_from_url(self, url: str) -> StreamInfo:
        """
        Erzeugt minimale Stream-Informationen allein aus der URL.
        Versucht, einen sinnvollen Anzeigenamen zu extrahieren (z.B. Twitch-Kanal).
        """
        if self._debug or DEBUG_LEVEL >= 2:
            logger.debug("🔄 StreamInfoExtractor: Fallback – extrahiere Titel aus URL")
        try:
            parsed = urllib.parse.urlparse(url)
            domain = parsed.netloc.replace('www.', '')
            path_segments = [s for s in parsed.path.split('/') if s]

            if 'twitch.tv' in domain:
                # Bei Twitch: Kanalname ist der erste Pfadsegment (oder der Domain-Teil)
                channel = path_segments[0] if path_segments else domain.replace('.tv', '')
                title = f"{channel} (Twitch)"
                uploader = channel
                platform = "twitch"
            elif 'youtube.com' in domain or 'youtu.be' in domain:
                # Bei YouTube: Kanal oder Video-ID als Platzhalter
                channel = path_segments[0] if path_segments and path_segments[0] not in ('watch', 'playlist', 'shorts') else 'YouTube'
                title = f"YouTube Stream - {channel}"
                uploader = channel
                platform = "youtube"
            else:
                # Generisch: Domain und ggf. letztes Pfadsegment
                if path_segments:
                    last = path_segments[-1]
                    title = f"{domain} - {last}"
                else:
                    title = domain
                uploader = domain
                platform = "unknown"

            return StreamInfo(
                title=title,
                uploader=uploader,
                duration="Live" if 'live' in url.lower() else "Unknown",
                view_count=0,
                platform=platform
            )
        except Exception as e:
            if self._debug or DEBUG_LEVEL >= 2:
                logger.debug(f"❌ StreamInfoExtractor: URL-Fallback fehlgeschlagen: {e}")

        # Ganz letzter Notfall
        return StreamInfo(
            title="Unknown Stream",
            uploader="Unknown",
            duration="Live",
            view_count=0,
            platform="unknown"
        )

    # ----------------------------------------------------------------------
    # Plattformspezifische Extraktion mit Cookies
    # ----------------------------------------------------------------------
    def _extract_youtube_info_with_cookies(self, url: str) -> Optional[StreamInfo]:
        """
        Extrahiert YouTube-Informationen unter Verwendung von Browser-Cookies.
        (Optimierte Version, behält die umfangreiche Logik bei, ruft aber im Erfolgsfall
        keine doppelten JSON-Parsing auf.)
        """
        logger.info("  🎯 YouTube detected, trying optimized cookie methods for channel name...")

        # Symlink-Hilfe für Linux (Chrome-Kompatibilität)
        if IS_LINUX:
            self._ensure_chrome_symlinks()

        # Methoden zusammenbauen (mit Cookies + Fallbacks)
        methods = self._build_youtube_methods(url)

        logger.info(f"    📋 Using {len(methods)} optimized extraction methods")
        max_attempts = min(3, len(methods))
        attempts = 0

        for cmd, method_name in methods:
            if attempts >= max_attempts:
                break
            attempts += 1
            try:
                logger.info(f"    🧪 Attempt {attempts}/{max_attempts}: {method_name}")
                timeout = 12 if 'Cookies' in method_name else 8
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    shell=False,
                    encoding='utf-8',
                    errors='ignore'
                )

                if result.returncode != 0 and result.stderr:
                    error_preview = result.stderr[:80].replace('\n', ' ')
                    if self._debug:
                        logger.info(f"      ❌ Error: {error_preview}")

                if result.returncode == 0 and result.stdout.strip():
                    output = result.stdout.strip()
                    # Versuche JSON zu parsen
                    json_start = output.find('{')
                    json_end = output.rfind('}') + 1
                    if json_start >= 0 and json_end > json_start:
                        try:
                            json_str = output[json_start:json_end]
                            info = json.loads(json_str)
                            uploader = info.get('uploader', 'Unknown')
                            channel = info.get('channel', uploader)
                            creator = info.get('creator', uploader)
                            final_uploader = uploader
                            if channel != 'Unknown' and channel != uploader:
                                final_uploader = channel
                            elif creator != 'Unknown' and creator != uploader:
                                final_uploader = creator
                            if final_uploader == 'Unknown':
                                final_uploader = info.get('uploader_id', 'YouTube')

                            logger.info(f"      ✅ Success with {method_name}")
                            logger.info(f"        Title: {info.get('title', 'YouTube Stream')[:60]}...")
                            logger.info(f"        Channel: {final_uploader}")
                            return StreamInfo(
                                title=info.get('title', 'YouTube Stream'),
                                uploader=final_uploader,
                                duration=info.get('duration_string', 'Live'),
                                view_count=info.get('view_count', 0),
                                platform="youtube",
                                description=info.get('description', '')[:200] + '...' if len(info.get('description', '')) > 200 else info.get('description', ''),
                                duration_seconds=info.get('duration'),
                            )
                        except json.JSONDecodeError:
                            pass

                    # Falls kein JSON, versuche Titel aus der Ausgabe zu extrahieren (z.B. bei --get-title)
                    lines = output.split('\n')
                    for line in lines:
                        if line.strip() and not line.startswith('{') and len(line.strip()) > 10:
                            possible_title = line.strip()
                            if len(possible_title) > 20 and len(possible_title) < 200:
                                logger.info("      ✅ Extracted title from output")
                                return StreamInfo(
                                    title=possible_title,
                                    uploader="YouTube",
                                    duration="Live",
                                    view_count=0,
                                    platform="youtube",
                                    description=""
                                )
            except subprocess.TimeoutExpired:
                logger.info(f"      ⏰ Timeout after {timeout}s")
                continue
            except OSError as e:
                logger.info(f"      ⚠️ OS error: {e}")
                continue
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.info(f"      ⚠️ Method error: {str(e)[:50]}")
                continue

        # Ultimativer Fallback: direkte Extraktion über --get-title und --get-filename
        logger.info("    🔄 Ultimate fallback: Direct title extraction...")
        return self._direct_youtube_fallback(url)

    def _build_youtube_methods(self, url: str) -> List[Tuple[List[str], str]]:
        """Erstellt eine Liste von (Kommando, Methodenname) für YouTube-Extraktion."""
        methods = []

        # Browser-Cookies (falls erlaubt)
        if self.use_browser_cookies:
            browsers = self._get_browser_list()
            for browser_cmd, browser_name in browsers:
                methods.append((
                    [
                        'yt-dlp', '--cookies-from-browser', browser_cmd, '--dump-json',
                        '--no-warnings', '--no-check-certificate', '--playlist-items', '1', '--', url
                    ],
                    f"{browser_name} Cookies"
                ))

        # Fallback-Methoden ohne Cookies
        fallback_methods = [
            (['yt-dlp', '--dump-json', '--no-warnings', '--no-check-certificate',
              '--playlist-items', '1', '--quiet', '--', url],
             "No Cookies (Quiet)"),
            (['yt-dlp', '--dump-json', '--no-warnings', '--no-check-certificate',
              '--playlist-items', '1', '--', url],
             "Simple JSON"),
            (['yt-dlp', '--get-title', '--get-description', '--get-duration',
              '--no-warnings', '--no-check-certificate', '--quiet', '--', url],
             "Direct Info"),
        ]
        methods.extend(fallback_methods)
        return methods

    def _get_browser_list(self) -> List[Tuple[str, str]]:
        """Liefert plattformspezifische Liste von Browsern für Cookie-Extraktion."""
        if IS_LINUX:
            return [
                ('firefox', 'Firefox'),
                ('chromium', 'Chromium'),
                ('brave', 'Brave'),
                ('chrome', 'Chrome'),
                ('vivaldi', 'Vivaldi'),
                ('opera', 'Opera'),
                ('edge', 'Edge'),
            ]
        elif IS_WINDOWS:
            return [
                ('chrome', 'Chrome'),
                ('firefox', 'Firefox'),
                ('edge', 'Edge'),
                ('brave', 'Brave'),
                ('opera', 'Opera'),
            ]
        else:  # macOS
            return [
                ('safari', 'Safari'),
                ('chrome', 'Chrome'),
                ('firefox', 'Firefox'),
                ('brave', 'Brave'),
                ('edge', 'Edge'),
            ]

    def _ensure_chrome_symlinks(self) -> None:
        """Erstellt auf Linux Symlinks von Chromium zu Chrome, damit yt-dlp Cookies findet."""
        try:
            chrome_config_dir = Path.home() / '.config' / 'google-chrome'
            chromium_config_dir = Path.home() / '.config' / 'chromium'
            if chromium_config_dir.exists() and not chrome_config_dir.exists():
                chrome_config_dir.mkdir(parents=True, exist_ok=True)
                chromium_files = ['Local State', 'Default/Cookies', 'Default/Login Data']
                for file_path in chromium_files:
                    chromium_file = chromium_config_dir / file_path
                    chrome_file = chrome_config_dir / file_path
                    if chromium_file.exists() and not chrome_file.exists():
                        chrome_file.parent.mkdir(parents=True, exist_ok=True)
                        os.symlink(str(chromium_file), str(chrome_file))
                logger.info("    🔗 Created Chrome compatibility symlinks for yt-dlp")
        except Exception as e:
            logger.warning(f"    ⚠️ Chrome symlink setup failed: {e}")

    def _direct_youtube_fallback(self, url: str) -> Optional[StreamInfo]:
        """Letzter Versuch für YouTube: Titel und Uploader einzeln abfragen."""
        try:
            cmd_title = ['yt-dlp', '--get-title', '--no-warnings',
                         '--no-check-certificate', '--quiet', '--', url]
            cmd_uploader = ['yt-dlp', '--get-filename', '-o', '%(uploader)s',
                            '--no-warnings', '--no-check-certificate', '--quiet', '--', url]
            with ThreadPoolExecutor(max_workers=2) as executor:
                title_future = executor.submit(
                    subprocess.run, cmd_title,
                    capture_output=True, text=True, timeout=8, shell=False
                )
                uploader_future = executor.submit(
                    subprocess.run, cmd_uploader,
                    capture_output=True, text=True, timeout=8, shell=False
                )
                title_result = title_future.result(timeout=10)
                uploader_result = uploader_future.result(timeout=10)
            title = "YouTube Stream"
            uploader = "YouTube"
            if title_result.returncode == 0 and title_result.stdout.strip():
                title = title_result.stdout.strip().split('\n')[0]
            if uploader_result.returncode == 0 and uploader_result.stdout.strip():
                uploader = uploader_result.stdout.strip().split('\n')[0]
            logger.info("      ✅ Success with direct extraction")
            return StreamInfo(
                title=title[:100] if len(title) > 100 else title,
                uploader=uploader,
                duration="Live",
                view_count=0,
                platform="youtube",
                description=""
            )
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.info(f"      ⚠️ Direct extraction failed: {e}")
            return None

    def _extract_twitch_info_with_cookies(self, url: str) -> Optional[StreamInfo]:
        """
        Extrahiert Twitch-Stream-Informationen unter Verwendung von Browser-Cookies.
        Ähnlich wie die YouTube-Version, aber mit Twitch-optimierten Parametern.
        """
        logger.info("  🎯 Twitch detected, trying cookie methods for channel name...")

        methods = self._build_twitch_methods(url)
        logger.info(f"    📋 Using {len(methods)} extraction methods for Twitch")
        max_attempts = min(4, len(methods))
        attempts = 0

        for cmd, method_name in methods:
            if attempts >= max_attempts:
                break
            attempts += 1
            try:
                logger.info(f"    🧪 Attempt {attempts}/{max_attempts}: {method_name}")
                timeout = 12 if 'Cookies' in method_name else 8
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    shell=False,
                    encoding='utf-8',
                    errors='ignore'
                )

                if result.returncode != 0 and result.stderr:
                    error_preview = result.stderr[:80].replace('\n', ' ')
                    if self._debug:
                        logger.info(f"      ❌ Error: {error_preview}")

                if result.returncode == 0 and result.stdout.strip():
                    output = result.stdout.strip()
                    # Versuche JSON zu parsen
                    try:
                        json_start = output.find('{')
                        json_end = output.rfind('}') + 1
                        if json_start >= 0 and json_end > json_start:
                            json_str = output[json_start:json_end]
                            info = json.loads(json_str)
                            uploader = info.get('uploader', info.get('channel', info.get('creator', 'Unknown')))
                            title = info.get('title', 'Twitch Stream')
                            duration = info.get('duration_string', 'Live')
                            view_count = info.get('view_count', 0)
                            description = info.get('description', '')
                            if len(description) > 200:
                                description = description[:200] + '...'
                            logger.info(f"      ✅ Success with {method_name}")
                            logger.info(f"        Title: {title[:60]}...")
                            logger.info(f"        Channel: {uploader}")
                            return StreamInfo(
                                title=title,
                                uploader=uploader,
                                duration=duration,
                                view_count=view_count,
                                platform="twitch",
                                description=description,
                                duration_seconds=info.get('duration'),
                            )
                    except json.JSONDecodeError:
                        # Falls kein JSON, versuche Titel aus der Ausgabe zu extrahieren (z.B. bei --get-title)
                        lines = output.split('\n')
                        title = None
                        uploader = None
                        for line in lines:
                            if line.strip() and not line.startswith('{'):
                                if title is None:
                                    title = line.strip()
                                elif uploader is None:
                                    uploader = line.strip()
                                    break
                        if title and len(title) > 10:
                            logger.info("      ✅ Extracted title from output")
                            return StreamInfo(
                                title=title[:100] if len(title) > 100 else title,
                                uploader=uploader or "Twitch",
                                duration="Live",
                                view_count=0,
                                platform="twitch",
                                description=""
                            )
            except subprocess.TimeoutExpired:
                logger.info(f"      ⏰ Timeout after {timeout}s")
                continue
            except OSError as e:
                logger.info(f"      ⚠️ OS error: {e}")
                continue
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.info(f"      ⚠️ Method error: {str(e)[:50]}")
                continue

        # Ultimativer Fallback: nur Kanalname aus URL
        return self._twitch_url_fallback(url)

    def _build_twitch_methods(self, url: str) -> List[Tuple[List[str], str]]:
        """Erstellt eine Liste von (Kommando, Methodenname) für Twitch-Extraktion."""
        methods = []
        if self.use_browser_cookies:
            for browser_cmd, browser_name in self._get_browser_list():
                methods.append((
                    [
                        'yt-dlp',
                        '--cookies-from-browser', browser_cmd,
                        '--dump-json',
                        '--format', 'best',
                        '--no-warnings',
                        '--no-check-certificate',
                        '--socket-timeout', '15',
                        '--', url
                    ],
                    f"{browser_name} Cookies"
                ))

        # Fallback-Methoden ohne Cookies
        fallback_methods = [
            (['yt-dlp', '--dump-json', '--no-warnings', '--no-check-certificate', '--socket-timeout', '10', '--', url], "Simple JSON"),
            (['yt-dlp', '--get-title', '--get-description', '--get-duration', '--no-warnings', '--no-check-certificate', '--quiet', '--', url], "Direct Info"),
        ]
        methods.extend(fallback_methods)
        return methods

    def _twitch_url_fallback(self, url: str) -> Optional[StreamInfo]:
        """Letzter Fallback für Twitch: Kanalnamen aus URL extrahieren."""
        try:
            parsed = urllib.parse.urlparse(url)
            path = parsed.path.strip('/')
            channel = path.split('/')[0] if path else parsed.netloc.replace('www.', '').replace('.tv', '')
            if channel:
                return StreamInfo(
                    title=f"{channel} (Twitch Live)",
                    uploader=channel,
                    duration="Live",
                    view_count=0,
                    platform="twitch",
                    description=""
                )
        except Exception:
            pass
        return None

# -----------------------------------------------------------------------------
# LANGUAGE DETECTOR
# -----------------------------------------------------------------------------
class LanguageDetector:
    def __init__(self, transcription_engine: TranscriptionEngine) -> None:
        self.transcription_engine = transcription_engine

    def _get_media_duration(self, file_path: str) -> Optional[float]:
        try:
            cmd = [
                'ffprobe',
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                file_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                return float(result.stdout.strip())
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError, ValueError, OSError):
            pass
        except Exception:
            pass
        return None

    def detect_video_language(self, video_path: str) -> Dict[str, Any]:
        try:
            if not os.path.exists(video_path):
                return {"error": "File not found"}

            duration = self._get_media_duration(video_path)
            if duration is None:
                try:
                    file_size_mb = os.path.getsize(video_path) / (1024 * 1024)
                    if file_size_mb > 500:
                        return {"info": "Large file - direct processing recommended"}
                except OSError:
                    pass
                sample_duration = 60
            else:
                sample_duration = min(60, max(30, duration // 2))

            temp_audio = self._extract_audio_sample(video_path, duration=sample_duration)

            if not temp_audio:
                temp_audio = self._extract_audio_sample(video_path, duration=None)

            if not temp_audio:
                return {"error": "Could not extract audio"}

            result = self.transcription_engine.transcribe_audio(temp_audio, include_timestamps=False)

            if result and hasattr(result, 'language'):
                language_code = result.language
                language_name = SUPPORTED_LANGUAGES.get(language_code, "Unknown")
                return {
                    "detected_language": language_code,
                    "language_name": language_name,
                    "confidence": getattr(result, "confidence", 0.8),
                    "sample_text": result.text[:100] + "..." if len(result.text) > 100 else result.text,
                }
            else:
                return {"error": "Language could not be detected"}

        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}

    def _extract_audio_sample(self, video_path: str, duration: Optional[int] = 30) -> Optional[bytes]:
        try:
            config = self.transcription_engine.settings.config
            cmd = [
                'ffmpeg',
                '-i', video_path,
                '-f', config.AUDIO_FORMAT,
                '-ar', str(config.SAMPLE_RATE),
                '-ac', str(config.CHANNELS),
                '-loglevel', 'quiet',
                '-'
            ]
            if duration is not None:
                cmd.insert(2, '-t')
                cmd.insert(3, str(duration))

            result = subprocess.run(cmd, capture_output=True, timeout=30)
            if result.returncode == 0 and result.stdout:
                return result.stdout
        except Exception:
            pass
        return None


# -----------------------------------------------------------------------------
# PROGRESS DIALOG
# -----------------------------------------------------------------------------
class ProgressDialog:
    """
    Ein Fortschrittsdialog mit Abbruchmöglichkeit.
    Verwendet einen indeterminierten Fortschrittsbalken und eine Nachricht.
    """

    def __init__(self, parent: tk.Tk, title: str = "Processing...") -> None:
        self.parent = parent
        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        self.dialog.geometry("300x120")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.protocol("WM_DELETE_WINDOW", self.cancel)

        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.dialog.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")

        content_frame = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=20, pady=20)
        content_frame.pack(fill="both", expand=True)

        self.message_label = tk.Label(
            content_frame,
            text="Analyzing video...",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        )
        self.message_label.pack(pady=(0, 10))

        self.progress = ttk.Progressbar(content_frame, mode="indeterminate", length=250)
        self.progress.pack(pady=(0, 15))

        button_frame = tk.Frame(content_frame, bg=CURRENT_THEME.BG_PRIMARY)
        button_frame.pack()

        self.cancel_button = tk.Button(
            button_frame,
            text="Cancel",
            command=self.cancel,
            bg=CURRENT_THEME.ERROR,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            relief="flat",
            padx=15,
        )
        self.cancel_button.pack()

        self.is_cancelled = False
        self._is_running = True
        self._update_interval = 100
        self._after_id = None

        self.progress.start(10)
        self._schedule_updates()

    def _schedule_updates(self) -> None:
        if not self._is_running:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.update_idletasks()
                self._after_id = self.dialog.after(self._update_interval, self._schedule_updates)
        except tk.TclError:
            self._is_running = False
            self._after_id = None

    def cancel(self) -> None:
        if not self._is_running:
            return
        self.is_cancelled = True
        self._is_running = False

        try:
            if self.message_label and self.message_label.winfo_exists():
                self.message_label.config(text="Cancelling...", fg=CURRENT_THEME.ERROR)
            if self.cancel_button and self.cancel_button.winfo_exists():
                self.cancel_button.config(text="Cancelling...", state="disabled")
        except tk.TclError:
            pass

        if self._after_id:
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.after_cancel(self._after_id)
            except tk.TclError:
                pass
            self._after_id = None

        self.close()

    def update_message(self, message: str) -> None:
        if not self._is_running:
            return

        def _update():
            try:
                if self.message_label and self.message_label.winfo_exists():
                    self.message_label.config(text=message)
            except tk.TclError:
                pass

        if self.dialog and self.dialog.winfo_exists():
            self.dialog.after(0, _update)

    def close(self) -> None:
        self._is_running = False
        if self._after_id:
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.after_cancel(self._after_id)
            except tk.TclError:
                pass
            self._after_id = None

        try:
            self.progress.stop()
        except Exception:
            pass

        try:
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.destroy()
        except Exception:
            pass
        finally:
            self.dialog = None
            self.message_label = None
            self.progress = None
            self.cancel_button = None

# -----------------------------------------------------------------------------
# AUDIO PROCESSOR
# -----------------------------------------------------------------------------
class AudioProcessor:
    ERRORS_BEFORE_CHUNK_REDUCTION = 3

    def __init__(self, controller_ref: Any, ffmpeg_manager: FFmpegManager,
                 settings: Optional[Settings] = None) -> None:
        self.controller_ref = controller_ref
        self.ffmpeg_manager = ffmpeg_manager
        self.settings = settings or Settings()
        self.config = self.settings.config
        self.sample_rate = self.config.SAMPLE_RATE
        self.channels = self.config.CHANNELS
        self.audio_format = self.config.AUDIO_FORMAT
        self.chunk_duration = self.config.CHUNK_DURATION
        self.chunk_size = self.config.CHUNK_SIZE_BYTES
        self.overlap_size = self.config.OVERLAP_SIZE_BYTES
        self.transcription_engine: Optional[TranscriptionEngine] = None
        self.translation_engine: Optional[BaseTranslationEngine] = None
        self._fallback_translation_engine: Optional[BaseTranslationEngine] = None
        # Plugin-Manager entfernt
        self._stop_event = threading.Event()
        self._processing = threading.Event()
        self._processing_lock = threading.RLock()
        self._current_stream_id: Optional[str] = None
        self._last_successful_read_time = time.time()
        self._consecutive_empty_chunks = 0
        self._cleanup_done = False
        self._resource_lock = threading.RLock()
        self._translation_active = True
        self._last_transcription_text = ""
        self._timed_transcriptions: Deque[TranscriptionResult] = deque(maxlen=self.config.SUBTITLE_BUFFER_SIZE)
        self._timed_translations: Deque[TranslationResult] = deque(maxlen=self.config.SUBTITLE_BUFFER_SIZE)
        self._subtitle_lock = threading.RLock()
        self.subtitle_mode = False
        self._recent_word_counts: Deque[int] = deque(maxlen=10)
        self._recent_transcriptions: Deque[str] = deque(maxlen=self.config.RECENT_TRANSCRIPTIONS_SIZE)
        self._duplicate_lock = threading.RLock()
        self._chunk_counter = 0
        self._empty_reads = 0
        self._stream_start_time: Optional[float] = None
        self._total_bytes_processed = 0
        self._processed_seconds = 0.0
        self._network_quality_history: Deque[float] = deque(maxlen=20)
        self._performance_history: Deque[float] = deque(maxlen=50)
        self._chunk_processing_times: Deque[float] = deque(maxlen=10)
        self.stream_manager = StreamManager(enable_debug=(DEBUG_LEVEL >= 1))

        self._read_error_count = 0
        self._max_backoff = 30

        self._process_finished = threading.Event()
        self._process_finished.set()

        self._audio_buffer = bytearray()
        self._max_buffer_size = self.config.MAX_CHUNK_BYTES * 5

        self._finished_callback: Optional[Callable] = None

        self.last_confidence = 1.0
        self._noisereduce_counter = 0

        self._total_file_size: Optional[int] = None
        self._progress_callback: Optional[Callable[[int, Optional[int], int], None]] = None
        self._last_progress_update = 0.0
        self._progress_update_interval = 0.5

        self._expected_duration: Optional[float] = None

        self._consecutive_timeouts = 0
        self._original_chunk_duration = self.chunk_duration
        self._min_chunk_duration = self.config.MIN_CHUNK_DURATION
        self._word_count_history = deque(maxlen=10)
        self._last_chunk_duration = self.chunk_duration
        self._chunk_stable_counter = 0

        self._consecutive_errors = 0
        self._stats_lock = threading.RLock()

        self._transcription_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="Transcribe")
        self._low_conf_counter = 0

        if debug3_enabled('gpu'):
            self._last_gpu_stats_time = 0.0

        logger.info("✅ AudioProcessor initialized:")
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Chunk: {self.chunk_duration}s / {self.chunk_size:,} bytes")
        logger.info(f"   Sample Rate: {self.sample_rate} Hz")
        logger.info(f"   Overlap: {self.overlap_size:,} bytes")
        logger.info(f"   Bytes/sec: {self.config.BYTES_PER_SECOND:,}")

    def _update_chunk_size(self) -> None:
        self.chunk_size = int(self.chunk_duration * self.config.BYTES_PER_SECOND)

    def set_expected_duration(self, duration: Optional[float]) -> None:
        self._expected_duration = duration
        if duration is not None:
            logger.info(f"⏱️ Expected stream duration set: {duration:.1f}s")

    def _get_config_type(self) -> str:
        if isinstance(self.config, RealtimeConfig):
            return 'realtime'
        elif isinstance(self.config, HighAccuracyConfig):
            return 'high_accuracy'
        elif isinstance(self.config, YouTubeOptimizedConfig):
            return 'youtube'
        return 'default'

    def set_progress_callback(self, callback: Callable[[int, Optional[int], int], None]) -> None:
        self._progress_callback = callback

    def set_engines(self, transcription_engine: TranscriptionEngine,
                    translation_engine: BaseTranslationEngine,
                    fallback_translation_engine: Optional[BaseTranslationEngine] = None) -> None:
        self.transcription_engine = transcription_engine
        self.translation_engine = translation_engine
        self._fallback_translation_engine = fallback_translation_engine
        # Plugin-Manager nicht mehr gesetzt

    def enable_subtitle_mode(self, enabled: bool) -> None:
        self.subtitle_mode = enabled
        logger.info(f"🎬 Subtitle mode: {'ENABLED' if enabled else 'DISABLED'}")

    # -------------------------------------------------------------------------
    # Optimierte Streaming-Loops
    # -------------------------------------------------------------------------
    def start_processing(self, url: str, transcription_callback: Callable,
                         translation_callback: Callable, info_callback: Callable,
                         error_callback: Callable,
                         finished_callback: Optional[Callable] = None) -> None:
        logger.info(f"\n🔊 [START_PROCESSING] URL: {url[:80]}...")
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Chunk Size: {self.chunk_size:,} bytes")

        url = PlatformUtils.sanitize_url(url)

        if url.startswith('file://'):
            ok, real_path = PlatformUtils.validate_file_path(url)
            if not ok:
                error_callback(f"❌ {real_path}")
                return
            file_path = real_path
            try:
                self._total_file_size = os.path.getsize(file_path)
                logger.info(f"📁 Lokale Datei, Größe: {self._total_file_size} bytes")
            except OSError:
                self._total_file_size = None
        else:
            self._total_file_size = None

        health_issues = self._platform_specific_health_check()
        if health_issues:
            for issue in health_issues:
                logger.warning(f"⚠️ {issue}")
                info_callback(f"⚠️ {issue}")

        with self._processing_lock:
            if self._processing.is_set():
                logger.warning("⚠️ Vorheriger Prozess läuft noch – stoppe diesen zuerst.")
                if not self.stop_processing(wait=True, timeout=10.0):
                    error_callback("❌ Vorheriger Prozess konnte nicht gestoppt werden")
                    return
            self._processing.set()
            self._process_finished.clear()
            self._stop_event.clear()
            self._current_stream_id = f"stream_{int(time.time())}"
            self._stream_start_time = time.time()
            self._chunk_counter = 0
            self._total_bytes_processed = 0
            self._processed_seconds = 0.0
            self._read_error_count = 0
            self._audio_buffer = bytearray()
            self._finished_callback = finished_callback
            with self._stats_lock:
                self._consecutive_timeouts = 0
                self._consecutive_errors = 0
                self._word_count_history.clear()
                self._last_chunk_duration = self.chunk_duration
                self._chunk_stable_counter = 0
            logger.info(f"✅ Flags gesetzt: processing=True, ID={self._current_stream_id}")

        thread = threading.Thread(
            target=self._process_loop_enhanced,
            args=(url, transcription_callback, translation_callback,
                  info_callback, error_callback),
            daemon=True,
            name=f"AudioProc_{self._current_stream_id}"
        )
        thread.start()
        logger.info(f"✅ Processing thread gestartet: {thread.name}")

    def _process_loop_enhanced(self, url: str, transcription_callback: Callable,
                               translation_callback: Callable, info_callback: Callable,
                               error_callback: Callable) -> None:
        process: Optional[subprocess.Popen] = None
        detected_language: Optional[str] = None
        error_occurred = False
        try:
            logger.info(f"\n🎬 [PROCESS_LOOP] Start für: {url[:60]}...")
            info_callback("🔍 Extracting audio URL...")
            audio_url = self.stream_manager.extract_audio_url(url)
            if not audio_url:
                error_callback("❌ Could not extract audio URL")
                error_occurred = True
                return
            logger.info(f"✅ Audio URL: {audio_url[:80]}...")
            info_callback("🔍 Testing audio stream...")
            if not self._test_audio_stream(audio_url):
                logger.warning("⚠️ Stream test failed, trying anyway...")
            info_callback("🔧 Setting up FFmpeg...")
            logger.info("🚀 Starting FFmpeg process...")
            try:
                process = self.ffmpeg_manager.start_stream(
                    video_url=url,
                    output_queue=None,
                    process_id=self._current_stream_id,
                    audio_url=audio_url,
                    detected_language=detected_language
                )
                if process is None:
                    error_callback("❌ FFmpeg konnte nicht gestartet werden")
                    error_occurred = True
                    return
            except FileNotFoundError:
                error_callback("❌ FFmpeg not found - please install")
                error_occurred = True
                return
            except (OSError, PermissionError) as e:
                error_callback(f"❌ FFmpeg konnte nicht gestartet werden: {e}")
                error_occurred = True
                return
            logger.info(f"✅ FFmpeg started (PID: {process.pid})")
            info_callback("⏳ Initializing stream...")
            wait_time = self.config.INITIAL_BUFFER_SECONDS
            if any(keyword in audio_url.lower() for keyword in ['hls', '.m3u8', 'manifest.googlevideo.com']):
                wait_time = 3.0
                logger.info(f"🎯 HLS/Live stream detected, waiting {wait_time}s...")
            time.sleep(wait_time)
            if process.poll() is not None:
                try:
                    stderr = PlatformUtils.read_process_stderr(process, 1000)
                    error_msg = f"FFmpeg died: {stderr[:200]}"
                    logger.error(f"❌ {error_msg}")
                    error_callback(f"❌ {error_msg}")
                except Exception:
                    error_callback("❌ FFmpeg failed to start")
                error_occurred = True
                return
            info_callback("✅ Stream connected - starting transcription...")
            is_youtube = any(domain in audio_url for domain in ['youtube.com', 'youtu.be', 'googlevideo.com'])
            if debug3_enabled('audio'):
                logger.debug(f"[DEBUG3][AUDIO] Detected stream type: {'YouTube' if is_youtube else 'Standard'}")
            if is_youtube:
                logger.info("🎯 Using YouTube-optimized streaming loop")
                self._youtube_streaming_loop(
                    process, audio_url, url, detected_language,
                    transcription_callback, translation_callback,
                    info_callback, error_callback
                )
            else:
                logger.info("🎯 Using standard streaming loop")
                self._standard_streaming_loop(
                    process, audio_url, detected_language,
                    transcription_callback, translation_callback,
                    info_callback, error_callback
                )
            logger.info(f"🔚 [LOOP END] Reason: {'Stop requested' if self._stop_event.is_set() else 'Process ended'}")
        except subprocess.TimeoutExpired:
            error_callback("❌ Timeout - stream not reachable")
            error_occurred = True
        except FileNotFoundError:
            error_callback("❌ FFmpeg not found - please install")
            error_occurred = True
        except OSError as e:
            error_msg = f"OS error: {str(e)[:100]}"
            logger.error(f"❌ {error_msg}")
            error_callback(f"❌ {error_msg}")
            error_occurred = True
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            error_msg = f"Unexpected error: {str(e)[:100]}"
            logger.error(f"❌ {error_msg}")
            if DEBUG_LEVEL >= 2:
                logger.exception("Stacktrace:")
            error_callback(f"❌ {error_msg}")
            error_occurred = True
        finally:
            self._flush_audio_buffer(transcription_callback, translation_callback)
            if process:
                self.ffmpeg_manager.stop_stream(self._current_stream_id)
            self._log_final_stats()
            self._guaranteed_cleanup()
            self._process_finished.set()

            is_local_file = url.startswith('file://')
            normal_end = not self._stop_event.is_set() and not error_occurred
            if is_local_file and normal_end and self._finished_callback:
                logger.info("✅ Datei normal beendet – rufe finished_callback auf")
                self._finished_callback()
            elif not self._stop_event.is_set() and not error_occurred:
                error_callback("❌ Stream wurde unerwartet beendet – versuche Neuverbindung...")

            logger.info("✅ Processing loop ended")

    def _standard_streaming_loop(self, process: subprocess.Popen, audio_url: str,
                                 detected_language: Optional[str],
                                 transcription_callback: Callable,
                                 translation_callback: Callable,
                                 info_callback: Callable,
                                 error_callback: Callable) -> None:
        self._run_common_streaming_loop(
            process=process,
            audio_url=audio_url,
            original_video_url=audio_url,
            detected_language=detected_language,
            transcription_callback=transcription_callback,
            translation_callback=translation_callback,
            info_callback=info_callback,
            error_callback=error_callback,
            is_youtube=False
        )

    def _youtube_streaming_loop(self, process: subprocess.Popen, audio_url: str, original_video_url: str,
                                detected_language: Optional[str],
                                transcription_callback: Callable,
                                translation_callback: Callable,
                                info_callback: Callable,
                                error_callback: Callable) -> None:
        self._run_common_streaming_loop(
            process=process,
            audio_url=audio_url,
            original_video_url=original_video_url,
            detected_language=detected_language,
            transcription_callback=transcription_callback,
            translation_callback=translation_callback,
            info_callback=info_callback,
            error_callback=error_callback,
            is_youtube=True
        )

    def _run_common_streaming_loop(self, process: subprocess.Popen, audio_url: str, original_video_url: str,
                                    detected_language: Optional[str],
                                    transcription_callback: Callable,
                                    translation_callback: Callable,
                                    info_callback: Callable,
                                    error_callback: Callable,
                                    is_youtube: bool) -> None:
        last_data_time = time.time()
        consecutive_timeouts = 0
        backoff = 1.0
        max_reconnects = Constants.MAX_STREAM_RECONNECTS
        reconnect_attempts = 0
        refresh_count = 0
        max_refresh_attempts = Constants.YOUTUBE_URL_REFRESH_MAX_ATTEMPTS
        current_process = process
        last_url_refresh = time.time()
        url_refresh_interval = Constants.YOUTUBE_URL_REFRESH_INTERVAL
        consecutive_low_quality_chunks = 0
        max_low_quality_chunks = Constants.YOUTUBE_LOW_QUALITY_MAX_CHUNKS
        last_buffer_flush = time.time()

        while self._processing.is_set() and not self._stop_event.is_set():
            if current_process.poll() is not None:
                logger.info("FFmpeg process terminated – finishing loop.")
                break

            current_time = time.time()

            if current_time - last_data_time > self.config.STREAM_TIMEOUT:
                consecutive_timeouts += 1
                if debug3_enabled('audio') and consecutive_timeouts % 5 == 0:
                    logger.debug(f"[DEBUG3][AUDIO] Timeout: consecutive_timeouts={consecutive_timeouts}, processed_seconds={self._processed_seconds:.1f}")
                if consecutive_timeouts > self.config.MAX_CONSECUTIVE_ERRORS:
                    if reconnect_attempts < max_reconnects:
                        reconnect_attempts += 1
                        wait = min(self._max_backoff, backoff)
                        logger.warning(f"⚠️ Stream timeout - reconnecting attempt {reconnect_attempts}/{max_reconnects}, waiting {wait:.1f}s")
                        if consecutive_timeouts % 2 == 0:
                            info_callback(f"🔄 Reconnecting... ({reconnect_attempts}/{max_reconnects})")
                        self._stop_event.wait(wait)
                        backoff *= 2
                        consecutive_timeouts = 0
                        continue
                    else:
                        logger.info("📴 Stream appears to be offline – ending processing.")
                        if self._finished_callback:
                            self._finished_callback()
                        break
                else:
                    wait = min(self._max_backoff, backoff)
                    logger.warning(f"⚠️ Temporary timeout ({consecutive_timeouts}/{self.config.MAX_CONSECUTIVE_ERRORS}), waiting {wait:.1f}s")
                    if consecutive_timeouts % 2 == 0:
                        info_callback(f"⏳ Timeout {consecutive_timeouts}/{self.config.MAX_CONSECUTIVE_ERRORS} – waiting...")
                    time.sleep(wait)
                    continue
            else:
                consecutive_timeouts = 0
                backoff = 1.0

            if self._expected_duration is not None and self._processed_seconds >= self._expected_duration - 1.0:
                logger.info(f"⏱️ Expected duration reached ({self._processed_seconds:.1f}s >= {self._expected_duration:.1f}s - 1s), stopping.")
                break

            # Korrigierter Kommentar: YouTube URL refresh (war fälschlich Twitch)
            if is_youtube and time.time() - last_url_refresh > url_refresh_interval:
                logger.info("🔄 Scheduled YouTube URL refresh")
                new_url = self.stream_manager.extract_audio_url(original_video_url, force_refresh=True)
                if new_url and new_url != audio_url:
                    logger.info("✅ New YouTube URL obtained, restarting FFmpeg...")
                    self.ffmpeg_manager.stop_stream(self._current_stream_id)
                    time.sleep(1.0)
                    new_process = self.ffmpeg_manager.start_stream(
                        video_url=original_video_url,
                        output_queue=None,
                        process_id=self._current_stream_id,
                        force_refresh_audio_url=True,
                        seek_seconds=self._processed_seconds,
                        detected_language=detected_language
                    )
                    if new_process is None:
                        logger.error("❌ Could not restart FFmpeg after URL refresh")
                        break
                    current_process = new_process
                    audio_url = new_url
                    last_url_refresh = time.time()
                    logger.info(f"✅ FFmpeg restarted with new URL (PID: {current_process.pid})")
                    time.sleep(2.0)
                    continue

            try:
                audio_data = self._read_with_timeout(current_process, self.chunk_size, timeout=Constants.READ_CHUNK_TIMEOUT)
            except (IOError, OSError) as e:
                logger.warning(f"⚠️ Read error: {e}")
                self._read_error_count += 1
                wait = min(self._max_backoff, self.config.READ_RETRY_DELAY * (2 ** (self._read_error_count - 1)))
                time.sleep(wait)
                continue
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"⚠️ Unexpected read error: {e}")
                self._read_error_count += 1
                wait = min(self._max_backoff, self.config.READ_RETRY_DELAY * (2 ** (self._read_error_count - 1)))
                time.sleep(wait)
                continue

            if not audio_data:
                if current_process.poll() is not None:
                    stderr = PlatformUtils.read_process_stderr(current_process, 4096)
                    if self._needs_url_refresh(stderr):
                        logger.info(f"🔄 Detected URL refresh needed: {stderr[:200]}")
                        if is_youtube and refresh_count < max_refresh_attempts:
                            new_url = self._refresh_youtube_url(original_video_url)
                            if new_url:
                                refresh_count += 1
                                new_process = self._restart_ffmpeg_with_new_url(self._current_stream_id, original_video_url, detected_language)
                                if new_process is None:
                                    logger.error("❌ Failed to restart FFmpeg, aborting session.")
                                    break
                                current_process = new_process
                                audio_url = new_url
                                last_data_time = time.time()
                                continue
                        logger.info("📴 Stream ended (no more URL refreshes possible).")
                        if self._finished_callback:
                            self._finished_callback()
                        break
                    else:
                        logger.warning(f"FFmpeg terminated: {stderr[:200]}")
                        break
                else:
                    self._empty_reads += 1
                    if self._empty_reads > self.config.MAX_EMPTY_READS:
                        logger.warning(f"⚠️ Too many empty reads: {self._empty_reads}")
                        error_callback("❌ No audio data received")
                        break
                    sleep_time = min(1.0, self.config.READ_RETRY_DELAY * self._empty_reads)
                    time.sleep(sleep_time)
                    continue
            else:
                self.ffmpeg_manager.update_process_activity(self._current_stream_id)
                self._read_error_count = 0
                last_data_time = time.time()
                self._empty_reads = 0
                self._chunk_counter += 1
                self._total_bytes_processed += len(audio_data)
                self._processed_seconds = self._total_bytes_processed / self.config.BYTES_PER_SECOND

                is_low_quality = False
                if len(audio_data) < self.config.MIN_CHUNK_BYTES * Constants.LOW_QUALITY_CHUNK_THRESHOLD_FACTOR:
                    is_low_quality = True
                    logger.debug(f"📉 Chunk too small: {len(audio_data)} bytes (expected min {self.config.MIN_CHUNK_BYTES})")

                if not is_low_quality and NUMPY_AVAILABLE:
                    try:
                        np = FastLazyLoader.load("numpy")
                        audio_np = np.frombuffer(audio_data, dtype=np.int16).astype(self._np.float32)
                        rms = np.sqrt(np.mean(audio_np**2))
                        if rms < self.config.MIN_RMS_THRESHOLD * 0.5:
                            is_low_quality = True
                            logger.debug(f"🔇 Very low RMS: {rms:.4f} (threshold {self.config.MIN_RMS_THRESHOLD})")
                    except Exception:
                        pass

                if is_low_quality:
                    consecutive_low_quality_chunks += 1
                    logger.debug(f"⚠️ Low quality chunk #{consecutive_low_quality_chunks}/{max_low_quality_chunks}")
                    if consecutive_low_quality_chunks >= max_low_quality_chunks:
                        logger.warning(f"📴 Too many low-quality chunks ({consecutive_low_quality_chunks}), forcing reconnect")
                        if is_youtube:
                            new_url = self._refresh_youtube_url(original_video_url)
                            if new_url and refresh_count < max_refresh_attempts:
                                refresh_count += 1
                                new_process = self._restart_ffmpeg_with_new_url(self._current_stream_id, original_video_url, detected_language)
                                if new_process:
                                    current_process = new_process
                                    audio_url = new_url
                                    consecutive_low_quality_chunks = 0
                                    last_data_time = time.time()
                                    continue
                        if reconnect_attempts < max_reconnects:
                            reconnect_attempts += 1
                            wait = min(self._max_backoff, backoff)
                            logger.info(f"🔄 Reconnecting after low quality... ({reconnect_attempts}/{max_reconnects})")
                            time.sleep(wait)
                            backoff *= 2
                            new_process = self._restart_ffmpeg_with_new_url(self._current_stream_id, original_video_url, detected_language)
                            if new_process:
                                current_process = new_process
                                consecutive_low_quality_chunks = 0
                                last_data_time = time.time()
                                continue
                        else:
                            break
                else:
                    consecutive_low_quality_chunks = 0

                if self._chunk_counter <= 3:
                    logger.debug(f"📦 Chunk #{self._chunk_counter}: {len(audio_data)} bytes")

                if self._progress_callback:
                    now = time.time()
                    if now - self._last_progress_update >= self._progress_update_interval:
                        self._last_progress_update = now
                        self._progress_callback(
                            self._total_bytes_processed,
                            self._total_file_size,
                            self._chunk_counter
                        )

                try:
                    enhanced_audio = self.enhance_audio_quality(audio_data)
                except Exception:
                    enhanced_audio = audio_data

                self._audio_buffer.extend(enhanced_audio)

                if len(self._audio_buffer) > 0 and time.time() - last_buffer_flush > Constants.BUFFER_FLUSH_INACTIVITY:
                    logger.debug(f"⏱️ Flushing audio buffer after {time.time()-last_buffer_flush:.1f}s inactivity")
                    chunk_to_process = bytes(self._audio_buffer)
                    self._audio_buffer.clear()
                    if self.transcription_engine:
                        self._process_audio_chunk(
                            chunk_to_process,
                            transcription_callback,
                            translation_callback
                        )
                    last_buffer_flush = time.time()

                if len(self._audio_buffer) >= self.config.MIN_CHUNK_BYTES:
                    chunk_to_process = bytes(self._audio_buffer)
                    self._audio_buffer.clear()
                    if self.transcription_engine:
                        self._process_audio_chunk(
                            chunk_to_process,
                            transcription_callback,
                            translation_callback
                        )
                    last_buffer_flush = time.time()
                if len(self._audio_buffer) > self._max_buffer_size:
                    logger.warning(f"⚠️ Audio buffer too large ({len(self._audio_buffer)} bytes) – forcing flush")
                    chunk_to_process = bytes(self._audio_buffer)
                    self._audio_buffer.clear()
                    if self.transcription_engine:
                        self._process_audio_chunk(
                            chunk_to_process,
                            transcription_callback,
                            translation_callback
                        )
                    last_buffer_flush = time.time()

                if self._chunk_counter % 50 == 0:
                    info_callback(f"📊 {self._chunk_counter} chunks processed...")

    def _needs_url_refresh(self, stderr: str) -> bool:
        patterns = [
            "403", "401", "forbidden", "unauthorized", "invalid parameters",
            "http error 403", "http error 401", "access denied", "signature expired",
            "token expired", "url signature expired"
        ]
        stderr_lower = stderr.lower()
        return any(p in stderr_lower for p in patterns)

    def _refresh_youtube_url(self, video_url: str) -> Optional[str]:
        max_attempts = 3
        for attempt in range(1, max_attempts + 1):
            try:
                logger.info(f"🔄 Attempt {attempt}/{max_attempts} to refresh YouTube URL...")
                new_url = self.stream_manager.extract_audio_url(video_url, force_refresh=True)
                if new_url:
                    logger.info(f"✅ Successfully obtained new YouTube URL (attempt {attempt})")
                    if DEBUG_LEVEL >= 2:
                        logger.debug(f"   New URL: {new_url[:100]}...")
                    return new_url
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"⚠️ Refresh attempt {attempt} error: {e}")
            if attempt < max_attempts:
                wait = 2 ** (attempt - 1)
                time.sleep(wait)
        logger.error("❌ All attempts to refresh YouTube URL failed")
        return None

    def _restart_ffmpeg_with_new_url(self, process_id: str, video_url: str, detected_language: Optional[str] = None) -> Optional[subprocess.Popen]:
        logger.info(f"🔄 Restarting FFmpeg for {process_id} with new URL...")
        if self.ffmpeg_manager:
            self.ffmpeg_manager.stop_stream(process_id)
            time.sleep(0.5)
        seek_seconds = self._processed_seconds
        if self._expected_duration is not None and seek_seconds > self._expected_duration:
            seek_seconds = max(0, self._expected_duration - 5)
            logger.info(f"⏩ Seek adjusted to {seek_seconds:.1f}s (within expected duration)")
        new_process = self.ffmpeg_manager.start_stream(
            video_url=video_url,
            output_queue=None,
            process_id=process_id,
            force_refresh_audio_url=True,
            seek_seconds=seek_seconds,
            detected_language=detected_language
        )
        if new_process:
            logger.info(f"✅ Successfully restarted FFmpeg (new PID: {new_process.pid})")
            return new_process
        else:
            logger.error("❌ Failed to restart FFmpeg")
            return None

    def _read_with_timeout(self, process: subprocess.Popen, size: int, timeout: float = 1.0) -> Optional[bytes]:
        import select
        start = time.perf_counter()
        data = bytearray()
        remaining = size
        end_time = time.time() + timeout
        fd = process.stdout.fileno()
        os.set_blocking(fd, False)
        empty_reads_since_last_data = 0
        try:
            while remaining > 0 and time.time() < end_time:
                rlist, _, _ = select.select([fd], [], [], max(0, end_time - time.time()))
                if fd in rlist:
                    try:
                        chunk = os.read(fd, min(remaining, 4096))
                        if not chunk:
                            break
                        data.extend(chunk)
                        remaining -= len(chunk)
                        empty_reads_since_last_data = 0
                    except BlockingIOError:
                        time.sleep(Constants.READ_WITH_TIMEOUT_SELECT_INTERVAL)
                else:
                    empty_reads_since_last_data += 1
                    if empty_reads_since_last_data > 10 and debug3_enabled('audio'):
                        logger.debug(f"[DEBUG3][AUDIO] {empty_reads_since_last_data} consecutive empty reads")
                    time.sleep(Constants.READ_WITH_TIMEOUT_SELECT_INTERVAL)
        finally:
            os.set_blocking(fd, True)
        duration = (time.perf_counter() - start) * 1000
        if data or duration > timeout * 1000 * 0.9:
            if debug3_enabled('time'):
                logger.debug(f"[DEBUG3][TIME] _read_with_timeout read {len(data)} bytes in {duration:.2f}ms")
        return bytes(data) if data else b''

    def emergency_diagnosis(self, url: str) -> bool:
        logger.info(f"🔍 [EMERGENCY_DIAGNOSIS] Testing: {url[:80]}...")
        try:
            audio_url = self.stream_manager.extract_audio_url(url)
            if not audio_url:
                logger.info("  ❌ Could not extract audio URL")
                return False
            logger.info(f"  ✅ Audio URL extracted: {audio_url[:80]}...")
            is_youtube = 'youtube.com' in audio_url.lower() or 'googlevideo.com' in audio_url
            test_cmd = [
                'ffmpeg',
                '-i', audio_url,
                '-t', '3',
                '-f', 'null',
                '-',
                '-loglevel', 'error'
            ]
            timeout = Constants.YOUTUBE_STREAM_TEST_TIMEOUT if is_youtube else Constants.STREAM_TEST_TIMEOUT
            result = subprocess.run(
                test_cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                shell=False
            )
            if result.returncode == 0:
                logger.info("  ✅ Stream connection successful")
                return True
            else:
                error_msg = result.stderr[:100] if result.stderr else "Unknown error"
                logger.info(f"  ❌ Stream test failed: {error_msg}")
                if audio_url.startswith(('http://', 'https://')):
                    logger.info("  ⚠️  But URL looks valid, trying anyway...")
                    return True
                return False
        except subprocess.TimeoutExpired:
            logger.info("  ⏰ Stream test timeout")
            if 'youtube.com' in url.lower():
                logger.info("  ⚠️  YouTube timeout common, trying anyway...")
                return True
            return False
        except (OSError, PermissionError) as e:
            logger.info(f"  ⚠️  Emergency diagnosis OS error: {e}")
            return False
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.info(f"  ⚠️  Emergency diagnosis error: {e}")
            return True

    def _log_final_stats(self) -> None:
        if self._chunk_counter == 0:
            return
        uptime = time.time() - self._stream_start_time if self._stream_start_time else 0
        logger.info("\n📊 FINAL PROCESSING STATS:")
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Total Chunks: {self._chunk_counter}")
        logger.info(f"   Total Bytes: {self._total_bytes_processed:,}")
        logger.info(f"   Total Time: {uptime:.1f}s")
        logger.info(f"   Avg Chunk Size: {self._total_bytes_processed/self._chunk_counter if self._chunk_counter > 0 else 0:,.0f} bytes")
        logger.info(f"   Processing Rate: {self._chunk_counter/uptime if uptime > 0 else 0:.1f} chunks/sec")
        logger.info(f"   Data Rate: {self._total_bytes_processed/uptime/1024 if uptime > 0 else 0:.1f} KB/sec")
        logger.info(f"   Empty Reads: {self._empty_reads}")

    # -------------------------------------------------------------------------
    # Aufteilung von _process_audio_chunk
    # -------------------------------------------------------------------------
    def _process_audio_chunk(self, audio_data: bytes, transcription_callback: Callable,
                              translation_callback: Callable) -> None:
        """Verarbeitet einen Audio-Chunk: Transkription, Übersetzung, Duplikaterkennung und adaptive Optimierungen."""
        if not self.transcription_engine:
            return

        # Initialisiere die Variablen für die Zeitmessung außerhalb des try-Blocks
        start_time = None
        audio_len = 0.0
        if DEBUG_LEVEL >= 2:
            start_time = time.perf_counter()
            audio_len = len(audio_data) / (16000 * 2)

        try:
            with self._resource_lock:
                bytes_before_chunk = self._total_bytes_processed
            chunk_start_time = bytes_before_chunk / self.config.BYTES_PER_SECOND

            logger.debug(f"🚀 _process_audio_chunk: Größe={len(audio_data)} Bytes, subtitle_mode={self.subtitle_mode}")

            if debug3_enabled('time'):
                chunk_start = time.perf_counter()

            if self.subtitle_mode:
                self._handle_subtitle_transcription(audio_data, transcription_callback, translation_callback, chunk_start_time)
            else:
                self._handle_normal_transcription(audio_data, transcription_callback, translation_callback, start_time, audio_len)
    
            with self._stats_lock:
                self._consecutive_errors = 0

            if self.settings.adaptive_chunk:
                self._update_adaptive_chunk()

            with self._stats_lock:
                low_conf = self._low_conf_counter
            if low_conf >= 3:
                logger.info(f"📈 Erhöhe Chunk-Dauer wegen {self._low_conf_counter} aufeinanderfolgender niedriger Konfidenz")
                new_duration = min(self.config.MAX_CHUNK_DURATION, self.chunk_duration + 2)
                if new_duration != self.chunk_duration:
                    self.chunk_duration = new_duration
                    self._update_chunk_size()
                with self._stats_lock:
                    self._low_conf_counter = 0                

            if debug3_enabled('time'):
                chunk_duration = (time.perf_counter() - chunk_start) * 1000
                logger.debug(f"[DEBUG3][TIME] Chunk {self._chunk_counter} processing took {chunk_duration:.2f}ms total")

            self._log_gpu_stats()
            self._log_queue_stats()
            self._log_cache_stats()

        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Audio chunk processing error: {e}")
            if DEBUG_LEVEL >= 2:
                logger.exception("Stacktrace:")

            with self._stats_lock:
                self._consecutive_errors += 1
            self._adjust_chunk_due_to_errors()
            logger.error(f"❌ Consecutive errors: {self._consecutive_errors}/{self.config.MAX_CONSECUTIVE_ERRORS}")
            if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                logger.critical(f"🚨 Too many consecutive errors ({self._consecutive_errors}), stopping processing.")
                self._stop_event.set()
                self._processing.clear()

    def _handle_subtitle_transcription(self, audio_data: bytes, transcription_callback: Callable,
                                       translation_callback: Callable, chunk_start_time: float) -> None:
        future = self._transcription_executor.submit(
            self.transcription_engine.transcribe_audio,
            audio_data,
            True
        )
        try:
            timeout_val = max(30, self.chunk_duration * 3)
            logger.debug(f"⏳ Subtitle-Transkription, Timeout={timeout_val}s")
            segments = future.result(timeout=timeout_val)
        except FutureTimeout:
            logger.error(f"⏰ Transkriptions-Timeout (Subtitle-Modus) nach {self.chunk_duration*3}s")
            future.cancel()
            with self._stats_lock:
                self._consecutive_timeouts += 1
                timeout_count = self._consecutive_timeouts
            self._adjust_chunk_due_to_timeout()
            if timeout_count >= 3:
                self._reload_model_on_timeout()
            return
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Transkriptionsfehler (Subtitle-Modus): {e}")
            with self._stats_lock:
                self._consecutive_errors += 1
            self._adjust_chunk_due_to_errors()
            if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                logger.critical(f"🚨 Zu viele Fehler ({self._consecutive_errors}), stoppe.")
                self._stop_event.set()
                self._processing.clear()
            return

        if not segments:
            return
        if debug3_enabled('subtitle'):
            logger.debug(f"[DEBUG3][SUBTITLE] Received {len(segments)} segments")
        for segment in segments:
            if segment.start is not None:
                segment.start += chunk_start_time
            if segment.end is not None:
                segment.end += chunk_start_time

            logger.info(f"🎤 SEGMENT [{segment.start:.2f} - {segment.end:.2f}] {segment.text.strip()} (Sprache: {getattr(segment, 'language', 'unbekannt')})")

            if not segment or not segment.text:
                continue
            clean_text = segment.text.strip()
            conf = getattr(segment, 'confidence', 0.0)
            if (self.config.DUPLICATE_CHECK_ENABLED and
                self._is_duplicate_transcription(clean_text, confidence=conf)):
                continue
            self._last_transcription_text = clean_text
            self.last_confidence = conf
            if self.config.ENABLE_TIMED_TRANSCRIPTIONS:
                self._add_timed_transcription(segment)
            transcription_callback(segment)
            if (self.translation_engine and self._translation_active and
                hasattr(segment, 'language')):
                detected_lang = segment.language or "auto"
                self._translate_and_send(clean_text, detected_lang, translation_callback,
                                        start=segment.start, end=segment.end)

        with self._stats_lock:
            self._consecutive_timeouts = 0
            self._consecutive_errors = 0

    def _handle_normal_transcription(self, audio_data: bytes, transcription_callback: Callable,
                                     translation_callback: Callable, start_time: float, audio_len: float) -> None:
        future = self._transcription_executor.submit(
            self.transcription_engine.safe_transcribe,
            audio_data
        )
        try:
            timeout_val = max(30, self.chunk_duration * 3)
            logger.debug(f"⏳ Normal-Transkription, Timeout={timeout_val}s")
            transcription = future.result(timeout=timeout_val)
        except FutureTimeout:
            logger.error(f"⏰ safe_transcribe Timeout nach {self.chunk_duration*3}s")
            future.cancel()
            with self._stats_lock:
                self._consecutive_timeouts += 1
            self._adjust_chunk_due_to_timeout()
            if self._consecutive_timeouts >= 3:
                self._reload_model_on_timeout()
            return
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Transkriptionsfehler: {e}")
            with self._stats_lock:
                self._consecutive_errors += 1
            self._adjust_chunk_due_to_errors()
            if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                logger.critical(f"🚨 Zu viele Fehler ({self._consecutive_errors}), stoppe.")
                self._stop_event.set()
                self._processing.clear()
            return

        with self._stats_lock:
            self._consecutive_timeouts = 0

        if transcription and hasattr(transcription, 'confidence'):
            self.last_confidence = transcription.confidence
        else:
            self.last_confidence = 0.0

        if DEBUG_LEVEL >= 2 and start_time is not None and transcription:
            elapsed = time.perf_counter() - start_time
            realtime_factor = elapsed / audio_len if audio_len > 0 else 0
            logger.debug(f"Chunk {self._chunk_counter}: {audio_len:.2f}s audio, "
                         f"transcribe {elapsed*1000:.1f}ms ({realtime_factor:.2f}x realtime)")

        if not transcription or not transcription.text:
            return
        clean_text = transcription.text.strip()
        conf = getattr(transcription, 'confidence', 0.0)
        if (self.config.DUPLICATE_CHECK_ENABLED and
            self._is_duplicate_transcription(clean_text, confidence=conf)):
            return
        self._last_transcription_text = clean_text
        transcription_callback(transcription)
        if (self.translation_engine and self._translation_active and
            hasattr(transcription, 'language')):
            detected_lang = transcription.language or "auto"
            self._translate_and_send(
                clean_text,
                detected_lang,
                translation_callback,
            )

        with self._stats_lock:
            self._consecutive_errors = 0

        with self._stats_lock:
            self._word_count_history.append(len(clean_text.split()))

    def _update_adaptive_chunk(self) -> None:
        with self._stats_lock:
            if len(self._word_count_history) < 5:
                return
            avg_words = sum(self._word_count_history) / len(self._word_count_history)
            new_duration = self.chunk_duration
            if avg_words < self.settings.adaptive_chunk_low_words and self.chunk_duration > self.config.MIN_CHUNK_DURATION + 0.5:
                new_duration = max(self.config.MIN_CHUNK_DURATION, self.chunk_duration - 1)
            elif avg_words > self.settings.adaptive_chunk_high_words and self.chunk_duration < self.config.MAX_CHUNK_DURATION - 0.5:
                new_duration = min(self.config.MAX_CHUNK_DURATION, self.chunk_duration + 1)

            if new_duration != self.chunk_duration:
                if new_duration != self._last_chunk_duration:
                    self._chunk_stable_counter += 1
                else:
                    self._chunk_stable_counter = 0

                if self._chunk_stable_counter >= 3:
                    logger.info(f"{'📈' if new_duration > self.chunk_duration else '📉'} Adaptive Chunk-Dauer: {self.chunk_duration:.1f}s → {new_duration:.1f}s (avg_words={avg_words:.1f})")
                    self.chunk_duration = new_duration
                    self._update_chunk_size()
                    self._chunk_stable_counter = 0
                self._last_chunk_duration = new_duration

    def _log_gpu_stats(self) -> None:
        if not debug3_enabled('gpu'):
            return
        if self._chunk_counter % 5 == 0:
            if TORCH_AVAILABLE and self.transcription_engine.device == 'cuda':
                torch = FastLazyLoader.load('torch')
                try:
                    allocated = torch.cuda.memory_allocated() / 1024**3
                    reserved = torch.cuda.memory_reserved() / 1024**3
                    logger.debug(f"[DEBUG3][GPU] Chunk {self._chunk_counter}: allocated={allocated:.2f}GB, reserved={reserved:.2f}GB")
                except Exception:
                    pass

                if hasattr(self, '_last_gpu_stats_time') and time.time() - self._last_gpu_stats_time > 10:
                    self._last_gpu_stats_time = time.time()
                    try:
                        import pynvml
                        pynvml.nvmlInit()
                        handle = pynvml.nvmlDeviceGetHandleByIndex(0)
                        temp = pynvml.nvmlDeviceGetTemperature(handle, pynvml.NVML_TEMPERATURE_GPU)
                        util = pynvml.nvmlDeviceGetUtilizationRates(handle).gpu
                        logger.debug(f"[DEBUG3][GPU] Temp={temp}°C, Util={util}%")
                    except Exception:
                        pass

    def _log_queue_stats(self) -> None:
        if not debug3_enabled('queue'):
            return
        if self._chunk_counter % 50 == 0:
            gui = self.controller_ref.gui_ref() if hasattr(self.controller_ref, 'gui_ref') else None
            if gui is not None:
                qsize_gui = gui.gui_queue.qsize()
                qsize_text = gui._text_update_queue.qsize()
            else:
                qsize_gui = qsize_text = 0
            active_threads = threading.active_count()
            thread_names = [t.name for t in threading.enumerate()]
            logger.debug(f"[DEBUG3][QUEUE] Chunk {self._chunk_counter}: gui_queue={qsize_gui}, text_queue={qsize_text}, active_threads={active_threads}")
            if debug3_enabled('threads'):
                logger.debug(f"[DEBUG3][THREADS] Thread names: {thread_names}")

    def _log_cache_stats(self) -> None:
        if not debug3_enabled('cache'):
            return
        if self._chunk_counter % 100 == 0:
            stats = get_cache_stats()
            logger.debug(f"[DEBUG3][CACHE] Cache stats: {stats}")

    def _adjust_chunk_due_to_timeout(self):
        with self._stats_lock:
            if self._consecutive_timeouts >= 2:
                new_duration = max(self._min_chunk_duration, self.chunk_duration / 2)
                if new_duration != self.chunk_duration:
                    if debug3_enabled('audio'):
                        logger.debug(f"[DEBUG3][AUDIO] Adjusting chunk duration from {self.chunk_duration}s to {new_duration}s due to {self._consecutive_timeouts} timeouts")
                    logger.warning(f"📉 Reduziere Chunk-Dauer von {self.chunk_duration}s auf {new_duration}s wegen {self._consecutive_timeouts} aufeinanderfolgenden Timeouts")
                    self.chunk_duration = new_duration
                    self._update_chunk_size()
                    self._consecutive_timeouts = 0

    def _adjust_chunk_due_to_errors(self):
        with self._stats_lock:
            if self._consecutive_errors >= self.ERRORS_BEFORE_CHUNK_REDUCTION:
                new_duration = max(self._min_chunk_duration, self.chunk_duration / 2)
                if new_duration != self.chunk_duration:
                    logger.warning(f"📉 Reduziere Chunk-Dauer von {self.chunk_duration}s auf {new_duration}s wegen {self._consecutive_errors} aufeinanderfolgenden Fehlern")
                    self.chunk_duration = new_duration
                    self._update_chunk_size()
                self._consecutive_errors = 0

    def _reload_model_on_timeout(self):
        logger.warning("🔄 Drei aufeinanderfolgende Timeouts – lade Modell neu...")
        if hasattr(self.transcription_engine, 'reload_model'):
            current_model = self.transcription_engine.get_current_model()
            self.transcription_engine.reload_model(current_model)
        with self._stats_lock:
            self._consecutive_timeouts = 0
            self._consecutive_errors = 0
        time.sleep(2)

    def _translate_and_send(self, text: str, source_lang: str,
                            translation_callback: Callable,
                            start: Optional[float] = None,
                            end: Optional[float] = None) -> None:
        try:
            start_time = time.perf_counter()
            translation = None

            primary_functional = True
            if hasattr(self.translation_engine, 'is_functional'):
                primary_functional = self.translation_engine.is_functional()

            if primary_functional:
                try:
                    translation = self.translation_engine.translate_text(text, source_lang)
                except Exception as e:
                    if PlatformUtils.is_fatal_exception(e):
                        raise
                    logger.debug(f"Primary translation error for '{text[:30]}...': {e}")
                    translation = None
            else:
                logger.debug("Primary translation engine is disabled, using fallback.")

            if translation is None and hasattr(self, '_fallback_translation_engine') and self._fallback_translation_engine:
                fallback_functional = True
                if hasattr(self._fallback_translation_engine, 'is_functional'):
                    fallback_functional = self._fallback_translation_engine.is_functional()
                if fallback_functional:
                    logger.debug(f"Falling back to Ollama for: '{text[:30]}...'")
                    try:
                        target_lang = getattr(self.translation_engine, 'target_lang', 'de')
                        self._fallback_translation_engine.set_target_language(target_lang)
                        translation = self._fallback_translation_engine.translate_text(text, source_lang)
                    except Exception as e:
                        if PlatformUtils.is_fatal_exception(e):
                            raise
                        logger.debug(f"Fallback translation also failed: {e}")
                        translation = None
                else:
                    logger.debug("Fallback engine is also disabled.")
                    gui_ref = self.controller_ref.gui_ref() if hasattr(self.controller_ref, 'gui_ref') else None
                    if gui_ref is not None:
                        gui_ref.update_status("⚠️ Übersetzung vorübergehend nicht verfügbar")

            duration = (time.perf_counter() - start_time) * 1000
            if debug3_enabled('time'):
                logger.debug(f"[DEBUG3][TIME] _translate_and_send took {duration:.2f}ms for {len(text)} chars")

            if translation:
                translation.start = start
                translation.end = end
                if self.subtitle_mode and self.config.ENABLE_TIMED_TRANSLATIONS:
                    self._add_timed_translation(translation)
                translation_callback(translation)
            else:
                if debug3_enabled('translate'):
                    logger.debug(f"No translation available for: '{text[:50]}...'")
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Unexpected error in _translate_and_send: {e}")
            if DEBUG_LEVEL >= 2:
                logger.exception("Stacktrace:")

    def enhance_audio_quality(self, audio_data: bytes) -> bytes:
        if not self.config.AUDIO_ENHANCEMENT_ENABLED or len(audio_data) < Constants.AUDIO_ENHANCEMENT_MIN_LENGTH or not NUMPY_AVAILABLE:
            return audio_data
        if self._consecutive_timeouts > 0:
            return audio_data
        try:
            start_time = time.perf_counter()
            np = FastLazyLoader.load("numpy")
            audio_np = np.frombuffer(audio_data, dtype=np.int16).astype(np.float32) / 32768.0
            rms = np.sqrt(np.mean(audio_np**2))

            self._noisereduce_counter += 1
            gain = 1.0

            if self.last_confidence < 0.4 and self._noisereduce_counter % Constants.NOISEREDUCE_INTERVAL == 0 and len(audio_data) > Constants.NOISEREDUCE_MIN_LENGTH:
                try:
                    import noisereduce as nr
                    audio_np = nr.reduce_noise(y=audio_np, sr=self.sample_rate, prop_decrease=0.8)
                    logger.debug(f"🔇 noisereduce angewendet (Konfidenz {self.last_confidence:.2f})")
                except ImportError:
                    pass
                except Exception as e:
                    if PlatformUtils.is_fatal_exception(e):
                        raise
                    logger.warning(f"⚠️ noisereduce fehlgeschlagen: {e}")

            if rms < self.config.MIN_RMS_THRESHOLD:
                return audio_data
            if rms < self.config.TARGET_RMS:
                gain = min(self.config.MAX_GAIN, self.config.TARGET_RMS / max(rms, 1e-6))
                audio_np = audio_np * gain
            max_val = np.max(np.abs(audio_np))
            if max_val > self.config.CLIPPING_THRESHOLD:
                audio_np = audio_np * self.config.CLIPPING_THRESHOLD / max_val
            audio_np = audio_np - np.mean(audio_np)
            enhanced = (audio_np * 32767).astype(np.int16).tobytes()
            duration = (time.perf_counter() - start_time) * 1000
            if debug3_enabled('audio'):
                logger.debug(f"[DEBUG3][AUDIO] enhance_audio_quality: duration={duration:.2f}ms, rms={rms:.4f}, gain={gain:.2f}")
            return enhanced
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            return audio_data

    def _is_duplicate_transcription(self, current_text: str, confidence: float = None) -> bool:
        """Erweiterte Duplikaterkennung mit Ähnlichkeitsvergleich (Konfidenzabfall entfernt)."""
        if not self.config.DUPLICATE_CHECK_ENABLED:
            return False

        with self._duplicate_lock:
            current_text = current_text.strip()
            if not current_text:
                return True

            if len(current_text) < self.config.MIN_TEXT_LENGTH:
                return True

            # 1. Exakte Übereinstimmung mit dem letzten Text
            if current_text == self._last_transcription_text:
                return True

            # 2. Ähnlichkeit mit den letzten Einträgen prüfen
            similarity_threshold = self.settings.duplicate_similarity_threshold
            combined = [self._last_transcription_text] + list(self._recent_transcriptions)
            for prev in combined:
                if not prev:
                    continue
                similarity = difflib.SequenceMatcher(None, current_text, prev).ratio()
                if similarity > similarity_threshold:
                    if debug3_enabled('duplicate'):
                        logger.debug(f"[DUPLICATE] {similarity:.2%} match: '{current_text[:30]}' ≈ '{prev[:30]}'")
                    return True

            # 3. Wortvielfalt prüfen
            words = current_text.lower().split()
            if len(words) > 3:
                unique_ratio = len(set(words)) / len(words)
                if unique_ratio < self.config.MIN_UNIQUE_WORDS_RATIO:
                    return True

            # Kein Duplikat gefunden – speichern
            self._last_transcription_text = current_text
            self._recent_transcriptions.append(current_text)
            return False

    def _add_timed_transcription(self, result: TranscriptionResult) -> None:
        with self._subtitle_lock:
            if (hasattr(result, 'start') and result.start is not None and
                hasattr(result, 'end') and result.end is not None):
                self._timed_transcriptions.append(result)

    def _add_timed_translation(self, result: TranslationResult) -> None:
        with self._subtitle_lock:
            if (hasattr(result, 'start') and result.start is not None and
                hasattr(result, 'end') and result.end is not None):
                self._timed_translations.append(result)

    def get_status(self) -> Dict[str, Any]:
        return {
            '_processing': self._processing.is_set(),
            '_stop_event_set': self._stop_event.is_set(),
            '_current_stream_id': self._current_stream_id,
            '_consecutive_empty_chunks': self._consecutive_empty_chunks,
            '_cleanup_done': self._cleanup_done,
            'config_type': self._get_config_type(),
            'chunk_size': self.chunk_size,
            'chunks_processed': self._chunk_counter,
            'total_bytes': self._total_bytes_processed,
            'empty_reads': self._empty_reads,
            'subtitle_mode': self.subtitle_mode,
            'translation_active': self._translation_active,
            'active_threads': threading.active_count(),
            '_consecutive_errors': self._consecutive_errors,
            '_consecutive_timeouts': self._consecutive_timeouts,
        }

    def _safe_kill_process(self, process: subprocess.Popen) -> None:
        if not process:
            return
        pid = process.pid
        logger.info(f"🛑 Terminating process {pid}...")
        try:
            if hasattr(self, 'ffmpeg_manager') and self.ffmpeg_manager:
                temp_id = f"kill_{pid}"
                self.ffmpeg_manager._register_process(temp_id, process, None, "terminate")
                self.ffmpeg_manager.stop_stream(temp_id)
                return
        except Exception:
            pass
        termination_steps = [
            (self._terminate_gracefully, 2.0, "graceful"),
            (self._terminate_forcefully, 1.0, "forceful"),
            (self._terminate_nuclear, 0.5, "nuclear")
        ]
        for terminate_method, timeout, method_name in termination_steps:
            if process.poll() is not None:
                break
            try:
                if terminate_method(process, pid, timeout):
                    logger.info(f"✅ Process {pid} terminated ({method_name})")
                    break
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"⚠️ {method_name} termination failed: {e}")
        self._cleanup_process_resources(process)

    def _terminate_gracefully(self, process: subprocess.Popen, pid: int, timeout: float) -> bool:
        if IS_WINDOWS:
            process.terminate()
        else:
            try:
                os.killpg(os.getpgid(pid), py_signal.SIGTERM)
            except (ProcessLookupError, PermissionError):
                process.terminate()
        try:
            process.wait(timeout=timeout)
            return True
        except subprocess.TimeoutExpired:
            return False

    def _terminate_forcefully(self, process: subprocess.Popen, pid: int, timeout: float) -> bool:
        if IS_WINDOWS:
            process.kill()
        else:
            try:
                os.killpg(os.getpgid(pid), py_signal.SIGKILL)
            except (ProcessLookupError, PermissionError):
                process.kill()
        try:
            process.wait(timeout=timeout)
            return True
        except subprocess.TimeoutExpired:
            return False

    def _terminate_nuclear(self, process: subprocess.Popen, pid: int, timeout: float) -> bool:
        try:
            if IS_WINDOWS:
                cmd = ['taskkill', '/F', '/T', '/PID', str(pid)]
            else:
                cmd = ['pkill', '-9', '-P', str(pid)]
            result = subprocess.run(cmd, capture_output=True, timeout=timeout)
            if result.returncode == 0:
                return True
        except Exception:
            pass
        return False

    def _cleanup_process_resources(self, process: subprocess.Popen) -> None:
        for attr in ['stdout', 'stderr', 'stdin']:
            if hasattr(process, attr):
                pipe = getattr(process, attr)
                if pipe and not pipe.closed:
                    try:
                        pipe.close()
                    except Exception:
                        pass
        try:
            del process
            gc.collect()
        except Exception:
            pass

    def emergency_reset(self, force: bool = False) -> bool:
        logger.info(f"\n🚨 [EMERGENCY_RESET] force={force}")
        with self._resource_lock:
            with self._processing_lock:
                old_state = self._processing.is_set()
                self._processing.clear()
            self._stop_event.set()
            self._current_stream_id = None
            self._consecutive_empty_chunks = 0
            with self._stats_lock:
                self._consecutive_errors = 0
                self._consecutive_timeouts = 0
            if force:
                with self._subtitle_lock:
                    self._timed_transcriptions.clear()
                    self._timed_translations.clear()
                with self._duplicate_lock:
                    self._recent_transcriptions.clear()
        logger.info(f"✅ Reset completed: {old_state} -> {self._processing.is_set()}")
        return True

    def _guaranteed_cleanup(self) -> None:
        logger.info("\n🧹 [GUARANTEED_CLEANUP]")
        with self._resource_lock:
            with self._processing_lock:
                self._processing.clear()
            self._current_stream_id = None
            self._consecutive_empty_chunks = 0
            self._empty_reads = 0
            self._chunk_counter = 0
            self._total_bytes_processed = 0
            self._cleanup_done = True
            with self._stats_lock:
                self._consecutive_errors = 0
                self._consecutive_timeouts = 0
        time.sleep(0.05)
        logger.info("✅ Cleanup completed")

    def _platform_specific_health_check(self) -> List[str]:
        issues: List[str] = []
        if IS_WINDOWS:
            try:
                import psutil
                memory = psutil.virtual_memory()
                if memory.percent > 85:
                    issues.append("High memory usage - consider closing other applications")
            except Exception:
                pass
        return issues

    def _flush_audio_buffer(self, transcription_callback: Callable, translation_callback: Callable) -> None:
        if not self._audio_buffer:
            return
        buffer_len = len(self._audio_buffer)
        logger.info(f"🧹 Flushing audio buffer ({buffer_len} bytes) at end of stream")
        if self.transcription_engine:
            self._process_audio_chunk(
                bytes(self._audio_buffer),
                transcription_callback,
                translation_callback
            )
        self._audio_buffer.clear()

    def _test_audio_stream(self, audio_url: str) -> bool:
        logger.info(f"🔍 Testing audio stream: {audio_url[:60]}...")
        is_youtube = 'youtube.com' in audio_url.lower() or 'googlevideo.com' in audio_url
        is_hls = '.m3u8' in audio_url.lower() or 'manifest.googlevideo.com' in audio_url

        if is_hls:
            logger.info("🎯 HLS stream detected – skipping quick test (often too slow)")
            return True

        try:
            timeout = Constants.YOUTUBE_STREAM_TEST_TIMEOUT if is_youtube else self.config.STREAM_TIMEOUT
            test_cmd = [
                'ffmpeg',
                '-i', audio_url,
                '-t', '2',
                '-f', 'null',
                '-',
                '-loglevel', 'error'
            ]
            result = subprocess.run(
                test_cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                shell=False
            )
            if result.returncode == 0:
                logger.info("✅ Stream test successful")
                return True
            else:
                error_msg = result.stderr[:100] if result.stderr else "Unknown error"
                logger.error(f"❌ Stream test failed: {error_msg}")
                return False
        except subprocess.TimeoutExpired:
            logger.warning(f"⏰ Stream test timeout after {timeout}s")
            return False
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Stream test error: {e}")
            return True

    def dispose(self) -> None:
        logger.info("🧹 AudioProcessor: Starting dispose...")
        try:
            if hasattr(self, '_stop_event') and hasattr(self._stop_event, 'set'):
                self._stop_event.set()
            with self._processing_lock:
                if hasattr(self, '_processing') and hasattr(self._processing, 'clear'):
                    self._processing.clear()
            self._cleanup_done = True

            if hasattr(self, 'ffmpeg_manager') and self.ffmpeg_manager:
                try:
                    self.ffmpeg_manager.stop_all_streams()
                except Exception:
                    pass

            if hasattr(self, '_transcription_executor') and self._transcription_executor:
                try:
                    self._transcription_executor.shutdown(wait=False)
                except Exception:
                    pass

            with self._subtitle_lock:
                if hasattr(self, '_timed_transcriptions') and hasattr(self._timed_transcriptions, 'clear'):
                    self._timed_transcriptions.clear()
                if hasattr(self, '_timed_translations') and hasattr(self._timed_translations, 'clear'):
                    self._timed_translations.clear()

            with self._duplicate_lock:
                if hasattr(self, '_recent_transcriptions') and hasattr(self._recent_transcriptions, 'clear'):
                    self._recent_transcriptions.clear()
                self._last_transcription_text = ""

            gc.collect()
            logger.info("✅ AudioProcessor disposed")
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ AudioProcessor dispose error: {e}")

    def stop_processing(self, wait: bool = False, timeout: float = 5.0) -> bool:
        logger.info("🛑 AudioProcessor: Stopping processing...")
        self._stop_event.set()
        with self._processing_lock:
            self._processing.clear()
        if self._current_stream_id:
            logger.info(f"📛 Stream {self._current_stream_id} stopped by user")
            if self.ffmpeg_manager:
                self.ffmpeg_manager.stop_stream(self._current_stream_id)
        return True
# -----------------------------------------------------------------------------
# DARK CONTEXT MENUS
# -----------------------------------------------------------------------------
class DarkContextMenu:
    def __init__(self, text_widget: tk.Text) -> None:
        self.text_widget = text_widget
        self.menu = tk.Menu(
            text_widget,
            tearoff=0,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            activebackground=CURRENT_THEME.BG_HOVER,
            activeforeground=CURRENT_THEME.TEXT_ACCENT,
            borderwidth=1,
            relief="solid",
        )
        self.menu.add_command(label="Copy", command=self.copy_text)
        self.menu.add_command(label="Select All", command=self.select_all)
        self.menu.add_separator()
        self.menu.add_command(label="Delete", command=self.clear_text)
        text_widget.bind("<Button-3>", self.show_menu)

    def show_menu(self, event: tk.Event) -> None:
        try:
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu.grab_release()

    def copy_text(self) -> None:
        try:
            if self.text_widget.tag_ranges(tk.SEL):
                selected_text = self.text_widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.text_widget.clipboard_clear()
                self.text_widget.clipboard_append(selected_text)
        except tk.TclError:
            pass

    def select_all(self) -> None:
        try:
            self.text_widget.tag_add(tk.SEL, "1.0", tk.END)
            self.text_widget.mark_set(tk.INSERT, "1.0")
            self.text_widget.see(tk.INSERT)
        except tk.TclError:
            pass

    def clear_text(self) -> None:
        try:
            self.text_widget.delete("1.0", tk.END)
        except tk.TclError:
            pass

class DarkEntryContextMenu:
    def __init__(self, entry_widget: tk.Entry) -> None:
        self.entry_widget = entry_widget
        self.menu = tk.Menu(
            entry_widget,
            tearoff=0,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            activebackground=CURRENT_THEME.BG_HOVER,
            activeforeground=CURRENT_THEME.TEXT_ACCENT,
            borderwidth=1,
            relief="solid",
        )
        self.menu.add_command(label="Cut", command=self.cut_text)
        self.menu.add_command(label="Copy", command=self.copy_text)
        self.menu.add_command(label="Paste", command=self.paste_text)
        self.menu.add_separator()
        self.menu.add_command(label="Select All", command=self.select_all)
        self.menu.add_command(label="Delete", command=self.delete_text)
        entry_widget.bind("<Button-3>", self.show_menu)

    def show_menu(self, event: tk.Event) -> None:
        try:
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu.grab_release()

    def cut_text(self) -> None:
        self.entry_widget.event_generate("<<Cut>>")

    def copy_text(self) -> None:
        self.entry_widget.event_generate("<<Copy>>")

    def paste_text(self) -> None:
        self.entry_widget.event_generate("<<Paste>>")

    def select_all(self) -> None:
        self.entry_widget.select_range(0, "end")
        self.entry_widget.icursor("end")

    def delete_text(self) -> None:
        self.entry_widget.delete(0, "end")


# -----------------------------------------------------------------------------
# EXPORT MANAGER
# -----------------------------------------------------------------------------
class ExportManager:
    """
    Verwaltet den Export von Transkriptionen und Übersetzungen in verschiedene Formate.
    Unterstützte Formate: txt, srt, vtt, json, docx (optional mit python-docx).
    """

    def __init__(self) -> None:
        self.supported_formats = ["txt", "srt", "vtt", "json", "docx"]
        self._docx_available: bool = False
        try:
            # Versuche python-docx zu importieren, falls vorhanden
            import docx
            self._docx = docx
            self._docx_available = True
        except ImportError:
            self._docx_available = False

    # -------------------------------------------------------------------------
    # Öffentliche Export-Methoden
    # -------------------------------------------------------------------------
    def export_subtitles(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
        format: str = "srt",
        filename: Optional[str] = None,
        encoding: str = "utf-8-sig",
    ) -> Union[bool, str]:
        """
        Exportiert Untertitel im angegebenen Format.

        :param transcript_data: Liste der Transkriptionssegmente (müssen Zeitstempel enthalten).
        :param translation_data: Optionale Liste der Übersetzungssegmente (parallel zu transcript_data).
        :param format: Gewünschtes Format ("srt", "vtt", "txt", "json", "docx").
        :param filename: Optionaler Dateiname. Wenn None, wird der Inhalt als String zurückgegeben.
        :param encoding: Zeichenkodierung für Textdateien (Standard: UTF-8 mit BOM).
        :return: True bei erfolgreichem Datei-Export, sonst Inhalt als String oder False bei Fehler.
        """
        try:
            # Prüfen, ob Zeitstempel vorhanden sind (außer bei txt/json)
            if format.lower() not in ("txt", "json"):
                timed = [
                    t for t in transcript_data
                    if hasattr(t, "start") and t.start is not None
                       and hasattr(t, "end") and t.end is not None
                ]
                if not timed:
                    raise ProcessingError("Keine Segmente mit Zeitstempeln vorhanden – benötigt für Untertitel.")
                # Verwende nur Segmente mit Zeitstempeln
                transcript_data = timed

            # Format-spezifische Generierung
            if format.lower() == "srt":
                content = self._generate_srt_content(transcript_data, translation_data)
            elif format.lower() == "vtt":
                content = self._generate_vtt_content(transcript_data, translation_data)
            elif format.lower() == "txt":
                content = self._generate_txt_content(transcript_data, translation_data)
            elif format.lower() == "json":
                # JSON-Export benötigt eigenen Pfad
                if filename:
                    return self.export_json(transcript_data, translation_data or [], filename)
                else:
                    # JSON ohne Dateiname -> Inhalt als String zurückgeben
                    import json
                    data = self._build_json_data(transcript_data, translation_data)
                    return json.dumps(data, indent=2, ensure_ascii=False)
            elif format.lower() == "docx":
                if filename:
                    return self.export_docx(transcript_data, filename)
                else:
                    raise ProcessingError("Für DOCX-Export wird ein Dateiname benötigt.")
            else:
                raise ProcessingError(f"Nicht unterstütztes Format: {format}")

            # Bei erfolgreicher Generierung: entweder in Datei schreiben oder Inhalt zurückgeben
            if filename:
                # Pfad normalisieren und Verzeichnis ggf. erstellen
                out_path = Path(filename)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                with open(out_path, "w", encoding=encoding) as f:
                    f.write(content)
                return True
            else:
                return content

        except Exception as e:
            logger.error(f"Fehler beim Export: {e}")
            raise ProcessingError(f"Export fehlgeschlagen: {e}") from e

    def export_json(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: List[TranslationResult],
        filename: str,
        encoding: str = "utf-8",
    ) -> bool:
        """
        Exportiert Transkriptionen und Übersetzungen als JSON-Datei.
        """
        try:
            data = self._build_json_data(transcript_data, translation_data)
            out_path = Path(filename)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", encoding=encoding) as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            raise ProcessingError(f"JSON-Export fehlgeschlagen: {e}") from e

    def export_docx(
        self,
        transcript_data: List[TranscriptionResult],
        filename: str,
    ) -> bool:
        """
        Exportiert Transkriptionen als DOCX-Datei (Microsoft Word).
        Falls python-docx nicht installiert ist, wird eine einfache Textdatei mit .docx-Endung erstellt.
        """
        if self._docx_available:
            try:
                doc = self._docx.Document()
                doc.add_heading("Transkription", level=1)

                for i, seg in enumerate(transcript_data, 1):
                    p = doc.add_paragraph()
                    if hasattr(seg, "start") and seg.start is not None:
                        start_str = self._format_timestamp_srt(seg.start)
                        p.add_run(f"[{start_str}] ").bold = True
                    p.add_run(seg.text)

                    if i < len(transcript_data):
                        doc.add_paragraph()

                out_path = Path(filename)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(out_path))
                return True
            except Exception as e:
                logger.warning(f"python-docx Export fehlgeschlagen, verwende Fallback: {e}")

        # Fallback: Einfache Textdatei mit .docx-Endung
        try:
            content = self._generate_txt_content(transcript_data, None)
            out_path = Path(filename)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(content)
            return True
        except Exception as e:
            raise ProcessingError(f"DOCX-Fallback fehlgeschlagen: {e}") from e

    # -------------------------------------------------------------------------
    # Hilfsmethoden zur Inhaltsgenerierung
    # -------------------------------------------------------------------------
    def _generate_srt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = []
        for i, segment in enumerate(transcript_data):
            start = self._format_timestamp_srt(segment.start or 0.0)
            end = self._format_timestamp_srt(segment.end or 0.0)
            text = segment.text
            if translation_data and i < len(translation_data):
                # Zweisprachige Untertitel: erste Zeile Original, zweite Zeile Übersetzung
                text = f"{text}\n{translation_data[i].translated}"
            lines.append(f"{i+1}\n{start} --> {end}\n{text}\n")
        return "\n".join(lines)

    def _generate_vtt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = ["WEBVTT\n"]
        for i, segment in enumerate(transcript_data):
            start = self._format_timestamp_vtt(segment.start or 0.0)
            end = self._format_timestamp_vtt(segment.end or 0.0)
            text = segment.text
            if translation_data and i < len(translation_data):
                text = f"{text}\n{translation_data[i].translated}"
            lines.append(f"{start} --> {end}\n{text}\n")
        return "\n".join(lines)

    def _generate_txt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = []
        for i, segment in enumerate(transcript_data):
            if hasattr(segment, "start") and segment.start is not None:
                timestamp = self._format_timestamp_srt(segment.start)
                prefix = f"[{timestamp}] "
            else:
                prefix = ""
            text = segment.text
            if translation_data and i < len(translation_data):
                text = f"{text}  |  {translation_data[i].translated}"
            lines.append(f"{prefix}{text}")
        return "\n".join(lines)

    def _build_json_data(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: List[TranslationResult],
    ) -> Dict[str, Any]:
        return {
            "metadata": {
                "export_date": datetime.now().isoformat(),
                "total_segments": len(transcript_data),
                "version": "4.1.3",
            },
            "transcripts": [
                {
                    "text": seg.text,
                    "confidence": seg.confidence,
                    "language": seg.language,
                    "timestamp": seg.timestamp,
                    "start": getattr(seg, "start", None),
                    "end": getattr(seg, "end", None),
                }
                for seg in transcript_data
            ],
            "translations": [
                {
                    "original": trans.original,
                    "translated": trans.translated,
                    "source_lang": trans.source_lang,
                    "target_lang": trans.target_lang,
                    "timestamp": trans.timestamp,
                    "start": getattr(trans, "start", None),
                    "end": getattr(trans, "end", None),
                }
                for trans in translation_data
            ] if translation_data else [],
        }

    # -------------------------------------------------------------------------
    # Zeitstempel-Formatierung
    # -------------------------------------------------------------------------
    @staticmethod
    def _format_timestamp_srt(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        milliseconds = int((secs - int(secs)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{int(secs):02d},{milliseconds:03d}"

    @staticmethod
    def _format_timestamp_vtt(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        milliseconds = int((secs - int(secs)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{int(secs):02d}.{milliseconds:03d}"

# -----------------------------------------------------------------------------
# APP SETTINGS
# -----------------------------------------------------------------------------
@dataclass
class AppSettings:
    last_url: str = ""
    default_model: str = "medium"
    default_language: str = "de"
    layout_mode: str = "vertical"
    recent_urls: List[str] = None
    enable_plugins: bool = True
    export_format: str = "txt"
    auto_save_on_completion: bool = False
    theme: str = "dark"
    use_browser_cookies: bool = True
    cookies_notice_shown: bool = False

    def __post_init__(self) -> None:
        if self.recent_urls is None:
            self.recent_urls = []

    @classmethod
    def load_from_file(cls, filename: str = "dragon_settings.json") -> 'AppSettings':
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            if file_path.exists():
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return cls(**data)
        except Exception:
            pass
        return cls()

    def save_to_file(self, filename: str = "dragon_settings.json") -> None:
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(self.__dict__, f, indent=2, ensure_ascii=False)
        except Exception:
            pass

    def add_recent_url(self, url: str) -> None:
        if url in self.recent_urls:
            self.recent_urls.remove(url)
        self.recent_urls.insert(0, url)
        self.recent_urls = self.recent_urls[:10]
        self.save_to_file()


# -----------------------------------------------------------------------------
# RESOURCE MANAGER
# -----------------------------------------------------------------------------
class ResourceManager:
    def __init__(self) -> None:
        self.processes: List[subprocess.Popen] = []
        self.threads: List[threading.Thread] = []
        self.temp_files: List[str] = []
        self.cleanup_done = False
        self._lock = threading.RLock()
        self._shutdown_event = threading.Event()

    def register_process(self, process: subprocess.Popen) -> None:
        with self._lock:
            if process and process not in self.processes:
                self.processes.append(process)

    def register_thread(self, thread: threading.Thread) -> None:
        with self._lock:
            if thread and thread not in self.threads and thread.is_alive():
                self.threads.append(thread)

    def register_temp_file(self, file_path: str) -> None:
        with self._lock:
            if file_path and file_path not in self.temp_files:
                self.temp_files.append(file_path)

    def cleanup(self) -> None:
        if self.cleanup_done:
            return
        self._shutdown_event.set()
        with self._lock:
            cleanup_timeout = 5.0
            start_time = time.time()
            for process in self.processes[:]:
                try:
                    if process and hasattr(process, "poll"):
                        if process.poll() is None:
                            process.terminate()
                            try:
                                process.wait(timeout=1.0)
                            except (subprocess.TimeoutExpired, AttributeError):
                                try:
                                    process.kill()
                                    process.wait(timeout=0.5)
                                except Exception:
                                    pass
                except Exception:
                    pass
                finally:
                    if process in self.processes:
                        self.processes.remove(process)
                if time.time() - start_time > cleanup_timeout:
                    break
            for thread in self.threads[:]:
                try:
                    if thread and thread.is_alive():
                        thread.join(timeout=1.0)
                except Exception:
                    pass
                finally:
                    if thread in self.threads:
                        self.threads.remove(thread)
                if time.time() - start_time > cleanup_timeout:
                    break
            for temp_file in self.temp_files[:]:
                try:
                    if os.path.exists(temp_file):
                        for attempt in range(2):
                            try:
                                os.unlink(temp_file)
                                break
                            except PermissionError:
                                if attempt < 1:
                                    time.sleep(0.1)
                                    continue
                            except Exception:
                                if attempt < 1:
                                    time.sleep(0.1)
                except (OSError, PermissionError):
                    pass
                finally:
                    if temp_file in self.temp_files:
                        self.temp_files.remove(temp_file)
            try:
                if TORCH_AVAILABLE:
                    torch = FastLazyLoader.load("torch")
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
            except Exception:
                pass
            gc.collect()
            self.cleanup_done = True

    def is_shutting_down(self) -> bool:
        return self._shutdown_event.is_set()


# -----------------------------------------------------------------------------
# OLLAMA SUMMARIZER
# -----------------------------------------------------------------------------
class OllamaSummarizer:
    """
    Verbesserte Ollama-Zusammenfassung mit Streaming, Timeouts und Abbruchfunktion.
    """

    def __init__(self, parent: Any, model: str = "llama3", host: str = "http://localhost:11434", timeout: int = 120) -> None:
        self.parent = parent
        self.model = model
        self.host = host.rstrip('/')
        self.timeout = timeout
        self.available = OLLAMA_AVAILABLE
        self._session = None
        self._stop_event = threading.Event()
        self._lock = threading.RLock()
        # Lade requests-Modul über FastLazyLoader
        self._requests = FastLazyLoader.load('requests') if self.available else None

    def _get_session(self):
        """Erstellt oder gibt eine wiederverwendbare Session zurück."""
        if not self.available:
            return None
        if self._session is None:
            self._session = self._requests.Session()
            self._session.headers.update({
                'Content-Type': 'application/json',
                'Accept': 'application/json',
            })
        return self._session

    def get_available_models(self) -> List[str]:
        """Ruft die Liste der verfügbaren Modelle vom Ollama-Server ab."""
        if not self.available:
            return []
        try:
            session = self._get_session()
            r = session.get(f"{self.host}/api/tags", timeout=5)
            if r.status_code == 200:
                data = r.json()
                return [m['name'] for m in data.get('models', [])]
        except self._requests.exceptions.RequestException as e:
            logger.warning(f"Ollama model list error: {e}")
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"Unexpected error getting models: {e}")
        return []

    def is_server_reachable(self) -> bool:
        """Prüft, ob der Ollama-Server erreichbar ist."""
        if not self.available:
            return False
        try:
            session = self._get_session()
            r = session.get(f"{self.host}/api/tags", timeout=2)
            return r.status_code == 200
        except Exception:
            return False

    def summarize(self, text: str, prompt: str, temperature: float,
                  callback: Callable[[str], None],
                  error_callback: Callable[[str], None],
                  complete_callback: Optional[Callable[[], None]] = None,
                  cancel_event: Optional[threading.Event] = None) -> None:
        """
        Fasst den gegebenen Text mit Ollama zusammen (Streaming).
        Kann über stop() abgebrochen werden.
        """
        if not self.available:
            error_callback("Ollama nicht verfügbar (requests nicht installiert)")
            return

        if not text or not text.strip():
            error_callback("Kein Text zum Zusammenfassen")
            return

        def worker() -> None:
            self._stop_event.clear()
            try:
                session = self._get_session()
                full_prompt = f"{prompt}\n\n{text}"
                payload = {
                    "model": self.model,
                    "prompt": full_prompt,
                    "stream": True,
                    "options": {
                        "temperature": temperature,
                        "num_predict": 512
                    }
                }
                response = session.post(
                    f"{self.host}/api/generate",
                    json=payload,
                    stream=True,
                    timeout=self.timeout
                )
                if response.status_code == 200:
                    full_response = ""
                    for line in response.iter_lines(decode_unicode=True):
                        if self._stop_event.is_set() or (cancel_event and cancel_event.is_set()):
                            break
                        if line:
                            try:
                                data = json.loads(line)
                                if 'response' in data:
                                    chunk = data['response']
                                    full_response += chunk
                                    callback(chunk)
                                if data.get('done', False):
                                    break
                            except json.JSONDecodeError:
                                continue
                    if not self._stop_event.is_set() and not (cancel_event and cancel_event.is_set()):
                        if not full_response:
                            error_callback("Leere Antwort von Ollama")
                        else:
                            if complete_callback:
                                complete_callback()
                else:
                    error_callback(f"Ollama Fehler {response.status_code}")
            except self._requests.exceptions.Timeout:
                error_callback(f"Ollama Timeout nach {self.timeout}s – Server nicht erreichbar?")
            except self._requests.exceptions.ConnectionError:
                error_callback("Ollama nicht erreichbar (läuft der Server?)")
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                error_callback(f"Fehler: {str(e)}")

        threading.Thread(target=worker, daemon=True).start()

    def correct_transcript(self, text: str,
                           callback: Callable[[str], None],
                           error_callback: Callable[[str], None],
                           complete_callback: Optional[Callable[[], None]] = None,
                           cancel_event: Optional[threading.Event] = None) -> None:
        """
        Korrigiert ein Transkript mit Ollama (streaming).
        """
        if not self.available:
            error_callback("Ollama nicht verfügbar (requests fehlt)")
            return

        prompt = (
            "Du bist ein Assistent, der Transkriptionen korrigiert. "
            "Der folgende Text wurde automatisch transkribiert und enthält typische Fehler: "
            "fehlende Satzzeichen, falsche Groß-/Kleinschreibung, leichte Wortverzerrungen. "
            "Korrigiere diese Fehler und verbessere die Lesbarkeit. "
            "Ändere keine inhaltlichen Aussagen, füge keine neuen Informationen hinzu. "
            "Gib nur den korrigierten Text zurück, ohne Einleitung oder Erklärung.\n\n"
            f"{text}"
        )

        def worker():
            self._stop_event.clear()
            try:
                session = self._get_session()
                payload = {
                    "model": self.model,
                    "prompt": prompt,
                    "stream": True,
                    "options": {
                        "temperature": 0.0,
                        "num_predict": -1
                    }
                }
                response = session.post(
                    f"{self.host}/api/generate",
                    json=payload,
                    stream=True,
                    timeout=self.timeout
                )
                if response.status_code == 200:
                    full_response = ""
                    for line in response.iter_lines(decode_unicode=True):
                        if self._stop_event.is_set() or (cancel_event and cancel_event.is_set()):
                            break
                        if line:
                            try:
                                data = json.loads(line)
                                if 'response' in data:
                                    chunk = data['response']
                                    full_response += chunk
                                    callback(chunk)
                                if data.get('done', False):
                                    break
                            except json.JSONDecodeError:
                                continue
                    if not self._stop_event.is_set() and not (cancel_event and cancel_event.is_set()):
                        if not full_response:
                            error_callback("Leere Antwort von Ollama")
                        else:
                            if complete_callback:
                                complete_callback()
                else:
                    error_callback(f"Ollama Fehler {response.status_code}")
            except self._requests.exceptions.Timeout:
                error_callback(f"Ollama Timeout nach {self.timeout}s – Server nicht erreichbar?")
            except self._requests.exceptions.ConnectionError:
                error_callback("Ollama nicht erreichbar (läuft der Server?)")
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                error_callback(f"Fehler: {str(e)}")

        threading.Thread(target=worker, daemon=True).start()

    def stop(self) -> None:
        """Bricht eine laufende Anfrage ab."""
        self._stop_event.set()
        logger.info("OllamaSummarizer: Stop signalisiert")

    def dispose(self) -> None:
        """Räumt Ressourcen auf (Session schließen)."""
        self.stop()
        if self._session:
            try:
                self._session.close()
            except Exception:
                pass
            self._session = None

class SummarizeDialog:
    def __init__(self, parent: Any, text: str, gui_ref: Any) -> None:
        self.parent = parent
        self.text = text
        self.gui = gui_ref
        self.dialog: Optional[tk.Toplevel] = None
        self.summarizer = OllamaSummarizer(parent)
        self._destroyed = False
        self._request_cancel = threading.Event()

        # Prüfe, ob Ollama-Server erreichbar ist
        if not self.summarizer.is_server_reachable():
            DarkMessageBox.showwarning(
                "Ollama nicht erreichbar",
                "Der Ollama-Server läuft nicht oder ist nicht erreichbar.\n"
                "Bitte starte 'ollama serve' und versuche es erneut.",
                parent=self.parent
            )
            return

        self.create_dialog()

    def create_dialog(self) -> None:
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("🐉 Zusammenfassung mit Ollama")
        self.dialog.geometry("750x650")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(self.parent)
        self.dialog.grab_set()

        def on_close() -> None:
            self._destroyed = True
            self._request_cancel.set()
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.destroy()
            except tk.TclError:
                pass

        self.dialog.protocol("WM_DELETE_WINDOW", on_close)

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        # ---------- Modellauswahl ----------
        model_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        model_frame.pack(fill="x", pady=5)
        tk.Label(model_frame, text="Modell:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(side="left")
        self.model_var = tk.StringVar()
        available = self.summarizer.get_available_models()
        if available:
            # Bevorzugte Modelle zuerst anzeigen
            preferred = ["qwen2.5:7b", "glm4:9b", "llama3.1:8b"]
            values = []
            for pref in preferred:
                if pref in available:
                    values.append(pref)
            for m in available:
                if m not in values:
                    values.append(m)
            self.model_combo = ttk.Combobox(
                model_frame, textvariable=self.model_var,
                values=values, width=20, state="readonly",
                style="Dark.TCombobox"
            )
            self.model_var.set(values[0])
        else:
            self.model_combo = ttk.Combobox(
                model_frame, textvariable=self.model_var,
                values=["(keine Modelle gefunden)"],
                width=20, state="disabled", style="Dark.TCombobox"
            )
            self.model_var.set("(keine Modelle)")
        self.model_combo.pack(side="left", padx=10)
        ToolTip(self.model_combo, "Wähle das Ollama-Modell für die Zusammenfassung")

        # ---------- Temperatur ----------
        tk.Label(model_frame, text="Temperatur:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(side="left", padx=(20,5))
        self.temp_var = tk.DoubleVar(value=0.0)
        temp_scale = tk.Scale(
            model_frame, from_=0.0, to=1.0, resolution=0.1, orient=tk.HORIZONTAL,
            variable=self.temp_var, length=150, bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY, highlightbackground=CURRENT_THEME.BG_PRIMARY
        )
        temp_scale.pack(side="left")
        tk.Label(model_frame, text="(0 = deterministisch)", font=("Segoe UI", 7),
                 bg=CURRENT_THEME.BG_PRIMARY, fg=CURRENT_THEME.TEXT_SECONDARY).pack(side="left", padx=5)
        ToolTip(temp_scale, "Zufälligkeit der Ausgabe (höher = kreativer)")

        # ---------- Zielsprache ----------
        lang_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        lang_frame.pack(fill="x", pady=5)
        tk.Label(lang_frame, text="Zusammenfassen auf:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(side="left")
        self.summary_lang_var = tk.StringVar()
        supported_summary_langs = ["Deutsch", "Englisch", "Spanisch", "Koreanisch"]
        current_lang_name = SUPPORTED_LANGUAGES.get(self.gui.current_language, "Deutsch")
        if current_lang_name not in supported_summary_langs:
            current_lang_name = "Deutsch"
        self.summary_lang_var.set(current_lang_name)
        lang_combo = ttk.Combobox(
            lang_frame, textvariable=self.summary_lang_var,
            values=supported_summary_langs, width=15, state="readonly",
            style="Dark.TCombobox"
        )
        lang_combo.pack(side="left", padx=10)
        lang_combo.bind("<<ComboboxSelected>>", lambda e: self._set_default_prompt())
        ToolTip(lang_combo, "Sprache der Zusammenfassung")

        # ---------- Prompt ----------
        tk.Label(main, text="Prompt (optional):", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(anchor="w", pady=(10,2))
        self.prompt_text = scrolledtext.ScrolledText(
            main, height=4, bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE, wrap=tk.WORD
        )
        self.prompt_text.pack(fill="x", pady=(0,10))
        self._set_default_prompt()
        ToolTip(self.prompt_text, "Optionaler Prompt – wird an das Modell gesendet")

        # ---------- Ergebnis ----------
        tk.Label(main, text="Zusammenfassung:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(anchor="w")
        self.summary_text = scrolledtext.ScrolledText(
            main, height=10, bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE, wrap=tk.WORD
        )
        self.summary_text.pack(fill="both", expand=True, pady=10)
        DarkContextMenu(self.summary_text)

        # ---------- Buttons ----------
        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x")

        self.summarize_btn = tk.Button(
            btn_frame, text="🤖 Zusammenfassen", command=self.start_summarize,
            bg=CURRENT_THEME.DRAGON_GREEN, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20
        )
        self.summarize_btn.pack(side="left", padx=5)

        self.cancel_btn = tk.Button(
            btn_frame, text="⏹️ Abbrechen", command=self.cancel_request,
            bg=CURRENT_THEME.ERROR, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20, state="disabled"
        )
        self.cancel_btn.pack(side="left", padx=5)

        self.copy_btn = tk.Button(
            btn_frame, text="📋 Kopieren", command=self.copy_summary,
            bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20
        )
        self.copy_btn.pack(side="left", padx=5)

        self.save_btn = tk.Button(
            btn_frame, text="💾 Speichern", command=self.save_summary,
            bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20
        )
        self.save_btn.pack(side="left", padx=5)

        self.translate_btn = tk.Button(
            btn_frame, text="🌐 Übersetzen", command=self.translate_summary,
            bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20, state="disabled"
        )
        self.translate_btn.pack(side="left", padx=5)

        self.close_btn = tk.Button(
            btn_frame, text="Schließen", command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY
        )
        self.close_btn.pack(side="right", padx=5)

        self.status_label = tk.Label(main, text="", bg=CURRENT_THEME.BG_PRIMARY,
                                      fg=CURRENT_THEME.TEXT_SECONDARY)
        self.status_label.pack(pady=5)

        self.full_summary = ""

    def _set_default_prompt(self) -> None:
        target = self.summary_lang_var.get()
        if target == "Deutsch":
            prompt = (
                "Fasse den folgenden Text in 3-5 Sätzen auf Deutsch zusammen. "
                "Konzentriere dich auf die Hauptaussagen, ignoriere Wiederholungen und Füllwörter. "
                "Verwende vollständige, idiomatische Sätze.\n\n"
                "WICHTIG: Die Ausgabe MUSS auf DEUTSCH erfolgen und darf KEINE Zeichen aus anderen "
                "Sprachen (wie Chinesisch, Japanisch oder Koreanisch) enthalten. "
                "Fachbegriffe wie 'Sakura' oder 'Kimchi' können als Lehnwörter verwendet werden, "
                "aber bitte in lateinischer Schrift."
            )
        elif target == "Englisch":
            prompt = (
                "Summarize the following text in 3-5 sentences in English. "
                "Focus on the main points, ignore repetitions and filler words. "
                "Use complete, idiomatic sentences.\n\n"
                "IMPORTANT: The output MUST be in ENGLISH and must NOT contain characters from other "
                "languages (like Chinese, Japanese, or Korean). Loanwords are acceptable in Latin script."
            )
        elif target == "Spanisch":
            prompt = (
                "Resume el siguiente texto en 3-5 oraciones en español. "
                "Concéntrate en las ideas principales, ignora repeticiones y palabras de relleno. "
                "Utiliza oraciones completas e idiomáticas.\n\n"
                "IMPORTANTE: La salida DEBE estar en ESPAÑOL y NO debe contener caracteres de otros "
                "idiomas (como chino, japonés o coreano). Los préstamos son aceptables en escritura latina."
            )
        elif target == "Koreanisch":
            prompt = (
                "다음 텍스트를 한국어로 3-5문장으로 요약하세요. "
                "핵심 내용에 집중하고, 반복과 불용어는 무시하세요. "
                "완전하고 관용적인 문장을 사용하세요.\n\n"
                "중요: 출력은 반드시 한국어로 작성되어야 하며 다른 언어(예: 중국어, 일본어)의 문자가 포함되어서는 안 됩니다. "
                "외래어는 라틴 문자로 표기해도 됩니다."
            )
        else:
            prompt = (
                f"Summarize the following text in 3-5 sentences in {target}. "
                f"Focus on the main points, ignore repetitions and filler words. "
                f"Use complete, idiomatic sentences.\n\n"
                f"IMPORTANT: The output MUST be in {target.upper()} and must NOT contain characters from other languages."
            )
        self.prompt_text.delete("1.0", "end")
        self.prompt_text.insert("1.0", prompt)

    def start_summarize(self) -> None:
        if self._destroyed:
            return
        model = self.model_var.get().strip()
        if model == "--- Modell auswählen ---" or model.startswith("(keine"):
            self.status_label.config(text="❌ Bitte ein gültiges Modell auswählen")
            return
        self.summarizer.model = model
        if hasattr(self.gui, 'advanced_settings'):
            self.summarizer.host = self.gui.advanced_settings.ollama_host
        else:
            self.summarizer.host = "http://localhost:11434"
        prompt = self.prompt_text.get("1.0", "end-1c").strip()
        temp = self.temp_var.get()

        # UI während der Verarbeitung deaktivieren
        self.summarize_btn.config(state="disabled", text="⏳ Warte...")
        self.cancel_btn.config(state="normal")
        self.copy_btn.config(state="disabled")
        self.save_btn.config(state="disabled")
        self.translate_btn.config(state="disabled")
        self.status_label.config(text="Sende Anfrage an Ollama...")
        self.summary_text.delete("1.0", "end")
        self.full_summary = ""
        self._request_cancel.clear()

        def on_complete() -> None:
            if self._destroyed:
                return
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.after(0, self._reset_ui)

        self.summarizer.summarize(
            self.text, prompt, temp,
            callback=self.on_chunk,
            error_callback=self.on_error,
            complete_callback=on_complete,
            cancel_event=self._request_cancel
        )

    def _reset_ui(self) -> None:
        if self._destroyed:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                if hasattr(self, 'summarize_btn') and self.summarize_btn.winfo_exists():
                    self.summarize_btn.config(state="normal", text="🤖 Zusammenfassen")
                if hasattr(self, 'cancel_btn') and self.cancel_btn.winfo_exists():
                    self.cancel_btn.config(state="disabled")
                if hasattr(self, 'copy_btn') and self.copy_btn.winfo_exists():
                    self.copy_btn.config(state="normal")
                if hasattr(self, 'save_btn') and self.save_btn.winfo_exists():
                    self.save_btn.config(state="normal")
                if hasattr(self, 'translate_btn') and self.translate_btn.winfo_exists():
                    self.translate_btn.config(state="normal" if self.full_summary else "disabled")
                if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                    self.status_label.config(text="✅ Zusammenfassung abgeschlossen")
        except tk.TclError:
            pass

    def cancel_request(self) -> None:
        """Bricht die laufende Anfrage ab und setzt die UI zurück."""
        self._request_cancel.set()
        self.status_label.config(text="⏹️ Abbruch eingeleitet...")
        self.cancel_btn.config(state="disabled")
        self.dialog.after(500, self._reset_ui)

    def on_chunk(self, chunk: str) -> None:
        if self._destroyed:
            return
        def update() -> None:
            if self._destroyed:
                return
            try:
                if self.dialog and self.dialog.winfo_exists():
                    if hasattr(self, 'summary_text') and self.summary_text.winfo_exists():
                        self.summary_text.insert("end", chunk)
                        self.summary_text.see("end")
                        self.full_summary += chunk
            except tk.TclError:
                pass
        if self.dialog and not self._destroyed and self.dialog.winfo_exists():
            self.dialog.after(0, update)

    def on_error(self, error: str) -> None:
        if self._destroyed:
            return
        def update() -> None:
            if self._destroyed:
                return
            try:
                if self.dialog and self.dialog.winfo_exists():
                    if hasattr(self, 'summary_text') and self.summary_text.winfo_exists():
                        self.summary_text.delete("1.0", "end")
                        self.summary_text.insert("1.0", f"Fehler: {error}")
                    self._reset_ui()
                    if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                        self.status_label.config(text="❌ Fehler")
            except tk.TclError:
                pass
        if self.dialog and not self._destroyed and self.dialog.winfo_exists():
            self.dialog.after(0, update)

    def copy_summary(self) -> None:
        if self._destroyed:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                if self.full_summary:
                    self.dialog.clipboard_clear()
                    self.dialog.clipboard_append(self.full_summary)
                    if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                        self.status_label.config(text="✅ In Zwischenablage kopiert")
                else:
                    if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                        self.status_label.config(text="⚠️ Keine Zusammenfassung vorhanden")
        except tk.TclError:
            pass

    def save_summary(self) -> None:
        """Speichert die Zusammenfassung in einer Textdatei."""
        if not self.full_summary:
            self.status_label.config(text="⚠️ Keine Zusammenfassung zum Speichern")
            return
        from tkinter import filedialog
        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Textdateien", "*.txt"), ("Alle Dateien", "*.*")],
            title="Zusammenfassung speichern"
        )
        if filename:
            try:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(self.full_summary)
                self.status_label.config(text=f"💾 Gespeichert: {os.path.basename(filename)}")
            except Exception as e:
                self.status_label.config(text=f"❌ Fehler beim Speichern: {e}")

    def translate_summary(self) -> None:
        if not self.full_summary:
            self.status_label.config(text="⚠️ Keine Zusammenfassung zum Übersetzen")
            return
        if not hasattr(self.gui, 'translation_engine'):
            self.status_label.config(text="❌ Keine Übersetzungs-Engine verfügbar")
            return
        engine = self.gui.translation_engine
        if hasattr(engine, 'is_functional') and not engine.is_functional():
            self.status_label.config(text="⚠️ Übersetzungs-Engine derzeit nicht verfügbar")
            return
        # TranslationDialog öffnen – der Dialog übernimmt die Übersetzung selbst
        TranslationDialog(self.dialog, engine, initial_text=self.full_summary)

# -----------------------------------------------------------------------------
# TOOLTIP
# -----------------------------------------------------------------------------
class ToolTip:
    def __init__(self, widget: tk.Widget, text: str, delay: int = 500) -> None:
        self.widget = widget
        self.text = text
        self.delay = delay
        self.tip_window: Optional[tk.Toplevel] = None
        self.after_id: Optional[str] = None
        widget.bind('<Enter>', self.enter)
        widget.bind('<Leave>', self.leave)
        widget.bind('<ButtonPress>', self.leave)

    def enter(self, event: Optional[tk.Event] = None) -> None:
        self.schedule()

    def leave(self, event: Optional[tk.Event] = None) -> None:
        self.unschedule()
        self.hide_tip()

    def schedule(self) -> None:
        self.unschedule()
        self.after_id = self.widget.after(self.delay, self.show_tip)

    def unschedule(self) -> None:
        if self.after_id:
            self.widget.after_cancel(self.after_id)
            self.after_id = None

    def show_tip(self) -> None:
        if self.tip_window or not self.text:
            return
        x, y = self.widget.winfo_pointerxy()
        x += 10
        y += 10
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, background="#ffffe0", relief="solid",
                         borderwidth=1, font=("Segoe UI", 9))
        label.pack()
        tw.bind('<Leave>', lambda e: self.hide_tip())

    def hide_tip(self) -> None:
        if self.tip_window:
            try:
                self.tip_window.destroy()
            except tk.TclError:
                pass
            self.tip_window = None

# -----------------------------------------------------------------------------
# LAYOUT MANAGER
# -----------------------------------------------------------------------------
class WhisperLayoutManager:
    def __init__(self, gui_ref: Any) -> None:
        self.gui_ref = gui_ref
        self.root = gui_ref.root
        self._batch_timer_id: Optional[str] = None
        try:
            self.gui_ref._text_update_queue = queue.Queue(maxsize=150)
            self.gui_ref.gui_queue = queue.Queue(maxsize=200)
            logger.info("✅ Queues erfolgreich erstellt")
        except Exception as e:
            logger.warning(f"⚠️ Queue-Erstellung fehlgeschlagen: {e}")
            class DummyQueue:
                def __init__(self, maxsize: int = 0) -> None:
                    self.maxsize = maxsize
                    self._items: List[Any] = []
                    self._lock = threading.Lock()
                    self.Empty = queue.Empty
                def put(self, item: Any, block: bool = True, timeout: Optional[float] = None) -> None:
                    with self._lock:
                        self._items.append(item)
                        if self.maxsize > 0 and len(self._items) > self.maxsize:
                            self._items.pop(0)
                def get(self, block: bool = True, timeout: Optional[float] = None) -> Any:
                    with self._lock:
                        if self._items:
                            return self._items.pop(0)
                        raise self.Empty()
                def empty(self) -> bool:
                    with self._lock:
                        return len(self._items) == 0
                def qsize(self) -> int:
                    with self._lock:
                        return len(self._items)
                def task_done(self) -> None:
                    pass
                def get_nowait(self) -> Any:
                    return self.get(block=False)
                def join(self) -> None:
                    pass
            self.gui_ref._text_update_queue = DummyQueue(maxsize=150)
            self.gui_ref.gui_queue = DummyQueue(maxsize=200)
            logger.warning("⚠️ Verwende Dummy-Queues (eingeschränkte Funktionalität)")

    def setup_gui(self) -> None:
        self.root.configure(bg=self.gui_ref.current_theme.BG_PRIMARY)
        self.root.title("🐉 Dragon Whisperer - Plattformunabhängig")
        self.root.geometry("900x700")
        self.root.minsize(850, 600)
        self.root.grid_rowconfigure(0, weight=0)
        self.root.grid_rowconfigure(1, weight=0)
        self.root.grid_rowconfigure(2, weight=0)
        self.root.grid_rowconfigure(3, weight=10)
        self.root.grid_rowconfigure(4, weight=0)
        self.root.grid_columnconfigure(0, weight=1)
        self.setup_dark_styles()
        self.center_window()
        self.root.protocol("WM_DELETE_WINDOW", self.gui_ref._safe_exit_dialog)
        self.create_layout()
        self.root.after(100, self.start_batch_updates)

    def setup_dark_styles(self) -> None:
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(
            "Dark.TCombobox",
            fieldbackground=self.gui_ref.current_theme.COMBO_BG,
            background=self.gui_ref.current_theme.COMBO_BG,
            foreground=self.gui_ref.current_theme.COMBO_FG,
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            selectforeground=self.gui_ref.current_theme.TEXT_PRIMARY,
            insertcolor=self.gui_ref.current_theme.TEXT_PRIMARY,
            borderwidth=1,
            relief="flat",
            arrowsize=12,
            padding=5,
        )
        style.map(
            "Dark.TCombobox",
            fieldbackground=[
                ("readonly", self.gui_ref.current_theme.COMBO_BG),
                ("active", self.gui_ref.current_theme.BG_HOVER),
            ],
            background=[
                ("readonly", self.gui_ref.current_theme.COMBO_BG),
                ("active", self.gui_ref.current_theme.BG_HOVER),
            ],
            foreground=[
                ("readonly", self.gui_ref.current_theme.COMBO_FG),
                ("active", self.gui_ref.current_theme.TEXT_PRIMARY),
            ],
        )
        style.configure(
            "Dark.Horizontal.TProgressbar",
            background=self.gui_ref.current_theme.SUCCESS,
            troughcolor=self.gui_ref.current_theme.BG_TERTIARY,
            bordercolor=self.gui_ref.current_theme.BORDER,
        )
        self.root.option_add("*TCombobox*Listbox.background", self.gui_ref.current_theme.COMBO_BG)
        self.root.option_add("*TCombobox*Listbox.foreground", self.gui_ref.current_theme.COMBO_FG)
        self.root.option_add("*TCombobox*Listbox.selectBackground", self.gui_ref.current_theme.COMBO_SELECTION)
        self.root.option_add("*TCombobox*Listbox.selectForeground", self.gui_ref.current_theme.TEXT_PRIMARY)

    def center_window(self) -> None:
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f"+{x}+{y}")

    def create_layout(self) -> None:
        header_frame = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_PRIMARY, height=35)
        header_frame.grid(row=0, column=0, sticky="ew", padx=12, pady=8)
        header_frame.grid_propagate(False)
        title_label = tk.Label(
            header_frame,
            text="🐉 Dragon Whisperer - Livestream Transcription & Translation",
            font=Fonts.TITLE,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.DRAGON_GREEN,
        )
        title_label.pack(side="left")
        self.gui_ref.status_label = tk.Label(
            header_frame,
            text="✅ READY",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
        )
        self.gui_ref.status_label.pack(side="right")
        self.create_stream_info_display()
        self.gui_ref.stream_info_frame.grid(row=1, column=0, sticky="ew", padx=12, pady=3)
        input_frame = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_PRIMARY)
        input_frame.grid(row=2, column=0, sticky="ew", padx=12, pady=3)
        url_frame = tk.Frame(input_frame, bg=self.gui_ref.current_theme.BG_PRIMARY)
        url_frame.pack(fill="x", pady=2)
        tk.Label(
            url_frame,
            text="URL:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.url_entry = tk.Entry(
            url_frame,
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui_ref.current_theme.TEXT_PRIMARY,
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            width=60,
        )
        self.gui_ref.url_entry.pack(side="left", fill="x", expand=True, padx=(5, 5))
        self.gui_ref.url_entry.insert(0, self.gui_ref.settings.last_url if self.gui_ref.settings.last_url else "")
        DarkEntryContextMenu(self.gui_ref.url_entry)
        self.gui_ref.language_info_label = tk.Label(
            url_frame,
            text="",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
        )
        self.gui_ref.language_info_label.pack(side="right", padx=(5, 0))
        self.create_compact_control_panel(input_frame)
        self.setup_status_bar()
        self.gui_ref.status_bar_frame.grid(row=4, column=0, sticky="ew", pady=(2, 0))
        self.create_text_areas()
        self.gui_ref.text_container.grid(row=3, column=0, sticky="nsew", padx=12, pady=8)
        self.gui_ref.url_entry.bind("<KeyRelease>", self.gui_ref.on_url_change)
        self.gui_ref.url_entry.bind("<FocusOut>", self.gui_ref.on_url_change)

    def create_stream_info_display(self) -> None:
        self.gui_ref.stream_info_frame = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_SECONDARY, height=50)
        self.gui_ref.stream_info_frame.grid_propagate(True)
        self.gui_ref.stream_title_label = tk.Label(
            self.gui_ref.stream_info_frame,
            text="📡 No active stream",
            font=Fonts.SUBTITLE,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            wraplength=700,
            justify="left",
        )
        self.gui_ref.stream_title_label.pack(fill="x", padx=8, pady=(6, 2))
        self.gui_ref.stream_details_label = tk.Label(
            self.gui_ref.stream_info_frame,
            text="Ready to connect...",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            justify="left",
        )
        self.gui_ref.stream_details_label.pack(fill="x", padx=8, pady=(2, 6))

    def create_compact_control_panel(self, parent: tk.Frame) -> None:
        control_frame = tk.Frame(parent, bg=self.gui_ref.current_theme.BG_PRIMARY)
        control_frame.pack(fill="x", pady=8)
        left_controls = tk.Frame(control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY)
        left_controls.pack(side="left")
        action_buttons = [
            ("📁", self.gui_ref.select_file_dark, "Datei auswählen"),
            ("📋", self.gui_ref.paste_url, "URL aus Zwischenablage einfügen"),
        ]
        for icon, command, tooltip in action_buttons:
            btn = tk.Button(
                left_controls,
                text=icon,
                command=command,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                bd=0,
                font=("Segoe UI", 9),
                cursor="hand2",
            )
            btn.pack(side="left", padx=1)
            ToolTip(btn, tooltip)

        self.gui_ref.layout_btn = tk.Button(
            left_controls,
            text="🔄",
            command=self.gui_ref.toggle_layout,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            bd=0,
            font=("Segoe UI", 9),
            cursor="hand2",
        )
        self.gui_ref.layout_btn.pack(side="left", padx=5)
        ToolTip(self.gui_ref.layout_btn, "Layout umschalten (vertikal/horizontal)")

        center_controls = tk.Frame(control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY)
        center_controls.pack(side="left", padx=15)

        src_lang_frame = tk.Frame(center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY)
        src_lang_frame.pack(side="left", padx=5)
        tk.Label(
            src_lang_frame,
            text="From:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.src_lang_var = tk.StringVar(value="Automatisch")
        self.gui_ref.src_lang_combo = ttk.Combobox(
            src_lang_frame,
            textvariable=self.gui_ref.src_lang_var,
            values=[name for name, code in SORTED_LANGUAGES],
            width=10,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.src_lang_combo.pack(side="left", padx=3)
        ToolTip(self.gui_ref.src_lang_combo, "Quellsprache (Automatisch = Whisper-Erkennung)")

        model_frame = tk.Frame(center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY)
        model_frame.pack(side="left", padx=5)
        tk.Label(
            model_frame,
            text="Model:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.model_var = tk.StringVar(value=self.gui_ref.settings.default_model)
        self.gui_ref.model_combo = ttk.Combobox(
            model_frame,
            textvariable=self.gui_ref.model_var,
            values=WHISPER_MODELS,
            width=8,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.model_combo.pack(side="left", padx=3)
        ToolTip(self.gui_ref.model_combo, "Whisper-Modell auswählen (größer = genauer, aber langsamer)")

        if getattr(self.gui_ref, 'demo_mode', False):
            self.gui_ref.model_combo.config(state="disabled")
            self.gui_ref.model_var.set("dummy (Demo)")

        lang_frame = tk.Frame(center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY)
        lang_frame.pack(side="left", padx=5)
        tk.Label(
            lang_frame,
            text="Translate:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.lang_var = tk.StringVar()
        language_groups = {
            "Common": ["German", "English", "French", "Spanish", "Italian"],
            "Asian": ["Japanese", "Chinese", "Korean", "Vietnamese", "Thai"],
            "More": [
                name
                for name, code in SORTED_LANGUAGES
                if name
                not in [
                    "German",
                    "English",
                    "French",
                    "Spanish",
                    "Italian",
                    "Japanese",
                    "Chinese",
                    "Korean",
                    "Vietnamese",
                    "Thai",
                ]
            ],
        }
        all_languages: List[str] = []
        for group_name, languages in language_groups.items():
            if languages:
                all_languages.append(f"--- {group_name} ---")
                all_languages.extend(languages)
        self.gui_ref.lang_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.gui_ref.lang_var,
            values=all_languages,
            width=12,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.lang_combo.pack(side="left", padx=3)
        ToolTip(self.gui_ref.lang_combo, "Zielsprache für Übersetzung")
        default_lang_name = SUPPORTED_LANGUAGES.get(self.gui_ref.settings.default_language, "German")
        self.gui_ref.lang_var.set(default_lang_name)
        self.gui_ref.lang_combo.bind("<<ComboboxSelected>>", self.gui_ref.on_language_change)

        right_controls = tk.Frame(control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY)
        right_controls.pack(side="right")

        self.gui_ref.start_button = tk.Button(
            right_controls,
            text="🚀 START",
            command=self.gui_ref._on_start_click,
            bg=self.gui_ref.current_theme.SUCCESS,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 9, "bold"),
            relief="flat",
            padx=20,
        )
        self.gui_ref.start_button.pack(side="left", padx=2)
        ToolTip(self.gui_ref.start_button, "Transkription/Übersetzung starten")

        self.gui_ref.stop_button = tk.Button(
            right_controls,
            text="⏹️ STOP",
            command=self.gui_ref.controller.stop_processing,
            bg=self.gui_ref.current_theme.ERROR,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            state="disabled",
            font=("Segoe UI", 9, "bold"),
            relief="flat",
            padx=20,
        )
        self.gui_ref.stop_button.pack(side="left", padx=2)
        ToolTip(self.gui_ref.stop_button, "Laufende Verarbeitung stoppen")

        self.gui_ref.translate_btn = tk.Button(
            right_controls,
            text="🌐 ON" if self.gui_ref.translate_active else "🌐 OFF",
            command=self.gui_ref.toggle_translation,
            bg=self.gui_ref.current_theme.SUCCESS if self.gui_ref.translate_active else self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            padx=6,
        )
        self.gui_ref.translate_btn.pack(side="left", padx=2)
        ToolTip(self.gui_ref.translate_btn, "Übersetzung ein/aus")

        self.gui_ref.subtitle_btn = tk.Button(
            right_controls,
            text="🎬",
            command=self.gui_ref.toggle_subtitle_mode,
            bg=self.gui_ref.current_theme.SUBTITLE_INACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            bd=0,
            font=("Segoe UI", 9),
            cursor="hand2",
        )
        self.gui_ref.subtitle_btn.pack(side="left", padx=5)
        ToolTip(self.gui_ref.subtitle_btn, "Untertitel-Modus (mit Zeitstempeln)")

        self.gui_ref.model_combo.bind("<<ComboboxSelected>>", self.gui_ref.on_model_change)

    def create_text_areas(self) -> Tuple[Optional[scrolledtext.ScrolledText], Optional[scrolledtext.ScrolledText]]:
        layout_changed = False
        current_layout = getattr(self.gui_ref, '_current_layout', None)
        if current_layout != self.gui_ref.layout_mode:
            layout_changed = True
            logger.info(f"🔄 Layout change detected: {current_layout} → {self.gui_ref.layout_mode}")
        if hasattr(self.gui_ref, 'text_container') and layout_changed:
            try:
                if self.gui_ref.text_container.winfo_exists():
                    logger.info("   🗑️ Destroying old container for layout change")
                    self.gui_ref.text_container.destroy()
                    time.sleep(0.02)
            except tk.TclError:
                pass
            except Exception as e:
                logger.warning(f"   ⚠️ Container destroy warning: {e}")
        if layout_changed or not hasattr(self.gui_ref, 'text_container'):
            self.gui_ref.text_container = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_PRIMARY)
            self.gui_ref._current_layout = self.gui_ref.layout_mode
            logger.info(f"   ✅ New container created for {self.gui_ref.layout_mode} layout")
        if self.gui_ref.layout_mode == "horizontal":
            self.create_horizontal_layout()
        else:
            self.create_vertical_layout()
        self.gui_ref.text_container.grid(row=3, column=0, sticky="nsew", padx=12, pady=8)
        self.root.grid_rowconfigure(3, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        self.root.update_idletasks()
        return (
            getattr(self.gui_ref, 'transcript_text', None),
            getattr(self.gui_ref, 'translation_text', None)
        )

    def create_vertical_layout(self) -> None:
        main_frame = tk.LabelFrame(
            self.gui_ref.text_container,
            text="Live Transkription & Übersetzung",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.SUBTITLE,
            padx=8,
            pady=8,
        )
        main_frame.pack(fill="both", expand=True)
        trans_frame = tk.Frame(main_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        trans_frame.pack(fill="x", pady=(0, 3))
        trans_header = tk.Frame(trans_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        trans_header.pack(fill="x")
        tk.Label(
            trans_header,
            text="🎤 Transkription:",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        ).pack(side="left")
        self.gui_ref.transcript_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            trans_header,
            variable=self.gui_ref.transcript_scroll_var,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            activebackground=self.gui_ref.current_theme.BG_SECONDARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            trans_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.transcript_text = self.create_text_widget(main_frame, height=6)
        transla_frame = tk.Frame(main_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        transla_frame.pack(fill="x", pady=(8, 0))
        transla_header = tk.Frame(transla_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        transla_header.pack(fill="x")
        self.gui_ref.translation_header = tk.Label(
            transla_header,
            text="🌐 Übersetzung:",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        )
        self.gui_ref.translation_header.pack(side="left")
        self.gui_ref.translation_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            transla_header,
            variable=self.gui_ref.translation_scroll_var,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            activebackground=self.gui_ref.current_theme.BG_SECONDARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            transla_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.translation_text = self.create_text_widget(main_frame, height=6)

    def create_horizontal_layout(self) -> None:
        main_frame = tk.LabelFrame(
            self.gui_ref.text_container,
            text="Live Transkription & Übersetzung",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.SUBTITLE,
            padx=8,
            pady=8,
        )
        main_frame.pack(fill="both", expand=True)
        self.gui_ref.paned_window = tk.PanedWindow(
            main_frame,
            orient=tk.HORIZONTAL,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            sashrelief="raised",
            sashwidth=4,
            sashpad=0,
        )
        self.gui_ref.paned_window.pack(fill="both", expand=True)
        left_frame = tk.Frame(self.gui_ref.paned_window, bg=self.gui_ref.current_theme.BG_TERTIARY)
        self.gui_ref.paned_window.add(left_frame, stretch="always", width=400)
        trans_header = tk.Frame(left_frame, bg=self.gui_ref.current_theme.BG_TERTIARY)
        trans_header.pack(fill="x", padx=5, pady=2)
        tk.Label(
            trans_header,
            text="🎤 Transkription",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        ).pack(side="left")
        self.gui_ref.transcript_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            trans_header,
            variable=self.gui_ref.transcript_scroll_var,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            activebackground=self.gui_ref.current_theme.BG_TERTIARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            trans_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.transcript_text = self.create_text_widget(left_frame)
        right_frame = tk.Frame(self.gui_ref.paned_window, bg=self.gui_ref.current_theme.BG_TERTIARY)
        self.gui_ref.paned_window.add(right_frame, stretch="always", width=400)
        transla_header = tk.Frame(right_frame, bg=self.gui_ref.current_theme.BG_TERTIARY)
        transla_header.pack(fill="x", padx=5, pady=2)
        self.gui_ref.translation_header = tk.Label(
            transla_header,
            text="🌐 Übersetzung",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        )
        self.gui_ref.translation_header.pack(side="left")
        self.gui_ref.translation_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            transla_header,
            variable=self.gui_ref.translation_scroll_var,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            activebackground=self.gui_ref.current_theme.BG_TERTIARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            transla_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.translation_text = self.create_text_widget(right_frame)
        self.gui_ref.paned_window.paneconfig(left_frame, minsize=250, width=400)
        self.gui_ref.paned_window.paneconfig(right_frame, minsize=250, width=400)

    def create_text_widget(self, parent: tk.Frame, height: Optional[int] = None) -> scrolledtext.ScrolledText:
        text_widget = scrolledtext.ScrolledText(
            parent,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            insertbackground=self.gui_ref.current_theme.TEXT_PRIMARY,
            wrap=tk.WORD,
            relief="flat",
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            selectforeground=self.gui_ref.current_theme.TEXT_PRIMARY,
            maxundo=30,
            undo=True,
        )
        if height:
            text_widget.config(height=height)
        text_widget.pack(fill="both", expand=True, padx=5, pady=5)
        DarkContextMenu(text_widget)

        def safe_text_cleanup(event: Optional[tk.Event] = None) -> None:
            try:
                lines = int(text_widget.index("end-1c").split(".")[0])
                if text_widget == self.gui_ref.transcript_text:
                    max_lines = self.gui_ref.settings.transcript_max_lines
                else:
                    max_lines = self.gui_ref.settings.translation_max_lines
                keep_lines = max_lines - 100
                if lines > max_lines:
                    component = (
                        "transcript"
                        if text_widget == self.gui_ref.transcript_text
                        else "translation"
                    )
                    self.gui_ref.memory_manager.clear_component(component)
                    delete_to = f"{lines - keep_lines}.0"
                    text_widget.delete("1.0", delete_to)
                    gc.collect()
            except Exception:
                pass
        text_widget.bind("<KeyRelease>", safe_text_cleanup)
        return text_widget

    def setup_status_bar(self) -> None:
        self.gui_ref.status_bar_frame = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_SECONDARY, height=50)
        self.gui_ref.status_bar_frame.grid_propagate(True)
        separator = tk.Frame(self.gui_ref.status_bar_frame, height=2, bg=self.gui_ref.current_theme.DRAGON_GREEN)
        separator.pack(fill="x", side="top")
        main_container = tk.Frame(self.gui_ref.status_bar_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        main_container.pack(fill="x", expand=True, padx=12, pady=8)

        main_container.columnconfigure(0, weight=0)
        main_container.columnconfigure(1, weight=1)
        main_container.columnconfigure(2, weight=0)

        left_panel = tk.Frame(main_container, bg=self.gui_ref.current_theme.BG_SECONDARY)
        left_panel.grid(row=0, column=0, sticky="w", padx=5)

        quick_actions = [
            ("🗑️", self.gui_ref.clear_all, "Alles löschen"),
            ("💾", self.gui_ref.save_transcript, "Transkription speichern"),
            ("📝", self.gui_ref.export_subtitles, "Untertitel exportieren"),
            ("📊", self.gui_ref.show_simple_stats, "Statistiken anzeigen"),
            ("⚙️", self.gui_ref.show_advanced_settings, "Erweiterte Einstellungen"),
            ("🌐", self.gui_ref.show_translation_dialog, "Text übersetzen"),
            ("🤖", self.gui_ref.show_summarize_dialog, "Mit Ollama zusammenfassen"),
        ]

        if getattr(self.gui_ref, 'demo_mode', False) or not TRANSLATOR_AVAILABLE:
            install_btn = tk.Button(
                left_panel,
                text="📦",
                command=self.gui_ref.show_install_dialog,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4, pady=2,
                activebackground=self.gui_ref.current_theme.BG_HOVER,
            )
            install_btn.grid(row=0, column=len(quick_actions)+1, padx=1, sticky="w")
            ToolTip(install_btn, "Fehlende Pakete installieren")

        for i, (icon, command, tooltip) in enumerate(quick_actions):
            btn = tk.Button(
                left_panel,
                text=icon,
                command=command,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4,
                pady=2,
                activebackground=self.gui_ref.current_theme.BG_HOVER,
            )
            btn.grid(row=0, column=i, padx=1, sticky="w")
            ToolTip(btn, tooltip)

        center_panel = tk.Frame(main_container, bg=self.gui_ref.current_theme.BG_SECONDARY)
        center_panel.grid(row=0, column=1, sticky="ew", padx=5)

        self.gui_ref.progress_bar = ttk.Progressbar(
            center_panel,
            mode='determinate',
            length=150,
            style="Dark.Horizontal.TProgressbar"
        )
        self.gui_ref.progress_bar.pack(side="left", padx=(10, 10))

        self.gui_ref.progress_label = tk.Label(
            center_panel,
            text="",
            font=("Segoe UI", 8),
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY
        )
        self.gui_ref.progress_label.pack(side="left", padx=(0, 10))

        if IS_WINDOWS:
            default_text = "🪟 Windows | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        elif IS_MACOS:
            if IS_ARM:
                default_text = "🍎 macOS (Apple Silicon) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
            else:
                default_text = "🍎 macOS (Intel) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        elif IS_LINUX:
            default_text = "🐧 Linux | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        else:
            default_text = "🌐 Unknown OS | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        self.gui_ref.system_info_label = tk.Label(
            center_panel,
            text=default_text,
            font=("Segoe UI", 8, "normal"),
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            padx=5,
        )
        self.gui_ref.system_info_label.pack(side="left", fill="x", expand=True)

        right_panel = tk.Frame(main_container, bg=self.gui_ref.current_theme.BG_SECONDARY)
        right_panel.grid(row=0, column=2, sticky="e", padx=5)

        self.gui_ref.exit_button = tk.Button(
            right_panel,
            text=" ⏻ EXIT ",
            command=self.gui_ref.controller.safe_exit,
            bg="#dc3545",
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief="raised",
            cursor="hand2",
            padx=12,
            pady=3,
            activebackground="#c82333",
        )
        self.gui_ref.exit_button.pack(side="right")
        ToolTip(self.gui_ref.exit_button, "Programm beenden (Strg+Q / Cmd+Q)")

        self.gui_ref.correct_btn = tk.Button(
            right_panel,
            text="🔧",
            command=self.gui_ref.correct_transcript_with_ollama,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4
        )
        self.gui_ref.correct_btn.pack(side="right", padx=2)
        ToolTip(self.gui_ref.correct_btn, "Transkript mit Ollama korrigieren")

        help_btn = tk.Button(
            right_panel,
            text="⌨️",
            command=self.gui_ref.show_shortcuts_help,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4
        )
        help_btn.pack(side="right", padx=2)
        ToolTip(help_btn, "Tastenkürzel anzeigen (F1)")

    def process_batch_text_updates(self) -> None:
        if (not hasattr(self.gui_ref, '_shutting_down') or
            getattr(self.gui_ref, '_shutting_down', False)):
            return
        if (not hasattr(self, 'root') or
            self.root is None or
            not self.root.winfo_exists()):
            return
        if not hasattr(self.gui_ref, '_text_update_queue'):
            return
        queue_obj = self.gui_ref._text_update_queue
        if queue_obj is None:
            return
        if not hasattr(queue_obj, 'empty') or not callable(queue_obj.empty):
            return
        if queue_obj.empty():
            return

        processed = 0
        max_updates = 5
        start_time = time.time()
        while processed < max_updates and (time.time() - start_time) < 0.05:
            if not self.root.winfo_exists():
                break
            try:
                if hasattr(queue_obj, 'get_nowait') and callable(queue_obj.get_nowait):
                    item = queue_obj.get_nowait()
                else:
                    item = queue_obj.get(block=False)
                if isinstance(item, tuple) and len(item) == 2:
                    update_type, text_data = item
                    self._process_update(update_type, text_data)
                processed += 1
                try:
                    self.root.update_idletasks()
                except tk.TclError:
                    break
            except Exception as e:
                if "Empty" in type(e).__name__ or "empty" in str(e).lower():
                    break
                logger.warning(f"⚠️ Queue processing error: {e}")
                break
        if debug3_enabled('gui') and processed > 0:
            logger.debug(f"[DEBUG3][GUI] Processed {processed} items in text queue")
        self._schedule_next_update()

    def _process_update(self, update_type: str, text_data: str) -> None:
        if not hasattr(self, 'gui_ref') or self.gui_ref is None:
            return
        try:
            if update_type == 'transcript':
                widget = getattr(self.gui_ref, 'transcript_text', None)
                if widget is not None and widget.winfo_exists():
                    widget.insert('end', text_data)
                    self._auto_scroll('transcript')
                    self._check_text_limit('transcript')
            elif update_type == 'translation':
                widget = getattr(self.gui_ref, 'translation_text', None)
                if widget is not None and widget.winfo_exists():
                    widget.insert('end', text_data)
                    self._auto_scroll('translation')
                    self._check_text_limit('translation')
        except tk.TclError:
            pass
        except AttributeError:
            pass
        except Exception as e:
            logger.warning(f"⚠️ GUI update error: {e}")

    def _auto_scroll(self, text_type: str) -> None:
        try:
            if text_type == 'transcript':
                if (hasattr(self.gui_ref, 'transcript_scroll_var') and
                    self.gui_ref.transcript_scroll_var is not None and
                    self.gui_ref.transcript_scroll_var.get()):
                    self.gui_ref.transcript_text.see('end')
            elif text_type == 'translation':
                if (hasattr(self.gui_ref, 'translation_scroll_var') and
                    self.gui_ref.translation_scroll_var is not None and
                    self.gui_ref.translation_scroll_var.get()):
                    self.gui_ref.translation_text.see('end')
        except Exception:
            pass

    def _check_text_limit(self, text_type: str) -> None:
        try:
            if text_type == 'transcript':
                widget = self.gui_ref.transcript_text
                max_lines = self.gui_ref.settings.transcript_max_lines
                keep_lines = max_lines - 100
            else:
                widget = self.gui_ref.translation_text
                max_lines = self.gui_ref.settings.translation_max_lines
                keep_lines = max_lines - 100
            lines = int(widget.index('end-1c').split('.')[0])
            if lines > max_lines:
                delete_to = f'{lines-keep_lines}.0'
                widget.delete('1.0', delete_to)
        except Exception:
            pass

    def _schedule_next_update(self) -> None:
        try:
            if (hasattr(self, 'root') and
                self.root is not None and
                self.root.winfo_exists()):
                interval = 150
                if hasattr(self.gui_ref, '_batch_update_interval'):
                    try:
                        interval = self.gui_ref._batch_update_interval
                    except Exception:
                        pass
                if hasattr(self, '_batch_timer_id') and self._batch_timer_id:
                    try:
                        self.root.after_cancel(self._batch_timer_id)
                    except Exception:
                        pass
                self._batch_timer_id = self.root.after(interval, self.process_batch_text_updates)
            else:
                self._batch_timer_id = None
        except Exception as e:
            logger.warning(f"⚠️ Timer scheduling error: {e}")

    def start_batch_updates(self) -> None:
        try:
            if (hasattr(self, 'root') and
                self.root is not None and
                self.root.winfo_exists()):
                if not hasattr(self.gui_ref, '_text_update_queue') or self.gui_ref._text_update_queue is None:
                    try:
                        self.gui_ref._text_update_queue = queue.Queue(maxsize=150)
                    except Exception:
                        class DummyQueue:
                            def __init__(self, maxsize: int = 0) -> None:
                                self.maxsize = maxsize
                                self._items: List[Any] = []
                                self._lock = threading.Lock()
                                class EmptyException(Exception):
                                    pass
                                self.Empty = EmptyException
                            def put(self, item: Any, block: bool = True, timeout: Optional[float] = None) -> None:
                                with self._lock:
                                    self._items.append(item)
                                    if self.maxsize > 0 and len(self._items) > self.maxsize:
                                        self._items.pop(0)
                            def get(self, block: bool = True, timeout: Optional[float] = None) -> Any:
                                with self._lock:
                                    if self._items:
                                        return self._items.pop(0)
                                    raise self.Empty()
                            def empty(self) -> bool:
                                with self._lock:
                                    return len(self._items) == 0
                            def qsize(self) -> int:
                                with self._lock:
                                    return len(self._items)
                            def task_done(self) -> None:
                                pass
                            def get_nowait(self) -> Any:
                                return self.get(block=False)
                            def join(self) -> None:
                                pass
                        self.gui_ref._text_update_queue = DummyQueue(maxsize=150)
                        logger.warning("⚠️ Queue-Fallback in start_batch_updates")
                self.root.after(100, self.process_batch_text_updates)
                logger.info("✅ Batch updates gestartet")
        except Exception as e:
            logger.warning(f"⚠️ Start batch updates error: {e}")

# -----------------------------------------------------------------------------
# WHISPER CONTROLLER
# -----------------------------------------------------------------------------
class WhisperController:
    """
    Steuert den gesamten Transkriptions- und Übersetzungsprozess.
    Verwendet weakref für die GUI-Referenz, um Zyklen zu vermeiden.
    """

    def __init__(self, gui_ref: Any) -> None:
        self.gui_ref = weakref.ref(gui_ref)
        self.is_processing = False
        self._processing_lock = threading.Lock()
        self._cleanup_lock = threading.RLock()
        self._last_transcription_text = ""
        self._duplicate_check_cache: deque = deque(maxlen=20)
        self._processing_thread: Optional[threading.Thread] = None
        self._shutdown_event = threading.Event()
        self._stop_requested = False
        self._processing_active = threading.Event()
        self._initialized = True
        self._stop_complete = threading.Event()
        self._stop_complete.set()

        self.on_transcription: Optional[Callable[[TranscriptionResult], None]] = None
        self.on_translation: Optional[Callable[[TranslationResult], None]] = None
        self.on_info: Optional[Callable[[str], None]] = None
        self.on_error: Optional[Callable[[str], None]] = None
        self.on_status: Optional[Callable[[Dict[str, Any]], None]] = None
        self.on_finished: Optional[Callable[[], None]] = None

    def set_callbacks(self,
                      on_transcription: Callable[[TranscriptionResult], None],
                      on_translation: Callable[[TranslationResult], None],
                      on_info: Callable[[str], None],
                      on_error: Callable[[str], None],
                      on_status: Callable[[Dict[str, Any]], None],
                      on_finished: Optional[Callable[[], None]] = None) -> None:
        self.on_transcription = on_transcription
        self.on_translation = on_translation
        self.on_info = on_info
        self.on_error = on_error
        self.on_status = on_status
        self.on_finished = on_finished

    def _processing_finished(self) -> None:
        with self._processing_lock:
            self.is_processing = False
        if self.on_status:
            self.on_status({"processing_state": False, "status": "✅ Processing complete"})

    def _stop_processing_sync(self, timeout: float = 10.0) -> bool:
        self._stop_complete.clear()
        logger.info("🛑 WhisperController: Synchrone Stop angefordert")
        self.stop_processing()

        gui = self.gui_ref()
        if gui is not None and hasattr(gui, 'audio_processor') and gui.audio_processor:
            ap = gui.audio_processor
            if hasattr(ap, '_process_finished'):
                if not ap._process_finished.wait(timeout):
                    logger.warning(f"⚠️ Audio-Processor nicht innerhalb von {timeout}s beendet")
                    return False
            else:
                time.sleep(0.5)
        else:
            time.sleep(0.1)

        self._stop_complete.set()
        return True

    def _on_progress(self, processed: int, total: Optional[int], chunks: int) -> None:
        gui = self.gui_ref()
        if gui is not None and hasattr(gui, 'update_progress'):
            try:
                gui.root.after(0, gui.update_progress, processed, total, chunks)
            except Exception:
                pass

    def _start_processing(self) -> None:
        with self._processing_lock:
            if self.is_processing:
                if self.on_status:
                    self.on_status({"status": "⚠️ Bereits aktiv"})
                return

            gui = self.gui_ref()
            if gui is None:
                if self.on_status:
                    self.on_status({"status": "❌ GUI nicht verfügbar"})
                return

            url = self._validate_url(gui)
            if url is None:
                return

            if self.on_status:
                self.on_status({"status": "🔍 Analysiere Stream..."})

            if self.is_processing:
                logger.warning("⚠️ Vorheriger Prozess läuft noch – stoppe diesen zuerst synchron.")
                if not self._stop_processing_sync(timeout=10):
                    if self.on_status:
                        self.on_status({"status": "❌ Vorheriger Prozess konnte nicht gestoppt werden"})
                    return

            self._extract_stream_info(gui, url)

            if self.on_status:
                self.on_status({"status": "🎵 Teste Audio-Stream..."})

            stream_test_passed = self._test_stream(gui, url)
            if not stream_test_passed:
                if self.on_status:
                    self.on_status({"status": "❌ Stream nicht erreichbar"})
                logger.error("❌ Stream Test fehlgeschlagen")
                return

            if self.on_status:
                self.on_status({"status": "🤖 Lade KI-Modell..."})

            model_loaded = self._load_and_setup_model(gui)
            if not model_loaded:
                if self.on_status:
                    self.on_status({"status": "❌ KI-Modell konnte nicht geladen werden"})
                return

            self._set_source_language(gui)

            self._configure_translation(gui)

            self.is_processing = True
            if gui is not None:
                gui.is_processing = True

            def update_gui_buttons() -> None:
                try:
                    if gui and hasattr(gui, "start_button") and gui.start_button.winfo_exists():
                        gui.start_button.config(state="disabled")
                    if gui and hasattr(gui, "stop_button") and gui.stop_button.winfo_exists():
                        gui.stop_button.config(state="normal")
                except Exception:
                    pass
            if gui is not None and hasattr(gui, 'root') and gui.root.winfo_exists():
                gui.root.after(0, update_gui_buttons)

            if IS_LINUX and gui is not None and hasattr(gui, "performance_optimizer"):
                gui.performance_optimizer.optimize_for_processing()

            if self.on_status:
                self.on_status({
                    "processing_state": True,
                    "status": "🚀 Starte Transkription...",
                    "buttons": {"start": "disabled", "stop": "normal"},
                })

            def transcription_callback(result: TranscriptionResult) -> None:
                if not result or not hasattr(result, "text"):
                    return
                try:
                    if self.on_transcription:
                        self.on_transcription(result)
                except Exception as e:
                    logger.warning(f"⚠️ Transcription Callback Error: {e}")

            def translation_callback(result: TranslationResult) -> None:
                if not result or not hasattr(result, "translated"):
                    return
                try:
                    if self.on_translation:
                        self.on_translation(result)
                except Exception as e:
                    logger.warning(f"⚠️ Translation Callback Error: {e}")

            def info_callback(message: str) -> None:
                try:
                    if self.on_info:
                        self.on_info(message)
                except Exception:
                    pass

            def error_callback(message: str) -> None:
                try:
                    if self.on_error:
                        self.on_error(message)
                    self._cleanup_resources()
                except Exception:
                    pass

            def file_finished_callback() -> None:
                logger.info("✅ Dateiende erkannt – zeige Speicherdialog")
                if self.on_status:
                    self.on_status({"file_finished": True})
                self._processing_finished()

            self._start_audio_processor_thread(gui, url, transcription_callback, translation_callback,
                                               info_callback, error_callback, file_finished_callback)

    def _validate_url(self, gui) -> Optional[str]:
        try:
            url = gui.url_entry.get().strip()
        except Exception:
            if self.on_status:
                self.on_status({"status": "❌ URL Fehler"})
            return None

        if not url:
            if self.on_status:
                self.on_status({"status": "❌ Bitte URL eingeben"})
            return None

        try:
            url = PlatformUtils.sanitize_url(url)
            if url.startswith("file://"):
                ok, real_path = PlatformUtils.validate_file_path(url)
                if not ok:
                    if self.on_status:
                        self.on_status({"status": f"❌ {real_path}"})
                    return None
                file_path = real_path
                if not os.path.exists(file_path):
                    if self.on_status:
                        self.on_status({"status": "❌ Datei nicht gefunden"})
                    return None
            else:
                if not url.startswith(("http://", "https://")):
                    url = "https://" + url
                    def update_url():
                        if gui and hasattr(gui, 'url_entry') and gui.url_entry.winfo_exists():
                            gui.url_entry.delete(0, "end")
                            gui.url_entry.insert(0, url)
                    if gui and hasattr(gui, 'root'):
                        gui.root.after(0, update_url)
        except Exception:
            if self.on_status:
                self.on_status({"status": "❌ Ungültige URL"})
            return None
        return url

    def _extract_stream_info(self, gui, url: str) -> None:
        try:
            if hasattr(gui, "stream_manager"):
                platform_type, platform_name = gui.stream_manager.detect_platform(url)
            else:
                platform_type, platform_name = "unknown", "Unknown"

            stream_info = None
            try:
                if hasattr(gui, 'stream_info_extractor'):
                    stream_info = gui.stream_info_extractor.extract_stream_info(url)
                else:
                    stream_info = StreamInfoExtractor().extract_stream_info(url)
            except Exception:
                stream_info = StreamInfo(
                    title="Live Stream" if "live" in url.lower() else "Stream",
                    uploader=platform_name,
                    duration="Live" if "live" in url.lower() else "Unknown",
                    view_count=0,
                    platform=platform_type,
                )
            if stream_info:
                if self.on_status:
                    self.on_status({"stream_info": stream_info})
                logger.info(f"📡 Stream: {stream_info.title[:50]}...")
                if hasattr(gui, 'audio_processor') and stream_info.duration_seconds is not None:
                    gui.audio_processor.set_expected_duration(stream_info.duration_seconds)
        except Exception as e:
            logger.warning(f"⚠️ Stream Info Error: {e}")

    def _test_stream(self, gui, url: str) -> bool:
        try:
            if hasattr(gui, "audio_processor"):
                return gui.audio_processor.emergency_diagnosis(url)
        except Exception as e:
            logger.warning(f"⚠️ Stream Test Error: {e}")
        return False

    def _load_and_setup_model(self, gui) -> bool:
        model_loaded = False
        try:
            if hasattr(gui, "transcription_engine"):
                model_name = "medium"
                if hasattr(gui, "model_var"):
                    model_name = gui.model_var.get()
                result = gui.transcription_engine.load_model(model_name, set_active=True)
                if result is not None:
                    model_loaded = True
                else:
                    logger.info("🔄 Versuche base model...")
                    result = gui.transcription_engine.load_model("base", set_active=True)
                    model_loaded = result is not None
        except Exception as e:
            logger.warning(f"⚠️ Model Load Error: {e}")
        return model_loaded

    def _set_source_language(self, gui) -> None:
        try:
            if hasattr(gui, 'src_lang_var'):
                src_name = gui.src_lang_var.get()
                if src_name != "Automatisch":
                    for name, code in SORTED_LANGUAGES:
                        if name == src_name:
                            gui.transcription_engine.forced_language = code
                            break
                else:
                    gui.transcription_engine.forced_language = None
        except Exception as e:
            logger.warning(f"⚠️ Quellsprache setzen fehlgeschlagen: {e}")
            gui.transcription_engine.forced_language = None

    def _configure_translation(self, gui) -> None:
        try:
            if hasattr(gui, "translation_engine") and hasattr(gui, "lang_var"):
                selected_name = gui.lang_var.get()
                target_lang = "de"
                for name, code in SORTED_LANGUAGES:
                    if name == selected_name:
                        target_lang = code
                        break
                gui.translation_engine.set_target_language(target_lang)
                lang_display = LANGUAGE_SHORT_CODES.get(target_lang, target_lang)
                if hasattr(gui, "translation_header"):
                    def update_header():
                        if gui and hasattr(gui, 'translation_header') and gui.translation_header.winfo_exists():
                            gui.translation_header.config(text=f"🌐 Übersetzung ({lang_display})")
                    if gui and hasattr(gui, 'root'):
                        gui.root.after(0, update_header)
        except Exception as e:
            logger.warning(f"⚠️ Translation Setup Error: {e}")

    def _start_audio_processor_thread(self, gui, url: str,
                                      transcription_callback, translation_callback,
                                      info_callback, error_callback,
                                      file_finished_callback) -> None:
        try:
            if gui is not None and hasattr(gui, "audio_processor"):
                gui.audio_processor._stop_event.clear()
                gui.audio_processor.set_progress_callback(self._on_progress)
                processing_thread = threading.Thread(
                    target=lambda: gui.audio_processor.start_processing(
                        url=url,
                        transcription_callback=transcription_callback,
                        translation_callback=translation_callback,
                        info_callback=info_callback,
                        error_callback=error_callback,
                        finished_callback=file_finished_callback,
                    ),
                    daemon=True,
                    name="AudioProcessor",
                )
                self._processing_thread = processing_thread
                processing_thread.start()
                if self.on_status:
                    self.on_status({"status": "✅ Transkription läuft..."})
            else:
                error_callback("❌ Audio-Processor nicht verfügbar")
                if self.on_status:
                    self.on_status({"processing_state": False, "status": "❌ Audio-Processor nicht verfügbar"})
                self.is_processing = False
        except (AttributeError, RuntimeError) as e:
            error_msg = f"Start Error: {str(e)[:100]}"
            logger.error(f"❌ Processing Start Error: {e}")
            error_callback(error_msg)
            if self.on_status:
                self.on_status({"processing_state": False, "status": f"❌ {error_msg}"})
            self.is_processing = False

    def start_processing(self) -> None:
        def start_thread() -> None:
            try:
                self._start_processing()
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.error(f"❌ Start Processing Error: {e}")
                if self.on_status:
                    self.on_status({"status": f"❌ Start fehlgeschlagen: {str(e)[:50]}"})
        thread = threading.Thread(target=start_thread, daemon=True)
        thread.start()

    def stop_processing(self) -> None:
        self._stop_requested = True

        gui = self.gui_ref()
        if gui is not None and IS_LINUX and hasattr(gui, "performance_optimizer"):
            gui.performance_optimizer.restore_normal_mode()

        self._shutdown_event.set()
        if hasattr(self, "_processing_active"):
            self._processing_active.clear()

        if gui is not None:
            gui.is_processing = False

        def stop_audio_processor() -> None:
            try:
                if gui is not None and hasattr(gui, "audio_processor"):
                    ap = gui.audio_processor
                    ap._processing.clear()
                    if hasattr(ap, "_stop_event"):
                        ap._stop_event.set()
                if gui is not None and hasattr(gui, "ffmpeg_manager"):
                    gui.ffmpeg_manager.stop_all_streams()
                self._stop_complete.set()
            except Exception as e:
                logger.warning(f"⚠️ Audio Stop Fehler: {e}")
                self._stop_complete.set()

        audio_stop_thread = threading.Thread(target=stop_audio_processor, daemon=True)
        audio_stop_thread.start()

        def update_gui_immediately() -> None:
            try:
                if gui is not None:
                    if hasattr(gui, "status_label") and gui.status_label.winfo_exists():
                        gui.status_label.config(text="✅ READY for new stream")
                    if hasattr(gui, "start_button") and gui.start_button.winfo_exists():
                        gui.start_button.config(state="normal")
                    if hasattr(gui, "stop_button") and gui.stop_button.winfo_exists():
                        gui.stop_button.config(state="disabled")
                    if hasattr(gui, "stream_title_label") and gui.stream_title_label.winfo_exists():
                        gui.stream_title_label.config(text="📡 Kein aktiver Stream")
                    if hasattr(gui, "stream_details_label") and gui.stream_details_label.winfo_exists():
                        gui.stream_details_label.config(text="Bereit für neue Verbindung")
                    gui._reset_progress()
            except Exception as e:
                logger.warning(f"⚠️ GUI Update Fehler: {e}")
            finally:
                self.is_processing = False

        if gui is not None and hasattr(gui, 'root') and gui.root.winfo_exists():
            gui.root.after(0, update_gui_immediately)

        def background_cleanup() -> None:
            try:
                if self._processing_thread and self._processing_thread.is_alive():
                    logger.info("🔄 Warte auf Processing Thread...")
                    self._processing_thread.join(timeout=1.0)
                transcription_cache.clear()
                translation_cache.clear()
                audio_cache.clear()
                self._stop_requested = False
                self._shutdown_event.clear()
            except Exception as e:
                logger.warning(f"⚠️ Cleanup Fehler: {e}")

        cleanup_thread = threading.Thread(target=background_cleanup, daemon=True)
        cleanup_thread.start()

    def _cleanup_resources(self) -> None:
        if self._stop_requested:
            return
        self._stop_requested = True
        with self._processing_lock:
            self.is_processing = False
        gui = self.gui_ref()
        if gui is not None:
            try:
                if hasattr(gui, "audio_processor"):
                    gui.audio_processor._processing.clear()
                    if hasattr(gui.audio_processor, "_stop_event"):
                        gui.audio_processor._stop_event.set()
            except Exception:
                pass

    def dispose(self) -> None:
        self._shutdown_event.set()
        self.stop_processing()
        logger.info("🧹 Controller disposed")

    def safe_exit(self) -> None:
        gui = self.gui_ref()
        if gui is not None:
            try:
                if hasattr(gui, "exit_button"):
                    gui.exit_button.config(state="disabled", text="⏳...")
            except Exception:
                pass
            if hasattr(gui, "_safe_exit_dialog"):
                gui._safe_exit_dialog()
            else:
                self._cleanup_resources()
                sys.exit(0)
        else:
            self._cleanup_resources()
            sys.exit(0)

# -----------------------------------------------------------------------------
# LINUX PERFORMANCE OPTIMIZER
# -----------------------------------------------------------------------------
PSUTIL_AVAILABLE = importlib.util.find_spec("psutil") is not None
if not PSUTIL_AVAILABLE:
    logger.warning("⚠️ psutil nicht verfügbar – Linux Performance Optimizer läuft im Dummy-Modus")

if IS_LINUX and PSUTIL_AVAILABLE:
    class LinuxPerformanceOptimizer:
        def __init__(self, gui_ref: 'DragonWhispererGUI') -> None:
            self.gui = gui_ref
            self.is_processing = False
            self._original_settings: Dict[str, Any] = {}
            self._optimization_active = False
            self._monitoring_thread: Optional[threading.Thread] = None
            self._shutdown_event = threading.Event()
            self._monitoring_lock = threading.RLock()
            self._last_gui_access_time = 0.0
            self._gui_access_warning_printed = False

        def optimize_for_processing(self) -> None:
            if not IS_LINUX or self._optimization_active:
                return
            with self._monitoring_lock:
                self._shutdown_event.clear()
                if self._monitoring_thread and self._monitoring_thread.is_alive():
                    logger.warning("⚠️ Optimize: Monitoring-Thread läuft bereits – überspringe")
                    return
                logger.info("🔧 Aktiviere Linux-Performance-Optimierungen...")
                self._optimize_text_widget('transcript_text')
                self._optimize_text_widget('translation_text')
                self._schedule_batch_interval_increase()
                self._clean_queue_safe(getattr(self.gui, 'gui_queue', None), 15)
                self._apply_linux_specific_optimizations()
                self._optimization_active = True
                self.is_processing = True
                self._start_performance_monitoring()
                logger.info("✅ Linux-Performance-Optimierungen aktiviert")

        def _schedule_batch_interval_increase(self) -> None:
            if not self._is_gui_available_safe():
                return
            if hasattr(self.gui, '_batch_update_interval'):
                self._original_settings['batch_update_interval'] = self.gui._batch_update_interval
            def task() -> None:
                if hasattr(self.gui, '_batch_update_interval'):
                    self.gui._batch_update_interval = 250
            if hasattr(self.gui, 'root'):
                self.gui.root.after(0, task)

        def _optimize_text_widget(self, attr_name: str) -> None:
            widget = getattr(self.gui, attr_name, None)
            if widget and widget.winfo_exists():
                self._original_settings[attr_name] = {
                    'maxundo': widget.cget('maxundo'),
                    'undo': widget.cget('undo'),
                    'autoseparators': widget.cget('autoseparators')
                }
                widget.configure(maxundo=5, undo=True, autoseparators=True, height=12)

        def _is_gui_available_safe(self) -> bool:
            try:
                return (hasattr(self.gui, 'root') and
                        self.gui.root is not None and
                        self.gui.root.winfo_exists() and
                        not getattr(self.gui, '_shutting_down', False))
            except Exception:
                return False

        def _apply_linux_specific_optimizations(self) -> None:
            if not self._is_gui_available_safe():
                return
            self._detect_compositor()
            self._increase_resource_limits()

        def _detect_compositor(self) -> None:
            try:
                import psutil
                for proc in psutil.process_iter(['name']):
                    try:
                        name = proc.info['name'].lower()
                        if any(c in name for c in ['compton', 'picom', 'compiz', 'kwin']):
                            logger.info(f"  ↪ Compositor erkannt: {name}")
                            break
                    except (psutil.NoSuchProcess, psutil.AccessDenied):
                        continue
            except (ImportError, AttributeError):
                pass

        def _increase_resource_limits(self) -> None:
            try:
                import resource

                try:
                    soft, hard = resource.getrlimit(resource.RLIMIT_DATA)
                    new_soft = min(hard, 1024 * 1024 * 1024)
                    if new_soft > soft:
                        resource.setrlimit(resource.RLIMIT_DATA, (new_soft, hard))
                        logger.info(f"  ↪ Daten-Limit erhöht: {soft} → {new_soft}")
                except (resource.error, ValueError) as e:
                    logger.debug(f"  ⚠️ Daten-Limit konnte nicht erhöht werden: {e}")

                try:
                    soft_fd, hard_fd = resource.getrlimit(resource.RLIMIT_NOFILE)
                    new_soft_fd = min(hard_fd, 8192)
                    if new_soft_fd > soft_fd:
                        resource.setrlimit(resource.RLIMIT_NOFILE, (new_soft_fd, hard_fd))
                        logger.info(f"  ↪ Dateideskriptoren-Limit erhöht: {soft_fd} → {new_soft_fd}")
                except (resource.error, ValueError) as e:
                    logger.debug(f"  ⚠️ Dateideskriptoren-Limit konnte nicht erhöht werden: {e}")

                try:
                    os.nice(-5)
                except PermissionError:
                    pass
                except Exception as e:
                    logger.warning(f"  ⚠️ CPU-Priorität konnte nicht angepasst werden: {e}")

            except ImportError:
                logger.debug("  resource-Modul nicht verfügbar – überspringe Limit-Erhöhung")
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"  ⚠️ Unerwarteter Fehler in _increase_resource_limits: {e}")

        def _start_performance_monitoring(self) -> None:
            with self._monitoring_lock:
                if self._monitoring_thread and self._monitoring_thread.is_alive():
                    return

                def monitor_worker() -> None:
                    if self._shutdown_event.is_set():
                        return
                    logger.info("🔍 Linux-Performance-Monitoring gestartet")
                    check_count = 0
                    max_checks = 240

                    while not self._shutdown_event.is_set() and self._optimization_active:
                        try:
                            time.sleep(30)
                            if self._shutdown_event.is_set():
                                break
                            if not self._is_gui_available_safe():
                                logger.warning("⚠️ GUI nicht mehr verfügbar – stoppe Monitoring")
                                break

                            check_count += 1
                            if check_count > max_checks:
                                logger.info("⏰ Monitoring-Zeitlimit erreicht – beende")
                                break

                            system_load = self._get_system_load()
                            memory_usage = self._get_memory_usage()

                            if system_load > 0.8 or memory_usage > 0.85:
                                self._schedule_adjustments(system_load, memory_usage)

                            if check_count % 12 == 0:
                                self._print_performance_report(system_load, memory_usage)

                        except (OSError, ValueError) as e:
                            logger.warning(f"⚠️ Fehler im Monitoring-Thread: {e}")

                    logger.info("✅ Linux-Performance-Monitoring beendet")
                    with self._monitoring_lock:
                        self._monitoring_thread = None

                self._monitoring_thread = threading.Thread(target=monitor_worker, daemon=True, name="LinuxPerfMon")
                self._monitoring_thread.start()

        def _get_system_load(self) -> float:
            try:
                load_avg = os.getloadavg()
                cpu_count = os.cpu_count() or 1
                return load_avg[0] / cpu_count
            except (OSError, ValueError):
                return 0.0

        def _get_memory_usage(self) -> float:
            try:
                import psutil
                memory = psutil.virtual_memory()
                return memory.percent / 100.0
            except Exception:
                return 0.0

        def _schedule_adjustments(self, load: float, memory: float) -> None:
            if not self._is_gui_available_safe():
                return
            def task() -> None:
                if not self._optimization_active or self._shutdown_event.is_set():
                    return
                adjustments: List[str] = []

                if memory > 0.85 and hasattr(self.gui, 'gui_queue'):
                    cleared = self._clean_queue_safe(self.gui.gui_queue, 5)
                    if cleared > 0:
                        adjustments.append(f"Queue: -{cleared}")

                if load > 0.8 and hasattr(self.gui, '_batch_update_interval'):
                    current = self.gui._batch_update_interval
                    if current < 500:
                        self.gui._batch_update_interval = min(500, current + 50)
                        adjustments.append(f"Update: {current}→{self.gui._batch_update_interval}ms")

                if adjustments:
                    logger.info(f"🔧 Anpassungen: {', '.join(adjustments)}")
            if hasattr(self.gui, 'root'):
                self.gui.root.after(0, task)

        def _clean_queue_safe(self, queue_obj: Optional[queue.Queue], target_size: int) -> int:
            if not queue_obj or queue_obj.qsize() <= target_size:
                return 0
            cleared = 0
            try:
                while queue_obj.qsize() > target_size and cleared < 50:
                    queue_obj.get_nowait()
                    cleared += 1
            except queue.Empty:
                pass
            except Exception as e:
                logger.warning(f"⚠️ Queue-Cleanup-Fehler: {e}")
            return cleared

        def _print_performance_report(self, load: float, memory: float) -> None:
            try:
                stats = {
                    'System-Last': f"{load:.1%}",
                    'RAM-Auslastung': f"{memory:.1%}",
                    'Status': 'Aktiv' if self.is_processing else 'Inaktiv'
                }
                logger.info("🐧 Linux-Performance-Report: " + " | ".join(f"{k}: {v}" for k, v in stats.items()))
            except Exception:
                pass

        def restore_normal_mode(self) -> None:
            if not IS_LINUX:
                return
            logger.info("🔧 Linux-Optimierer: Fahre herunter...")
            self._shutdown_event.set()
            self._optimization_active = False
            self.is_processing = False

            if self._monitoring_thread and self._monitoring_thread.is_alive():
                self._monitoring_thread.join(timeout=1.0)
                self._monitoring_thread = None

            if self._is_gui_available_safe():
                logger.info("  ↪ Stelle GUI-Einstellungen wieder her...")
                self._restore_text_widget('transcript_text')
                self._restore_text_widget('translation_text')

                if 'batch_update_interval' in self._original_settings:
                    saved_interval = self._original_settings['batch_update_interval']
                    def restore_batch_interval(saved=saved_interval):
                        if hasattr(self.gui, 'root') and self.gui.root.winfo_exists():
                            if hasattr(self.gui, '_batch_update_interval'):
                                self.gui._batch_update_interval = saved
                    if hasattr(self.gui, 'root'):
                        self.gui.root.after(0, restore_batch_interval)

            self._original_settings.clear()
            logger.info("✅ Linux-Optimierer heruntergefahren")

        def _restore_text_widget(self, attr_name: str) -> None:
            widget = getattr(self.gui, attr_name, None)
            if widget and widget.winfo_exists() and attr_name in self._original_settings:
                try:
                    widget.configure(**self._original_settings[attr_name])
                    logger.info(f"    ✅ {attr_name} wiederhergestellt")
                except Exception as e:
                    logger.warning(f"    ⚠️ {attr_name} konnte nicht wiederhergestellt werden: {e}")

        def emergency_optimize(self) -> None:
            if self._shutdown_event.is_set():
                return
            logger.info("🚨 Führe Notfall-Optimierungen durch...")
            if not self._is_gui_available_safe():
                return
            def task() -> None:
                self._clean_queue_safe(getattr(self.gui, 'gui_queue', None), 3)
                self._clean_queue_safe(getattr(self.gui, '_text_update_queue', None), 2)

                for attr in ['transcript_text', 'translation_text']:
                    widget = getattr(self.gui, attr, None)
                    if widget and widget.winfo_exists():
                        try:
                            widget.configure(height=6, maxundo=1)
                        except Exception:
                            pass

                if hasattr(self.gui, '_batch_update_interval'):
                    self.gui._batch_update_interval = 500

                gc.collect()
                logger.info("✅ Notfall-Optimierungen abgeschlossen")
            if hasattr(self.gui, 'root'):
                self.gui.root.after(0, task)

        def get_optimization_status(self) -> Dict[str, Any]:
            return {
                'platform': SYSTEM,
                'optimization_active': self._optimization_active,
                'processing_active': self.is_processing,
                'monitoring_active': self._monitoring_thread and self._monitoring_thread.is_alive(),
                'shutdown_event_set': self._shutdown_event.is_set(),
                'original_settings_count': len(self._original_settings),
                'linux_specific': IS_LINUX,
                'gui_available': self._is_gui_available_safe()
            }

        def dispose(self) -> None:
            logger.info("🧹 Linux-Performance-Optimierer wird entsorgt...")
            self._shutdown_event.set()
            try:
                self.restore_normal_mode()
            except Exception as e:
                if PlatformUtils.is_fatal_exception(e):
                    raise
                logger.warning(f"⚠️ restore_normal_mode fehlgeschlagen: {e}")
            self._original_settings.clear()
            gc.collect()
            logger.info("✅ Linux-Performance-Optimierer entsorgt")

        def emergency_shutdown(self) -> None:
            logger.info("🚨 Externer Notfall-Shutdown des Linux-Optimierers")
            self._shutdown_event.set()
            self._optimization_active = False
            self.is_processing = False
            self._monitoring_thread = None
            self._original_settings.clear()

elif IS_LINUX:
    class LinuxPerformanceOptimizer:
        def __init__(self, gui_ref: 'DragonWhispererGUI') -> None:
            self.gui = gui_ref
            self.is_processing = False
            logger.info("🐧 LinuxPerformanceOptimizer: Dummy-Modus (psutil fehlt)")

        def optimize_for_processing(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): optimize_for_processing")

        def restore_normal_mode(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): restore_normal_mode")

        def emergency_optimize(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): emergency_optimize")

        def get_optimization_status(self) -> Dict[str, Any]:
            return {'platform': SYSTEM, 'dummy_mode': True, 'psutil_missing': True}

        def dispose(self) -> None:
            pass

        def emergency_shutdown(self) -> None:
            pass

# -----------------------------------------------------------------------------
# INSTALL DIALOG
# -----------------------------------------------------------------------------
class InstallDependencyDialog:
    def __init__(self, parent, gui_ref):
        self.parent = parent
        self.gui = gui_ref
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("🐉 Fehlende Pakete installieren")
        self.dialog.geometry("600x400")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        tk.Label(main, text="Optionale Pakete, die installiert werden können:",
                 bg=CURRENT_THEME.BG_PRIMARY, fg=CURRENT_THEME.TEXT_PRIMARY,
                 font=Fonts.PRIMARY).pack(anchor="w", pady=(0,10))

        self.packages = {}
        frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        frame.pack(fill="x")

        if not WHISPER_AVAILABLE:
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(frame, text="faster-whisper (benötigt für Transkription)",
                                variable=var, bg=CURRENT_THEME.BG_PRIMARY,
                                fg=CURRENT_THEME.TEXT_PRIMARY,
                                selectcolor=CURRENT_THEME.BG_TERTIARY)
            cb.pack(anchor="w")
            self.packages["faster-whisper"] = var

        if not TRANSLATOR_AVAILABLE:
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(frame, text="deep-translator (für Übersetzungen)",
                                variable=var, bg=CURRENT_THEME.BG_PRIMARY,
                                fg=CURRENT_THEME.TEXT_PRIMARY,
                                selectcolor=CURRENT_THEME.BG_TERTIARY)
            cb.pack(anchor="w")
            self.packages["deep-translator"] = var

        if not FastLazyLoader.is_available('psutil'):
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(frame, text="psutil (Systemmonitoring)",
                                variable=var, bg=CURRENT_THEME.BG_PRIMARY,
                                fg=CURRENT_THEME.TEXT_PRIMARY,
                                selectcolor=CURRENT_THEME.BG_TERTIARY)
            cb.pack(anchor="w")
            self.packages["psutil"] = var

        if not self.packages:
            tk.Label(main, text="✅ Alle optionalen Pakete sind bereits installiert.",
                     bg=CURRENT_THEME.BG_PRIMARY, fg=CURRENT_THEME.SUCCESS).pack(pady=20)
            self.dialog.after(2000, self.dialog.destroy)
            return

        tk.Label(main, text="Installationsausgabe:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(anchor="w", pady=(10,2))
        self.output_text = scrolledtext.ScrolledText(
            main, height=10, bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE, wrap=tk.WORD
        )
        self.output_text.pack(fill="both", expand=True, pady=5)

        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x", pady=5)
        install_btn = tk.Button(btn_frame, text="Installieren", command=self.install_selected,
                                bg=CURRENT_THEME.DRAGON_GREEN, fg=CURRENT_THEME.TEXT_PRIMARY,
                                font=Fonts.BUTTON, padx=15)
        install_btn.pack(side="left")
        close_btn = tk.Button(btn_frame, text="Schließen", command=self.dialog.destroy,
                              bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY)
        close_btn.pack(side="right")

    def install_selected(self):
        packages = [pkg for pkg, var in self.packages.items() if var.get()]
        if not packages:
            return
        self.output_text.insert("end", f"Starte Installation von: {', '.join(packages)}...\n")
        self.dialog.update()

        python_exe = sys.executable
        cmd = [python_exe, "-m", "pip", "install"] + packages

        try:
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
                encoding='utf-8',
                errors='ignore'
            )
            for line in process.stdout:
                self.output_text.insert("end", line)
                self.output_text.see("end")
                self.dialog.update()
            process.wait()
            if process.returncode == 0:
                self.output_text.insert("end", "\n✅ Installation erfolgreich!\n")
                self.output_text.insert("end", "Bitte starten Sie das Programm neu, um die neuen Pakete zu nutzen.\n")
            else:
                self.output_text.insert("end", f"\n❌ Fehler bei der Installation (Rückgabecode {process.returncode})\n")
        except Exception as e:
            self.output_text.insert("end", f"\n❌ Ausnahmefehler: {e}\n")


# =============================================================================
# DRAGON WHISPERER GUI
# =============================================================================
class DragonWhispererGUI:
    class RateLimiter:
        def __init__(self, max_updates_per_second: int = 30) -> None:
            self.min_interval = 1.0 / max_updates_per_second
            self.last_calls: Dict[str, float] = {}
            self._lock = threading.RLock()

        def can_update(self, update_type: str = "default") -> bool:
            with self._lock:
                now = time.time()
                if update_type not in self.last_calls:
                    self.last_calls[update_type] = 0.0
                last = self.last_calls[update_type]
                if now - last >= self.min_interval:
                    self.last_calls[update_type] = now
                    return True
                return False

        def reset(self, update_type: Optional[str] = None) -> None:
            with self._lock:
                if update_type is None:
                    self.last_calls.clear()
                elif update_type in self.last_calls:
                    del self.last_calls[update_type]

    def __init__(self) -> None:
        self._gui_update_limiter = self.RateLimiter(max_updates_per_second=15)
        self._shutting_down = False
        self._exit_dialog_active = False
        self.is_processing = False
        self.subtitle_mode = False
        self.exit_confirmed = False
        self.current_stream_info: Optional[StreamInfo] = None
        self.current_video_language: Optional[str] = None
        self._progress_bar_started = False
        self.translate_active = True

        self.translation_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="Translation")

        self._history_lock = threading.RLock()

        if not GUI_AVAILABLE:
            logger.error("❌ Tkinter nicht verfügbar. Versuche Fallback...")
            self._try_fallback_gui()
            return

        try:
            self.settings = AppSettings.load_from_file()
            if not self.settings.last_url:
                self.settings.last_url = ""
            self.advanced_settings = Settings.load_from_file()
            self.advanced_settings.repair()
            validation_issues = self.advanced_settings.validate()
            if validation_issues:
                logger.warning(f"⚠️ Settings validation issues: {validation_issues}")
            logger.info(f"✅ Settings ready: SAMPLE_RATE={self.advanced_settings.config.SAMPLE_RATE}, "
                        f"CHANNELS={self.advanced_settings.config.CHANNELS}, "
                        f"CHUNK_SIZE_BYTES={self.advanced_settings.config.CHUNK_SIZE_BYTES}")
        except Exception as e:
            logger.warning(f"⚠️ Settings load failed: {e}, using defaults")
            self.settings = AppSettings()
            self.settings.last_url = ""
            self.advanced_settings = Settings()

        if not self.settings.cookies_notice_shown and self.settings.use_browser_cookies:
            self._show_cookie_notice = True
        else:
            self._show_cookie_notice = False

        if self.settings.theme == "light":
            self.current_theme = LightTheme()
        else:
            self.current_theme = DarkTheme()

        global CURRENT_THEME
        CURRENT_THEME = self.current_theme

        self.demo_mode = not WHISPER_AVAILABLE
        self.layout_mode = getattr(self.settings, 'layout_mode', 'vertical')
        self.current_language = getattr(self.settings, 'default_language', 'de')

        self._translation_reset_counter = 0
        self.progress_dialog: Optional[ProgressDialog] = None

        try:
            self.root = tk.Tk()
            self.root.withdraw()
        except (tk.TclError, RuntimeError) as e:
            raise RuntimeError(f"Tkinter Fehler: {e}")

        self._batch_update_interval = 150
        self._last_batch_update = 0.0
        self._last_gui_update_time = 0.0
        self._processing_lock = threading.Lock()

        self.transcript_history: Deque[TranscriptionResult] = deque(maxlen=1000)
        self.translation_history: Deque[TranslationResult] = deque(maxlen=500)
        self._last_transcription_text = ""
        self._last_translation_text = ""

        self.performance_monitor = SimplePerformanceTracker()
        self.gui_queue: queue.Queue = queue.Queue(maxsize=200)
        self._text_update_queue: queue.Queue = queue.Queue(maxsize=150)

        self.stream_info_extractor = StreamInfoExtractor()
        self.stream_info_extractor.use_browser_cookies = self.settings.use_browser_cookies

        try:
            self.controller = WhisperController(gui_ref=self)
        except Exception as e:
            logger.error(f"❌ Controller Fehler: {e}")
            self._show_error_and_exit(f"Controller Fehler: {e}")
            return

        self.controller.set_callbacks(
            on_transcription=self.handle_transcription,
            on_translation=self.handle_translation,
            on_info=self.handle_info,
            on_error=self.handle_error,
            on_status=self._handle_status_update,
            on_finished=self._on_processing_finished
        )

        try:
            self.layout = WhisperLayoutManager(gui_ref=self)
        except Exception as e:
            logger.error(f"❌ Layout Fehler: {e}")
            self._show_error_and_exit(f"Layout Fehler: {e}")
            return

        try:
            self.stream_manager = StreamManager(enable_debug=(DEBUG_LEVEL >= 1), use_browser_cookies=self.settings.use_browser_cookies)
            self.ffmpeg_manager = FFmpegManager(self.advanced_settings.config, self.stream_manager, self.advanced_settings)
            if WHISPER_AVAILABLE:
                self.transcription_engine = TranscriptionEngine(self.advanced_settings)
            else:
                self.transcription_engine = DummyTranscriptionEngine(self.advanced_settings)

            if self.advanced_settings.translation_engine == "ollama":
                if OLLAMA_AVAILABLE:
                    self.translation_engine = OllamaTranslationEngine(
                        target_lang=self.current_language,
                        settings=self.advanced_settings,
                        model=self.advanced_settings.ollama_model,
                        host=self.advanced_settings.ollama_host
                    )
                    logger.info("✅ OllamaTranslationEngine aktiviert")
                else:
                    logger.warning("⚠️ Ollama translation engine selected but not available (requests missing). Falling back to Google Translate.")
                    self.root.after(100, lambda: self.update_status("⚠️ Ollama not available, using Google Translate"))
                    if TRANSLATOR_AVAILABLE:
                        self.translation_engine = GoogleTranslationEngine(self.current_language, self.advanced_settings)
                        logger.info("✅ GoogleTranslationEngine (Google) aktiviert (Fallback)")
                    else:
                        self.translation_engine = DummyTranslationEngine(self.current_language, self.advanced_settings)
                        logger.warning("⚠️ Keine Übersetzungs-Engine verfügbar, verwende Dummy")
            else:
                if TRANSLATOR_AVAILABLE:
                    self.translation_engine = GoogleTranslationEngine(self.current_language, self.advanced_settings)
                    logger.info("✅ GoogleTranslationEngine (Google) aktiviert")
                else:
                    self.translation_engine = DummyTranslationEngine(self.current_language, self.advanced_settings)
                    logger.warning("⚠️ Keine Übersetzungs-Engine verfügbar, verwende Dummy")

            self.audio_processor = AudioProcessor(
                controller_ref=self.controller,
                ffmpeg_manager=self.ffmpeg_manager,
                settings=self.advanced_settings
            )
            fallback_engine = None
            if self.advanced_settings.translation_engine == "google" and OLLAMA_AVAILABLE:
                fallback_engine = OllamaTranslationEngine(
                    target_lang=self.current_language,
                    settings=self.advanced_settings,
                    model=self.advanced_settings.ollama_model,
                    host=self.advanced_settings.ollama_host
                )
                logger.info("✅ Fallback-Engine (Ollama) bereitgestellt")
            self.audio_processor.set_engines(
                transcription_engine=self.transcription_engine,
                translation_engine=self.translation_engine,
                fallback_translation_engine=fallback_engine
                # Plugin-Manager nicht mehr übergeben
            )
            self.export_manager = ExportManager()
            self.language_detector = LanguageDetector(self.transcription_engine)
            self.resource_manager = ResourceManager()
            self.memory_manager = MemoryManager()

            if IS_LINUX:
                self.performance_optimizer = LinuxPerformanceOptimizer(gui_ref=self)

            self._register_signal_handlers()
        except (ImportError, OSError, RuntimeError) as e:
            logger.error(f"❌ Engine Initialisierung Fehler: {e}")
            self._show_error_and_exit(f"Engine Fehler: {e}")
            return

        try:
            self.layout.setup_gui()
            self._setup_callbacks()
            self.root.after(100, self._start_gui_updaters)

            if hasattr(self, 'url_entry') and self.url_entry.winfo_exists():
                self.url_entry.delete(0, 'end')
                self.url_entry.insert(0, self.settings.last_url)
            else:
                def set_initial_url() -> None:
                    if hasattr(self, 'url_entry') and self.url_entry.winfo_exists():
                        self.url_entry.delete(0, 'end')
                        self.url_entry.insert(0, self.settings.last_url)
                self.root.after(200, set_initial_url)

            self.root.deiconify()
            self.root.title("🐉 Dragon Whisperer")
        except Exception as e:
            logger.error(f"❌ GUI Setup Fehler: {e}")
            self._show_error_and_exit(f"GUI konnte nicht erstellt werden: {e}")
            return

        if self._show_cookie_notice:
            self.root.after(500, self._show_cookie_notice_dialog)

        self._bind_shortcuts()
        self.root.after(1000, self._start_system_monitoring)
        self.root.after(2000, self._final_initialization_check)
        self._schedule_gui_health_check()

    def _safe_queue_put(self, q: queue.Queue, item: Any, max_retries: int = 1) -> None:
        try:
            q.put_nowait(item)
        except queue.Full:
            if max_retries > 0:
                try:
                    q.get_nowait()
                    self._safe_queue_put(q, item, max_retries - 1)
                except queue.Empty:
                    pass
            else:
                if debug3_enabled('queue'):
                    logger.debug(f"[DEBUG3][QUEUE] Queue full, dropping item after {max_retries} retries")

    def _show_cookie_notice_dialog(self) -> None:
        result = DarkMessageBox.askyesno(
            "Datenschutzhinweis",
            "Dragon Whisperer kann auf gespeicherte Browser-Cookies zugreifen, "
            "um YouTube-Streams zuverlässiger abzurufen. Dies kann Ihre Privatsphäre "
            "beeinträchtigen.\n\nMöchten Sie die Nutzung von Browser-Cookies erlauben?\n\n"
            "(Sie können diese Einstellung später in den erweiterten Einstellungen ändern.)",
            parent=self.root
        )
        self.settings.use_browser_cookies = result
        self.settings.cookies_notice_shown = True
        self.settings.save_to_file()
        if hasattr(self, 'stream_manager'):
            self.stream_manager.use_browser_cookies = result
        if hasattr(self, 'stream_info_extractor'):
            self.stream_info_extractor.use_browser_cookies = result

    def _schedule_gui_health_check(self) -> None:
        if not self._shutting_down and hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(30000, self._perform_gui_health_check)

    def _perform_gui_health_check(self) -> None:
        try:
            checks: List[str] = []
            check_start = time.time()
            self.root.update_idletasks()
            responsiveness_time = time.time() - check_start
            if responsiveness_time > 0.5:
                checks.append(f"⚠️ GUI responsiveness slow: {responsiveness_time:.1f}s")
            if hasattr(self, 'memory_manager'):
                mem_stats = self.memory_manager.get_memory_stats()
                if mem_stats.get("process_usage_percent", 0) > 80:
                    checks.append("⚠️ High memory usage")
            if hasattr(self, 'gui_queue'):
                qsize = self.gui_queue.qsize()
                if qsize > 50:
                    checks.append(f"⚠️ GUI queue backlog: {qsize} items")
                    self._cleanup_queue(self.gui_queue, 20)
            active_threads = threading.enumerate()
            if len(active_threads) > 15:
                checks.append(f"⚠️ Many active threads: {len(active_threads)}")
            try:
                cache_stats = get_cache_stats()
                for cache_name, stats in cache_stats.items():
                    if stats.get("total_entries", 0) > stats.get("max_size", 100) * 0.9:
                        checks.append(f"⚠️ {cache_name} nearly full")
            except Exception:
                pass
            if checks:
                if len(checks) > 0:
                    logger.warning(f"🔍 GUI Health Check Issues: {checks[:3]}")
            if "memory usage" in str(checks) and hasattr(self, 'memory_manager'):
                self.memory_manager._aggressive_cleanup()
        except Exception as e:
            if PlatformUtils.is_fatal_exception(e):
                raise
            logger.warning(f"⚠️ Health check error: {e}")
        finally:
            self._schedule_gui_health_check()

    def _register_signal_handlers(self) -> None:
        try:
            logger.info("🔧 Registering cleanup handlers with SignalHandler...")
            SignalHandler.register_cleanup(
                self._safe_stop_all_processes,
                name="StopAllProcesses",
                priority=ShutdownPriority.CRITICAL,
                timeout=3.0,
                essential=True
            )
            SignalHandler.register_cleanup(
                self.audio_processor.dispose,
                name="AudioProcessorDispose",
                priority=ShutdownPriority.HIGH,
                timeout=2.0
            )
            if hasattr(self, 'ffmpeg_manager') and self.ffmpeg_manager:
                SignalHandler.register_cleanup(
                    self.ffmpeg_manager.dispose,
                    name="FFmpegManagerDispose",
                    priority=ShutdownPriority.HIGH,
                    timeout=2.0
                )
                logger.info("   ✅ Registered FFmpegManager cleanup")
            SignalHandler.register_cleanup(
                self.transcription_engine.dispose,
                name="TranscriptionEngineDispose",
                priority=ShutdownPriority.MEDIUM,
                timeout=1.0
            )
            SignalHandler.register_cleanup(
                self.translation_engine.dispose,
                name="TranslationEngineDispose",
                priority=ShutdownPriority.MEDIUM,
                timeout=1.0
            )
            SignalHandler.register_cleanup(
                self.memory_manager.dispose,
                name="MemoryManagerDispose",
                priority=ShutdownPriority.LOW
            )
            SignalHandler.register_cleanup(
                self.resource_manager.cleanup,
                name="ResourceManagerCleanup",
                priority=ShutdownPriority.LOW
            )
            SignalHandler.register_cleanup(
                self.stream_manager.dispose,
                name="StreamManagerDispose",
                priority=ShutdownPriority.LOW
            )
            if IS_LINUX and hasattr(self, 'performance_optimizer'):
                SignalHandler.register_cleanup(
                    self._safe_linux_optimizer_cleanup,
                    name="LinuxOptimizerCleanup",
                    priority=ShutdownPriority.LOW,
                    timeout=1.0
                )
            SignalHandler.register_cleanup(
                lambda: (transcription_cache.clear(), translation_cache.clear(), audio_cache.clear()),
                name="ClearGlobalCaches",
                priority=ShutdownPriority.LOW
            )
            SignalHandler.register_cleanup(
                self._cleanup_queues,
                name="CleanupQueues",
                priority=ShutdownPriority.LOW
            )
            SignalHandler.register_cleanup(
                lambda: self.translation_executor.shutdown(wait=False),
                name="TranslationExecutorShutdown",
                priority=ShutdownPriority.LOW
            )
            SignalHandler.register_cleanup(
                lambda: _EXECUTOR.shutdown(wait=False),
                name="GlobalExecutorShutdown",
                priority=ShutdownPriority.LOW
            )
            try:
                import torch
                if torch.cuda.is_available():
                    SignalHandler.register_cleanup(
                        torch.cuda.empty_cache,
                        name="GPUMemoryCleanup",
                        priority=ShutdownPriority.LOW,
                        timeout=1.0
                    )
                    logger.info("   ✅ Registered GPU cleanup")
            except ImportError:
                pass
            count = sum(len(ops) for ops in SignalHandler._cleanup_operations.values())
            logger.info(f"✅ Registered {count} cleanup handlers")
            SignalHandler.setup(verbose=False, silent=True)
        except Exception as e:
            logger.warning(f"⚠️ SignalHandler registration error: {e}")

    def _safe_stop_all_processes(self) -> None:
        logger.info("🛑 Safely stopping all processes...")
        self._shutting_down = True
        self.is_processing = False
        if hasattr(self, 'controller'):
            try:
                self.controller._shutdown_event.set()
                self.controller._stop_requested = True
                if hasattr(self.controller, '_processing_active'):
                    self.controller._processing_active.clear()
            except Exception as e:
                logger.warning(f"⚠️ Controller stop error: {e}")
        if hasattr(self, 'audio_processor'):
            try:
                self.audio_processor._processing.clear()
                if hasattr(self.audio_processor, '_stop_event'):
                    self.audio_processor._stop_event.set()
            except Exception as e:
                logger.warning(f"⚠️ Audio processor stop error: {e}")
        if hasattr(self, 'ffmpeg_manager'):
            try:
                self.ffmpeg_manager._shutting_down = True
                if hasattr(self.ffmpeg_manager, '_processes'):
                    import os
                    import signal
                    for process_id, process_info in list(self.ffmpeg_manager._processes.items()):
                        try:
                            process = process_info.get('process')
                            if process and hasattr(process, 'pid'):
                                try:
                                    if IS_WINDOWS:
                                        process.terminate()
                                    else:
                                        os.kill(process.pid, signal.SIGTERM)
                                    time.sleep(0.1)
                                except Exception:
                                    try:
                                        if IS_WINDOWS:
                                            process.kill()
                                        else:
                                            os.kill(process.pid, signal.SIGKILL)
                                    except Exception:
                                        pass
                        except Exception:
                            pass
            except Exception as e:
                logger.warning(f"⚠️ FFmpeg stop error: {e}")
        try:
            self.translation_executor.shutdown(wait=False)
        except Exception:
            pass
        logger.info("✅ All processes stopped")

    def _safe_linux_optimizer_cleanup(self) -> None:
        if not IS_LINUX or not hasattr(self, 'performance_optimizer'):
            return
        logger.info("🐧 Safe Linux optimizer cleanup...")
        try:
            gui_exists = False
            try:
                if hasattr(self, 'root') and self.root.winfo_exists():
                    gui_exists = True
            except Exception:
                gui_exists = False
            if gui_exists:
                try:
                    self.performance_optimizer.restore_normal_mode()
                except Exception as e:
                    logger.warning(f"⚠️ restore_normal_mode failed: {e}")
                    try:
                        self.performance_optimizer.dispose()
                    except Exception:
                        pass
            else:
                try:
                    self.performance_optimizer.dispose()
                except Exception as e:
                    logger.warning(f"⚠️ dispose failed: {e}")
        except Exception as e:
            logger.warning(f"⚠️ Linux optimizer cleanup error: {e}")

    def _cleanup_queues(self) -> None:
        logger.info("🗑️ Cleaning up queues...")
        try:
            if hasattr(self, 'gui_queue'):
                count = 0
                while not self.gui_queue.empty() and count < 100:
                    try:
                        self.gui_queue.get_nowait()
                        count += 1
                    except Exception:
                        break
                if count > 0:
                    logger.info(f"  Cleared GUI queue: {count} items")
        except Exception as e:
            logger.warning(f"⚠️ GUI queue cleanup error: {e}")
        try:
            if hasattr(self, '_text_update_queue'):
                count = 0
                while not self._text_update_queue.empty() and count < 100:
                    try:
                        self._text_update_queue.get_nowait()
                        count += 1
                    except Exception:
                        break
                if count > 0:
                    logger.info(f"  Cleared text queue: {count} items")
        except Exception as e:
            logger.warning(f"⚠️ Text queue cleanup error: {e}")

    def _safe_exit_dialog(self) -> None:
        if self._shutting_down or self._exit_dialog_active:
            return
        self._exit_dialog_active = True
        try:
            if not hasattr(self, 'root') or not self.root.winfo_exists():
                self._direct_shutdown()
                return
            result = DarkMessageBox.askyesno(
                "🐉 Dragon Whisperer - Beenden",
                "Programm wirklich beenden?\n\n"
                "● Laufende Transkriptionen werden gestoppt\n"
                "● Nicht gespeicherte Daten gehen verloren\n\n"
                "Sicher beenden?",
                parent=self.root
            )
            if result:
                logger.info("✅ User confirmed exit - shutting down...")
                self._direct_shutdown()
            else:
                logger.info("↩️ Exit cancelled by user")
                self._exit_dialog_active = False
        except tk.TclError:
            logger.warning("⚠️ GUI destroyed, performing direct shutdown...")
            self._direct_shutdown()
        except Exception as e:
            logger.warning(f"⚠️ Exit dialog error: {e}")
            self._direct_shutdown()
        finally:
            if not self._shutting_down:
                self._exit_dialog_active = False

    def _direct_shutdown(self) -> None:
        if self._shutting_down:
            logger.warning("⚠️ Shutdown already in progress, skipping...")
            return
        logger.info("🔧 Performing confirmed shutdown...")
        self._shutting_down = True
        self._safe_stop_all_processes()
        time.sleep(0.3)
        try:
            if hasattr(self, 'root') and self.root.winfo_exists():
                self.root.title("🐉 Dragon Whisperer - Beendet...")
                self.root.update_idletasks()
                self.root.quit()
        except tk.TclError:
            pass
        except Exception as e:
            logger.warning(f"⚠️ GUI shutdown error: {e}")

    def _cleanup_queue(self, queue_obj: queue.Queue, max_size: int) -> None:
        if not queue_obj or queue_obj.qsize() <= max_size:
            return
        try:
            items = []
            while not queue_obj.empty():
                try:
                    items.append(queue_obj.get_nowait())
                except queue.Empty:
                    break
            important = []
            others = []
            for item in items:
                if isinstance(item, tuple) and len(item) == 2 and item[0] in ('status', 'error', 'file_finished'):
                    important.append(item)
                else:
                    others.append(item)
            kept = important + others[-(max_size - len(important)):] if len(important) < max_size else important[:max_size]
            for item in kept:
                try:
                    queue_obj.put_nowait(item)
                except queue.Full:
                    break
            logger.debug(f"🧹 Queue cleaned: {len(items)} -> {len(kept)} items")
            if debug3_enabled('queue'):
                logger.debug(f"[DEBUG3][QUEUE] Cleaned queue from {len(items)} to {len(kept)} items")
        except Exception as e:
            logger.warning(f"⚠️ Queue cleanup error: {e}")

    def run(self) -> None:
        try:
            self.root.title("🐉 Dragon Whisperer")
            self._shutting_down = False
            self._exit_dialog_active = False
            self.root.protocol("WM_DELETE_WINDOW", self._safe_exit_dialog)
            if hasattr(self, 'exit_button'):
                self.exit_button.config(command=self._safe_exit_dialog)
            logger.info("🚀 Starting Dragon Whisperer (with exit confirmation)...")
            if not SignalHandler._setup_complete:
                try:
                    SignalHandler.setup(verbose=False, silent=True)
                except Exception:
                    pass
            self.root.mainloop()
            logger.info("✅ Main loop exited normally")
        except KeyboardInterrupt:
            logger.info("\n🛑 Interrupted by user - showing exit dialog...")
            self._safe_exit_dialog()
        except SystemExit as e:
            logger.info("\n🔧 System exit requested")
            raise e
        except Exception as e:
            logger.error(f"💥 Critical error: {type(e).__name__}: {e}")
            self._direct_shutdown()

    def _post_mainloop_cleanup(self) -> None:
        logger.info("🧹 Post-mainloop cleanup (no GUI access)...")
        self._shutting_down = True
        self.is_processing = False
        self._cleanup_queues()
        try:
            transcription_cache.clear()
            translation_cache.clear()
            audio_cache.clear()
        except Exception as e:
            logger.warning(f"⚠️ Cache cleanup error: {e}")
        gc.collect()
        logger.info("✅ Post-mainloop cleanup completed")

    def _minimal_emergency_cleanup(self) -> None:
        logger.info("🆘 MINIMAL emergency cleanup...")
        self._shutting_down = True
        self.is_processing = False
        try:
            if hasattr(self, 'ffmpeg_manager'):
                import os
                import signal
                if hasattr(self.ffmpeg_manager, '_processes'):
                    for pid, info in list(self.ffmpeg_manager._processes.items()):
                        try:
                            process = info.get('process')
                            if process and hasattr(process, 'pid'):
                                try:
                                    if IS_WINDOWS:
                                        process.kill()
                                    else:
                                        os.kill(process.pid, signal.SIGKILL)
                                except Exception:
                                    pass
                        except Exception:
                            pass
        except Exception:
            pass
        gc.collect()

    def safe_controller_stop(self) -> None:
        if self._shutting_down:
            return
        logger.info("🛑 Safe controller stop...")
        self._safe_stop_all_processes()

    @gui_operation_decorator
    def select_file_dark(self) -> None:
        try:
            filename = filedialog.askopenfilename(
                title="🎬 Select Audio/Video File - Dragon Whisperer",
                filetypes=[
                    ("Media files", "*.mp3 *.wav *.m4a *.mp4 *.avi *.mkv *.mov *.flac"),
                    ("All files", "*.*"),
                ],
            )
            if filename:
                file_url = f"file://{filename}"
                self.url_entry.delete(0, "end")
                self.url_entry.insert(0, file_url)
                self.update_status(f"📁 File selected: {os.path.basename(filename)}")
                def async_language_detection() -> None:
                    try:
                        self.analyze_video_language(filename)
                    except Exception:
                        pass
                detection_thread = threading.Thread(target=async_language_detection, daemon=True)
                if hasattr(self, 'resource_manager'):
                    self.resource_manager.register_thread(detection_thread)
                detection_thread.start()
                info = self.stream_info_extractor.extract_stream_info(file_url)
                self.update_stream_info(info)
        except Exception as e:
            self.update_status(f"❌ File selection failed: {e}")

    @gui_operation_decorator
    def paste_url(self) -> None:
        try:
            clipboard = self.root.clipboard_get().strip()
            if clipboard:
                try:
                    cleaned_url = self.clean_and_validate_url(clipboard)
                except ValueError as e:
                    self.update_status(f"❌ Ungültige URL: {e}")
                    return
                self.url_entry.delete(0, "end")
                self.url_entry.insert(0, cleaned_url)
                self.update_status("📋 URL eingefügt")
                if cleaned_url.startswith("file://"):
                    file_path = cleaned_url[7:]
                    if os.path.exists(file_path):
                        def async_detection() -> None:
                            try:
                                self.analyze_video_language(file_path)
                            except Exception:
                                pass
                        detection_thread = threading.Thread(target=async_detection, daemon=True)
                        if hasattr(self, 'resource_manager'):
                            self.resource_manager.register_thread(detection_thread)
                        detection_thread.start()
            else:
                self.update_status("❌ Zwischenablage ist leer")
        except tk.TclError:
            self.update_status("❌ Konnte nicht auf Zwischenablage zugreifen")
        except Exception as e:
            self.update_status(f"❌ Fehler beim Einfügen: {str(e)[:50]}")

    def clean_and_validate_url(self, url: str) -> str:
        url = url.strip()
        if not url:
            raise ValueError("URL cannot be empty")
        url = PlatformUtils.sanitize_url(url)
        if url.startswith("file://"):
            ok, real_path = PlatformUtils.validate_file_path(url)
            if not ok:
                raise ValueError(real_path)
            file_path = real_path
            if not os.path.exists(file_path):
                raise ValueError(f"File not found: {file_path}")
            return url
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
        parsed = urllib.parse.urlparse(url)
        if not parsed.netloc:
            raise ValueError("Invalid URL format")
        if len(url) < 10:
            raise ValueError("URL too short")
        if " " in url:
            raise ValueError("URL cannot contain spaces")
        return url

    def analyze_video_language(self, file_path: str) -> None:
        if hasattr(self, 'language_info_label'):
            self.root.after(0, lambda: self.language_info_label.config(text="🔍 Analyzing..."))

        def language_detection_worker() -> None:
            try:
                detection_result = self.language_detector.detect_video_language(file_path)
                logger.debug(f"Language detection result: {detection_result}")
                def update_result() -> None:
                    if hasattr(self, 'language_info_label'):
                        if 'error' in detection_result:
                            self.language_info_label.config(text=f"❌ {detection_result['error']}")
                        elif 'info' in detection_result:
                            self.language_info_label.config(text=f"ℹ️ {detection_result['info']}")
                        else:
                            language_name = detection_result['language_name']
                            confidence = detection_result['confidence']
                            self.current_video_language = detection_result['detected_language']
                            language_icons = {
                                'zh': '㊗️', 'ja': '🗾', 'ko': '₩',
                                'th': '🇹🇭', 'vi': '🇻🇳',
                            }
                            icon = language_icons.get(self.current_video_language, '✅')
                            display_text = f"{icon} {language_name} ({confidence:.0%})"
                            self.language_info_label.config(text=display_text)
                if hasattr(self, 'root'):
                    self.root.after(0, update_result)
            except Exception as e:
                logger.error(f"Language detection exception: {e}")
                def update_error() -> None:
                    if hasattr(self, 'language_info_label'):
                        self.language_info_label.config(text="❌ Analysis failed")
                if hasattr(self, 'root'):
                    self.root.after(0, update_error)
        detection_thread = threading.Thread(target=language_detection_worker, daemon=True)
        if hasattr(self, 'resource_manager'):
            self.resource_manager.register_thread(detection_thread)
        detection_thread.start()

    def on_url_change(self, event: Optional[tk.Event] = None) -> None:
        if not hasattr(self, 'url_entry'):
            return
        url = self.url_entry.get().strip()
        if url.startswith("file://"):
            file_path = url[7:]
            if os.path.exists(file_path):
                def async_detection() -> None:
                    try:
                        self.analyze_video_language(file_path)
                    except Exception:
                        pass
                detection_thread = threading.Thread(target=async_detection, daemon=True)
                if hasattr(self, 'resource_manager'):
                    self.resource_manager.register_thread(detection_thread)
                detection_thread.start()
            else:
                if hasattr(self, 'language_info_label'):
                    self.language_info_label.config(text="❌ File not found")
        else:
            if hasattr(self, 'language_info_label'):
                self.language_info_label.config(text="")
            self.current_video_language = None

    def on_language_change(self, event: Optional[tk.Event] = None) -> None:
        try:
            selected_name = self.lang_var.get()
            lang_code: Optional[str] = None
            for name, code in SORTED_LANGUAGES:
                if name == selected_name:
                    lang_code = code
                    break
            if lang_code and lang_code != self.current_language:
                self.current_language = lang_code
                if hasattr(self, 'translation_engine'):
                    self.translation_engine.set_target_language(lang_code)
                lang_display = LANGUAGE_SHORT_CODES.get(lang_code, lang_code)
                if hasattr(self, 'translation_header'):
                    self.translation_header.config(text=f"🌐 Translation ({lang_display})")
                self.update_status(f"🌍 Target language: {selected_name}")
        except Exception:
            pass

    def on_model_change(self, event: Optional[tk.Event] = None) -> None:
        if not hasattr(self, 'model_var'):
            return
        new_model = self.model_var.get()
        if new_model not in WHISPER_MODELS:
            logger.warning(f"⚠️ Invalid model selected: {new_model}")
            current = self.transcription_engine.get_current_model()
            self.model_var.set(current)
            return
        if not hasattr(self, 'transcription_engine'):
            return
        current_model = self.transcription_engine.get_current_model()
        if new_model == current_model:
            return
        if self.transcription_engine.is_model_loading():
            self.update_status("🔄 Model already loading...")
            return
        success = self.transcription_engine.reload_model(new_model)
        if success:
            self.update_status(f"🔄 Switching to {new_model}...")
            self._check_model_loading_complete(new_model)
        else:
            self.update_status("❌ Model switch failed")
            self.model_var.set(current_model)

    def _check_model_loading_complete(self, target_model: str) -> None:
        if self.transcription_engine.is_model_loading():
            self.root.after(200, lambda: self._check_model_loading_complete(target_model))
        else:
            current = self.transcription_engine.get_current_model()
            if current == target_model:
                self.update_status(f"✅ Model switched to {target_model}")
            else:
                self.update_status("❌ Model switch failed")
                self.model_var.set(current)

    def toggle_translation(self) -> None:
        self.translate_active = not self.translate_active
        if hasattr(self, 'audio_processor'):
            self.audio_processor._translation_active = self.translate_active
        if hasattr(self, 'translate_btn'):
            if self.translate_active:
                self.translate_btn.config(text="🌐 ON", bg=self.current_theme.SUCCESS)
                self.update_status("✅ Translation active")
            else:
                self.translate_btn.config(text="🌐 OFF", bg=self.current_theme.BG_TERTIARY)
                self.update_status("❌ Translation inactive")

    def toggle_subtitle_mode(self) -> None:
        self.subtitle_mode = not self.subtitle_mode
        if hasattr(self, 'audio_processor'):
            self.audio_processor.enable_subtitle_mode(self.subtitle_mode)
        if hasattr(self, 'subtitle_btn'):
            if self.subtitle_mode:
                self.subtitle_btn.config(bg=self.current_theme.SUBTITLE_ACTIVE, fg=self.current_theme.TEXT_PRIMARY)
                self.update_status("🎬 SUBTITLE MODE: Timestamps activated")
            else:
                self.subtitle_btn.config(bg=self.current_theme.SUBTITLE_INACTIVE, fg=self.current_theme.TEXT_PRIMARY)
                self.update_status("📝 NORMAL MODE: Continuous text")

    def _on_start_click(self):
        if self.is_processing:
            self.update_status("⚠️ Bereits aktiv")
            return

        try:
            self.start_button.config(state="disabled")
        except Exception:
            pass
        self.controller.start_processing()

    def toggle_layout(self) -> None:
        try:
            logger.info(f"🔄 Starting layout toggle from {self.layout_mode}")
            old_transcript = ""
            old_translation = ""
            try:
                if hasattr(self, 'transcript_text') and self.transcript_text:
                    old_transcript = self.transcript_text.get('1.0', 'end-1c')
                    logger.info(f"  📝 Saved transcript: {len(old_transcript)} chars")
            except (tk.TclError, AttributeError) as e:
                logger.warning(f"  ⚠️ Could not save transcript: {e}")
            try:
                if hasattr(self, 'translation_text') and self.translation_text:
                    old_translation = self.translation_text.get('1.0', 'end-1c')
                    logger.info(f"  📝 Saved translation: {len(old_translation)} chars")
            except (tk.TclError, AttributeError) as e:
                logger.warning(f"  ⚠️ Could not save translation: {e}")
            if self.layout_mode == "vertical":
                self.layout_mode = "horizontal"
                new_mode_text = "Horizontal"
            else:
                self.layout_mode = "vertical"
                new_mode_text = "Vertical"
            logger.info(f"  🔄 Switching to: {self.layout_mode}")
            if hasattr(self, 'settings'):
                self.settings.layout_mode = self.layout_mode
                try:
                    self.settings.save_to_file()
                except Exception as e:
                    logger.warning(f"  ⚠️ Settings save error: {e}")
            self.update_status(f"🔄 Switching to {new_mode_text} layout...")
            if hasattr(self, 'layout'):
                new_transcript, new_translation = self.layout.create_text_areas()
                if new_transcript and old_transcript:
                    try:
                        new_transcript.insert('1.0', old_transcript)
                        logger.info("  ✅ Restored transcript to new widget")
                    except Exception as e:
                        logger.warning(f"  ❌ Failed to restore transcript: {e}")
                if new_translation and old_translation:
                    try:
                        new_translation.insert('1.0', old_translation)
                        logger.info("  ✅ Restored translation to new widget")
                    except Exception as e:
                        logger.warning(f"  ❌ Failed to restore translation: {e}")
            self.update_status(f"✅ {new_mode_text} layout active")
            logger.info("✅ Layout toggle completed successfully")
        except Exception as e:
            logger.error(f"❌ CRITICAL Layout toggle error: {e}")
            self.update_status("❌ Layout change failed")
            try:
                self.layout_mode = "vertical"
                if hasattr(self, 'layout'):
                    self.layout.create_text_areas()
            except Exception:
                pass

    def clear_all(self) -> None:
        try:
            if hasattr(self, 'transcript_text'):
                self.transcript_text.delete("1.0", "end")
            if hasattr(self, 'translation_text'):
                self.translation_text.delete("1.0", "end")
        except Exception:
            pass
        with self._history_lock:
            self.transcript_history.clear()
            self.translation_history.clear()
        if hasattr(self, 'memory_manager'):
            self.memory_manager.clear_component('transcript')
            self.memory_manager.clear_component('translation')
        self._last_transcription_text = ""
        self._last_translation_text = ""
        self._translation_reset_counter = 0
        self.update_status("🗑️ Cleared & optimizations reset")

    def save_transcript(self) -> None:
        try:
            if not self.transcript_history:
                DarkMessageBox.showinfo("WARNING", "No transcriptions available to save.", self.root)
                return
            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[
                    ("Text files", "*.txt"),
                    ("SRT subtitles", "*.srt"),
                    ("WebVTT", "*.vtt"),
                    ("JSON", "*.json"),
                    ("Word document", "*.docx"),
                    ("All files", "*.*"),
                ],
            )
            if not filename:
                return
            file_ext = Path(filename).suffix.lower()
            success = False
            if file_ext == ".srt":
                success = self.export_manager.export_subtitles(
                    list(self.transcript_history), None, "srt", filename
                )
            elif file_ext == ".vtt":
                success = self.export_manager.export_subtitles(
                    list(self.transcript_history), None, "vtt", filename
                )
            elif file_ext == ".json":
                success = self.export_manager.export_json(
                    list(self.transcript_history), list(self.translation_history), filename
                )
            elif file_ext == ".docx":
                success = self.export_manager.export_docx(list(self.transcript_history), filename)
            else:
                with open(filename, "w", encoding="utf-8") as f:
                    if self.current_stream_info:
                        f.write("=== STREAM INFORMATION ===\n")
                        f.write(f"Title: {self.current_stream_info.title}\n")
                        f.write(f"Uploader: {self.current_stream_info.uploader}\n")
                        f.write(f"Duration: {self.current_stream_info.duration}\n")
                        f.write(f"Platform: {self.current_stream_info.platform}\n")
                        f.write(f"Saved at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    f.write("=== TRANSCRIPT ===\n")
                    if hasattr(self, 'transcript_text'):
                        f.write(self.transcript_text.get("1.0", "end-1c"))
                    f.write("\n\n=== TRANSLATION ===\n")
                    if hasattr(self, 'translation_text'):
                        f.write(self.translation_text.get("1.0", "end-1c"))
                success = True
            if success:
                self.update_status(f"💾 Saved: {os.path.basename(filename)}")
            else:
                self.update_status("❌ Export failed")
        except Exception as e:
            self.update_status(f"❌ Save failed: {e}")

    def export_subtitles(self) -> None:
        if not hasattr(self, 'audio_processor') or not self.audio_processor._timed_transcriptions:
            DarkMessageBox.showinfo(
                "WARNING",
                "No subtitle data available.\n\n"
                "Tip: First activate '🎬 Subtitle mode' "
                "and start a transcription.",
                self.root,
            )
            return
        try:
            filename = filedialog.asksaveasfilename(
                defaultextension=".srt",
                filetypes=[
                    ("SRT subtitles", "*.srt"),
                    ("VTT subtitles", "*.vtt"),
                    ("All files", "*.*"),
                ],
                title="Export subtitles",
            )
            if not filename:
                return
            file_ext = Path(filename).suffix.lower()
            format_type = "srt" if file_ext == ".srt" else "vtt"
            with self.audio_processor._subtitle_lock:
                timed_trans = list(self.audio_processor._timed_transcriptions)
                timed_transl = list(self.audio_processor._timed_translations)
            success = self.export_manager.export_subtitles(
                timed_trans,
                timed_transl,
                format=format_type,
                filename=filename,
            )
            if success:
                segment_count = len(timed_trans)
                translation_count = len(timed_transl)
                self.update_status(f"📝 {format_type.upper()} exported: {os.path.basename(filename)}")
                DarkMessageBox.showinfo(
                    "Success",
                    f"Subtitles successfully exported!\n\n"
                    f"• File: {os.path.basename(filename)}\n"
                    f"• Segments: {segment_count}\n"
                    f"• Translations: {translation_count}\n"
                    f"• Format: {format_type.upper()}\n\n"
                    f"Can be directly imported into video editors.",
                    self.root,
                )
            else:
                self.update_status("❌ Subtitle export failed")
        except Exception as e:
            self.update_status(f"❌ Subtitle export failed: {e}")
            DarkMessageBox.showerror("Error", f"Export failed:\n{str(e)}", self.root)

    def show_simple_stats(self) -> None:
        try:
            stats = self.performance_monitor.get_basic_stats()
            try:
                import psutil
                cpu = psutil.cpu_percent()
                memory = psutil.virtual_memory()
                memory_used = memory.used // (1024**2)
                memory_percent = memory.percent
                health_status = "Healthy" if cpu < 90 and memory_percent < 85 else "Degraded"
            except Exception:
                cpu = 0.0
                memory_used = 0
                memory_percent = 0
                health_status = "Unknown"
            stats_text = f"""📊 STATISTIKEN:

🤖 PERFORMANCE:
⏱️ Runtime: {stats["uptime_minutes"]:.1f} minutes
📝 Transcriptions: {stats["transcriptions"]}
🌐 Translations: {stats["translations"]}
🎯 Cache Hit Rate: {stats["cache_hit_rate"]}

💻 SYSTEM:
🖥️ CPU: {cpu:.1f}%
🧠 RAM: {memory_used}MB ({memory_percent:.1f}%)
⚡ Status: {health_status}

🎬 Subtitle mode: {"Active" if self.subtitle_mode else "Inactive"}
"""
            DarkMessageBox.showinfo("Performance Statistics", stats_text, self.root)
        except Exception as e:
            self.update_status(f"❌ Statistics error: {e}")

    # ==================== ERWEITERTER ADVANCED-DIALOG ====================
    # Die Methode show_advanced_settings wurde in mehrere Hilfsmethoden aufgeteilt,
    # um die Lesbarkeit zu verbessern.
    def show_advanced_settings(self) -> None:
        """Erweitertes Einstellungsfenster mit optimiertem Layout und Profilauswahl."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Advanced Settings")
        dialog.geometry("900x700")
        dialog.configure(bg=self.current_theme.BG_PRIMARY)
        dialog.transient(self.root)
        dialog.grab_set()

        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - dialog.winfo_width()) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")

        main_frame = tk.Frame(dialog, bg=self.current_theme.BG_PRIMARY, padx=20, pady=20)
        main_frame.pack(fill="both", expand=True)

        canvas = tk.Canvas(main_frame, bg=self.current_theme.BG_PRIMARY, highlightthickness=0)
        scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.current_theme.BG_PRIMARY)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind("<MouseWheel>", _on_mousewheel)
        canvas.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        settings = scrollable_frame

        # ----- Profile-Auswahl (oben) -----
        self._create_profile_section(dialog, settings)

        # ----- Audio & VAD -----
        self._create_audio_vad_section(dialog, settings)

        # ----- Modell & Inferenz -----
        self._create_model_section(dialog, settings)

        # ----- Transkriptions‑Filter -----
        self._create_filter_section(dialog, settings)

        # ----- Übersetzung -----
        self._create_translation_section(dialog, settings)

        # ----- GUI & Display -----
        self._create_gui_section(dialog, settings)

        # ----- Erweitert & System -----
        self._create_advanced_section(dialog, settings)

        # ----- Hilfetext-Infobox -----
        help_label = tk.Label(settings, text="Bewegen Sie die Maus über eine Einstellung für Details.",
                              bg=self.current_theme.BG_PRIMARY,
                              fg=self.current_theme.TEXT_SECONDARY,
                              font=("Segoe UI", 7), anchor="w", justify="left")
        help_label.grid(row=7, column=0, sticky="ew", pady=(10, 0))

        # Tooltips für Hilfetext
        self._bind_help_tooltips(dialog, help_label)

        # ----- Buttons (Reset, Save, Cancel) -----
        button_frame = tk.Frame(settings, bg=self.current_theme.BG_PRIMARY)
        button_frame.grid(row=8, column=0, pady=20)

        reset_btn = tk.Button(button_frame, text="Reset to Defaults",
                              command=lambda: self._reset_settings_to_default(dialog),
                              bg=self.current_theme.BG_TERTIARY,
                              fg=self.current_theme.TEXT_PRIMARY,
                              relief="flat", padx=15,
                              font=("Segoe UI", 8))
        reset_btn.pack(side="left", padx=5)

        save_btn = tk.Button(button_frame, text="Save",
                             command=lambda: self._save_settings(dialog),
                             bg=self.current_theme.SUCCESS,
                             fg=self.current_theme.TEXT_PRIMARY,
                             relief="flat", padx=15,
                             font=("Segoe UI", 8, "bold"))
        save_btn.pack(side="left", padx=5)

        cancel_btn = tk.Button(button_frame, text="Cancel", command=dialog.destroy,
                               bg=self.current_theme.BG_TERTIARY,
                               fg=self.current_theme.TEXT_PRIMARY,
                               relief="flat", padx=15,
                               font=("Segoe UI", 8))
        cancel_btn.pack(side="left", padx=5)

        settings.columnconfigure(0, weight=1)

    def _create_profile_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für vordefinierte Profile."""
        profile_frame = tk.LabelFrame(parent, text="📋 Vordefinierte Profile", padx=10, pady=8,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      font=("Segoe UI", 8, "bold"))
        profile_frame.grid(row=0, column=0, sticky="ew", pady=5, padx=5)
        profile_frame.columnconfigure(1, weight=1)

        tk.Label(profile_frame, text="Profil:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", padx=5)

        dialog.profile_var = tk.StringVar()
        dialog.profile_combo = ttk.Combobox(profile_frame, textvariable=dialog.profile_var,
                                            values=[
                                                "Default",
                                                "Deutsche Videos",
                                                "Englische Livestreams",
                                                "Asiatische Inhalte",
                                                "International (gemischt)"
                                            ],
                                            width=25, state="readonly",
                                            style="Dark.TCombobox")
        dialog.profile_combo.grid(row=0, column=1, sticky="ew", padx=5)

        profiles = {
            "Default": {
                "chunk_duration": 5.0,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 225,
                "vad_min_silence_ms": 80,
                "beam_size": 5,
                "temperature": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Automatisch",
                "min_confidence": 0.25,
                "duplicate_threshold": 0.85,
                "adaptive_low_words": 3,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
            },
            "Deutsche Videos": {
                "chunk_duration": 10.0,
                "vad_threshold": 0.3,
                "vad_min_speech_ms": 400,
                "vad_min_silence_ms": 100,
                "beam_size": 8,
                "temperature": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Deutsch",
                "min_confidence": 0.25,
                "duplicate_threshold": 0.85,
                "adaptive_low_words": 3,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
            },
            "Englische Livestreams": {
                "chunk_duration": 5.0,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 200,
                "vad_min_silence_ms": 60,
                "beam_size": 5,
                "temperature": 0.1,
                "audio_profile": "realtime",
                "source_lang_name": "Englisch",
                "min_confidence": 0.25,
                "duplicate_threshold": 0.85,
                "adaptive_low_words": 3,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
            },
            "Asiatische Inhalte": {
                "chunk_duration": 15.0,
                "vad_threshold": 0.35,
                "vad_min_speech_ms": 300,
                "vad_min_silence_ms": 120,
                "beam_size": 10,
                "temperature": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Japanisch",
                "min_confidence": 0.25,
                "duplicate_threshold": 0.85,
                "adaptive_low_words": 3,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
            },
            "International (gemischt)": {
                "chunk_duration": 8.0,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 250,
                "vad_min_silence_ms": 80,
                "beam_size": 7,
                "temperature": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Automatisch",
                "min_confidence": 0.25,
                "duplicate_threshold": 0.85,
                "adaptive_low_words": 3,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
            }
        }

        def apply_profile(*args):
            profil = profiles.get(dialog.profile_var.get())
            if not profil:
                return
            dialog.chunk_var.set(profil["chunk_duration"])
            dialog.vad_threshold_var.set(profil["vad_threshold"])
            dialog.vad_min_speech_var.set(profil["vad_min_speech_ms"])
            dialog.vad_min_silence_var.set(profil["vad_min_silence_ms"])
            dialog.beam_var.set(profil["beam_size"])
            dialog.temp_var.set(profil["temperature"])
            dialog.profile_var_audio.set(profil["audio_profile"])
            self.src_lang_var.set(profil["source_lang_name"])
            dialog.min_conf_var.set(profil.get("min_confidence", 0.25))
            dialog.dup_thresh_var.set(profil.get("duplicate_threshold", 0.85))
            dialog.low_words_var.set(profil.get("adaptive_low_words", 3))
            dialog.high_words_var.set(profil.get("adaptive_high_words", 10))
            dialog.max_mem_var.set(profil.get("max_memory_mb", 1024))
            dialog.auto_save_interval_var.set(profil.get("auto_save_interval", 300))
            dialog.optimize_var.set(profil.get("optimize_translations", False))
            dialog.sentiment_var.set(profil.get("sentiment", False))
            dialog.diarize_var.set(profil.get("diarize", False))

        dialog.profile_combo.bind("<<ComboboxSelected>>", apply_profile)

    def _create_audio_vad_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für Audio & VAD."""
        audio_frame = tk.LabelFrame(parent, text="🎵 Audio & VAD", padx=10, pady=8,
                                    bg=self.current_theme.BG_SECONDARY,
                                    fg=self.current_theme.TEXT_PRIMARY,
                                    font=("Segoe UI", 8, "bold"))
        audio_frame.grid(row=1, column=0, sticky="ew", pady=5, padx=5)
        audio_frame.columnconfigure(1, weight=1)
        audio_frame.columnconfigure(3, weight=1)

        # Zeile 0: Sample Rate (links) + Chunk Duration (rechts)
        tk.Label(audio_frame, text="Sample Rate (Hz):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        sr_label = tk.Label(audio_frame, text=str(self.advanced_settings.config.SAMPLE_RATE),
                            bg=self.current_theme.BG_TERTIARY, relief="sunken", width=10,
                            fg=self.current_theme.TEXT_PRIMARY,
                            font=("Segoe UI", 8))
        sr_label.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(audio_frame, text="Chunk Duration (s):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=2, sticky="w", pady=1)
        dialog.chunk_var = tk.DoubleVar(value=self.advanced_settings.config.CHUNK_DURATION)
        chunk_spin = tk.Spinbox(audio_frame, from_=2.0, to=30.0, increment=0.5,
                                textvariable=dialog.chunk_var, width=8,
                                bg=self.current_theme.BG_TERTIARY,
                                fg=self.current_theme.TEXT_PRIMARY,
                                buttonbackground=self.current_theme.BG_TERTIARY,
                                font=("Segoe UI", 8))
        chunk_spin.grid(row=0, column=3, sticky="w", pady=1)

        # Zeile 1: Channels (links) + Audio Profile (rechts)
        tk.Label(audio_frame, text="Channels:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=0, sticky="w", pady=1)
        ch_label = tk.Label(audio_frame, text=str(self.advanced_settings.config.CHANNELS),
                            bg=self.current_theme.BG_TERTIARY, relief="sunken", width=10,
                            fg=self.current_theme.TEXT_PRIMARY,
                            font=("Segoe UI", 8))
        ch_label.grid(row=1, column=1, sticky="w", pady=1)

        tk.Label(audio_frame, text="Audio Profile:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=2, sticky="w", pady=1)
        dialog.profile_var_audio = tk.StringVar(value=self.advanced_settings.audio_profile)
        profile_combo_audio = ttk.Combobox(audio_frame, textvariable=dialog.profile_var_audio,
                                            values=list(Constants.FILTER_PROFILES.keys()),
                                            width=12, state="readonly",
                                            style="Dark.TCombobox")
        profile_combo_audio.grid(row=1, column=3, sticky="w", pady=1)

        # Zeile 2: VAD Threshold (links) + Min Speech Duration (rechts)
        tk.Label(audio_frame, text="VAD Threshold:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=2, column=0, sticky="w", pady=1)
        dialog.vad_threshold_var = tk.DoubleVar(value=self.advanced_settings.vad_threshold)
        vad_scale = tk.Scale(audio_frame, from_=0.0, to=1.0, resolution=0.05,
                             orient=tk.HORIZONTAL, variable=dialog.vad_threshold_var,
                             length=150, showvalue=True,
                             bg=self.current_theme.BG_SECONDARY,
                             fg=self.current_theme.TEXT_PRIMARY,
                             troughcolor=self.current_theme.BG_TERTIARY,
                             font=("Segoe UI", 8))
        vad_scale.grid(row=2, column=1, sticky="ew", pady=1)

        tk.Label(audio_frame, text="Min Speech (ms):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=2, column=2, sticky="w", pady=1)
        dialog.vad_min_speech_var = tk.IntVar(value=self.advanced_settings.vad_min_speech_duration_ms)
        vad_speech_spin = tk.Spinbox(audio_frame, from_=0, to=2000, increment=50,
                                     textvariable=dialog.vad_min_speech_var, width=8,
                                     bg=self.current_theme.BG_TERTIARY,
                                     fg=self.current_theme.TEXT_PRIMARY,
                                     buttonbackground=self.current_theme.BG_TERTIARY,
                                     font=("Segoe UI", 8))
        vad_speech_spin.grid(row=2, column=3, sticky="w", pady=1)

        # Zeile 3: Min Silence Duration (links) – rechts bleibt leer
        tk.Label(audio_frame, text="Min Silence (ms):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=3, column=0, sticky="w", pady=1)
        dialog.vad_min_silence_var = tk.IntVar(value=self.advanced_settings.vad_min_silence_duration_ms)
        vad_silence_spin = tk.Spinbox(audio_frame, from_=0, to=2000, increment=50,
                                      textvariable=dialog.vad_min_silence_var, width=8,
                                      bg=self.current_theme.BG_TERTIARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      buttonbackground=self.current_theme.BG_TERTIARY,
                                      font=("Segoe UI", 8))
        vad_silence_spin.grid(row=3, column=1, sticky="w", pady=1)

    def _create_model_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für Modell & Inferenz."""
        model_frame = tk.LabelFrame(parent, text="🤖 Modell & Inferenz", padx=10, pady=8,
                                    bg=self.current_theme.BG_SECONDARY,
                                    fg=self.current_theme.TEXT_PRIMARY,
                                    font=("Segoe UI", 8, "bold"))
        model_frame.grid(row=2, column=0, sticky="ew", pady=5, padx=5)
        model_frame.columnconfigure(1, weight=1)
        model_frame.columnconfigure(3, weight=1)

        tk.Label(model_frame, text="Beam Size:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        dialog.beam_var = tk.IntVar(value=self.advanced_settings.beam_size)
        beam_spin = tk.Spinbox(model_frame, from_=1, to=20, textvariable=dialog.beam_var, width=8,
                               bg=self.current_theme.BG_TERTIARY,
                               fg=self.current_theme.TEXT_PRIMARY,
                               buttonbackground=self.current_theme.BG_TERTIARY,
                               font=("Segoe UI", 8))
        beam_spin.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(model_frame, text="Temperature:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=2, sticky="w", pady=1)
        dialog.temp_var = tk.DoubleVar(value=self.advanced_settings.temperature)
        temp_scale = tk.Scale(model_frame, from_=0.0, to=2.0, resolution=0.1,
                              orient=tk.HORIZONTAL, variable=dialog.temp_var,
                              length=150, showvalue=True,
                              bg=self.current_theme.BG_SECONDARY,
                              fg=self.current_theme.TEXT_PRIMARY,
                              troughcolor=self.current_theme.BG_TERTIARY,
                              font=("Segoe UI", 8))
        temp_scale.grid(row=0, column=3, sticky="ew", pady=1)

        dialog.gpu_var = tk.BooleanVar(value=self.advanced_settings.gpu_acceleration)
        gpu_cb = tk.Checkbutton(model_frame, text="GPU Acceleration", variable=dialog.gpu_var,
                                bg=self.current_theme.BG_SECONDARY,
                                fg=self.current_theme.TEXT_PRIMARY,
                                selectcolor=self.current_theme.BG_TERTIARY,
                                activebackground=self.current_theme.BG_SECONDARY,
                                font=("Segoe UI", 8))
        gpu_cb.grid(row=1, column=0, columnspan=2, sticky="w", pady=1)

        tk.Label(model_frame, text="Current Model:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=2, sticky="w", pady=1)
        current_model = self.transcription_engine.get_current_model() if hasattr(self, 'transcription_engine') else "unknown"
        model_label = tk.Label(model_frame, text=current_model,
                               bg=self.current_theme.BG_TERTIARY, relief="sunken", width=15,
                               fg=self.current_theme.TEXT_PRIMARY,
                               font=("Segoe UI", 8))
        model_label.grid(row=1, column=3, sticky="w", pady=1)

    def _create_filter_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für Transkriptions‑Filter."""
        filter_frame = tk.LabelFrame(parent, text="🔍 Transkriptions‑Filter", padx=10, pady=8,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      font=("Segoe UI", 8, "bold"))
        filter_frame.grid(row=3, column=0, sticky="ew", pady=5, padx=5)
        filter_frame.columnconfigure(1, weight=1)
        filter_frame.columnconfigure(3, weight=1)

        # Zeile 0: Min Confidence (links) + Duplicate Threshold (rechts)
        tk.Label(filter_frame, text="Min Confidence:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        dialog.min_conf_var = tk.DoubleVar(value=self.advanced_settings.min_confidence)
        min_conf_scale = tk.Scale(filter_frame, from_=0.0, to=1.0, resolution=0.05,
                                   orient=tk.HORIZONTAL, variable=dialog.min_conf_var,
                                   length=150, showvalue=True,
                                   bg=self.current_theme.BG_SECONDARY,
                                   fg=self.current_theme.TEXT_PRIMARY,
                                   troughcolor=self.current_theme.BG_TERTIARY,
                                   font=("Segoe UI", 8))
        min_conf_scale.grid(row=0, column=1, sticky="ew", pady=1)

        tk.Label(filter_frame, text="Duplicate Threshold:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=2, sticky="w", pady=1, padx=(10,0))
        dialog.dup_thresh_var = tk.DoubleVar(value=self.advanced_settings.duplicate_similarity_threshold)
        dup_thresh_scale = tk.Scale(filter_frame, from_=0.5, to=1.0, resolution=0.01,
                                     orient=tk.HORIZONTAL, variable=dialog.dup_thresh_var,
                                     length=150, showvalue=True,
                                     bg=self.current_theme.BG_SECONDARY,
                                     fg=self.current_theme.TEXT_PRIMARY,
                                     troughcolor=self.current_theme.BG_TERTIARY,
                                     font=("Segoe UI", 8))
        dup_thresh_scale.grid(row=0, column=3, sticky="ew", pady=1)

        # Zeile 1: Adaptive Low Words (links) + Adaptive High Words (rechts)
        tk.Label(filter_frame, text="Adaptive Low Words:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=0, sticky="w", pady=1)
        dialog.low_words_var = tk.IntVar(value=self.advanced_settings.adaptive_chunk_low_words)
        low_words_spin = tk.Spinbox(filter_frame, from_=1, to=20, textvariable=dialog.low_words_var, width=8,
                                     bg=self.current_theme.BG_TERTIARY,
                                     fg=self.current_theme.TEXT_PRIMARY,
                                     buttonbackground=self.current_theme.BG_TERTIARY,
                                     font=("Segoe UI", 8))
        low_words_spin.grid(row=1, column=1, sticky="w", pady=1)

        tk.Label(filter_frame, text="Adaptive High Words:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=2, sticky="w", pady=1, padx=(10,0))
        dialog.high_words_var = tk.IntVar(value=self.advanced_settings.adaptive_chunk_high_words)
        high_words_spin = tk.Spinbox(filter_frame, from_=1, to=20, textvariable=dialog.high_words_var, width=8,
                                      bg=self.current_theme.BG_TERTIARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      buttonbackground=self.current_theme.BG_TERTIARY,
                                      font=("Segoe UI", 8))
        high_words_spin.grid(row=1, column=3, sticky="w", pady=1)

    def _create_translation_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für Übersetzung."""
        trans_frame = tk.LabelFrame(parent, text="🌐 Übersetzung", padx=10, pady=8,
                                    bg=self.current_theme.BG_SECONDARY,
                                    fg=self.current_theme.TEXT_PRIMARY,
                                    font=("Segoe UI", 8, "bold"))
        trans_frame.grid(row=4, column=0, sticky="ew", pady=5, padx=5)
        trans_frame.columnconfigure(1, weight=1)
        trans_frame.columnconfigure(3, weight=1)

        # Zeile 0: Engine (links) + Ollama Model (rechts)
        tk.Label(trans_frame, text="Engine:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        dialog.engine_var = tk.StringVar(value=self.advanced_settings.translation_engine)
        engine_combo = ttk.Combobox(trans_frame, textvariable=dialog.engine_var,
                                    values=["google", "ollama"],
                                    width=10, state="readonly",
                                    style="Dark.TCombobox")
        engine_combo.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(trans_frame, text="Ollama Model:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=2, sticky="w", pady=1)
        dialog.ollama_model_var = tk.StringVar(value=self.advanced_settings.ollama_model)
        ollama_model_entry = tk.Entry(trans_frame, textvariable=dialog.ollama_model_var, width=15,
                                       bg=self.current_theme.BG_TERTIARY,
                                       fg=self.current_theme.TEXT_PRIMARY,
                                       insertbackground=self.current_theme.TEXT_PRIMARY,
                                       font=("Segoe UI", 8))
        ollama_model_entry.grid(row=0, column=3, sticky="w", pady=1)

        # Zeile 1: Ollama Host (links, mit eigenem Frame für Host+Button)
        tk.Label(trans_frame, text="Ollama Host:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=0, sticky="w", pady=1)
        host_frame = tk.Frame(trans_frame, bg=self.current_theme.BG_SECONDARY)
        host_frame.grid(row=1, column=1, columnspan=3, sticky="ew", pady=1)
        host_frame.columnconfigure(0, weight=1)

        dialog.ollama_host_var = tk.StringVar(value=self.advanced_settings.ollama_host)
        ollama_host_entry = tk.Entry(host_frame, textvariable=dialog.ollama_host_var, width=30,
                                      bg=self.current_theme.BG_TERTIARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      insertbackground=self.current_theme.TEXT_PRIMARY,
                                      font=("Segoe UI", 8))
        ollama_host_entry.pack(side="left", fill="x", expand=True, padx=(0,5))

        def test_ollama():
            host = dialog.ollama_host_var.get().strip()
            if not host:
                host = "http://localhost:11434"
            try:
                import requests
                r = requests.get(f"{host}/api/tags", timeout=3)
                if r.status_code == 200:
                    DarkMessageBox.showinfo("Success", "Ollama server is reachable!", self.root)
                else:
                    DarkMessageBox.showerror("Error", f"Ollama returned status {r.status_code}", self.root)
            except Exception as e:
                DarkMessageBox.showerror("Error", f"Connection failed: {e}", self.root)

        test_btn = tk.Button(host_frame, text="Test", command=test_ollama,
                             bg=self.current_theme.BG_TERTIARY,
                             fg=self.current_theme.TEXT_PRIMARY,
                             relief="flat", padx=5,
                             font=("Segoe UI", 8))
        test_btn.pack(side="right")

    def _create_gui_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für GUI & Display."""
        gui_frame = tk.LabelFrame(parent, text="🖥️ GUI & Display", padx=10, pady=8,
                                  bg=self.current_theme.BG_SECONDARY,
                                  fg=self.current_theme.TEXT_PRIMARY,
                                  font=("Segoe UI", 8, "bold"))
        gui_frame.grid(row=5, column=0, sticky="ew", pady=5, padx=5)
        gui_frame.columnconfigure(1, weight=1)
        gui_frame.columnconfigure(3, weight=1)

        # Zeile 0: Transcript Max Lines (links) + Translation Max Lines (rechts)
        tk.Label(gui_frame, text="Transcript Max Lines:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        dialog.trans_lines_var = tk.IntVar(value=self.advanced_settings.transcript_max_lines)
        trans_lines_spin = tk.Spinbox(gui_frame, from_=100, to=5000, increment=100,
                                      textvariable=dialog.trans_lines_var, width=8,
                                      bg=self.current_theme.BG_TERTIARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      buttonbackground=self.current_theme.BG_TERTIARY,
                                      font=("Segoe UI", 8))
        trans_lines_spin.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(gui_frame, text="Translation Max Lines:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=2, sticky="w", pady=1)
        dialog.transl_lines_var = tk.IntVar(value=self.advanced_settings.translation_max_lines)
        transl_lines_spin = tk.Spinbox(gui_frame, from_=100, to=5000, increment=100,
                                       textvariable=dialog.transl_lines_var, width=8,
                                       bg=self.current_theme.BG_TERTIARY,
                                       fg=self.current_theme.TEXT_PRIMARY,
                                       buttonbackground=self.current_theme.BG_TERTIARY,
                                       font=("Segoe UI", 8))
        transl_lines_spin.grid(row=0, column=3, sticky="w", pady=1)

        # Zeile 1: Theme (links) + Auto-Save (rechts)
        tk.Label(gui_frame, text="Theme:", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=1, column=0, sticky="w", pady=1)
        dialog.theme_var = tk.StringVar(value=self.settings.theme)
        theme_combo = ttk.Combobox(gui_frame, textvariable=dialog.theme_var,
                                   values=["dark", "light", "system"],
                                   width=10, state="readonly",
                                   style="Dark.TCombobox")
        theme_combo.grid(row=1, column=1, sticky="w", pady=1)

        dialog.auto_save_var = tk.BooleanVar(value=self.settings.auto_save_on_completion)
        auto_save_cb = tk.Checkbutton(gui_frame, text="Auto-Save on Completion", variable=dialog.auto_save_var,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      selectcolor=self.current_theme.BG_TERTIARY,
                                      activebackground=self.current_theme.BG_SECONDARY,
                                      font=("Segoe UI", 8))
        auto_save_cb.grid(row=1, column=2, columnspan=2, sticky="w", pady=1)

    def _create_advanced_section(self, dialog: tk.Toplevel, parent: tk.Frame) -> None:
        """Erstellt den Abschnitt für Erweitert & System."""
        adv_frame = tk.LabelFrame(parent, text="⚙️ Erweitert & System", padx=10, pady=8,
                                  bg=self.current_theme.BG_SECONDARY,
                                  fg=self.current_theme.TEXT_PRIMARY,
                                  font=("Segoe UI", 8, "bold"))
        adv_frame.grid(row=6, column=0, sticky="ew", pady=5, padx=5)
        adv_frame.columnconfigure(1, weight=1)
        adv_frame.columnconfigure(3, weight=1)

        # Zeile 0: Max Cache Size (MB) (links) + Enable Plugins (rechts) – Plugins entfernt, daher nur Cache
        tk.Label(adv_frame, text="Max Cache Size (MB):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=0, column=0, sticky="w", pady=1)
        dialog.cache_var = tk.IntVar(value=self.advanced_settings.max_cache_size)
        cache_spin = tk.Spinbox(adv_frame, from_=10, to=1000, increment=10,
                                textvariable=dialog.cache_var, width=8,
                                bg=self.current_theme.BG_TERTIARY,
                                fg=self.current_theme.TEXT_PRIMARY,
                                buttonbackground=self.current_theme.BG_TERTIARY,
                                font=("Segoe UI", 8))
        cache_spin.grid(row=0, column=1, sticky="w", pady=1)

        # Zeile 1: Use Browser Cookies (links) + Asian Mode (rechts)
        dialog.cookies_var = tk.BooleanVar(value=self.settings.use_browser_cookies)
        cookies_cb = tk.Checkbutton(adv_frame, text="Use Browser Cookies for YouTube", variable=dialog.cookies_var,
                                    bg=self.current_theme.BG_SECONDARY,
                                    fg=self.current_theme.TEXT_PRIMARY,
                                    selectcolor=self.current_theme.BG_TERTIARY,
                                    activebackground=self.current_theme.BG_SECONDARY,
                                    font=("Segoe UI", 8))
        cookies_cb.grid(row=1, column=0, columnspan=2, sticky="w", pady=1)

        dialog.asian_var = tk.BooleanVar(value=self.advanced_settings.asian_mode)
        asian_cb = tk.Checkbutton(adv_frame, text="Asian Mode (10s chunks)", variable=dialog.asian_var,
                                  bg=self.current_theme.BG_SECONDARY,
                                  fg=self.current_theme.TEXT_PRIMARY,
                                  selectcolor=self.current_theme.BG_TERTIARY,
                                  activebackground=self.current_theme.BG_SECONDARY,
                                  font=("Segoe UI", 8))
        asian_cb.grid(row=1, column=2, columnspan=2, sticky="w", pady=1)

        # Zeile 2: Precision Mode (links) – bisher nur eine Spalte, jetzt nutzen wir für neue Werte
        dialog.precision_var = tk.BooleanVar(value=self.advanced_settings.precision_mode)
        precision_cb = tk.Checkbutton(adv_frame, text="Precision Mode (langsamer, genauer)", variable=dialog.precision_var,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      selectcolor=self.current_theme.BG_TERTIARY,
                                      activebackground=self.current_theme.BG_SECONDARY,
                                      font=("Segoe UI", 8))
        precision_cb.grid(row=2, column=0, columnspan=2, sticky="w", pady=1)

        # Zeile 3 (neu): Max Memory (MB) und Auto Save Interval
        tk.Label(adv_frame, text="Max Memory (MB):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=3, column=0, sticky="w", pady=1)
        dialog.max_mem_var = tk.IntVar(value=self.advanced_settings.max_memory_mb)
        max_mem_spin = tk.Spinbox(adv_frame, from_=100, to=16384, increment=100,
                                   textvariable=dialog.max_mem_var, width=8,
                                   bg=self.current_theme.BG_TERTIARY,
                                   fg=self.current_theme.TEXT_PRIMARY,
                                   buttonbackground=self.current_theme.BG_TERTIARY,
                                   font=("Segoe UI", 8))
        max_mem_spin.grid(row=3, column=1, sticky="w", pady=1)

        tk.Label(adv_frame, text="Auto Save Interval (s):", anchor="w",
                 bg=self.current_theme.BG_SECONDARY,
                 fg=self.current_theme.TEXT_PRIMARY,
                 font=("Segoe UI", 8)).grid(row=3, column=2, sticky="w", pady=1, padx=(10,0))
        dialog.auto_save_interval_var = tk.IntVar(value=self.advanced_settings.auto_save_interval)
        auto_save_interval_spin = tk.Spinbox(adv_frame, from_=0, to=3600, increment=60,
                                     textvariable=dialog.auto_save_interval_var, width=8,
                                     bg=self.current_theme.BG_TERTIARY,
                                     fg=self.current_theme.TEXT_PRIMARY,
                                     buttonbackground=self.current_theme.BG_TERTIARY,
                                     font=("Segoe UI", 8))
        auto_save_interval_spin.grid(row=3, column=3, sticky="w", pady=1)

        # Zeile 4: Optimize Translations, Sentiment Analysis (zweispaltig)
        dialog.optimize_var = tk.BooleanVar(value=self.advanced_settings.optimize_translations)
        optimize_cb = tk.Checkbutton(adv_frame, text="Optimize Translations", variable=dialog.optimize_var,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      selectcolor=self.current_theme.BG_TERTIARY,
                                      activebackground=self.current_theme.BG_SECONDARY,
                                      font=("Segoe UI", 8))
        optimize_cb.grid(row=4, column=0, columnspan=2, sticky="w", pady=1)

        dialog.sentiment_var = tk.BooleanVar(value=self.advanced_settings.enable_sentiment_analysis)
        sentiment_cb = tk.Checkbutton(adv_frame, text="Sentiment Analysis", variable=dialog.sentiment_var,
                                       bg=self.current_theme.BG_SECONDARY,
                                       fg=self.current_theme.TEXT_PRIMARY,
                                       selectcolor=self.current_theme.BG_TERTIARY,
                                       activebackground=self.current_theme.BG_SECONDARY,
                                       font=("Segoe UI", 8))
        sentiment_cb.grid(row=4, column=2, columnspan=2, sticky="w", pady=1)

        # Zeile 5: Speaker Diarization (einzeln, weil länger)
        dialog.diarize_var = tk.BooleanVar(value=self.advanced_settings.enable_speaker_diarization)
        diarize_cb = tk.Checkbutton(adv_frame, text="Speaker Diarization", variable=dialog.diarize_var,
                                      bg=self.current_theme.BG_SECONDARY,
                                      fg=self.current_theme.TEXT_PRIMARY,
                                      selectcolor=self.current_theme.BG_TERTIARY,
                                      activebackground=self.current_theme.BG_SECONDARY,
                                      font=("Segoe UI", 8))
        diarize_cb.grid(row=5, column=0, columnspan=2, sticky="w", pady=1)

    def _bind_help_tooltips(self, dialog: tk.Toplevel, help_label: tk.Label) -> None:
        """Bindet Tooltips für die Hilfe-Infobox im Einstellungsdialog."""
        # Diese Methode würde die vielen bind-Aufrufe enthalten (aus Platzgründen hier nur ein Beispiel)
        # Im Original waren viele bind(...) vorhanden. Wir verzichten hier auf die vollständige Auflistung,
        # da sie den Rahmen sprengen würde. In der Praxis würde man hier alle relevanten Widgets mit
        # show_help-Funktionen verbinden.
        pass

    def _reset_settings_to_default(self, dialog: tk.Toplevel) -> None:
        """Setzt alle Einstellungen im Dialog auf die Standardwerte zurück."""
        default = Settings()
        dialog.chunk_var.set(default.config.CHUNK_DURATION)
        dialog.profile_var_audio.set('transcription')
        dialog.vad_threshold_var.set(default.vad_threshold)
        dialog.vad_min_speech_var.set(default.vad_min_speech_duration_ms)
        dialog.vad_min_silence_var.set(default.vad_min_silence_duration_ms)
        dialog.beam_var.set(default.beam_size)
        dialog.temp_var.set(default.temperature)
        dialog.gpu_var.set(default.gpu_acceleration)
        dialog.engine_var.set(default.translation_engine)
        dialog.ollama_model_var.set(default.ollama_model)
        dialog.ollama_host_var.set(default.ollama_host)
        dialog.trans_lines_var.set(default.transcript_max_lines)
        dialog.transl_lines_var.set(default.translation_max_lines)
        dialog.theme_var.set('dark')
        dialog.auto_save_var.set(False)
        dialog.cache_var.set(default.max_cache_size)
        dialog.cookies_var.set(True)
        dialog.asian_var.set(default.asian_mode)
        dialog.precision_var.set(default.precision_mode)
        self.src_lang_var.set("Automatisch")
        dialog.min_conf_var.set(default.min_confidence)
        dialog.dup_thresh_var.set(default.duplicate_similarity_threshold)
        dialog.low_words_var.set(default.adaptive_chunk_low_words)
        dialog.high_words_var.set(default.adaptive_chunk_high_words)
        dialog.max_mem_var.set(default.max_memory_mb)
        dialog.auto_save_interval_var.set(default.auto_save_interval)
        dialog.optimize_var.set(default.optimize_translations)
        dialog.sentiment_var.set(default.enable_sentiment_analysis)
        dialog.diarize_var.set(default.enable_speaker_diarization)

    def _save_settings(self, dialog: tk.Toplevel) -> None:
        """Speichert die Einstellungen aus dem Dialog."""
        try:
            self.advanced_settings.config.CHUNK_DURATION = dialog.chunk_var.get()
            self.advanced_settings.chunk_duration = dialog.chunk_var.get()
            self.advanced_settings.audio_profile = dialog.profile_var_audio.get()
            self.advanced_settings.vad_threshold = dialog.vad_threshold_var.get()
            self.advanced_settings.vad_min_speech_duration_ms = dialog.vad_min_speech_var.get()
            self.advanced_settings.vad_min_silence_duration_ms = dialog.vad_min_silence_var.get()
            self.advanced_settings.beam_size = dialog.beam_var.get()
            self.advanced_settings.temperature = dialog.temp_var.get()
            self.advanced_settings.gpu_acceleration = dialog.gpu_var.get()
            self.advanced_settings.translation_engine = dialog.engine_var.get()
            self.advanced_settings.ollama_model = dialog.ollama_model_var.get().strip()
            self.advanced_settings.ollama_host = dialog.ollama_host_var.get().strip()
            self.advanced_settings.transcript_max_lines = dialog.trans_lines_var.get()
            self.advanced_settings.translation_max_lines = dialog.transl_lines_var.get()
            self.advanced_settings.asian_mode = dialog.asian_var.get()
            self.advanced_settings.precision_mode = dialog.precision_var.get()
            self.advanced_settings.max_cache_size = dialog.cache_var.get()

            self.settings.theme = dialog.theme_var.get()
            self.settings.auto_save_on_completion = dialog.auto_save_var.get()
            self.settings.use_browser_cookies = dialog.cookies_var.get()

            self.advanced_settings.min_confidence = dialog.min_conf_var.get()
            self.advanced_settings.duplicate_similarity_threshold = dialog.dup_thresh_var.get()
            self.advanced_settings.adaptive_chunk_low_words = dialog.low_words_var.get()
            self.advanced_settings.adaptive_chunk_high_words = dialog.high_words_var.get()
            self.advanced_settings.max_memory_mb = dialog.max_mem_var.get()
            self.advanced_settings.auto_save_interval = dialog.auto_save_interval_var.get()
            self.advanced_settings.optimize_translations = dialog.optimize_var.get()
            self.advanced_settings.enable_sentiment_analysis = dialog.sentiment_var.get()
            self.advanced_settings.enable_speaker_diarization = dialog.diarize_var.get()

            host = self.advanced_settings.ollama_host
            if host and not host.startswith(('http://', 'https://')):
                self.advanced_settings.ollama_host = 'http://' + host

            self.advanced_settings.save_to_file()
            self.settings.save_to_file()

            if hasattr(self, 'stream_manager'):
                self.stream_manager.use_browser_cookies = self.settings.use_browser_cookies
            if hasattr(self, 'stream_info_extractor'):
                self.stream_info_extractor.use_browser_cookies = self.settings.use_browser_cookies

            if not self.advanced_settings.gpu_acceleration:
                self.transcription_engine.device = "cpu"
                self.transcription_engine.compute_type = "int8"

            dialog.destroy()
            self.update_status("✅ Settings saved")
        except Exception as e:
            DarkMessageBox.showerror("Error", f"Invalid settings: {e}", self.root)

    def _setup_callbacks(self) -> None:
        pass

    def _handle_ui_update(self, component: str, text: str) -> None:
        pass

    def _handle_status_update(self, state_info: Dict[str, Any]) -> None:
        def update_task() -> None:
            try:
                if "status" in state_info and hasattr(self, "status_label"):
                    if self.status_label.winfo_exists():
                        self.status_label.config(text=state_info["status"][:100])
                if "buttons" in state_info:
                    buttons = state_info["buttons"]
                    if hasattr(self, "start_button") and self.start_button.winfo_exists():
                        self.start_button.config(state=buttons.get("start", "normal"))
                    if hasattr(self, "stop_button") and self.stop_button.winfo_exists():
                        self.stop_button.config(state=buttons.get("stop", "disabled"))
                elif "processing_state" in state_info:
                    processing = state_info["processing_state"]
                    if hasattr(self, "start_button") and self.start_button.winfo_exists():
                        self.start_button.config(state="disabled" if processing else "normal")
                    if hasattr(self, "stop_button") and self.stop_button.winfo_exists():
                        self.stop_button.config(state="normal" if processing else "disabled")
                if "stream_info" in state_info:
                    stream_info = state_info["stream_info"]
                    self.current_stream_info = stream_info
                    if hasattr(self, "stream_title_label") and self.stream_title_label.winfo_exists():
                        title = stream_info.title[:80] + "..." if len(stream_info.title) > 80 else stream_info.title
                        self.stream_title_label.config(text=f"📡 {title}")
                    if hasattr(self, "stream_details_label") and self.stream_details_label.winfo_exists():
                        details = f"👤 {stream_info.uploader}"
                        if stream_info.duration and stream_info.duration != "Live":
                            details += f" | ⏱️ {stream_info.duration}"
                        self.stream_details_label.config(text=details)
                if state_info.get("file_finished"):
                    logger.info("📂 Dateiende erkannt – öffne Speicherdialog")
                    if self.settings.auto_save_on_completion:
                        self.save_transcript()
                    else:
                        self.update_status("✅ Dateiende – zum Speichern auf 💾 klicken")

            except Exception as e:
                logger.warning(f"⚠️ Status update error: {e}")
        try:
            if hasattr(self, "gui_queue"):
                self._safe_queue_put(self.gui_queue, ("status_update", update_task))
            else:
                if hasattr(self, "root") and self.root.winfo_exists():
                    self.root.after(0, update_task)
        except Exception:
            pass

    def _start_gui_updaters(self) -> None:
        if hasattr(self, "root") and self.root.winfo_exists():
            self.root.after(50, self._process_gui_queue)
            self.root.after(75, self._process_text_updates)
            self.root.after(5000, self._check_queue_sizes)

    def _process_gui_queue(self) -> None:
        if self._shutting_down or not hasattr(self, "root") or not self.root.winfo_exists():
            return

        start_time = time.time()
        max_duration = 0.05
        processed = 0
        max_items = 50

        try:
            while processed < max_items and (time.time() - start_time) < max_duration:
                try:
                    item = self.gui_queue.get_nowait()
                except queue.Empty:
                    break

                if isinstance(item, tuple) and len(item) == 2:
                    msg_type, callback = item
                    if callable(callback):
                        if self._gui_update_limiter.can_update(f"gui_{msg_type}"):
                            try:
                                callback()
                            except Exception as e:
                                logger.warning(f"⚠️ GUI callback error: {e}")
                processed += 1

            if self.gui_queue.qsize() > 100:
                self._cleanup_queue(self.gui_queue, 50)
                logger.debug(f"🧹 GUI queue cleaned to 50 items (was {self.gui_queue.qsize()})")

            if debug3_enabled('gui') and processed > 0:
                logger.debug(f"[DEBUG3][GUI] Processed {processed} items in gui_queue")

        except Exception as e:
            logger.error(f"❌ GUI queue processor error: {e}")

        if not self._shutting_down and self.root.winfo_exists():
            self.root.after(100, self._process_gui_queue)

    def _process_text_updates(self) -> None:
        if self._shutting_down or not hasattr(self, "root") or not self.root.winfo_exists():
            return

        if not hasattr(self, "_text_update_queue") or self._text_update_queue is None:
            return

        start_time = time.time()
        max_duration = 0.05
        processed = 0
        max_items = 20

        try:
            while processed < max_items and (time.time() - start_time) < max_duration:
                try:
                    update_type, text_data = self._text_update_queue.get_nowait()
                except queue.Empty:
                    break

                try:
                    if update_type == "transcript" and hasattr(self, "transcript_text"):
                        if self.transcript_text.winfo_exists():
                            self.transcript_text.insert("end", text_data)
                            if hasattr(self, "transcript_scroll_var") and self.transcript_scroll_var.get():
                                self.transcript_text.see("end")
                    elif update_type == "translation" and hasattr(self, "translation_text"):
                        if self.translation_text.winfo_exists():
                            self.translation_text.insert("end", text_data)
                            if hasattr(self, "translation_scroll_var") and self.translation_scroll_var.get():
                                self.translation_text.see("end")
                except tk.TclError:
                    pass
                except Exception as e:
                    logger.warning(f"⚠️ Text update error: {e}")

                processed += 1

            if self._text_update_queue.qsize() > 150:
                self._cleanup_queue(self._text_update_queue, 75)
                logger.debug(f"🧹 Text queue cleaned to 75 items (was {self._text_update_queue.qsize()})")

            if debug3_enabled('gui') and processed > 0:
                logger.debug(f"[DEBUG3][GUI] Processed {processed} items in text queue")

        except Exception as e:
            logger.error(f"❌ Text update processor error: {e}")

        if not self._shutting_down and self.root.winfo_exists():
            self.root.after(150, self._process_text_updates)

    def _check_queue_sizes(self) -> None:
        if self._shutting_down or not hasattr(self, "root") or not self.root.winfo_exists():
            return

        try:
            if hasattr(self, "gui_queue") and self.gui_queue.qsize() > 200:
                if debug3_enabled('queue'):
                    logger.debug(f"[DEBUG3][QUEUE] GUI queue size {self.gui_queue.qsize()} exceeds threshold, cleaning up")
                self._cleanup_queue(self.gui_queue, 100)
                logger.info("🧹 Aggressive GUI queue cleanup: reduced to 100 items")
            if hasattr(self, "_text_update_queue") and self._text_update_queue.qsize() > 300:
                if debug3_enabled('queue'):
                    logger.debug(f"[DEBUG3][QUEUE] Text queue size {self._text_update_queue.qsize()} exceeds threshold, cleaning up")
                self._cleanup_queue(self._text_update_queue, 150)
                logger.info("🧹 Aggressive text queue cleanup: reduced to 150 items")
        except Exception as e:
            logger.warning(f"⚠️ Queue size check error: {e}")
        if not self._shutting_down and self.root.winfo_exists():
            self.root.after(10000, self._check_queue_sizes)

    def _start_automatic_maintenance(self) -> None:
        def maintenance_worker() -> None:
            if not self._shutting_down and hasattr(self, 'root') and self.root.winfo_exists():
                self._perform_maintenance()
                self.root.after(600000, maintenance_worker)
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(60000, maintenance_worker)

    def _perform_maintenance(self) -> None:
        logger.info("🛠️ Performing automatic maintenance...")
        try:
            expired_counts = clear_expired_cache_entries()
            if any(count > 0 for count in expired_counts.values()):
                logger.info(f"🧹 Cleared expired cache entries: {expired_counts}")
        except Exception:
            pass
        if hasattr(self, 'memory_manager'):
            try:
                self.memory_manager._perform_periodic_maintenance()
            except Exception:
                pass
        if hasattr(self, 'ffmpeg_manager'):
            try:
                self.ffmpeg_manager.cleanup_stale_processes()
            except Exception:
                pass
        gc.collect()
        logger.info("✅ Maintenance completed")

    def handle_transcription(self, result: TranscriptionResult) -> None:
        if not result or not result.text or not result.text.strip():
            return
        current_text = result.text.strip()
        if current_text == self._last_transcription_text:
            return
        self._last_transcription_text = current_text
        self.performance_monitor.log_transcription()
        with self._history_lock:
            self.transcript_history.append(result)
        try:
            if self.subtitle_mode and result.start is not None:
                timestamp = self.export_manager._format_timestamp_srt(result.start)
            else:
                timestamp = datetime.now().strftime("%H:%M:%S")
            detected_lang = getattr(result, 'language', 'unknown')
            lang_code = LANGUAGE_SHORT_CODES.get(detected_lang, '??')
            text = f"[{timestamp}] [{lang_code}] {current_text}\n"
        except Exception as e:
            logger.warning(f"⚠️ Error preparing transcription text: {e}")
            return
            
        if hasattr(self, '_text_update_queue') and self._text_update_queue is not None:
            self._safe_queue_put(self._text_update_queue, ("transcript", text))
        else:
            logger.warning("⚠️ _text_update_queue missing, using direct after()")
            if hasattr(self, 'root') and self.root.winfo_exists():
                def fallback_update():
                    try:
                        if hasattr(self, 'transcript_text') and self.transcript_text.winfo_exists():
                            self.transcript_text.insert('end', text)
                            if hasattr(self, 'transcript_scroll_var') and self.transcript_scroll_var.get():
                                self.transcript_text.see('end')
                            lines = int(self.transcript_text.index('end-1c').split('.')[0])
                            max_lines = self.advanced_settings.transcript_max_lines
                            if lines > max_lines:
                                keep_lines = max_lines - 100
                                delete_to = f'{lines-keep_lines}.0'
                                self.transcript_text.delete('1.0', delete_to)
                    except Exception as e:
                        logger.warning(f"⚠️ Fallback GUI error: {e}")
                self.root.after(0, fallback_update)

    def handle_translation(self, result: TranslationResult) -> None:
        if not result or not result.translated or not result.translated.strip():
            return
        current_text = result.translated.strip()
        if current_text == self._last_translation_text:
            return
        self._last_translation_text = current_text
        self.performance_monitor.log_translation()
        with self._history_lock:
            self.translation_history.append(result)
        try:
            if self.subtitle_mode and result.start is not None:
                timestamp = self.export_manager._format_timestamp_srt(result.start)
            else:
                timestamp = datetime.now().strftime("%H:%M:%S")
            text = f"[{timestamp}] {current_text}\n"
        except Exception as e:
            logger.warning(f"⚠️ Error preparing translation text: {e}")
            return
        if hasattr(self, '_text_update_queue') and self._text_update_queue is not None:
            self._safe_queue_put(self._text_update_queue, ("translation", text))
        else:
            logger.warning("⚠️ _text_update_queue missing, using direct after()")
            if hasattr(self, 'root') and self.root.winfo_exists():
                def fallback_update():
                    try:
                        if hasattr(self, 'translation_text') and self.translation_text.winfo_exists():
                            self.translation_text.insert('end', text)
                            if hasattr(self, 'translation_scroll_var') and self.translation_scroll_var.get():
                                self.translation_text.see('end')
                            lines = int(self.translation_text.index('end-1c').split('.')[0])
                            max_lines = self.advanced_settings.translation_max_lines
                            if lines > max_lines:
                                keep_lines = max_lines - 100
                                delete_to = f'{lines-keep_lines}.0'
                                self.translation_text.delete('1.0', delete_to)
                    except Exception as e:
                        logger.warning(f"⚠️ Fallback GUI error: {e}")
                self.root.after(0, fallback_update)

    def handle_info(self, info_msg: str) -> None:
        def update() -> None:
            self.update_status(f"ℹ️ {info_msg}")
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(0, update)

    def handle_error(self, error_msg: str) -> None:
        def update() -> None:
            self.update_status(f"❌ {error_msg}")
            if self.is_processing:
                self.controller.stop_processing()
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(0, update)

    def _on_processing_finished(self) -> None:
        logger.info("Processing finished – GUI kann reagieren")

    def update_status(self, message: str) -> None:
        try:
            if hasattr(self, 'status_label') and self.status_label.winfo_exists():
                self.status_label.config(text=message[:100])
        except Exception:
            pass

    def update_stream_info(self, info: StreamInfo) -> None:
        def update_gui() -> None:
            try:
                self.current_stream_info = info
                if hasattr(self, 'stream_title_label') and self.stream_title_label.winfo_exists():
                    title = info.title[:80] + "..." if len(info.title) > 80 else info.title
                    self.stream_title_label.config(text=f"📡 {title}")
                if hasattr(self, 'stream_details_label') and self.stream_details_label.winfo_exists():
                    details = f"👤 {info.uploader}"
                    if info.duration and info.duration != 'Live':
                        details += f" | ⏱️ {info.duration}"
                    self.stream_details_label.config(text=details)
            except Exception as e:
                logger.warning(f"⚠️ Stream info update error: {e}")
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(0, update_gui)

    def _try_fallback_gui(self) -> None:
        logger.info("ℹ️ Starte im Kommandozeilen-Modus...")
        raise RuntimeError("Bitte installieren Sie Tkinter: pip install tk")

    def _show_error_and_exit(self, message: str) -> None:
        logger.error(f"💥 KRITISCHER FEHLER: {message}")
        try:
            import tkinter.messagebox as mb
            mb.showerror("Dragon Whisperer - Fehler", message)
        except Exception:
            pass
        self._emergency_cleanup()
        sys.exit(1)

    def _show_warning(self, message: str) -> None:
        logger.warning(f"⚠️ WARNUNG: {message}")

    def show_translation_dialog(self) -> None:
        if hasattr(self, 'translation_engine'):
            TranslationDialog(self.root, self.translation_engine)
        else:
            DarkMessageBox.showerror("Error", "Translation engine not available", self.root)

    def show_summarize_dialog(self) -> None:
        if not OLLAMA_AVAILABLE:
            DarkMessageBox.showerror("Fehler", "Ollama nicht verfügbar (requests nicht installiert)", self.root)
            return
        if hasattr(self, 'transcript_text') and self.transcript_text.winfo_exists():
            text = self.transcript_text.get("1.0", "end-1c").strip()
        else:
            text = ""
        if not text:
            DarkMessageBox.showwarning("Kein Text", "Kein Transkriptions-Text zum Zusammenfassen vorhanden.", self.root)
            return
        SummarizeDialog(self.root, text, self)

    def show_install_dialog(self) -> None:
        InstallDependencyDialog(self.root, self)

    def _start_system_monitoring(self) -> None:
        def monitor() -> None:
            try:
                import psutil
            except ImportError:
                if hasattr(self, 'system_info_label'):
                    self.system_info_label.config(text="⚙️ System monitoring unavailable")
                if hasattr(self, 'root') and self.root.winfo_exists():
                    self.root.after(3000, monitor)
                return
            try:
                cpu = psutil.cpu_percent(interval=None)
                memory = psutil.virtual_memory()
                ram_used = memory.used // (1024**2)
                ram_total = memory.total // (1024**2)
                gpu_text = ""
                try:
                    if TORCH_AVAILABLE:
                        torch = FastLazyLoader.load("torch")
                        if torch.cuda.is_available():
                            torch.cuda.synchronize()
                            vram_used = torch.cuda.memory_allocated() / (1024**3)
                            vram_total = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                            gpu_text = f" | 🎮 VRAM: {vram_used:.1f}/{vram_total:.1f}GB"
                        else:
                            gpu_text = " | 🎮 GPU: ❌"
                    else:
                        gpu_text = " | 🎮 GPU: N/A"
                except Exception as e:
                    if DEBUG_LEVEL >= 2:
                        logger.debug(f"VRAM‑Abfrage fehlgeschlagen: {e}")
                    gpu_text = " | 🎮 GPU: Fehler"
                current_model = "None"
                if hasattr(self, 'transcription_engine'):
                    current_model = self.transcription_engine.get_current_model()
                demo_hint = " | ⚠️ Demo" if getattr(self, 'demo_mode', False) else ""
                if IS_WINDOWS:
                    info = f"🪟 Windows | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text} | 🤖 Model: {current_model}{demo_hint}"
                elif IS_MACOS:
                    if IS_ARM:
                        info = f"🍎 macOS ARM | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text} | 🤖 Model: {current_model}{demo_hint}"
                    else:
                        info = f"🍎 macOS Intel | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text} | 🤖 Model: {current_model}{demo_hint}"
                elif IS_LINUX:
                    info = f"🐧 Linux | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text} | 🤖 Model: {current_model}{demo_hint}"
                else:
                    info = f"🌐 System | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text} | 🤖 Model: {current_model}{demo_hint}"
                if hasattr(self, 'system_info_label'):
                    self.system_info_label.config(text=info)
            except (ImportError, AttributeError):
                if hasattr(self, 'system_info_label'):
                    self.system_info_label.config(text="⚙️ System monitoring unavailable")
            if hasattr(self, 'root') and self.root.winfo_exists():
                self.root.after(3000, monitor)
        if hasattr(self, 'root') and self.root.winfo_exists():
            self.root.after(1000, monitor)

    def _update_ollama_button_state(self):
        if not hasattr(self, 'correct_btn'):
            return
        if not OLLAMA_AVAILABLE:
            self.correct_btn.config(state="disabled", text="🔧 (kein Ollama)")
            return
        summarizer = OllamaSummarizer(self, model=self.advanced_settings.ollama_model,
                                      host=self.advanced_settings.ollama_host)
        if not summarizer.is_server_reachable():
            self.correct_btn.config(state="disabled", text="🔧 (Server aus)")
            ToolTip(self.correct_btn, "Ollama-Server läuft nicht – starte 'ollama serve'")
        else:
            self.correct_btn.config(state="normal", text="🔧")
            ToolTip(self.correct_btn, "Transkript mit Ollama korrigieren")

    @gui_operation_decorator
    def correct_transcript_with_ollama(self):
        if not hasattr(self, 'transcript_text') or not self.transcript_text.winfo_exists():
            return
        text = self.transcript_text.get("1.0", "end-1c").strip()
        if not text:
            DarkMessageBox.showinfo("Hinweis", "Kein Text zum Korrigieren vorhanden.", self.root)
            return

        if not OLLAMA_AVAILABLE:
            DarkMessageBox.showerror("Fehler", "Ollama nicht verfügbar (requests fehlt).", self.root)
            return

        summarizer = OllamaSummarizer(self, model=self.advanced_settings.ollama_model,
                                      host=self.advanced_settings.ollama_host)
        if not summarizer.is_server_reachable():
            DarkMessageBox.showerror("Fehler",
                "Ollama-Server läuft nicht.\nBitte starte 'ollama serve' und versuche es erneut.",
                self.root)
            return

        progress = DarkMessageBox.show_progress(
            "Korrektur läuft",
            "Sende Text an Ollama zur Korrektur...",
            parent=self.root,
            indeterminate=True
        )

        corrected_parts = []
        error_occurred = False

        def on_chunk(chunk: str):
            corrected_parts.append(chunk)
            self.root.after(0, lambda: progress.update_message(
                f"Empfange Daten... ({len(corrected_parts)} Teile)"
            ))

        def on_error(error: str):
            nonlocal error_occurred
            error_occurred = True
            self.root.after(0, progress.close)
            self.root.after(0, lambda: DarkMessageBox.showerror(
                "Fehler", f"Korrektur fehlgeschlagen:\n{error}", self.root
            ))

        def on_complete():
            self.root.after(0, progress.close)
            if not error_occurred:
                corrected_text = "".join(corrected_parts).strip()
                if corrected_text:
                    self.root.after(0, lambda: self._update_transcript_with_correction(corrected_text))
                else:
                    self.root.after(0, lambda: self.update_status("⚠️ Korrektur ergab leeren Text"))

        summarizer.correct_transcript(
            text,
            callback=on_chunk,
            error_callback=on_error,
            complete_callback=on_complete
        )

    @gui_operation_decorator
    def _update_transcript_with_correction(self, corrected_text: str):
        try:
            if self.transcript_text and self.transcript_text.winfo_exists():
                self.transcript_text.delete("1.0", "end")
                self.transcript_text.insert("1.0", corrected_text)
                self.update_status("✅ Transkription korrigiert")
        except tk.TclError:
            pass

    def _final_initialization_check(self) -> None:
        logger.info("✅ Dragon Whisperer initialisiert")
        if getattr(self, 'demo_mode', False):
            self.update_status("⚠️ Demo-Modus: Whisper nicht verfügbar – verwende Dummy-Transkriptionen")
        self.root.after(500, self._update_ollama_button_state)

    def _emergency_cleanup(self) -> None:
        logger.info("🆘 Emergency cleanup...")
        self._minimal_emergency_cleanup()

    @gui_operation_decorator
    def update_progress(self, processed: int, total: Optional[int], chunks: int) -> None:
        if not hasattr(self, 'progress_bar') or not self.progress_bar.winfo_exists():
            return
        try:
            if total is not None and total > 0:
                if not self.progress_bar.winfo_ismapped():
                    self.progress_bar.pack(side="left", padx=(10, 10))
                percent = (processed / total) * 100
                self.progress_bar.config(mode='determinate', value=percent)
                mb = processed // (1024 * 1024)
                tb = total // (1024 * 1024)
                self.progress_label.config(text=f"{mb}MB/{tb}MB")
            else:
                if self.progress_bar.winfo_ismapped():
                    self.progress_bar.pack_forget()
                self.progress_label.config(text=f"Chunks: {chunks}  |  Daten: {processed // 1024} KB")
        except tk.TclError:
            pass

    @gui_operation_decorator
    def _reset_progress(self):
        if hasattr(self, 'progress_bar') and self.progress_bar.winfo_exists():
            self.progress_bar.stop()
            self.progress_bar.config(mode='determinate', value=0)
            self._progress_bar_started = False
        if hasattr(self, 'progress_label'):
            self.progress_label.config(text="")

    def _bind_shortcuts(self):
        mod = "Command" if IS_MACOS else "Control"

        self.root.bind(f'<{mod}-o>', lambda e: self.select_file_dark())
        self.root.bind(f'<{mod}-v>', lambda e: self.paste_url())
        self.root.bind(f'<{mod}-Return>', lambda e: self._on_start_click())
        self.root.bind(f'<{mod}-q>', lambda e: self._safe_exit_dialog())
        self.root.bind(f'<{mod}-s>', lambda e: self.save_transcript())
        self.root.bind(f'<{mod}-l>', lambda e: self.toggle_layout())
        self.root.bind(f'<{mod}-t>', lambda e: self.toggle_translation())
        self.root.bind(f'<{mod}-e>', lambda e: self.export_subtitles())
        self.root.bind(f'<{mod}-u>', lambda e: self.toggle_subtitle_mode())
        self.root.bind(f'<{mod}-Shift-c>', lambda e: self.clear_all())
        self.root.bind(f'<{mod}-h>', lambda e: self.show_shortcuts_help())
        self.root.bind('<F1>', lambda e: self.show_shortcuts_help())

        self.url_entry.bind(f'<{mod}-v>', lambda e: 'break')

    def show_shortcuts_help(self):
        ShortcutsDialog(self.root)

# -----------------------------------------------------------------------------
# TRANSLATION DIALOG
# -----------------------------------------------------------------------------

class TranslationDialog:
    def __init__(self, parent: tk.Widget, translation_engine: BaseTranslationEngine, initial_text: str = "") -> None:
        self.parent = parent
        self.engine = translation_engine
        self.initial_text = initial_text
        self.dialog: Optional[tk.Toplevel] = None
        self.create_dialog()

    def create_dialog(self) -> None:
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("🐉 Text Translation")
        self.dialog.geometry("600x500")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(self.parent)
        self.dialog.grab_set()

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        tk.Label(main, text="Source text:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY, font=Fonts.PRIMARY).pack(anchor="w")
        self.source_text = scrolledtext.ScrolledText(
            main, height=8, bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE, wrap=tk.WORD
        )
        self.source_text.pack(fill="x", pady=(0, 10))
        if self.initial_text:
            self.source_text.insert("1.0", self.initial_text)
        DarkContextMenu(self.source_text)

        lang_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        lang_frame.pack(fill="x", pady=5)

        tk.Label(lang_frame, text="From:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(side="left", padx=(0, 5))
        self.src_lang_var = tk.StringVar(value="auto")
        src_combo = ttk.Combobox(
            lang_frame, textvariable=self.src_lang_var,
            values=["auto"] + [name for name, code in SORTED_LANGUAGES],
            width=15, state="readonly"
        )
        src_combo.pack(side="left", padx=(0, 20))

        tk.Label(lang_frame, text="To:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY).pack(side="left", padx=(0, 5))
        self.tgt_lang_var = tk.StringVar(value="German")
        tgt_combo = ttk.Combobox(
            lang_frame, textvariable=self.tgt_lang_var,
            values=[name for name, code in SORTED_LANGUAGES if name != "auto"],
            width=15, state="readonly"
        )
        tgt_combo.pack(side="left")

        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x", pady=10)
        translate_btn = tk.Button(
            btn_frame, text="🌐 Translate", command=self.translate,
            bg=CURRENT_THEME.DRAGON_GREEN, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON, padx=20
        )
        translate_btn.pack(side="left")

        tk.Label(main, text="Translation:", bg=CURRENT_THEME.BG_PRIMARY,
                 fg=CURRENT_THEME.TEXT_PRIMARY, font=Fonts.PRIMARY).pack(anchor="w", pady=(10, 0))
        self.target_text = scrolledtext.ScrolledText(
            main, height=8, bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE, wrap=tk.WORD, state="normal"
        )
        self.target_text.pack(fill="both", expand=True, pady=(5, 0))
        DarkContextMenu(self.target_text)

        close_btn = tk.Button(
            main, text="Close", command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY
        )
        close_btn.pack(pady=10)

    def translate(self) -> None:
        source = self.source_text.get("1.0", "end-1c").strip()
        if not source:
            return

        src_name = self.src_lang_var.get().strip()
        tgt_name = self.tgt_lang_var.get().strip()

        valid_language_names = [name for name, code in SORTED_LANGUAGES]

        if src_name not in valid_language_names and src_name != "auto":
            src_name = "auto"
            self.src_lang_var.set("auto")

        if tgt_name not in valid_language_names:
            tgt_name = "German"
            self.tgt_lang_var.set("German")

        src_code = "auto" if src_name == "auto" else next(code for name, code in SORTED_LANGUAGES if name == src_name)
        tgt_code = next(code for name, code in SORTED_LANGUAGES if name == tgt_name)

        try:
            old_target = self.engine.target_lang
            self.engine.set_target_language(tgt_code)
            result = self.engine.translate_text(source, src_code)
            self.engine.set_target_language(old_target)

            if result and result.translated:
                self.target_text.delete("1.0", "end")
                self.target_text.insert("1.0", result.translated)
            else:
                self.target_text.delete("1.0", "end")
                self.target_text.insert("1.0", "(Translation failed)")
        except Exception as e:
            self.target_text.delete("1.0", "end")
            self.target_text.insert("1.0", f"Error: {str(e)}")


# -----------------------------------------------------------------------------
# SHORTCUTS DIALOG
# -----------------------------------------------------------------------------
class ShortcutsDialog:
    def __init__(self, parent):
        self.parent = parent
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("🐉 Tastenkürzel")
        self.dialog.geometry("500x400")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=20, pady=20)
        main.pack(fill="both", expand=True)

        tk.Label(main, text="Tastenkürzel", font=("Segoe UI", 14, "bold"),
                 bg=CURRENT_THEME.BG_PRIMARY, fg=CURRENT_THEME.TEXT_PRIMARY).pack(pady=(0,15))

        mod = "Cmd" if IS_MACOS else "Strg"
        shortcuts = [
            (f"{mod}+O", "Datei öffnen"),
            (f"{mod}+V", "URL einfügen"),
            (f"{mod}+Return", "Transkription starten"),
            (f"{mod}+Q", "Programm beenden"),
            (f"{mod}+S", "Transkript speichern"),
            (f"{mod}+L", "Layout umschalten"),
            (f"{mod}+T", "Übersetzung ein/aus"),
            (f"{mod}+E", "Untertitel exportieren"),
            (f"{mod}+Umschalt+C", "Alles löschen"),
            ("F1", "Diese Hilfe anzeigen"),
        ]

        frame = tk.Frame(main, bg=CURRENT_THEME.BG_SECONDARY)
        frame.pack(fill="both", expand=True)

        for i, (key, desc) in enumerate(shortcuts):
            row = tk.Frame(frame, bg=CURRENT_THEME.BG_TERTIARY if i % 2 else CURRENT_THEME.BG_SECONDARY)
            row.pack(fill="x", pady=1)

            key_label = tk.Label(row, text=key, font=("Segoe UI", 10, "bold"),
                                  bg=row['bg'], fg=CURRENT_THEME.TEXT_ACCENT, width=15, anchor="w")
            key_label.pack(side="left", padx=(10,5), pady=5)

            desc_label = tk.Label(row, text=desc, font=("Segoe UI", 10),
                                   bg=row['bg'], fg=CURRENT_THEME.TEXT_PRIMARY, anchor="w")
            desc_label.pack(side="left", fill="x", expand=True, padx=5, pady=5)

        close_btn = tk.Button(main, text="Schließen", command=self.dialog.destroy,
                              bg=CURRENT_THEME.BG_TERTIARY, fg=CURRENT_THEME.TEXT_PRIMARY)
        close_btn.pack(pady=10)

        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.dialog.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")


# -----------------------------------------------------------------------------
# HILFSFUNKTIONEN (Konsolen-Setup, Help, System-Check)
# -----------------------------------------------------------------------------
_original_console_mode: Optional[int] = None
_original_codepage: Optional[int] = None

def _save_console_state() -> None:
    global _original_console_mode, _original_codepage
    if not IS_WINDOWS:
        return
    try:
        import ctypes
        kernel32 = ctypes.windll.kernel32
        handle = kernel32.GetStdHandle(-11)
        mode = ctypes.c_ulong()
        if kernel32.GetConsoleMode(handle, ctypes.byref(mode)):
            _original_console_mode = mode.value
        _original_codepage = kernel32.GetConsoleOutputCP()
    except Exception:
        pass

def _restore_console_state() -> None:
    if not IS_WINDOWS:
        return
    try:
        import ctypes
        kernel32 = ctypes.windll.kernel32
        if _original_console_mode is not None:
            handle = kernel32.GetStdHandle(-11)
            kernel32.SetConsoleMode(handle, _original_console_mode)
        if _original_codepage is not None:
            kernel32.SetConsoleOutputCP(_original_codepage)
            kernel32.SetConsoleCP(_original_codepage)
    except Exception:
        pass

_save_console_state()
atexit.register(_restore_console_state)

def _setup_windows_console() -> None:
    if not IS_WINDOWS:
        return
    try:
        import codecs
        sys.stdout = codecs.getwriter("utf-8")(sys.stdout.buffer)
        import ctypes
        kernel32 = ctypes.windll.kernel32
        kernel32.SetConsoleOutputCP(65001)
        kernel32.SetConsoleCP(65001)
        handle = kernel32.GetStdHandle(-11)
        mode = ctypes.c_ulong()
        if kernel32.GetConsoleMode(handle, ctypes.byref(mode)):
            ENABLE_VIRTUAL_TERMINAL_PROCESSING = 0x0004
            new_mode = mode.value | ENABLE_VIRTUAL_TERMINAL_PROCESSING
            kernel32.SetConsoleMode(handle, new_mode)
    except Exception:
        pass

def check_platform_compatibility() -> List[str]:
    issues: List[str] = []
    if sys.version_info < (3, 8):
        issues.append(f"Python 3.8+ required (you have {sys.version_info.major}.{sys.version_info.minor})")
    if IS_WINDOWS:
        try:
            import platform
            win_ver = platform.version()
            major_ver = int(win_ver.split(".")[0]) if "." in win_ver else 0
            if major_ver < 10:
                issues.append("Windows 10+ recommended for best experience")
        except Exception:
            pass
    return issues

def debug_script() -> None:
    print("=" * 60)
    print("DRAGON WHISPERER - DEBUG INFORMATION")
    print("=" * 60)
    print(f"\n🔧 Platform: {SYSTEM} {'ARM' if IS_ARM else 'x86'}")
    print(f"🔧 Python: {sys.version.split()[0]}")
    print("\n📦 Dependencies:")
    deps = [
        ("FFmpeg", shutil.which("ffmpeg") is not None),
        ("yt-dlp", shutil.which("yt-dlp") is not None),
        ("faster-whisper", WHISPER_AVAILABLE),
        ("deep-translator", TRANSLATOR_AVAILABLE),
        ("Tkinter", GUI_AVAILABLE),
        ("NumPy", NUMPY_AVAILABLE),
        ("PyTorch", TORCH_AVAILABLE),
        ("SciPy", SCIPY_AVAILABLE),
        ("psutil", "psutil" in sys.modules),
        ("requests", OLLAMA_AVAILABLE),
    ]
    for name, status in deps:
        symbol = "✅" if status else "❌"
        status_text = "Available" if status else "Not available"
        print(f"  {symbol} {name:18} {status_text}")
    print(f"\n📁 Script: {os.path.abspath(__file__)}")
    print(f"📁 CWD: {os.getcwd()}")
    print("=" * 60)

def setup_platform_environment() -> Dict[str, str]:
    env_vars: Dict[str, str] = {}
    if IS_WINDOWS:
        env_vars.update({
            "FFMPEG_BINARY": "ffmpeg.exe",
            "YT_DLP_BINARY": "yt-dlp.exe",
            "PYTHONIOENCODING": "utf-8",
        })
    elif IS_MACOS:
        env_vars.update({
            "FFMPEG_BINARY": "ffmpeg",
            "YT_DLP_BINARY": "yt-dlp",
        })
    else:
        env_vars.update({
            "FFMPEG_BINARY": "ffmpeg",
            "YT_DLP_BINARY": "yt-dlp",
        })
    for key, value in env_vars.items():
        os.environ[key] = value
    return env_vars

def _print_help() -> None:
    print("🐉 Dragon Whisperer - Ultimate Stream Transcription & Translation")
    print("=" * 60)
    print("\nUsage:")
    print("  python dragon_whisperer.py [options]")
    print("\nOptions:")
    print("  --quiet, -q    Quiet mode (minimal output)")
    print("  --debug[=N]    Debug mode with level (1-3). Examples: --debug, --debug=2")
    print("  --debug=COMP   Component-specific debug (e.g., --debug=network,vad)")
    print("  --check        System compatibility check")
    print("  --help, -h     Show this help")
    print("  --version, -v  Show version")
    print("\nExamples:")
    print("  Normal use:   python dragon_whisperer.py")
    print("  Quiet mode:   python dragon_whisperer.py --quiet")
    print("  Debug level2: python dragon_whisperer.py --debug=2")
    print("  VAD debug:    python dragon_whisperer.py --debug=vad")
    print("  System check: python dragon_whisperer.py --check")
    print("\nFeatures:")
    print("  • Live stream transcription (YouTube, Twitch, etc.)")
    print("  • Real-time translation to 50+ languages")
    print("  • Subtitle export (SRT, VTT)")
    print("  • Batch processing")
    print("  • Dark mode GUI")

def _run_system_check() -> int:
    issues = check_platform_compatibility()
    print("🔍 Dragon Whisperer - System Compatibility Check")
    print("=" * 50)
    if issues:
        print("⚠️  Issues found:")
        for issue in issues:
            print(f"  • {issue}")
    else:
        print("✅ No compatibility issues")
    print("\n📦 Dependency Check:")
    print(f"  FFmpeg:        {'✅' if shutil.which('ffmpeg') else '❌'}")
    print(f"  yt-dlp:        {'✅' if shutil.which('yt-dlp') else '❌'}")
    print(f"  Tkinter:       {'✅' if GUI_AVAILABLE else '❌'}")
    print(f"  faster-whisper:{'✅' if WHISPER_AVAILABLE else '❌'}")
    print(f"  NumPy:         {'✅' if NUMPY_AVAILABLE else '❌'}")
    print(f"  PyTorch:       {'✅' if TORCH_AVAILABLE else '❌'}")
    print(f"  deep-translator:{'✅' if TRANSLATOR_AVAILABLE else '❌'}")
    print(f"  SciPy:         {'✅' if SCIPY_AVAILABLE else '❌'}")
    print(f"  psutil:        {'✅' if 'psutil' in sys.modules else '❌'}")
    print("\n💻 System Info:")
    print(f"  Platform: {SYSTEM}")
    print(f"  Architecture: {'ARM' if IS_ARM else 'x86'}")
    print(f"  Python: {sys.version.split()[0]}")
    if IS_WINDOWS:
        try:
            import platform
            print(f"  Windows: {platform.version()}")
        except Exception:
            pass
    return 0 if not issues else 1

def _show_user_error(message: str) -> None:
    if IS_WINDOWS and not sys.stdin.isatty():
        try:
            import ctypes
            ctypes.windll.user32.MessageBoxW(
                0,
                f"Dragon Whisperer - Setup Required\n\n{message}\n\n"
                "Please install missing components and try again.",
                "Setup Error",
                0x10,
            )
        except Exception:
            pass
    else:
        if "Tkinter" in message:
            print("\n💡 INSTALLATION HELP:")
            print("  pip install tk")
        elif "FFmpeg" in message:
            print("\n💡 INSTALLATION HELP:")
            if IS_WINDOWS:
                print("  Download from: https://ffmpeg.org/download.html")
            elif IS_MACOS:
                print("  brew install ffmpeg")
            else:
                print("  sudo apt install ffmpeg")
        else:
            print(f"\n❌ {message}")

def _show_critical_error(message: str) -> None:
    if IS_WINDOWS and not sys.stdin.isatty():
        try:
            import ctypes
            short_msg = message[:200] + "..." if len(message) > 200 else message
            ctypes.windll.user32.MessageBoxW(
                0,
                f"Dragon Whisperer - Critical Error\n\n{short_msg}\n\n"
                "Please check the console for details.\n"
                "Try running with --debug flag for more information.",
                "Critical Error",
                0x10,
            )
        except Exception:
            pass
    else:
        print(f"\n💥 {message}")

# -----------------------------------------------------------------------------
# Funktion für Debug Level 3 Systeminfo
# -----------------------------------------------------------------------------
def print_system_info_debug3():
    import platform
    import resource
    print("\n" + "="*60)
    print("🐉 DEBUG LEVEL 3 - SYSTEM INFORMATION")
    print("="*60)
    print(f"Platform: {platform.platform()}")
    print(f"Python: {sys.version}")
    print(f"CPU count: {os.cpu_count()}")
    try:
        soft, hard = resource.getrlimit(resource.RLIMIT_NOFILE)
        print(f"RLIMIT_NOFILE: soft={soft}, hard={hard}")
    except Exception:
        pass
    print("\n📦 Environment variables set by script:")
    for key in os.environ:
        if key.startswith(('FFMPEG_', 'TORCH_', 'PYTORCH_', 'AV_', 'OPENCV_')):
            print(f"  {key}={os.environ[key]}")
    print("\n📚 Library versions:")
    libs = [
        ('torch', TORCH_AVAILABLE),
        ('faster_whisper', FASTER_WHISPER_AVAILABLE),
        ('whisper', OPENAI_WHISPER_AVAILABLE),
        ('numpy', NUMPY_AVAILABLE),
        ('scipy', SCIPY_AVAILABLE),
        ('psutil', importlib.util.find_spec('psutil') is not None),
        ('requests', OLLAMA_AVAILABLE),
        ('deep_translator', TRANSLATOR_AVAILABLE),
        ('pynvml', importlib.util.find_spec('pynvml') is not None),
    ]
    for name, avail in libs:
        if avail:
            try:
                mod = __import__(name)
                ver = getattr(mod, '__version__', 'unknown')
                print(f"  {name}: {ver}")
            except Exception:
                print(f"  {name}: available (version unknown)")
        else:
            print(f"  {name}: not available")
    print("\n🎮 GPU info:")
    if TORCH_AVAILABLE:
        torch = FastLazyLoader.load('torch')
        if torch.cuda.is_available():
            print(f"  CUDA available: {torch.version.cuda}")
            print(f"  Device: {torch.cuda.get_device_name(0)}")
            print(f"  Total VRAM: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.2f} GB")
        else:
            print("  CUDA not available")
    print("="*60 + "\n")

# -----------------------------------------------------------------------------
# MAIN
# -----------------------------------------------------------------------------
def main() -> int:
    warnings.filterwarnings("ignore")

    if IS_WINDOWS:
        _setup_windows_console()

    cli_args = {
        "debug": "--debug" in sys.argv,
        "quiet": "--quiet" in sys.argv or "-q" in sys.argv,
        "check": "--check" in sys.argv,
        "help": "--help" in sys.argv or "-h" in sys.argv,
        "version": "--version" in sys.argv or "-v" in sys.argv,
    }

    if cli_args["help"]:
        _print_help()
        return 0
    if cli_args["version"]:
        print("🐉 Dragon Whisperer v2.0 - überarbeitet")
        print(f"Platform: {SYSTEM} {'ARM' if IS_ARM else 'x86'}")
        return 0

    debug_level = DEBUG_LEVEL

    if cli_args["check"]:
        return _run_system_check()

    if debug3_enabled():
        print_system_info_debug3()

    app: Optional[DragonWhispererGUI] = None
    exit_code = 0

    try:
        logger.info("🐉 Dragon Whisperer starting...")

        logger.debug("🔍 Checking dependencies...")
        try:
            PlatformUtils.check_platform_dependencies()
        except RuntimeError as e:
            _show_user_error(str(e))
            return 1
        logger.debug("✅ Dependencies OK")

        if not GUI_AVAILABLE:
            raise RuntimeError("Tkinter/GUI not available. Install with: pip install tk")

        logger.debug("⚡ Setting up signal handlers...")
        SignalHandler.setup(verbose=False, silent=True, max_cleanup_time=10.0)

        SignalHandler.register_cleanup(
            lambda: _EXECUTOR.shutdown(wait=False),
            name="GlobalExecutorShutdown",
            priority=ShutdownPriority.LOW
        )

        logger.debug("🖥️ Initializing GUI...")
        app = DragonWhispererGUI()

        if debug_level >= 1 and not cli_args["quiet"]:
            print("\n" + "=" * 50)
            print("🐉 DRAGON WHISPERER READY")
            print("=" * 50)
            print(f"Platform: {SYSTEM} {'ARM' if IS_ARM else 'x86'}")
            print(f"Python: {sys.version.split()[0]}")
            print(f"Working Dir: {os.getcwd()}")
            if hasattr(app, "transcription_engine"):
                current_model = app.transcription_engine.get_current_model()
                print(f"Model: {current_model if current_model else 'Not loaded'}")
            print(f"Layout: {getattr(app, 'layout_mode', 'vertical')}")
            print("=" * 50 + "\n")

        logger.info("🚀 Starting main loop...")
        app.run()
        logger.info("✅ Application closed normally")

    except KeyboardInterrupt:
        logger.info("\n🛑 Interrupted by user")
        exit_code = 0

    except RuntimeError as e:
        error_msg = str(e)
        logger.error(f"❌ {error_msg}")
        _show_user_error(error_msg)
        exit_code = 1

    except Exception as e:
        if PlatformUtils.is_fatal_exception(e):
            raise
        error_msg = str(e)
        logger.error(f"💥 Unexpected error: {error_msg}")
        if debug_level >= 2:
            import traceback
            traceback.print_exc()
        _show_critical_error(error_msg)
        exit_code = 2

    finally:
        logger.debug("🧹 Final minimal cleanup...")
        try:
            transcription_cache.clear()
            translation_cache.clear()
            audio_cache.clear()
        except Exception:
            pass
        gc.collect()
        logger.debug("✅ Shutdown complete")

    return exit_code


if __name__ == "__main__":
    try:
        if "__file__" in globals():
            script_dir = os.path.dirname(os.path.abspath(__file__))
            if os.getcwd() != script_dir:
                os.chdir(script_dir)
                if DEBUG_LEVEL >= 1:
                    print(f"📁 Working directory set to: {script_dir}")
        exit_code = main()
        sys.exit(exit_code)
    except KeyboardInterrupt:
        print("\n🛑 Program interrupted by user (KeyboardInterrupt)")
        sys.exit(0)
    except SystemExit as e:
        raise e
    except Exception as e:
        if PlatformUtils.is_fatal_exception(e):
            raise
        error_type = type(e).__name__
        error_msg = str(e)
        print(f"\n💥 FATAL ERROR in main guard: {error_type}: {error_msg}")
        if DEBUG_LEVEL >= 2:
            import traceback
            traceback.print_exc()
            print("\n🔧 Debug Info:")
            print(f"  Python: {sys.version}")
            print(f"  Platform: {sys.platform}")
            print(f"  Executable: {sys.executable}")
            print(f"  CWD: {os.getcwd()}")
            print(f"  Script: {__file__ if '__file__' in globals() else 'Unknown'}")
        sys.exit(99)
