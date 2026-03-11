#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""🐉 DRAGON WHISPERER – (v2.1)"""

# =============================================================================
# 1. IMPORTS – Standardbibliothek, Drittanbieter
# =============================================================================
import atexit
import gc
import hashlib
import importlib
import importlib.util
import json
import logging
import os
import platform
import queue
import re
import resource
import shutil
import signal as py_signal
import subprocess
import sys
import tempfile
import threading
import time
import urllib.parse
import urllib.request
import warnings
import weakref
from abc import ABC, abstractmethod
from collections import OrderedDict, deque
from concurrent.futures import Future, ThreadPoolExecutor
from concurrent.futures import TimeoutError as FutureTimeout
from concurrent.futures import as_completed
from dataclasses import asdict, dataclass, field, fields
from datetime import datetime
from enum import Enum, auto
from functools import wraps
from pathlib import Path
from typing import (Any, Callable, Deque, Dict, List, Optional, Set, Tuple,
                    TypeVar, Union, Literal)

# Drittanbieter (optional)
try:
    import requests
except ImportError:
    requests = None

# Tkinter (GUI)
try:
    import tkinter as tk
    from tkinter import filedialog, scrolledtext, ttk

    GUI_AVAILABLE = True
except ImportError:
    GUI_AVAILABLE = False
    tk = None
    ttk = None
    scrolledtext = None
    filedialog = None

# =============================================================================
# 2. GLOBALE DEBUG-KONFIGURATION
# =============================================================================
DEBUG_LEVEL = 0
DEBUG_COMPONENTS = []

for arg in sys.argv:
    if arg == "--debug":
        DEBUG_LEVEL = max(DEBUG_LEVEL, 1)
    elif arg.startswith("--debug="):
        value = arg.split("=", 1)[1]
        if value.isdigit():
            DEBUG_LEVEL = max(DEBUG_LEVEL, int(value))
        else:
            DEBUG_COMPONENTS.extend(value.split(","))

QUIET_MODE = "--quiet" in sys.argv or "-q" in sys.argv

# -----------------------------------------------------------------------------
# Logging Setup
# -----------------------------------------------------------------------------
logging.basicConfig(
    level=logging.WARNING,
    format="[%(asctime)s.%(msecs)03d] [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("dragon")

logging.getLogger("huggingface_hub").setLevel(logging.INFO)
logging.getLogger("faster_whisper").setLevel(logging.INFO)

for lib in ["httpx", "urllib3", "httpcore"]:
    logging.getLogger(lib).setLevel(logging.WARNING)

if DEBUG_LEVEL >= 1:
    logger.setLevel(logging.DEBUG)
    logging.getLogger().setLevel(logging.DEBUG)
if QUIET_MODE:
    logger.setLevel(logging.ERROR)

warnings.filterwarnings("ignore", message=".*pynvml.*")
warnings.filterwarnings("ignore", message=".*The pynvml package is deprecated.*")

os.environ.update(
    {
        "PYTHONWARNINGS": "default",
        "TORCH_DISABLE_CUDA_WARNINGS": "1",
        "TORCH_CPP_LOG_LEVEL": "0",
        "PYTORCH_JIT": "0",
    }
)

SYSTEM = platform.system()
IS_WINDOWS = SYSTEM == "Windows"
IS_MACOS = SYSTEM == "Darwin"
IS_LINUX = SYSTEM == "Linux"

try:
    machine = platform.machine().lower()
    IS_ARM = machine in ("arm64", "aarch64", "armv8l", "arm")
    IS_X86 = machine in ("x86_64", "amd64", "i386", "i686", "x86")
except Exception:
    IS_ARM = False
    IS_X86 = True

logger.info(
    f"🐉 Dragon Whisperer - Platform: {SYSTEM} "
    f"{'ARM' if IS_ARM else 'x86'} (Debug-Level: {DEBUG_LEVEL})"
)

os.environ.update(
    {
        "FFMPEG_DISABLE_RKMPP": "1",
        "AV_DISABLE_RKMPP": "1",
        "FFMPEG_DISABLE_VAAPI": "0" if IS_LINUX else "1",
        "FFMPEG_DISABLE_VDPAU": "0" if IS_LINUX else "1",
        "OPENCV_LOG_LEVEL": "ERROR",
        "GST_DEBUG": "0",
    }
)

if IS_WINDOWS:
    os.environ.update({"PYTHONIOENCODING": "utf-8"})


# =============================================================================
# 3. KONSTANTEN & KONFIGURATION (Optimiert für maximale Präzision)
# =============================================================================
class Constants:
    """Zentrale Konstanten für das gesamte Programm."""

    # Audio
    SAMPLE_RATE: int = 16000
    CHANNELS: int = 1
    AUDIO_FORMAT: str = "s16le"
    BYTES_PER_SAMPLE: int = 2

    # Chunking – optimiert: größere Chunks und mehr Overlap für bessere Genauigkeit
    BASE_CHUNK_DURATION: int = 20
    CHUNK_OVERLAP: float = 2.0
    MIN_CHUNK_DURATION: int = 2
    MAX_CHUNK_DURATION: int = 30

    # Prozesse & Timeouts
    MAX_SUBPROCESSES: int = 8
    SUBPROCESS_TIMEOUT: int = 60
    GUI_OPERATION_TIMEOUT: float = 10.0
    MEMORY_CHECK_INTERVAL: int = 15
    MAX_GUI_UPDATES_PER_SECOND: int = 40
    MAX_MEMORY_USAGE: int = 8 * 1024 * 1024 * 1024
    MAX_CACHE_SIZE: int = 500
    MAX_TEXT_LINES: int = 2000
    DEFAULT_BEAM_SIZE: int = 10        # leicht erhöht für mehr Genauigkeit
    DEFAULT_TEMPERATURE: float = 0.0
    ENABLE_VAD_FILTER: bool = True
    MAX_CONSECUTIVE_ERRORS: int = 5

    # Stream
    STREAM_TIMEOUT: int = 25
    INITIAL_BUFFER_SECONDS: float = 2.0
    MAX_EMPTY_READS: int = 30
    RECONNECT_DELAY: int = 2
    READ_RETRY_DELAY: float = 0.1
    YOUTUBE_TIMEOUT: int = 10000000
    NORMAL_TIMEOUT: int = 30000000
    MAX_STREAM_RECONNECTS: int = 5

    # FFmpeg
    FFMPEG_BUFSIZE: str = "2048k"
    FFMPEG_THREADS: int = 1
    FFMPEG_PROBESIZE: str = "32"
    FFMPEG_ANALYZE_DURATION: str = "0"

    # Audio-Filter (Sprachoptimierung) – für Präzision minimal gehalten
    AUDIO_FILTER: str = "aresample=16000,volume=1.5,dynaudnorm"
    LANGUAGE_FILTERS: Dict[str, str] = {
        "ko": "aresample=16000,volume=2.0,highpass=f=80,lowpass=f=3800,afftdn=nf=-15",
        "ja": "aresample=16000,volume=2.0,highpass=f=90,lowpass=f=3700,afftdn=nf=-15",
        "zh": "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3500,afftdn=nf=-20",
        "de": "aresample=16000,volume=1.8,highpass=f=100,lowpass=f=3200,dynaudnorm",
        "en": "aresample=16000,volume=1.8,highpass=f=80,lowpass=f=3400,dynaudnorm",
        "fr": "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3300,dynaudnorm",
        "es": "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3400,dynaudnorm",
    }
    FILTER_PROFILES: Dict[str, str] = {
        "transcription": "aresample=16000,volume=1.5,dynaudnorm",
        "translation": "aresample=16000,volume=2.0,highpass=f=100,lowpass=f=3400",
        "realtime": "aresample=16000,volume=1.8,dynaudnorm",
        "noisy": "aresample=16000,volume=2.5,highpass=f=150,lowpass=f=3000,afftdn=nf=-30",
        "music": "aresample=16000,volume=1.5,highpass=f=50,lowpass=f=5000",
        "podcast": "aresample=16000,volume=2.0,highpass=f=80,lowpass=f=3500",
    }

    # Audio-Enhancement
    AUDIO_ENHANCEMENT_ENABLED: bool = True
    MIN_RMS_THRESHOLD: float = 0.002
    TARGET_RMS: float = 0.2
    MAX_GAIN: float = 5.0
    CLIPPING_THRESHOLD: float = 0.9

    # Duplikaterkennung (deaktiviert für maximale Präzision)
    DUPLICATE_CHECK_ENABLED: bool = False
    RECENT_TRANSCRIPTIONS_SIZE: int = 10
    MIN_TEXT_LENGTH: int = 3
    MIN_UNIQUE_WORDS_RATIO: float = 0.3

    # Untertitel
    SUBTITLE_BUFFER_SIZE: int = 1000
    ENABLE_TIMED_TRANSCRIPTIONS: bool = True
    ENABLE_TIMED_TRANSLATIONS: bool = True

    # Logging
    ENABLE_DEBUG_LOGGING: bool = True
    LOG_CHUNK_PROCESSING: bool = False
    LOG_AUDIO_STATS: bool = True
    LOG_PERFORMANCE: bool = True
    LOG_STREAM_EVENTS: bool = True
    PERFORMANCE_LOG_INTERVAL: int = 50

    # Cache
    MAX_CACHE_SIZE_MB: int = 100
    CACHE_ENABLED: bool = True

    # Cache-Größen (TTL in Sekunden)
    TRANSCRIPTION_CACHE_SIZE: int = 200
    TRANSCRIPTION_CACHE_TTL: int = 300
    TRANSLATION_CACHE_SIZE: int = 500
    TRANSLATION_CACHE_TTL: int = 3600
    AUDIO_CACHE_SIZE: int = 128
    AUDIO_CACHE_TTL: int = 1800

    # VAD (Standardwerte) – deaktiviert für maximale Präzision
    VAD_THRESHOLD: float = 0.25
    VAD_MIN_SPEECH_DURATION_MS: int = 225
    VAD_MIN_SILENCE_DURATION_MS: int = 80

    # Sprachspezifische VAD-Anpassungen (für asiatische Sprachen) – ignoriert, da VAD aus
    LANGUAGE_VAD: Dict[str, Dict[str, Any]] = {
        "ja": {"threshold": 0.3, "min_speech_ms": 300, "min_silence_ms": 100},
        "ko": {"threshold": 0.3, "min_speech_ms": 300, "min_silence_ms": 100},
        "zh": {"threshold": 0.3, "min_speech_ms": 300, "min_silence_ms": 100},
        "th": {"threshold": 0.3, "min_speech_ms": 300, "min_silence_ms": 100},
        "vi": {"threshold": 0.3, "min_speech_ms": 250, "min_silence_ms": 90},
    }

    # Pfad-Validierung (Sicherheit)
    ALLOWED_FILE_SCHEME_PREFIX: str = "file://"
    ALLOWED_FILE_BASE_DIRS: List[str] = [
        str(Path.home()),
        os.getcwd(),
    ]

    # URL-Zeichen Whitelist (für subprocess-Aufrufe)
    URL_ALLOWED_CHARS: str = r"a-zA-Z0-9\-._~:/?#\[\]@!$&'()*+,;=%"

    # YouTube-Header
    YOUTUBE_HEADERS: Dict[str, str] = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": "https://www.youtube.com/",
        "Origin": "https://www.youtube.com",
        "Accept": "*/*",
        "Accept-Language": "en-US,en;q=0.9",
    }

    # Plattform-Konfiguration
    PLATFORM_CONFIG: Dict[str, Dict[str, Any]] = {
        "windows": {
            "ffmpeg_flags": ["-reconnect", "1", "-reconnect_streamed", "1"],
            "process_creation_flags": 0x08000000,
        },
        "macos": {
            "ffmpeg_flags": ["-reconnect", "1", "-reconnect_on_network_error", "1"],
            "start_new_session": True,
        },
        "linux": {
            "ffmpeg_flags": ["-reconnect", "1", "-reconnect_streamed", "1"],
            "start_new_session": True,
        },
    }

    # -------------------------------------------------------------------------
    # Konstanten für Timeouts
    # -------------------------------------------------------------------------
    YOUTUBE_IDLE_TIMEOUT: int = 25
    READ_CHUNK_TIMEOUT: float = 0.5
    STREAM_RECONNECT_DELAY: int = 2
    STREAM_RECONNECT_BACKOFF_FACTOR: float = 2.0
    MAX_STREAM_RECONNECT_ATTEMPTS: int = 5
    READ_EMPTY_SLEEP_BASE: float = 0.1
    READ_ERROR_BACKOFF_BASE: float = 0.1
    YOUTUBE_SESSION_IDLE_TIMEOUT: int = 25
    YOUTUBE_INITIAL_TIMEOUT: int = 25
    YOUTUBE_URL_REFRESH_INTERVAL: int = 600
    YOUTUBE_URL_REFRESH_MAX_ATTEMPTS: int = 3
    YOUTUBE_LOW_QUALITY_MAX_CHUNKS: int = 5
    YOUTUBE_INITIAL_WAIT: float = 3.0
    FFMPEG_START_WAIT: float = 1.5
    STREAM_TEST_TIMEOUT: int = 8
    YOUTUBE_STREAM_TEST_TIMEOUT: int = 8
    AUDIO_ENHANCEMENT_MIN_LENGTH: int = 1600
    NOISEREDUCE_INTERVAL: int = 10
    NOISEREDUCE_MIN_LENGTH: int = 32000
    LOW_QUALITY_CHUNK_THRESHOLD_FACTOR: float = 0.25
    BUFFER_FLUSH_INACTIVITY: float = 10.0
    READ_WITH_TIMEOUT_SELECT_INTERVAL: float = 0.001
    READ_WITH_TIMEOUT_EMPTY_LOG_INTERVAL: int = 10
    MAX_BACKOFF: int = 30
    PROGRESS_UPDATE_INTERVAL: float = 0.5
    BUFFER_FLUSH_TIMEOUT: float = 30.0
    ADAPTIVE_CHUNK_MIN_SAMPLES: int = 5
    ADAPTIVE_CHUNK_SMOOTHING_ALPHA: float = (
        0.3  # Glättungsfaktor für gleitenden Durchschnitt
    )
    ADAPTIVE_CHUNK_STABLE_THRESHOLD: int = (
        2  # wie oft gleiche Abweichung für Änderung nötig
    )
    ERRORS_BEFORE_CHUNK_REDUCTION: int = 3  # Fehler, bevor Chunk-Dauer halbiert wird
    SUCCESSES_BEFORE_CHUNK_INCREASE: int = (
        5  # Erfolge, bevor Chunk-Dauer erhöht werden kann
    )

    # -------------------------------------------------------------------------
    # Queue-Management
    # -------------------------------------------------------------------------
    GUI_QUEUE_MAX_SIZE: int = 300
    GUI_QUEUE_CLEANUP_TARGET: int = 100
    TEXT_QUEUE_MAX_SIZE: int = 150
    TEXT_QUEUE_CLEANUP_TARGET: int = 75

    # -------------------------------------------------------------------------
    # Logging-Schwellwerte
    # -------------------------------------------------------------------------
    LOW_QUALITY_CHUNK_LOG_LEVEL: int = (
        4  # Ab DEBUG_LEVEL 4 werden „Chunk too small“-Meldungen geloggt
    )
    LOW_QUALITY_CHUNK_LOG_INTERVAL: int = (
        10  # Jede n-te Meldung loggen, wenn Debug-Level nicht ausreicht
    )


# -----------------------------------------------------------------------------
# Config
# -----------------------------------------------------------------------------
@dataclass
class Config:
    """Optimierte Konfigurationsklasse für Dragon Whisperer – maximale Präzision (ohne __slots__)."""

    # Audio-Grundlagen
    SAMPLE_RATE: int = Constants.SAMPLE_RATE
    CHANNELS: int = Constants.CHANNELS
    AUDIO_FORMAT: str = Constants.AUDIO_FORMAT
    BYTES_PER_SAMPLE: int = Constants.BYTES_PER_SAMPLE

    # Prozesse & Limits
    MAX_SUBPROCESSES: int = Constants.MAX_SUBPROCESSES
    SUBPROCESS_TIMEOUT: int = Constants.SUBPROCESS_TIMEOUT
    GUI_OPERATION_TIMEOUT: float = Constants.GUI_OPERATION_TIMEOUT
    MEMORY_CHECK_INTERVAL: int = Constants.MEMORY_CHECK_INTERVAL
    MAX_GUI_UPDATES_PER_SECOND: int = Constants.MAX_GUI_UPDATES_PER_SECOND
    MAX_MEMORY_USAGE: int = Constants.MAX_MEMORY_USAGE
    MAX_CACHE_SIZE: int = Constants.MAX_CACHE_SIZE
    MAX_TEXT_LINES: int = Constants.MAX_TEXT_LINES
    DEFAULT_BEAM_SIZE: int = Constants.DEFAULT_BEAM_SIZE
    DEFAULT_TEMPERATURE: float = Constants.DEFAULT_TEMPERATURE
    ENABLE_VAD_FILTER: bool = Constants.ENABLE_VAD_FILTER
    MAX_CONSECUTIVE_ERRORS: int = Constants.MAX_CONSECUTIVE_ERRORS

    # Chunking – Basiswerte
    _base_chunk_duration: int = Constants.BASE_CHUNK_DURATION
    CHUNK_OVERLAP: float = Constants.CHUNK_OVERLAP
    MIN_CHUNK_DURATION: int = Constants.MIN_CHUNK_DURATION
    MAX_CHUNK_DURATION: int = Constants.MAX_CHUNK_DURATION
    LANGUAGE_VAD: Dict[str, Dict[str, Any]] = field(
        default_factory=lambda: Constants.LANGUAGE_VAD.copy()
    )
    _actual_chunk_duration: float = field(init=False, default=float(Constants.BASE_CHUNK_DURATION))

    # Stream-Parameter
    STREAM_TIMEOUT: int = Constants.STREAM_TIMEOUT
    INITIAL_BUFFER_SECONDS: float = Constants.INITIAL_BUFFER_SECONDS
    MAX_EMPTY_READS: int = Constants.MAX_EMPTY_READS
    RECONNECT_DELAY: int = Constants.RECONNECT_DELAY
    READ_RETRY_DELAY: float = Constants.READ_RETRY_DELAY

    # FFmpeg
    FFMPEG_BUFSIZE: str = Constants.FFMPEG_BUFSIZE
    FFMPEG_THREADS: int = Constants.FFMPEG_THREADS
    FFMPEG_PROBESIZE: str = Constants.FFMPEG_PROBESIZE
    FFMPEG_ANALYZE_DURATION: str = Constants.FFMPEG_ANALYZE_DURATION

    # Timeouts (Mikrosekunden für FFmpeg)
    YOUTUBE_TIMEOUT: int = Constants.YOUTUBE_TIMEOUT
    NORMAL_TIMEOUT: int = Constants.NORMAL_TIMEOUT

    # Audio-Filter – für maximale Präzision auf reines Resample reduziert
    AUDIO_FILTER: str = "aresample=16000"
    LANGUAGE_FILTERS: Dict[str, str] = field(
        default_factory=lambda: {
            k: "aresample=16000" for k in Constants.LANGUAGE_FILTERS
        }
    )
    FILTER_PROFILES: Dict[str, str] = field(
        default_factory=lambda: {
            "transcription": "aresample=16000",
            "translation": "aresample=16000",
            "realtime": "aresample=16000",
            "noisy": "aresample=16000,volume=2.5,highpass=f=150,lowpass=f=3000",
            "music": "aresample=16000",
            "podcast": "aresample=16000",
        }
    )

    # YouTube-Header
    YOUTUBE_HEADERS: Dict[str, str] = field(
        default_factory=lambda: Constants.YOUTUBE_HEADERS.copy()
    )

    # Plattform-Konfiguration
    PLATFORM_CONFIG: Dict[str, Dict[str, Any]] = field(
        default_factory=lambda: Constants.PLATFORM_CONFIG.copy()
    )

    # Audio-Enhancement (standardmäßig deaktiviert für maximale Präzision)
    AUDIO_ENHANCEMENT_ENABLED: bool = False
    MIN_RMS_THRESHOLD: float = Constants.MIN_RMS_THRESHOLD
    TARGET_RMS: float = Constants.TARGET_RMS
    MAX_GAIN: float = Constants.MAX_GAIN
    CLIPPING_THRESHOLD: float = Constants.CLIPPING_THRESHOLD

    # Duplikaterkennung (deaktiviert für maximale Präzision)
    DUPLICATE_CHECK_ENABLED: bool = False
    RECENT_TRANSCRIPTIONS_SIZE: int = Constants.RECENT_TRANSCRIPTIONS_SIZE
    MIN_TEXT_LENGTH: int = Constants.MIN_TEXT_LENGTH
    MIN_UNIQUE_WORDS_RATIO: float = Constants.MIN_UNIQUE_WORDS_RATIO

    # Untertitel
    SUBTITLE_BUFFER_SIZE: int = Constants.SUBTITLE_BUFFER_SIZE
    ENABLE_TIMED_TRANSCRIPTIONS: bool = Constants.ENABLE_TIMED_TRANSCRIPTIONS
    ENABLE_TIMED_TRANSLATIONS: bool = Constants.ENABLE_TIMED_TRANSLATIONS

    # Logging
    ENABLE_DEBUG_LOGGING: bool = Constants.ENABLE_DEBUG_LOGGING
    LOG_CHUNK_PROCESSING: bool = Constants.LOG_CHUNK_PROCESSING
    LOG_AUDIO_STATS: bool = Constants.LOG_AUDIO_STATS
    LOG_PERFORMANCE: bool = Constants.LOG_PERFORMANCE
    LOG_STREAM_EVENTS: bool = Constants.LOG_STREAM_EVENTS
    PERFORMANCE_LOG_INTERVAL: int = Constants.PERFORMANCE_LOG_INTERVAL

    # Cache
    MAX_CACHE_SIZE_MB: int = Constants.MAX_CACHE_SIZE_MB
    CACHE_ENABLED: bool = Constants.CACHE_ENABLED

    # ----------------------------------------------------------------------
    # Dynamische Eigenschaften (berechnet)
    # ----------------------------------------------------------------------
    @property
    def CHUNK_DURATION(self) -> float:
        return self._actual_chunk_duration

    @CHUNK_DURATION.setter
    def CHUNK_DURATION(self, value: float) -> None:
        if self.MIN_CHUNK_DURATION <= value <= self.MAX_CHUNK_DURATION:
            self._actual_chunk_duration = float(value)
        else:
            logger.warning(
                f"⚠️ Chunk duration {value}s out of range, clamping to [{self.MIN_CHUNK_DURATION}, {self.MAX_CHUNK_DURATION}]"
            )
            self._actual_chunk_duration = max(
                self.MIN_CHUNK_DURATION, min(value, self.MAX_CHUNK_DURATION)
            )

    @property
    def CHUNK_SIZE_BYTES(self) -> int:
        return int(
            self.CHUNK_DURATION
            * self.SAMPLE_RATE
            * self.CHANNELS
            * self.BYTES_PER_SAMPLE
        )

    @property
    def OVERLAP_SIZE_BYTES(self) -> int:
        return int(
            self.CHUNK_OVERLAP
            * self.SAMPLE_RATE
            * self.CHANNELS
            * self.BYTES_PER_SAMPLE
        )

    @property
    def BYTES_PER_SECOND(self) -> int:
        return self.SAMPLE_RATE * self.CHANNELS * self.BYTES_PER_SAMPLE

    @property
    def MIN_CHUNK_BYTES(self) -> int:
        return int(self.MIN_CHUNK_DURATION * self.BYTES_PER_SECOND)

    @property
    def MAX_CHUNK_BYTES(self) -> int:
        return int(self.MAX_CHUNK_DURATION * self.BYTES_PER_SECOND)

    @property
    def INITIAL_BUFFER_BYTES(self) -> int:
        return int(self.INITIAL_BUFFER_SECONDS * self.BYTES_PER_SECOND)

    # ----------------------------------------------------------------------
    # Methoden
    # ----------------------------------------------------------------------
    def get_timeout_microseconds(self, is_youtube: bool = False) -> int:
        return self.YOUTUBE_TIMEOUT if is_youtube else self.NORMAL_TIMEOUT

    def get_audio_filter(
        self, language: Optional[str] = None, profile: Optional[str] = None
    ) -> str:
        """Liefert den optimalen Audio-Filter – für maximale Präzision nur Resampling."""
        return "aresample=16000"

    def get_youtube_headers(self, is_manifest: bool = False) -> Dict[str, str]:
        headers = self.YOUTUBE_HEADERS.copy()
        if is_manifest:
            headers.update(
                {"X-Client-Data": "CI22yQE=", "Content-Type": "application/x-mpegURL"}
            )
        return headers

    def get_platform_config(self, platform: Optional[str] = None) -> Dict[str, Any]:
        if not platform:
            platform = SYSTEM.lower()
        return self.PLATFORM_CONFIG.get(platform, self.PLATFORM_CONFIG["linux"])

    def __post_init__(self) -> None:
        self._actual_chunk_duration = float(self._base_chunk_duration)

    def calculate_optimal_chunk_duration(
        self, model_size: str = "medium", is_realtime: bool = False
    ) -> int:
        if is_realtime:
            return self.MIN_CHUNK_DURATION
        model_durations = {
            "tiny": 3,
            "tiny.en": 3,
            "base": 4,
            "base.en": 4,
            "small": 5,
            "small.en": 5,
            "medium": 5,
            "medium.en": 5,
            "large": 6,
            "large-v2": 6,
            "large-v3": 6,
        }
        return model_durations.get(model_size.lower(), self._base_chunk_duration)

    def validate_config(self) -> bool:
        try:
            return (
                self.SAMPLE_RATE in [8000, 16000, 22050, 44100, 48000]
                and self.CHANNELS in [1, 2]
                and self.MIN_CHUNK_DURATION
                <= self.CHUNK_DURATION
                <= self.MAX_CHUNK_DURATION
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Config validation error: {e}")
            return False

    def print_summary(self) -> None:
        logger.info("\n" + "=" * 60)
        logger.info("🎵 CONFIGURATION (optimized for precision)")
        logger.info("=" * 60)
        logger.info(f"📊 Audio: {self.SAMPLE_RATE}Hz, {self.CHANNELS}ch")
        logger.info(f"📦 Chunk: {self.CHUNK_DURATION}s ({self.CHUNK_SIZE_BYTES:,}B)")
        logger.info(f"⚡ Bytes/sec: {self.BYTES_PER_SECOND:,}")
        logger.info(f"🎛️ Filter: {self.AUDIO_FILTER}")
        logger.info(f"✅ Valid: {self.validate_config()}")
        logger.info("=" * 60)

    def __str__(self) -> str:
        return (
            f"Config(chunk={self.CHUNK_DURATION}s, "
            f"filter_profiles={len(self.FILTER_PROFILES)})"
        )


class RealtimeConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.CHUNK_DURATION = 5
        self.CHUNK_OVERLAP = 0.3
        self.STREAM_TIMEOUT = 5
        self.AUDIO_FILTER = self.FILTER_PROFILES["realtime"]


class HighAccuracyConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.CHUNK_DURATION = 25
        self.CHUNK_OVERLAP = 0.8
        self.AUDIO_FILTER = (
            "aresample=16000,volume=1.8,highpass=f=80,"
            "lowpass=f=3800,dynaudnorm=p=0.3:s=3:g=20"
        )


class YouTubeOptimizedConfig(Config):
    def __init__(self) -> None:
        super().__init__()
        self.FFMPEG_THREADS = 1
        self.FFMPEG_BUFSIZE = "1024k"
        self.YOUTUBE_TIMEOUT = 5000000
        self.RECONNECT_DELAY = 1
        self.AUDIO_FILTER = (
            "aresample=16000,volume=2.2,highpass=f=120,"
            "lowpass=f=3200,compand=attacks=0:decays=0.3"
        )


def get_config(config_type: str = "default") -> Config:
    configs = {
        "default": Config,
        "realtime": RealtimeConfig,
        "high_accuracy": HighAccuracyConfig,
        "youtube": YouTubeOptimizedConfig,
    }
    config_class = configs.get(config_type, Config)
    config = config_class()
    config.validate_config()
    return config


# =============================================================================
# 4. HILFSKLASSEN UND -FUNKTIONEN (UTILS)
# =============================================================================
class FastLazyLoader:
    _loaded_modules: Dict[str, Any] = {}
    _module_locks: Dict[str, threading.RLock] = {}
    _class_lock = threading.RLock()
    MAX_CACHE_SIZE = 10

    @classmethod
    def _get_lock(cls, module_name: str) -> threading.RLock:
        with cls._class_lock:
            if module_name not in cls._module_locks:
                cls._module_locks[module_name] = threading.RLock()
        return cls._module_locks[module_name]

    @classmethod
    def _prune_cache(cls) -> None:
        with cls._class_lock:
            while len(cls._loaded_modules) > cls.MAX_CACHE_SIZE:
                cls._loaded_modules.popitem(last=False)

    @classmethod
    def load(cls, module_name: str, import_path: Optional[str] = None) -> Any:
        if module_name in cls._loaded_modules:
            return cls._loaded_modules[module_name]
        lock = cls._get_lock(module_name)
        with lock:
            if module_name in cls._loaded_modules:
                return cls._loaded_modules[module_name]
            try:
                if module_name == "torch":
                    import torch

                    try:
                        import torch._logging

                        torch._logging.set_logs(all=logging.ERROR)
                    except Exception:
                        pass
                    cls._loaded_modules["torch"] = torch
                elif module_name == "faster_whisper":
                    from faster_whisper import WhisperModel

                    cls._loaded_modules["faster_whisper"] = WhisperModel
                elif module_name == "numpy":
                    import numpy as np

                    cls._loaded_modules["numpy"] = np
                elif module_name == "deep_translator":
                    from deep_translator import GoogleTranslator

                    cls._loaded_modules["deep_translator"] = GoogleTranslator
                elif module_name == "scipy.signal":
                    import scipy.signal

                    cls._loaded_modules["scipy.signal"] = scipy.signal
                else:
                    module = importlib.import_module(import_path or module_name)
                    cls._loaded_modules[module_name] = module
                cls._prune_cache()
                return cls._loaded_modules[module_name]
            except ImportError as e:
                logger.warning(f"⚠️ Module {module_name} not available: {e}")

                class MockModule:
                    _is_mock = True

                    def __init__(self, name: str):
                        self.__name__ = name

                    def __getattr__(self, name: str) -> Any:
                        def mock_method(*args: Any, **kwargs: Any) -> Any:
                            raise ImportError(f"Module {self.__name__} not available")

                        return mock_method

                mock = MockModule(module_name)
                cls._loaded_modules[module_name] = mock
                cls._prune_cache()
                return mock

    @classmethod
    def is_available(cls, module_name: str) -> bool:
        if importlib.util.find_spec(module_name) is None:
            return False
        if module_name in cls._loaded_modules:
            module = cls._loaded_modules[module_name]
            return not getattr(module, "_is_mock", False)
        return True

    @classmethod
    def clear_cache(cls) -> None:
        with cls._class_lock:
            cls._loaded_modules.clear()
            cls._module_locks.clear()


# -----------------------------------------------------------------------------
# Plattform-Stderr-Filter
# -----------------------------------------------------------------------------
class PlatformStderrFilter:
    def __init__(self, original_stderr: Any) -> None:
        self.original_stderr = original_stderr
        self.filter_patterns = [
            "mpp_soc:",
            "mpp_platform:",
            "can not found match soc name",
            "/proc/device-tree/compatible",
            "rockchip",
            "ffmpeg",
            "TORCH_NCCL",
        ]
        if IS_WINDOWS:
            self.filter_patterns.extend(
                [
                    "Failed to set direct console mode",
                    "Console code page",
                    "chcp",
                    "win32api",
                ]
            )

    def write(self, text: str) -> None:
        if text and any(p in text for p in self.filter_patterns):
            return
        self.original_stderr.write(text)

    def flush(self) -> None:
        self.original_stderr.flush()


sys.stderr = PlatformStderrFilter(sys.stderr)


# -----------------------------------------------------------------------------
# Terminal-Einstellungen speichern/wiederherstellen
# -----------------------------------------------------------------------------
_original_stty_settings: Optional[str] = None


def _save_terminal_settings() -> None:
    global _original_stty_settings
    if not IS_LINUX:
        return
    try:
        result = subprocess.run(
            ["stty", "-g"], capture_output=True, text=True, check=False, timeout=2
        )
        if result.returncode == 0 and result.stdout:
            _original_stty_settings = result.stdout.strip()
    except (subprocess.TimeoutExpired, OSError) as e:
        if isinstance(e, (KeyboardInterrupt, SystemExit)):
            raise
        if DEBUG_LEVEL >= 2:
            logger.debug(f"stty -g fehlgeschlagen: {e}")


def _restore_terminal_settings() -> None:
    if not IS_LINUX or _original_stty_settings is None:
        return
    try:
        subprocess.run(["stty", _original_stty_settings], check=False, timeout=2)
    except Exception:
        pass


_save_terminal_settings()
atexit.register(_restore_terminal_settings)


# -----------------------------------------------------------------------------
# SignalHandler
# -----------------------------------------------------------------------------
class ShutdownPriority(Enum):
    """Prioritäten für Cleanup-Operationen (höhere Priorität = früher ausgeführt)."""

    CRITICAL = auto()
    HIGH = auto()
    MEDIUM = auto()
    LOW = auto()


class SignalHandler:
    """
    Zentraler Signal- und Shutdown-Handler mit Registrierung von Cleanup-Funktionen.
    Unterstützt gestaffelte, parallele Ausführung und Notfall-Shutdown.
    """

    _instance: Optional["SignalHandler"] = None
    _lock = threading.RLock()
    _shutdown_requested = False
    _shutdown_in_progress = False
    _signal_count = 0
    _original_handlers: Dict[int, Any] = {}
    _atexit_registered = False

    # Konfiguration (kann über setup() geändert werden)
    _config = {
        "verbose": False,
        "silent": True,
        "max_cleanup_time": 20.0,
        "emergency_timeout": 2.0,
        "atexit_enabled": True,
        "hybrid_shutdown": True,
        "parallel_cleanup": True,
        "parallel_workers": 4,
    }

    _cleanup_operations: Dict[ShutdownPriority, List["_CleanupOperation"]] = {
        ShutdownPriority.CRITICAL: [],
        ShutdownPriority.HIGH: [],
        ShutdownPriority.MEDIUM: [],
        ShutdownPriority.LOW: [],
    }

    class _CleanupOperation:
        """Repräsentiert eine einzelne Cleanup-Funktion mit Metadaten."""

        __slots__ = (
            "func",
            "name",
            "priority",
            "timeout",
            "essential",
            "attempts",
            "last_error",
        )

        def __init__(
            self,
            func: Callable[[], Any],
            name: str,
            priority: ShutdownPriority = ShutdownPriority.MEDIUM,
            timeout: float = 3.0,
            essential: bool = False,
        ):
            self.func = func
            self.name = name
            self.priority = priority
            self.timeout = timeout
            self.essential = essential
            self.attempts = 0
            self.last_error: Optional[str] = None

        def run(self) -> bool:
            """Führt die Funktion aus (blockierend) und gibt Erfolg zurück."""
            self.attempts += 1
            try:
                start = time.perf_counter()
                self.func()
                duration = (time.perf_counter() - start) * 1000
                if SignalHandler._config["verbose"] and duration > 10:
                    print(f"  ✅ {self.name} – {duration:.2f}ms")
                self.last_error = None
                return True
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                self.last_error = str(e)
                if SignalHandler._config["verbose"]:
                    print(f"  ❌ {self.name} fehlgeschlagen: {e}")
                return False

        def execute_with_timeout(self) -> bool:
            """Führt die Funktion mit Timeout aus (in eigenem Thread)."""
            self.attempts += 1
            result = None
            exception = None

            def target():
                nonlocal result, exception
                try:
                    result = self.func()
                except Exception as e:
                    exception = e

            thread = threading.Thread(target=target, daemon=True)
            thread.start()
            thread.join(timeout=self.timeout)

            if thread.is_alive():
                self.last_error = f"Timeout nach {self.timeout}s"
                return False
            if exception:
                self.last_error = str(exception)
                return False
            self.last_error = None
            return True

    def __new__(cls) -> "SignalHandler":
        with cls._lock:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
        return cls._instance

    # ----------------------------------------------------------------------
    # Öffentliche API
    # ----------------------------------------------------------------------
    @classmethod
    def setup(
        cls,
        verbose: bool = False,
        silent: bool = True,
        hybrid_shutdown: bool = True,
        **kwargs: Any,
    ) -> "SignalHandler":
        """
        Initialisiert den SignalHandler (nur einmalig).
        - verbose: Gibt detaillierte Ausgaben aus
        - silent: Unterdrückt Benutzerhinweise
        - hybrid_shutdown: Bei True wird sys.exit() im Hauptthread verwendet, sonst os._exit()
        - Weitere Optionen: max_cleanup_time, emergency_timeout, atexit_enabled,
          parallel_cleanup, parallel_workers
        """
        with cls._lock:
            if hasattr(cls, "_setup_complete") and cls._setup_complete:
                return cls._instance

            cls._config.update(
                {
                    "verbose": verbose,
                    "silent": silent,
                    "hybrid_shutdown": hybrid_shutdown,
                    "max_cleanup_time": kwargs.get("max_cleanup_time", 20.0),
                    "emergency_timeout": kwargs.get("emergency_timeout", 2.0),
                    "atexit_enabled": kwargs.get("atexit_enabled", True),
                    "parallel_cleanup": kwargs.get("parallel_cleanup", True),
                    "parallel_workers": kwargs.get("parallel_workers", 4),
                }
            )

            cls._save_original_handlers()
            cls._install_signal_handlers()

            if cls._config["atexit_enabled"]:
                cls._register_atexit()

            cls._setup_complete = True
            if cls._config["verbose"]:
                print("✅ SignalHandler bereit")
            return cls._instance

    @classmethod
    def register_cleanup(
        cls,
        func: Callable[[], Any],
        name: Optional[str] = None,
        priority: ShutdownPriority = ShutdownPriority.MEDIUM,
        timeout: float = 3.0,
        essential: bool = False,
    ) -> None:
        """
        Registriert eine Cleanup-Funktion.
        - name: Anzeigename (falls None, wird der Funktionsname verwendet)
        - priority: Priorität (CRITICAL, HIGH, MEDIUM, LOW)
        - timeout: Max. Ausführungszeit in Sekunden (für nicht-CRITICAL)
        - essential: Bei True führt ein Fehler zum Abbruch des geordneten Shutdowns
        """
        if name is None:
            name = getattr(func, "__name__", "Anonymous")

        op = cls._CleanupOperation(
            func=func,
            name=name,
            priority=priority,
            timeout=timeout,
            essential=essential,
        )

        with cls._lock:
            for existing in cls._cleanup_operations[priority]:
                if existing.func == func:
                    return
            cls._cleanup_operations[priority].append(op)

        if cls._config["verbose"]:
            print(f"✅ Cleanup registriert: {name} (Priority: {priority.name})")

    @classmethod
    def unregister_cleanup(cls, func: Callable) -> bool:
        """Entfernt eine registrierte Cleanup-Funktion."""
        with cls._lock:
            for priority in ShutdownPriority:
                for i, op in enumerate(cls._cleanup_operations[priority]):
                    if op.func == func:
                        del cls._cleanup_operations[priority][i]
                        return True
        return False

    @classmethod
    def should_shutdown(cls) -> bool:
        """Gibt zurück, ob ein Shutdown angefordert wurde."""
        with cls._lock:
            return cls._shutdown_requested

    @classmethod
    def get_status(cls) -> Dict[str, Any]:
        """Liefert einen Statusbericht."""
        with cls._lock:
            return {
                "shutdown_requested": cls._shutdown_requested,
                "shutdown_in_progress": cls._shutdown_in_progress,
                "signal_count": cls._signal_count,
                "setup_complete": getattr(cls, "_setup_complete", False),
                "cleanup_operations": {
                    p.name: len(ops) for p, ops in cls._cleanup_operations.items()
                },
                "atexit_registered": cls._atexit_registered,
                "hybrid_mode": cls._config["hybrid_shutdown"],
                "config": {
                    k: v for k, v in cls._config.items() if not k.startswith("_")
                },
            }

    @classmethod
    def emergency_shutdown(cls, reason: str = "Emergency", exit_code: int = 1) -> None:
        """Sofortiger Abbruch ohne weitere Cleanups."""
        if not cls._config["silent"]:
            print(f"🚨 NOTFALL-SHUTDOWN: {reason}")
        with cls._lock:
            cls._shutdown_requested = True
            cls._shutdown_in_progress = True
        os._exit(exit_code)

    @classmethod
    def reset(cls) -> None:
        """Setzt den Handler zurück (nur für Tests)."""
        with cls._lock:
            cls._instance = None
            cls._shutdown_requested = False
            cls._shutdown_in_progress = False
            cls._signal_count = 0
            cls._setup_complete = False
            cls._original_handlers.clear()
            cls._atexit_registered = False
            cls._cleanup_operations = {
                ShutdownPriority.CRITICAL: [],
                ShutdownPriority.HIGH: [],
                ShutdownPriority.MEDIUM: [],
                ShutdownPriority.LOW: [],
            }

    # ----------------------------------------------------------------------
    # Private Methoden
    # ----------------------------------------------------------------------
    @classmethod
    def _save_original_handlers(cls) -> None:
        """Sichert die ursprünglichen Signal-Handler (nicht Windows)."""
        if sys.platform == "win32":
            return
        for sig in (py_signal.SIGINT, py_signal.SIGTERM):
            try:
                cls._original_handlers[sig] = py_signal.getsignal(sig)
            except (ValueError, OSError):
                pass

    @classmethod
    def _install_signal_handlers(cls) -> None:
        """Installiert die eigenen Signal-Handler."""

        def signal_handler(signum: int, frame: Any) -> None:
            with cls._lock:
                cls._signal_count += 1
                if cls._signal_count == 1:
                    if not cls._config["silent"]:
                        print("\n⚠️ Shutdown angefordert...")
                    cls._shutdown_requested = True
                    cls._initiate_graceful_shutdown()
                elif cls._signal_count >= 2:
                    if not cls._config["silent"]:
                        print("\n🛑 Forcierter Shutdown...")
                    cls._force_shutdown()

        if sys.platform == "win32":
            try:
                import win32api

                def win32_handler(ctrl_type: int) -> bool:
                    if ctrl_type in (0, 1, 2, 5):
                        signal_handler(ctrl_type, None)
                        return True
                    return False

                win32api.SetConsoleCtrlHandler(win32_handler, True)
            except ImportError:
                try:
                    py_signal.signal(py_signal.SIGINT, signal_handler)
                except (ValueError, OSError):
                    pass
        else:
            try:
                py_signal.signal(py_signal.SIGINT, signal_handler)
                py_signal.signal(py_signal.SIGTERM, signal_handler)
            except (ValueError, OSError):
                pass

    @classmethod
    def _initiate_graceful_shutdown(cls) -> None:
        """Startet den geordneten Shutdown-Prozess."""
        with cls._lock:
            if cls._shutdown_in_progress:
                return
            cls._shutdown_in_progress = True

        if not cls._config["silent"]:
            print("🧹 Starte geordneten Shutdown...")

        success = cls._execute_cleanup()
        cls._restore_original_handlers()

        # Hybrid-Shutdown: sys.exit() im Hauptthread, sonst os._exit()
        if (
            cls._config["hybrid_shutdown"]
            and threading.current_thread() is threading.main_thread()
        ):
            if cls._config["verbose"]:
                print("💡 Sauberes Exit im Hauptthread")
            sys.exit(0 if success else 1)
        else:
            if cls._config["verbose"]:
                print("💡 Sofort-Exit in Thread")
            os._exit(0 if success else 1)

    @classmethod
    def _force_shutdown(cls) -> None:
        """Notfall-Shutdown ohne Cleanups."""
        try:
            cls._handle_atexit_cleanup()  # Nur kritische GPU/Memory-Cleanups
        except Exception:
            pass
        os._exit(1)

    @classmethod
    def _execute_cleanup(cls) -> bool:
        """
        Führt alle registrierten Cleanup-Funktionen aus.
        CRITICAL: sequenziell, essenziell.
        HIGH/MEDIUM/LOW: parallel (wenn aktiviert) mit Timeout.
        Gibt True zurück, wenn alle essenziellen Cleanups erfolgreich waren.
        """
        start_time = time.time()
        max_time = cls._config["max_cleanup_time"]
        overall_success = True
        completed = 0
        failed = 0

        # 1. CRITICAL – sequenziell, essenziell
        critical_ops = cls._cleanup_operations[ShutdownPriority.CRITICAL]
        if critical_ops:
            if cls._config["verbose"]:
                print(f"🔴 {len(critical_ops)} CRITICAL Cleanups (sequenziell)")
            for op in critical_ops:
                if time.time() - start_time > max_time:
                    if cls._config["verbose"]:
                        print("⏰ Zeitlimit für CRITICAL Cleanups überschritten")
                    break
                success = op.run()
                completed += 1
                if not success:
                    failed += 1
                    if op.essential:
                        overall_success = False
                        if cls._config["verbose"]:
                            print(f"❌ ESSENTIAL fehlgeschlagen: {op.name}")
                    elif cls._config["verbose"]:
                        print(f"⚠️ Cleanup fehlgeschlagen: {op.name}")
                elif cls._config["verbose"]:
                    print(f"✅ {op.name}")

        # 2. HIGH, MEDIUM, LOW – optional parallel
        if cls._config["parallel_cleanup"]:
            non_critical_ops = []
            for prio in (
                ShutdownPriority.HIGH,
                ShutdownPriority.MEDIUM,
                ShutdownPriority.LOW,
            ):
                non_critical_ops.extend(cls._cleanup_operations[prio])

            if non_critical_ops:
                if cls._config["verbose"]:
                    print(
                        f"🟡 {len(non_critical_ops)} Cleanups (parallel, max {cls._config['parallel_workers']} Worker)"
                    )
                with ThreadPoolExecutor(
                    max_workers=cls._config["parallel_workers"]
                ) as executor:
                    future_to_op = {}
                    for op in non_critical_ops:
                        if time.time() - start_time > max_time:
                            break
                        future = executor.submit(op.execute_with_timeout)
                        future_to_op[future] = op

                    for future in as_completed(future_to_op):
                        op = future_to_op[future]
                        try:
                            success = future.result()
                        except Exception as e:
                            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                raise
                            success = False
                            op.last_error = str(e)
                        completed += 1
                        if not success:
                            failed += 1
                            if op.essential:
                                overall_success = False
                                if cls._config["verbose"]:
                                    print(
                                        f"❌ ESSENTIAL parallel fehlgeschlagen: {op.name}"
                                    )
                            elif cls._config["verbose"]:
                                print(f"⚠️ Cleanup parallel fehlgeschlagen: {op.name}")
                        elif cls._config["verbose"]:
                            print(f"✅ {op.name}")
                        if time.time() - start_time > max_time:
                            if cls._config["verbose"]:
                                print(
                                    "⏰ Zeitlimit erreicht – weitere parallele Cleanups abbrechen"
                                )
                            break
        else:
            for prio in (
                ShutdownPriority.HIGH,
                ShutdownPriority.MEDIUM,
                ShutdownPriority.LOW,
            ):
                for op in cls._cleanup_operations[prio]:
                    if time.time() - start_time > max_time:
                        break
                    success = op.execute_with_timeout()
                    completed += 1
                    if not success:
                        failed += 1
                        if op.essential:
                            overall_success = False
                        if cls._config["verbose"]:
                            print(f"⚠️ Cleanup fehlgeschlagen: {op.name}")
                    elif cls._config["verbose"]:
                        print(f"✅ {op.name}")
                if time.time() - start_time > max_time:
                    break

        if cls._config["verbose"]:
            print(f"📊 Cleanup abgeschlossen: {completed} ops, {failed} fehlgeschlagen")
        return overall_success

    @classmethod
    def _register_atexit(cls) -> None:
        """Registriert einen atexit-Handler für den Fall, dass das Programm normal endet."""
        with cls._lock:
            if cls._atexit_registered:
                return

            def safe_atexit_handler() -> None:
                try:
                    if (
                        threading.current_thread() is threading.main_thread()
                        and not cls._shutdown_in_progress
                    ):
                        cls._handle_atexit_cleanup()
                except Exception:
                    pass

            atexit.register(safe_atexit_handler)
            cls._atexit_registered = True
            if cls._config["verbose"]:
                print("✅ AtExit-Handler registriert")

    @classmethod
    def _handle_atexit_cleanup(cls) -> None:
        """Führt im atexit-Fall nur die kritischsten Cleanups aus (GPU, Memory)."""
        if cls._config["verbose"]:
            print("🔧 AtExit-Cleanup...")
        emergency_ops = []
        for op in cls._cleanup_operations[ShutdownPriority.CRITICAL]:
            if op.essential and ("GPU" in op.name or "Memory" in op.name):
                emergency_ops.append(op)
                if len(emergency_ops) >= 3:
                    break

        start = time.time()
        for op in emergency_ops:
            if time.time() - start > cls._config["emergency_timeout"]:
                break
            try:
                op.func()
            except Exception:
                pass
        gc.collect()

    @classmethod
    def _restore_original_handlers(cls) -> None:
        """Stellt die ursprünglichen Signal-Handler wieder her (nicht Windows)."""
        if sys.platform == "win32":
            return
        for sig, handler in cls._original_handlers.items():
            if handler is not None:
                try:
                    py_signal.signal(sig, handler)
                except (ValueError, OSError):
                    pass


# -----------------------------------------------------------------------------
# PlatformUtils
# -----------------------------------------------------------------------------
class PlatformUtils:
    _environment_setup_done = False
    _environment_setup_lock = threading.RLock()
    _dependencies_checked = False
    _dependencies_lock = threading.RLock()

    @staticmethod
    def get_platform_config_dir() -> Path:
        try:
            if IS_WINDOWS:
                config_dir = Path(os.environ.get("APPDATA", "")) / "DragonWhisperer"
            elif IS_MACOS:
                config_dir = (
                    Path.home() / "Library" / "Application Support" / "DragonWhisperer"
                )
            else:
                config_dir = Path.home() / ".config" / "dragonwhisperer"
            config_dir.mkdir(parents=True, exist_ok=True)
            return config_dir
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Config directory error: {e}")
            fallback_dir = Path.home() / ".dragonwhisperer"
            fallback_dir.mkdir(parents=True, exist_ok=True)
            return fallback_dir

    @staticmethod
    def kill_process_tree(pid: int) -> bool:
        try:
            if IS_WINDOWS:
                subprocess.run(
                    ["taskkill", "/F", "/T", "/PID", str(pid)],
                    capture_output=True,
                    timeout=5,
                    check=False,
                    creationflags=subprocess.CREATE_NO_WINDOW,
                )
            else:
                try:
                    os.killpg(os.getpgid(pid), py_signal.SIGKILL)
                except (ProcessLookupError, PermissionError):
                    subprocess.run(
                        ["pkill", "-9", "-P", str(pid)],
                        capture_output=True,
                        timeout=5,
                        check=False,
                    )
            return True
        except subprocess.TimeoutExpired:
            logger.warning(f"⚠️ Timeout beim Beenden des Prozessbaums {pid}")
            return False
        except (OSError, subprocess.CalledProcessError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Error killing process tree {pid}: {e}")
            return False
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(
                f"⚠️ Unerwarteter Fehler beim Beenden von Prozessbaum {pid}: {e}"
            )
            return False

    @staticmethod
    def check_platform_dependencies() -> bool:
        with PlatformUtils._dependencies_lock:
            if PlatformUtils._dependencies_checked:
                return True
            missing: List[str] = []
            issues: List[str] = []
            logger.info("🔍 Checking platform dependencies...")
            ffmpeg_found = shutil.which("ffmpeg") is not None
            if not ffmpeg_found:
                missing.append("ffmpeg")
                issues.append("FFmpeg not found in PATH or standard locations")
            ytdlp_found = shutil.which("yt-dlp") is not None
            if not ytdlp_found:
                missing.append("yt-dlp")
                issues.append("yt-dlp not found in PATH")
            psutil_found = FastLazyLoader.is_available("psutil")
            if not psutil_found:
                issues.append("psutil not available (system monitoring)")
            critical_missing = []
            if not ffmpeg_found:
                critical_missing.append("ffmpeg")
            if not ytdlp_found:
                critical_missing.append("yt-dlp")
            if not GUI_AVAILABLE:
                critical_missing.append("tkinter")
                issues.append("Tkinter not available – required for GUI")
            if not WHISPER_AVAILABLE:
                logger.warning(
                    "⚠️ Kein Whisper-Backend verfügbar. Starte im Demo-Modus."
                )
            if not TRANSLATOR_AVAILABLE:
                issues.append(
                    "deep-translator not available (translation will be limited)"
                )
            if not TORCH_AVAILABLE:
                issues.append("PyTorch not available (optional for GPU acceleration)")
            if not psutil_found:
                issues.append("psutil not available (system monitoring limited)")
            if critical_missing:
                error_msg = (
                    f"❌ Fehlende kritische Abhängigkeiten: "
                    f"{', '.join(critical_missing)}\n\n"
                )
                error_msg += "\n".join(issues) + "\n"
                if "ffmpeg" in critical_missing:
                    error_msg += "FFmpeg Installation:\n"
                    if IS_WINDOWS:
                        error_msg += (
                            "  • Download from: https://ffmpeg.org/download.html\n"
                        )
                    elif IS_MACOS:
                        error_msg += "  • brew install ffmpeg\n"
                    else:
                        error_msg += "  • sudo apt install ffmpeg\n"
                if "yt-dlp" in critical_missing:
                    error_msg += "yt-dlp Installation:\n"
                    error_msg += "  • pip install yt-dlp\n"
                if "tkinter" in critical_missing:
                    error_msg += "Tkinter Installation:\n"
                    error_msg += "  • Usually included with Python. On Linux: sudo apt-get install python3-tk\n"
                if issues:
                    error_msg += "\nAdditional issues:\n"
                    for issue in issues:
                        error_msg += f"  • {issue}\n"
                error_msg += "\n💡 After installing, restart Dragon Whisperer."
                PlatformUtils._dependencies_checked = False
                raise RuntimeError(error_msg)
            logger.info("✅ Alle kritischen Abhängigkeiten gefunden")
            PlatformUtils._dependencies_checked = True
            return True

    @staticmethod
    def setup_platform_environment() -> None:
        with PlatformUtils._environment_setup_lock:
            if PlatformUtils._environment_setup_done:
                return
            logger.info("🔧 Setting up platform environment...")
            if IS_WINDOWS:
                try:
                    import ctypes

                    ctypes.windll.kernel32.SetConsoleOutputCP(65001)
                    os.system("chcp 65001 > nul 2>&1")
                    os.system("color")
                except (OSError, AttributeError) as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Windows console setup failed: {e}")
            elif IS_MACOS:
                temp_dir = Path(tempfile.gettempdir()) / "dragonwhisperer"
                try:
                    temp_dir.mkdir(exist_ok=True)
                except (OSError, PermissionError) as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ macOS temp dir creation failed: {e}")
            PlatformUtils._environment_setup_done = True
            logger.info("✅ Platform environment setup complete")

    @staticmethod
    def get_ffmpeg_path() -> Optional[str]:
        env_path = os.environ.get('FFMPEG_PATH')
        if env_path and os.path.exists(env_path):
            return env_path

        ffmpeg_path = shutil.which("ffmpeg")
        if ffmpeg_path:
            return ffmpeg_path
        if IS_WINDOWS:
            paths = [
                "C:\\ffmpeg\\bin\\ffmpeg.exe",
                "C:\\Program Files\\ffmpeg\\bin\\ffmpeg.exe",
                "C:\\Program Files (x86)\\ffmpeg\\bin\\ffmpeg.exe",
            ]
        elif IS_MACOS:
            paths = [
                "/usr/local/bin/ffmpeg",
                "/opt/homebrew/bin/ffmpeg",
                "/usr/bin/ffmpeg",
            ]
        else:
            paths = [
                "/usr/bin/ffmpeg",
                "/usr/local/bin/ffmpeg",
            ]
        for path in paths:
            if os.path.exists(path):
                return path
        return None

    @staticmethod
    def get_platform_info() -> Dict[str, Any]:
        info: Dict[str, Any] = {
            "system": SYSTEM,
            "is_windows": IS_WINDOWS,
            "is_macos": IS_MACOS,
            "is_linux": IS_LINUX,
            "is_arm": IS_ARM,
            "is_x86": IS_X86,
            "python_version": sys.version,
            "python_executable": sys.executable,
            "current_directory": os.getcwd(),
            "environment_setup": PlatformUtils._environment_setup_done,
            "dependencies_checked": PlatformUtils._dependencies_checked,
        }
        try:
            import psutil

            info["cpu_count"] = psutil.cpu_count()
            info["memory_total_gb"] = psutil.virtual_memory().total / (1024**3)
        except ImportError:
            info["cpu_count"] = "unknown"
            info["memory_total_gb"] = "unknown"
        except Exception:
            info["cpu_count"] = "error"
            info["memory_total_gb"] = "error"
        return info

    @staticmethod
    def print_platform_info() -> None:
        info = PlatformUtils.get_platform_info()
        logger.info("\n" + "=" * 60)
        logger.info("🐉 PLATFORM INFORMATION")
        logger.info("=" * 60)
        for key, value in info.items():
            if key not in ["environment_setup", "dependencies_checked"]:
                logger.info(f"{key:25} {value}")
        logger.info("-" * 60)
        logger.info(
            f"{'Environment Setup':25} {'✅' if info['environment_setup'] else '❌'}"
        )
        logger.info(
            f"{'Dependencies Checked':25} {'✅' if info['dependencies_checked'] else '❌'}"
        )
        logger.info("=" * 60)

    @staticmethod
    def sanitize_url(url: str) -> str:
        if not url:
            return ""
        return url.strip()

    @staticmethod
    def validate_file_path(file_url: str) -> Tuple[bool, str]:
        if not file_url.startswith(Constants.ALLOWED_FILE_SCHEME_PREFIX):
            return False, "Keine file://-URL"
        try:
            if IS_WINDOWS and file_url.startswith("file:///"):
                path_part = file_url[8:]
                path_part = urllib.request.url2pathname(path_part)
            else:
                path_part = file_url[len(Constants.ALLOWED_FILE_SCHEME_PREFIX) :]

            if IS_WINDOWS and path_part.startswith("\\\\"):
                pass
            else:
                path_part = os.path.normpath(path_part)

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            return False, f"Pfad kann nicht extrahiert werden: {e}"
        try:
            real_path = Path(path_part).resolve()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            return False, f"Pfad kann nicht normalisiert werden: {e}"
        if not real_path.exists():
            return False, f"Datei existiert nicht: {real_path}"
        if not real_path.is_file():
            return False, "Keine gültige Datei (möglicherweise ein Verzeichnis)"
        allowed_bases = [Path(p).resolve() for p in Constants.ALLOWED_FILE_BASE_DIRS]
        for base in allowed_bases:
            try:
                if real_path.is_relative_to(base):
                    return True, str(real_path)
            except AttributeError:
                if str(real_path).startswith(str(base)):
                    return True, str(real_path)
        temp_dir = Path(tempfile.gettempdir()).resolve()
        try:
            if real_path.is_relative_to(temp_dir):
                return True, str(real_path)
        except AttributeError:
            if str(real_path).startswith(str(temp_dir)):
                return True, str(real_path)
        logger.warning(f"⚠️ Datei außerhalb erlaubter Verzeichnisse: {real_path}")
        return True, str(real_path)


PlatformUtils.setup_platform_environment()


# -----------------------------------------------------------------------------
# Verfügbarkeiten
# -----------------------------------------------------------------------------
TORCH_AVAILABLE = FastLazyLoader.is_available("torch")
NUMPY_AVAILABLE = FastLazyLoader.is_available("numpy")
TRANSLATOR_AVAILABLE = FastLazyLoader.is_available("deep_translator")
SCIPY_AVAILABLE = FastLazyLoader.is_available("scipy.signal")
FASTER_WHISPER_AVAILABLE = importlib.util.find_spec("faster_whisper") is not None
OPENAI_WHISPER_AVAILABLE = importlib.util.find_spec("whisper") is not None
WHISPER_AVAILABLE = FASTER_WHISPER_AVAILABLE or OPENAI_WHISPER_AVAILABLE
OLLAMA_AVAILABLE = importlib.util.find_spec("requests") is not None

ARGOS_AVAILABLE = False
try:
    if importlib.util.find_spec("argostranslate") is not None:
        ARGOS_AVAILABLE = True
except ModuleNotFoundError:
    pass

if ARGOS_AVAILABLE:
    logger.info("✅ argos-translate verfügbar")
else:
    logger.warning(
        "⚠️ argos-translate nicht installiert – für Offline‑Übersetzung: pip install argostranslate"
    )

if FASTER_WHISPER_AVAILABLE:
    logger.info("✅ faster-whisper verfügbar")
else:
    logger.warning("⚠️ faster-whisper nicht verfügbar")

if OPENAI_WHISPER_AVAILABLE:
    logger.info("✅ openai-whisper verfügbar")
else:
    logger.warning("⚠️ openai-whisper nicht verfügbar")

if not WHISPER_AVAILABLE:
    logger.warning("⚠️ KEINE Whisper-Bibliothek verfügbar! Starte im Demo-Modus.")

if not GUI_AVAILABLE:
    logger.info("📟 Terminal-Modus (kein GUI)")
else:
    logger.info("✅ GUI verfügbar")


# -----------------------------------------------------------------------------
# DummyQueue
# -----------------------------------------------------------------------------
class DummyQueue:
    def __init__(self, maxsize: int = 0) -> None:
        self.maxsize = maxsize
        self._items: List[Any] = []
        self._lock = threading.Lock()
        self.Empty = queue.Empty

    def put(
        self, item: Any, block: bool = True, timeout: Optional[float] = None
    ) -> None:
        with self._lock:
            self._items.append(item)
            if self.maxsize > 0 and len(self._items) > self.maxsize:
                self._items.pop(0)

    def get(self, block: bool = True, timeout: Optional[float] = None) -> Any:
        with self._lock:
            if self._items:
                return self._items.pop(0)
            raise self.Empty()

    def put_nowait(self, item: Any) -> None:
        self.put(item, block=False)

    def get_nowait(self) -> Any:
        return self.get(block=False)

    def empty(self) -> bool:
        with self._lock:
            return len(self._items) == 0

    def qsize(self) -> int:
        with self._lock:
            return len(self._items)


# -----------------------------------------------------------------------------
# Datenklassen für Ergebnisse
# -----------------------------------------------------------------------------
@dataclass
class TranscriptionResult:
    text: str
    confidence: float
    language: str = "unknown"
    timestamp: float = field(default_factory=time.time)
    start: Optional[float] = None
    end: Optional[float] = None


@dataclass
class TranslationResult:
    original: str
    translated: str
    source_lang: str = "auto"
    target_lang: str = "de"
    timestamp: float = field(default_factory=time.time)
    start: Optional[float] = None
    end: Optional[float] = None


@dataclass
class StreamInfo:
    title: str
    uploader: str
    duration: str
    view_count: int
    platform: str
    description: str = ""
    duration_seconds: Optional[float] = None
    is_live: bool = False
    thumbnail: str = ""
    original_url: str = ""
    stream_url: Optional[str] = None


# -----------------------------------------------------------------------------
# Hilfsklassen für die Transkription (Wrapper für verschiedene Backends)
# -----------------------------------------------------------------------------
class _EmptyInfo:
    language = "unknown"
    duration = 0.0


class _UniversalInfo:
    def __init__(self, result_dict: Dict[str, Any]) -> None:
        self.language = result_dict.get("language", "unknown")
        self.duration = result_dict.get("duration", 0.0)


class _UniversalSegment:
    def __init__(self, seg_dict: Dict[str, Any]) -> None:
        self.text = seg_dict.get("text", "").strip()
        self.start = seg_dict.get("start", 0.0)
        self.end = seg_dict.get("end", 0.0)
        self.confidence = seg_dict.get("confidence", 0.9)


class _EmergencySegment:
    def __init__(self, seg_dict: Dict[str, Any]) -> None:
        self.text = seg_dict.get("text", "")
        self.start = seg_dict.get("start", 0.0)
        self.end = seg_dict.get("end", 0.0)
        self.confidence = 0.5


# -----------------------------------------------------------------------------
# SimplePerformanceTracker
# -----------------------------------------------------------------------------
class SimplePerformanceTracker:
    def __init__(self) -> None:
        self.transcription_count = 0
        self.translation_count = 0
        self.start_time = time.time()
        self.cache_hits = 0
        self.cache_misses = 0
        self._lock = threading.RLock()

    def log_transcription(self) -> None:
        with self._lock:
            self.transcription_count += 1

    def log_translation(self) -> None:
        with self._lock:
            self.translation_count += 1

    def log_cache_hit(self) -> None:
        with self._lock:
            self.cache_hits += 1

    def log_cache_miss(self) -> None:
        with self._lock:
            self.cache_misses += 1

    def get_basic_stats(self) -> Dict[str, Any]:
        with self._lock:
            uptime_minutes = (time.time() - self.start_time) / 60
            total_cache = self.cache_hits + self.cache_misses
            cache_hit_rate = self.cache_hits / total_cache if total_cache > 0 else 0
            return {
                "transcriptions": self.transcription_count,
                "translations": self.translation_count,
                "uptime_minutes": uptime_minutes,
                "cache_hits": self.cache_hits,
                "cache_misses": self.cache_misses,
                "cache_hit_rate": f"{cache_hit_rate:.1%}",
                "timestamp": datetime.now().isoformat(),
            }


T = TypeVar("T")


class OptimizedThreadPoolExecutor:
    """
    Ein optimierter ThreadPoolExecutor mit erweiterten Funktionen:
    - Timeout‑fähige Aufgaben (submit_with_timeout)
    - Automatische Wiederholungen bei Fehlern (submit_with_retry)
    - Parallele Verarbeitung mit item‑bezogenem Timeout (map_with_timeout)
    """

    def __init__(self, max_workers: int = 4, thread_name_prefix: str = "OptExec"):
        self._max_workers = max_workers
        self._thread_name_prefix = thread_name_prefix
        self._executor = ThreadPoolExecutor(
            max_workers=max_workers, thread_name_prefix=thread_name_prefix
        )
        self._shutdown_lock = threading.RLock()
        self._shutdown = False

    @property
    def max_workers(self) -> int:
        return self._max_workers

    def submit(self, fn: Callable[..., T], *args: Any, **kwargs: Any) -> Future[T]:
        with self._shutdown_lock:
            if self._shutdown:
                raise RuntimeError("Executor wurde bereits heruntergefahren")
            return self._executor.submit(fn, *args, **kwargs)

    def submit_with_timeout(
        self,
        fn: Callable[..., T],
        timeout: Optional[float] = None,
        *args: Any,
        **kwargs: Any,
    ) -> T:
        future = self.submit(fn, *args, **kwargs)
        try:
            return future.result(timeout=timeout)
        except FutureTimeout:
            future.cancel()
            raise TimeoutError(
                f"Timeout nach {timeout}s bei Ausführung von {fn.__name__}"
            )

    def submit_with_retry(
        self,
        fn: Callable[..., T],
        max_retries: int = 3,
        retry_delay_base: float = 1.0,
        retry_delay_max: float = 10.0,
        timeout_per_attempt: Optional[float] = None,
        *args: Any,
        **kwargs: Any,
    ) -> T:
        last_exception = None
        for attempt in range(1, max_retries + 1):
            try:
                if timeout_per_attempt is not None:
                    result = self.submit_with_timeout(
                        fn, timeout_per_attempt, *args, **kwargs
                    )
                else:
                    result = self.submit(fn, *args, **kwargs).result()
                logger.debug(
                    f"Funktion {fn.__name__} erfolgreich nach {attempt} Versuchen"
                )
                return result
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                last_exception = e
                if attempt < max_retries:
                    delay = min(retry_delay_max, retry_delay_base * (2 ** (attempt - 1)))
                    logger.warning(
                        f"Funktion {fn.__name__} fehlgeschlagen (Versuch {attempt}/{max_retries}): {e}. "
                        f"Nächster Versuch in {delay:.1f}s"
                    )
                    time.sleep(delay)
                else:
                    logger.error(
                        f"Funktion {fn.__name__} auch nach {max_retries} Versuchen fehlgeschlagen: {e}"
                    )
        raise last_exception if last_exception is not None else RuntimeError("Unbekannter Fehler")

    def map_with_timeout(
        self,
        fn: Callable[..., T],
        iterable: Union[list, tuple],
        timeout_per_item: Optional[float] = None,
        *args: Any,
        **kwargs: Any,
    ) -> list[T]:
        futures = [self.submit(fn, item, *args, **kwargs) for item in iterable]
        results = []
        for i, future in enumerate(futures):
            try:
                results.append(future.result(timeout=timeout_per_item))
            except FutureTimeout:
                for f in futures[i:]:
                    f.cancel()
                raise TimeoutError(
                    f"Timeout nach {timeout_per_item}s bei Element {i} in {fn.__name__}"
                )
            except Exception:
                for f in futures[i:]:
                    f.cancel()
                raise
        return results

    def shutdown(self, wait: bool = True, cancel_futures: bool = False) -> None:
        with self._shutdown_lock:
            if self._shutdown:
                return
            self._shutdown = True
        try:
            if cancel_futures and hasattr(self._executor, "shutdown"):
                self._executor.shutdown(wait=wait, cancel_futures=cancel_futures)
            else:
                self._executor.shutdown(wait=wait)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"Fehler beim Herunterfahren des Executors: {e}")

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.shutdown(wait=True)
        return False


# -----------------------------------------------------------------------------
# TTLCache (Thread-sicher, optimiert)
# -----------------------------------------------------------------------------
class TTLCache:
    def __init__(
        self, maxsize: int = 128, ttl: float = 300.0, cleanup_interval: int = 100
    ):
        self.maxsize = maxsize
        self.ttl = ttl
        self._cleanup_interval = cleanup_interval
        self._cache: OrderedDict[str, Tuple[Any, float]] = OrderedDict()
        self._lock = threading.RLock()
        self._access_counter = 0

    def _clear_expired(self) -> None:
        now = time.time()
        expired = [k for k, (v, ts) in self._cache.items() if now - ts > self.ttl]
        for k in expired:
            del self._cache[k]

    def get(self, key: str) -> Optional[Any]:
        with self._lock:
            self._access_counter += 1
            if self._access_counter % self._cleanup_interval == 0:
                self._clear_expired()

            if key in self._cache:
                self._cache.move_to_end(key)
                return self._cache[key][0]
            return None

    def put(self, key: str, value: Any) -> None:
        with self._lock:
            self._access_counter += 1
            if self._access_counter % self._cleanup_interval == 0:
                self._clear_expired()

            now = time.time()
            if key in self._cache:
                self._cache.move_to_end(key)
                self._cache[key] = (value, now)
                return

            if len(self._cache) >= self.maxsize:
                self._cache.popitem(last=False)
            self._cache[key] = (value, now)

    def clear(self) -> None:
        with self._lock:
            self._cache.clear()

    def clear_expired(self) -> int:
        with self._lock:
            before = len(self._cache)
            self._clear_expired()
            return before - len(self._cache)

    def get_stats(self) -> Dict[str, Any]:
        with self._lock:
            self._clear_expired()
            return {
                "size": len(self._cache),
                "maxsize": self.maxsize,
                "ttl": self.ttl,
            }


# -----------------------------------------------------------------------------
# CacheManager – kapselt alle globalen Caches
# -----------------------------------------------------------------------------
class CacheManager:
    def __init__(self) -> None:
        self.transcription_cache = TTLCache(
            maxsize=Constants.TRANSCRIPTION_CACHE_SIZE,
            ttl=Constants.TRANSCRIPTION_CACHE_TTL,
        )
        self.translation_cache = TTLCache(
            maxsize=Constants.TRANSLATION_CACHE_SIZE,
            ttl=Constants.TRANSLATION_CACHE_TTL,
        )
        self.audio_cache = TTLCache(
            maxsize=Constants.AUDIO_CACHE_SIZE, ttl=Constants.AUDIO_CACHE_TTL
        )

    def clear_expired_entries(self) -> Dict[str, int]:
        return {
            "transcription_expired": self.transcription_cache.clear_expired(),
            "translation_expired": self.translation_cache.clear_expired(),
            "audio_expired": self.audio_cache.clear_expired(),
        }

    def get_stats(self) -> Dict[str, Any]:
        return {
            "transcription_cache": self.transcription_cache.get_stats(),
            "translation_cache": self.translation_cache.get_stats(),
            "audio_cache": self.audio_cache.get_stats(),
        }

    def cache_transcription(self, result: TranscriptionResult) -> str:
        key = hashlib.sha256(f"{result.text}:{result.language}".encode()).hexdigest()
        self.transcription_cache.put(key, result)
        return key

    def get_cached_transcription(
        self, text: str, language: str = "unknown"
    ) -> Optional[TranscriptionResult]:
        key = hashlib.sha256(f"{text}:{language}".encode()).hexdigest()
        return self.transcription_cache.get(key)

    def cache_translation(self, result: TranslationResult) -> str:
        key = hashlib.sha256(
            (result.original + result.target_lang).encode()
        ).hexdigest()
        self.translation_cache.put(key, result)
        return key

    def get_cached_translation(
        self, original: str, target_lang: str
    ) -> Optional[TranslationResult]:
        key = hashlib.sha256((original + target_lang).encode()).hexdigest()
        return self.translation_cache.get(key)


# -----------------------------------------------------------------------------
# AppContext – zentraler Zustandscontainer (ersetzt globale Variablen)
# -----------------------------------------------------------------------------
class AppContext:
    _instance: Optional["AppContext"] = None
    _lock = threading.Lock()

    def __new__(cls) -> "AppContext":
        with cls._lock:
            if cls._instance is None:
                cls._instance = super().__new__(cls)
                cls._instance._initialized = False
            return cls._instance

    def __init__(self) -> None:
        if self._initialized:
            return
        self._initialized = True
        self.theme = DarkTheme()
        self.cache_manager = CacheManager()
        self._debug_filter = DebugFilter()
        logger.addFilter(self._debug_filter)

    def set_theme(self, theme_name: str) -> None:
        if theme_name == "dark":
            self.theme = DarkTheme()
        elif theme_name == "light":
            self.theme = LightTheme()
        else:
            self.theme = DarkTheme()


# -----------------------------------------------------------------------------
# DebugFilter
# -----------------------------------------------------------------------------
class DebugFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        if DEBUG_LEVEL >= 3:
            return True
        component = getattr(record, "component", None)
        if component and component in DEBUG_COMPONENTS:
            return True
        return record.levelno >= logging.WARNING


def log_debug(component: str, msg: str, *args, **kwargs):
    logger.debug(msg, *args, extra={"component": component}, **kwargs)


# -----------------------------------------------------------------------------
# ThreadPoolExecutor und Decorator
# -----------------------------------------------------------------------------
_EXECUTOR = ThreadPoolExecutor(max_workers=4, thread_name_prefix="ProcExec")


def execution_decorator(timeout: int = 60, max_retries: int = 3) -> Callable:
    def decorator(func: Callable) -> Callable:
        @wraps(func)
        def wrapper(*args: Any, **kwargs: Any) -> Any:
            last_exception: Optional[Exception] = None
            for attempt in range(max_retries + 1):
                try:
                    future = _EXECUTOR.submit(func, *args, **kwargs)
                    return future.result(timeout=timeout)
                except FutureTimeout as e:
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(
                            f"⏰ Timeout attempt {attempt + 1}/{max_retries + 1} for {func.__name__}"
                        )
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(
                            f"⚠️ Exception in {func.__name__}: {str(e)[:100]}"
                        )
                if attempt < max_retries:
                    wait_time = min(30, 2**attempt)
                    time.sleep(wait_time)
                    continue
            if last_exception is not None:
                raise last_exception
            raise RuntimeError(f"Decorator logic failed for {func.__name__}.")

        return wrapper

    return decorator


class ProcessingError(Exception):
    pass


def gui_operation_decorator(func: Callable) -> Callable:
    @wraps(func)
    def wrapper(*args: Any, **kwargs: Any) -> Any:
        try:
            return func(*args, **kwargs)
        except (tk.TclError, RuntimeError):
            return None
        except Exception:
            return None

    return wrapper


# =============================================================================
# MemoryManager (optimiert)
# =============================================================================
class MemoryManager:
    __slots__ = (
        "_buffers", "_buffer_sizes", "_lock", "_max_memory_per_component",
        "_last_gc_time", "_gc_interval", "_ring_buffers", "_ring_buffer_pointers",
        "_ring_buffer_sizes", "_memory_warning_threshold", "_long_term_monitor",
        "_monitoring_active", "_maintenance_thread", "_maintenance_stop",
        "_psutil", "_cache_stats"
    )

    def __init__(self) -> None:
        self._buffers: Dict[str, Deque[str]] = {}
        self._buffer_sizes: Dict[str, int] = {}
        self._lock = threading.RLock()
        self._max_memory_per_component = 100 * 1024 * 1024
        self._last_gc_time = time.time()
        self._gc_interval = 300

        self._ring_buffers: Dict[str, List[Optional[Tuple[str, int]]]] = {}
        self._ring_buffer_pointers: Dict[str, int] = {}
        self._ring_buffer_sizes: Dict[str, int] = {}

        self._memory_warning_threshold = 0.8
        self._long_term_monitor: Deque[Dict[str, Any]] = deque(maxlen=1000)
        self._monitoring_active = True
        self._maintenance_thread: Optional[threading.Thread] = None
        self._maintenance_stop = threading.Event()
        self._psutil = None
        self._start_maintenance()

        self._cache_stats = {"total_allocated": 0, "total_freed": 0}

    def _get_psutil(self):
        if self._psutil is None:
            try:
                import psutil
                self._psutil = psutil
            except ImportError:
                self._psutil = False
        return self._psutil if self._psutil is not False else None

    def _start_maintenance(self) -> None:
        def maintenance_worker() -> None:
            while not self._maintenance_stop.is_set():
                try:
                    if self._maintenance_stop.wait(60):
                        break
                    self._perform_periodic_maintenance()
                    self._perform_memory_health_check()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    log_debug("memory", f"Maintenance worker error: {e}")

        self._maintenance_thread = threading.Thread(
            target=maintenance_worker, daemon=True, name="MemoryMaintenance"
        )
        self._maintenance_thread.start()

    def _perform_periodic_maintenance(self) -> None:
        with self._lock:
            total_memory = sum(self._buffer_sizes.values())
            memory_usage_percent = total_memory / self._max_memory_per_component
            current_time = time.time()
            do_gc = memory_usage_percent > 0.8 or (current_time - self._last_gc_time > self._gc_interval)
            ring_keys = list(self._ring_buffers.keys())

        if do_gc:
            gc.collect()
            log_debug("memory", f"GC executed, memory used: {self.get_total_memory_usage()/1024/1024:.2f} MB")
            with self._lock:
                self._last_gc_time = time.time()

        if memory_usage_percent > 0.8:
            log_debug("memory", f"High buffer memory: {memory_usage_percent:.1%}, starting aggressive cleanup")
            cleanup_thread = threading.Thread(target=self.aggressive_cleanup, daemon=True)
            cleanup_thread.start()

        for comp in ring_keys:
            with self._lock:
                if comp in self._ring_buffers:
                    current_size = self._ring_buffer_sizes.get(comp, 0)
                    if current_size > Constants.MAX_TEXT_LINES // 2:
                        new_size = Constants.MAX_TEXT_LINES // 2
                        self._resize_ring_buffer(comp, new_size)

    def _perform_memory_health_check(self) -> None:
        psutil = self._get_psutil()
        if psutil is None:
            return

        try:
            system_memory = psutil.virtual_memory()
            system_usage_percent = system_memory.percent / 100.0
            process = psutil.Process()
            process_memory = process.memory_info().rss
            process_usage_percent = process_memory / Constants.MAX_MEMORY_USAGE

            sample = {
                "timestamp": time.time(),
                "system_usage": system_usage_percent,
                "process_usage": process_usage_percent,
                "system_mb": system_memory.used // (1024 * 1024),
                "process_mb": process_memory // (1024 * 1024),
            }
            self._long_term_monitor.append(sample)

            log_debug("memory", f"Health: system={system_usage_percent:.1%}, process={process_usage_percent:.1%}")

            if system_usage_percent > self._memory_warning_threshold:
                logger.warning(f"⚠️ High system memory usage: {system_memory.percent:.1f}%")

            if process_usage_percent > self._memory_warning_threshold:
                logger.warning(f"⚠️ High process memory usage: {process_usage_percent:.1%}")
                self.aggressive_cleanup()

            if len(self._long_term_monitor) >= 10:
                recent = list(self._long_term_monitor)[-10:]
                avg_sys = sum(s["system_usage"] for s in recent) / 10
                if avg_sys > 0.75:
                    logger.warning(f"⚠️ Sustained high memory usage: {avg_sys:.1%}")

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            log_debug("memory", f"Health check error: {e}")

    def add_text(self, component: str, text: str) -> None:
        if not text or not text.strip():
            return

        with self._lock:
            if component in self._ring_buffers:
                self._add_to_ring_buffer(component, text)
                return

            if component not in self._buffers:
                self._buffers[component] = deque(maxlen=Constants.MAX_TEXT_LINES)
                self._buffer_sizes[component] = 0

            text_size = len(text.encode("utf-8"))
            current_size = self._buffer_sizes[component]

            if current_size + text_size > self._max_memory_per_component:
                self._optimize_buffer(component)

            self._buffers[component].append(text)
            self._buffer_sizes[component] += text_size
            self._cache_stats["total_allocated"] += text_size

    def _add_to_ring_buffer(self, component: str, text: str) -> None:
        with self._lock:
            if component not in self._ring_buffers:
                buffer_size = Constants.MAX_TEXT_LINES
                self._ring_buffers[component] = [None] * buffer_size
                self._ring_buffer_pointers[component] = 0
                self._ring_buffer_sizes[component] = 0
                self._buffer_sizes[component] = 0

            ring = self._ring_buffers[component]
            ptr = self._ring_buffer_pointers[component]
            text_size = len(text.encode("utf-8"))

            old_entry = ring[ptr]
            if old_entry is not None:
                old_text, old_size = old_entry
                self._buffer_sizes[component] -= old_size
                self._cache_stats["total_freed"] += old_size

            ring[ptr] = (text, text_size)
            self._buffer_sizes[component] += text_size
            self._cache_stats["total_allocated"] += text_size

            self._ring_buffer_pointers[component] = (ptr + 1) % len(ring)
            if self._ring_buffer_sizes[component] < len(ring):
                self._ring_buffer_sizes[component] += 1

    def _optimize_buffer(self, component: str) -> None:
        if component in self._ring_buffers:
            current_size = self._ring_buffer_sizes[component]
            if current_size > Constants.MAX_TEXT_LINES // 2:
                new_size = Constants.MAX_TEXT_LINES // 2
                self._resize_ring_buffer(component, new_size)
            return

        if component in self._buffers:
            keep_ratio = 0.7
            keep_count = int(len(self._buffers[component]) * keep_ratio)
            if keep_count > 0:
                dq = self._buffers[component]
                while len(dq) > keep_count:
                    removed = dq.popleft()
                    self._buffer_sizes[component] -= len(removed.encode("utf-8"))
                log_debug("memory", f"Buffer {component} optimized: {keep_count} entries kept")

    def _resize_ring_buffer(self, component: str, new_size: int) -> None:
        if component not in self._ring_buffers:
            return

        with self._lock:
            old_buffer = self._ring_buffers[component]
            old_pointer = self._ring_buffer_pointers[component]
            old_filled = self._ring_buffer_sizes[component]
            old_capacity = len(old_buffer)

            start_idx = (old_pointer - min(old_filled, new_size)) % old_capacity
            if start_idx < 0:
                start_idx += old_capacity

            new_buffer = [None] * new_size
            new_pointer = 0
            new_filled = 0
            new_total_size = 0

            for i in range(min(old_filled, new_size)):
                idx = (start_idx + i) % old_capacity
                entry = old_buffer[idx]
                if entry is not None:
                    text, size = entry
                    new_buffer[new_pointer] = (text, size)
                    new_total_size += size
                    new_pointer = (new_pointer + 1) % new_size
                    new_filled += 1

            self._ring_buffers[component] = new_buffer
            self._ring_buffer_pointers[component] = new_pointer
            self._ring_buffer_sizes[component] = new_filled
            self._buffer_sizes[component] = new_total_size

            log_debug("memory", f"Ring buffer {component} resized: {old_filled} → {new_filled} entries")

    def get_text(self, component: str) -> str:
        with self._lock:
            if component in self._ring_buffers:
                return self._get_from_ring_buffer(component)
            if component in self._buffers:
                return "\n".join(self._buffers[component])
            return ""

    def _get_from_ring_buffer(self, component: str) -> str:
        with self._lock:
            if component not in self._ring_buffers:
                return ""
            ring = self._ring_buffers[component]
            pointer = self._ring_buffer_pointers[component]
            filled = self._ring_buffer_sizes[component]
            capacity = len(ring)

            if filled == 0:
                return ""

            texts = []
            start = (pointer - filled) % capacity
            for i in range(filled):
                idx = (start + i) % capacity
                entry = ring[idx]
                if entry is not None:
                    texts.append(entry[0])
            return "\n".join(texts)

    def clear_component(self, component: str) -> None:
        with self._lock:
            if component in self._buffers:
                del self._buffers[component]
            if component in self._buffer_sizes:
                del self._buffer_sizes[component]
            if component in self._ring_buffers:
                del self._ring_buffers[component]
            if component in self._ring_buffer_pointers:
                del self._ring_buffer_pointers[component]
            if component in self._ring_buffer_sizes:
                del self._ring_buffer_sizes[component]
            logger.info(f"🧹 Component {component} cleared")

    def aggressive_cleanup(self) -> None:
        logger.info("🧹 Starting aggressive memory cleanup...")

        with self._lock:
            components = list(self._buffers.keys()) + list(self._ring_buffers.keys())

        for component in components:
            try:
                with self._lock:
                    if component in self._buffers:
                        buffer = self._buffers[component]
                        current_len = len(buffer)
                        if current_len > 100:
                            keep_count = max(50, current_len // 2)
                            while len(buffer) > keep_count:
                                removed = buffer.popleft()
                                self._buffer_sizes[component] -= len(removed.encode("utf-8"))
                            logger.info(f"  ↪ {component}: {current_len} → {len(buffer)} entries")
                    elif component in self._ring_buffers:
                        current_size = self._ring_buffer_sizes[component]
                        if current_size > 100:
                            keep_count = max(50, current_size // 2)
                            self._resize_ring_buffer(component, keep_count)
                            logger.info(f"  ↪ {component} (ring): {current_size} → {keep_count} entries")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Cleanup error for {component}: {e}")

        threading.Thread(target=gc.collect, daemon=True).start()
        logger.info("✅ Aggressive cleanup completed")

    def get_memory_stats(self) -> Dict[str, Any]:
        psutil = self._get_psutil()
        if psutil is None:
            return {}

        try:
            system_memory = psutil.virtual_memory()
            process = psutil.Process()
            process_memory = process.memory_info().rss

            with self._lock:
                buffer_count = len(self._buffers)
                ring_buffer_count = len(self._ring_buffers)
                total_buffer_size = sum(self._buffer_sizes.values())

            return {
                "system_usage_percent": system_memory.percent,
                "system_used_mb": system_memory.used // (1024 * 1024),
                "system_total_mb": system_memory.total // (1024 * 1024),
                "process_usage_percent": (process_memory / Constants.MAX_MEMORY_USAGE) * 100,
                "process_used_mb": process_memory // (1024 * 1024),
                "process_peak_mb": self._get_peak_memory() // (1024 * 1024),
                "long_term_samples": len(self._long_term_monitor),
                "buffer_components": buffer_count,
                "ring_buffer_components": ring_buffer_count,
                "total_buffer_size_mb": total_buffer_size // (1024 * 1024),
                "active_monitoring": self._monitoring_active,
            }
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            log_debug("memory", f"Stats error: {e}")
            return {}

    def _get_peak_memory(self) -> int:
        psutil = self._get_psutil()
        if psutil is None:
            return 0
        try:
            process = psutil.Process()
            return process.memory_info().rss
        except Exception:
            return 0

    def get_buffer_stats(self, component: str) -> Dict[str, Any]:
        with self._lock:
            if component in self._ring_buffers:
                return {
                    "type": "ring_buffer",
                    "size": self._ring_buffer_sizes.get(component, 0),
                    "capacity": len(self._ring_buffers[component]),
                    "memory_bytes": self._buffer_sizes.get(component, 0),
                    "pointer": self._ring_buffer_pointers.get(component, 0),
                }
            if component in self._buffers:
                return {
                    "type": "deque",
                    "size": len(self._buffers[component]),
                    "capacity": Constants.MAX_TEXT_LINES,
                    "memory_bytes": self._buffer_sizes.get(component, 0),
                    "maxlen": Constants.MAX_TEXT_LINES,
                }
            return {"type": "not_found"}

    def list_components(self) -> List[str]:
        with self._lock:
            return list(set(self._buffers.keys()) | set(self._ring_buffers.keys()))

    def get_total_memory_usage(self) -> int:
        with self._lock:
            return sum(self._buffer_sizes.values())

    def optimize_all_buffers(self) -> None:
        logger.info("🧹 Optimizing all buffers...")
        with self._lock:
            components = list(self._buffers.keys()) + list(self._ring_buffers.keys())
        for comp in components:
            self._optimize_buffer(comp)
        gc.collect()
        logger.info("✅ All buffers optimized")

    def print_debug_info(self) -> None:
        stats = self.get_memory_stats()
        logger.info("\n" + "=" * 50)
        logger.info("🧠 MEMORY MANAGER DEBUG INFO")
        logger.info("=" * 50)
        logger.info(f"System Memory: {stats.get('system_used_mb', 0)}MB / "
                    f"{stats.get('system_total_mb', 0)}MB ({stats.get('system_usage_percent', 0):.1f}%)")
        logger.info(f"Process Memory: {stats.get('process_used_mb', 0)}MB "
                    f"({stats.get('process_usage_percent', 0):.1f}%)")
        logger.info(f"Buffer Components: {stats.get('buffer_components', 0)}")
        logger.info(f"Ring Buffer Components: {stats.get('ring_buffer_components', 0)}")
        logger.info(f"Total Buffer Size: {stats.get('total_buffer_size_mb', 0)}MB")
        logger.info(f"Long Term Samples: {stats.get('long_term_samples', 0)}")
        components = self.list_components()
        if components:
            logger.info(f"\nActive Components ({len(components)}):")
            for comp in components[:5]:
                comp_stats = self.get_buffer_stats(comp)
                logger.info(f"  • {comp}: {comp_stats['type']}, "
                            f"size: {comp_stats.get('size', 0)}")
            if len(components) > 5:
                logger.info(f"  ... and {len(components) - 5} more")
        logger.info("=" * 50)

    def dispose(self) -> None:
        self._monitoring_active = False
        self._maintenance_stop.set()
        if self._maintenance_thread and self._maintenance_thread.is_alive():
            self._maintenance_thread.join(timeout=1.0)

        with self._lock:
            self._buffers.clear()
            self._buffer_sizes.clear()
            self._ring_buffers.clear()
            self._ring_buffer_pointers.clear()
            self._ring_buffer_sizes.clear()
            self._long_term_monitor.clear()

        gc.collect()
        logger.info("✅ MemoryManager disposed")


# =============================================================================
# 5. BASISKLASSEN UND MIXINS
# =============================================================================


class BaseTranslationEngine(ABC):
    @abstractmethod
    def set_target_language(self, target_lang: str) -> None:
        pass

    @abstractmethod
    def translate_text(
        self, text: str, source_lang: str = "auto"
    ) -> Optional[TranslationResult]:
        pass

    @abstractmethod
    def dispose(self) -> None:
        pass


# -----------------------------------------------------------------------------
# Mixin für Fehlerzähler und Deaktivierung
# -----------------------------------------------------------------------------
class ErrorHandlingMixin:
    def __init__(self, max_errors: int = 5, disable_duration: float = 300.0):
        self._error_count = 0
        self._disabled_until = 0.0
        self._max_errors = max_errors
        self._disable_duration = disable_duration
        self._error_lock = threading.RLock()

    def _check_disable(self):
        with self._error_lock:
            if self._error_count >= self._max_errors and self._disabled_until == 0.0:
                self._disabled_until = time.time() + self._disable_duration
                logger.warning(
                    f"⚠️ {self.__class__.__name__} vorübergehend deaktiviert für {self._disable_duration}s wegen {self._error_count} Fehlern"
                )

    def is_functional(self) -> bool:
        with self._error_lock:
            if self._disabled_until > time.time():
                return False
            if self._error_count >= self._max_errors:
                if time.time() >= self._disabled_until:
                    self._error_count = 0
                    self._disabled_until = 0.0
                    self._reinitialize()
                else:
                    return False
            return True

    def _reinitialize(self):
        pass

    def _record_success(self):
        with self._error_lock:
            self._error_count = 0
            self._disabled_until = 0.0

    def _record_error(self):
        with self._error_lock:
            self._error_count += 1
            self._check_disable()


# -----------------------------------------------------------------------------
# Gemeinsame Basisklasse für alle Übersetzungs-Engines mit Caching und Retry
# -----------------------------------------------------------------------------
class BaseCachedTranslationEngine(BaseTranslationEngine, ErrorHandlingMixin):
    def __init__(
        self,
        default_target_lang: str = "de",
        settings: Optional["AdvancedSettings"] = None,
        max_retries: int = 3,
        retry_delay_base: float = 1.0,
        retry_delay_max: float = 5.0,
        cache_manager: Optional[CacheManager] = None,
    ):
        BaseTranslationEngine.__init__(self)
        ErrorHandlingMixin.__init__(self, max_errors=5, disable_duration=300.0)
        self.default_target_lang = default_target_lang
        self.settings = settings or AdvancedSettings()
        self._max_retries = max_retries
        self._retry_delay_base = retry_delay_base
        self._retry_delay_max = retry_delay_max
        self._cache_manager = cache_manager or AppContext().cache_manager
        self._lock = threading.RLock()
        self._last_translations: Deque[str] = deque(maxlen=15)
        self.last_detected_language = "auto"

        self._low_quality_count = 0
        self._max_low_quality = 3

    def set_target_language(self, target_lang: str) -> None:
        if target_lang != self.default_target_lang:
            self.default_target_lang = target_lang
            with self._lock:
                self._cache_manager.translation_cache.clear()
                self._last_translations.clear()
                self._error_count = 0
                self._disabled_until = 0.0

    @abstractmethod
    def _call_translation_api(
        self, text: str, source_lang: str, target_lang: str
    ) -> Optional[str]:
        pass

    def _is_valid_translation(self, original: str, translated: str) -> bool:
        if not translated or len(translated) < 2:
            return False
        if translated.lower() == original.lower():
            return False
        if len(translated) > len(original) * 8:
            return False
        return True

    def _rate_translation_quality(self, original: str, translated: str) -> float:
        if not translated or len(translated) < 2:
            return 0.0
        orig_words = len(original.split())
        trans_words = len(translated.split())
        if trans_words < 2 and orig_words > 2:
            return 0.0
        if orig_words > 0:
            ratio = trans_words / orig_words
            if ratio < 0.2 or ratio > 5.0:
                return 0.0
        alpha_orig = sum(c.isalpha() for c in original) / max(len(original), 1)
        alpha_trans = sum(c.isalpha() for c in translated) / max(len(translated), 1)
        if alpha_trans < 0.3 and alpha_orig > 0.5:
            return 0.0
        return 1.0

    def translate_text(
        self, text: str, source_lang: str = "auto", target_lang: Optional[str] = None
    ) -> Optional[TranslationResult]:
        if not self.is_functional():
            log_debug(
                "translate",
                f"{self.__class__.__name__} derzeit deaktiviert – überspringe",
            )
            return None

        target = target_lang or self.default_target_lang

        if source_lang != "auto" and source_lang == target:
            return None

        if not text:
            return None
        original_text = text.strip()
        if len(original_text) < 2:
            return None

        with self._lock:
            if original_text in self._last_translations:
                return None
            cached = self._cache_manager.get_cached_translation(original_text, target)
            if cached is not None:
                return cached

        logger.debug(f"translate_text: Original text length = {len(original_text)}")

        last_exception = None
        for attempt in range(self._max_retries):
            try:
                translated_text = self._call_translation_api(
                    original_text, source_lang, target
                )
                logger.debug(f"translate_text: Translated text length = {len(translated_text) if translated_text else 0}")
                if not translated_text or not translated_text.strip():
                    time.sleep(self._retry_delay_base * (2**attempt))
                    continue

                if not self._is_valid_translation(original_text, translated_text):
                    continue

                result = TranslationResult(
                    original=original_text,
                    translated=translated_text.strip(),
                    source_lang=source_lang,
                    target_lang=target,
                )

                quality = self._rate_translation_quality(original_text, translated_text)
                if quality < 0.5:
                    self._record_error()
                else:
                    self._record_success()

                with self._lock:
                    self._cache_manager.cache_translation(result)
                    self._last_translations.append(original_text)
                return result

            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                last_exception = e
                self._record_error()
                if attempt < self._max_retries - 1:
                    delay = min(
                        self._retry_delay_max, self._retry_delay_base * (2**attempt)
                    )
                    time.sleep(delay)

        if last_exception:
            logger.warning(
                f"Übersetzung fehlgeschlagen nach {self._max_retries} Versuchen: {last_exception}"
            )
        return None

    def dispose(self) -> None:
        with self._lock:
            self._cache_manager.translation_cache.clear()
            self._last_translations.clear()
            self._error_count = 0
            self._disabled_until = 0.0
            self._low_quality_count = 0
            gc.collect()


# -----------------------------------------------------------------------------
# Basisklasse für Plugins
# -----------------------------------------------------------------------------
class Plugin(ABC):
    def __init__(self, name: str, version: str = "1.0.0", max_errors: int = 3):
        self.name = name
        self.version = version
        self.enabled = True
        self._disabled = False
        self._error_count = 0
        self._max_errors = max_errors
        self.config: Dict[str, Any] = {}
        self._lock = threading.RLock()

    def _handle_error(self, method_name: str, error: Exception) -> None:
        with self._lock:
            self._error_count += 1
            logger.warning(
                f"Plugin '{self.name}' Fehler in {method_name}: {error} "
                f"({self._error_count}/{self._max_errors})"
            )
            if self._error_count >= self._max_errors:
                self._disabled = True
                logger.error(
                    f"❌ Plugin '{self.name}' wurde wegen zu vieler Fehler deaktiviert. "
                    f"Bitte überprüfen Sie das Plugin oder laden Sie es neu."
                )

    def is_functional(self) -> bool:
        with self._lock:
            return self.enabled and not self._disabled

    def reset_errors(self) -> None:
        with self._lock:
            self._error_count = 0
            self._disabled = False
            logger.info(f"Plugin '{self.name}' Fehlerzähler zurückgesetzt.")

    def on_load(self, manager: "PluginManager") -> None:
        try:
            self._on_load_impl(manager)
        except Exception as e:
            self._handle_error("on_load", e)

    def _on_load_impl(self, manager: "PluginManager") -> None:
        pass

    def on_unload(self) -> None:
        try:
            self._on_unload_impl()
        except Exception as e:
            self._handle_error("on_unload", e)

    def _on_unload_impl(self) -> None:
        pass

    def on_config_change(self, new_config: Dict[str, Any]) -> None:
        try:
            with self._lock:
                self.config.update(new_config)
            self._on_config_change_impl(new_config)
        except Exception as e:
            self._handle_error("on_config_change", e)

    def _on_config_change_impl(self, new_config: Dict[str, Any]) -> None:
        pass

    def on_start(self) -> None:
        try:
            self._on_start_impl()
        except Exception as e:
            self._handle_error("on_start", e)

    def _on_start_impl(self) -> None:
        pass

    def on_stop(self) -> None:
        try:
            self._on_stop_impl()
        except Exception as e:
            self._handle_error("on_stop", e)

    def _on_stop_impl(self) -> None:
        pass

    def on_transcription(self, result: TranscriptionResult) -> TranscriptionResult:
        if not self.is_functional():
            return result
        try:
            return self._on_transcription_impl(result)
        except Exception as e:
            self._handle_error("on_transcription", e)
            return result

    def _on_transcription_impl(
        self, result: TranscriptionResult
    ) -> TranscriptionResult:
        return result

    def on_translation(self, result: TranslationResult) -> TranslationResult:
        if not self.is_functional():
            return result
        try:
            return self._on_translation_impl(result)
        except Exception as e:
            self._handle_error("on_translation", e)
            return result

    def _on_translation_impl(self, result: TranslationResult) -> TranslationResult:
        return result

    def get_config_ui(self, parent) -> Optional[Any]:
        try:
            return self._get_config_ui_impl(parent)
        except Exception as e:
            self._handle_error("get_config_ui", e)
            return None

    def _get_config_ui_impl(self, parent) -> Optional[Any]:
        return None

    def __repr__(self) -> str:
        status = "enabled" if self.is_functional() else "disabled"
        return f"<Plugin {self.name} v{self.version} {status}>"


# =============================================================================
# 6. ÜBERSETZUNGS-ENGINES (abgeleitet von BaseCachedTranslationEngine)
# =============================================================================


class GoogleTranslationEngine(BaseCachedTranslationEngine):
    __slots__ = ("translator",)

    _preprocess_rules = [
        (r"\s+", " "),
        (r"[ ]+([.,!?])", r"\1"),
        (r"([.,!?])[ ]*", r"\1 "),
        ("bass communi", "best community"),
        (" ,", ","),
        (" .", "."),
        ("„", '"'),
        ("“", '"'),
        (r"(?<=[a-zA-Z])\.(?=[a-zA-Z])", " "),
    ]
    _postprocess_rules = [
        (r"\s+\.", "."),
        (r"\s+,", ","),
        (r"\s+\?", "?"),
        (r"\s+!", "!"),
        (" ,", ","),
        (r" \.", "."),
    ]
    _preprocess_patterns = [
        (re.compile(p, re.UNICODE), r) for p, r in _preprocess_rules
    ]
    _postprocess_patterns = [
        (re.compile(p, re.UNICODE), r) for p, r in _postprocess_rules
    ]

    _GOOGLE_LANG_MAP = {
        'zh': 'zh-CN',
        'he': 'iw',
    }

    def _map_to_google_code(self, lang_code: str) -> str:
        if lang_code == 'auto':
            return 'auto'
        return self._GOOGLE_LANG_MAP.get(lang_code, lang_code)

    def __init__(
        self,
        target_lang: str = "de",
        settings: Optional["AdvancedSettings"] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        super().__init__(
            target_lang,
            settings,
            max_retries=3,
            retry_delay_base=1.0,
            retry_delay_max=5.0,
            cache_manager=cache_manager,
        )
        self.translator = None
        self._setup_translator()

    def _setup_translator(self) -> None:
        with self._lock:
            try:
                if TRANSLATOR_AVAILABLE:
                    GoogleTranslator = FastLazyLoader.load("deep_translator")
                    target = self._map_to_google_code(self.default_target_lang)
                    self.translator = GoogleTranslator(
                        source="auto", target=target, timeout=10
                    )
                    self._record_success()
                else:
                    self.translator = None
            except ImportError as e:
                logger.warning(f"deep_translator nicht verfügbar: {e}")
                self.translator = None
                self._record_error()

    def _reinitialize(self):
        self._setup_translator()

    def set_target_language(self, target_lang: str) -> None:
        super().set_target_language(target_lang)
        self._setup_translator()

    def _contains_asian(self, text: str) -> bool:
        asian_ranges = [
            (0x4E00, 0x9FFF),
            (0x3400, 0x4DBF),
            (0x20000, 0x2A6DF),
            (0x2A700, 0x2B73F),
            (0x2B740, 0x2B81F),
            (0x2B820, 0x2CEAF),
            (0xF900, 0xFAFF),
            (0xAC00, 0xD7AF),
            (0x1100, 0x11FF),
            (0x3130, 0x318F),
            (0x3040, 0x309F),
            (0x30A0, 0x30FF),
            (0x31F0, 0x31FF),
            (0x0E00, 0x0E7F),
            (0x0E80, 0x0EFF),
            (0x1000, 0x109F),
            (0x1780, 0x17FF),
            (0x1950, 0x19DF),
            (0x1980, 0x19DF),
            (0x1A20, 0x1AAF),
            (0xAA60, 0xAA7F),
        ]
        for char in text:
            code = ord(char)
            for low, high in asian_ranges:
                if low <= code <= high:
                    return True
        return False

    def _is_valid_translation(self, original: str, translated: str) -> bool:
        if not translated or not translated.strip():
            return False
        orig_clean = original.strip()
        trans_clean = translated.strip()
        if len(trans_clean) < 1:
            return False
        if trans_clean.isspace():
            return False
        is_asian = self._contains_asian(orig_clean) or self._contains_asian(trans_clean)
        if is_asian:
            if len(trans_clean) <= 1:
                return True
            orig_len = len(orig_clean)
            trans_len = len(trans_clean)
            if orig_len == 0 or trans_len == 0:
                return False
            ratio = trans_len / max(orig_len, 1)
            return 0.05 <= ratio <= 15.0
        else:
            if len(trans_clean) <= 3:
                if len(set(trans_clean)) == 1 and len(trans_clean) > 1:
                    return False
            else:
                if len(set(trans_clean)) < 3:
                    return False
            orig_len = len(orig_clean)
            trans_len = len(trans_clean)
            if orig_len == 0 or trans_len == 0:
                return False
            ratio_valid = 0.1 <= trans_len / max(orig_len, 1) <= 8.0
            punkt_valid = True
            if not is_asian:
                punkt_original = orig_clean.count(".") / max(orig_len, 1)
                punkt_trans = trans_clean.count(".") / max(trans_len, 1)
                if punkt_trans > 0.3 and punkt_original < 0.05:
                    punkt_valid = False
            return ratio_valid and punkt_valid

    def _clean_common_errors(self, text: str) -> str:
        corrections = {
            "bass communi": "best community",
            "thc": "the",
            "thc ": "the ",
            " thc": " the",
        }
        for wrong, right in corrections.items():
            text = text.replace(wrong, right)
        return text

    def _preprocess_text(self, text: str) -> str:
        if not text:
            return ""
        clean_text = text.strip()
        clean_text = re.sub(r"\.{2,}", ".", clean_text)
        for pattern, repl in self._preprocess_patterns:
            clean_text = pattern.sub(repl, clean_text)
        clean_text = self._clean_common_errors(clean_text)
        if len(clean_text.split()) < 1:
            return ""
        return clean_text.strip()

    def _postprocess_translation(self, translated: str, original: str) -> str:
        if not translated:
            return ""
        result = translated.strip()
        if result and result[-1] not in (".", "!", "?", ":", ";"):
            result += "."
        if result and result[0].islower():
            result = result[0].upper() + result[1:]
        result = re.sub(r"\s+", " ", result)
        for pattern, repl in self._postprocess_patterns:
            result = pattern.sub(repl, result)
        result = re.sub(r"(?<=[a-zA-Z])\.(?=[a-zA-Z])", " ", result)
        result = re.sub(r"\s+", " ", result)

        # Alpha-Ratio-Prüfung (deaktiviert für maximale Präzision)
        # alpha_count = sum(c.isalpha() for c in result)
        # total_len = len(result)
        # alpha_ratio = alpha_count / max(total_len, 1)
        # if alpha_ratio < 0.5 and total_len > 10:
        #     return ""

        return result.strip()

    def _call_translation_api(
        self, text: str, source_lang: str, target_lang: str
    ) -> Optional[str]:
        with self._lock:
            if not self.translator:
                self._setup_translator()
                if not self.translator:
                    return None

            clean_text = self._preprocess_text(text)
            if not clean_text:
                return None
            logger.debug(f"Google _call: preprocessed text length = {len(clean_text)}")

            try:
                src = self._map_to_google_code(source_lang)
                translated = self.translator.translate(clean_text, source=src)
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("translate", f"Google raw translation: {translated}")
                if not translated:
                    return None
                final = self._postprocess_translation(translated, clean_text)
                return final
            except ValueError as e:
                if "language" in str(e).lower() and src != 'auto':
                    logger.warning(f"Ungültiger Quellcode {src}, versuche mit 'auto'")
                    try:
                        translated = self.translator.translate(clean_text, source='auto')
                        if translated:
                            final = self._postprocess_translation(translated, clean_text)
                            return final
                    except Exception as fallback_e:
                        if isinstance(fallback_e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.warning(f"Auch Fallback fehlgeschlagen: {fallback_e}")
                raise
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("translate", f"Google translation error: {e}")
                raise

    def dispose(self) -> None:
        super().dispose()
        self.translator = None


# -----------------------------------------------------------------------------
# OllamaTranslationEngine (optimiert: Formatierung erhalten)
# -----------------------------------------------------------------------------
class OllamaTranslationEngine(BaseCachedTranslationEngine):
    def __init__(
        self,
        target_lang: str = "de",
        settings: Optional["AdvancedSettings"] = None,
        model: str = "llama3.1:8b",
        host: str = "http://localhost:11434",
        temperature: float = 0.1,
        timeout: int = 30,
        system_prompt: Optional[str] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        super().__init__(
            target_lang,
            settings,
            max_retries=3,
            retry_delay_base=1.0,
            retry_delay_max=5.0,
            cache_manager=cache_manager,
        )
        self.model = model
        self.host = host.rstrip("/")
        self.temperature = temperature
        self.timeout = timeout
        self.system_prompt = system_prompt
        self.available = OLLAMA_AVAILABLE and (requests is not None)
        self._session = None
        if self.available:
            self._session = requests.Session()
            self._session.headers.update(
                {
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                }
            )
        self._available_models: List[str] = []
        self._models_cache_time = 0.0
        self._models_cache_ttl = 300
        self._fetch_available_models()

    def _fetch_available_models(self) -> List[str]:
        if not self.available:
            return []
        with self._lock:
            now = time.time()
            if (
                now - self._models_cache_time > self._models_cache_ttl
                or not self._available_models
            ):
                try:
                    r = self._session.get(f"{self.host}/api/tags", timeout=3)
                    if r.status_code == 200:
                        data = r.json()
                        self._available_models = [
                            m["name"] for m in data.get("models", [])
                        ]
                        self._models_cache_time = now
                    else:
                        if not self._available_models:
                            self._available_models = []
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    log_debug("ollama", f"Model list fetch failed: {e}")
                    pass
            return self._available_models.copy()

    def is_model_available(self, model: Optional[str] = None) -> bool:
        check_model = model or self.model
        available = self._fetch_available_models()
        if not available:
            return True
        return check_model in available

    def set_target_language(self, target_lang: str) -> None:
        super().set_target_language(target_lang)

    def set_model(self, model: str) -> None:
        with self._lock:
            self.model = model
            self._models_cache_time = 0
            self._available_models = []
        self._fetch_available_models()
        logger.info(f"Ollama model geändert zu: {model}")

    def _call_ollama(self, prompt: str) -> Optional[str]:
        if not self.available or self._session is None:
            logger.error("Ollama nicht verfügbar (requests nicht installiert)")
            return None
        if not self.is_model_available():
            logger.warning(f"Ollama-Modell '{self.model}' nicht auf Server gefunden.")
            return None
        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": self.temperature},
        }
        if self.system_prompt:
            payload["system"] = self.system_prompt
        try:
            response = self._session.post(
                f"{self.host}/api/generate", json=payload, timeout=self.timeout
            )
            if logger.isEnabledFor(logging.DEBUG):
                log_debug("ollama", f"Response status: {response.status_code}")
            if response.status_code == 200:
                data = response.json()
                translated = data.get("response", "").strip()
                return translated if translated else None
            else:
                logger.warning(f"Ollama Fehler {response.status_code}: {response.text}")
                return None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if requests is not None:
                if isinstance(e, requests.exceptions.Timeout):
                    logger.warning(f"Ollama Timeout nach {self.timeout}s")
                    return None
                if isinstance(e, requests.exceptions.ConnectionError):
                    logger.warning("Ollama nicht erreichbar (läuft der Server?)")
                    return None
            logger.warning(f"Ollama Fehler: {e}")
            return None

    def _reinitialize(self):
        logger.info("🔄 OllamaTranslationEngine wird reinitialisiert...")
        with self._lock:
            self._models_cache_time = 0
            self._available_models = []
        self._fetch_available_models()
        if self.is_model_available():
            test_result = self._call_ollama("Translate 'Hello' to German.")
            if test_result:
                logger.info(
                    "✅ OllamaTranslationEngine nach Reinitialisierung funktionsfähig."
                )
            else:
                logger.warning(
                    "⚠️ OllamaTranslationEngine reagiert nicht – bleibt möglicherweise deaktiviert."
                )

    def _is_valid_translation(self, original: str, translated: str) -> bool:
        if not translated or len(translated) < 2:
            return False
        if translated.lower() == original.lower():
            return False
        if len(translated) > len(original) * 5:
            return False
        return True

    def _call_translation_api(
        self, text: str, source_lang: str, target_lang: str
    ) -> Optional[str]:
        if not self.available:
            return None
        if not self.is_model_available():
            log_debug("ollama", f"Modell '{self.model}' nicht verfügbar")
            self._record_error()
            return None
        source_lang_name = "auto"
        if source_lang != "auto":
            source_lang_name = SUPPORTED_LANGUAGES.get(source_lang, source_lang)
        target_lang_name = SUPPORTED_LANGUAGES.get(target_lang, target_lang)
        prompt = (
            f"Translate the following text from {source_lang_name} to {target_lang_name}. "
            f"Preserve the original formatting, including line breaks, paragraphs, and bullet points. "
            f"Output only the translation, without any additional commentary.\n\n{text}"
        )
        if logger.isEnabledFor(logging.DEBUG):
            log_debug("ollama", f"Prompt: {prompt}")
        return self._call_ollama(prompt)

    def dispose(self) -> None:
        super().dispose()
        if self._session:
            try:
                self._session.close()
            except Exception:
                pass
            self._session = None
        self._available_models.clear()


# -----------------------------------------------------------------------------
# ArgosTranslateEngine
# -----------------------------------------------------------------------------
class ArgosTranslateEngine(BaseCachedTranslationEngine):
    __slots__ = ("_translators", "_package_lock")

    _ARGOS_LANG_MAP = {}

    def _normalize_argos_code(self, lang_code: str) -> str:
        if lang_code == 'auto':
            return 'auto'
        normalized = lang_code.lower().replace('_', '-')
        return self._ARGOS_LANG_MAP.get(normalized, normalized)

    def __init__(
        self,
        target_lang: str = "de",
        settings: Optional["AdvancedSettings"] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        super().__init__(
            target_lang,
            settings,
            max_retries=3,
            retry_delay_base=1.0,
            retry_delay_max=5.0,
            cache_manager=cache_manager,
        )
        self._translators: Dict[str, Any] = {}
        self._package_lock = threading.RLock()
        if not ARGOS_AVAILABLE:
            logger.warning(
                "⚠️ argos-translate nicht installiert – Engine wird nicht funktionieren"
            )
            self._disabled_until = time.time() + self._disable_duration
            self._error_count = self._max_errors
            return
        norm_target = self._normalize_argos_code(self.default_target_lang)
        if norm_target != self.default_target_lang:
            logger.debug(f"Argos: Zielcode normalisiert: {self.default_target_lang} -> {norm_target}")
            self.default_target_lang = norm_target
        self._load_translator(self.default_target_lang)

    def set_target_language(self, target_lang: str) -> None:
        norm_target = self._normalize_argos_code(target_lang)
        if norm_target != target_lang:
            logger.debug(f"Argos: Zielcode normalisiert: {target_lang} -> {norm_target}")
        if norm_target not in SUPPORTED_LANGUAGES and norm_target != 'auto':
            logger.warning(f"Unbekannter Sprachcode '{norm_target}', verwende 'en'")
            norm_target = 'en'
        super().set_target_language(norm_target)
        self._load_translator(self.default_target_lang)

    def _load_translator(self, target_lang: str) -> Optional[Any]:
        if not ARGOS_AVAILABLE:
            return None
        with self._package_lock:
            if target_lang in self._translators:
                return self._translators[target_lang]
            try:
                from argostranslate import package, translate

                installed_languages = translate.get_installed_languages()
                source_lang_obj = None
                target_lang_obj = None
                for lang in installed_languages:
                    if lang.code == target_lang:
                        target_lang_obj = lang
                if target_lang_obj is None:
                    available_packages = package.get_available_packages()
                    matching_packages = [
                        p for p in available_packages if p.to_code == target_lang
                    ]
                    if not matching_packages:
                        logger.warning(
                            f"⚠️ Kein argos-Paket für Zielsprache '{target_lang}' verfügbar"
                        )
                        return None
                    pkg = matching_packages[0]
                    logger.info(f"📦 Installiere argos-Paket: {pkg} ...")
                    try:
                        package.install_from_path(pkg.download())
                        logger.info(f"✅ argos-Paket für {target_lang} installiert")
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.error(f"❌ Installation fehlgeschlagen: {e}")
                        return None
                    translate.load_installed_languages()
                    installed_languages = translate.get_installed_languages()
                    for lang in installed_languages:
                        if lang.code == target_lang:
                            target_lang_obj = lang
                            break
                source_lang_obj = None
                for lang in installed_languages:
                    if lang.code == "en":
                        source_lang_obj = lang
                        break
                if source_lang_obj is None or target_lang_obj is None:
                    logger.warning("⚠️ argos: Quell‑ oder Zielsprache nicht verfügbar")
                    return None
                translator = target_lang_obj.get_translator(source_lang_obj)
                if translator:
                    self._translators[target_lang] = translator
                    return translator
                else:
                    logger.warning(
                        f"⚠️ argos: Kein direkter Übersetzer von 'en' nach '{target_lang}' gefunden"
                    )
                    return None
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Fehler beim Laden des argos-Übersetzers: {e}")
                return None

    def _reinitialize(self) -> None:
        self._load_translator(self.default_target_lang)

    def _load_specific_translator(self, from_lang: str, to_lang: str) -> Optional[Any]:
        norm_from = self._normalize_argos_code(from_lang)
        norm_to = self._normalize_argos_code(to_lang)
        try:
            from argostranslate import translate

            installed_languages = translate.get_installed_languages()
            from_lang_obj = None
            to_lang_obj = None
            for lang in installed_languages:
                if lang.code == norm_from:
                    from_lang_obj = lang
                if lang.code == norm_to:
                    to_lang_obj = lang
            if from_lang_obj and to_lang_obj:
                translator = from_lang_obj.get_translator(to_lang_obj)
                if translator:
                    key = f"{norm_from}_{norm_to}"
                    self._translators[key] = translator
                    return translator
        except Exception:
            pass
        return None

    def _call_translation_api(
        self, text: str, source_lang: str, target_lang: str
    ) -> Optional[str]:
        norm_source = self._normalize_argos_code(source_lang)
        norm_target = self._normalize_argos_code(target_lang)

        translator = self._translators.get(norm_target)
        if translator is None:
            translator = self._load_translator(norm_target)
            if translator is None:
                self._record_error()
                return None

        effective_translator = translator
        if norm_source != "auto" and norm_source != "en":
            specific = self._load_specific_translator(norm_source, norm_target)
            if specific:
                effective_translator = specific

        try:
            translated = effective_translator.translate(text)
            return translated.strip() if translated else None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            log_debug("translate", f"Argos translation error: {e}")
            raise

    def dispose(self) -> None:
        super().dispose()
        with self._lock:
            self._translators.clear()


# -----------------------------------------------------------------------------
# DummyTranslationEngine (für Demo-Modus)
# -----------------------------------------------------------------------------
class DummyTranslationEngine(BaseTranslationEngine):
    def __init__(
        self,
        target_lang: str = "de",
        settings: Optional["AdvancedSettings"] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        self.target_lang = target_lang
        self.settings = settings or AdvancedSettings()
        self._cache_manager = cache_manager
        self._lock = threading.RLock()
        self._last_translations: Deque[str] = deque(maxlen=5)
        self.last_detected_language = "auto"

    def set_target_language(self, target_lang: str) -> None:
        self.target_lang = target_lang

    def translate_text(
        self, text: str, source_lang: str = "auto"
    ) -> Optional[TranslationResult]:
        return TranslationResult(
            original=text,
            translated="[Übersetzung nicht verfügbar]",
            source_lang=source_lang,
            target_lang=self.target_lang,
        )

    def dispose(self) -> None:
        pass


# =============================================================================
# 7. TRANSKRIPTIONS-ENGINE
# =============================================================================
class TranscriptionEngine:
    __slots__ = (
        "model",
        "model_size",
        "whisper_backend",
        "settings",
        "config",
        "device",
        "compute_type",
        "_cache",
        "_lock",
        "_model_loading",
        "_max_cached_models",
        "_model_cache",
        "_model_usage_lock",
        "_performance_monitor",
        "_last_transcription_text",
        "_active_model_loads",
        "_model_loaded_flag",
        "_disposing",
        "forced_language",
        "_last_detected_language",
        "_torch",
        "_np",
        "_scipy_signal",
        "_last_confidence_threshold",
        "_audio_enhancer",
        "_cache_manager",
        "_model_locks",
        "_model_loading_status",
        "_reloading",
        "_vad_fallback_enabled",
    )

    MODEL_SIZE_ORDER = [
        "tiny",
        "tiny.en",
        "base",
        "base.en",
        "small",
        "small.en",
        "medium",
        "medium.en",
        "large-v1",
        "large-v2",
        "large-v3",
    ]

    _ALLOWED_FASTER = {
        "language", "task", "beam_size", "best_of", "temperature",
        "compression_ratio_threshold", "log_prob_threshold", "no_speech_threshold",
        "condition_on_previous_text", "vad_filter", "vad_parameters",
        "without_timestamps", "word_timestamps", "initial_prompt", "prefix",
        "suppress_blank", "suppress_tokens", "hotwords",
        "language_detection_threshold", "language_detection_segments",
        "multilingual", "repetition_penalty", "no_repeat_ngram_size",
        "prompt_reset_on_temperature"
    }

    _ALLOWED_OPENAI = {
        "language", "task", "temperature", "best_of", "beam_size", "patience",
        "length_penalty", "repetition_penalty", "no_repeat_ngram_size",
        "initial_prompt", "prefix", "suppress_tokens", "without_timestamps",
        "max_initial_timestamp", "word_timestamps", "prepend_punctuations",
        "append_punctuations", "max_new_tokens", "clip_timestamps",
        "hallucination_silence_threshold", "compression_ratio_threshold",
        "log_prob_threshold", "no_speech_threshold", "condition_on_previous_text"
    }

    def __init__(
        self,
        settings: Optional["AdvancedSettings"] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        self.settings = settings or AdvancedSettings()
        self.config = self.settings.config
        self.model: Any = None
        self.model_size: Optional[str] = None
        self.whisper_backend: Optional[str] = None
        self._lock = threading.RLock()
        self._model_usage_lock = threading.RLock()
        self._model_loading = False
        self._max_cached_models = 1
        self._cache_manager = cache_manager or AppContext().cache_manager
        self._cache = self._cache_manager.transcription_cache
        self._performance_monitor = SimplePerformanceTracker()
        self._last_transcription_text = ""
        self._active_model_loads: Set[str] = set()
        self._model_loaded_flag = False
        self._disposing = False
        self._model_cache: OrderedDict[Tuple[str, str], Any] = OrderedDict()
        self.forced_language: Optional[str] = None
        self._last_detected_language: Optional[str] = None
        self._torch = None
        self._np = None
        self._scipy_signal = None
        self._last_confidence_threshold = 0.6
        self._audio_enhancer = AudioEnhancer(self.config, self.settings)
        self._model_locks: Dict[str, threading.Lock] = {}
        self._model_loading_status: Dict[str, bool] = {}
        self._reloading = False
        self._vad_fallback_enabled = True

        if TORCH_AVAILABLE:
            self._torch = FastLazyLoader.load("torch")
        if NUMPY_AVAILABLE:
            self._np = FastLazyLoader.load("numpy")
        if SCIPY_AVAILABLE:
            self._scipy_signal = FastLazyLoader.load("scipy.signal")
        self.device, self.compute_type = self._detect_optimal_device()

    def set_vad_fallback_enabled(self, enabled: bool) -> None:
        self._vad_fallback_enabled = enabled
        logger.debug(f"VAD-Fallback in TranscriptionEngine {'aktiviert' if enabled else 'deaktiviert'}")

    def _parse_suppress_tokens(self, token_str: str) -> List[int]:
        if not token_str:
            return [-1]
        try:
            return [int(x.strip()) for x in token_str.split(",")]
        except ValueError:
            logger.warning(f"Ungültiges suppress_tokens Format: {token_str}, verwende [-1]")
            return [-1]

    def _detect_optimal_device(self) -> Tuple[str, str]:
        device = "cpu"
        compute_type = "int8"
        if self._torch is not None:
            torch = self._torch
            if torch.cuda.is_available():
                try:
                    torch.tensor([1.0]).cuda()
                    device = "cuda"
                    compute_type = (
                        "float16" if self.settings.gpu_acceleration else "int8"
                    )
                    logger.info(
                        f"✅ NVIDIA GPU detected: {torch.cuda.get_device_name(0)}"
                    )
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.warning(f"⚠️ CUDA test failed, falling back: {e}")
            if hasattr(torch.version, "hip") and torch.version.hip:
                try:
                    if torch.cuda.device_count() > 0:
                        device = "cuda"
                        compute_type = (
                            "float16" if self.settings.gpu_acceleration else "int8"
                        )
                        logger.info("✅ AMD GPU (ROCm) detected")
                except (AttributeError, RuntimeError) as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.warning(f"⚠️ ROCm test failed: {e}")
            if hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
                device = "mps"
                compute_type = "float16"
                logger.info("✅ Apple Silicon GPU (MPS) detected")
        logger.info(f"✅ Verwende Device: {device} (compute_type={compute_type})")
        return device, compute_type

    def _estimate_model_memory(self, model_size: str) -> float:
        estimates = {
            "tiny": 1.0,
            "tiny.en": 1.0,
            "base": 1.5,
            "base.en": 1.5,
            "small": 2.5,
            "small.en": 2.5,
            "medium": 4.0,
            "medium.en": 4.0,
            "large": 6.0,
            "large-v1": 6.0,
            "large-v2": 6.0,
            "large-v3": 6.0,
        }
        return estimates.get(model_size.lower(), 3.0)

    def _get_free_gpu_memory(self) -> Optional[float]:
        if self.device != "cuda" or self._torch is None:
            return None
        try:
            torch = self._torch
            if not torch.cuda.is_available():
                return None
            device_idx = torch.cuda.current_device()
            allocated = torch.cuda.memory_allocated(device_idx)
            total = torch.cuda.get_device_properties(device_idx).total_memory
            free = (total - allocated) / (1024**3)
            return free
        except Exception:
            return None

    @execution_decorator(timeout=1800.0)
    def load_model(
        self, model_size: str, set_active: bool = False
    ) -> Optional[Tuple[Any, str]]:
        if set_active:
            self._force_model_cleanup()

        if FASTER_WHISPER_AVAILABLE:
            backend = "faster_whisper"
        elif OPENAI_WHISPER_AVAILABLE:
            backend = "openai_whisper"
        else:
            logger.error("❌ Kein Whisper-Backend verfügbar")
            return None

        cache_key = (model_size, backend)

        with self._lock:
            if cache_key in self._model_cache:
                model = self._model_cache[cache_key]
                self._model_cache.move_to_end(cache_key)
                if set_active:
                    with self._model_usage_lock:
                        self.model = model
                        self.model_size = model_size
                        self.whisper_backend = backend
                return model, backend

        with self._lock:
            if model_size not in self._model_locks:
                self._model_locks[model_size] = threading.Lock()
            model_lock = self._model_locks[model_size]

        with model_lock:
            with self._lock:
                if cache_key in self._model_cache:
                    model = self._model_cache[cache_key]
                    self._model_cache.move_to_end(cache_key)
                    if set_active:
                        with self._model_usage_lock:
                            self.model = model
                            self.model_size = model_size
                            self.whisper_backend = backend
                    return model, backend
                if self._model_loading_status.get(model_size, False):
                    logger.info(
                        f"⏳ Modell {model_size} wird bereits in einem anderen Thread geladen – breche ab."
                    )
                    return None
                self._model_loading_status[model_size] = True

        free_gb = self._get_free_gpu_memory()
        need_fallback = None
        if free_gb is not None:
            estimated = self._estimate_model_memory(model_size)
            buffer_gb = max(1.0, free_gb * 0.2)
            required_gb = estimated + buffer_gb

            if free_gb < required_gb:
                logger.warning(
                    f"⚠️ Nur {free_gb:.1f} GB VRAM frei, {model_size} benötigt ~{estimated:.1f} GB + Puffer. "
                    "Versuche zuerst, GPU-Cache zu leeren..."
                )
                if self._torch is not None and self.device == "cuda":
                    self._torch.cuda.empty_cache()
                    time.sleep(0.2)
                    free_gb_after = self._get_free_gpu_memory()
                    if free_gb_after is not None and free_gb_after >= required_gb:
                        logger.info(
                            f"✅ Nach Cache-Leerung: {free_gb_after:.1f} GB frei – genug für {model_size}"
                        )
                        free_gb = free_gb_after
                    else:
                        logger.warning(
                            f"⚠️ Nach Cache-Leerung immer noch nicht genug VRAM: {free_gb_after:.1f} GB. "
                            "Versuche kleineres Modell..."
                        )
                        if model_size in self.MODEL_SIZE_ORDER:
                            current_idx = self.MODEL_SIZE_ORDER.index(model_size)
                            for idx in range(current_idx - 1, -1, -1):
                                smaller = self.MODEL_SIZE_ORDER[idx]
                                smaller_est = self._estimate_model_memory(smaller)
                                smaller_required = smaller_est + buffer_gb
                                if free_gb_after >= smaller_required:
                                    need_fallback = smaller
                                    break
                            if need_fallback is None:
                                logger.error("❌ Nicht genug VRAM für irgendein Modell – breche ab.")
                                with self._lock:
                                    self._model_loading_status[model_size] = False
                                return None
                        else:
                            logger.warning(
                                f"⚠️ Modell {model_size} nicht in MODEL_SIZE_ORDER – Fallback übersprungen."
                            )
                else:
                    if model_size in self.MODEL_SIZE_ORDER:
                        current_idx = self.MODEL_SIZE_ORDER.index(model_size)
                        for idx in range(current_idx - 1, -1, -1):
                            smaller = self.MODEL_SIZE_ORDER[idx]
                            smaller_est = self._estimate_model_memory(smaller)
                            smaller_required = smaller_est + buffer_gb
                            if free_gb >= smaller_required:
                                need_fallback = smaller
                                break
                        if need_fallback is None:
                            logger.error("❌ Nicht genug VRAM für irgendein Modell – breche ab.")
                            with self._lock:
                                self._model_loading_status[model_size] = False
                            return None
                    else:
                        logger.warning(
                            f"⚠️ Modell {model_size} nicht in MODEL_SIZE_ORDER – Fallback übersprungen."
                        )

        if need_fallback is not None:
            logger.info(f"🔄 Fallback auf {need_fallback}")
            result = self.load_model(need_fallback, set_active=set_active)
            with self._lock:
                self._model_loading_status[model_size] = False
            return result

        try:
            logger.info("📁 Modell wird im Standard-Cache von Hugging Face gespeichert.")

            model = None
            if backend == "faster_whisper":
                model = self._load_faster_whisper(model_size)
                if model is None and OPENAI_WHISPER_AVAILABLE:
                    backend = "openai_whisper"
                    cache_key = (model_size, backend)
            if backend == "openai_whisper" and model is None:
                model = self._load_openai_whisper(model_size)

            if model is None:
                logger.error(
                    f"❌ Konnte Modell {model_size} mit Backend {backend} nicht laden."
                )
                return None

            with self._lock:
                self._model_cache[cache_key] = model
                if len(self._model_cache) > self._max_cached_models:
                    oldest_key, old_model = next(iter(self._model_cache.items()))
                    del self._model_cache[oldest_key]
                    logger.info(
                        f"🧹 Entferne ältestes Modell '{oldest_key[0]}' aus Cache"
                    )
                    self._unload_model(old_model)
                if set_active:
                    with self._model_usage_lock:
                        self.model = model
                        self.model_size = model_size
                        self.whisper_backend = backend

            return model, backend

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Unerwarteter Fehler beim Laden von {model_size}: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return None

        finally:
            with self._lock:
                self._model_loading_status[model_size] = False

    def _load_faster_whisper(self, model_size: str) -> Any:
        try:
            from faster_whisper import WhisperModel

            model = WhisperModel(
                model_size,
                device=self.device,
                compute_type=self.compute_type,
                cpu_threads=4,
                num_workers=1,
            )
            if self._np is None:
                self._np = FastLazyLoader.load("numpy")

            logger.info(f"✅ faster-whisper '{model_size}' erfolgreich geladen")
            return model
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ faster-whisper konnte nicht geladen werden: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return None

    def _load_openai_whisper(self, model_size: str) -> Any:
        try:
            import whisper

            device = "cuda" if self.device == "cuda" else "cpu"
            model = whisper.load_model(
                model_size,
                device=device,
                download_root=None,
            )
            logger.info(f"✅ openai-whisper '{model_size}' erfolgreich geladen")
            return model
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ openai-whisper fehlgeschlagen: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return None

    def _unload_model(self, model: Any) -> None:
        if hasattr(model, "unload_model"):
            try:
                model.unload_model()
            except Exception:
                pass

    def reload_model(self, model_size: str) -> bool:
        with self._lock:
            if self._model_loading:
                logger.warning("⚠️ Model loading already in progress")
                return False
            self._model_loading = True

        def _load_in_background():
            try:
                result = self.load_model(model_size, set_active=False)
                if result is None:
                    logger.error(
                        "❌ Background model loading failed (load_model returned None)"
                    )
                    return
                new_model, new_backend = result
                cache_key = (model_size, new_backend)
                with self._lock:
                    with self._model_usage_lock:
                        old_model = self.model
                        old_key = (
                            (self.model_size, self.whisper_backend)
                            if self.model_size
                            else None
                        )
                        self.model = new_model
                        self.model_size = model_size
                        self.whisper_backend = new_backend
                        self._model_loaded_flag = True
                        if old_model is not None and old_key is not None:
                            if old_key in self._model_cache:
                                del self._model_cache[old_key]
                            self._unload_model(old_model)
                            logger.info(
                                f"🧹 Unloaded old model {old_key[0]} ({old_key[1]})"
                            )
                        if cache_key not in self._model_cache:
                            self._model_cache[cache_key] = new_model
                    logger.info(f"✅ Model switched to {model_size} ({new_backend})")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"❌ Background model loading error: {e}")
                if logger.isEnabledFor(logging.DEBUG):
                    logger.exception("Stacktrace:")
            finally:
                with self._lock:
                    self._model_loading = False
                if self.device == "cuda" and self._torch is not None:
                    try:
                        self._torch.cuda.empty_cache()
                        logger.info("🧹 GPU cache emptied after model switch")
                    except Exception:
                        pass

        thread = threading.Thread(
            target=_load_in_background, daemon=True, name=f"ModelLoader-{model_size}"
        )
        thread.start()
        return True

    def is_model_loading(self) -> bool:
        return self._model_loading

    def _force_model_cleanup(self) -> None:
        with self._model_usage_lock:
            old_model = self.model
            old_size = self.model_size
            old_backend = self.whisper_backend
            self.model = None
            self.model_size = None
            self._model_loaded_flag = False
        if old_model is not None:
            self._unload_model(old_model)
            logger.debug(f"🧹 Active model {old_size} ({old_backend}) unloaded")
        if old_size is not None and old_backend is not None:
            with self._lock:
                cache_key = (old_size, old_backend)
                if cache_key in self._model_cache:
                    del self._model_cache[cache_key]
                    logger.debug(f"🧹 Model {old_size} removed from cache")
        gc.collect()
        if self.device == "cuda" and self._torch is not None:
            try:
                self._torch.cuda.empty_cache()
                logger.info("🧹 GPU memory cache cleared")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Failed to clear GPU cache: {e}")

    def _universal_transcribe(
        self, model: Any, audio_np: Any, **kwargs: Any
    ) -> Tuple[List[Any], Any]:
        if model is None:
            raise ValueError("Kein Modell geladen")
        backend = self.whisper_backend
        if logger.isEnabledFor(logging.DEBUG):
            log_debug("transcribe", f"Backend: {backend}, Parameter: {kwargs}")

        if backend == "faster_whisper":
            allowed = self._ALLOWED_FASTER
        else:
            allowed = self._ALLOWED_OPENAI

        filtered = {k: v for k, v in kwargs.items() if k in allowed}

        if DEBUG_LEVEL >= 3:
            logger.debug(f"Effektive Whisper-Parameter: {filtered}")

        for attempt in range(2):
            try:
                if backend == "faster_whisper":
                    return self._faster_whisper_transcribe(model, audio_np, **filtered)
                else:
                    return self._openai_whisper_transcribe(model, audio_np, **filtered)
            except RuntimeError as e:
                if "out of memory" in str(e).lower():
                    logger.error(
                        f"🚨 CUDA out of memory (Versuch {attempt+1}/2) – leere GPU-Cache"
                    )
                    if self._torch is not None and self.device == "cuda":
                        self._torch.cuda.empty_cache()
                    time.sleep(0.5)
                    if attempt == 0:
                        continue
                    else:
                        logger.critical(
                            "❌ CUDA out of memory auch nach Wiederholung – Abbruch"
                        )
                        raise
                else:
                    raise
            except Exception:
                raise
        return [], _EmptyInfo()

    def _faster_whisper_transcribe(
        self, model: Any, audio_np: Any, **kwargs: Any
    ) -> Tuple[List[Any], Any]:
        try:
            vad_params = kwargs.pop("vad_parameters", None)
            if kwargs.get("vad_filter", False) and vad_params is None:
                vad_params = {
                    "threshold": self.settings.vad_threshold,
                    "min_speech_duration_ms": self.settings.vad_min_speech_duration_ms,
                    "min_silence_duration_ms": self.settings.vad_min_silence_duration_ms,
                }
                language = kwargs.get("language")
                if language and language in self.config.LANGUAGE_VAD:
                    lang_vad = self.config.LANGUAGE_VAD[language]
                    vad_params.update({
                        "threshold": lang_vad["threshold"],
                        "min_speech_duration_ms": lang_vad["min_speech_ms"],
                        "min_silence_duration_ms": lang_vad["min_silence_ms"],
                    })
            if vad_params is not None:
                kwargs["vad_parameters"] = vad_params

            segments, info = model.transcribe(audio_np, **kwargs)
            segments_list = list(segments)
            if logger.isEnabledFor(logging.DEBUG) and segments_list:
                log_debug(
                    "transcribe",
                    f"{len(segments_list)} Segmente erhalten, erste: {segments_list[0].text[:50]}..."
                )

            if DEBUG_LEVEL >= 3:
                logger.debug(f"VAD-Info: Sprache={info.language}, Dauer={info.duration}")

            return segments_list, info
        except (TypeError, ValueError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(
                f"⚠️ faster-whisper Parameterfehler: {e} – verwende minimale Parameter"
            )
            minimal_kwargs = {
                k: v
                for k, v in kwargs.items()
                if k in ["language", "task", "temperature", "beam_size", "best_of"]
            }
            try:
                segments, info = model.transcribe(audio_np, **minimal_kwargs)
                segments_list = list(segments)
                return segments_list, info
            except Exception as e2:
                if isinstance(e2, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"❌ faster-whisper auch mit minimalen Parametern fehlgeschlagen: {e2}"
                )
                if logger.isEnabledFor(logging.DEBUG):
                    logger.exception("Stacktrace:")
                return [], _EmptyInfo()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ faster-whisper Fehler: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return [], _EmptyInfo()

    def _openai_whisper_transcribe(
        self, model: Any, audio_np: Any, **kwargs: Any
    ) -> Tuple[List[Any], Any]:
        filtered_kwargs = {k: v for k, v in kwargs.items() if k in self._ALLOWED_OPENAI}
        filtered_kwargs.setdefault("language", None)
        filtered_kwargs.setdefault("task", "transcribe")
        filtered_kwargs.setdefault("temperature", 0.0)
        try:
            result = model.transcribe(audio_np, **filtered_kwargs)
            segments = result.get("segments", [])
            converted = []
            for seg in segments:
                if seg.get("text", "").strip():
                    converted.append(_UniversalSegment(seg))
            info = _UniversalInfo(result)
            if logger.isEnabledFor(logging.DEBUG) and converted:
                log_debug(
                    "transcribe",
                    f"{len(converted)} Segmente erhalten, erste: {converted[0].text[:50]}..."
                )
            return converted, info
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ openai-whisper Fehler: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            try:
                minimal_result = model.transcribe(
                    audio_np, language=None, task="transcribe", temperature=0.1
                )
                emergency = []
                for seg in minimal_result.get("segments", []):
                    emergency.append(_EmergencySegment(seg))
                logger.debug(f"[TRANSCRIBE] Fallback: {len(emergency)} Segmente")
                return emergency, _UniversalInfo(minimal_result)
            except Exception as fallback_error:
                if isinstance(fallback_error, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"💥 Auch Fallback fehlgeschlagen: {fallback_error}")
                return [], _EmptyInfo()

    def validate_audio_data(self, audio_data: bytes) -> Tuple[bool, str]:
        if not isinstance(audio_data, bytes):
            return False, "Audio data must be bytes"
        if len(audio_data) == 0:
            return False, "Audio data is empty"
        if len(audio_data) < 1600:
            return False, f"Audio data too short: {len(audio_data)} bytes"
        if self._np is not None:
            try:
                audio_np = self._np.frombuffer(audio_data, dtype=self._np.int16)
                if self._np.all(audio_np == 0):
                    return False, "Audio data is completely silent"
                if self._np.var(audio_np) < 100:
                    return False, "Audio variance too low (likely silent)"
            except Exception:
                pass
        return True, "Valid"

    def is_valid_segment(self, text: str, confidence: float) -> bool:
        if not text or len(text.strip()) < 2:
            return False
        clean = text.strip()
        if clean.isspace():
            return False
        if len(clean) > 500:
            return False
        if not any(c.isalpha() for c in clean):
            return False
        if confidence < self.settings.min_confidence:
            return False
        return True

    def safe_transcribe(
        self, audio_data: bytes, max_retries: int = 2
    ) -> Optional[TranscriptionResult]:
        with self._model_usage_lock:
            if self._reloading:
                logger.debug("Reloading in progress – skip transcription")
                return None

        is_valid, msg = self.validate_audio_data(audio_data)
        if not is_valid:
            if logger.isEnabledFor(logging.DEBUG):
                log_debug("transcribe", f"Audio validation failed: {msg}")
            return None
        for attempt in range(max_retries + 1):
            try:
                if attempt == 0 or self._last_confidence_threshold < 0.5:
                    processed = self._audio_enhancer.enhance_audio(
                        audio_data, self._last_confidence_threshold, 0
                    )
                else:
                    processed = audio_data
                result = self.transcribe_audio(processed)
                if result and result.text and result.text.strip():
                    self._last_confidence_threshold = getattr(result, "confidence", 0.5)
                    return result
                logger.debug(f"Transcription returned empty text (attempt {attempt+1})")
            except RuntimeError as e:
                if "out of memory" in str(e).lower() and self.device == "cuda":
                    logger.error(
                        f"🚨 CUDA out of memory (attempt {attempt+1}) – versuche Gegenmaßnahmen"
                    )
                    self._handle_cuda_oom()
                    continue
                else:
                    logger.warning(
                        f"RuntimeError in Transkription (Versuch {attempt+1}): {e}"
                    )
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.exception("Stacktrace:")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Transkriptionsfehler (Versuch {attempt+1}): {e}")
                if logger.isEnabledFor(logging.DEBUG):
                    logger.exception("Stacktrace:")
            if attempt < max_retries:
                wait_time = 0.5 * (attempt + 1)
                logger.debug(f"Warte {wait_time:.1f}s vor nächstem Versuch")
                time.sleep(wait_time)
        logger.error(
            f"❌ safe_transcribe fehlgeschlagen nach {max_retries+1} Versuchen"
        )
        return None

    def _handle_cuda_oom(self) -> None:
        logger.info("🧹 CUDA OOM: Bereinige GPU-Speicher...")
        if self.device == "cpu":
            logger.warning("⚠️ Bereits auf CPU, kann OOM nicht beheben.")
            return

        with self._model_usage_lock:
            self._reloading = True

        try:
            if self._torch and self.device == "cuda":
                self._torch.cuda.empty_cache()
                time.sleep(0.2)
                free_gb = self._get_free_gpu_memory()
                if free_gb is not None and free_gb > 1.0:
                    logger.info(f"✅ Nach Cache-Leerung: {free_gb:.1f} GB frei – versuche weiter mit gleichem Modell")
                    return

                with self._model_usage_lock:
                    if self.model is not None:
                        self._unload_model(self.model)
                        self.model = None
                        self.model_size = None
                self._torch.cuda.empty_cache()
                time.sleep(0.1)

            current = self.model_size or "medium"
            if current not in self.MODEL_SIZE_ORDER:
                for size in self.MODEL_SIZE_ORDER:
                    if size.startswith(current.split(".")[0]):
                        current = size
                        break
                else:
                    current = "medium"
            current_idx = self.MODEL_SIZE_ORDER.index(current)

            smaller_models = self.MODEL_SIZE_ORDER[:current_idx]

            if smaller_models:
                for smaller in reversed(smaller_models):
                    logger.info(f"🔄 Versuche kleineres Modell: {smaller}")
                    result = self.load_model(smaller, set_active=True)
                    if result is not None:
                        break
                else:
                    logger.warning(
                        "⚠️ Auch kleinere Modelle fehlgeschlagen, wechsle auf CPU"
                    )
                    with self._model_usage_lock:
                        self.device = "cpu"
                        self.compute_type = "int8"
                    self.load_model(current, set_active=True)
            else:
                logger.warning("⚠️ Schalte wegen OOM auf CPU um")
                with self._model_usage_lock:
                    self.device = "cpu"
                    self.compute_type = "int8"
                self.load_model(current, set_active=True)

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Fehler in _handle_cuda_oom: {e}")
        finally:
            with self._model_usage_lock:
                self._reloading = False

    def _calculate_enhanced_confidence(self, segment: Any, text: str) -> float:
        base = max(getattr(segment, "confidence", 0.0), 0.1)
        words = text.split()
        word_count = len(words)
        text_len = len(text.strip())
        unique_ratio = len(set(words)) / max(word_count, 1)
        boosts = (
            min(0.2, text_len / 300.0)
            + min(0.15, word_count * 0.03)
            + (0.08 if any(c in text for c in ".!?,;:") else 0.0)
            + (0.1 if any(c.isalpha() for c in text) else 0.0)
            + min(0.1, unique_ratio * 0.1)
        )
        return min(0.95, base + boosts)

    def transcribe_audio(
        self, audio_data: bytes, include_timestamps: bool = False
    ) -> Any:
        with self._model_usage_lock:
            model = self.model
            if not model:
                return None if not include_timestamps else []
        try:
            processed = self._audio_enhancer.enhance_audio(
                audio_data, self._last_confidence_threshold, 0
            )
            if self._np is None:
                self._np = FastLazyLoader.load("numpy")
            audio_np = (
                self._np.frombuffer(processed, dtype=self._np.int16).astype(
                    self._np.float32
                )
                / 32768.0
            )
            beam_size = self.settings.beam_size
            language = self.forced_language if self.forced_language else None
            vad_language = self.forced_language or self._last_detected_language
            if vad_language and vad_language in self.config.LANGUAGE_VAD:
                lang_vad = self.config.LANGUAGE_VAD[vad_language]
                vad_params = {
                    "threshold": lang_vad["threshold"],
                    "min_speech_duration_ms": lang_vad["min_speech_ms"],
                    "min_silence_duration_ms": lang_vad["min_silence_ms"],
                }
            else:
                vad_params = {
                    "threshold": self.settings.vad_threshold,
                    "min_speech_duration_ms": self.settings.vad_min_speech_duration_ms,
                    "min_silence_duration_ms": self.settings.vad_min_silence_duration_ms,
                }
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler bei Audio-Vorbereitung: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return [] if include_timestamps else None

        with self._model_usage_lock:
            if self.model is not model:
                logger.warning(
                    "⚠️ Modell wurde gewechselt – Transkription abgebrochen."
                )
                return [] if include_timestamps else None

        return self._transcribe_worker(
            model, audio_np, language, beam_size, vad_params, include_timestamps, hotwords=self.settings.hotwords
        )

    def _transcribe_worker(
        self,
        model: Any,
        audio_np: Any,
        language: Optional[str],
        beam_size: int,
        vad_params: Dict[str, Any],
        include_timestamps: bool,
        hotwords: str = "",
    ) -> Any:
        try:
            if logger.isEnabledFor(logging.DEBUG):
                log_debug(
                    "vad", f"Parameter: {vad_params}, aktiv: {self.settings.vad_filter}"
                )
            kwargs = {
                "language": language,
                "task": "transcribe",
                "temperature": self.settings.temperature,
                "best_of": self.settings.best_of,
                "beam_size": beam_size,
                "patience": self.settings.patience,
                "no_speech_threshold": self.settings.no_speech_threshold,
                "log_prob_threshold": self.settings.log_prob_threshold,
                "compression_ratio_threshold": self.settings.compression_ratio_threshold,
                "condition_on_previous_text": self.settings.condition_on_previous_text,
                "suppress_tokens": self._parse_suppress_tokens(self.settings.suppress_tokens),
                "without_timestamps": False,
                "word_timestamps": include_timestamps,
                "vad_filter": self.settings.vad_filter,
                "vad_parameters": vad_params,
            }
            if hotwords:
                kwargs["hotwords"] = hotwords

            segments, info = self._universal_transcribe(
                model,
                audio_np,
                **kwargs
            )
            if not segments:
                if self._vad_fallback_enabled:
                    if logger.isEnabledFor(logging.DEBUG):
                        log_debug("vad", "Keine Segmente mit VAD – Versuch ohne VAD...")
                    if DEBUG_LEVEL >= 3:
                        logger.debug("Fallback ohne VAD wird ausgeführt")
                    segments, info = self._universal_transcribe(
                        model,
                        audio_np,
                        language=language,
                        task="transcribe",
                        temperature=0.0,
                        best_of=5,
                        beam_size=beam_size,
                        vad_filter=False,
                        without_timestamps=False,
                    )
                else:
                    if logger.isEnabledFor(logging.DEBUG):
                        log_debug("vad", "Keine Segmente mit VAD – Fallback deaktiviert, ignoriere Chunk")
                    return [] if include_timestamps else None

            if hasattr(info, "language") and info.language != "unknown":
                if getattr(info, "language_probability", 1.0) < 0.4:
                    logger.debug("Low language confidence, using fallback")
                    # Kein automatischer Fallback, Benutzer kann manuell Sprache setzen
                else:
                    self._last_detected_language = info.language

            valid_segments = []
            total_confidence = 0.0
            for seg in segments:
                text = seg.text.strip()
                conf = self._calculate_enhanced_confidence(seg, text)
                if self.is_valid_segment(text, conf):
                    valid_segments.append(seg)
                    total_confidence += conf

            if not valid_segments:
                logger.debug("🔄 Keine validen Segmente – minimaler Fallback")
                minimal = self._transcribe_minimal(model, audio_np, language)
                if minimal:
                    if include_timestamps:
                        duration = audio_np.shape[0] / self.config.SAMPLE_RATE
                        return [
                            TranscriptionResult(
                                text=minimal.text,
                                confidence=minimal.confidence,
                                language=minimal.language,
                                start=0.0,
                                end=duration,
                            )
                        ]
                    return minimal
                return [] if include_timestamps else None

            if include_timestamps:
                return [
                    TranscriptionResult(
                        text=seg.text.strip(),
                        confidence=self._calculate_enhanced_confidence(
                            seg, seg.text.strip()
                        ),
                        language=getattr(info, "language", "unknown"),
                        start=getattr(seg, "start", 0.0),
                        end=getattr(seg, "end", 0.0),
                    )
                    for seg in valid_segments
                ]
            else:
                full_text = " ".join(seg.text.strip() for seg in valid_segments)
                avg_conf = total_confidence / len(valid_segments)
                return TranscriptionResult(
                    text=full_text,
                    confidence=avg_conf,
                    language=getattr(info, "language", "unknown"),
                )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ _transcribe_worker Fehler: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            return [] if include_timestamps else None

    def _transcribe_minimal(
        self, model: Any, audio_np: Any, language: Optional[str]
    ) -> Optional[TranscriptionResult]:
        try:
            segments, info = self._universal_transcribe(
                model,
                audio_np,
                language=language,
                task="transcribe",
                temperature=0.0,
                best_of=1,
                beam_size=1,
                no_speech_threshold=0.9,
                log_prob_threshold=-2.0,
                compression_ratio_threshold=3.5,
                condition_on_previous_text=False,
                without_timestamps=False,
                vad_filter=False,
            )
            if segments:
                seg = segments[0]
                text = seg.text.strip()
                if text:
                    conf = self._calculate_enhanced_confidence(seg, text)
                    return TranscriptionResult(
                        text=text,
                        confidence=conf,
                        language=getattr(info, "language", "unknown"),
                    )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.debug(f"Minimal transcription failed: {e}")
        return None

    def emergency_fallback_transcription(
        self, audio_data: Union[bytes, Any]
    ) -> Optional[TranscriptionResult]:
        with self._model_usage_lock:
            model = self.model
            if not model:
                return None
            try:
                if isinstance(audio_data, self._np.ndarray):
                    audio_np = audio_data.astype(self._np.float32)
                    if audio_np.dtype == self._np.int16:
                        audio_np = audio_np / 32768.0
                else:
                    if self._np is None:
                        self._np = FastLazyLoader.load("numpy")
                    audio_np = (
                        self._np.frombuffer(audio_data, dtype=self._np.int16).astype(
                            self._np.float32
                        )
                        / 32768.0
                    )
                return self._transcribe_minimal(model, audio_np, None)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"❌ Emergency fallback exception: {e}")
                return None

    def clear_cache(self) -> None:
        with self._lock:
            self._cache.clear()
            self._last_transcription_text = ""
        gc.collect()
        if self.device == "cuda" and self._torch is not None:
            try:
                self._torch.cuda.empty_cache()
            except Exception:
                pass

    def get_current_model(self) -> str:
        with self._model_usage_lock:
            return self.model_size if self.model_size else "None"

    def test_model_functionality(self) -> bool:
        with self._model_usage_lock:
            if not self.model:
                return False
            try:
                if self._np is None:
                    self._np = FastLazyLoader.load("numpy")
                test_audio = self._np.random.randn(16000).astype(self._np.float32) * 0.1
                segments, info = self._universal_transcribe(
                    self.model,
                    test_audio,
                    language=None,
                    task="transcribe",
                    temperature=0.0,
                    best_of=1,
                    beam_size=1,
                    without_timestamps=False,
                )
                list(segments)
                return True
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"❌ Model-Test fehlgeschlagen: {e}")
                return False

    def dispose(self) -> None:
        logger.info("🧹 Transcription Engine Dispose...")
        self._disposing = True
        with self._lock:
            self._cache.clear()
            self._last_transcription_text = ""
            for (size, backend), model in list(self._model_cache.items()):
                self._unload_model(model)
            self._model_cache.clear()
        self._force_model_cleanup()
        gc.collect()
        logger.info("✅ Transcription Engine disposed")


# -----------------------------------------------------------------------------
# DummyTranscriptionEngine (für Demo-Modus)
# -----------------------------------------------------------------------------
class DummyTranscriptionEngine:
    def __init__(
        self,
        settings: Optional["AdvancedSettings"] = None,
        cache_manager: Optional[CacheManager] = None,
    ):
        self.settings = settings or AdvancedSettings()
        self.model = None
        self.model_size = "dummy"
        self.whisper_backend = None
        self.demo_mode = True

    def load_model(
        self, model_size: str, set_active: bool = False
    ) -> Optional[Tuple[Any, str]]:
        logger.info("Dummy-Modus: Laden eines Modells nicht erforderlich.")
        return (None, "dummy")

    def transcribe_audio(
        self, audio_data: bytes, include_timestamps: bool = False
    ) -> Any:
        if include_timestamps:
            dummy = TranscriptionResult(
                text="[Whisper nicht verfügbar]",
                confidence=0.5,
                language="de",
                start=0.0,
                end=5.0,
            )
            return [dummy]
        else:
            return TranscriptionResult(
                text="[Whisper nicht verfügbar]", confidence=0.5, language="de"
            )

    def safe_transcribe(
        self, audio_data: bytes, max_retries: int = 2
    ) -> Optional[TranscriptionResult]:
        return self.transcribe_audio(audio_data)

    def get_current_model(self) -> str:
        return "dummy"

    def is_model_loading(self) -> bool:
        return False

    def reload_model(self, model_size: str) -> bool:
        return False

    def test_model_functionality(self) -> bool:
        return True

    def dispose(self) -> None:
        pass


# =============================================================================
# 8. MANAGER-KLASSEN
# =============================================================================


# -----------------------------------------------------------------------------
# PluginManager
# -----------------------------------------------------------------------------
class PluginManager:
    def __init__(self):
        self._plugins: List[Plugin] = []
        self._plugin_map: Dict[str, Plugin] = {}
        self._lock = threading.RLock()
        self._global_enabled = True
        self._event_handlers = {
            "transcription": [],
            "translation": [],
            "start": [],
            "stop": [],
        }

    def register_plugin(self, plugin: Plugin) -> bool:
        with self._lock:
            if plugin.name in self._plugin_map:
                logger.warning(
                    f"Plugin '{plugin.name}' bereits registriert – überspringe."
                )
                return False
            self._plugins.append(plugin)
            self._plugin_map[plugin.name] = plugin
            try:
                plugin.on_load(self)
                logger.info(f"✅ Plugin geladen: {plugin.name} v{plugin.version}")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"❌ Fehler beim Laden von Plugin {plugin.name}: {e}")
                plugin.enabled = False
            return True

    def unregister_plugin(self, plugin_name: str) -> bool:
        with self._lock:
            if plugin_name not in self._plugin_map:
                return False
            plugin = self._plugin_map[plugin_name]
            try:
                plugin.on_unload()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Entladen von {plugin_name}: {e}")
            self._plugins.remove(plugin)
            del self._plugin_map[plugin_name]
            logger.info(f"Plugin entfernt: {plugin_name}")
            return True

    def get_plugin(self, name: str) -> Optional[Plugin]:
        return self._plugin_map.get(name)

    def list_plugins(self) -> List[Plugin]:
        with self._lock:
            return self._plugins.copy()

    def set_plugin_enabled(self, name: str, enabled: bool) -> bool:
        plugin = self.get_plugin(name)
        if plugin:
            plugin.enabled = enabled
            logger.info(f"Plugin {name} {'aktiviert' if enabled else 'deaktiviert'}")
            return True
        return False

    def set_global_enabled(self, enabled: bool) -> None:
        with self._lock:
            self._global_enabled = enabled

    def process_transcription(self, result: TranscriptionResult) -> TranscriptionResult:
        if not self._global_enabled:
            return result
        with self._lock:
            plugins = self._plugins.copy()
        for plugin in plugins:
            if not plugin.enabled:
                continue
            try:
                result = plugin.on_transcription(result)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Plugin {plugin.name} Fehler in on_transcription: {e}")
        return result

    def process_translation(self, result: TranslationResult) -> TranslationResult:
        if not self._global_enabled:
            return result
        with self._lock:
            plugins = self._plugins.copy()
        for plugin in plugins:
            if not plugin.enabled:
                continue
            try:
                result = plugin.on_translation(result)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Plugin {plugin.name} Fehler in on_translation: {e}")
        return result

    def on_start(self) -> None:
        if not self._global_enabled:
            return
        with self._lock:
            plugins = self._plugins.copy()
        for plugin in plugins:
            if not plugin.enabled:
                continue
            try:
                plugin.on_start()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Plugin {plugin.name} Fehler in on_start: {e}")

    def on_stop(self) -> None:
        if not self._global_enabled:
            return
        with self._lock:
            plugins = self._plugins.copy()
        for plugin in plugins:
            if not plugin.enabled:
                continue
            try:
                plugin.on_stop()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Plugin {plugin.name} Fehler in on_stop: {e}")

    def load_config(self, config_data: Dict[str, Dict[str, Any]]) -> None:
        for name, data in config_data.items():
            plugin = self.get_plugin(name)
            if plugin:
                enabled = data.get("enabled", True)
                plugin.enabled = enabled
                plugin_config = data.get("config", {})
                plugin.on_config_change(plugin_config)

    def save_config(self) -> Dict[str, Dict[str, Any]]:
        config = {}
        with self._lock:
            for plugin in self._plugins:
                config[plugin.name] = {
                    "enabled": plugin.enabled,
                    "config": plugin.config.copy(),
                }
        return config

    def dispose(self) -> None:
        with self._lock:
            plugins = self._plugins.copy()
        for plugin in plugins:
            try:
                plugin.on_unload()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Fehler beim Entladen von {plugin.name}: {e}")
        with self._lock:
            self._plugins.clear()
            self._plugin_map.clear()


# -----------------------------------------------------------------------------
# StreamManager
# -----------------------------------------------------------------------------
class YtDlpHelper:
    @staticmethod
    def run_command(cmd: List[str], timeout: int = 15, method_name: str = "unknown") -> Optional[str]:
        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"  Ausführen: {' '.join(cmd)}")

        start = time.perf_counter()
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
                shell=False,
                encoding="utf-8",
                errors="replace",
            )
            duration = (time.perf_counter() - start) * 1000
            if result.returncode == 0 and result.stdout.strip():
                if logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"  ✅ {method_name} erfolgreich in {duration:.2f}ms")
                return result.stdout.strip()
            else:
                if logger.isEnabledFor(logging.DEBUG) and result.stderr:
                    stderr_preview = result.stderr[:200].replace("\n", " ")
                    logger.debug(f"  ⚠️ {method_name} fehlgeschlagen (Code {result.returncode}): {stderr_preview}")
        except subprocess.TimeoutExpired:
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"  ⏰ {method_name} Timeout nach {timeout}s")
        except (OSError, ValueError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"  ⚠️ {method_name} Fehler: {e}")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Unerwarteter Fehler in {method_name}: {e}", exc_info=True)
        return None

    @staticmethod
    def get_json(url: str, timeout: int = 20, use_cookies: bool = False, browser: Optional[str] = None) -> Optional[Dict[str, Any]]:
        cmd = [
            "yt-dlp",
            "--dump-json",
            "--no-warnings",
            "--no-check-certificate",
            "--socket-timeout", str(timeout),
        ]
        if use_cookies and browser:
            cmd.extend(["--cookies-from-browser", browser])
        cmd.extend(["--", url])

        stdout = YtDlpHelper.run_command(cmd, timeout=timeout, method_name="get_json")
        if stdout:
            try:
                return json.loads(stdout)
            except json.JSONDecodeError:
                pass
        return None

    @staticmethod
    def get_audio_url(url: str, format_str: str = "bestaudio/best", timeout: int = 15, use_cookies: bool = False, browser: Optional[str] = None) -> Optional[str]:
        cmd = [
            "yt-dlp",
            "-g",
            "-f", format_str,
            "--no-warnings",
            "--no-check-certificate",
            "--socket-timeout", str(timeout),
        ]
        if use_cookies and browser:
            cmd.extend(["--cookies-from-browser", browser])
        cmd.extend(["--", url])

        stdout = YtDlpHelper.run_command(cmd, timeout=timeout, method_name=f"get_audio_url_{format_str}")
        if stdout:
            for line in stdout.splitlines():
                line = line.strip()
                if line.startswith(("http://", "https://")):
                    return line
        return None


class StreamManager:
    def __init__(self, enable_debug: bool = False, use_browser_cookies: bool = True) -> None:
        self._platform_cache = TTLCache(maxsize=50, ttl=3600)
        self._audio_url_cache = TTLCache(maxsize=50, ttl=1800)
        self._audio_url_fail_cache = TTLCache(maxsize=50, ttl=300)
        self._live_status_cache = TTLCache(maxsize=30, ttl=300)
        self._stream_info_cache = TTLCache(maxsize=30, ttl=600)

        self._debug = enable_debug
        self.use_browser_cookies = use_browser_cookies
        self._last_error: Optional[str] = None
        self._last_method: Optional[str] = None

        self._stats = {
            "extraction_attempts": 0,
            "successful_extractions": 0,
            "cache_hits": 0,
            "errors": 0,
            "start_time": time.time(),
        }
        self._stats_lock = threading.RLock()

        self._format_priorities = {
            "youtube": ["bestaudio[ext=m4a]/bestaudio/best", "bestaudio/best", "ba"],
            "youtube_live": ["bestaudio/best", "ba"],
            "twitch": ["bestaudio/best", "audio_only"],
            "tiktok": ["bestaudio/best"],
            "facebook": ["bestaudio/best"],
            "hls": ["bestaudio/best"],
            "dash": ["bestaudio/best"],
            "generic": ["bestaudio/best", "ba"],
            "kick": ["bestaudio/best", "ba"],
            "rumble": ["bestaudio/best", "ba"],
            "dailymotion": ["bestaudio/best", "ba"],
            "vimeo": ["bestaudio/best", "ba"],
            "twitter": ["bestaudio/best", "ba"],
        }

        self._user_agents = {
            "desktop": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "mobile": "Mozilla/5.0 (Linux; Android 10; SM-G975F) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36",
        }

        self._browsers = [
            ("firefox", "Firefox"),
            ("chrome", "Chrome"),
            ("brave", "Brave"),
            ("edge", "Edge"),
            ("chromium", "Chromium"),
            ("opera", "Opera"),
            ("vivaldi", "Vivaldi"),
        ]

        self._dvb_processes: List[subprocess.Popen] = []
        self._dvb_lock = threading.RLock()

    def detect_platform(self, url: str) -> Tuple[str, str]:
        if not url:
            return ("unknown", "Invalid URL")

        url = PlatformUtils.sanitize_url(url)
        cached = self._platform_cache.get(url)
        if cached is not None:
            with self._stats_lock:
                self._stats["cache_hits"] += 1
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"🔍 detect_platform: Cache-Treffer für {url[:50]}...")
            return cached

        result, reason_str = self._detect_platform_impl(url)
        self._platform_cache.put(url, result)

        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"🔍 detect_platform: {url[:50]}... -> {result}, reason: {reason_str}")
        return result

    def extract_audio_url(self, url: str, force_refresh: bool = False) -> Optional[str]:
        with self._stats_lock:
            self._stats["extraction_attempts"] += 1

        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"\n🎵 [EXTRACT_AUDIO_URL] Start für: {url[:80]}...")

        self._last_error = None
        self._last_method = None

        if not shutil.which("yt-dlp"):
            self._last_error = "yt-dlp not found in PATH"
            logger.error(self._last_error)
            with self._stats_lock:
                self._stats["errors"] += 1
            return None

        url = PlatformUtils.sanitize_url(url)
        if not url:
            self._last_error = "Empty URL"
            with self._stats_lock:
                self._stats["errors"] += 1
            return None

        if url.startswith(("dvb://", "dvb-s://")):
            logger.info("📡 DVB-Stream erkannt, starte VLC-Server...")
            result = self._start_dvb_stream(url)
            if result:
                with self._stats_lock:
                    self._stats["successful_extractions"] += 1
                self._last_method = "dvb_vlc"
            else:
                with self._stats_lock:
                    self._stats["errors"] += 1
                self._last_error = "DVB stream could not be started"
            return result

        if not force_refresh:
            cached = self._audio_url_cache.get(url)
            if cached:
                with self._stats_lock:
                    self._stats["cache_hits"] += 1
                if self._debug or logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"📦 Audio-URL Cache-Treffer für {url[:50]}...")
                return cached
            if self._audio_url_fail_cache.get(url):
                if self._debug or logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"📦 Audio-URL Fail-Cache-Treffer für {url[:50]}...")
                return None

        platform_id, platform_name = self.detect_platform(url)
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"🔍 Plattform erkannt: {platform_id} ({platform_name})")

        result = None
        extraction_method = "unknown"

        if url.startswith("file://"):
            ok, real_path = PlatformUtils.validate_file_path(url)
            if ok:
                result = url
                extraction_method = "local_file"
            else:
                self._last_error = real_path
        elif self._is_direct_media_url(url):
            result = url
            extraction_method = "direct_link"

        if not result and platform_id in ("youtube", "youtube_live"):
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug("🎯 YouTube erkannt, verwende optimierte Extraktion...")
            result = self._extract_youtube_audio_optimized(url, platform_id)
            extraction_method = "youtube_optimized"

        if not result:
            result, method = self._extract_generic_audio(url, platform_id)
            if method:
                extraction_method = method

        if not result:
            result = self._json_extraction_fallback(url)
            if result:
                extraction_method = "json_fallback"

        if result:
            self._audio_url_cache.put(url, result)
            with self._stats_lock:
                self._stats["successful_extractions"] += 1
        else:
            self._audio_url_fail_cache.put(url, True)
            if not self._last_error:
                self._last_error = "No audio URL could be extracted"
            with self._stats_lock:
                self._stats["errors"] += 1

        self._last_method = extraction_method
        if (self._debug or logger.isEnabledFor(logging.DEBUG)) and result:
            logger.debug(f"🎵 EXTRACT_AUDIO_URL ENDE - Ergebnis: {result[:80]}...")
        return result

    def dispose(self) -> None:
        self.clear_caches()
        with self._dvb_lock:
            for proc in self._dvb_processes:
                try:
                    proc.terminate()
                    proc.wait(timeout=1)
                except subprocess.TimeoutExpired:
                    try:
                        proc.kill()
                        proc.wait(timeout=0.5)
                    except Exception:
                        pass
                except Exception:
                    pass
            self._dvb_processes.clear()
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug("🔌 StreamManager disposed")

    def clear_caches(self) -> None:
        self._platform_cache.clear()
        self._audio_url_cache.clear()
        self._audio_url_fail_cache.clear()
        self._live_status_cache.clear()
        self._stream_info_cache.clear()
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug("🗑️ Alle Caches geleert")

    def _detect_platform_impl(self, url: str) -> Tuple[Tuple[str, str], str]:
        url_lower = url.lower()
        reasons = []

        if url_lower.startswith(("dvb://", "dvb-s://")):
            result = ("dvb", "DVB-S Stream")
            reasons.append("dvb protocol")
            return result, ", ".join(reasons)

        if url_lower.startswith("file://"):
            ok, _ = PlatformUtils.validate_file_path(url)
            if not ok:
                result = ("invalid", "Invalid file path")
                reasons.append("invalid file")
            else:
                result = ("local", "Local File")
                reasons.append("file://")
            return result, ", ".join(reasons)

        audio_ext = (".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac", ".opus", ".webm")
        video_ext = (".mp4", ".avi", ".mkv", ".mov", ".webm", ".m4v", ".wmv", ".flv")

        if any(url_lower.endswith(ext) for ext in audio_ext):
            result = ("direct_audio", "Direct Audio")
            reasons.append("audio extension")
            return result, ", ".join(reasons)
        if any(url_lower.endswith(ext) for ext in video_ext):
            result = ("direct_video", "Direct Video")
            reasons.append("video extension")
            return result, ", ".join(reasons)

        domain_map = {
            "youtube.com": ("youtube", "YouTube Video"),
            "youtu.be": ("youtube", "YouTube Video"),
            "twitch.tv": ("twitch", "Twitch"),
            "kick.com": ("kick", "Kick"),
            "rumble.com": ("rumble", "Rumble"),
            "dailymotion.com": ("dailymotion", "Dailymotion"),
            "vimeo.com": ("vimeo", "Vimeo"),
            "twitter.com": ("twitter", "Twitter/X"),
            "x.com": ("twitter", "Twitter/X"),
            "tiktok.com": ("tiktok", "TikTok"),
            "facebook.com": ("facebook", "Facebook"),
            "fb.watch": ("facebook", "Facebook"),
        }
        for domain, (plat, name) in domain_map.items():
            if domain in url_lower:
                reasons.append(domain)
                if plat == "youtube" and self._check_youtube_live_status(url):
                    result = ("youtube_live", "YouTube Live")
                    reasons.append("live pattern")
                else:
                    result = (plat, name)
                return result, ", ".join(reasons)

        if ".m3u8" in url_lower:
            result = ("hls", "HLS Stream")
            reasons.append(".m3u8")
            return result, ", ".join(reasons)
        if ".mpd" in url_lower:
            result = ("dash", "DASH Stream")
            reasons.append(".mpd")
            return result, ", ".join(reasons)

        if url_lower.startswith(("http://", "https://")):
            result = ("generic", "Website/Stream")
            reasons.append("http(s) fallback")
            return result, ", ".join(reasons)

        result = ("unknown", "Unknown Source")
        reasons.append("no pattern matched")
        return result, ", ".join(reasons)

    def _is_direct_media_url(self, url: str) -> bool:
        url_lower = url.lower()
        audio_ext = (".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac", ".opus", ".webm")
        video_ext = (".mp4", ".avi", ".mkv", ".mov", ".webm", ".m4v", ".wmv", ".flv")
        return url_lower.endswith(audio_ext + video_ext)

    def _extract_generic_audio(self, url: str, platform_id: str) -> Tuple[Optional[str], str]:
        format_list = self._format_priorities.get(platform_id, self._format_priorities["generic"])

        for i, format_str in enumerate(format_list[:2]):
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"  🔄 Versuche Format {i+1}: {format_str}")
            audio_url = YtDlpHelper.get_audio_url(url, format_str, timeout=15)
            if audio_url:
                if self._debug or logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"  ✅ Erfolg mit Format {format_str}")
                return audio_url, f"generic_{format_str}"
        return None, ""

    def _extract_youtube_video_id(self, url: str) -> Optional[str]:
        patterns = [
            r"(?:youtube\.com/watch\?v=|youtu\.be/)([a-zA-Z0-9_-]{11})",
            r"youtube\.com/embed/([a-zA-Z0-9_-]{11})",
            r"youtube\.com/v/([a-zA-Z0-9_-]{11})",
        ]
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None

    def _check_youtube_live_status(self, url: str) -> bool:
        cached = self._live_status_cache.get(url)
        if cached is not None:
            return cached
        url_lower = url.lower()
        live_patterns = ["/live", "live=1", "/stream", "livestream"]
        is_live = any(pattern in url_lower for pattern in live_patterns)
        self._live_status_cache.put(url, is_live)
        return is_live

    def _extract_youtube_audio_optimized(self, url: str, platform_id: str) -> Optional[str]:
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"  🔍 Optimierte YouTube-Extraktion für: {url[:60]}...")

        video_id = self._extract_youtube_video_id(url)
        if not video_id or len(video_id) != 11:
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug("  ❌ Ungültige Video-ID")
            return None

        if self.use_browser_cookies:
            result = self._try_browser_cookies(url)
            if result:
                return result

        result = self._try_standard_methods(url)
        if result:
            return result

        result = self._try_json_fallback(url)
        if result:
            return result

        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug("    🔄 Alle Methoden fehlgeschlagen")
        return None

    def _try_browser_cookies(self, url: str) -> Optional[str]:
        for browser_cmd, browser_name in self._browsers:
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"    🧪 Teste mit {browser_name}-Cookies...")
            audio_url = YtDlpHelper.get_audio_url(
                url,
                format_str="bestaudio[ext=m4a]/bestaudio/best",
                timeout=20,
                use_cookies=True,
                browser=browser_cmd
            )
            if audio_url:
                if self._debug or logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"    ✅ Erfolg mit {browser_name}-Cookies")
                return audio_url
        return None

    def _try_standard_methods(self, url: str) -> Optional[str]:
        methods = [
            {
                "name": "Standard yt-dlp",
                "format": "bestaudio[ext=m4a]/bestaudio/best",
                "timeout": 20,
            },
            {
                "name": "Mobile User-Agent",
                "format": "bestaudio/best",
                "timeout": 20,
            },
            {
                "name": "Lowest Quality",
                "format": "worstaudio",
                "timeout": 20,
            },
        ]

        for method in methods:
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"    🧪 Teste: {method['name']}")
            audio_url = YtDlpHelper.get_audio_url(url, method["format"], timeout=method["timeout"])
            if audio_url:
                if self._debug or logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"    ✅ Erfolg mit {method['name']}")
                return audio_url
        return None

    def _try_json_fallback(self, url: str) -> Optional[str]:
        try:
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug("    🔄 Versuche JSON-Fallback...")
            data = YtDlpHelper.get_json(url, timeout=25)
            if not data:
                return None

            best_audio = None
            best_score = 0
            for fmt in data.get("formats", []):
                if fmt.get("acodec") != "none" and fmt.get("url"):
                    score = fmt.get("abr", 0) or fmt.get("tbr", 0) or 0
                    if fmt.get("vcodec") == "none":
                        score += 1000
                    ext = fmt.get("ext", "").lower()
                    if ext in ["m4a", "mp4"]:
                        score += 500
                    elif ext in ["webm", "opus"]:
                        score += 300
                    if score > best_score:
                        best_score = score
                        best_audio = fmt["url"]
            if best_audio and (self._debug or logger.isEnabledFor(logging.DEBUG)):
                logger.debug("    ✅ JSON-Fallback erfolgreich")
            return best_audio
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"    ⚠️ JSON-Fallback Fehler: {e}")
            return None

    def _start_dvb_stream(self, dvb_url: str) -> Optional[str]:
        if not shutil.which("vlc"):
            logger.error("❌ VLC nicht gefunden – für DVB wird VLC benötigt.")
            return None

        if not dvb_url.startswith("dvb-s://"):
            logger.error("❌ Bitte vollständige Parameter-URL angeben, z.B. dvb-s://frequency=...")
            return None

        params = dvb_url[8:]
        if not re.match(r"^[a-zA-Z0-9=&\-_.]+$", params):
            logger.error("❌ Ungültige Zeichen in DVB-URL-Parametern.")
            return None

        port = 8080
        mount = "/dvb"
        http_url = f"http://localhost:{port}{mount}"

        vlc_cmd = [
            "vlc",
            f"dvb-s://{params}",
            "--dvb-adapter=0",
            "--sout", f"#standard{{access=http,mux=ts,dst=0.0.0.0:{port}{mount}}}",
            "--intf", "dummy",
            "--no-audio",
            "--no-video",
            "--live-caching", "300",
        ]

        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(f"🚀 Starte VLC: {' '.join(vlc_cmd)}")

        try:
            process = subprocess.Popen(
                vlc_cmd,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                stdin=subprocess.DEVNULL,
            )
            time.sleep(2)

            if process.poll() is None:
                with self._dvb_lock:
                    self._dvb_processes.append(process)
                logger.info(f"✅ VLC-DVB-Server gestartet: {http_url}")
                return http_url
            else:
                logger.error("❌ VLC-Prozess für DVB ist sofort beendet.")
                return None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler beim Starten von VLC für DVB: {e}")
            return None

    def get_stats(self) -> Dict[str, Any]:
        with self._stats_lock:
            stats = self._stats.copy()
            stats["uptime_seconds"] = time.time() - stats["start_time"]
            if stats["extraction_attempts"] > 0:
                stats["success_rate"] = stats["successful_extractions"] / stats["extraction_attempts"] * 100
            else:
                stats["success_rate"] = 0.0
            stats["last_error"] = self._last_error
            stats["last_method"] = self._last_method
            return stats


# =============================================================================
# FFmpegManager
# =============================================================================
class FFmpegManager:
    class _ProcessInfo:
        __slots__ = (
            "process",
            "output_queue",
            "start_time",
            "url",
            "stopping",
            "bytes_read",
            "platform",
            "is_live",
            "chunks_processed",
            "last_activity",
            "headers_used",
            "process_id",
        )

        def __init__(
            self,
            process_id: str,
            process: subprocess.Popen,
            output_queue: Optional[queue.Queue],
            url: str,
            platform: str,
            is_live: bool,
            headers_used: bool,
        ):
            self.process_id = process_id
            self.process = process
            self.output_queue = output_queue
            self.start_time = time.time()
            self.url = url
            self.stopping = False
            self.bytes_read = 0
            self.platform = platform
            self.is_live = is_live
            self.chunks_processed = 0
            self.last_activity = time.time()
            self.headers_used = headers_used

    FIRST_DATA_TIMEOUT = 20.0
    HLS_INITIAL_WAIT = 1.5
    CLEANUP_INTERVAL = 60

    def __init__(
        self,
        config: Optional[Config] = None,
        stream_manager: Optional[StreamManager] = None,
        settings: Optional["AdvancedSettings"] = None,
    ) -> None:
        self._processes: Dict[str, FFmpegManager._ProcessInfo] = {}
        self._lock = threading.RLock()
        self._shutting_down = False
        self.config = config or Config()
        self.stream_manager = stream_manager or StreamManager()
        self.settings = settings
        self._pid_tracking: Dict[int, Dict[str, Any]] = {}
        self._live_detection_cache: Dict[str, Dict[str, Any]] = {}
        self._cleanup_running = True
        self._cleanup_thread = threading.Thread(
            target=self._cleanup_worker, daemon=True, name="FFmpegCleanup"
        )
        self._cleanup_thread.start()
        self._stats = {
            "extraction_attempts": 0,
            "successful_extractions": 0,
            "failed_extractions": 0,
            "cache_hits": 0,
            "start_time": time.time(),
        }
        logger.info(f"✅ FFmpeg Manager initialized (Platform: {SYSTEM})")

    def set_stream_manager(self, stream_manager: StreamManager) -> "FFmpegManager":
        if stream_manager:
            self.stream_manager = stream_manager
            logger.info("✅ FFmpegManager: StreamManager linked")
        return self

    def _build_ffmpeg_command_optimized(
        self,
        url: str,
        seek_seconds: Optional[float] = None,
        detected_language: Optional[str] = None,
    ) -> List[str]:

        is_live, platform = self._detect_stream_type(url)
        stream_type = "LIVE" if is_live else "VIDEO"
        logger.info(f"\n🎬 Building FFmpeg command for {platform} ({stream_type})")
        logger.info(f"  📍 URL: {url[:80]}...")

        cmd = ["ffmpeg", "-hide_banner", "-loglevel", "warning"]

        if "youtube.com" in url.lower() or "youtu.be" in url.lower():
            logger.info("  🎯 Adding YouTube-specific headers")
            headers_dict = self.config.get_youtube_headers(
                is_manifest="manifest.googlevideo.com" in url
            )
            headers_list = [f"{k}: {v}" for k, v in headers_dict.items()]
            headers_string = "\r\n".join(headers_list)
            cmd.extend(["-headers", headers_string])

        if is_live:
            logger.info("  📡 LIVE: Using optimized HLS/Live parameters")
            cmd.extend(
                [
                    "-reconnect",
                    "1",
                    "-reconnect_streamed",
                    "1",
                    "-reconnect_delay_max",
                    "10",
                    "-reconnect_on_network_error",
                    "1",
                    "-timeout",
                    "10000000",
                    "-rw_timeout",
                    "30000000",
                    "-multiple_requests",
                    "1",
                    "-seekable",
                    "0",
                    "-fflags",
                    "+discardcorrupt+fastseek+genpts",
                    "-analyzeduration",
                    "10M",
                    "-probesize",
                    "10M",
                ]
            )
            if seek_seconds is not None:
                logger.warning(
                    f"⚠️ seek_seconds={seek_seconds} wird bei Live-Stream ignoriert"
                )
        else:
            logger.info("  🎬 VIDEO: Fast access for non-live content")
            cmd.extend(
                [
                    "-rw_timeout",
                    "10000000",
                    "-accurate_seek",
                    "-fflags",
                    "+genpts+discardcorrupt+fastseek",
                ]
            )
            if seek_seconds is not None and seek_seconds > 0:
                logger.info(f"  ⏩ Seeking to {seek_seconds}s")
                cmd.extend(["-ss", str(seek_seconds)])

        cmd.extend(["-i", url])

        profile = "realtime" if is_live else "transcription"
        if self.settings and hasattr(self.settings, "audio_profile"):
            profile = self.settings.audio_profile

        audio_filter = self.config.get_audio_filter(
            language=detected_language, profile=profile
        )
        logger.info(f"  🎚️ Using audio filter (profile={profile}): {audio_filter}")

        cmd.extend(
            [
                "-vn",
                "-f",
                "s16le",
                "-acodec",
                "pcm_s16le",
                "-ar",
                str(Constants.SAMPLE_RATE),
                "-ac",
                str(Constants.CHANNELS),
                "-af",
                audio_filter,
                "-fflags",
                "+genpts+discardcorrupt",
                "-avoid_negative_ts",
                "make_zero",
                "-max_interleave_delta",
                "0",
                "-threads",
                "2",
                "-bufsize",
                "2048k",
                "pipe:1",
            ]
        )

        if logger.isEnabledFor(logging.DEBUG):
            log_debug("ffmpeg", f"Kommando: {' '.join(cmd)}")
        return cmd

    def start_stream(
        self,
        video_url: str,
        output_queue: Optional[queue.Queue],
        process_id: str,
        force_refresh_audio_url: bool = False,
        audio_url: Optional[str] = None,
        seek_seconds: Optional[float] = None,
        detected_language: Optional[str] = None,
    ) -> Optional[subprocess.Popen]:
        logger.info(f"\n🎬 FFmpegManager: Starting stream for: {video_url[:80]}...")
        with self._lock:
            if self.is_active(process_id):
                logger.warning(f"⚠️ Stream {process_id} already active")
                return None

        if audio_url is None:
            logger.info("🎵 Resolving audio URL...")
            audio_url = self.stream_manager.extract_audio_url(
                video_url, force_refresh=force_refresh_audio_url
            )
            if not audio_url:
                logger.error("❌ Audio URL resolution failed")
                return None
            logger.info(f"✅ Resolved URL: {audio_url[:100]}...")
        else:
            logger.info(f"✅ Using pre-resolved audio URL: {audio_url[:100]}...")

        cmd = self._build_ffmpeg_command_optimized(
            audio_url, seek_seconds=seek_seconds, detected_language=detected_language
        )

        try:
            process_kwargs = {
                "stdout": subprocess.PIPE,
                "stderr": subprocess.PIPE,
                "stdin": subprocess.DEVNULL,
                "bufsize": 10 * 1024 * 1024,
            }
            if IS_WINDOWS:
                process_kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
                process_kwargs["encoding"] = "utf-8"
                process_kwargs["errors"] = "ignore"
            elif IS_MACOS or IS_LINUX:
                process_kwargs["start_new_session"] = True

            logger.info("🚀 Starting FFmpeg process...")
            process = subprocess.Popen(cmd, **process_kwargs)
            logger.info(f"✅ FFmpeg process started (PID: {process.pid})")
            logger.info("⏳ Waiting for first audio data...")

            wait_timeout = self.FIRST_DATA_TIMEOUT
            end_time = time.time() + wait_timeout
            data_received = False

            if IS_WINDOWS:
                while time.time() < end_time and not data_received:
                    if process.poll() is not None:
                        break
                    try:
                        time.sleep(0.2)
                        import msvcrt
                        fd = process.stdout.fileno()
                        if msvcrt.kbhit():
                            try:
                                os.set_blocking(fd, False)
                                chunk = os.read(fd, 1)
                                if chunk:
                                    data_received = True
                            except (OSError, BlockingIOError):
                                pass
                    except Exception:
                        pass
            else:
                import select
                fd = process.stdout.fileno()
                while time.time() < end_time and not data_received:
                    if process.poll() is not None:
                        break
                    rlist, _, _ = select.select([fd], [], [], 0.2)
                    if fd in rlist:
                        time.sleep(0.2)
                        data_received = True
                        break

            if not data_received:
                if process.poll() is not None:
                    try:
                        stderr_output = process.stderr.read(1000).decode(
                            "utf-8", errors="ignore"
                        )
                        if logger.isEnabledFor(logging.DEBUG):
                            log_debug(
                                "ffmpeg",
                                f"FFmpeg died immediately, stderr: {stderr_output}",
                            )
                        logger.error(
                            f"❌ FFmpeg died immediately. Exit code: {process.poll()}"
                        )
                        if stderr_output:
                            logger.error("📋 FFMPEG STDERR (first 200 chars):")
                            logger.error(stderr_output[:200])
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.warning(f"⚠️ Could not read stderr: {e}")
                    return None
                else:
                    logger.warning(
                        f"⚠️ No data from FFmpeg within {wait_timeout}s, but process still running – continuing anyway."
                    )

            if any(
                keyword in audio_url.lower()
                for keyword in ["hls", ".m3u8", "manifest.googlevideo.com"]
            ):
                time.sleep(self.HLS_INITIAL_WAIT)

            logger.info(f"✅ FFmpeg is running (PID: {process.pid})")
            self._register_process(process_id, process, output_queue, audio_url)
            return process

        except FileNotFoundError:
            logger.error("❌ FFmpeg not found! Please install FFmpeg.")
            return None
        except PermissionError:
            logger.error("❌ Permission denied - cannot execute FFmpeg")
            return None
        except subprocess.TimeoutExpired:
            logger.error("❌ Timeout beim Start von FFmpeg")
            return None
        except OSError as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ OS-Fehler beim Start von FFmpeg: {e}")
            return None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Failed to start FFmpeg: {e}")
            return None

    def _register_process(
        self,
        process_id: str,
        process: subprocess.Popen,
        output_queue: Optional[queue.Queue],
        url: str,
    ) -> None:
        with self._lock:
            is_live, platform = self._detect_stream_type(url)
            headers_used = (
                "youtube.com" in url.lower() or "googlevideo.com" in url.lower()
            )
            pinfo = self._ProcessInfo(
                process_id=process_id,
                process=process,
                output_queue=output_queue,
                url=url,
                platform=platform,
                is_live=is_live,
                headers_used=headers_used,
            )
            self._processes[process_id] = pinfo
            self._pid_tracking[process.pid] = {
                "process_id": process_id,
                "start_time": pinfo.start_time,
                "url": url[:100],
                "platform": platform,
                "is_live": is_live,
            }
            logger.info(f"📊 Process registered: {process_id} (PID: {process.pid})")

    def update_process_activity(self, process_id: str) -> None:
        with self._lock:
            if process_id in self._processes:
                self._processes[process_id].last_activity = time.time()

    def _detect_stream_type(self, url: str) -> Tuple[bool, str]:
        cache_key = hashlib.md5(url.encode()).hexdigest()[:16]
        with self._lock:
            if cache_key in self._live_detection_cache:
                cached = self._live_detection_cache[cache_key]
                if time.time() - cached["timestamp"] < 300:
                    return cached["is_live"], cached["platform"]

        is_live = False
        platform = "unknown"
        try:
            url_lower = url.lower()
            if "youtube.com" in url_lower or "youtu.be" in url_lower:
                platform = "YouTube"
                is_live = any(
                    indicator in url_lower
                    for indicator in [
                        "/live", "live=1", "/stream", "livestream",
                        "live/", "&live", "?live", "/watch_live"
                    ]
                )
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug(
                        "network",
                        f"YouTube detection: {url_lower[:100]}, is_live={is_live}",
                    )
            elif "twitch.tv" in url_lower:
                platform = "Twitch"
                is_live = True
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("network", "Twitch detected -> is_live=True")
            elif "tiktok.com" in url_lower:
                platform = "TikTok"
                is_live = "live" in url_lower
            elif "facebook.com" in url_lower or "fb.watch" in url_lower:
                platform = "Facebook"
                is_live = "live" in url_lower or "watch/live" in url_lower
            elif "kick.com" in url_lower:
                platform = "Kick"
                is_live = True
            elif "rumble.com" in url_lower:
                platform = "Rumble"
                is_live = "live" in url_lower
            elif "dailymotion.com" in url_lower:
                platform = "Dailymotion"
                is_live = "live" in url_lower
            elif "vimeo.com" in url_lower:
                platform = "Vimeo"
                is_live = "live" in url_lower
            elif "twitter.com" in url_lower or "x.com" in url_lower:
                platform = "Twitter/X"
                is_live = "live" in url_lower
            elif url_lower.startswith("file://"):
                platform = "Local File"
                is_live = False
            elif ".m3u8" in url_lower:
                platform = "HLS Stream"
                is_live = True
            elif ".mpd" in url_lower:
                platform = "DASH Stream"
                is_live = True
            else:
                platform = "HTTP Stream"

            with self._lock:
                self._live_detection_cache[cache_key] = {
                    "is_live": is_live,
                    "platform": platform,
                    "timestamp": time.time(),
                    "url": url[:50],
                }
                if len(self._live_detection_cache) > 100:
                    oldest = min(
                        self._live_detection_cache.items(),
                        key=lambda x: x[1]["timestamp"],
                    )[0]
                    del self._live_detection_cache[oldest]

            return is_live, platform
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Stream type detection error: {e}")
            return False, "unknown"

    def get_stats(self) -> Dict[str, Any]:
        with self._lock:
            stats = self._stats.copy()
            stats["uptime_seconds"] = time.time() - stats["start_time"]
            if stats["extraction_attempts"] > 0:
                stats["success_rate"] = (
                    stats["successful_extractions"] / stats["extraction_attempts"] * 100
                )
                stats["failure_rate"] = (
                    stats["failed_extractions"] / stats["extraction_attempts"] * 100
                )
            else:
                stats["success_rate"] = stats["failure_rate"] = 0
            stats["active_processes"] = len(
                [p for p in self._processes.values() if p.process.poll() is None]
            )
            stats["total_processes"] = len(self._processes)
            stats["live_detection_cache_size"] = len(self._live_detection_cache)
            return stats

    def read_audio_data(self, process_id: str, size: int) -> Optional[bytes]:
        with self._lock:
            if process_id not in self._processes:
                return None
            pinfo = self._processes[process_id]
            if pinfo.stopping:
                return None
            process = pinfo.process

        try:
            audio_data = process.stdout.read(size)
            if audio_data:
                with self._lock:
                    pinfo.bytes_read += len(audio_data)
                    pinfo.chunks_processed += 1
                    pinfo.last_activity = time.time()
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug(
                        "ffmpeg",
                        f"{len(audio_data)} Bytes gelesen von PID {process.pid}",
                    )
                return audio_data
            else:
                if process.poll() is not None:
                    exit_code = process.poll()
                    logger.warning(
                        f"⚠️ Process {process_id} terminated (exit: {exit_code})"
                    )
                    try:
                        stderr = process.stderr.read(300).decode(
                            "utf-8", errors="ignore"
                        )
                        if stderr:
                            logger.info(f"📝 Last error: {stderr[:150]}")
                            if logger.isEnabledFor(logging.DEBUG):
                                log_debug(
                                    "ffmpeg", f"Prozess beendet, stderr: {stderr}"
                                )
                    except Exception:
                        pass
                    self.stop_stream(process_id)
                    return None
                return None
        except (IOError, OSError, ValueError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Read error for {process_id}: {e}")
            self.stop_stream(process_id)
            return None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Unexpected read error for {process_id}: {e}")
            self.stop_stream(process_id)
            return None

    def stop_stream(self, process_id: str) -> bool:
        with self._lock:
            if process_id not in self._processes:
                return True
            pinfo = self._processes[process_id]
            if pinfo.stopping:
                return True
            pinfo.stopping = True
            process = pinfo.process
            termination_success = False

        try:
            if process.poll() is None:
                logger.info(f"🔄 Stopping process {process_id} ({process.pid})...")
                try:
                    process.terminate()
                    process.wait(timeout=1.0)
                    termination_success = True
                    logger.info(f"✅ Process {process_id} terminated gracefully")
                except subprocess.TimeoutExpired:
                    try:
                        process.kill()
                        process.wait(timeout=1.0)
                        termination_success = True
                        logger.info(f"✅ Process {process_id} killed")
                    except subprocess.TimeoutExpired:
                        try:
                            import psutil

                            p = psutil.Process(process.pid)
                            p.terminate()
                            gone, alive = psutil.wait_procs([p], timeout=1.0)
                            if p in alive:
                                p.kill()
                                psutil.wait_procs([p], timeout=1.0)
                            termination_success = True
                            if logger.isEnabledFor(logging.DEBUG):
                                log_debug(
                                    "subprocess",
                                    f"Prozess {process.pid} mit psutil beendet",
                                )
                        except ImportError:
                            termination_success = False
                            logger.error(f"❌ Could not terminate {process_id}")
            else:
                termination_success = True
                logger.info(f"✅ Process {process_id} already terminated")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Error stopping {process_id}: {e}")
            termination_success = False
        finally:
            self._cleanup_process_resources(process_id, process)

        if logger.isEnabledFor(logging.DEBUG):
            log_debug(
                "subprocess",
                f"stop_stream({process_id}) -> success={termination_success}",
            )
        return termination_success

    def _cleanup_process_resources(
        self, process_id: str, process: subprocess.Popen
    ) -> None:
        with self._lock:
            if process_id not in self._processes:
                return
            del self._processes[process_id]
            if process.pid in self._pid_tracking:
                del self._pid_tracking[process.pid]

        for pipe_name in ("stdout", "stderr", "stdin"):
            pipe = getattr(process, pipe_name, None)
            if pipe and not pipe.closed:
                try:
                    pipe.close()
                except Exception:
                    pass

        if process.poll() is None:
            try:
                process.kill()
                time.sleep(0.1)
            except Exception:
                pass

        logger.debug(f"🧹 Resources cleaned for: {process_id}")

    def stop_all_streams(self) -> None:
        logger.info("🛑 Stopping all streams...")
        with self._lock:
            self._shutting_down = True
            process_ids = list(self._processes.keys())
            success_count = fail_count = 0
            for process_id in process_ids:
                try:
                    if self.stop_stream(process_id):
                        success_count += 1
                    else:
                        fail_count += 1
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Error stopping {process_id}: {e}")
                    fail_count += 1
            self._shutting_down = False
            logger.info(
                f"✅ Streams stopped: {success_count} successful, {fail_count} failed"
            )

    def is_active(self, process_id: str) -> bool:
        with self._lock:
            if process_id not in self._processes:
                return False
            pinfo = self._processes[process_id]
            process = pinfo.process
            if process.poll() is not None:
                return False
            if time.time() - pinfo.last_activity > 30:
                return False
            return True

    def _cleanup_worker(self) -> None:
        while self._cleanup_running:
            try:
                time.sleep(self.CLEANUP_INTERVAL)
                if not self._cleanup_running:
                    break
                self.cleanup_stale_processes()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Cleanup worker error: {e}")

    def cleanup_stale_processes(self) -> None:
        with self._lock:
            stale_pids = [
                pid
                for pid, pinfo in self._processes.items()
                if pinfo.process.poll() is not None
            ]
        for pid in stale_pids:
            with self._lock:
                if pid not in self._processes:
                    continue
                pinfo = self._processes[pid]
                process = pinfo.process
            logger.info(f"🧹 Cleaning terminated process: {pid}")
            self._cleanup_process_resources(pid, process)

    def dispose(self) -> None:
        logger.info("🧹 Shutting down FFmpeg Manager...")
        self._cleanup_running = False
        self.stop_all_streams()
        if self._cleanup_thread and self._cleanup_thread.is_alive():
            self._cleanup_thread.join(timeout=2.0)
        self._live_detection_cache.clear()
        self._pid_tracking.clear()
        self._processes.clear()
        gc.collect()
        logger.info("✅ FFmpeg Manager disposed")


# -----------------------------------------------------------------------------
# StreamInfoExtractor
# -----------------------------------------------------------------------------
class StreamInfoExtractor:
    def __init__(self) -> None:
        self.current_info = StreamInfo(
            title="Unknown Stream",
            uploader="Unknown",
            duration="Live",
            view_count=0,
            platform="Unknown",
        )
        self._lock = threading.RLock()
        self._debug = DEBUG_LEVEL >= 1
        self.use_browser_cookies = True

    def extract_stream_info(self, url: str) -> StreamInfo:
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug(
                f"🔍 StreamInfoExtractor.extract_stream_info für: {url[:80]}..."
            )
        url = PlatformUtils.sanitize_url(url)
        if url.startswith("file://"):
            return self._handle_local_file(url)
        if "youtube.com" in url.lower() or "youtu.be" in url.lower():
            if self.use_browser_cookies:
                info = self._extract_youtube_info_with_cookies(url)
                if info:
                    self.current_info = info
                    return info
            info = self._run_ytdlp_json(url, platform="youtube")
            if info:
                self.current_info = info
                return info
        if "twitch.tv" in url.lower():
            if self.use_browser_cookies:
                info = self._extract_twitch_info_with_cookies(url)
                if info:
                    self.current_info = info
                    return info
            info = self._run_ytdlp_json(url, platform="twitch")
            if info:
                self.current_info = info
                return info
        info = self._run_ytdlp_json(url)
        if info:
            self.current_info = info
            return info
        info = self._fallback_from_url(url)
        self.current_info = info
        return info

    def _handle_local_file(self, url: str) -> StreamInfo:
        ok, real_path = PlatformUtils.validate_file_path(url)
        if not ok:
            return StreamInfo(
                title="Invalid file",
                uploader="Error",
                duration="",
                view_count=0,
                platform="invalid",
            )
        file_path = real_path
        return StreamInfo(
            title=os.path.basename(file_path),
            uploader="Local File",
            duration="File",
            view_count=0,
            platform="local",
        )

    def _run_ytdlp_json(
        self, url: str, platform: str = "generic"
    ) -> Optional[StreamInfo]:
        timeout = 15 if platform in ("youtube", "twitch") else 10
        try:
            data = YtDlpHelper.get_json(url, timeout=timeout)
            if not data:
                return None
            extractor = data.get("extractor", "").lower()
            platform_map = {
                "youtube": "youtube",
                "twitch": "twitch",
                "tiktok": "tiktok",
                "facebook": "facebook",
                "kick": "kick",
                "rumble": "rumble",
                "dailymotion": "dailymotion",
                "vimeo": "vimeo",
                "twitter": "twitter",
                "x": "twitter",
            }
            detected_platform = "unknown"
            for key, value in platform_map.items():
                if key in extractor:
                    detected_platform = value
                    break
            description = data.get("description", "")
            if len(description) > 200:
                description = description[:200] + "..."
            return StreamInfo(
                title=data.get("title", "Unknown Title"),
                uploader=data.get(
                    "uploader", data.get("channel", data.get("creator", "Unknown"))
                ),
                duration=data.get("duration_string", "Live"),
                view_count=data.get("view_count", 0),
                platform=detected_platform,
                description=description,
                duration_seconds=data.get("duration"),
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(f"⚠️ yt-dlp JSON fehlgeschlagen für {url[:50]}: {e}")
        return None

    def _fallback_from_url(self, url: str) -> StreamInfo:
        if self._debug or logger.isEnabledFor(logging.DEBUG):
            logger.debug("🔄 StreamInfoExtractor: Fallback – extrahiere Titel aus URL")
        try:
            parsed = urllib.parse.urlparse(url)
            domain = parsed.netloc.replace("www.", "")
            path_segments = [s for s in parsed.path.split("/") if s]
            if "twitch.tv" in domain:
                channel = (
                    path_segments[0] if path_segments else domain.replace(".tv", "")
                )
                title = f"{channel} (Twitch)"
                uploader = channel
                platform = "twitch"
            elif "youtube.com" in domain or "youtu.be" in domain:
                channel = (
                    path_segments[0]
                    if path_segments
                    and path_segments[0] not in ("watch", "playlist", "shorts")
                    else "YouTube"
                )
                title = f"YouTube Stream - {channel}"
                uploader = channel
                platform = "youtube"
            else:
                if path_segments:
                    last = path_segments[-1]
                    title = f"{domain} - {last}"
                else:
                    title = domain
                uploader = domain
                platform = "unknown"
            return StreamInfo(
                title=title,
                uploader=uploader,
                duration="Live" if "live" in url.lower() else "Unknown",
                view_count=0,
                platform=platform,
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if self._debug or logger.isEnabledFor(logging.DEBUG):
                logger.debug(
                    f"❌ StreamInfoExtractor: URL-Fallback fehlgeschlagen: {e}"
                )
        return StreamInfo(
            title="Unknown Stream",
            uploader="Unknown",
            duration="Live",
            view_count=0,
            platform="unknown",
        )

    def _extract_youtube_info_with_cookies(self, url: str) -> Optional[StreamInfo]:
        logger.info(
            "  🎯 YouTube detected, trying optimized cookie methods for channel name..."
        )
        if IS_LINUX:
            self._ensure_chrome_symlinks()
        methods = self._build_youtube_methods(url)
        logger.info(f"    📋 Using {len(methods)} optimized extraction methods")
        max_attempts = min(3, len(methods))
        attempts = 0
        for cmd, method_name in methods:
            if attempts >= max_attempts:
                break
            attempts += 1
            try:
                logger.info(f"    🧪 Attempt {attempts}/{max_attempts}: {method_name}")
                timeout = 12 if "Cookies" in method_name else 8
                stdout = YtDlpHelper.run_command(cmd, timeout=timeout, method_name=method_name)
                if not stdout:
                    continue
                output = stdout
                json_start = output.find("{")
                json_end = output.rfind("}") + 1
                if json_start >= 0 and json_end > json_start:
                    try:
                        json_str = output[json_start:json_end]
                        info = json.loads(json_str)
                        uploader = info.get("uploader", "Unknown")
                        channel = info.get("channel", uploader)
                        creator = info.get("creator", uploader)
                        final_uploader = uploader
                        if channel != "Unknown" and channel != uploader:
                            final_uploader = channel
                        elif creator != "Unknown" and creator != uploader:
                            final_uploader = creator
                        if final_uploader == "Unknown":
                            final_uploader = info.get("uploader_id", "YouTube")
                        logger.info(f"      ✅ Success with {method_name}")
                        logger.info(
                            f"        Title: {info.get('title', 'YouTube Stream')[:60]}..."
                        )
                        logger.info(f"        Channel: {final_uploader}")
                        return StreamInfo(
                            title=info.get("title", "YouTube Stream"),
                            uploader=final_uploader,
                            duration=info.get("duration_string", "Live"),
                            view_count=info.get("view_count", 0),
                            platform="youtube",
                            description=(
                                info.get("description", "")[:200] + "..."
                                if len(info.get("description", "")) > 200
                                else info.get("description", "")
                            ),
                            duration_seconds=info.get("duration"),
                        )
                    except json.JSONDecodeError:
                        pass
                lines = output.split("\n")
                for line in lines:
                    if (
                        line.strip()
                        and not line.startswith("{")
                        and len(line.strip()) > 10
                    ):
                        possible_title = line.strip()
                        if len(possible_title) > 20 and len(possible_title) < 200:
                            logger.info("      ✅ Extracted title from output")
                            return StreamInfo(
                                title=possible_title,
                                uploader="YouTube",
                                duration="Live",
                                view_count=0,
                                platform="youtube",
                                description="",
                            )
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.info(f"      ⚠️ Method error: {str(e)[:50]}")
                continue
        logger.info("    🔄 Ultimate fallback: Direct title extraction...")
        return self._direct_youtube_fallback(url)

    def _build_youtube_methods(self, url: str) -> List[Tuple[List[str], str]]:
        methods = []
        if self.use_browser_cookies:
            browsers = self._get_browser_list()
            for browser_cmd, browser_name in browsers:
                methods.append(
                    (
                        [
                            "yt-dlp",
                            "--cookies-from-browser",
                            browser_cmd,
                            "--dump-json",
                            "--no-warnings",
                            "--no-check-certificate",
                            "--playlist-items",
                            "1",
                            "--",
                            url,
                        ],
                        f"{browser_name} Cookies",
                    )
                )
        fallback_methods = [
            (
                [
                    "yt-dlp",
                    "--dump-json",
                    "--no-warnings",
                    "--no-check-certificate",
                    "--playlist-items",
                    "1",
                    "--quiet",
                    "--",
                    url,
                ],
                "No Cookies (Quiet)",
            ),
            (
                [
                    "yt-dlp",
                    "--dump-json",
                    "--no-warnings",
                    "--no-check-certificate",
                    "--playlist-items",
                    "1",
                    "--",
                    url,
                ],
                "Simple JSON",
            ),
            (
                [
                    "yt-dlp",
                    "--get-title",
                    "--get-description",
                    "--get-duration",
                    "--no-warnings",
                    "--no-check-certificate",
                    "--quiet",
                    "--",
                    url,
                ],
                "Direct Info",
            ),
        ]
        methods.extend(fallback_methods)
        return methods

    def _get_browser_list(self) -> List[Tuple[str, str]]:
        if IS_LINUX:
            return [
                ("firefox", "Firefox"),
                ("chromium", "Chromium"),
                ("brave", "Brave"),
                ("chrome", "Chrome"),
                ("vivaldi", "Vivaldi"),
                ("opera", "Opera"),
                ("edge", "Edge"),
            ]
        elif IS_WINDOWS:
            return [
                ("chrome", "Chrome"),
                ("firefox", "Firefox"),
                ("edge", "Edge"),
                ("brave", "Brave"),
                ("opera", "Opera"),
            ]
        else:
            return [
                ("safari", "Safari"),
                ("chrome", "Chrome"),
                ("firefox", "Firefox"),
                ("brave", "Brave"),
                ("edge", "Edge"),
            ]

    def _ensure_chrome_symlinks(self) -> None:
        try:
            chrome_config_dir = Path.home() / ".config" / "google-chrome"
            chromium_config_dir = Path.home() / ".config" / "chromium"
            if chromium_config_dir.exists() and not chrome_config_dir.exists():
                chrome_config_dir.mkdir(parents=True, exist_ok=True)
                chromium_files = [
                    "Local State",
                    "Default/Cookies",
                    "Default/Login Data",
                ]
                for file_path in chromium_files:
                    chromium_file = chromium_config_dir / file_path
                    chrome_file = chrome_config_dir / file_path
                    if chromium_file.exists() and not chrome_file.exists():
                        chrome_file.parent.mkdir(parents=True, exist_ok=True)
                        os.symlink(str(chromium_file), str(chrome_file))
                logger.info("    🔗 Created Chrome compatibility symlinks for yt-dlp")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"    ⚠️ Chrome symlink setup failed: {e}")

    def _direct_youtube_fallback(self, url: str) -> Optional[StreamInfo]:
        try:
            cmd_title = [
                "yt-dlp",
                "--get-title",
                "--no-warnings",
                "--no-check-certificate",
                "--quiet",
                "--",
                url,
            ]
            cmd_uploader = [
                "yt-dlp",
                "--get-filename",
                "-o",
                "%(uploader)s",
                "--no-warnings",
                "--no-check-certificate",
                "--quiet",
                "--",
                url,
            ]
            with ThreadPoolExecutor(max_workers=2) as executor:
                title_future = executor.submit(
                    YtDlpHelper.run_command, cmd_title, 8, "direct_title"
                )
                uploader_future = executor.submit(
                    YtDlpHelper.run_command, cmd_uploader, 8, "direct_uploader"
                )
                title_result = title_future.result(timeout=10)
                uploader_result = uploader_future.result(timeout=10)
            title = "YouTube Stream"
            uploader = "YouTube"
            if title_result and title_result.strip():
                title = title_result.strip().split("\n")[0]
            if uploader_result and uploader_result.strip():
                uploader = uploader_result.strip().split("\n")[0]
            logger.info("      ✅ Success with direct extraction")
            return StreamInfo(
                title=title[:100] if len(title) > 100 else title,
                uploader=uploader,
                duration="Live",
                view_count=0,
                platform="youtube",
                description="",
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.info(f"      ⚠️ Direct extraction failed: {e}")
            return None

    def _extract_twitch_info_with_cookies(self, url: str) -> Optional[StreamInfo]:
        logger.info("  🎯 Twitch detected, trying cookie methods for channel name...")
        methods = self._build_twitch_methods(url)
        logger.info(f"    📋 Using {len(methods)} extraction methods for Twitch")
        max_attempts = min(4, len(methods))
        attempts = 0
        for cmd, method_name in methods:
            if attempts >= max_attempts:
                break
            attempts += 1
            try:
                logger.info(f"    🧪 Attempt {attempts}/{max_attempts}: {method_name}")
                timeout = 12 if "Cookies" in method_name else 8
                stdout = YtDlpHelper.run_command(cmd, timeout=timeout, method_name=method_name)
                if not stdout:
                    continue
                output = stdout
                try:
                    json_start = output.find("{")
                    json_end = output.rfind("}") + 1
                    if json_start >= 0 and json_end > json_start:
                        json_str = output[json_start:json_end]
                        info = json.loads(json_str)
                        uploader = info.get(
                            "uploader",
                            info.get("channel", info.get("creator", "Unknown")),
                        )
                        title = info.get("title", "Twitch Stream")
                        duration = info.get("duration_string", "Live")
                        view_count = info.get("view_count", 0)
                        description = info.get("description", "")
                        if len(description) > 200:
                            description = description[:200] + "..."
                        logger.info(f"      ✅ Success with {method_name}")
                        logger.info(f"        Title: {title[:60]}...")
                        logger.info(f"        Channel: {uploader}")
                        return StreamInfo(
                            title=title,
                            uploader=uploader,
                            duration=duration,
                            view_count=view_count,
                            platform="twitch",
                            description=description,
                            duration_seconds=info.get("duration"),
                        )
                except json.JSONDecodeError:
                    lines = output.split("\n")
                    title = None
                    uploader = None
                    for line in lines:
                        if line.strip() and not line.startswith("{"):
                            if title is None:
                                title = line.strip()
                            elif uploader is None:
                                uploader = line.strip()
                                break
                    if title and len(title) > 10:
                        logger.info("      ✅ Extracted title from output")
                        return StreamInfo(
                            title=title[:100] if len(title) > 100 else title,
                            uploader=uploader or "Twitch",
                            duration="Live",
                            view_count=0,
                            platform="twitch",
                            description="",
                        )
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.info(f"      ⚠️ Method error: {str(e)[:50]}")
                continue
        return self._twitch_url_fallback(url)

    def _build_twitch_methods(self, url: str) -> List[Tuple[List[str], str]]:
        methods = []
        if self.use_browser_cookies:
            for browser_cmd, browser_name in self._get_browser_list():
                methods.append(
                    (
                        [
                            "yt-dlp",
                            "--cookies-from-browser",
                            browser_cmd,
                            "--dump-json",
                            "--format",
                            "best",
                            "--no-warnings",
                            "--no-check-certificate",
                            "--socket-timeout",
                            "15",
                            "--",
                            url,
                        ],
                        f"{browser_name} Cookies",
                    )
                )
        fallback_methods = [
            (
                [
                    "yt-dlp",
                    "--dump-json",
                    "--no-warnings",
                    "--no-check-certificate",
                    "--socket-timeout",
                    "10",
                    "--",
                    url,
                ],
                "Simple JSON",
            ),
            (
                [
                    "yt-dlp",
                    "--get-title",
                    "--get-description",
                    "--get-duration",
                    "--no-warnings",
                    "--no-check-certificate",
                    "--quiet",
                    "--",
                    url,
                ],
                "Direct Info",
            ),
        ]
        methods.extend(fallback_methods)
        return methods

    def _twitch_url_fallback(self, url: str) -> Optional[StreamInfo]:
        try:
            parsed = urllib.parse.urlparse(url)
            path = parsed.path.strip("/")
            channel = (
                path.split("/")[0]
                if path
                else parsed.netloc.replace("www.", "").replace(".tv", "")
            )
            if channel:
                return StreamInfo(
                    title=f"{channel} (Twitch Live)",
                    uploader=channel,
                    duration="Live",
                    view_count=0,
                    platform="twitch",
                    description="",
                )
        except Exception:
            pass
        return None


# -----------------------------------------------------------------------------
# ExportManager
# -----------------------------------------------------------------------------
class ExportManager:
    def __init__(self) -> None:
        self.supported_formats = ["txt", "srt", "vtt", "json", "docx"]
        self._docx_available: bool = False
        try:
            import docx

            self._docx = docx
            self._docx_available = True
        except ImportError:
            self._docx_available = False

    def export_subtitles(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
        format: str = "srt",
        filename: Optional[str] = None,
        encoding: str = "utf-8-sig",
    ) -> Union[bool, str]:
        try:
            if format.lower() not in ("txt", "json"):
                timed = [
                    t
                    for t in transcript_data
                    if hasattr(t, "start")
                    and t.start is not None
                    and hasattr(t, "end")
                    and t.end is not None
                ]
                if not timed:
                    raise ProcessingError(
                        "Keine Segmente mit Zeitstempeln vorhanden – benötigt für Untertitel."
                    )
                transcript_data = timed
            if format.lower() == "srt":
                content = self._generate_srt_content(transcript_data, translation_data)
            elif format.lower() == "vtt":
                content = self._generate_vtt_content(transcript_data, translation_data)
            elif format.lower() == "txt":
                content = self._generate_txt_content(transcript_data, translation_data)
            elif format.lower() == "json":
                if filename:
                    return self.export_json(
                        transcript_data, translation_data or [], filename
                    )
                else:
                    import json

                    data = self._build_json_data(transcript_data, translation_data)
                    return json.dumps(data, indent=2, ensure_ascii=False)
            elif format.lower() == "docx":
                if filename:
                    return self.export_docx(transcript_data, filename)
                else:
                    raise ProcessingError(
                        "Für DOCX-Export wird ein Dateiname benötigt."
                    )
            else:
                raise ProcessingError(f"Nicht unterstütztes Format: {format}")
            if filename:
                out_path = Path(filename)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                with open(out_path, "w", encoding=encoding) as f:
                    f.write(content)
                return True
            else:
                return content
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Fehler beim Export: {e}")
            raise ProcessingError(f"Export fehlgeschlagen: {e}") from e

    def export_json(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: List[TranslationResult],
        filename: str,
        encoding: str = "utf-8",
    ) -> bool:
        try:
            data = self._build_json_data(transcript_data, translation_data)
            out_path = Path(filename)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", encoding=encoding) as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return True
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            raise ProcessingError(f"JSON-Export fehlgeschlagen: {e}") from e

    def export_docx(
        self,
        transcript_data: List[TranscriptionResult],
        filename: str,
    ) -> bool:
        if self._docx_available:
            try:
                doc = self._docx.Document()
                doc.add_heading("Transkription", level=1)
                for i, seg in enumerate(transcript_data, 1):
                    p = doc.add_paragraph()
                    if hasattr(seg, "start") and seg.start is not None:
                        start_str = self._format_timestamp_srt(seg.start)
                        p.add_run(f"[{start_str}] ").bold = True
                    p.add_run(seg.text)
                    if i < len(transcript_data):
                        doc.add_paragraph()
                out_path = Path(filename)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(out_path))
                return True
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(
                    f"python-docx Export fehlgeschlagen, verwende Fallback: {e}"
                )
        try:
            content = self._generate_txt_content(transcript_data, None)
            out_path = Path(filename)
            out_path.parent.mkdir(parents=True, exist_ok=True)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(content)
            return True
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            raise ProcessingError(f"DOCX-Fallback fehlgeschlagen: {e}") from e

    def _generate_srt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = []
        for i, segment in enumerate(transcript_data):
            start = self._format_timestamp_srt(segment.start or 0.0)
            end = self._format_timestamp_srt(segment.end or 0.0)
            text = segment.text
            if translation_data and i < len(translation_data):
                text = f"{text}\n{translation_data[i].translated}"
            lines.append(f"{i+1}\n{start} --> {end}\n{text}\n")
        return "\n".join(lines)

    def _generate_vtt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = ["WEBVTT\n"]
        for i, segment in enumerate(transcript_data):
            start = self._format_timestamp_vtt(segment.start or 0.0)
            end = self._format_timestamp_vtt(segment.end or 0.0)
            text = segment.text
            if translation_data and i < len(translation_data):
                text = f"{text}\n{translation_data[i].translated}"
            lines.append(f"{start} --> {end}\n{text}\n")
        return "\n".join(lines)

    def _generate_txt_content(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: Optional[List[TranslationResult]] = None,
    ) -> str:
        lines = []
        for i, segment in enumerate(transcript_data):
            if hasattr(segment, "start") and segment.start is not None:
                timestamp = self._format_timestamp_srt(segment.start)
                prefix = f"[{timestamp}] "
            else:
                prefix = ""
            text = segment.text
            if translation_data and i < len(translation_data):
                text = f"{text}  |  {translation_data[i].translated}"
            lines.append(f"{prefix}{text}")
        return "\n".join(lines)

    def _build_json_data(
        self,
        transcript_data: List[TranscriptionResult],
        translation_data: List[TranslationResult],
    ) -> Dict[str, Any]:
        return {
            "metadata": {
                "export_date": datetime.now().isoformat(),
                "total_segments": len(transcript_data),
                "version": "4.1.3",
            },
            "transcripts": [
                {
                    "text": seg.text,
                    "confidence": seg.confidence,
                    "language": seg.language,
                    "timestamp": seg.timestamp,
                    "start": getattr(seg, "start", None),
                    "end": getattr(seg, "end", None),
                }
                for seg in transcript_data
            ],
            "translations": (
                [
                    {
                        "original": trans.original,
                        "translated": trans.translated,
                        "source_lang": trans.source_lang,
                        "target_lang": trans.target_lang,
                        "timestamp": trans.timestamp,
                        "start": getattr(trans, "start", None),
                        "end": getattr(trans, "end", None),
                    }
                    for trans in translation_data
                ]
                if translation_data
                else []
            ),
        }

    @staticmethod
    def _format_timestamp_srt(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        milliseconds = int((secs - int(secs)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{int(secs):02d},{milliseconds:03d}"

    @staticmethod
    def _format_timestamp_vtt(seconds: float) -> str:
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        milliseconds = int((secs - int(secs)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{int(secs):02d}.{milliseconds:03d}"


# -----------------------------------------------------------------------------
# ResourceManager
# -----------------------------------------------------------------------------
class ResourceManager:
    def __init__(self) -> None:
        self.processes: List[subprocess.Popen] = []
        self.threads: List[threading.Thread] = []
        self.temp_files: List[str] = []
        self.cleanup_done = False
        self._lock = threading.RLock()
        self._shutdown_event = threading.Event()
        self._atexit_registered = False

        self._psutil = None
        try:
            import psutil

            self._psutil = psutil
        except ImportError:
            pass

    def register_process(self, process: subprocess.Popen) -> None:
        with self._lock:
            if process and process not in self.processes:
                self.processes.append(process)

    def register_thread(self, thread: threading.Thread) -> None:
        with self._lock:
            if thread and thread not in self.threads and thread.is_alive():
                self.threads.append(thread)

    def register_temp_file(self, file_path: str) -> None:
        with self._lock:
            if file_path and file_path not in self.temp_files:
                self.temp_files.append(file_path)

    def is_shutting_down(self) -> bool:
        return self._shutdown_event.is_set()

    def cleanup(self, timeout: float = 5.0) -> None:
        if self.cleanup_done:
            return

        self._shutdown_event.set()
        start_time = time.time()

        with self._lock:
            for proc in self.processes[:]:
                if time.time() - start_time > timeout:
                    logger.warning(
                        "⚠️ Cleanup timeout – breche weitere Prozessbereinigung ab"
                    )
                    break
                try:
                    self._terminate_process(proc)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(
                        f"⚠️ Fehler beim Beenden von Prozess {proc.pid}: {e}"
                    )
                finally:
                    if proc in self.processes:
                        self.processes.remove(proc)

            for thread in self.threads[:]:
                if time.time() - start_time > timeout:
                    logger.warning("⚠️ Cleanup timeout – breche Thread-Join ab")
                    break
                try:
                    if thread.is_alive():
                        thread.join(
                            timeout=max(0.1, timeout - (time.time() - start_time))
                        )
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Fehler beim Join von Thread {thread.name}: {e}")
                finally:
                    if thread in self.threads:
                        self.threads.remove(thread)

            for temp_file in self.temp_files[:]:
                if time.time() - start_time > timeout:
                    logger.warning(
                        "⚠️ Cleanup timeout – breche Löschen temp. Dateien ab"
                    )
                    break
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except (OSError, PermissionError) as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(
                        f"⚠️ Konnte temp. Datei {temp_file} nicht löschen: {e}"
                    )
                finally:
                    if temp_file in self.temp_files:
                        self.temp_files.remove(temp_file)

            try:
                import torch

                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                    logger.debug("🧹 GPU-Cache geleert")
            except ImportError:
                pass

            gc.collect()

            self.cleanup_done = True
            logger.info("✅ ResourceManager: Alle Ressourcen bereinigt")

    def _terminate_process(self, proc: subprocess.Popen) -> None:
        if proc.poll() is not None:
            return

        pid = proc.pid
        logger.debug(f"🛑 Beende Prozess {pid}...")

        if self._psutil:
            try:
                parent = self._psutil.Process(pid)
                children = parent.children(recursive=True)
                for child in children:
                    try:
                        child.terminate()
                    except self._psutil.NoSuchProcess:
                        pass
                gone, alive = self._psutil.wait_procs(children, timeout=1.0)
                for p in alive:
                    try:
                        p.kill()
                    except self._psutil.NoSuchProcess:
                        pass
                parent.terminate()
                parent.wait(timeout=1.0)
                logger.debug(f"✅ Prozess {pid} mit psutil beendet")
                return
            except (
                self._psutil.NoSuchProcess,
                self._psutil.AccessDenied,
                AttributeError,
            ):
                pass

        try:
            proc.terminate()
            proc.wait(timeout=1.0)
        except subprocess.TimeoutExpired:
            try:
                proc.kill()
                proc.wait(timeout=0.5)
            except Exception:
                pass

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cleanup()


# -----------------------------------------------------------------------------
# LanguageDetector
# -----------------------------------------------------------------------------
class LanguageDetector:
    def __init__(self, transcription_engine: TranscriptionEngine) -> None:
        self.transcription_engine = transcription_engine

    def _get_media_duration(self, file_path: str) -> Optional[float]:
        try:
            cmd = [
                "ffprobe",
                "-v",
                "error",
                "-show_entries",
                "format=duration",
                "-of",
                "default=noprint_wrappers=1:nokey=1",
                file_path,
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0 and result.stdout.strip():
                return float(result.stdout.strip())
        except (
            subprocess.TimeoutExpired,
            subprocess.CalledProcessError,
            ValueError,
            OSError,
        ):
            pass
        except Exception:
            pass
        return None

    def detect_video_language(self, video_path: str) -> Dict[str, Any]:
        try:
            if not os.path.exists(video_path):
                return {"error": "File not found"}
            duration = self._get_media_duration(video_path)
            if duration is None:
                try:
                    file_size_mb = os.path.getsize(video_path) / (1024 * 1024)
                    if file_size_mb > 500:
                        return {"info": "Large file - direct processing recommended"}
                except OSError:
                    pass
                sample_duration = 60
            else:
                sample_duration = min(60, max(30, duration // 2))
            temp_audio = self._extract_audio_sample(
                video_path, duration=sample_duration
            )
            if not temp_audio:
                temp_audio = self._extract_audio_sample(video_path, duration=None)
            if not temp_audio:
                return {"error": "Could not extract audio"}
            result = self.transcription_engine.transcribe_audio(
                temp_audio, include_timestamps=False
            )
            if result and hasattr(result, "language"):
                language_code = result.language
                language_name = SUPPORTED_LANGUAGES.get(language_code, "Unknown")
                return {
                    "detected_language": language_code,
                    "language_name": language_name,
                    "confidence": getattr(result, "confidence", 0.8),
                    "sample_text": (
                        result.text[:100] + "..."
                        if len(result.text) > 100
                        else result.text
                    ),
                }
            else:
                return {"error": "Language could not be detected"}
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            return {"error": f"Analysis failed: {str(e)}"}

    def _extract_audio_sample(
        self, video_path: str, duration: Optional[int] = 30
    ) -> Optional[bytes]:
        try:
            config = self.transcription_engine.settings.config
            cmd = [
                "ffmpeg",
                "-i",
                video_path,
                "-f",
                config.AUDIO_FORMAT,
                "-ar",
                str(config.SAMPLE_RATE),
                "-ac",
                str(config.CHANNELS),
                "-loglevel",
                "quiet",
                "-",
            ]
            if duration is not None:
                cmd.insert(2, "-t")
                cmd.insert(3, str(duration))
            result = subprocess.run(cmd, capture_output=True, timeout=30)
            if result.returncode == 0 and result.stdout:
                return result.stdout
        except Exception:
            pass
        return None


# -----------------------------------------------------------------------------
# OllamaSummarizer
# -----------------------------------------------------------------------------
class OllamaSummarizer:
    def __init__(
        self,
        parent: Any,
        model: str = "llama3.1:8b",
        host: str = "http://localhost:11434",
        timeout: int = 120,
        cache_ttl: int = 300,
        system_prompt: Optional[str] = None,
    ) -> None:
        self.parent = parent
        self.model = model
        self.host = host.rstrip("/")
        self.timeout = timeout
        self.cache_ttl = cache_ttl
        self.system_prompt = system_prompt
        self.available = OLLAMA_AVAILABLE
        self._session = None
        self._stop_event = threading.Event()
        self._lock = threading.RLock()
        self._requests = FastLazyLoader.load("requests") if self.available else None

        self._models_cache: List[str] = []
        self._models_cache_time = 0.0

        self.last_result: Optional[str] = None

    def _get_session(self):
        if not self.available:
            return None
        with self._lock:
            if self._session is None:
                self._session = self._requests.Session()
                self._session.headers.update(
                    {
                        "Content-Type": "application/json",
                        "Accept": "application/json",
                    }
                )
        return self._session

    def get_available_models(self) -> List[str]:
        if not self.available:
            return []

        now = time.time()
        with self._lock:
            if now - self._models_cache_time < self.cache_ttl and self._models_cache:
                return self._models_cache.copy()

        try:
            session = self._get_session()
            r = session.get(f"{self.host}/api/tags", timeout=5)
            if r.status_code == 200:
                data = r.json()
                models = [m["name"] for m in data.get("models", [])]
                with self._lock:
                    self._models_cache = models
                    self._models_cache_time = now
                return models.copy()
            else:
                logger.warning(f"Ollama model list error: {r.status_code}")
        except self._requests.exceptions.RequestException as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"Ollama model list error: {e}")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"Unexpected error getting models: {e}")
        return []

    def is_model_available(self, model: Optional[str] = None) -> bool:
        check_model = model or self.model
        available = self.get_available_models()
        if not available:
            return True
        return check_model in available

    def is_server_reachable(self) -> bool:
        if not self.available:
            return False
        try:
            session = self._get_session()
            r = session.get(f"{self.host}/api/tags", timeout=2)
            return r.status_code == 200
        except Exception:
            return False

    def summarize(
        self,
        text: str,
        prompt: str,
        temperature: float,
        callback: Callable[[str], None],
        error_callback: Callable[[str], None],
        complete_callback: Optional[Callable[[], None]] = None,
        cancel_event: Optional[threading.Event] = None,
    ) -> None:
        if not self.available:
            error_callback("Ollama nicht verfügbar (requests nicht installiert)")
            return
        if not text or not text.strip():
            error_callback("Kein Text zum Zusammenfassen")
            return
        if not self.is_model_available():
            error_callback(f"Ollama-Modell '{self.model}' nicht auf Server gefunden.")
            return

        self._stop_event.clear()
        cancel = cancel_event or threading.Event()

        def worker() -> None:
            full_response = ""
            try:
                session = self._get_session()
                full_prompt = f"{prompt}\n\n{text}"
                payload = {
                    "model": self.model,
                    "prompt": full_prompt,
                    "stream": True,
                    "options": {"temperature": temperature, "num_predict": 512},
                }
                if self.system_prompt:
                    payload["system"] = self.system_prompt
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("ollama", f"Ollama request payload: {json.dumps(payload, indent=2)}")
                response = session.post(
                    f"{self.host}/api/generate",
                    json=payload,
                    stream=True,
                    timeout=self.timeout,
                )
                if response.status_code == 200:
                    for line in response.iter_lines(decode_unicode=True):
                        if self._stop_event.is_set() or cancel.is_set():
                            break
                        if line:
                            try:
                                data = json.loads(line)
                                if "response" in data:
                                    chunk = data["response"]
                                    full_response += chunk
                                    callback(chunk)
                                if data.get("done", False):
                                    break
                            except json.JSONDecodeError:
                                continue
                    if not (self._stop_event.is_set() or cancel.is_set()):
                        if not full_response:
                            error_callback("Leere Antwort von Ollama")
                        else:
                            self.last_result = full_response
                            if complete_callback:
                                complete_callback()
                else:
                    error_callback(f"Ollama Fehler {response.status_code}")
            except self._requests.exceptions.Timeout:
                error_callback(
                    f"Ollama Timeout nach {self.timeout}s – Server nicht erreichbar?"
                )
            except self._requests.exceptions.ConnectionError:
                error_callback("Ollama nicht erreichbar (läuft der Server?)")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                error_callback(f"Fehler: {str(e)}")

        threading.Thread(target=worker, daemon=True).start()

    def correct_transcript(
        self,
        text: str,
        callback: Callable[[str], None],
        error_callback: Callable[[str], None],
        complete_callback: Optional[Callable[[], None]] = None,
    ) -> None:
        if not self.available:
            error_callback("Ollama nicht verfügbar (requests fehlt)")
            return
        if not self.is_model_available():
            error_callback(f"Ollama-Modell '{self.model}' nicht auf Server gefunden.")
            return

        prompt = (
            "Du bist ein Assistent, der Transkriptionen korrigiert. "
            "Der folgende Text wurde automatisch transkribiert und enthält typische Fehler: "
            "fehlende Satzzeichen, falsche Groß-/Kleinschreibung, leichte Wortverzerrungen. "
            "Korrigiere diese Fehler und verbessere die Lesbarkeit. "
            "Ändere keine inhaltlichen Aussagen, füge keine neuen Informationen hinzu. "
            "Gib nur den korrigierten Text zurück, ohne Einleitung oder Erklärung.\n\n"
            f"{text}"
        )

        self.summarize(
            text=text,
            prompt=prompt,
            temperature=0.0,
            callback=callback,
            error_callback=error_callback,
            complete_callback=complete_callback,
            cancel_event=None,
        )

    def stop(self) -> None:
        self._stop_event.set()
        logger.info("OllamaSummarizer: Stop signalisiert")

    def dispose(self) -> None:
        self.stop()
        if self._session:
            try:
                self._session.close()
            except Exception:
                pass
            self._session = None
        with self._lock:
            self._models_cache = []
            self._models_cache_time = 0.0


# =============================================================================
# QueueManager
# =============================================================================
class QueueManager:
    GUI_QUEUE_MAX_SIZE = 200
    GUI_QUEUE_CLEANUP_TARGET = 100
    TEXT_QUEUE_MAX_SIZE = 150
    TEXT_QUEUE_CLEANUP_TARGET = 75
    GUI_PROCESS_MAX_ITEMS = 50
    TEXT_PROCESS_MAX_ITEMS = 20
    GUI_PROCESS_MAX_DURATION = 0.05
    TEXT_PROCESS_MAX_DURATION = 0.05

    IMPORTANT_MSG_TYPES = {"status", "error", "file_finished"}

    def __init__(self, gui: "DragonWhispererGUI", **kwargs):
        self.gui = gui
        self.root = gui.root
        self._last_gui_update_time = 0.0
        self._gui_update_limiter = gui._gui_update_limiter

        self.GUI_QUEUE_MAX_SIZE = kwargs.get("gui_queue_max_size", self.GUI_QUEUE_MAX_SIZE)
        self.GUI_QUEUE_CLEANUP_TARGET = kwargs.get("gui_queue_cleanup_target", self.GUI_QUEUE_CLEANUP_TARGET)
        self.TEXT_QUEUE_MAX_SIZE = kwargs.get("text_queue_max_size", self.TEXT_QUEUE_MAX_SIZE)
        self.TEXT_QUEUE_CLEANUP_TARGET = kwargs.get("text_queue_cleanup_target", self.TEXT_QUEUE_CLEANUP_TARGET)
        self.GUI_PROCESS_MAX_ITEMS = kwargs.get("gui_process_max_items", self.GUI_PROCESS_MAX_ITEMS)
        self.TEXT_PROCESS_MAX_ITEMS = kwargs.get("text_process_max_items", self.TEXT_PROCESS_MAX_ITEMS)
        self.GUI_PROCESS_MAX_DURATION = kwargs.get("gui_process_max_duration", self.GUI_PROCESS_MAX_DURATION)
        self.TEXT_PROCESS_MAX_DURATION = kwargs.get("text_process_max_duration", self.TEXT_PROCESS_MAX_DURATION)

        self._gui_queue_lock = threading.RLock()
        self._text_queue_lock = threading.RLock()

        self.gui_queue = getattr(gui, "gui_queue", None)
        self.text_queue = getattr(gui, "_text_update_queue", None)

        if self.gui_queue is None:
            logger.warning("QueueManager: gui_queue fehlt – verwende DummyQueue")
            self.gui_queue = DummyQueue(maxsize=self.GUI_QUEUE_MAX_SIZE)
        if self.text_queue is None:
            logger.warning("QueueManager: text_queue fehlt – verwende DummyQueue")
            self.text_queue = DummyQueue(maxsize=self.TEXT_QUEUE_MAX_SIZE)

    def start(self) -> None:
        self.root.after(50, self._process_gui_queue_dynamic)
        self.root.after(75, self._process_text_updates)
        self.root.after(5000, self._check_queue_sizes)

    def safe_put(self, queue_type: Literal["gui", "text"], item: Any) -> bool:
        if queue_type == "gui":
            q = self.gui_queue
            lock = self._gui_queue_lock
        elif queue_type == "text":
            q = self.text_queue
            lock = self._text_queue_lock
        else:
            logger.error(f"safe_put: unbekannter queue_type '{queue_type}'")
            return False

        with lock:
            try:
                q.put_nowait(item)
                return True
            except queue.Full:
                try:
                    q.get_nowait()
                    q.put_nowait(item)
                    if logger.isEnabledFor(logging.DEBUG):
                        log_debug("queue", f"Queue {queue_type} war voll – ein Element verworfen")
                    return True
                except (queue.Empty, queue.Full):
                    return False
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"safe_put error ({queue_type}): {e}")
                return False

    def add_important_type(self, msg_type: str) -> None:
        self.IMPORTANT_MSG_TYPES.add(msg_type)

    def get_queue_sizes(self) -> Tuple[int, int]:
        gui_size = self.gui_queue.qsize() if self.gui_queue else 0
        text_size = self.text_queue.qsize() if self.text_queue else 0
        return gui_size, text_size

    def _process_gui_queue_dynamic(self) -> None:
        if self.gui._shutting_down or not self.root.winfo_exists():
            return
        if self.gui_queue is None:
            return

        now = time.time()
        diff = now - self._last_gui_update_time
        if diff > 2.0 and self._last_gui_update_time > 0:
            if logger.isEnabledFor(logging.DEBUG):
                log_debug("gui", f"Letztes Update vor {diff:.1f}s – möglicher Freeze!")
        self._last_gui_update_time = now

        start_time = time.time()
        processed = 0
        max_items = self.GUI_PROCESS_MAX_ITEMS
        max_duration = self.GUI_PROCESS_MAX_DURATION
        queue_size = self.gui_queue.qsize()

        try:
            while processed < max_items and (time.time() - start_time) < max_duration:
                try:
                    item = self.gui_queue.get_nowait()
                except queue.Empty:
                    break

                if isinstance(item, tuple) and len(item) == 2:
                    msg_type, callback = item
                    if callable(callback):
                        if self._gui_update_limiter.can_update(f"gui_{msg_type}"):
                            try:
                                callback()
                            except (tk.TclError, RuntimeError) as e:
                                if logger.isEnabledFor(logging.DEBUG):
                                    log_debug("gui", f"GUI callback error ({msg_type}): {e}")
                            except Exception as e:
                                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                    raise
                                logger.warning(f"⚠️ GUI callback error ({msg_type}): {e}")
                        else:
                            if logger.isEnabledFor(logging.DEBUG):
                                log_debug("gui", f"Rate-Limit für {msg_type} – Element verworfen")
                else:
                    logger.warning(f"⚠️ Unbekanntes Element in GUI-Queue entfernt: {type(item)}")
                processed += 1

            if queue_size > self.GUI_QUEUE_MAX_SIZE:
                if logger.isEnabledFor(logging.DEBUG) or queue_size > self.GUI_QUEUE_MAX_SIZE * 2:
                    log_debug("queue", f"GUI queue zu groß ({queue_size}) – bereinige")
                self._cleanup_queue(self.gui_queue, self.GUI_QUEUE_CLEANUP_TARGET, self._gui_queue_lock)

            if logger.isEnabledFor(logging.DEBUG) and processed > 0:
                log_debug("gui", f"Verarbeitet: {processed} Items, Queue-Größe: {self.gui_queue.qsize()}")

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Kritischer Fehler in _process_gui_queue_dynamic: {e}", exc_info=True)

        if not self.gui._shutting_down:
            try:
                if self.root.winfo_exists():
                    next_interval = max(10, min(100, 100 - queue_size)) if queue_size > 0 else 100
                    self.root.after(next_interval, self._process_gui_queue_dynamic)
            except tk.TclError:
                pass

    def _process_text_updates(self) -> None:
        if self.gui._shutting_down or not self.root.winfo_exists():
            return
        if self.text_queue is None:
            return

        start_time = time.time()
        processed = 0
        max_items = self.TEXT_PROCESS_MAX_ITEMS
        max_duration = self.TEXT_PROCESS_MAX_DURATION
        queue_size_before = self.text_queue.qsize()

        if logger.isEnabledFor(logging.DEBUG) and queue_size_before > 0:
            log_debug("queue", f"_process_text_updates: Queue-Größe vor Verarbeitung: {queue_size_before}")

        try:
            while processed < max_items and (time.time() - start_time) < max_duration:
                try:
                    update_type, text_data = self.text_queue.get_nowait()
                except queue.Empty:
                    break

                try:
                    if update_type == "transcript" and hasattr(self.gui, "transcript_text"):
                        if self.gui.transcript_text.winfo_exists():
                            self.gui.transcript_text.insert("end", text_data)
                            if hasattr(self.gui, "transcript_scroll_var") and self.gui.transcript_scroll_var.get():
                                self.gui.transcript_text.see("end")
                    elif update_type == "translation" and hasattr(self.gui, "translation_text"):
                        if self.gui.translation_text.winfo_exists():
                            self.gui.translation_text.insert("end", text_data)
                            if hasattr(self.gui, "translation_scroll_var") and self.gui.translation_scroll_var.get():
                                self.gui.translation_text.see("end")
                except tk.TclError:
                    pass
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Text update error: {e}")

                processed += 1

            if self.text_queue.qsize() > self.TEXT_QUEUE_MAX_SIZE:
                self._cleanup_queue(self.text_queue, self.TEXT_QUEUE_CLEANUP_TARGET, self._text_queue_lock)
                log_debug("queue", f"Text queue cleaned to {self.TEXT_QUEUE_CLEANUP_TARGET} items (was {self.text_queue.qsize()})")

            if logger.isEnabledFor(logging.DEBUG) and processed > 0:
                log_debug("queue", f"Verarbeitet: {processed} Elemente in text queue, verbleibende Größe: {self.text_queue.qsize()}")

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Text update processor error: {e}", exc_info=True)

        if not self.gui._shutting_down and self.root.winfo_exists():
            self.root.after(150, self._process_text_updates)

    def _check_queue_sizes(self) -> None:
        if self.gui._shutting_down or not self.root.winfo_exists():
            return

        try:
            if self.gui_queue and self.gui_queue.qsize() > self.GUI_QUEUE_MAX_SIZE:
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("queue", f"GUI queue size {self.gui_queue.qsize()} exceeds threshold, cleaning up")
                self._cleanup_queue(self.gui_queue, self.GUI_QUEUE_CLEANUP_TARGET, self._gui_queue_lock)
                logger.info(f"🧹 Aggressive GUI queue cleanup: reduced to {self.GUI_QUEUE_CLEANUP_TARGET} items")

            if self.text_queue and self.text_queue.qsize() > self.TEXT_QUEUE_MAX_SIZE:
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("queue", f"Text queue size {self.text_queue.qsize()} exceeds threshold, cleaning up")
                self._cleanup_queue(self.text_queue, self.TEXT_QUEUE_CLEANUP_TARGET, self._text_queue_lock)
                logger.info(f"🧹 Aggressive text queue cleanup: reduced to {self.TEXT_QUEUE_CLEANUP_TARGET} items")

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Queue size check error: {e}")

        if not self.gui._shutting_down and self.root.winfo_exists():
            self.root.after(10000, self._check_queue_sizes)

    def _cleanup_queue(self, queue_obj: queue.Queue, max_size: int, lock: threading.RLock) -> None:
        if not queue_obj or queue_obj.qsize() <= max_size:
            return

        with lock:
            try:
                items = []
                while not queue_obj.empty():
                    try:
                        items.append(queue_obj.get_nowait())
                    except queue.Empty:
                        break

                important = []
                others = []
                for item in items:
                    if (isinstance(item, tuple) and len(item) == 2 and
                            item[0] in self.IMPORTANT_MSG_TYPES):
                        important.append(item)
                    else:
                        others.append(item)

                if len(important) < max_size:
                    kept = important + others[-(max_size - len(important)):]
                else:
                    kept = important[:max_size]

                for item in kept:
                    try:
                        queue_obj.put_nowait(item)
                    except queue.Full:
                        break

                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("queue", f"Cleaned queue from {len(items)} to {len(kept)} items")

            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Queue cleanup error: {e}")


# -----------------------------------------------------------------------------
# StatusBar – eigenständige Komponente für die untere Statusleiste
# -----------------------------------------------------------------------------
class StatusBar:
    def __init__(self, parent: tk.Frame, gui: "DragonWhispererGUI"):
        self.gui = gui
        self.root = gui.root
        self.frame = tk.Frame(parent, bg=gui.current_theme.BG_SECONDARY, height=50)
        self.frame.grid_propagate(True)
        separator = tk.Frame(self.frame, height=2, bg=gui.current_theme.DRAGON_GREEN)
        separator.pack(fill="x", side="top")
        main_container = tk.Frame(self.frame, bg=gui.current_theme.BG_SECONDARY)
        main_container.pack(fill="x", expand=True, padx=12, pady=8)
        main_container.columnconfigure(0, weight=0)
        main_container.columnconfigure(1, weight=1)
        main_container.columnconfigure(2, weight=0)

        left_panel = tk.Frame(main_container, bg=gui.current_theme.BG_SECONDARY)
        left_panel.grid(row=0, column=0, sticky="w", padx=5)
        self._create_left_buttons(left_panel)

        center_panel = tk.Frame(main_container, bg=gui.current_theme.BG_SECONDARY)
        center_panel.grid(row=0, column=1, sticky="ew", padx=5)
        self._create_center_panel(center_panel)

        right_panel = tk.Frame(main_container, bg=gui.current_theme.BG_SECONDARY)
        right_panel.grid(row=0, column=2, sticky="e", padx=5)
        self._create_right_buttons(right_panel)

    def _create_left_buttons(self, parent: tk.Frame):
        quick_actions = [
            ("🗑️", self.gui.clear_all, "Alles löschen"),
            ("💾", self.gui.save_transcript, "Transkription speichern"),
            ("📝", self.gui.export_subtitles, "Untertitel exportieren"),
            ("📊", self.gui.show_simple_stats, "Statistiken anzeigen"),
            ("⚙️", self.gui.show_advanced_settings, "Erweiterte Einstellungen"),
            ("🌐", self.gui.show_translation_dialog, "Text übersetzen"),
            ("🤖", self.gui.show_summarize_dialog, "Mit Ollama zusammenfassen"),
        ]
        for i, (icon, command, tooltip) in enumerate(quick_actions):
            btn = tk.Button(
                parent,
                text=icon,
                command=command,
                bg=self.gui.current_theme.BG_TERTIARY,
                fg=self.gui.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4,
                pady=2,
                activebackground=self.gui.current_theme.BG_HOVER,
            )
            btn.grid(row=0, column=i, padx=1, sticky="w")
            ToolTip(btn, tooltip)
        if getattr(self.gui, "demo_mode", False) or not TRANSLATOR_AVAILABLE:
            install_btn = tk.Button(
                parent,
                text="📦",
                command=self.gui.show_install_dialog,
                bg=self.gui.current_theme.BG_TERTIARY,
                fg=self.gui.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4,
                pady=2,
                activebackground=self.gui.current_theme.BG_HOVER,
            )
            install_btn.grid(row=0, column=len(quick_actions) + 1, padx=1, sticky="w")
            ToolTip(install_btn, "Fehlende Pakete installieren")

    def _create_center_panel(self, parent: tk.Frame):
        self.gui.progress_bar = ttk.Progressbar(
            parent, mode="determinate", length=150, style="Dark.Horizontal.TProgressbar"
        )
        self.gui.progress_bar.pack(side="left", padx=(10, 10))
        self.gui.progress_label = tk.Label(
            parent,
            text="",
            font=("Segoe UI", 8),
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_SECONDARY,
        )
        self.gui.progress_label.pack(side="left", padx=(0, 10))

        if IS_WINDOWS:
            default_text = "🪟 Windows | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        elif IS_MACOS:
            if IS_ARM:
                default_text = "🍎 macOS (Apple Silicon) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
            else:
                default_text = (
                    "🍎 macOS (Intel) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
                )
        elif IS_LINUX:
            default_text = "🐧 Linux | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        else:
            default_text = "🌐 Unknown OS | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        self.gui.system_info_label = tk.Label(
            parent,
            text=default_text,
            font=("Segoe UI", 8, "normal"),
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_SECONDARY,
            padx=5,
        )
        self.gui.system_info_label.pack(side="left", fill="x", expand=True)

    def _create_right_buttons(self, parent: tk.Frame):
        self.gui.exit_button = tk.Button(
            parent,
            text=" ⏻ EXIT ",
            command=self.gui.controller.safe_exit,
            bg="#dc3545",
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief="raised",
            cursor="hand2",
            padx=12,
            pady=3,
            activebackground="#c82333",
        )
        self.gui.exit_button.pack(side="right")
        ToolTip(self.gui.exit_button, "Programm beenden (Strg+Q / Cmd+Q)")

        self.gui.correct_btn = tk.Button(
            parent,
            text="🔧",
            command=self.gui.correct_transcript_with_ollama,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        self.gui.correct_btn.pack(side="right", padx=2)
        ToolTip(self.gui.correct_btn, "Transkript mit Ollama korrigieren")

        help_btn = tk.Button(
            parent,
            text="⌨️",
            command=self.gui.show_shortcuts_help,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        help_btn.pack(side="right", padx=2)
        ToolTip(help_btn, "Tastenkürzel anzeigen (F1)")

        self.gui.tts_btn = tk.Button(
            parent,
            text="🔊",
            command=self.gui.speak_current_text,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        self.gui.tts_btn.pack(side="right", padx=2)
        ToolTip(self.gui.tts_btn, "Ausgewählten Text vorlesen (TTS)")

        # --- Änderung: VAD-Fallback-Button mit Zustandsspeicherung
        self.gui.vad_fallback_btn = tk.Button(
            parent,
            text="🔁 VAD-Fallback ON" if self.gui.vad_fallback_enabled.get() else "🔁 VAD-Fallback OFF",
            command=self.gui.toggle_vad_fallback,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 8),
            padx=4,
        )
        self.gui.vad_fallback_btn.pack(side="right", padx=2)
        ToolTip(self.gui.vad_fallback_btn, "VAD-Fallback aktivieren/deaktivieren (wenn aus, werden leere Chunks ignoriert)")

        # --- Änderung: Live-Mode-Umschalter (20s/10s)
        self.gui.live_mode_btn = tk.Button(
            parent,
            text="⏱️ 20s",
            command=self.gui.toggle_live_mode,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 8),
            padx=4,
        )
        self.gui.live_mode_btn.pack(side="right", padx=2)
        ToolTip(self.gui.live_mode_btn, "Chunk-Dauer umschalten (20s/10s)")


# -----------------------------------------------------------------------------
# AdvancedSettingsDialog
# -----------------------------------------------------------------------------
class AdvancedSettingsDialog:
    def __init__(self, parent: tk.Tk, gui: "DragonWhispererGUI"):
        self.parent = parent
        self.gui = gui
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Advanced Settings")
        self.dialog.geometry("900x750")
        self.dialog.configure(bg=gui.current_theme.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.dialog.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")

        if hasattr(self.gui, '_open_dialogs'):
            self.gui._open_dialogs.append(self.dialog)
            self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)

        self.profiles = {
            "Default": {
                "chunk_duration": 15.0,
                "vad_filter": False,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 225,
                "vad_min_silence_ms": 80,
                "beam_size": 10,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Automatisch",
                "min_confidence": 0.1,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 3,
                "adaptive_high_words": 12,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Deutsches Video (Präzision)": {
                "chunk_duration": 20.0,
                "vad_filter": True,
                "vad_threshold": 0.15,
                "vad_min_speech_ms": 500,
                "vad_min_silence_ms": 300,
                "beam_size": 15,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.2,
                "length_penalty": 0.0,
                "audio_profile": "noisy",
                "source_lang_name": "Deutsch",
                "min_confidence": 0.2,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 3,
                "adaptive_high_words": 15,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "hotwords": "",
                "blacklist_mode": "word",
                "tts_engine": "piper",
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Englisches Video (Präzision)": {
                "chunk_duration": 20.0,
                "vad_filter": False,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 200,
                "vad_min_silence_ms": 60,
                "beam_size": 12,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Englisch",
                "min_confidence": 0.1,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 3,
                "adaptive_high_words": 15,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Asiatisches Video (Präzision)": {
                "chunk_duration": 25.0,
                "vad_filter": False,
                "vad_threshold": 0.35,
                "vad_min_speech_ms": 300,
                "vad_min_silence_ms": 120,
                "beam_size": 15,
                "temperature": 0.0,
                "no_speech_threshold": 0.55,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.2,
                "length_penalty": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Japanisch",
                "min_confidence": 0.1,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 3,
                "adaptive_high_words": 18,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Deutscher Livestream (Optimiert)": {
                "chunk_duration": 8.0,
                "vad_filter": False,
                "vad_threshold": 0.3,
                "vad_min_speech_ms": 400,
                "vad_min_silence_ms": 100,
                "beam_size": 8,
                "temperature": 0.0,
                "no_speech_threshold": 0.65,
                "log_prob_threshold": -0.8,
                "compression_ratio_threshold": 2.8,
                "condition_on_previous_text": False,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "realtime",
                "source_lang_name": "Deutsch",
                "min_confidence": 0.2,
                "duplicate_threshold": 0.95,
                "adaptive_low_words": 2,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Englischer Livestream (Optimiert)": {
                "chunk_duration": 8.0,
                "vad_filter": False,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 200,
                "vad_min_silence_ms": 60,
                "beam_size": 8,
                "temperature": 0.0,
                "no_speech_threshold": 0.65,
                "log_prob_threshold": -0.8,
                "compression_ratio_threshold": 2.8,
                "condition_on_previous_text": False,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "realtime",
                "source_lang_name": "Englisch",
                "min_confidence": 0.2,
                "duplicate_threshold": 0.95,
                "adaptive_low_words": 2,
                "adaptive_high_words": 10,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Asiatischer Livestream (Optimiert)": {
                "chunk_duration": 10.0,
                "vad_filter": False,
                "vad_threshold": 0.35,
                "vad_min_speech_ms": 300,
                "vad_min_silence_ms": 120,
                "beam_size": 10,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.6,
                "condition_on_previous_text": False,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "realtime",
                "source_lang_name": "Japanisch",
                "min_confidence": 0.2,
                "duplicate_threshold": 0.95,
                "adaptive_low_words": 2,
                "adaptive_high_words": 12,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "Podcast / Interview": {
                "chunk_duration": 15.0,
                "vad_filter": False,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 225,
                "vad_min_silence_ms": 80,
                "beam_size": 10,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "podcast",
                "source_lang_name": "Automatisch",
                "min_confidence": 0.1,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 4,
                "adaptive_high_words": 14,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
            "International (gemischt)": {
                "chunk_duration": 12.0,
                "vad_filter": False,
                "vad_threshold": 0.25,
                "vad_min_speech_ms": 250,
                "vad_min_silence_ms": 80,
                "beam_size": 8,
                "temperature": 0.0,
                "no_speech_threshold": 0.6,
                "log_prob_threshold": -1.0,
                "compression_ratio_threshold": 2.4,
                "condition_on_previous_text": True,
                "patience": 1.0,
                "length_penalty": 0.0,
                "audio_profile": "transcription",
                "source_lang_name": "Automatisch",
                "min_confidence": 0.1,
                "duplicate_threshold": 0.98,
                "adaptive_low_words": 3,
                "adaptive_high_words": 12,
                "max_memory_mb": 1024,
                "auto_save_interval": 300,
                "optimize_translations": False,
                "sentiment": False,
                "diarize": False,
                "best_of": 5,
                "suppress_tokens": "-1",
            },
        }

        self.custom_profiles = self._load_custom_profiles()
        all_profiles = list(self.profiles.keys()) + list(self.custom_profiles.keys())
        self.profile_list = all_profiles

        self._create_widgets()
        self._bind_events()

    def _on_close(self):
        if hasattr(self.gui, '_open_dialogs') and self.dialog in self.gui._open_dialogs:
            self.gui._open_dialogs.remove(self.dialog)
        self.dialog.destroy()

    def _create_widgets(self):
        main_frame = tk.Frame(
            self.dialog, bg=self.gui.current_theme.BG_PRIMARY, padx=20, pady=20
        )
        main_frame.pack(fill="both", expand=True)

        canvas = tk.Canvas(
            main_frame, bg=self.gui.current_theme.BG_PRIMARY, highlightthickness=0
        )
        scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.gui.current_theme.BG_PRIMARY)

        scrollable_frame.bind(
            "<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.bind("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1 * (e.delta / 120)), "units"))
        canvas.bind("<Button-4>", lambda e: canvas.yview_scroll(-1, "units"))
        canvas.bind("<Button-5>", lambda e: canvas.yview_scroll(1, "units"))

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.settings_frame = scrollable_frame
        self._build_content()

    def _build_content(self):
        profile_frame = tk.LabelFrame(
            self.settings_frame,
            text="📋 Vordefinierte Profile",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        profile_frame.grid(row=0, column=0, sticky="ew", pady=5, padx=5)
        profile_frame.columnconfigure(1, weight=1)

        tk.Label(
            profile_frame,
            text="Profil:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", padx=5)

        self.profile_var = tk.StringVar()
        self.profile_combo = ttk.Combobox(
            profile_frame,
            textvariable=self.profile_var,
            values=self.profile_list,
            width=25,
            state="readonly",
            style="Dark.TCombobox",
        )
        self.profile_combo.grid(row=0, column=1, sticky="ew", padx=5)

        self.save_profile_btn = tk.Button(
            profile_frame,
            text="💾 Profil speichern",
            command=self.save_custom_profile,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
            padx=5,
        )
        self.save_profile_btn.grid(row=0, column=2, padx=5)

        audio_frame = tk.LabelFrame(
            self.settings_frame,
            text="🎵 Audio & VAD",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        audio_frame.grid(row=1, column=0, sticky="ew", pady=5, padx=5)
        audio_frame.columnconfigure(1, weight=1)
        audio_frame.columnconfigure(3, weight=1)

        tk.Label(
            audio_frame,
            text="Sample Rate (Hz):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        sr_label = tk.Label(
            audio_frame,
            text=str(self.gui.advanced_settings.config.SAMPLE_RATE),
            bg=self.gui.current_theme.BG_TERTIARY,
            relief="sunken",
            width=10,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        )
        sr_label.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(
            audio_frame,
            text="Chunk Duration (s):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1)
        self.chunk_var = tk.DoubleVar(
            value=self.gui.advanced_settings.chunk_duration
        )
        self.chunk_spin = tk.Spinbox(
            audio_frame,
            from_=2.0,
            to=30.0,
            increment=0.5,
            textvariable=self.chunk_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.chunk_spin.grid(row=0, column=3, sticky="w", pady=1)

        tk.Label(
            audio_frame,
            text="Channels:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=0, sticky="w", pady=1)
        ch_label = tk.Label(
            audio_frame,
            text=str(self.gui.advanced_settings.config.CHANNELS),
            bg=self.gui.current_theme.BG_TERTIARY,
            relief="sunken",
            width=10,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        )
        ch_label.grid(row=1, column=1, sticky="w", pady=1)

        tk.Label(
            audio_frame,
            text="Audio Profile:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=2, sticky="w", pady=1)
        self.profile_var_audio = tk.StringVar(
            value=self.gui.advanced_settings.audio_profile
        )
        self.profile_combo_audio = ttk.Combobox(
            audio_frame,
            textvariable=self.profile_var_audio,
            values=list(Constants.FILTER_PROFILES.keys()),
            width=12,
            state="readonly",
            style="Dark.TCombobox",
        )
        self.profile_combo_audio.grid(row=1, column=3, sticky="w", pady=1)

        self.vad_filter_var = tk.BooleanVar(value=self.gui.advanced_settings.vad_filter)
        self.vad_filter_cb = tk.Checkbutton(
            audio_frame,
            text="VAD Filter aktivieren",
            variable=self.vad_filter_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.vad_filter_cb.grid(row=2, column=0, columnspan=2, sticky="w", pady=1)

        tk.Label(
            audio_frame,
            text="VAD Threshold:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=3, column=0, sticky="w", pady=1)
        self.vad_threshold_var = tk.DoubleVar(
            value=self.gui.advanced_settings.vad_threshold
        )
        self.vad_scale = tk.Scale(
            audio_frame,
            from_=0.0,
            to=1.0,
            resolution=0.05,
            orient=tk.HORIZONTAL,
            variable=self.vad_threshold_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.vad_scale.grid(row=3, column=1, sticky="ew", pady=1)

        tk.Label(
            audio_frame,
            text="Min Speech (ms):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=3, column=2, sticky="w", pady=1)
        self.vad_min_speech_var = tk.IntVar(
            value=self.gui.advanced_settings.vad_min_speech_duration_ms
        )
        self.vad_speech_spin = tk.Spinbox(
            audio_frame,
            from_=0,
            to=2000,
            increment=50,
            textvariable=self.vad_min_speech_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.vad_speech_spin.grid(row=3, column=3, sticky="w", pady=1)

        tk.Label(
            audio_frame,
            text="Min Silence (ms):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=4, column=0, sticky="w", pady=1)
        self.vad_min_silence_var = tk.IntVar(
            value=self.gui.advanced_settings.vad_min_silence_duration_ms
        )
        self.vad_silence_spin = tk.Spinbox(
            audio_frame,
            from_=0,
            to=2000,
            increment=50,
            textvariable=self.vad_min_silence_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.vad_silence_spin.grid(row=4, column=1, sticky="w", pady=1)

        model_frame = tk.LabelFrame(
            self.settings_frame,
            text="🤖 Modell & Inferenz",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        model_frame.grid(row=2, column=0, sticky="ew", pady=5, padx=5)
        model_frame.columnconfigure(1, weight=1)
        model_frame.columnconfigure(3, weight=1)

        tk.Label(
            model_frame,
            text="Beam Size:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.beam_var = tk.IntVar(value=self.gui.advanced_settings.beam_size)
        self.beam_spin = tk.Spinbox(
            model_frame,
            from_=1,
            to=20,
            textvariable=self.beam_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.beam_spin.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(
            model_frame,
            text="Temperature:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1)
        self.temp_var = tk.DoubleVar(value=self.gui.advanced_settings.temperature)
        self.temp_scale = tk.Scale(
            model_frame,
            from_=0.0,
            to=2.0,
            resolution=0.1,
            orient=tk.HORIZONTAL,
            variable=self.temp_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.temp_scale.grid(row=0, column=3, sticky="ew", pady=1)

        self.gpu_var = tk.BooleanVar(value=self.gui.advanced_settings.gpu_acceleration)
        self.gpu_cb = tk.Checkbutton(
            model_frame,
            text="GPU Acceleration",
            variable=self.gpu_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.gpu_cb.grid(row=1, column=0, columnspan=2, sticky="w", pady=1)

        tk.Label(
            model_frame,
            text="Current Model:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=2, sticky="w", pady=1)
        current_model = (
            self.gui.transcription_engine.get_current_model()
            if hasattr(self.gui, "transcription_engine")
            else "unknown"
        )
        model_label = tk.Label(
            model_frame,
            text=current_model,
            bg=self.gui.current_theme.BG_TERTIARY,
            relief="sunken",
            width=15,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        )
        model_label.grid(row=1, column=3, sticky="w", pady=1)

        tk.Label(
            model_frame,
            text="Hotwords:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=2, column=0, sticky="w", pady=1)
        self.hotwords_var = tk.StringVar(value=getattr(self.gui.advanced_settings, "hotwords", ""))
        self.hotwords_entry = tk.Entry(
            model_frame,
            textvariable=self.hotwords_var,
            width=30,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui.current_theme.TEXT_PRIMARY,
        )
        self.hotwords_entry.grid(row=2, column=1, columnspan=3, sticky="ew", pady=1, padx=5)
        ToolTip(self.hotwords_entry, "Kommagetrennte Hotwords für faster-whisper")

        filter_frame = tk.LabelFrame(
            self.settings_frame,
            text="🔍 Transkriptions‑Filter",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        filter_frame.grid(row=3, column=0, sticky="ew", pady=5, padx=5)
        filter_frame.columnconfigure(1, weight=1)
        filter_frame.columnconfigure(3, weight=1)

        tk.Label(
            filter_frame,
            text="Min Confidence:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.min_conf_var = tk.DoubleVar(
            value=self.gui.advanced_settings.min_confidence
        )
        self.min_conf_scale = tk.Scale(
            filter_frame,
            from_=0.0,
            to=1.0,
            resolution=0.05,
            orient=tk.HORIZONTAL,
            variable=self.min_conf_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.min_conf_scale.grid(row=0, column=1, sticky="ew", pady=1)

        tk.Label(
            filter_frame,
            text="Duplicate Threshold:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1, padx=(10, 0))
        self.dup_thresh_var = tk.DoubleVar(
            value=self.gui.advanced_settings.duplicate_similarity_threshold
        )
        self.dup_thresh_scale = tk.Scale(
            filter_frame,
            from_=0.5,
            to=1.0,
            resolution=0.01,
            orient=tk.HORIZONTAL,
            variable=self.dup_thresh_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.dup_thresh_scale.grid(row=0, column=3, sticky="ew", pady=1)

        tk.Label(
            filter_frame,
            text="Adaptive Low Words:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=0, sticky="w", pady=1)
        self.low_words_var = tk.IntVar(
            value=self.gui.advanced_settings.adaptive_chunk_low_words
        )
        self.low_words_spin = tk.Spinbox(
            filter_frame,
            from_=1,
            to=20,
            textvariable=self.low_words_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.low_words_spin.grid(row=1, column=1, sticky="w", pady=1)

        tk.Label(
            filter_frame,
            text="Adaptive High Words:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=2, sticky="w", pady=1, padx=(10, 0))
        self.high_words_var = tk.IntVar(
            value=self.gui.advanced_settings.adaptive_chunk_high_words
        )
        self.high_words_spin = tk.Spinbox(
            filter_frame,
            from_=1,
            to=20,
            textvariable=self.high_words_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.high_words_spin.grid(row=1, column=3, sticky="w", pady=1)

        trans_frame = tk.LabelFrame(
            self.settings_frame,
            text="🌐 Übersetzung",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        trans_frame.grid(row=4, column=0, sticky="ew", pady=5, padx=5)
        trans_frame.columnconfigure(1, weight=1)
        trans_frame.columnconfigure(3, weight=1)

        tk.Label(
            trans_frame,
            text="Engine:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.engine_var = tk.StringVar(
            value=self.gui.advanced_settings.translation_engine
        )
        self.engine_combo = ttk.Combobox(
            trans_frame,
            textvariable=self.engine_var,
            values=["google", "ollama", "argos"],
            width=10,
            state="readonly",
            style="Dark.TCombobox",
        )
        self.engine_combo.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(
            trans_frame,
            text="Ollama Model:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1)
        self.ollama_model_var = tk.StringVar(
            value=self.gui.advanced_settings.ollama_model
        )
        self.ollama_model_entry = tk.Entry(
            trans_frame,
            textvariable=self.ollama_model_var,
            width=15,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        )
        self.ollama_model_entry.grid(row=0, column=3, sticky="w", pady=1)

        tk.Label(
            trans_frame,
            text="Ollama Host:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=0, sticky="w", pady=1)
        host_frame = tk.Frame(trans_frame, bg=self.gui.current_theme.BG_SECONDARY)
        host_frame.grid(row=1, column=1, columnspan=3, sticky="ew", pady=1)
        host_frame.columnconfigure(0, weight=1)

        self.ollama_host_var = tk.StringVar(
            value=self.gui.advanced_settings.ollama_host
        )
        self.ollama_host_entry = tk.Entry(
            host_frame,
            textvariable=self.ollama_host_var,
            width=30,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        )
        self.ollama_host_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        def test_ollama():
            host = self.ollama_host_var.get().strip()
            if not host:
                host = "http://localhost:11434"
            try:
                import requests
                r = requests.get(f"{host}/api/tags", timeout=3)
                if r.status_code == 200:
                    DarkMessageBox.showinfo("Success", "Ollama server is reachable!", self.parent)
                else:
                    DarkMessageBox.showerror("Error", f"Ollama returned status {r.status_code}", self.parent)
            except Exception as e:
                DarkMessageBox.showerror("Error", f"Connection failed: {e}", self.parent)

        self.test_btn = tk.Button(
            host_frame,
            text="Test",
            command=test_ollama,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            padx=5,
            font=("Segoe UI", 8),
        )
        self.test_btn.pack(side="right")

        gui_frame = tk.LabelFrame(
            self.settings_frame,
            text="🖥️ GUI & Display",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        gui_frame.grid(row=5, column=0, sticky="ew", pady=5, padx=5)
        gui_frame.columnconfigure(1, weight=1)
        gui_frame.columnconfigure(3, weight=1)

        tk.Label(
            gui_frame,
            text="Transcript Max Lines:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.trans_lines_var = tk.IntVar(
            value=self.gui.advanced_settings.transcript_max_lines
        )
        self.trans_lines_spin = tk.Spinbox(
            gui_frame,
            from_=100,
            to=5000,
            increment=100,
            textvariable=self.trans_lines_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.trans_lines_spin.grid(row=0, column=1, sticky="w", pady=1)

        tk.Label(
            gui_frame,
            text="Translation Max Lines:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1)
        self.transl_lines_var = tk.IntVar(
            value=self.gui.advanced_settings.translation_max_lines
        )
        self.transl_lines_spin = tk.Spinbox(
            gui_frame,
            from_=100,
            to=5000,
            increment=100,
            textvariable=self.transl_lines_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.transl_lines_spin.grid(row=0, column=3, sticky="w", pady=1)

        tk.Label(
            gui_frame,
            text="Theme:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=0, sticky="w", pady=1)
        self.theme_var = tk.StringVar(value=self.gui.settings.theme)
        self.theme_combo = ttk.Combobox(
            gui_frame,
            textvariable=self.theme_var,
            values=["dark", "light", "pastel", "system", "highcontrast"],
            width=10,
            state="readonly",
            style="Dark.TCombobox",
        )
        self.theme_combo.grid(row=1, column=1, sticky="w", pady=1)

        def on_theme_change(*args):
            new_theme = self.theme_var.get()
            if new_theme != self.gui.settings.theme:
                self.gui.settings.theme = new_theme
                self.gui._apply_theme(new_theme)

        self.theme_var.trace_add("write", on_theme_change)

        self.auto_save_var = tk.BooleanVar(
            value=self.gui.settings.auto_save_on_completion
        )
        self.auto_save_cb = tk.Checkbutton(
            gui_frame,
            text="Auto-Save on Completion",
            variable=self.auto_save_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.auto_save_cb.grid(row=1, column=2, columnspan=2, sticky="w", pady=1)

        adv_frame = tk.LabelFrame(
            self.settings_frame,
            text="⚙️ Erweitert & System",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        adv_frame.grid(row=6, column=0, sticky="ew", pady=5, padx=5)
        adv_frame.columnconfigure(1, weight=1)
        adv_frame.columnconfigure(3, weight=1)

        tk.Label(
            adv_frame,
            text="Max Cache Size (MB):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.cache_var = tk.IntVar(value=self.gui.advanced_settings.max_cache_size)
        self.cache_spin = tk.Spinbox(
            adv_frame,
            from_=10,
            to=1000,
            increment=10,
            textvariable=self.cache_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.cache_spin.grid(row=0, column=1, sticky="w", pady=1)

        self.plugin_var = tk.BooleanVar(value=self.gui.settings.enable_plugins)
        self.plugin_cb = tk.Checkbutton(
            adv_frame,
            text="Enable Plugins",
            variable=self.plugin_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.plugin_cb.grid(row=0, column=2, columnspan=2, sticky="w", pady=1)

        self.cookies_var = tk.BooleanVar(value=self.gui.settings.use_browser_cookies)
        self.cookies_cb = tk.Checkbutton(
            adv_frame,
            text="Use Browser Cookies for YouTube",
            variable=self.cookies_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.cookies_cb.grid(row=1, column=0, columnspan=2, sticky="w", pady=1)

        self.asian_var = tk.BooleanVar(value=self.gui.advanced_settings.asian_mode)
        self.asian_cb = tk.Checkbutton(
            adv_frame,
            text="Asian Mode (10s chunks)",
            variable=self.asian_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.asian_cb.grid(row=1, column=2, columnspan=2, sticky="w", pady=1)

        self.precision_var = tk.BooleanVar(
            value=self.gui.advanced_settings.precision_mode
        )
        self.precision_cb = tk.Checkbutton(
            adv_frame,
            text="Precision Mode (langsamer, genauer)",
            variable=self.precision_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.precision_cb.grid(row=2, column=0, columnspan=2, sticky="w", pady=1)

        tk.Label(
            adv_frame,
            text="Max Memory (MB):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=3, column=0, sticky="w", pady=1)
        self.max_mem_var = tk.IntVar(value=self.gui.advanced_settings.max_memory_mb)
        self.max_mem_spin = tk.Spinbox(
            adv_frame,
            from_=100,
            to=16384,
            increment=100,
            textvariable=self.max_mem_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.max_mem_spin.grid(row=3, column=1, sticky="w", pady=1)

        tk.Label(
            adv_frame,
            text="Auto Save Interval (s):",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=3, column=2, sticky="w", pady=1, padx=(10, 0))
        self.auto_save_interval_var = tk.IntVar(
            value=self.gui.advanced_settings.auto_save_interval
        )
        self.auto_save_interval_spin = tk.Spinbox(
            adv_frame,
            from_=0,
            to=3600,
            increment=60,
            textvariable=self.auto_save_interval_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.auto_save_interval_spin.grid(row=3, column=3, sticky="w", pady=1)

        self.optimize_var = tk.BooleanVar(
            value=self.gui.advanced_settings.optimize_translations
        )
        self.optimize_cb = tk.Checkbutton(
            adv_frame,
            text="Optimize Translations",
            variable=self.optimize_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.optimize_cb.grid(row=4, column=0, columnspan=2, sticky="w", pady=1)

        self.sentiment_var = tk.BooleanVar(
            value=self.gui.advanced_settings.enable_sentiment_analysis
        )
        self.sentiment_cb = tk.Checkbutton(
            adv_frame,
            text="Sentiment Analysis",
            variable=self.sentiment_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.sentiment_cb.grid(row=4, column=2, columnspan=2, sticky="w", pady=1)

        self.diarize_var = tk.BooleanVar(
            value=self.gui.advanced_settings.enable_speaker_diarization
        )
        self.diarize_cb = tk.Checkbutton(
            adv_frame,
            text="Speaker Diarization",
            variable=self.diarize_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.diarize_cb.grid(row=5, column=0, columnspan=2, sticky="w", pady=1)

        blacklist_frame = tk.LabelFrame(
            self.settings_frame,
            text="🚫 Blacklist (Phrasen entfernen)",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        blacklist_frame.grid(row=7, column=0, sticky="ew", pady=5, padx=5)
        blacklist_frame.columnconfigure(0, weight=1)

        mode_frame = tk.Frame(blacklist_frame, bg=self.gui.current_theme.BG_SECONDARY)
        mode_frame.grid(row=0, column=0, sticky="ew", pady=2)
        tk.Label(
            mode_frame,
            text="Modus:",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).pack(side="left")
        self.blacklist_mode_var = tk.StringVar(value=getattr(self.gui.advanced_settings, "blacklist_mode", "word"))
        mode_combo = ttk.Combobox(
            mode_frame,
            textvariable=self.blacklist_mode_var,
            values=["word", "substring"],
            width=10,
            state="readonly",
            style="Dark.TCombobox",
        )
        mode_combo.pack(side="left", padx=5)
        ToolTip(mode_combo, "word = ganze Wörter, substring = beliebige Teilzeichenkette")

        tk.Label(
            blacklist_frame,
            text="Eine Phrase pro Zeile – wird nach der Transkription entfernt.",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
            justify="left",
        ).grid(row=1, column=0, sticky="w", pady=(0, 2))

        self.blacklist_text = scrolledtext.ScrolledText(
            blacklist_frame,
            height=5,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
        )
        self.blacklist_text.grid(row=2, column=0, sticky="ew", pady=2)

        blacklist = getattr(self.gui.advanced_settings, "blacklist", [])
        self.blacklist_text.insert("1.0", "\n".join(blacklist))

        tts_frame = tk.LabelFrame(
            self.settings_frame,
            text="🔊 Text-to-Speech",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        tts_frame.grid(row=8, column=0, sticky="ew", pady=5, padx=5)
        tts_frame.columnconfigure(1, weight=1)

        tk.Label(
            tts_frame,
            text="TTS Engine:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.tts_engine_var = tk.StringVar(
            value=self.gui.advanced_settings.tts_engine
        )
        self.tts_engine_combo = ttk.Combobox(
            tts_frame,
            textvariable=self.tts_engine_var,
            values=["piper", "pyttsx3"],
            width=15,
            state="readonly",
            style="Dark.TCombobox",
        )
        self.tts_engine_combo.grid(row=0, column=1, sticky="w", pady=1)
        ToolTip(self.tts_engine_combo, "Text-to-Speech Engine (piper empfohlen)")

        whisper_frame = tk.LabelFrame(
            self.settings_frame,
            text="🔧 Erweiterte Whisper-Parameter",
            padx=10,
            pady=8,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8, "bold"),
        )
        whisper_frame.grid(row=9, column=0, sticky="ew", pady=5, padx=5)
        whisper_frame.columnconfigure(1, weight=1)
        whisper_frame.columnconfigure(3, weight=1)

        tk.Label(
            whisper_frame,
            text="Best of:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=0, sticky="w", pady=1)
        self.best_of_var = tk.IntVar(value=self.gui.advanced_settings.best_of)
        self.best_of_spin = tk.Spinbox(
            whisper_frame,
            from_=1,
            to=20,
            textvariable=self.best_of_var,
            width=8,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            buttonbackground=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.best_of_spin.grid(row=0, column=1, sticky="w", pady=1)
        ToolTip(self.best_of_spin, "Anzahl der Suchpfade (größer = besser, aber langsamer)")

        tk.Label(
            whisper_frame,
            text="Patience:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=0, column=2, sticky="w", pady=1, padx=(10,0))
        self.patience_var = tk.DoubleVar(value=self.gui.advanced_settings.patience)
        self.patience_scale = tk.Scale(
            whisper_frame,
            from_=0.0,
            to=2.0,
            resolution=0.1,
            orient=tk.HORIZONTAL,
            variable=self.patience_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.patience_scale.grid(row=0, column=3, sticky="ew", pady=1)
        ToolTip(self.patience_scale, "Geduld bei der Beam-Suche (höher = genauer, aber langsamer)")

        tk.Label(
            whisper_frame,
            text="No Speech Threshold:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=0, sticky="w", pady=1)
        self.no_speech_var = tk.DoubleVar(value=self.gui.advanced_settings.no_speech_threshold)
        self.no_speech_scale = tk.Scale(
            whisper_frame,
            from_=0.0,
            to=1.0,
            resolution=0.05,
            orient=tk.HORIZONTAL,
            variable=self.no_speech_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.no_speech_scale.grid(row=1, column=1, sticky="ew", pady=1)
        ToolTip(self.no_speech_scale, "Schwellwert für ‚Keine Sprache‘ (niedriger = mehr Segmente)")

        tk.Label(
            whisper_frame,
            text="Log Prob Threshold:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=1, column=2, sticky="w", pady=1, padx=(10,0))
        self.log_prob_var = tk.DoubleVar(value=self.gui.advanced_settings.log_prob_threshold)
        self.log_prob_scale = tk.Scale(
            whisper_frame,
            from_=-5.0,
            to=0.0,
            resolution=0.1,
            orient=tk.HORIZONTAL,
            variable=self.log_prob_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.log_prob_scale.grid(row=1, column=3, sticky="ew", pady=1)
        ToolTip(self.log_prob_scale, "Log‑Wahrscheinlichkeits‑Schwelle (höher = weniger Segmente)")

        tk.Label(
            whisper_frame,
            text="Compression Ratio:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=2, column=0, sticky="w", pady=1)
        self.comp_ratio_var = tk.DoubleVar(value=self.gui.advanced_settings.compression_ratio_threshold)
        self.comp_ratio_scale = tk.Scale(
            whisper_frame,
            from_=1.0,
            to=5.0,
            resolution=0.1,
            orient=tk.HORIZONTAL,
            variable=self.comp_ratio_var,
            length=150,
            showvalue=True,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            troughcolor=self.gui.current_theme.BG_TERTIARY,
            font=("Segoe UI", 8),
        )
        self.comp_ratio_scale.grid(row=2, column=1, sticky="ew", pady=1)
        ToolTip(self.comp_ratio_scale, "Maximales Kompressionsverhältnis (höher = mehr Segmente)")

        self.condition_prev_var = tk.BooleanVar(value=self.gui.advanced_settings.condition_on_previous_text)
        self.condition_prev_cb = tk.Checkbutton(
            whisper_frame,
            text="Condition on previous text",
            variable=self.condition_prev_var,
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            selectcolor=self.gui.current_theme.BG_TERTIARY,
            activebackground=self.gui.current_theme.BG_SECONDARY,
            font=("Segoe UI", 8),
        )
        self.condition_prev_cb.grid(row=3, column=0, columnspan=2, sticky="w", pady=1)
        ToolTip(self.condition_prev_cb, "Vorherigen Text als Kontext verwenden (ausschalten reduziert Wiederholungen)")

        tk.Label(
            whisper_frame,
            text="Suppress Tokens:",
            anchor="w",
            bg=self.gui.current_theme.BG_SECONDARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 8),
        ).grid(row=3, column=2, sticky="w", pady=1, padx=(10,0))
        self.suppress_tokens_var = tk.StringVar(value=self.gui.advanced_settings.suppress_tokens)
        self.suppress_tokens_entry = tk.Entry(
            whisper_frame,
            textvariable=self.suppress_tokens_var,
            width=15,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui.current_theme.TEXT_PRIMARY,
        )
        self.suppress_tokens_entry.grid(row=3, column=3, sticky="w", pady=1)
        ToolTip(self.suppress_tokens_entry, "Komma-getrennte Token-IDs, die unterdrückt werden (z.B. '-1,0,1')")

        self.help_label = tk.Label(
            self.settings_frame,
            text="Bewegen Sie die Maus über eine Einstellung für Details.",
            bg=self.gui.current_theme.BG_PRIMARY,
            fg=self.gui.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
            anchor="w",
            justify="left",
        )
        self.help_label.grid(row=10, column=0, sticky="ew", pady=(10, 0))

        def show_help(text):
            self.help_label.config(text=text)

        self.best_of_spin.bind(
            "<Enter>",
            lambda e: show_help(
                "Anzahl der Suchpfade (größer = besser, aber langsamer)"
            ),
        )
        self.patience_scale.bind(
            "<Enter>",
            lambda e: show_help(
                "Geduld bei der Beam-Suche (höher = genauer, aber langsamer)"
            ),
        )
        self.no_speech_scale.bind(
            "<Enter>",
            lambda e: show_help(
                "Schwellwert für ‚Keine Sprache‘ (niedriger = mehr Segmente)"
            ),
        )
        self.log_prob_scale.bind(
            "<Enter>",
            lambda e: show_help(
                "Log‑Wahrscheinlichkeits‑Schwelle (höher = weniger Segmente)"
            ),
        )
        self.comp_ratio_scale.bind(
            "<Enter>",
            lambda e: show_help(
                "Maximales Kompressionsverhältnis (höher = mehr Segmente)"
            ),
        )
        self.condition_prev_cb.bind(
            "<Enter>",
            lambda e: show_help(
                "Vorherigen Text als Kontext verwenden (ausschalten reduziert Wiederholungen)"
            ),
        )
        self.suppress_tokens_entry.bind(
            "<Enter>",
            lambda e: show_help(
                "Komma-getrennte Token-IDs, die unterdrückt werden (z.B. '-1,0,1')"
            ),
        )

        button_frame = tk.Frame(
            self.settings_frame, bg=self.gui.current_theme.BG_PRIMARY
        )
        button_frame.grid(row=11, column=0, pady=20)

        reset_btn = tk.Button(
            button_frame,
            text="Reset to Defaults",
            command=self.reset_to_defaults,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            padx=15,
            font=("Segoe UI", 8),
        )
        reset_btn.pack(side="left", padx=5)

        save_btn = tk.Button(
            button_frame,
            text="Save",
            command=self.save_settings,
            bg=self.gui.current_theme.SUCCESS,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            padx=15,
            font=("Segoe UI", 8, "bold"),
        )
        save_btn.pack(side="left", padx=5)

        cancel_btn = tk.Button(
            button_frame,
            text="Cancel",
            command=self.dialog.destroy,
            bg=self.gui.current_theme.BG_TERTIARY,
            fg=self.gui.current_theme.TEXT_PRIMARY,
            relief="flat",
            padx=15,
            font=("Segoe UI", 8),
        )
        cancel_btn.pack(side="left", padx=5)

        self.settings_frame.columnconfigure(0, weight=1)

    def _load_custom_profiles(self) -> Dict[str, Dict[str, Any]]:
        profiles_dir = PlatformUtils.get_platform_config_dir() / "profiles"
        profiles = {}
        if not profiles_dir.exists():
            return profiles
        for file in profiles_dir.glob("*.json"):
            try:
                with open(file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                profiles[file.stem] = data
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"Fehler beim Laden des Profils {file.name}: {e}")
        return profiles

    def save_custom_profile(self):
        from tkinter.simpledialog import askstring
        name = askstring("Profil speichern", "Name des Profils:", parent=self.dialog)
        if not name:
            return
        profile_data = {
            "chunk_duration": self.chunk_var.get(),
            "vad_filter": self.vad_filter_var.get(),
            "vad_threshold": self.vad_threshold_var.get(),
            "vad_min_speech_ms": self.vad_min_speech_var.get(),
            "vad_min_silence_ms": self.vad_min_silence_var.get(),
            "beam_size": self.beam_var.get(),
            "temperature": self.temp_var.get(),
            "no_speech_threshold": self.no_speech_var.get(),
            "log_prob_threshold": self.log_prob_var.get(),
            "compression_ratio_threshold": self.comp_ratio_var.get(),
            "condition_on_previous_text": self.condition_prev_var.get(),
            "patience": self.patience_var.get(),
            "length_penalty": self.length_penalty_var.get(),
            "audio_profile": self.profile_var_audio.get(),
            "min_confidence": self.min_conf_var.get(),
            "duplicate_threshold": self.dup_thresh_var.get(),
            "adaptive_low_words": self.low_words_var.get(),
            "adaptive_high_words": self.high_words_var.get(),
            "max_memory_mb": self.max_mem_var.get(),
            "auto_save_interval": self.auto_save_interval_var.get(),
            "optimize_translations": self.optimize_var.get(),
            "sentiment": self.sentiment_var.get(),
            "diarize": self.diarize_var.get(),
            "hotwords": self.hotwords_var.get(),
            "blacklist_mode": self.blacklist_mode_var.get(),
            "tts_engine": self.tts_engine_var.get(),
            "best_of": self.best_of_var.get(),
            "suppress_tokens": self.suppress_tokens_var.get(),
        }
        blacklist_text = self.blacklist_text.get("1.0", "end-1c").strip()
        profile_data["blacklist"] = [line.strip() for line in blacklist_text.split("\n") if line.strip()]

        profiles_dir = PlatformUtils.get_platform_config_dir() / "profiles"
        profiles_dir.mkdir(parents=True, exist_ok=True)
        file_path = profiles_dir / f"{name}.json"
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(profile_data, f, indent=2)
            self.custom_profiles[name] = profile_data
            self.profile_list.append(name)
            self.profile_combo['values'] = self.profile_list
            DarkMessageBox.showinfo("Erfolg", f"Profil '{name}' gespeichert.", self.parent)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            DarkMessageBox.showerror("Fehler", f"Speichern fehlgeschlagen: {e}", self.parent)

    def reset_to_defaults(self):
        default = AdvancedSettings()
        self.chunk_var.set(default.chunk_duration)
        self.profile_var_audio.set("transcription")
        self.vad_filter_var.set(default.vad_filter)
        self.vad_threshold_var.set(default.vad_threshold)
        self.vad_min_speech_var.set(default.vad_min_speech_duration_ms)
        self.vad_min_silence_var.set(default.vad_min_silence_duration_ms)
        self.beam_var.set(default.beam_size)
        self.temp_var.set(default.temperature)
        self.gpu_var.set(default.gpu_acceleration)
        self.engine_var.set(default.translation_engine)
        self.ollama_model_var.set(default.ollama_model)
        self.ollama_host_var.set(default.ollama_host)
        self.trans_lines_var.set(default.transcript_max_lines)
        self.transl_lines_var.set(default.translation_max_lines)
        self.theme_var.set("dark")
        self.auto_save_var.set(False)
        self.cache_var.set(default.max_cache_size)
        self.plugin_var.set(True)
        self.cookies_var.set(True)
        self.asian_var.set(default.asian_mode)
        self.precision_var.set(default.precision_mode)
        self.gui.src_lang_var.set("Automatisch")
        self.min_conf_var.set(default.min_confidence)
        self.dup_thresh_var.set(default.duplicate_similarity_threshold)
        self.low_words_var.set(default.adaptive_chunk_low_words)
        self.high_words_var.set(default.adaptive_chunk_high_words)
        self.max_mem_var.set(default.max_memory_mb)
        self.auto_save_interval_var.set(default.auto_save_interval)
        self.optimize_var.set(default.optimize_translations)
        self.sentiment_var.set(default.enable_sentiment_analysis)
        self.diarize_var.set(default.enable_speaker_diarization)
        self.no_speech_var.set(default.no_speech_threshold)
        self.log_prob_var.set(default.log_prob_threshold)
        self.comp_ratio_var.set(default.compression_ratio_threshold)
        self.patience_var.set(default.patience)
        self.condition_prev_var.set(default.condition_on_previous_text)
        self.hotwords_var.set("")
        self.blacklist_mode_var.set("word")
        self.blacklist_text.delete("1.0", "end")
        self.blacklist_text.insert("1.0", "\n".join(default.blacklist))
        self.tts_engine_var.set(default.tts_engine)
        self.best_of_var.set(default.best_of)
        self.patience_var.set(default.patience)
        self.no_speech_var.set(default.no_speech_threshold)
        self.log_prob_var.set(default.log_prob_threshold)
        self.comp_ratio_var.set(default.compression_ratio_threshold)
        self.condition_prev_var.set(default.condition_on_previous_text)
        self.suppress_tokens_var.set(default.suppress_tokens)
        # Länge-Penalty wurde entfernt, daher keine Zeile mehr

    def save_settings(self):
        try:
            self.gui.advanced_settings.chunk_duration = self.chunk_var.get()
            self.gui.advanced_settings.audio_profile = self.profile_var_audio.get()
            self.gui.advanced_settings.vad_filter = self.vad_filter_var.get()
            self.gui.advanced_settings.vad_threshold = self.vad_threshold_var.get()
            self.gui.advanced_settings.vad_min_speech_duration_ms = (
                self.vad_min_speech_var.get()
            )
            self.gui.advanced_settings.vad_min_silence_duration_ms = (
                self.vad_min_silence_var.get()
            )
            self.gui.advanced_settings.beam_size = self.beam_var.get()
            self.gui.advanced_settings.temperature = self.temp_var.get()
            self.gui.advanced_settings.gpu_acceleration = self.gpu_var.get()
            self.gui.advanced_settings.translation_engine = self.engine_var.get()
            self.gui.advanced_settings.ollama_model = (
                self.ollama_model_var.get().strip()
            )
            self.gui.advanced_settings.ollama_host = self.ollama_host_var.get().strip()
            self.gui.advanced_settings.transcript_max_lines = self.trans_lines_var.get()
            self.gui.advanced_settings.translation_max_lines = (
                self.transl_lines_var.get()
            )
            self.gui.advanced_settings.asian_mode = self.asian_var.get()
            self.gui.advanced_settings.precision_mode = self.precision_var.get()
            self.gui.advanced_settings.max_cache_size = self.cache_var.get()

            self.gui.settings.theme = self.theme_var.get()
            self.gui.settings.auto_save_on_completion = self.auto_save_var.get()
            self.gui.settings.enable_plugins = self.plugin_var.get()
            self.gui.settings.use_browser_cookies = self.cookies_var.get()

            self.gui.advanced_settings.min_confidence = self.min_conf_var.get()
            self.gui.advanced_settings.duplicate_similarity_threshold = (
                self.dup_thresh_var.get()
            )
            self.gui.advanced_settings.adaptive_chunk_low_words = (
                self.low_words_var.get()
            )
            self.gui.advanced_settings.adaptive_chunk_high_words = (
                self.high_words_var.get()
            )
            self.gui.advanced_settings.max_memory_mb = self.max_mem_var.get()
            self.gui.advanced_settings.auto_save_interval = (
                self.auto_save_interval_var.get()
            )
            self.gui.advanced_settings.optimize_translations = self.optimize_var.get()
            self.gui.advanced_settings.enable_sentiment_analysis = (
                self.sentiment_var.get()
            )
            self.gui.advanced_settings.enable_speaker_diarization = (
                self.diarize_var.get()
            )
            self.gui.advanced_settings.no_speech_threshold = self.no_speech_var.get()
            self.gui.advanced_settings.log_prob_threshold = self.log_prob_var.get()
            self.gui.advanced_settings.compression_ratio_threshold = self.comp_ratio_var.get()
            self.gui.advanced_settings.patience = self.patience_var.get()
            self.gui.advanced_settings.condition_on_previous_text = self.condition_prev_var.get()
            self.gui.advanced_settings.hotwords = self.hotwords_var.get().strip()
            self.gui.advanced_settings.blacklist_mode = self.blacklist_mode_var.get()
            self.gui.advanced_settings.tts_engine = self.tts_engine_var.get()

            self.gui.advanced_settings.best_of = self.best_of_var.get()
            self.gui.advanced_settings.patience = self.patience_var.get()
            self.gui.advanced_settings.no_speech_threshold = self.no_speech_var.get()
            self.gui.advanced_settings.log_prob_threshold = self.log_prob_var.get()
            self.gui.advanced_settings.compression_ratio_threshold = self.comp_ratio_var.get()
            self.gui.advanced_settings.condition_on_previous_text = self.condition_prev_var.get()
            self.gui.advanced_settings.suppress_tokens = self.suppress_tokens_var.get().strip()

            blacklist_text = self.blacklist_text.get("1.0", "end-1c").strip()
            blacklist = [line.strip() for line in blacklist_text.split("\n") if line.strip()]
            self.gui.advanced_settings.blacklist = blacklist

            host = self.gui.advanced_settings.ollama_host
            if host and not host.startswith(("http://", "https://")):
                self.gui.advanced_settings.ollama_host = "http://" + host

            self.gui.advanced_settings.save_to_file()
            self.gui.settings.save_to_file()

            if hasattr(self.gui, "stream_manager"):
                self.gui.stream_manager.use_browser_cookies = (
                    self.gui.settings.use_browser_cookies
                )
            if hasattr(self.gui, "stream_info_extractor"):
                self.gui.stream_info_extractor.use_browser_cookies = (
                    self.gui.settings.use_browser_cookies
                )

            if not self.gui.advanced_settings.gpu_acceleration:
                self.gui.transcription_engine.device = "cpu"
                self.gui.transcription_engine.compute_type = "int8"

            if hasattr(self.gui, 'update_translation_engine'):
                self.gui.update_translation_engine()

            if hasattr(self.gui, 'tts_manager'):
                self.gui.tts_manager.set_engine(self.gui.advanced_settings.tts_engine)

            self.dialog.destroy()
            self.gui.update_status("✅ Settings saved")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            DarkMessageBox.showerror("Error", f"Invalid settings: {e}", self.parent)

    def _bind_events(self):
        self.profile_combo.bind("<<ComboboxSelected>>", self.apply_profile)

    def apply_profile(self, event=None):
        profil = self.profiles.get(self.profile_var.get()) or self.custom_profiles.get(self.profile_var.get())
        if not profil:
            logger.warning(f"Profil {self.profile_var.get()} nicht gefunden")
            return
        self.chunk_var.set(profil["chunk_duration"])
        self.vad_filter_var.set(profil.get("vad_filter", True))
        self.vad_threshold_var.set(profil["vad_threshold"])
        self.vad_min_speech_var.set(profil["vad_min_speech_ms"])
        self.vad_min_silence_var.set(profil["vad_min_silence_ms"])
        self.beam_var.set(profil["beam_size"])
        self.temp_var.set(profil["temperature"])
        self.profile_var_audio.set(profil["audio_profile"])
        self.gui.src_lang_var.set(profil["source_lang_name"])
        self.min_conf_var.set(profil.get("min_confidence", 0.25))
        self.dup_thresh_var.set(profil.get("duplicate_threshold", 0.85))
        self.low_words_var.set(profil.get("adaptive_low_words", 3))
        self.high_words_var.set(profil.get("adaptive_high_words", 10))
        self.max_mem_var.set(profil.get("max_memory_mb", 1024))
        self.auto_save_interval_var.set(profil.get("auto_save_interval", 300))
        self.optimize_var.set(profil.get("optimize_translations", False))
        self.sentiment_var.set(profil.get("sentiment", False))
        self.diarize_var.set(profil.get("diarize", False))
        self.no_speech_var.set(profil.get("no_speech_threshold", 0.5))
        self.log_prob_var.set(profil.get("log_prob_threshold", -1.0))
        self.comp_ratio_var.set(profil.get("compression_ratio_threshold", 2.4))
        self.patience_var.set(profil.get("patience", 1.0))
        self.condition_prev_var.set(profil.get("condition_on_previous_text", True))
        self.hotwords_var.set(profil.get("hotwords", ""))
        self.blacklist_mode_var.set(profil.get("blacklist_mode", "word"))
        self.tts_engine_var.set(profil.get("tts_engine", "piper"))
        self.best_of_var.set(profil.get("best_of", 5))
        self.patience_var.set(profil.get("patience", 1.0))
        self.no_speech_var.set(profil.get("no_speech_threshold", 0.6))
        self.log_prob_var.set(profil.get("log_prob_threshold", -1.2))
        self.comp_ratio_var.set(profil.get("compression_ratio_threshold", 2.8))
        self.condition_prev_var.set(profil.get("condition_on_previous_text", True))
        self.suppress_tokens_var.set(profil.get("suppress_tokens", "-1"))
        blacklist = profil.get("blacklist", [])
        self.blacklist_text.delete("1.0", "end")
        self.blacklist_text.insert("1.0", "\n".join(blacklist))


# -----------------------------------------------------------------------------
# TTSManager
# -----------------------------------------------------------------------------
class TTSManager:
    def __init__(self, settings: "AdvancedSettings"):
        self.settings = settings
        self._engine = None
        self._engine_name = settings.tts_engine
        self._lock = threading.RLock()
        self._speaking_thread: Optional[threading.Thread] = None
        self._stop_requested = False
        self._piper_available = False
        self._pyttsx3_available = False

    def set_engine(self, engine_name: str) -> None:
        with self._lock:
            self._engine_name = engine_name
            self._engine = None

    def is_available(self) -> bool:
        if self._piper_available:
            return True
        if importlib.util.find_spec("dimits") is not None:
            self._piper_available = True
            return True

        if self._pyttsx3_available:
            return True
        if importlib.util.find_spec("pyttsx3") is not None:
            self._pyttsx3_available = True
            return True

        return False

    def _load_engine(self) -> bool:
        with self._lock:
            if self._engine is not None:
                return True

            if self._engine_name == "piper":
                try:
                    from dimits import Dimits
                    self._engine = Dimits("de_DE-thorsten-medium")
                    logger.info("Piper TTS geladen (de_DE-thorsten-medium)")
                    self._piper_available = True
                    return True
                except ImportError:
                    logger.warning("Piper (dimits) nicht installiert, versuche pyttsx3")
                    self._engine_name = "pyttsx3"
                except Exception as e:
                    logger.warning(f"Fehler beim Laden von Piper: {e}")

            if self._engine_name == "pyttsx3":
                try:
                    import pyttsx3
                    self._engine = pyttsx3.init()
                    self._engine.setProperty('rate', 150)
                    self._engine.setProperty('volume', 0.9)
                    logger.info("pyttsx3 TTS geladen")
                    self._pyttsx3_available = True
                    return True
                except ImportError:
                    logger.error("pyttsx3 nicht installiert – TTS nicht verfügbar")
                except Exception as e:
                    logger.error(f"Fehler beim Laden von pyttsx3: {e}")

            self._engine = None
            return False

    def speak(self, text: str, callback: Optional[Callable[[bool, str], None]] = None) -> None:
        if not text or not text.strip():
            if callback:
                callback(False, "Leerer Text")
            return

        if self._engine is None:
            if not self._load_engine():
                if callback:
                    callback(False, "Keine TTS-Engine verfügbar")
                return

        self.stop()

        def _speak_worker():
            self._stop_requested = False
            success = False
            message = ""
            try:
                if self._engine_name == "piper":
                    self._engine.text_2_speech(text, engine='aplay')
                    success = True
                else:
                    self._engine.say(text)
                    self._engine.runAndWait()
                    success = True
            except Exception as e:
                message = str(e)
                logger.error(f"Fehler bei TTS: {e}")
            finally:
                if callback:
                    callback(success, message)

        with self._lock:
            self._speaking_thread = threading.Thread(target=_speak_worker, daemon=True)
            self._speaking_thread.start()

    def stop(self) -> None:
        self._stop_requested = True
        with self._lock:
            if self._speaking_thread and self._speaking_thread.is_alive():
                pass


# -----------------------------------------------------------------------------
# StreamHandler
# -----------------------------------------------------------------------------
class StreamHandler:
    def __init__(
        self, audio_processor: "AudioProcessor", stream_manager: "StreamManager"
    ):

        self._ap_ref = weakref.ref(audio_processor)
        self.config = audio_processor.config
        self.stream_manager = stream_manager
        self._ffmpeg_manager_ref = weakref.ref(audio_processor.ffmpeg_manager)
        self._max_backoff = Constants.MAX_BACKOFF
        self._last_stderr_read = 0.0
        self._stderr_read_interval = 5.0

    def _get_ap(self):
        return self._ap_ref()

    def _get_ffmpeg(self):
        return self._ffmpeg_manager_ref()

    def run_loop(
        self,
        process: subprocess.Popen,
        audio_url: str,
        original_video_url: str,
        detected_language: Optional[str],
        transcription_callback: Callable,
        translation_callback: Callable,
        info_callback: Callable,
        error_callback: Callable,
        is_youtube: bool,
    ) -> None:
        ap = self._get_ap()
        if ap is None:
            logger.error(
                "StreamHandler: AudioProcessor nicht mehr verfügbar – breche ab."
            )
            return

        error_occurred = False
        normal_ending = False

        platform_id, platform_name = self.stream_manager.detect_platform(
            original_video_url
        )
        logger.info(f"🎯 StreamHandler loop started for {platform_name}")

        last_data_time = time.time()
        consecutive_timeouts = 0
        backoff = 1.0
        max_reconnects = Constants.MAX_STREAM_RECONNECTS
        reconnect_attempts = 0
        refresh_count = 0
        max_refresh_attempts = Constants.YOUTUBE_URL_REFRESH_MAX_ATTEMPTS
        current_process = process
        last_url_refresh = time.time()
        url_refresh_interval = Constants.YOUTUBE_URL_REFRESH_INTERVAL
        consecutive_low_quality_chunks = 0
        max_low_quality_chunks = Constants.YOUTUBE_LOW_QUALITY_MAX_CHUNKS

        refresh_platforms = ("youtube", "youtube_live", "twitch")

        effective_stream_timeout = 40 if (is_youtube or platform_id in refresh_platforms) else self.config.STREAM_TIMEOUT

        while True:
            ap = self._get_ap()
            if ap is None:
                logger.info(
                    "AudioProcessor nicht mehr verfügbar – beende StreamHandler."
                )
                break

            if not ap._processing.is_set() or ap._stop_event.is_set():
                logger.info("Processing stopped by user.")
                break

            if current_process.poll() is not None:
                logger.info("FFmpeg process terminated – finishing loop.")
                try:
                    stderr = self._read_stderr_nonblocking(current_process, 4096)
                    if stderr:
                        logger.warning(f"FFmpeg exit stderr: {stderr[:200]}")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Lesen von FFmpeg stderr: {e}", exc_info=True
                    )
                break

            current_time = time.time()

            if current_time - self._last_stderr_read > self._stderr_read_interval:
                try:
                    stderr = self._read_stderr_nonblocking(current_process, 1024)
                    if stderr:
                        logger.debug(f"FFmpeg stderr: {stderr.strip()}")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Lesen von FFmpeg stderr: {e}", exc_info=True
                    )
                self._last_stderr_read = current_time

            if current_time - last_data_time > effective_stream_timeout:
                consecutive_timeouts += 1
                if logger.isEnabledFor(logging.DEBUG) and consecutive_timeouts % 5 == 0:
                    log_debug(
                        "audio",
                        f"Timeout: consecutive_timeouts={consecutive_timeouts}, processed_seconds={ap._processed_seconds:.1f}",
                    )
                if consecutive_timeouts > self.config.MAX_CONSECUTIVE_ERRORS:
                    if reconnect_attempts < max_reconnects:
                        reconnect_attempts += 1
                        wait = min(self._max_backoff, backoff)
                        logger.warning(
                            f"⚠️ Stream timeout - reconnecting attempt {reconnect_attempts}/{max_reconnects}, waiting {wait:.1f}s"
                        )
                        if consecutive_timeouts % 2 == 0:
                            info_callback(
                                f"🔄 Reconnecting... ({reconnect_attempts}/{max_reconnects})"
                            )
                        ap._stop_event.wait(wait)
                        backoff *= 2
                        consecutive_timeouts = 0
                        continue
                    else:
                        logger.info(
                            "📴 Stream appears to be offline – ending processing."
                        )
                        if ap._finished_callback:
                            try:
                                ap._finished_callback()
                            except Exception as e:
                                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                    raise
                                logger.error(
                                    f"Fehler im finished_callback: {e}", exc_info=True
                                )
                        break
                else:
                    wait = min(self._max_backoff, backoff)
                    logger.warning(
                        f"⚠️ Temporary timeout ({consecutive_timeouts}/{self.config.MAX_CONSECUTIVE_ERRORS}), waiting {wait:.1f}s"
                    )
                    if consecutive_timeouts % 2 == 0:
                        info_callback(
                            f"⏳ Timeout {consecutive_timeouts}/{self.config.MAX_CONSECUTIVE_ERRORS} – waiting..."
                        )
                    time.sleep(wait)
                    continue
            else:
                consecutive_timeouts = 0
                backoff = 1.0

            if (
                ap._expected_duration is not None
                and ap._processed_seconds >= ap._expected_duration - 1.0
            ):
                logger.info(
                    f"⏱️ Expected duration reached ({ap._processed_seconds:.1f}s >= {ap._expected_duration:.1f}s - 1s), stopping."
                )
                normal_ending = True
                break

            if (
                platform_id in refresh_platforms
                and time.time() - last_url_refresh > url_refresh_interval
            ):
                logger.info(f"🔄 Scheduled URL refresh for {platform_name}")
                try:
                    new_url = self.stream_manager.extract_audio_url(
                        original_video_url, force_refresh=True
                    )
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(f"Fehler beim URL-Refresh: {e}", exc_info=True)
                    new_url = None
                if new_url and new_url != audio_url:
                    logger.info(
                        f"✅ New {platform_name} URL obtained, restarting FFmpeg..."
                    )
                    ffmpeg = self._get_ffmpeg()
                    if ffmpeg is None:
                        logger.error("FFmpegManager nicht mehr verfügbar – breche ab.")
                        break
                    ffmpeg.stop_stream(ap._current_stream_id)
                    time.sleep(1.0)
                    try:
                        new_process = ffmpeg.start_stream(
                            video_url=original_video_url,
                            output_queue=None,
                            process_id=ap._current_stream_id,
                            force_refresh_audio_url=True,
                            seek_seconds=ap._processed_seconds,
                            detected_language=detected_language,
                        )
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.error(
                            f"Fehler beim Neustart von FFmpeg: {e}", exc_info=True
                        )
                        new_process = None
                    if new_process is None:
                        logger.error("❌ Could not restart FFmpeg after URL refresh")
                        break
                    current_process = new_process
                    audio_url = new_url
                    last_url_refresh = time.time()
                    logger.info(
                        f"✅ FFmpeg restarted with new URL (PID: {current_process.pid})"
                    )
                    time.sleep(2.0)
                    continue

            try:
                audio_data = self._read_with_timeout(current_process)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Lesen der Audiodaten: {e}", exc_info=True)
                audio_data = None

            if audio_data is None:
                if current_process.poll() is not None:
                    try:
                        stderr = self._read_stderr_nonblocking(current_process, 4096)
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.error(
                            f"Fehler beim Lesen von stderr: {e}", exc_info=True
                        )
                        stderr = ""
                    if self._needs_url_refresh(stderr):
                        logger.info(f"🔄 Detected URL refresh needed: {stderr[:200]}")
                        if (
                            platform_id in refresh_platforms
                            and refresh_count < max_refresh_attempts
                        ):
                            new_url = self._refresh_platform_url(
                                original_video_url, platform_id
                            )
                            if new_url:
                                refresh_count += 1
                                try:
                                    new_process = self._restart_ffmpeg_with_new_url(
                                        ap._current_stream_id,
                                        original_video_url,
                                        detected_language,
                                    )
                                except Exception as e:
                                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                        raise
                                    logger.error(
                                        f"Fehler beim Neustart von FFmpeg: {e}",
                                        exc_info=True,
                                    )
                                    new_process = None
                                if new_process is None:
                                    logger.error(
                                        "❌ Failed to restart FFmpeg, aborting session."
                                    )
                                    break
                                current_process = new_process
                                audio_url = new_url
                                last_data_time = time.time()
                                continue
                        logger.info("📴 Stream ended (no more URL refreshes possible).")
                        if ap._finished_callback:
                            try:
                                ap._finished_callback()
                            except Exception as e:
                                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                    raise
                                logger.error(
                                    f"Fehler im finished_callback: {e}", exc_info=True
                                )
                        break
                    else:
                        logger.warning(f"FFmpeg terminated: {stderr[:200]}")
                        break
                else:
                    ap._empty_reads += 1
                    if ap._empty_reads > self.config.MAX_EMPTY_READS:
                        logger.warning(f"⚠️ Too many empty reads: {ap._empty_reads}")
                        error_callback("❌ No audio data received")
                        break
                    sleep_time = min(
                        1.0, self.config.READ_RETRY_DELAY * ap._empty_reads
                    )
                    time.sleep(sleep_time)
                    continue

            if len(audio_data) == 0:
                ap._empty_reads += 1
                if ap._empty_reads > self.config.MAX_EMPTY_READS:
                    logger.warning(f"⚠️ Too many empty reads: {ap._empty_reads}")
                    error_callback("❌ No audio data received")
                    break
                sleep_time = min(1.0, self.config.READ_RETRY_DELAY * ap._empty_reads)
                time.sleep(sleep_time)
                continue

            ffmpeg = self._get_ffmpeg()
            if ffmpeg is not None:
                ffmpeg.update_process_activity(ap._current_stream_id)
            last_data_time = time.time()
            ap._empty_reads = 0
            ap._chunk_counter += 1
            ap._total_bytes_processed += len(audio_data)
            ap._processed_seconds = (
                ap._total_bytes_processed / self.config.BYTES_PER_SECOND
            )

            min_expected = int(
                self.config.MIN_CHUNK_BYTES
                * Constants.LOW_QUALITY_CHUNK_THRESHOLD_FACTOR
            )
            if len(audio_data) < min_expected:
                consecutive_low_quality_chunks += 1
                log_level = Constants.LOW_QUALITY_CHUNK_LOG_LEVEL
                log_interval = Constants.LOW_QUALITY_CHUNK_LOG_INTERVAL
                if logger.isEnabledFor(logging.DEBUG) and log_level <= logging.DEBUG:
                    logger.debug(
                        f"📉 Chunk too small: {len(audio_data)} bytes (expected min {min_expected})"
                    )
                elif consecutive_low_quality_chunks % log_interval == 0:
                    logger.debug(
                        f"📉 Chunk too small (every {log_interval}th): {len(audio_data)} bytes"
                    )
                if consecutive_low_quality_chunks >= max_low_quality_chunks:
                    logger.warning(
                        f"📴 Too many low-quality chunks ({consecutive_low_quality_chunks}), forcing reconnect"
                    )
                    if (
                        platform_id in refresh_platforms
                        and refresh_count < max_refresh_attempts
                    ):
                        new_url = self._refresh_platform_url(
                            original_video_url, platform_id
                        )
                        if new_url:
                            refresh_count += 1
                            try:
                                new_process = self._restart_ffmpeg_with_new_url(
                                    ap._current_stream_id,
                                    original_video_url,
                                    detected_language,
                                )
                            except Exception as e:
                                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                    raise
                                logger.error(
                                    f"Fehler beim Neustart von FFmpeg: {e}",
                                    exc_info=True,
                                )
                                new_process = None
                            if new_process:
                                current_process = new_process
                                audio_url = new_url
                                consecutive_low_quality_chunks = 0
                                last_data_time = time.time()
                                continue
                    if reconnect_attempts < max_reconnects:
                        reconnect_attempts += 1
                        wait = min(self._max_backoff, backoff)
                        logger.info(
                            f"🔄 Reconnecting after low quality... ({reconnect_attempts}/{max_reconnects})"
                        )
                        time.sleep(wait)
                        backoff *= 2
                        try:
                            new_process = self._restart_ffmpeg_with_new_url(
                                ap._current_stream_id,
                                original_video_url,
                                detected_language,
                            )
                        except Exception as e:
                            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                raise
                            logger.error(
                                f"Fehler beim Neustart von FFmpeg: {e}", exc_info=True
                            )
                            new_process = None
                        if new_process:
                            current_process = new_process
                            consecutive_low_quality_chunks = 0
                            last_data_time = time.time()
                            continue
                    else:
                        break
            else:
                consecutive_low_quality_chunks = 0

            try:
                ap._process_audio_data(
                    audio_data,
                    transcription_callback,
                    translation_callback,
                    info_callback,
                    error_callback,
                )
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"Fehler bei der Verarbeitung der Audiodaten: {e}", exc_info=True
                )
                error_callback("❌ Fehler bei der Verarbeitung")
        if normal_ending:
            if ap._finished_callback:
                try:
                    ap._finished_callback()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler im finished_callback: {e}", exc_info=True
                    )
        elif not ap._stop_event.is_set() and not error_occurred:
            error_callback(
                "❌ Stream wurde unerwartet beendet – versuche Neuverbindung..."
            )

    def _read_with_timeout(self, process: subprocess.Popen) -> Optional[bytes]:
        size = self.config.CHUNK_SIZE_BYTES
        if IS_WINDOWS:
            data = bytearray()
            end_time = time.time() + Constants.READ_CHUNK_TIMEOUT
            while len(data) < size and time.time() < end_time:
                try:
                    fd = process.stdout.fileno()
                    try:
                        os.set_blocking(fd, False)
                    except OSError as e:
                        logger.debug(f"set_blocking fehlgeschlagen: {e}")
                        pass
                    try:
                        chunk = os.read(fd, min(size - len(data), 4096))
                        if not chunk:
                            time.sleep(0.01)
                            continue
                        data.extend(chunk)
                    except BlockingIOError:
                        time.sleep(0.01)
                    except OSError as e:
                        if e.errno == 9:
                            break
                        raise
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(f"Fehler beim Lesen (Windows): {e}", exc_info=True)
                time.sleep(0.005)
            return bytes(data) if data else b""
        else:
            import select

            data = bytearray()
            remaining = size
            end_time = time.time() + Constants.READ_CHUNK_TIMEOUT
            try:
                fd = process.stdout.fileno()
                os.set_blocking(fd, False)
            except OSError as e:
                logger.debug(f"set_blocking fehlgeschlagen: {e}")
            try:
                while remaining > 0 and time.time() < end_time:
                    try:
                        rlist, _, _ = select.select(
                            [fd], [], [], max(0, end_time - time.time())
                        )
                    except (select.error, ValueError) as e:
                        logger.debug(f"select fehlgeschlagen: {e}")
                        time.sleep(0.01)
                        continue
                    if fd in rlist:
                        try:
                            chunk = os.read(fd, min(remaining, 4096))
                            if not chunk:
                                break
                            data.extend(chunk)
                            remaining -= len(chunk)
                        except BlockingIOError:
                            time.sleep(Constants.READ_WITH_TIMEOUT_SELECT_INTERVAL)
                        except OSError as e:
                            if e.errno == 9:
                                break
                            raise
                    else:
                        time.sleep(Constants.READ_WITH_TIMEOUT_SELECT_INTERVAL)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Lesen (Unix): {e}", exc_info=True)
            finally:
                try:
                    os.set_blocking(fd, True)
                except Exception:
                    pass
            return bytes(data) if data else b""

    def _read_stderr_nonblocking(
        self, process: subprocess.Popen, max_bytes: int = 4096
    ) -> str:
        if not process.stderr:
            return ""
        try:
            fd = process.stderr.fileno()
        except Exception:
            return ""

        if not IS_WINDOWS:
            import fcntl

            flags = None
            try:
                flags = fcntl.fcntl(fd, fcntl.F_GETFL)
                fcntl.fcntl(fd, fcntl.F_SETFL, flags | os.O_NONBLOCK)
            except Exception as e:
                logger.debug(f"fcntl fehlgeschlagen: {e}")
                flags = None

        try:
            data = os.read(fd, max_bytes)
            if data:
                return data.decode("utf-8", errors="ignore")
        except (BlockingIOError, OSError):
            pass
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.debug(f"Fehler beim Lesen von stderr: {e}")
        finally:
            if not IS_WINDOWS and flags is not None:
                try:
                    fcntl.fcntl(fd, fcntl.F_SETFL, flags)
                except Exception:
                    pass
        return ""

    def _needs_url_refresh(self, stderr: str) -> bool:
        patterns = [
            "403",
            "401",
            "forbidden",
            "unauthorized",
            "invalid parameters",
            "http error 403",
            "http error 401",
            "access denied",
            "signature expired",
            "token expired",
            "url signature expired",
        ]
        stderr_lower = stderr.lower()
        return any(p in stderr_lower for p in patterns)

    def _refresh_platform_url(self, video_url: str, platform: str) -> Optional[str]:
        max_attempts = 3
        for attempt in range(1, max_attempts + 1):
            try:
                logger.info(
                    f"🔄 Attempt {attempt}/{max_attempts} to refresh {platform} URL..."
                )
                new_url = self.stream_manager.extract_audio_url(
                    video_url, force_refresh=True
                )
                if new_url:
                    logger.info(
                        f"✅ Successfully obtained new {platform} URL (attempt {attempt})"
                    )
                    return new_url
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Refresh attempt {attempt} error: {e}")
                if logger.isEnabledFor(logging.DEBUG):
                    logger.exception("Stacktrace:")
            if attempt < max_attempts:
                wait = 2 ** (attempt - 1)
                time.sleep(wait)
        logger.error(f"❌ All attempts to refresh {platform} URL failed")
        return None

    def _restart_ffmpeg_with_new_url(
        self, process_id: str, video_url: str, detected_language: Optional[str] = None
    ) -> Optional[subprocess.Popen]:
        ap = self._get_ap()
        if ap is None:
            logger.error("AudioProcessor nicht verfügbar – Neustart abgebrochen.")
            return None

        logger.info(f"🔄 Restarting FFmpeg for {process_id} with new URL...")
        ffmpeg = self._get_ffmpeg()
        if ffmpeg is not None:
            try:
                ffmpeg.stop_stream(process_id)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Stoppen des Streams: {e}", exc_info=True)
            time.sleep(0.5)
        else:
            logger.error("FFmpegManager nicht verfügbar – Neustart abgebrochen.")
            return None

        seek_seconds = ap._processed_seconds
        if ap._expected_duration is not None and seek_seconds > ap._expected_duration:
            seek_seconds = max(0, ap._expected_duration - 5)
            logger.info(
                f"⏩ Seek adjusted to {seek_seconds:.1f}s (within expected duration)"
            )
        try:
            new_process = ffmpeg.start_stream(
                video_url=video_url,
                output_queue=None,
                process_id=process_id,
                force_refresh_audio_url=True,
                seek_seconds=seek_seconds,
                detected_language=detected_language,
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Fehler beim Starten von FFmpeg: {e}", exc_info=True)
            new_process = None
        if new_process:
            logger.info(
                f"✅ Successfully restarted FFmpeg (new PID: {new_process.pid})"
            )
            return new_process
        else:
            logger.error("❌ Failed to restart FFmpeg")
            return None


# -----------------------------------------------------------------------------
# AudioEnhancer
# -----------------------------------------------------------------------------
class AudioEnhancer:
    MAX_COMPARE_LEN = 150
    MAX_LEN_RATIO_DEVIATION = 0.5
    PUNCTUATION = str.maketrans('', '', '.,!?;:')

    def __init__(self, config: Config, settings: "AdvancedSettings"):
        self.config = config
        self.settings = settings
        self._np = None
        self._scipy_signal = None
        self._load_modules()
        self._rapidfuzz_available = False
        try:
            from rapidfuzz import fuzz
            self._fuzz = fuzz
            self._rapidfuzz_available = True
        except ImportError:
            import difflib
            self._difflib = difflib

    def _load_modules(self):
        if NUMPY_AVAILABLE:
            self._np = FastLazyLoader.load("numpy")
        if SCIPY_AVAILABLE:
            self._scipy_signal = FastLazyLoader.load("scipy.signal")

    def enhance_audio(
        self, audio_data: bytes, last_confidence: float, noisereduce_counter: int
    ) -> bytes:
        if (
            not self.config.AUDIO_ENHANCEMENT_ENABLED
            or len(audio_data) < Constants.AUDIO_ENHANCEMENT_MIN_LENGTH
            or self._np is None
        ):
            return audio_data

        try:
            np = self._np
            audio_np = (
                np.frombuffer(audio_data, dtype=np.int16).astype(np.float32) / 32768.0
            )
            if np.isnan(audio_np).any() or np.isinf(audio_np).any():
                return audio_data

            rms = np.sqrt(np.mean(audio_np**2))
            gain = 1.0

            if rms < 1e-8:
                gain = 1.5
            elif rms < 0.005:
                gain = min(1.5, 0.02 / max(rms, 1e-6))
            elif rms > 0.5:
                gain = 0.5 / rms
            elif rms > 0.3:
                gain = 0.3 / rms

            audio_np *= gain

            max_val = np.max(np.abs(audio_np))
            if max_val > 0.99:
                audio_np /= max_val * 1.01

            audio_np -= np.mean(audio_np)

            if (
                self.settings.enable_noise_reduction
                and last_confidence < 0.3
                and rms < 0.003
                and noisereduce_counter % 10 == 0
                and len(audio_data) > Constants.NOISEREDUCE_MIN_LENGTH
            ):
                try:
                    import noisereduce as nr
                    audio_np = nr.reduce_noise(
                        y=audio_np, sr=self.config.SAMPLE_RATE, prop_decrease=0.6
                    )
                    if logger.isEnabledFor(logging.DEBUG):
                        logger.debug(
                            f"🔇 noisereduce angewendet (Konfidenz {last_confidence:.2f}, RMS {rms:.4f})"
                        )
                except ImportError:
                    pass
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ noisereduce fehlgeschlagen: {e}")

            if self._scipy_signal is not None and len(audio_np) > 100:
                try:
                    b, a = self._scipy_signal.butter(
                        2, 80 / (self.config.SAMPLE_RATE / 2), btype="high"
                    )
                    audio_np = self._scipy_signal.filtfilt(b, a, audio_np)
                except Exception:
                    pass

            audio_np = np.clip(audio_np, -0.99, 0.99)

            enhanced = (audio_np * 32767).astype(np.int16).tobytes()
            return enhanced if len(enhanced) == len(audio_data) else audio_data

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"Audio-Enhancement fehlgeschlagen: {e}")
            return audio_data

    def _normalize_text(self, text: str) -> str:
        return text.translate(self.PUNCTUATION).strip().lower()

    def is_duplicate(
        self,
        current_text: str,
        last_text: str,
        recent_texts: Deque[str],
        confidence: float = None,
        last_confidence: float = 0.0,
    ) -> bool:
        if not self.config.DUPLICATE_CHECK_ENABLED:
            return False

        current_text = current_text.strip()
        if not current_text or len(current_text) < self.config.MIN_TEXT_LENGTH:
            return True

        if current_text == last_text:
            return True

        curr_norm = self._normalize_text(current_text)
        last_norm = self._normalize_text(last_text)

        curr_short = curr_norm[: self.MAX_COMPARE_LEN]

        combined = [last_norm] + [self._normalize_text(t) for t in recent_texts if t]

        similarity_threshold = self.settings.duplicate_similarity_threshold

        for prev in combined:
            if not prev:
                continue

            prev_short = prev[: self.MAX_COMPARE_LEN]

            len_curr = len(curr_short)
            len_prev = len(prev_short)
            if len_curr > 0 and len_prev > 0:
                ratio = max(len_curr, len_prev) / min(len_curr, len_prev)
                if ratio > 1.0 + self.MAX_LEN_RATIO_DEVIATION:
                    continue

            if self._rapidfuzz_available:
                sim = self._fuzz.ratio(curr_short, prev_short) / 100.0
            else:
                sim = self._difflib.SequenceMatcher(
                    None, curr_short, prev_short
                ).ratio()

            if sim > similarity_threshold:
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug(
                        "duplicate",
                        f"{sim:.2%} match: '{curr_short[:30]}' ≈ '{prev_short[:30]}'",
                    )
                return True

        words = current_text.lower().split()
        if len(words) > 3:
            unique_ratio = len(set(words)) / len(words)
            if unique_ratio < self.config.MIN_UNIQUE_WORDS_RATIO:
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug(
                        "duplicate",
                        f"Geringe Wortvielfalt: {unique_ratio:.2%} < {self.config.MIN_UNIQUE_WORDS_RATIO:.2%}",
                    )
                return True

        return False


# =============================================================================
# 10. WHISPER CONTROLLER
# =============================================================================
TranscriptionCallback = Callable[[TranscriptionResult], None]
TranslationCallback = Callable[[TranslationResult], None]
InfoCallback = Callable[[str], None]
ErrorCallback = Callable[[str], None]
StatusCallback = Callable[[Dict[str, Any]], None]
FinishedCallback = Callable[[], None]


class WhisperController:
    class State(Enum):
        IDLE = 0
        STARTING = 1
        PROCESSING = 2
        STOPPING = 3
        ERROR = 4

    __slots__ = (
        "gui_ref", "_state", "_state_lock", "_shutdown_event", "_processing_thread",
        "_stop_complete", "_stop_thread", "_stop_in_progress", "_stop_lock",
        "on_transcription", "on_translation", "on_info", "on_error", "on_status", "on_finished",
        "_last_transcription_text", "_duplicate_check_cache"
    )

    def __init__(self, gui_ref: Any) -> None:
        self.gui_ref = weakref.ref(gui_ref)
        self._state = WhisperController.State.IDLE
        self._state_lock = threading.RLock()
        self._shutdown_event = threading.Event()
        self._processing_thread: Optional[threading.Thread] = None
        self._stop_complete = threading.Event()
        self._stop_complete.set()
        self._stop_thread: Optional[threading.Thread] = None
        self._stop_in_progress = False
        self._stop_lock = threading.RLock()

        self.on_transcription: Optional[TranscriptionCallback] = None
        self.on_translation: Optional[TranslationCallback] = None
        self.on_info: Optional[InfoCallback] = None
        self.on_error: Optional[ErrorCallback] = None
        self.on_status: Optional[StatusCallback] = None
        self.on_finished: Optional[FinishedCallback] = None

        self._last_transcription_text = ""
        self._duplicate_check_cache: deque = deque(maxlen=20)

    @property
    def state(self) -> "WhisperController.State":
        with self._state_lock:
            return self._state

    @property
    def is_processing(self) -> bool:
        with self._state_lock:
            return self._state in (
                WhisperController.State.STARTING,
                WhisperController.State.PROCESSING,
            )

    @property
    def is_stopping(self) -> bool:
        with self._state_lock:
            return self._state == WhisperController.State.STOPPING

    def set_callbacks(
        self,
        on_transcription: TranscriptionCallback,
        on_translation: TranslationCallback,
        on_info: InfoCallback,
        on_error: ErrorCallback,
        on_status: StatusCallback,
        on_finished: Optional[FinishedCallback] = None,
    ) -> None:
        self.on_transcription = on_transcription
        self.on_translation = on_translation
        self.on_info = on_info
        self.on_error = on_error
        self.on_status = on_status
        self.on_finished = on_finished

    def start_processing(self) -> None:
        with self._state_lock:
            if self._state != WhisperController.State.IDLE:
                if self.on_status:
                    self.on_status({"status": f"⚠️ Bereits im Zustand {self._state.name}"})
                return
            self._set_state(WhisperController.State.STARTING)

        def start_target():
            try:
                self._start_processing()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"❌ Start Processing Error: {e}", exc_info=True)
                self._handle_error(f"Start fehlgeschlagen: {str(e)[:50]}")
                with self._state_lock:
                    self._set_state(WhisperController.State.ERROR)

        thread = threading.Thread(
            target=start_target, daemon=True, name="ControllerStarter"
        )
        thread.start()

    def stop_processing(self, wait: bool = False, timeout: float = 10.0) -> bool:
        with self._stop_lock:
            if self._stop_in_progress:
                logger.debug("stop_processing bereits in Gang – überspringe")
                return True
            self._stop_in_progress = True

        with self._state_lock:
            if self._state == WhisperController.State.IDLE:
                with self._stop_lock:
                    self._stop_in_progress = False
                return True
            if self._state == WhisperController.State.ERROR:
                logger.debug("stop_processing: switching from ERROR to IDLE")
                self._set_state(WhisperController.State.IDLE)
                with self._stop_lock:
                    self._stop_in_progress = False
                return True
            if self._state not in (
                WhisperController.State.STOPPING,
                WhisperController.State.ERROR,
            ):
                self._set_state(WhisperController.State.STOPPING)

        self._shutdown_event.set()
        gui = self.gui_ref()
        if gui is not None and IS_LINUX and hasattr(gui, "performance_optimizer"):
            gui.performance_optimizer.restore_normal_mode()

        def stop_audio():
            try:
                if gui is not None:
                    if hasattr(gui, "audio_processor"):
                        ap = gui.audio_processor
                        ap._processing.clear()
                        if hasattr(ap, "_stop_event"):
                            ap._stop_event.set()
                    if hasattr(gui, "ffmpeg_manager"):
                        gui.ffmpeg_manager.stop_all_streams()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Audio Stop Fehler: {e}")
            finally:
                self._stop_complete.set()
                if (
                    wait is False
                    and gui
                    and hasattr(gui, "root")
                    and gui.root.winfo_exists()
                ):
                    gui.queue_manager.safe_put("gui", ("status", self._set_state_idle))
                with self._stop_lock:
                    self._stop_in_progress = False

        self._stop_thread = threading.Thread(target=stop_audio, daemon=True)
        self._stop_thread.start()
        self._reset_gui_state()

        if wait:
            if not self._stop_complete.wait(timeout):
                logger.warning(f"⚠️ Stop-Thread nicht innerhalb von {timeout}s beendet")
                with self._stop_lock:
                    self._stop_in_progress = False
                return False
            if self._processing_thread and self._processing_thread.is_alive():
                self._processing_thread.join(timeout=1.0)
            with self._state_lock:
                if self._state != WhisperController.State.ERROR:
                    self._set_state(WhisperController.State.IDLE)
            with self._stop_lock:
                self._stop_in_progress = False
            return True
        else:
            return True

    def safe_exit(self) -> None:
        gui = self.gui_ref()
        if gui is not None:
            try:
                if hasattr(gui, "exit_button"):
                    gui.exit_button.config(state="disabled", text="⏳...")
            except Exception:
                pass
            if hasattr(gui, "_safe_exit_dialog"):
                gui._safe_exit_dialog()
            else:
                self._emergency_cleanup()
                sys.exit(0)
        else:
            self._emergency_cleanup()
            sys.exit(0)

    def dispose(self) -> None:
        self._shutdown_event.set()
        self.stop_processing(wait=False)
        if self._stop_thread and self._stop_thread.is_alive():
            self._stop_thread.join(timeout=2.0)
        logger.info("🧹 Controller disposed")

    def _set_state(self, new_state: "WhisperController.State") -> None:
        with self._state_lock:
            old = self._state
            self._state = new_state
        logger.debug(f"Controller state: {old.name} -> {new_state.name}")
        if self.on_status:
            self.on_status({"controller_state": new_state.name})

    def _set_state_idle(self):
        with self._state_lock:
            if self._state != WhisperController.State.IDLE:
                self._set_state(WhisperController.State.IDLE)

    def _emergency_cleanup(self) -> None:
        if self._shutdown_event.is_set():
            return
        self._shutdown_event.set()
        gui = self.gui_ref()
        if gui is not None:
            if hasattr(gui, "audio_processor"):
                try:
                    gui.audio_processor._processing.clear()
                    if hasattr(gui.audio_processor, "_stop_event"):
                        gui.audio_processor._stop_event.set()
                except Exception:
                    pass
            if hasattr(gui, "ffmpeg_manager"):
                try:
                    gui.ffmpeg_manager.stop_all_streams()
                except Exception:
                    pass

    def _handle_error(self, message: str) -> None:
        logger.error(f"❌ Controller error: {message}")
        if self.on_error:
            self.on_error(message)
        if self.on_status:
            self.on_status({"processing_state": False, "status": f"❌ {message}"})
        self._emergency_cleanup()
        self._reset_gui_state()
        with self._state_lock:
            self._set_state(WhisperController.State.ERROR)

    def _reset_gui_state(self) -> None:
        gui = self.gui_ref()
        if gui is None:
            return

        def update():
            try:
                if hasattr(gui, "status_label") and gui.status_label.winfo_exists():
                    gui.status_label.config(text="✅ READY for new stream")
                if hasattr(gui, "start_button") and gui.start_button.winfo_exists():
                    gui.start_button.config(state="normal")
                if hasattr(gui, "stop_button") and gui.stop_button.winfo_exists():
                    gui.stop_button.config(state="disabled")
                if (
                    hasattr(gui, "stream_title_label")
                    and gui.stream_title_label.winfo_exists()
                ):
                    gui.stream_title_label.config(text="📡 Kein aktiver Stream")
                if (
                    hasattr(gui, "stream_details_label")
                    and gui.stream_details_label.winfo_exists()
                ):
                    gui.stream_details_label.config(text="Bereit für neue Verbindung")
                gui._reset_progress()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ GUI Update Fehler: {e}")

        gui.queue_manager.safe_put("gui", ("status", update))

    def _start_processing(self) -> None:
        gui = self.gui_ref()
        if gui is None:
            self._handle_error("GUI nicht verfügbar")
            return

        url = self._validate_url(gui)
        if url is None:
            self._set_state(WhisperController.State.IDLE)
            return

        if self.on_status:
            self.on_status({"status": "🔍 Analysiere Stream..."})

        self._extract_stream_info(gui, url)

        if self.on_status:
            self.on_status({"status": "🎵 Teste Audio-Stream..."})

        if not self._test_stream(gui, url):
            self._handle_error("Stream nicht erreichbar")
            return

        if self.on_status:
            self.on_status({"status": "🤖 Lade KI-Modell..."})

        if not self._load_and_setup_model(gui):
            self._handle_error("KI-Modell konnte nicht geladen werden")
            return

        self._set_source_language(gui)
        self._configure_translation(gui)

        with self._state_lock:
            if self._state != WhisperController.State.STARTING:
                logger.info("Start abgebrochen – Zustand nicht mehr STARTING")
                return
            self._set_state(WhisperController.State.PROCESSING)

        self._update_gui_buttons(gui, processing=True)

        if IS_LINUX and hasattr(gui, "performance_optimizer"):
            gui.performance_optimizer.optimize_for_processing()

        if self.on_status:
            self.on_status(
                {"processing_state": True, "status": "🚀 Starte Transkription..."}
            )

        self._run_audio_processor(gui, url)

    def _validate_url(self, gui) -> Optional[str]:
        try:
            url = gui.url_entry.get().strip()
        except Exception:
            if self.on_status:
                self.on_status({"status": "❌ URL Fehler"})
            return None

        if not url:
            if self.on_status:
                self.on_status({"status": "❌ Bitte URL eingeben"})
            return None

        try:
            url = PlatformUtils.sanitize_url(url)
            if url.startswith("file://"):
                ok, real_path = PlatformUtils.validate_file_path(url)
                if not ok:
                    if self.on_status:
                        self.on_status({"status": f"❌ {real_path}"})
                    return None
                if not os.path.exists(real_path):
                    if self.on_status:
                        self.on_status({"status": "❌ Datei nicht gefunden"})
                    return None
            else:
                if not url.startswith(("http://", "https://")):
                    url = "https://" + url

                    def update_url():
                        if (
                            gui
                            and hasattr(gui, "url_entry")
                            and gui.url_entry.winfo_exists()
                        ):
                            gui.url_entry.delete(0, "end")
                            gui.url_entry.insert(0, url)

                    if gui and hasattr(gui, "root"):
                        gui.queue_manager.safe_put("gui", ("url_update", update_url))
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            if self.on_status:
                self.on_status({"status": f"❌ Ungültige URL: {e}"})
            return None

        return url

    def _extract_stream_info(self, gui, url: str) -> None:
        try:
            if hasattr(gui, "stream_manager"):
                platform_type, platform_name = gui.stream_manager.detect_platform(url)
            else:
                platform_type, platform_name = "unknown", "Unknown"

            stream_info = None
            try:
                if hasattr(gui, "stream_info_extractor"):
                    stream_info = gui.stream_info_extractor.extract_stream_info(url)
                else:
                    stream_info = StreamInfoExtractor().extract_stream_info(url)
            except Exception:
                stream_info = StreamInfo(
                    title="Live Stream" if "live" in url.lower() else "Stream",
                    uploader=platform_name,
                    duration="Live" if "live" in url.lower() else "Unknown",
                    view_count=0,
                    platform=platform_type,
                )

            if stream_info:
                if self.on_status:
                    self.on_status({"stream_info": stream_info})
                logger.info(f"📡 Stream: {stream_info.title[:50]}...")
                if (
                    hasattr(gui, "audio_processor")
                    and stream_info.duration_seconds is not None
                ):
                    gui.audio_processor.set_expected_duration(
                        stream_info.duration_seconds
                    )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Stream Info Error: {e}")

    def _test_stream(self, gui, url: str) -> bool:
        try:
            if hasattr(gui, "audio_processor"):
                return gui.audio_processor.emergency_diagnosis(url)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Stream Test Error: {e}")
        return False

    def _load_and_setup_model(self, gui) -> bool:
        try:
            if hasattr(gui, "transcription_engine"):
                model_name = (
                    gui.model_var.get() if hasattr(gui, "model_var") else "medium"
                )
                result = gui.transcription_engine.load_model(
                    model_name, set_active=True
                )
                if result is not None:
                    return True
                else:
                    logger.info("🔄 Versuche base model...")
                    result = gui.transcription_engine.load_model(
                        "base", set_active=True
                    )
                    return result is not None
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Model Load Error: {e}")
        return False

    def _set_source_language(self, gui) -> None:
        try:
            if hasattr(gui, "src_lang_var") and hasattr(gui, "transcription_engine"):
                src_name = gui.src_lang_var.get()
                if src_name != "Automatisch":
                    for name, code in SORTED_LANGUAGES:
                        if name == src_name:
                            gui.transcription_engine.forced_language = code
                            logger.info(f"🔤 Quellsprache manuell gesetzt: {code}")
                            break
                else:
                    gui.transcription_engine.forced_language = None
                    logger.info("🔤 Quellsprache: Automatisch (Whisper-Erkennung)")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Quellsprache setzen fehlgeschlagen: {e}")
            if hasattr(gui, "transcription_engine"):
                gui.transcription_engine.forced_language = None

    def _configure_translation(self, gui) -> None:
        try:
            if hasattr(gui, "translation_engine") and hasattr(gui, "lang_var"):
                selected_name = gui.lang_var.get()
                target_lang = "de"
                for name, code in SORTED_LANGUAGES:
                    if name == selected_name:
                        target_lang = code
                        break
                gui.translation_engine.set_target_language(target_lang)
                lang_display = LANGUAGE_SHORT_CODES.get(target_lang, target_lang)
                if hasattr(gui, "translation_header"):

                    def update_header():
                        if (
                            gui
                            and hasattr(gui, "translation_header")
                            and gui.translation_header.winfo_exists()
                        ):
                            gui.translation_header.config(
                                text=f"🌐 Übersetzung ({lang_display})"
                            )

                    if gui and hasattr(gui, "root"):
                        gui.queue_manager.safe_put(
                            "gui", ("translation_header", update_header)
                        )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Translation Setup Error: {e}")

    def _update_gui_buttons(self, gui, processing: bool) -> None:
        def task():
            try:
                if hasattr(gui, "start_button") and gui.start_button.winfo_exists():
                    gui.start_button.config(
                        state="disabled" if processing else "normal"
                    )
                if hasattr(gui, "stop_button") and gui.stop_button.winfo_exists():
                    gui.stop_button.config(state="normal" if processing else "disabled")
            except Exception:
                pass

        if gui and hasattr(gui, "root") and gui.root.winfo_exists():
            gui.queue_manager.safe_put("gui", ("button_update", task))

    def _run_audio_processor(self, gui, url: str) -> None:
        if not hasattr(gui, "audio_processor"):
            self._handle_error("Audio-Processor nicht verfügbar")
            return

        ap = gui.audio_processor
        ap._stop_event.clear()
        ap.set_progress_callback(self._on_progress)

        def transcription_callback(result: TranscriptionResult) -> None:
            if result and self.on_transcription:
                try:
                    self.on_transcription(result)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Transcription Callback Error: {e}")

        def translation_callback(result: TranslationResult) -> None:
            if result and self.on_translation:
                try:
                    self.on_translation(result)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ Translation Callback Error: {e}")

        def info_callback(message: str) -> None:
            if self.on_info:
                try:
                    self.on_info(message)
                except Exception:
                    pass

        def error_callback(message: str) -> None:
            if self.on_error:
                try:
                    self.on_error(message)
                except Exception:
                    pass
            self._emergency_cleanup()

        def file_finished_callback() -> None:
            logger.info("✅ Dateiende erkannt")
            if self.on_status:
                self.on_status({"file_finished": True})
            self._processing_finished()

        try:
            ap.start_processing(
                url=url,
                transcription_callback=transcription_callback,
                translation_callback=translation_callback,
                info_callback=info_callback,
                error_callback=error_callback,
                finished_callback=file_finished_callback,
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler im AudioProcessor: {e}", exc_info=True)
            self._handle_error(str(e)[:100])

    def _on_progress(self, processed: int, total: Optional[int], chunks: int) -> None:
        gui = self.gui_ref()
        if gui is not None and hasattr(gui, "update_progress"):
            gui.queue_manager.safe_put(
                "gui",
                (
                    "progress",
                    lambda: gui.update_progress(processed, total, chunks),
                ),
            )

    def _processing_finished(self) -> None:
        gui = self.gui_ref()
        if gui is not None and hasattr(gui, "audio_processor"):
            if gui.audio_processor._processing.is_set():
                logger.warning(
                    "⚠️ _processing_finished aufgerufen, aber AudioProcessor läuft noch – erzwinge Reset"
                )
        with self._state_lock:
            if self._state == WhisperController.State.PROCESSING:
                self._set_state(WhisperController.State.IDLE)
            else:
                logger.debug(
                    f"_processing_finished: aktueller Zustand {self._state.name} – kein Wechsel nötig"
                )
        self._reset_gui_state()
        if self.on_status:
            self.on_status(
                {"processing_state": False, "status": "✅ Verarbeitung beendet"}
            )
        if self.on_finished:
            try:
                self.on_finished()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Finished callback error: {e}")


# =============================================================================
# AUDIO PROCESSOR
# =============================================================================
class AudioProcessor:
    MAX_BUFFER_SECONDS = 30

    def __init__(
        self,
        controller_ref: Any,
        ffmpeg_manager: FFmpegManager,
        settings: Optional["AdvancedSettings"] = None,
        use_browser_cookies: bool = True,
    ):
        self._silent_chunk_counter = 0
        self.controller_ref = controller_ref
        self.ffmpeg_manager = ffmpeg_manager
        self.settings = settings or AdvancedSettings()
        self.use_browser_cookies = use_browser_cookies
        self.config = self.settings.config
        self.sample_rate = self.config.SAMPLE_RATE
        self.channels = self.config.CHANNELS
        self.audio_format = self.config.AUDIO_FORMAT
        self.chunk_size = self.config.CHUNK_SIZE_BYTES
        self.overlap_size = self.config.OVERLAP_SIZE_BYTES

        self._max_buffer_bytes = self.MAX_BUFFER_SECONDS * self.config.BYTES_PER_SECOND

        self.transcription_engine: Optional["TranscriptionEngine"] = None
        self.translation_engine: Optional["BaseTranslationEngine"] = None
        self._fallback_translation_engine: Optional["BaseTranslationEngine"] = None
        self.plugin_manager: Optional["PluginManager"] = None

        self._stop_event = threading.Event()
        self._processing = threading.Event()
        self._processing_lock = threading.RLock()
        self._current_stream_id: Optional[str] = None
        self._last_successful_read_time = time.time()
        self._consecutive_empty_chunks = 0
        self._cleanup_done = False
        self._resource_lock = threading.RLock()

        self._translation_enabled = threading.Event()
        self._translation_enabled.set()

        self._last_transcription_text = ""
        self._recent_transcriptions: Deque[str] = deque(
            maxlen=self.config.RECENT_TRANSCRIPTIONS_SIZE
        )
        self._duplicate_lock = threading.RLock()

        self._timed_transcriptions: Deque["TranscriptionResult"] = deque(
            maxlen=self.config.SUBTITLE_BUFFER_SIZE
        )
        self._timed_translations: Deque["TranslationResult"] = deque(
            maxlen=self.config.SUBTITLE_BUFFER_SIZE
        )
        self._subtitle_lock = threading.RLock()
        self.subtitle_mode = False

        self._word_count_history: Deque[float] = deque(maxlen=10)
        self._word_count_lock = threading.RLock()
        self._smoothed_word_count: Optional[float] = None
        self._last_chunk_duration = self.config.CHUNK_DURATION
        self._chunk_stable_counter = 0

        self._slow_chunks = 0
        self._last_realtime_factor = 0.0

        self._stats_lock = threading.RLock()
        self._chunk_counter = 0
        self._empty_reads = 0
        self._stream_start_time: Optional[float] = None
        self._total_bytes_processed = 0
        self._processed_seconds = 0.0
        self._consecutive_errors = 0
        self._consecutive_successes = 0
        self._consecutive_timeouts = 0
        self._low_conf_counter = 0
        self._read_error_count = 0
        self._max_backoff = Constants.MAX_BACKOFF

        self._audio_buffer = bytearray()
        self._max_buffer_size = self.config.MAX_CHUNK_BYTES * 5
        self._last_buffer_flush = time.time()
        self._buffer_lock = threading.RLock()

        self._sentence_buffer = ""
        self._sentence_segments: List[TranscriptionResult] = []
        self._sentence_lock = threading.RLock()

        transcribe_workers = getattr(self.settings, "transcription_workers", 2)
        translate_workers = getattr(self.settings, "translation_workers", 1)
        self._transcription_executor = OptimizedThreadPoolExecutor(
            max_workers=transcribe_workers, thread_name_prefix="Transcribe"
        )
        self._translation_executor = OptimizedThreadPoolExecutor(
            max_workers=translate_workers, thread_name_prefix="Translate"
        )

        self._total_file_size: Optional[int] = None
        self._progress_callback: Optional[Callable[[int, Optional[int], int], None]] = None
        self._last_progress_update = 0.0
        self._progress_update_interval = Constants.PROGRESS_UPDATE_INTERVAL
        self._expected_duration: Optional[float] = None
        self._finished_callback: Optional[Callable] = None
        self._min_chunk_duration = self.config.MIN_CHUNK_DURATION

        self._audio_enhancer = AudioEnhancer(self.config, self.settings)
        self.stream_manager = StreamManager(
            enable_debug=(DEBUG_LEVEL >= 1),
            use_browser_cookies=self.use_browser_cookies,
        )
        self._stream_handler = StreamHandler(self, self.stream_manager)

        self.last_confidence = 1.0
        self._last_confidence_lock = threading.RLock()
        self._noisereduce_counter = 0
        self._noisereduce_lock = threading.RLock()

        if logger.isEnabledFor(logging.DEBUG):
            self._last_gpu_stats_time = 0.0

        self._vad_fallback_enabled = True

        logger.info(
            "✅ AudioProcessor initialized (optimized with OptimizedThreadPoolExecutor):"
        )
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Chunk: {self.config.CHUNK_DURATION}s / {self.chunk_size:,} bytes")
        logger.info(f"   Sample Rate: {self.sample_rate} Hz")
        logger.info(f"   Overlap: {self.overlap_size:,} bytes")
        logger.info(f"   Bytes/sec: {self.config.BYTES_PER_SECOND:,}")
        logger.info(f"   Max Buffer: {self._max_buffer_bytes:,} bytes ({self.MAX_BUFFER_SECONDS}s)")
        logger.info(f"   Transcribe Workers: {transcribe_workers}, Translate Workers: {translate_workers}")

    def _update_chunk_size(self) -> None:
        self.chunk_size = int(self.config.CHUNK_DURATION * self.config.BYTES_PER_SECOND)

    def _get_config_type(self) -> str:
        if isinstance(self.config, RealtimeConfig):
            return "realtime"
        elif isinstance(self.config, HighAccuracyConfig):
            return "high_accuracy"
        elif isinstance(self.config, YouTubeOptimizedConfig):
            return "youtube"
        return "default"

    def set_expected_duration(self, duration: Optional[float]) -> None:
        self._expected_duration = duration
        if duration is not None:
            logger.info(f"⏱️ Expected stream duration set: {duration:.1f}s")

    def set_progress_callback(
        self, callback: Callable[[int, Optional[int], int], None]
    ) -> None:
        self._progress_callback = callback

    def set_engines(
        self,
        transcription_engine: TranscriptionEngine,
        translation_engine: BaseTranslationEngine,
        fallback_translation_engine: Optional[BaseTranslationEngine] = None,
        plugin_manager: Optional[PluginManager] = None,
    ) -> None:
        self.transcription_engine = transcription_engine
        self.translation_engine = translation_engine
        self._fallback_translation_engine = fallback_translation_engine
        self.plugin_manager = plugin_manager
        if (
            hasattr(translation_engine, "is_functional")
            and not translation_engine.is_functional()
        ):
            self._translation_enabled.clear()

    def set_vad_fallback_enabled(self, enabled: bool) -> None:
        self._vad_fallback_enabled = enabled
        if self.transcription_engine:
            self.transcription_engine.set_vad_fallback_enabled(enabled)
        logger.debug(f"VAD-Fallback im AudioProcessor {'aktiviert' if enabled else 'deaktiviert'}")

    def enable_subtitle_mode(self, enabled: bool) -> None:
        self.subtitle_mode = enabled
        logger.info(f"🎬 Subtitle mode: {'ENABLED' if enabled else 'DISABLED'}")

    def start_processing(
        self,
        url: str,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
        info_callback: InfoCallback,
        error_callback: ErrorCallback,
        finished_callback: Optional[FinishedCallback] = None,
    ) -> None:
        logger.info(f"\n🔊 [START_PROCESSING] URL: {url[:80]}...")
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Chunk Size: {self.chunk_size:,} bytes")

        url = PlatformUtils.sanitize_url(url)
        if url.startswith("file://"):
            try:
                ok, real_path = PlatformUtils.validate_file_path(url)
                if not ok:
                    error_callback(f"❌ {real_path}")
                    return
                file_path = real_path
                self._total_file_size = os.path.getsize(file_path)
                logger.info(f"📁 Lokale Datei, Größe: {self._total_file_size} bytes")
            except OSError as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                error_callback(f"❌ Dateizugriffsfehler: {e}")
                return
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"Unerwarteter Fehler bei Dateiprüfung: {e}", exc_info=True
                )
                error_callback("❌ Fehler bei der Dateiprüfung")
                return
        else:
            self._total_file_size = None

        try:
            health_issues = self._platform_specific_health_check()
            if health_issues:
                for issue in health_issues:
                    logger.warning(f"⚠️ {issue}")
                    info_callback(f"⚠️ {issue}")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Fehler beim Health-Check: {e}", exc_info=True)

        with self._processing_lock:
            if self._processing.is_set():
                logger.warning(
                    "⚠️ Vorheriger Prozess läuft noch – stoppe diesen zuerst."
                )
                if not self.stop_processing(wait=True, timeout=10.0):
                    error_callback("❌ Vorheriger Prozess konnte nicht gestoppt werden")
                    return
            self._processing.set()
            self._process_finished = threading.Event()
            self._stop_event.clear()
            self._current_stream_id = f"stream_{int(time.time())}"
            self._stream_start_time = time.time()
            with self._stats_lock:
                self._chunk_counter = 0
                self._total_bytes_processed = 0
                self._processed_seconds = 0.0
                self._read_error_count = 0
                self._consecutive_timeouts = 0
                self._consecutive_errors = 0
                self._consecutive_successes = 0
                self._low_conf_counter = 0
                self._slow_chunks = 0
                self._last_realtime_factor = 0.0
            with self._buffer_lock:
                self._audio_buffer = bytearray()
            self._finished_callback = finished_callback
            with self._word_count_lock:
                self._word_count_history.clear()
            self._smoothed_word_count = None
            self._last_chunk_duration = self.config.CHUNK_DURATION
            self._chunk_stable_counter = 0
            logger.info(
                f"✅ Flags gesetzt: processing=True, ID={self._current_stream_id}"
            )

        thread = threading.Thread(
            target=self._process_loop_enhanced,
            args=(
                url,
                transcription_callback,
                translation_callback,
                info_callback,
                error_callback,
            ),
            daemon=True,
            name=f"AudioProc_{self._current_stream_id}",
        )
        thread.start()
        logger.info(f"✅ Processing thread gestartet: {thread.name}")

    def _process_loop_enhanced(
        self,
        url: str,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
        info_callback: InfoCallback,
        error_callback: ErrorCallback,
    ) -> None:
        process: Optional[subprocess.Popen] = None
        detected_language: Optional[str] = None
        error_occurred = False
        stderr_thread: Optional[threading.Thread] = None
        stop_stderr = threading.Event()

        try:
            logger.info(f"\n🎬 [PROCESS_LOOP] Start für: {url[:60]}...")
            info_callback("🔍 Extracting audio URL...")
            try:
                audio_url = self.stream_manager.extract_audio_url(url)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler bei der Audio-URL-Extraktion: {e}", exc_info=True)
                error_callback("❌ Audio-URL konnte nicht extrahiert werden")
                return
            if not audio_url:
                error_callback("❌ Could not extract audio URL")
                error_occurred = True
                return
            logger.info(f"✅ Audio URL: {audio_url[:80]}...")
            info_callback("🔍 Testing audio stream...")
            if not self._test_audio_stream(audio_url):
                logger.warning("⚠️ Stream test failed, trying anyway...")
            info_callback("🔧 Setting up FFmpeg...")
            logger.info("🚀 Starting FFmpeg process...")
            try:
                process = self.ffmpeg_manager.start_stream(
                    video_url=url,
                    output_queue=None,
                    process_id=self._current_stream_id,
                    audio_url=audio_url,
                    detected_language=detected_language,
                )
                if process is None:
                    error_callback("❌ FFmpeg konnte nicht gestartet werden")
                    error_occurred = True
                    return
            except FileNotFoundError:
                error_callback("❌ FFmpeg not found - please install")
                error_occurred = True
                return
            except (OSError, PermissionError) as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                error_callback(f"❌ FFmpeg konnte nicht gestartet werden: {e}")
                error_occurred = True
                return
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"Unerwarteter Fehler beim Starten von FFmpeg: {e}", exc_info=True
                )
                error_callback("❌ Unerwarteter Fehler beim Starten von FFmpeg")
                error_occurred = True
                return
            logger.info(f"✅ FFmpeg started (PID: {process.pid})")

            def log_stderr():
                import select

                while not stop_stderr.is_set():
                    try:
                        if process and process.stderr:
                            rlist, _, _ = select.select([process.stderr], [], [], 0.2)
                            if process.stderr in rlist:
                                line = process.stderr.readline()
                                if line:
                                    logger.debug(
                                        f"FFmpeg stderr: {line.decode('utf-8', errors='ignore').strip()}"
                                    )
                                else:
                                    break
                        else:
                            break
                    except (OSError, ValueError) as e:
                        logger.debug(f"stderr-Lese-Fehler: {e}")
                        break
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.error(
                            f"Unerwarteter Fehler im stderr-Thread: {e}", exc_info=True
                        )
                        break

            stderr_thread = threading.Thread(
                target=log_stderr, daemon=True, name="FFmpegStderr"
            )
            stderr_thread.start()

            info_callback("⏳ Initializing stream...")
            wait_time = self.config.INITIAL_BUFFER_SECONDS
            if any(
                keyword in audio_url.lower()
                for keyword in ["hls", ".m3u8", "manifest.googlevideo.com"]
            ):
                wait_time = 3.0
                logger.info(f"🎯 HLS/Live stream detected, waiting {wait_time}s...")
            time.sleep(wait_time)

            if process.poll() is not None:
                try:
                    stderr = self._read_stderr(process, 1000)
                    error_msg = f"FFmpeg died: {stderr[:200]}"
                    logger.error(f"❌ {error_msg}")
                    error_callback(f"❌ {error_msg}")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Lesen von FFmpeg stderr: {e}", exc_info=True
                    )
                    error_callback("❌ FFmpeg failed to start")
                error_occurred = True
                return

            info_callback("✅ Stream connected - starting transcription...")
            is_youtube = any(
                domain in audio_url
                for domain in ["youtube.com", "youtu.be", "googlevideo.com"]
            )
            if logger.isEnabledFor(logging.DEBUG):
                log_debug(
                    "audio",
                    f"Detected stream type: {'YouTube' if is_youtube else 'Standard'}",
                )

            if is_youtube:
                logger.info("🎯 Using YouTube-optimized streaming loop")
                self._stream_handler.run_loop(
                    process,
                    audio_url,
                    url,
                    detected_language,
                    transcription_callback,
                    translation_callback,
                    info_callback,
                    error_callback,
                    is_youtube=True,
                )
            else:
                logger.info("🎯 Using standard streaming loop")
                self._stream_handler.run_loop(
                    process,
                    audio_url,
                    url,
                    detected_language,
                    transcription_callback,
                    translation_callback,
                    info_callback,
                    error_callback,
                    is_youtube=False,
                )
            logger.info(
                f"🔚 [LOOP END] Reason: {'Stop requested' if self._stop_event.is_set() else 'Process ended'}"
            )
        except subprocess.TimeoutExpired as e:
            error_callback(f"❌ Timeout - stream not reachable: {e}")
            error_occurred = True
        except FileNotFoundError:
            error_callback("❌ FFmpeg not found - please install")
            error_occurred = True
        except OSError as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            error_msg = f"OS error: {str(e)[:100]}"
            logger.error(f"❌ {error_msg}")
            error_callback(f"❌ {error_msg}")
            error_occurred = True
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            error_msg = f"Unexpected error: {str(e)[:100]}"
            logger.error(f"❌ {error_msg}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            error_callback(f"❌ {error_msg}")
            error_occurred = True
        finally:
            if stderr_thread and stderr_thread.is_alive():
                stop_stderr.set()
                stderr_thread.join(timeout=1.0)
            self._flush_audio_buffer(transcription_callback, translation_callback)
            if process:
                self.ffmpeg_manager.stop_stream(self._current_stream_id)
            self._log_final_stats()
            self._guaranteed_cleanup()
            is_local_file = url.startswith("file://")
            normal_end = not self._stop_event.is_set() and not error_occurred
            if is_local_file and normal_end and self._finished_callback:
                logger.info("✅ Datei normal beendet – rufe finished_callback auf")
                try:
                    self._finished_callback()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(f"Fehler im finished_callback: {e}", exc_info=True)
            elif not self._stop_event.is_set() and not error_occurred:
                error_callback(
                    "❌ Stream wurde unerwartet beendet – versuche Neuverbindung..."
                )
            logger.info("✅ Processing loop ended")

    def _process_audio_chunk(
        self,
        audio_data: bytes,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
    ) -> None:
        if not self.transcription_engine:
            return

        start_time = None
        audio_len = 0.0
        if logger.isEnabledFor(logging.DEBUG):
            start_time = time.perf_counter()
            audio_len = len(audio_data) / (16000 * 2)

        if DEBUG_LEVEL >= 3:
            rms = self._calculate_rms(audio_data)
            logger.debug(f"Chunk {self._chunk_counter}: {len(audio_data)} bytes, RMS={rms:.4f}")

        try:
            with self._resource_lock:
                bytes_before_chunk = self._total_bytes_processed
            chunk_start_time = bytes_before_chunk / self.config.BYTES_PER_SECOND

            if logger.isEnabledFor(logging.DEBUG):
                chunk_start = time.perf_counter()

            transcribe_start = time.perf_counter()

            if self.subtitle_mode:
                self._handle_subtitle_transcription(
                    audio_data,
                    transcription_callback,
                    translation_callback,
                    chunk_start_time,
                )
            else:
                self._handle_normal_transcription(
                    audio_data,
                    transcription_callback,
                    translation_callback,
                    start_time,
                    audio_len,
                )

            transcribe_duration = time.perf_counter() - transcribe_start
            if audio_len > 0:
                self._last_realtime_factor = transcribe_duration / audio_len
            else:
                self._last_realtime_factor = 0.0

            if self._last_realtime_factor > 1.5:
                self._slow_chunks += 1
            else:
                self._slow_chunks = 0

            with self._stats_lock:
                self._consecutive_errors = 0
                self._consecutive_successes += 1

            if self.settings.adaptive_chunk:
                self._update_adaptive_chunk()

            if logger.isEnabledFor(logging.DEBUG):
                chunk_duration = (time.perf_counter() - chunk_start) * 1000
                log_debug(
                    "time",
                    f"Chunk {self._chunk_counter} processing took {chunk_duration:.2f}ms total (realtime factor: {self._last_realtime_factor:.2f})",
                )

            self._log_gpu_stats()
            self._log_queue_stats()
            self._log_cache_stats()
        except FutureTimeout as e:
            logger.warning(f"⏰ Timeout in audio chunk processing: {e}")
            with self._stats_lock:
                self._consecutive_errors += 1
                self._consecutive_successes = 0
        except (ValueError, TypeError, RuntimeError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Audio chunk processing error: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            with self._stats_lock:
                self._consecutive_errors += 1
                self._consecutive_successes = 0
                if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                    logger.critical(
                        f"🚨 Too many consecutive errors ({self._consecutive_errors}), stopping processing."
                    )
                    self._stop_event.set()
                    self._processing.clear()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(
                f"❌ Unexpected error in audio chunk processing: {e}", exc_info=True
            )
            with self._stats_lock:
                self._consecutive_errors += 1
                self._consecutive_successes = 0
                if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                    logger.critical(
                        f"🚨 Too many consecutive errors ({self._consecutive_errors}), stopping processing."
                    )
                    self._stop_event.set()
                    self._processing.clear()

    def _calculate_rms(self, audio_data: bytes) -> float:
        if self.transcription_engine and self.transcription_engine._np is not None:
            np = self.transcription_engine._np
            try:
                audio_np = np.frombuffer(audio_data, dtype=np.int16).astype(np.float32)
                rms = np.sqrt(np.mean(audio_np**2))
                return float(rms)
            except Exception:
                pass
        return 0.0

    def _handle_normal_transcription(
        self,
        audio_data: bytes,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
        start_time: float,
        audio_len: float,
    ) -> None:
        if logger.isEnabledFor(logging.DEBUG):
            log_debug(
                "thread",
                f"Starte Transkription in Thread {threading.current_thread().name}",
            )

        try:
            timeout_val = max(30, self.config.CHUNK_DURATION * 3)
            logger.debug(f"⏳ Normal-Transkription, Timeout={timeout_val}s")
            transcription = self._transcription_executor.submit_with_timeout(
                self.transcription_engine.safe_transcribe, timeout_val, audio_data
            )
        except TimeoutError as e:
            logger.error(
                f"⏰ safe_transcribe Timeout nach {self.config.CHUNK_DURATION*3}s: {e}"
            )
            with self._stats_lock:
                self._consecutive_timeouts += 1
                if self._consecutive_timeouts >= 3:
                    self._reload_model_on_timeout()
            return
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Transkriptionsfehler: {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            with self._stats_lock:
                self._consecutive_errors += 1
                if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                    logger.critical(
                        f"🚨 Zu viele Fehler ({self._consecutive_errors}), stoppe."
                    )
                    self._stop_event.set()
                    self._processing.clear()
            return

        with self._stats_lock:
            self._consecutive_timeouts = 0

        if transcription and hasattr(transcription, "confidence"):
            with self._last_confidence_lock:
                self.last_confidence = transcription.confidence
        else:
            with self._last_confidence_lock:
                self.last_confidence = 0.0

        if (
            logger.isEnabledFor(logging.DEBUG)
            and start_time is not None
            and transcription
        ):
            elapsed = time.perf_counter() - start_time
            realtime_factor = elapsed / audio_len if audio_len > 0 else 0
            logger.debug(
                f"Chunk {self._chunk_counter}: {audio_len:.2f}s audio, "
                f"transcribe {elapsed*1000:.1f}ms ({realtime_factor:.2f}x realtime)"
            )

        if not transcription or not transcription.text:
            return

        clean_text = transcription.text.strip()
        conf = getattr(transcription, "confidence", 0.0)

        with self._stats_lock:
            if conf < 0.4:
                self._low_conf_counter += 1
            else:
                self._low_conf_counter = 0

        if self.settings.enable_duplicate_check and self.config.DUPLICATE_CHECK_ENABLED:
            with self._duplicate_lock:
                if self._audio_enhancer.is_duplicate(
                    clean_text,
                    self._last_transcription_text,
                    list(self._recent_transcriptions),
                    confidence=conf,
                    last_confidence=self.last_confidence,
                ):
                    return
                self._last_transcription_text = clean_text
                self._recent_transcriptions.append(clean_text)

        with self._sentence_lock:
            self._sentence_buffer += " " + clean_text if self._sentence_buffer else clean_text
            self._sentence_segments.append(transcription)
            if clean_text and clean_text[-1] in ".!?":
                sentence = self._sentence_buffer.strip()
                if sentence:
                    if (
                        self.translation_engine
                        and self._translation_enabled.is_set()
                        and hasattr(transcription, "language")
                    ):
                        detected_lang = transcription.language or "auto"
                        if self._sentence_segments:
                            first = self._sentence_segments[0]
                            last = self._sentence_segments[-1]
                            self._translate_and_send_async(
                                sentence,
                                detected_lang,
                                translation_callback,
                                start=first.start if hasattr(first, "start") else None,
                                end=last.end if hasattr(last, "end") else None,
                            )
                self._sentence_buffer = ""
                self._sentence_segments.clear()
            else:
                if len(self._sentence_buffer.split()) > 50:
                    sentence = self._sentence_buffer.strip()
                    if sentence and self.translation_engine and self._translation_enabled.is_set():
                        detected_lang = transcription.language or "auto"
                        first = self._sentence_segments[0] if self._sentence_segments else None
                        last = transcription
                        self._translate_and_send_async(
                            sentence,
                            detected_lang,
                            translation_callback,
                            start=first.start if first and hasattr(first, "start") else None,
                            end=last.end if hasattr(last, "end") else None,
                        )
                    self._sentence_buffer = ""
                    self._sentence_segments.clear()

        transcription_callback(transcription)

        with self._stats_lock:
            self._consecutive_errors = 0

        word_count = len(clean_text.split())
        with self._word_count_lock:
            self._word_count_history.append(word_count)

    def _handle_subtitle_transcription(
        self,
        audio_data: bytes,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
        chunk_start_time: float,
    ) -> None:
        if logger.isEnabledFor(logging.DEBUG):
            log_debug(
                "thread",
                f"Starte Transkription in Thread {threading.current_thread().name}",
            )

        try:
            timeout_val = max(30, self.config.CHUNK_DURATION * 3)
            logger.debug(f"⏳ Subtitle-Transkription, Timeout={timeout_val}s")
            segments = self._transcription_executor.submit_with_timeout(
                self.transcription_engine.transcribe_audio,
                timeout_val,
                audio_data,
                True,
            )
        except TimeoutError as e:
            logger.error(
                f"⏰ Transkriptions-Timeout (Subtitle-Modus) nach {self.config.CHUNK_DURATION*3}s: {e}"
            )
            with self._stats_lock:
                self._consecutive_timeouts += 1
                if self._consecutive_timeouts >= 3:
                    self._reload_model_on_timeout()
            return
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Transkriptionsfehler (Subtitle-Modus): {e}")
            if logger.isEnabledFor(logging.DEBUG):
                logger.exception("Stacktrace:")
            with self._stats_lock:
                self._consecutive_errors += 1
                if self._consecutive_errors >= self.config.MAX_CONSECUTIVE_ERRORS:
                    logger.critical(
                        f"🚨 Zu viele Fehler ({self._consecutive_errors}), stoppe."
                    )
                    self._stop_event.set()
                    self._processing.clear()
            return

        if not segments:
            return

        if logger.isEnabledFor(logging.DEBUG):
            log_debug("subtitle", f"Received {len(segments)} segments")

        for segment in segments:
            if segment.start is not None:
                segment.start += chunk_start_time
            if segment.end is not None:
                segment.end += chunk_start_time

            logger.info(
                f"🎤 SEGMENT [{segment.start:.2f} - {segment.end:.2f}] {segment.text.strip()} "
                f"(Sprache: {getattr(segment, 'language', 'unbekannt')})"
            )
            if not segment or not segment.text:
                continue

            clean_text = segment.text.strip()
            conf = getattr(segment, "confidence", 0.0)

            with self._last_confidence_lock:
                self.last_confidence = conf

            with self._stats_lock:
                if conf < 0.4:
                    self._low_conf_counter += 1
                else:
                    self._low_conf_counter = 0

            if self.settings.enable_duplicate_check and self.config.DUPLICATE_CHECK_ENABLED and not self.subtitle_mode:
                with self._duplicate_lock:
                    if self._audio_enhancer.is_duplicate(
                        clean_text,
                        self._last_transcription_text,
                        list(self._recent_transcriptions),
                        confidence=conf,
                        last_confidence=self.last_confidence,
                    ):
                        continue
                    self._last_transcription_text = clean_text
                    self._recent_transcriptions.append(clean_text)

            if not self.transcription_engine.is_valid_segment(clean_text, conf):
                continue

            if self.config.ENABLE_TIMED_TRANSCRIPTIONS:
                with self._subtitle_lock:
                    self._timed_transcriptions.append(segment)

            transcription_callback(segment)

            word_count = len(clean_text.split())
            with self._word_count_lock:
                self._word_count_history.append(word_count)

            with self._sentence_lock:
                self._sentence_buffer += " " + clean_text if self._sentence_buffer else clean_text
                self._sentence_segments.append(segment)
                if clean_text and clean_text[-1] in ".!?":
                    sentence = self._sentence_buffer.strip()
                    if sentence and self.translation_engine and self._translation_enabled.is_set():
                        detected_lang = segment.language or "auto"
                        first = self._sentence_segments[0]
                        last = segment
                        self._translate_and_send_async(
                            sentence,
                            detected_lang,
                            translation_callback,
                            start=first.start,
                            end=last.end,
                        )
                    self._sentence_buffer = ""
                    self._sentence_segments.clear()
                else:
                    if len(self._sentence_buffer.split()) > 50:
                        sentence = self._sentence_buffer.strip()
                        if sentence and self.translation_engine and self._translation_enabled.is_set():
                            detected_lang = segment.language or "auto"
                            first = self._sentence_segments[0]
                            last = segment
                            self._translate_and_send_async(
                                sentence,
                                detected_lang,
                                translation_callback,
                                start=first.start,
                                end=last.end,
                            )
                        self._sentence_buffer = ""
                        self._sentence_segments.clear()

        with self._stats_lock:
            self._consecutive_timeouts = 0
            self._consecutive_errors = 0

    def _translate_and_send_async(
        self,
        text: str,
        source_lang: str,
        translation_callback: TranslationCallback,
        start: Optional[float] = None,
        end: Optional[float] = None,
    ) -> None:
        if not self._translation_enabled.is_set():
            return

        if (
            source_lang != "auto"
            and source_lang == self.translation_engine.default_target_lang
        ):
            return

        def task():
            try:
                start_time = time.perf_counter()
                translation = None
                primary_functional = True
                if hasattr(self.translation_engine, "is_functional"):
                    primary_functional = self.translation_engine.is_functional()
                if primary_functional:
                    try:
                        translation = self.translation_engine.translate_text(
                            text, source_lang
                        )
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        log_debug("translate", f"Primary translation error: {e}")
                        if logger.isEnabledFor(logging.DEBUG):
                            logger.exception("Stacktrace:")
                if translation is None and self._fallback_translation_engine:
                    fallback_functional = True
                    if hasattr(self._fallback_translation_engine, "is_functional"):
                        fallback_functional = (
                            self._fallback_translation_engine.is_functional()
                        )
                    if fallback_functional:
                        try:
                            translation = (
                                self._fallback_translation_engine.translate_text(
                                    text, source_lang
                                )
                            )
                        except Exception as e:
                            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                                raise
                            log_debug("translate", f"Fallback error: {e}")
                duration = (time.perf_counter() - start_time) * 1000
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("time", f"_translate_and_send took {duration:.2f}ms")
                if translation:
                    translation.start = start
                    translation.end = end
                    if self.subtitle_mode and self.config.ENABLE_TIMED_TRANSLATIONS:
                        with self._subtitle_lock:
                            self._timed_translations.append(translation)
                    translation_callback(translation)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"❌ Unerwarteter Fehler in _translate_and_send: {e}", exc_info=True
                )

        try:
            self._translation_executor.submit(task)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"Fehler beim Einreichen der Übersetzungsaufgabe: {e}")

    def _process_audio_data(
        self,
        audio_data: bytes,
        transcription_callback: TranscriptionCallback,
        translation_callback: TranslationCallback,
        info_callback: InfoCallback,
        error_callback: ErrorCallback,
    ) -> None:
        with self._stats_lock:
            chunk_num = self._chunk_counter
        if chunk_num <= 3:
            logger.debug(f"📦 Chunk #{chunk_num}: {len(audio_data)} bytes")

        if len(audio_data) > 0:
            is_silent = self._is_silent(audio_data)
            if is_silent:
                self._silent_chunk_counter += 1
                if self._silent_chunk_counter % 10 == 0:
                    logger.debug(f"Still receiving silent chunks ({self._silent_chunk_counter} in a row)")
            else:
                self._silent_chunk_counter = 0

        if self._progress_callback:
            now = time.time()
            if now - self._last_progress_update >= self._progress_update_interval:
                self._last_progress_update = now
                try:
                    with self._stats_lock:
                        total_bytes = self._total_bytes_processed
                    self._progress_callback(
                        total_bytes,
                        self._total_file_size,
                        chunk_num,
                    )
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(f"Fehler im progress_callback: {e}", exc_info=True)

        with self._last_confidence_lock:
            last_conf = self.last_confidence

        apply_enhancement = (last_conf < 0.3)
        if apply_enhancement:
            try:
                with self._noisereduce_lock:
                    enhanced_audio = self._audio_enhancer.enhance_audio(
                        audio_data, last_conf, self._noisereduce_counter
                    )
                    self._noisereduce_counter = (self._noisereduce_counter + 1) % 10000
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Audio enhancement failed: {e}")
                enhanced_audio = audio_data
        else:
            enhanced_audio = audio_data

        with self._buffer_lock:
            self._audio_buffer.extend(enhanced_audio)
            if len(self._audio_buffer) > self._max_buffer_bytes:
                excess = len(self._audio_buffer) - self._max_buffer_bytes
                self._audio_buffer = self._audio_buffer[excess:]
                logger.debug(f"⚠️ Audio buffer truncated by {excess} bytes (max {self._max_buffer_bytes} bytes)")

            if (
                len(self._audio_buffer) > 0
                and time.time() - self._last_buffer_flush > 30
            ):
                logger.debug("⏱️ Flushing audio buffer after inactivity")
                chunk_to_process = bytes(self._audio_buffer)
                self._audio_buffer.clear()
                if self.transcription_engine:
                    self._process_audio_chunk(
                        chunk_to_process, transcription_callback, translation_callback
                    )
                self._last_buffer_flush = time.time()

            if len(self._audio_buffer) >= self.config.MIN_CHUNK_BYTES:
                chunk_to_process = bytes(self._audio_buffer)
                self._audio_buffer.clear()
                if self.transcription_engine:
                    self._process_audio_chunk(
                        chunk_to_process, transcription_callback, translation_callback
                    )
                self._last_buffer_flush = time.time()

            if len(self._audio_buffer) > self._max_buffer_size:
                logger.warning(
                    f"⚠️ Audio buffer too large ({len(self._audio_buffer)} bytes) – forcing flush"
                )
                chunk_to_process = bytes(self._audio_buffer)
                self._audio_buffer.clear()
                if self.transcription_engine:
                    self._process_audio_chunk(
                        chunk_to_process, transcription_callback, translation_callback
                    )
                self._last_buffer_flush = time.time()

        with self._stats_lock:
            self._chunk_counter += 1
            self._total_bytes_processed += len(audio_data)
            self._processed_seconds = (
                self._total_bytes_processed / self.config.BYTES_PER_SECOND
            )

        if chunk_num % 50 == 0:
            info_callback(f"📊 {chunk_num} chunks processed...")

    def _is_silent(self, audio_data: bytes, threshold: float = 0.001) -> bool:
        if len(audio_data) < 160:
            return False
        if self.transcription_engine and self.transcription_engine._np is not None:
            np = self.transcription_engine._np
            try:
                audio_np = np.frombuffer(audio_data, dtype=np.int16).astype(np.float32)
                rms = np.sqrt(np.mean(audio_np**2))
                return rms < threshold
            except Exception:
                pass
        return False

    def _update_adaptive_chunk(self) -> None:
        if self.subtitle_mode:
            return
        with self._stats_lock:
            with self._word_count_lock:
                if len(self._word_count_history) < Constants.ADAPTIVE_CHUNK_MIN_SAMPLES:
                    return
                avg_words = sum(self._word_count_history) / len(
                    self._word_count_history
                )
            if self._smoothed_word_count is None:
                self._smoothed_word_count = avg_words
            else:
                alpha = Constants.ADAPTIVE_CHUNK_SMOOTHING_ALPHA
                self._smoothed_word_count = (
                    alpha * avg_words + (1 - alpha) * self._smoothed_word_count
                )
            smoothed = self._smoothed_word_count
            new_duration = self.config.CHUNK_DURATION
            low_thresh = self.settings.adaptive_chunk_low_words
            high_thresh = self.settings.adaptive_chunk_high_words
            min_dur = self.config.MIN_CHUNK_DURATION
            max_dur = self.config.MAX_CHUNK_DURATION

            if self._last_realtime_factor > 1.5 and self.config.CHUNK_DURATION > min_dur + 0.5:
                new_duration = max(min_dur, self.config.CHUNK_DURATION - 1)
                logger.debug(f"Adaptive: Realtime-Faktor {self._last_realtime_factor:.2f} > 1.5, reduziere Chunk")
            elif smoothed < low_thresh and self.config.CHUNK_DURATION > min_dur + 0.5:
                new_duration = max(min_dur, self.config.CHUNK_DURATION - 1)
            elif smoothed > high_thresh and self.config.CHUNK_DURATION < max_dur - 0.5:
                new_duration = min(max_dur, self.config.CHUNK_DURATION + 1)

            if new_duration != self.config.CHUNK_DURATION:
                if new_duration != self._last_chunk_duration:
                    self._chunk_stable_counter += 1
                else:
                    self._chunk_stable_counter = 0
                if (
                    self._chunk_stable_counter
                    >= Constants.ADAPTIVE_CHUNK_STABLE_THRESHOLD
                ):
                    logger.info(
                        f"{'📈' if new_duration > self.config.CHUNK_DURATION else '📉'} Adaptive Chunk-Dauer: "
                        f"{self.config.CHUNK_DURATION:.1f}s → {new_duration:.1f}s (avg_words={smoothed:.1f}, realtime={self._last_realtime_factor:.2f})"
                    )
                    self.config.CHUNK_DURATION = new_duration
                    self._update_chunk_size()
                    self._chunk_stable_counter = 0
                self._last_chunk_duration = new_duration

    def _reload_model_on_timeout(self):
        logger.warning("🔄 Drei aufeinanderfolgende Timeouts – lade Modell neu...")
        if hasattr(self.transcription_engine, "reload_model"):
            current_model = self.transcription_engine.get_current_model()
            try:
                self.transcription_engine.reload_model(current_model)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Neuladen des Modells: {e}", exc_info=True)
        with self._stats_lock:
            self._consecutive_timeouts = 0
            self._consecutive_errors = 0
        time.sleep(2)

    def _read_stderr(self, process: subprocess.Popen, max_bytes: int = 4096) -> str:
        try:
            if process.stderr:
                return process.stderr.read(max_bytes).decode("utf-8", errors="ignore")
        except (OSError, ValueError) as e:
            logger.debug(f"Fehler beim Lesen von stderr: {e}")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(
                f"Unerwarteter Fehler beim Lesen von stderr: {e}", exc_info=True
            )
        return ""

    def emergency_diagnosis(self, url: str) -> bool:
        logger.info(f"🔍 [EMERGENCY_DIAGNOSIS] Testing: {url[:80]}...")
        try:
            audio_url = self.stream_manager.extract_audio_url(url)
            if not audio_url:
                logger.info("  ❌ Could not extract audio URL")
                return False
            logger.info(f"  ✅ Audio URL extracted: {audio_url[:80]}...")
            is_youtube = (
                "youtube.com" in audio_url.lower() or "googlevideo.com" in audio_url
            )
            test_cmd = [
                "ffmpeg",
                "-i",
                audio_url,
                "-t",
                "3",
                "-f",
                "null",
                "-",
                "-loglevel",
                "error",
            ]
            timeout = (
                Constants.YOUTUBE_STREAM_TEST_TIMEOUT
                if is_youtube
                else Constants.STREAM_TEST_TIMEOUT
            )
            try:
                result = subprocess.run(
                    test_cmd,
                    capture_output=True,
                    text=True,
                    timeout=timeout,
                    shell=False,
                )
            except subprocess.TimeoutExpired as e:
                logger.info(f"  ⏰ Stream test timeout: {e}")
                if "youtube.com" in url.lower():
                    logger.info("  ⚠️  YouTube timeout common, trying anyway...")
                    return True
                return False
            except (OSError, PermissionError) as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.info(f"  ⚠️  Emergency diagnosis OS error: {e}")
                return False
            if result.returncode == 0:
                logger.info("  ✅ Stream connection successful")
                return True
            else:
                error_msg = result.stderr[:100] if result.stderr else "Unknown error"
                logger.info(f"  ❌ Stream test failed: {error_msg}")
                if audio_url.startswith(("http://", "https://")):
                    logger.info("  ⚠️  But URL looks valid, trying anyway...")
                    return True
                return False
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.info(f"  ⚠️  Emergency diagnosis error: {e}")
            return True

    def _log_final_stats(self) -> None:
        if self._chunk_counter == 0:
            return
        uptime = time.time() - self._stream_start_time if self._stream_start_time else 0
        logger.info("\n📊 FINAL PROCESSING STATS:")
        logger.info(f"   Config Type: {self._get_config_type()}")
        logger.info(f"   Total Chunks: {self._chunk_counter}")
        logger.info(f"   Total Bytes: {self._total_bytes_processed:,}")
        logger.info(f"   Total Time: {uptime:.1f}s")
        logger.info(
            f"   Avg Chunk Size: {self._total_bytes_processed/self._chunk_counter if self._chunk_counter > 0 else 0:,.0f} bytes"
        )
        logger.info(
            f"   Processing Rate: {self._chunk_counter/uptime if uptime > 0 else 0:.1f} chunks/sec"
        )
        logger.info(
            f"   Data Rate: {self._total_bytes_processed/uptime/1024 if uptime > 0 else 0:.1f} KB/sec"
        )
        logger.info(f"   Empty Reads: {self._empty_reads}")
        logger.info(f"   Last Realtime Factor: {self._last_realtime_factor:.2f}")

    def _log_gpu_stats(self) -> None:
        if not logger.isEnabledFor(logging.DEBUG):
            return
        if self._chunk_counter % 5 == 0:
            if TORCH_AVAILABLE and self.transcription_engine.device == "cuda":
                torch = FastLazyLoader.load("torch")
                try:
                    allocated = torch.cuda.memory_allocated() / 1024**3
                    reserved = torch.cuda.memory_reserved() / 1024**3
                    log_debug(
                        "gpu",
                        f"Chunk {self._chunk_counter}: allocated={allocated:.2f}GB, reserved={reserved:.2f}GB",
                    )
                except (RuntimeError, AttributeError) as e:
                    log_debug("gpu", f"GPU-Statistik-Fehler: {e}")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Unerwarteter Fehler bei GPU-Statistik: {e}", exc_info=True
                    )
                if (
                    hasattr(self, "_last_gpu_stats_time")
                    and time.time() - self._last_gpu_stats_time > 10
                ):
                    self._last_gpu_stats_time = time.time()
                    try:
                        import pynvml

                        pynvml.nvmlInit()
                        handle = pynvml.nvmlDeviceGetHandleByIndex(0)
                        temp = pynvml.nvmlDeviceGetTemperature(
                            handle, pynvml.NVML_TEMPERATURE_GPU
                        )
                        util = pynvml.nvmlDeviceGetUtilizationRates(handle).gpu
                        log_debug("gpu", f"Temp={temp}°C, Util={util}%")
                    except (ImportError, pynvml.NVMLError) as e:
                        log_debug("gpu", f"pynvml-Fehler: {e}")
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.error(
                            f"Unerwarteter Fehler bei pynvml: {e}", exc_info=True
                        )

    def _log_queue_stats(self) -> None:
        if not logger.isEnabledFor(logging.DEBUG):
            return
        if self._chunk_counter % 50 == 0:
            gui = (
                self.controller_ref.gui_ref()
                if hasattr(self.controller_ref, "gui_ref")
                else None
            )
            if gui is not None:
                try:
                    qsize_gui = gui.gui_queue.qsize()
                    qsize_text = gui._text_update_queue.qsize()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Lesen der Queue-Größen: {e}", exc_info=True
                    )
                    qsize_gui = qsize_text = 0
            else:
                qsize_gui = qsize_text = 0
            active_threads = threading.active_count()
            thread_names = [t.name for t in threading.enumerate()]
            log_debug(
                "queue",
                f"Chunk {self._chunk_counter}: gui_queue={qsize_gui}, text_queue={qsize_text}, active_threads={active_threads}",
            )
            if logger.isEnabledFor(logging.DEBUG):
                log_debug("threads", f"Thread names: {thread_names}")

    def _log_cache_stats(self) -> None:
        if not logger.isEnabledFor(logging.DEBUG):
            return
        if self._chunk_counter % 100 == 0:
            try:
                stats = AppContext().cache_manager.get_stats()
                log_debug("cache", f"Cache stats: {stats}")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(
                    f"Fehler beim Abrufen der Cache-Statistiken: {e}", exc_info=True
                )

    def _platform_specific_health_check(self) -> List[str]:
        issues: List[str] = []
        if IS_WINDOWS:
            try:
                import psutil

                memory = psutil.virtual_memory()
                if memory.percent > 85:
                    issues.append(
                        "High memory usage - consider closing other applications"
                    )
            except ImportError:
                pass
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Fehler beim Health-Check: {e}", exc_info=True)
        return issues

    def _flush_audio_buffer(
        self, transcription_callback: TranscriptionCallback, translation_callback: TranslationCallback
    ) -> None:
        with self._buffer_lock:
            if not self._audio_buffer:
                return
            buffer_len = len(self._audio_buffer)
            logger.info(
                f"🧹 Flushing audio buffer ({buffer_len} bytes) at end of stream"
            )
            if self.transcription_engine:
                self._process_audio_chunk(
                    bytes(self._audio_buffer),
                    transcription_callback,
                    translation_callback,
                )
            self._audio_buffer.clear()

    def _test_audio_stream(self, audio_url: str) -> bool:
        logger.info(f"🔍 Testing audio stream: {audio_url[:60]}...")
        is_youtube = (
            "youtube.com" in audio_url.lower() or "googlevideo.com" in audio_url
        )
        is_hls = ".m3u8" in audio_url.lower() or "manifest.googlevideo.com" in audio_url
        if is_hls:
            logger.info("🎯 HLS stream detected – skipping quick test (often too slow)")
            return True
        try:
            timeout = (
                Constants.YOUTUBE_STREAM_TEST_TIMEOUT
                if is_youtube
                else self.config.STREAM_TIMEOUT
            )
            test_cmd = [
                "ffmpeg",
                "-i",
                audio_url,
                "-t",
                "2",
                "-f",
                "null",
                "-",
                "-loglevel",
                "error",
            ]
            result = subprocess.run(
                test_cmd, capture_output=True, text=True, timeout=timeout, shell=False
            )
            if result.returncode == 0:
                logger.info("✅ Stream test successful")
                return True
            else:
                error_msg = result.stderr[:100] if result.stderr else "Unknown error"
                logger.error(f"❌ Stream test failed: {error_msg}")
                return False
        except subprocess.TimeoutExpired as e:
            logger.warning(f"⏰ Stream test timeout after {timeout}s: {e}")
            return False
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Stream test error: {e}")
            return True

    def get_timed_transcriptions(self) -> List["TranscriptionResult"]:
        with self._subtitle_lock:
            return list(self._timed_transcriptions)

    def get_timed_translations(self) -> List["TranslationResult"]:
        with self._subtitle_lock:
            return list(self._timed_translations)

    def get_status(self) -> Dict[str, Any]:
        with self._stats_lock:
            return {
                "processing": self._processing.is_set(),
                "stop_event_set": self._stop_event.is_set(),
                "current_stream_id": self._current_stream_id,
                "consecutive_empty_chunks": self._consecutive_empty_chunks,
                "cleanup_done": self._cleanup_done,
                "config_type": self._get_config_type(),
                "chunk_size": self.chunk_size,
                "chunks_processed": self._chunk_counter,
                "total_bytes": self._total_bytes_processed,
                "empty_reads": self._empty_reads,
                "subtitle_mode": self.subtitle_mode,
                "translation_active": self._translation_enabled.is_set(),
                "active_threads": threading.active_count(),
                "consecutive_errors": self._consecutive_errors,
                "consecutive_timeouts": self._consecutive_timeouts,
                "low_conf_counter": self._low_conf_counter,
                "slow_chunks": self._slow_chunks,
                "last_realtime_factor": self._last_realtime_factor,
            }

    def stop_processing(self, wait: bool = False, timeout: float = 5.0) -> bool:
        logger.info("🛑 AudioProcessor: Stopping processing...")
        with self._processing_lock:
            self._stop_event.set()
            self._processing.clear()
        if self._current_stream_id:
            logger.info(f"📛 Stream {self._current_stream_id} stopped by user")
            if self.ffmpeg_manager:
                try:
                    self.ffmpeg_manager.stop_stream(self._current_stream_id)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(f"Fehler beim Stoppen des Streams: {e}", exc_info=True)
        return True

    def emergency_reset(self, force: bool = False) -> bool:
        logger.info(f"\n🚨 [EMERGENCY_RESET] force={force}")
        with self._resource_lock:
            with self._processing_lock:
                self._processing.clear()
            self._stop_event.set()
            self._current_stream_id = None
            self._consecutive_empty_chunks = 0
            with self._stats_lock:
                self._consecutive_errors = 0
                self._consecutive_timeouts = 0
                self._low_conf_counter = 0
                self._slow_chunks = 0
                self._last_realtime_factor = 0.0
            if force:
                with self._subtitle_lock:
                    self._timed_transcriptions.clear()
                    self._timed_translations.clear()
                with self._duplicate_lock:
                    self._recent_transcriptions.clear()
                    self._last_transcription_text = ""
            with self._buffer_lock:
                self._audio_buffer.clear()
        logger.info("✅ Reset completed")
        return True

    def _guaranteed_cleanup(self) -> None:
        logger.info("\n🧹 [GUARANTEED_CLEANUP]")
        with self._resource_lock:
            with self._processing_lock:
                self._processing.clear()
            self._current_stream_id = None
            self._consecutive_empty_chunks = 0
            self._empty_reads = 0
            with self._stats_lock:
                self._chunk_counter = 0
                self._total_bytes_processed = 0
                self._consecutive_errors = 0
                self._consecutive_timeouts = 0
                self._low_conf_counter = 0
                self._slow_chunks = 0
                self._last_realtime_factor = 0.0
            with self._buffer_lock:
                self._audio_buffer.clear()
            self._cleanup_done = True
        time.sleep(0.05)
        logger.info("✅ Cleanup completed")

    def dispose(self) -> None:
        logger.info("🧹 AudioProcessor: Starting dispose...")
        try:
            if hasattr(self, "_stop_event") and hasattr(self._stop_event, "set"):
                self._stop_event.set()
            with self._processing_lock:
                if hasattr(self, "_processing") and hasattr(self._processing, "clear"):
                    self._processing.clear()
            self._cleanup_done = True
            if hasattr(self, "ffmpeg_manager") and self.ffmpeg_manager:
                try:
                    self.ffmpeg_manager.stop_all_streams()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Stoppen aller Streams: {e}", exc_info=True
                    )
            if hasattr(self, "_transcription_executor"):
                try:
                    self._transcription_executor.shutdown(wait=True)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Shutdown des Transcription-Executors: {e}",
                        exc_info=True,
                    )
            if hasattr(self, "_translation_executor"):
                try:
                    self._translation_executor.shutdown(wait=True)
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.error(
                        f"Fehler beim Shutdown des Translation-Executors: {e}",
                        exc_info=True,
                    )
            with self._subtitle_lock:
                self._timed_transcriptions.clear()
                self._timed_translations.clear()
            with self._duplicate_lock:
                self._recent_transcriptions.clear()
                self._last_transcription_text = ""
            with self._buffer_lock:
                self._audio_buffer.clear()
            gc.collect()
            logger.info("✅ AudioProcessor disposed")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"⚠️ AudioProcessor dispose error: {e}", exc_info=True)


# -----------------------------------------------------------------------------
# Themes
# -----------------------------------------------------------------------------
class DarkTheme:
    BG_PRIMARY = "#0f1419"
    BG_SECONDARY = "#1a2129"
    BG_TERTIARY = "#242d38"
    BG_HOVER = "#2a3645"
    BG_CARD = "#1c252f"
    TEXT_PRIMARY = "#e6edf3"
    TEXT_SECONDARY = "#8b949e"
    TEXT_ACCENT = "#58a6ff"
    TEXT_MUTED = "#6e7681"
    DRAGON_GREEN = "#238636"
    DRAGON_GREEN_LIGHT = "#2ea043"
    DRAGON_BLUE = "#1f6feb"
    DRAGON_PURPLE = "#8957e5"
    SUCCESS = "#238636"
    WARNING = "#d29922"
    ERROR = "#f85149"
    INFO = "#58a6ff"
    BORDER = "#30363d"
    SCROLLBAR = "#3c444d"
    SCROLLBAR_HOVER = "#4c5560"
    INPUT_BG = "#161b22"
    INPUT_BORDER = "#30363d"
    INPUT_FOCUS = "#1f6feb"
    COMBO_BG = "#161b22"
    COMBO_FG = "#e6edf3"
    COMBO_BORDER = "#30363d"
    COMBO_SELECTION = "#1f6feb"
    CHECKBOX_BG = "#0d1117"
    CHECKBOX_FG = "#e6edf3"
    CHECKBOX_SELECTED = "#238636"
    CHECKBOX_ACTIVE = "#000000"
    SUBTITLE_ACTIVE = "#8957e5"
    SUBTITLE_INACTIVE = "#30363d"
    STATUS_BAR_BG = "#0d1117"
    STATUS_BAR_FG = "#8b949e"
    STATUS_BAR_ACCENT = "#58a6ff"


class PastelTheme:
    BG_PRIMARY = "#e6f0fa"
    BG_SECONDARY = "#d4e3f0"
    BG_TERTIARY = "#c2d6e8"
    BG_HOVER = "#b0c9e0"
    BG_CARD = "#deecf5"
    TEXT_PRIMARY = "#1a2b3c"
    TEXT_SECONDARY = "#2c3e50"
    TEXT_ACCENT = "#1f6feb"
    TEXT_MUTED = "#5a6b7a"
    DRAGON_GREEN = "#2ecc71"
    DRAGON_GREEN_LIGHT = "#27ae60"
    DRAGON_BLUE = "#3498db"
    DRAGON_PURPLE = "#9b59b6"
    SUCCESS = "#2ecc71"
    WARNING = "#f39c12"
    ERROR = "#e74c3c"
    INFO = "#3498db"
    BORDER = "#b0c4de"
    SCROLLBAR = "#a0b8d0"
    SCROLLBAR_HOVER = "#90a8c0"
    INPUT_BG = "#ffffff"
    INPUT_BORDER = "#b0c4de"
    INPUT_FOCUS = "#1f6feb"
    COMBO_BG = "#ffffff"
    COMBO_FG = "#1a2b3c"
    COMBO_BORDER = "#b0c4de"
    COMBO_SELECTION = "#1f6feb"
    CHECKBOX_BG = "#f0f8ff"
    CHECKBOX_FG = "#1a2b3c"
    CHECKBOX_SELECTED = "#2ecc71"
    CHECKBOX_ACTIVE = "#ffffff"
    SUBTITLE_ACTIVE = "#8957e5"
    SUBTITLE_INACTIVE = "#b0c4de"
    STATUS_BAR_BG = "#d4e3f0"
    STATUS_BAR_FG = "#2c3e50"
    STATUS_BAR_ACCENT = "#1f6feb"


class LightTheme:
    BG_PRIMARY = "#ffffff"
    BG_SECONDARY = "#f0f0f0"
    BG_TERTIARY = "#e5e5e5"
    BG_HOVER = "#d9d9d9"
    BG_CARD = "#f5f5f5"
    TEXT_PRIMARY = "#000000"
    TEXT_SECONDARY = "#4a4a4a"
    TEXT_ACCENT = "#1f6feb"
    TEXT_MUTED = "#666666"
    DRAGON_GREEN = "#2ecc71"
    DRAGON_GREEN_LIGHT = "#27ae60"
    DRAGON_BLUE = "#3498db"
    DRAGON_PURPLE = "#9b59b6"
    SUCCESS = "#2ecc71"
    WARNING = "#f39c12"
    ERROR = "#e74c3c"
    INFO = "#3498db"
    BORDER = "#cccccc"
    SCROLLBAR = "#b3b3b3"
    SCROLLBAR_HOVER = "#999999"
    INPUT_BG = "#ffffff"
    INPUT_BORDER = "#cccccc"
    INPUT_FOCUS = "#1f6feb"
    COMBO_BG = "#ffffff"
    COMBO_FG = "#000000"
    COMBO_BORDER = "#cccccc"
    COMBO_SELECTION = "#1f6feb"
    CHECKBOX_BG = "#f0f0f0"
    CHECKBOX_FG = "#000000"
    CHECKBOX_SELECTED = "#2ecc71"
    CHECKBOX_ACTIVE = "#ffffff"
    SUBTITLE_ACTIVE = "#8957e5"
    SUBTITLE_INACTIVE = "#b3b3b3"
    STATUS_BAR_BG = "#f0f0f0"
    STATUS_BAR_FG = "#4a4a4a"
    STATUS_BAR_ACCENT = "#1f6feb"


class HighContrastTheme:
    BG_PRIMARY = "#000000"
    BG_SECONDARY = "#1a1a1a"
    BG_TERTIARY = "#333333"
    BG_HOVER = "#444444"
    BG_CARD = "#222222"
    TEXT_PRIMARY = "#ffff00"
    TEXT_SECONDARY = "#ffffff"
    TEXT_ACCENT = "#00ffff"
    TEXT_MUTED = "#cccccc"
    DRAGON_GREEN = "#00ff00"
    DRAGON_GREEN_LIGHT = "#00ff00"
    DRAGON_BLUE = "#0000ff"
    DRAGON_PURPLE = "#ff00ff"
    SUCCESS = "#00ff00"
    WARNING = "#ffff00"
    ERROR = "#ff0000"
    INFO = "#00ffff"
    BORDER = "#ffffff"
    SCROLLBAR = "#ffff00"
    SCROLLBAR_HOVER = "#ffffff"
    INPUT_BG = "#000000"
    INPUT_BORDER = "#ffff00"
    INPUT_FOCUS = "#00ffff"
    COMBO_BG = "#000000"
    COMBO_FG = "#ffff00"
    COMBO_BORDER = "#ffffff"
    COMBO_SELECTION = "#00ffff"
    CHECKBOX_BG = "#000000"
    CHECKBOX_FG = "#ffff00"
    CHECKBOX_SELECTED = "#00ff00"
    CHECKBOX_ACTIVE = "#ffffff"
    SUBTITLE_ACTIVE = "#00ffff"
    SUBTITLE_INACTIVE = "#666666"
    STATUS_BAR_BG = "#000000"
    STATUS_BAR_FG = "#ffff00"
    STATUS_BAR_ACCENT = "#00ffff"


CURRENT_THEME = DarkTheme()


class Fonts:
    TITLE = ("Segoe UI", 12, "bold")
    SUBTITLE = ("Segoe UI", 10, "bold")
    PRIMARY = ("Segoe UI", 9)
    SECONDARY = ("Segoe UI", 8)
    MONOSPACE = ("Cascadia Code", 9)
    BUTTON = ("Segoe UI", 9, "bold")
    STATUS = ("Segoe UI", 8)
    SMALL = ("Segoe UI", 7)


# -----------------------------------------------------------------------------
# DarkMessageBox
# -----------------------------------------------------------------------------
class DarkMessageBox:
    __slots__ = ()

    @classmethod
    def showinfo(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        return cls._show_dialog(title, message, "info", parent)

    @classmethod
    def showwarning(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        return cls._show_dialog(title, message, "warning", parent)

    @classmethod
    def showerror(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        return cls._show_dialog(title, message, "error", parent)

    @classmethod
    def askokcancel(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        return cls._show_dialog(title, message, "question", parent, buttons=True)

    @classmethod
    def askyesno(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        return cls._ask_yesno(title, message, parent)

    @classmethod
    def show_progress(
        cls,
        title: str,
        message: str,
        parent: Optional[tk.Tk] = None,
        indeterminate: bool = True,
    ) -> Any:
        if threading.current_thread() is not threading.main_thread():
            warnings.warn(
                "DarkMessageBox.show_progress wurde nicht im Hauptthread aufgerufen. "
                "GUI-Updates können fehlschlagen.",
                RuntimeWarning,
                stacklevel=2,
            )

        def _create():
            nonlocal dialog, progress, message_label
            try:
                root = parent if parent and parent.winfo_exists() else tk._default_root
                if not root or not root.winfo_exists():
                    logger.error("Kein gültiges root-Fenster für ProgressDialog")
                    return

                dlg = tk.Toplevel(root)
                dlg.title(f"🐉 {title}")
                dlg.configure(bg=CURRENT_THEME.BG_PRIMARY)
                dlg.resizable(False, False)
                dlg.transient(root)
                dlg.grab_set()

                main = tk.Frame(dlg, bg=CURRENT_THEME.BG_PRIMARY, padx=30, pady=25)
                main.pack(fill="both", expand=True)

                msg_lbl = tk.Label(
                    main,
                    text=message,
                    font=Fonts.PRIMARY,
                    bg=CURRENT_THEME.BG_PRIMARY,
                    fg=CURRENT_THEME.TEXT_PRIMARY,
                    justify="center",
                )
                msg_lbl.pack(pady=(0, 20))

                prog = ttk.Progressbar(
                    main,
                    mode="indeterminate" if indeterminate else "determinate",
                    length=300,
                )
                prog.pack(pady=(0, 10))
                prog.start(10)

                cls._center_dialog(dlg, root)

                dialog = dlg
                progress = prog
                message_label = msg_lbl
            except (tk.TclError, RuntimeError) as e:
                logger.warning(f"⚠️ Progress Dialog Error: {e}")

        dialog = None
        progress = None
        message_label = None

        _create()

        class ProgressController:
            __slots__ = ("dialog", "progress", "_message_label")

            def __init__(self):
                self.dialog = dialog
                self.progress = progress
                self._message_label = message_label

            def close(self) -> None:
                if self.dialog and self.dialog.winfo_exists():
                    try:
                        if self.progress:
                            self.progress.stop()
                        self.dialog.destroy()
                    except Exception:
                        pass
                self.dialog = None
                self.progress = None
                self._message_label = None

            def update_message(self, new_message: str) -> None:
                if self.dialog and self.dialog.winfo_exists() and self._message_label:
                    try:
                        self._message_label.config(text=new_message)
                    except Exception:
                        pass

        return ProgressController()

    @classmethod
    def _show_dialog(
        cls,
        title: str,
        message: str,
        msg_type: str,
        parent: Optional[tk.Tk] = None,
        buttons: bool = False,
    ) -> Optional[bool]:
        try:
            parent = cls._resolve_parent(parent)
            if parent is None:
                return cls._fallback_messagebox(title, message, msg_type, buttons)

            dialog, result, timeout_id = cls._create_base_dialog(
                parent, title, msg_type, message
            )

            if buttons:
                cls._add_ok_cancel_buttons(dialog, result, timeout_id, title)
            else:
                cls._add_ok_button(dialog, result, timeout_id)

            dialog.protocol(
                "WM_DELETE_WINDOW",
                lambda: (
                    cls._cancel_timeout(dialog, timeout_id),
                    cls._on_closing(result, False if buttons else True),
                ),
            )
            cls._center_dialog(dialog, parent)
            parent.wait_window(dialog)
            return result["value"]

        except (tk.TclError, RuntimeError, AttributeError) as e:
            logger.warning(f"⚠️ DarkMessageBox Error: {e}")
            return cls._fallback_messagebox(title, message, msg_type, buttons)

    @classmethod
    def _ask_yesno(
        cls, title: str, message: str, parent: Optional[tk.Tk] = None
    ) -> Optional[bool]:
        try:
            parent = cls._resolve_parent(parent)
            if parent is None:
                import tkinter.messagebox as mb

                return mb.askyesno(title, message)

            dialog, result, timeout_id = cls._create_base_dialog(
                parent, title, "question", message, icon_only=True
            )

            btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
            btn_frame.pack(fill="x", pady=(10, 0))

            def set_res(val: bool):
                cls._cancel_timeout(dialog, timeout_id)
                result["value"] = val
                if dialog.winfo_exists():
                    dialog.destroy()

            yes_btn = tk.Button(
                btn_frame,
                text="  👍 Ja  ",
                command=lambda: set_res(True),
                bg=CURRENT_THEME.SUCCESS,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                font=("Segoe UI", 10, "bold"),
                relief="flat",
                padx=25,
                pady=10,
                cursor="hand2",
            )
            yes_btn.pack(side="left", expand=True, padx=(0, 10))

            no_btn = tk.Button(
                btn_frame,
                text="  👎 Nein  ",
                command=lambda: set_res(False),
                bg=CURRENT_THEME.ERROR,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                font=("Segoe UI", 10, "bold"),
                relief="flat",
                padx=25,
                pady=10,
                cursor="hand2",
            )
            no_btn.pack(side="right", expand=True)

            dialog.bind("<Return>", lambda e: set_res(True))
            dialog.bind("<Escape>", lambda e: set_res(False))
            dialog.bind("y", lambda e: set_res(True))
            dialog.bind("n", lambda e: set_res(False))
            yes_btn.focus_set()

            dialog.protocol(
                "WM_DELETE_WINDOW",
                lambda: (cls._cancel_timeout(dialog, timeout_id), set_res(False)),
            )
            cls._center_dialog(dialog, parent)
            parent.wait_window(dialog)
            return result["value"]

        except (tk.TclError, RuntimeError, AttributeError):
            import tkinter.messagebox as mb

            return mb.askyesno(title, message)

    @classmethod
    def _create_base_dialog(
        cls,
        parent: tk.Tk,
        title: str,
        msg_type: str,
        message: str,
        icon_only: bool = False,
    ) -> tuple:
        dialog = tk.Toplevel(parent)
        dialog.title(f"🐉 {title}" if not title.startswith("🐉") else title)
        dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        dialog.resizable(False, False)
        dialog.transient(parent)
        dialog.grab_set()

        timeout_seconds = (
            60
            if any(w in title.lower() for w in ["beenden", "exit", "quit", "schließen"])
            else 10
        )
        timeout_id = dialog.after(
            timeout_seconds * 1000, lambda: cls._auto_close_dialog(dialog, title)
        )

        main = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=25, pady=25)
        main.pack(fill="both", expand=True)

        if icon_only:
            icon_label = tk.Label(
                main,
                text="❓",
                font=("Segoe UI", 28),
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_ACCENT,
            )
            icon_label.pack(pady=(0, 20))

            msg = tk.Label(
                main,
                text=message,
                font=Fonts.PRIMARY,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                justify="center",
                wraplength=350,
            )
            msg.pack(pady=(0, 30))
        else:
            content = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
            content.pack(fill="both", expand=True, pady=(0, 20))

            icon_frame = tk.Frame(content, bg=CURRENT_THEME.BG_PRIMARY, width=60)
            icon_frame.pack(side="left", fill="y")
            icon_frame.pack_propagate(False)

            icons = {
                "info": ("ℹ️", CURRENT_THEME.TEXT_ACCENT),
                "warning": ("⚠️", CURRENT_THEME.WARNING),
                "error": ("❌", CURRENT_THEME.ERROR),
                "question": ("❓", CURRENT_THEME.TEXT_ACCENT),
                "success": ("✅", CURRENT_THEME.SUCCESS),
            }
            icon_char, icon_color = icons.get(
                msg_type, ("💬", CURRENT_THEME.TEXT_PRIMARY)
            )

            tk.Label(
                icon_frame,
                text=icon_char,
                font=("Segoe UI", 24),
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=icon_color,
            ).pack(expand=True)

            msg_frame = tk.Frame(content, bg=CURRENT_THEME.BG_PRIMARY)
            msg_frame.pack(side="left", fill="both", expand=True, padx=(20, 0))

            if len(title) > 30:
                tk.Label(
                    msg_frame,
                    text=title,
                    font=Fonts.SUBTITLE,
                    bg=CURRENT_THEME.BG_PRIMARY,
                    fg=CURRENT_THEME.TEXT_PRIMARY,
                    justify="left",
                    anchor="w",
                ).pack(anchor="w", pady=(0, 10))

            tk.Label(
                msg_frame,
                text=message,
                font=Fonts.PRIMARY,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                justify="left",
                wraplength=350,
                anchor="w",
            ).pack(fill="x", expand=True, anchor="w")

        result = {"value": None}
        return dialog, result, timeout_id

    @classmethod
    def _add_ok_button(
        cls, dialog: tk.Toplevel, result: dict, timeout_id: str
    ) -> tk.Frame:
        btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x")

        def on_ok():
            cls._cancel_timeout(dialog, timeout_id)
            result["value"] = True
            if dialog.winfo_exists():
                dialog.destroy()

        ok_btn = tk.Button(
            btn_frame,
            text="OK",
            command=on_ok,
            bg=CURRENT_THEME.SUCCESS,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=25,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        ok_btn.pack(side="right")
        ok_btn.focus_set()

        dialog.bind("<Return>", lambda e: on_ok())
        dialog.bind("<Escape>", lambda e: on_ok())
        dialog.bind("<space>", lambda e: on_ok())
        return btn_frame

    @classmethod
    def _add_ok_cancel_buttons(
        cls, dialog: tk.Toplevel, result: dict, timeout_id: str, title: str
    ) -> tk.Frame:
        btn_frame = tk.Frame(dialog, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x")

        def set_result(val: bool):
            cls._cancel_timeout(dialog, timeout_id)
            result["value"] = val
            if dialog.winfo_exists():
                dialog.destroy()

        cancel_btn = tk.Button(
            btn_frame,
            text="Abbrechen",
            command=lambda: set_result(False),
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=22,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        cancel_btn.pack(side="right", padx=(10, 0))

        ok_btn = tk.Button(
            btn_frame,
            text="OK",
            command=lambda: set_result(True),
            bg=CURRENT_THEME.SUCCESS,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            relief="flat",
            padx=25,
            pady=8,
            cursor="hand2",
            takefocus=True,
        )
        ok_btn.pack(side="right")

        dialog.bind("<Return>", lambda e: set_result(True))
        dialog.bind("<Escape>", lambda e: set_result(False))
        dialog.bind("<space>", lambda e: cancel_btn.focus_set())

        is_exit = any(
            w in title.lower() for w in ["beenden", "exit", "quit", "schließen"]
        )
        if is_exit:
            cancel_btn.focus_set()
        else:
            ok_btn.focus_set()

        return btn_frame

    @staticmethod
    def _resolve_parent(parent: Optional[tk.Tk]) -> Optional[tk.Tk]:
        if parent and parent.winfo_exists():
            return parent
        if tk._default_root and tk._default_root.winfo_exists():
            return tk._default_root
        return None

    @staticmethod
    def _center_dialog(dialog: tk.Toplevel, parent: tk.Tk) -> None:
        try:
            dialog.update_idletasks()
            if parent and parent.winfo_exists():
                parent_x = parent.winfo_rootx()
                parent_y = parent.winfo_rooty()
                parent_w = parent.winfo_width()
                parent_h = parent.winfo_height()
                dlg_w = dialog.winfo_width()
                dlg_h = dialog.winfo_height()
                x = parent_x + (parent_w - dlg_w) // 2
                y = parent_y + (parent_h - dlg_h) // 2
                screen_w = parent.winfo_screenwidth()
                screen_h = parent.winfo_screenheight()
                x = max(10, min(x, screen_w - dlg_w - 10))
                y = max(10, min(y, screen_h - dlg_h - 10))
                dialog.geometry(f"+{x}+{y}")
                dialog.lift()
                dialog.focus_force()
            else:
                screen_w = dialog.winfo_screenwidth()
                screen_h = dialog.winfo_screenheight()
                dlg_w = dialog.winfo_width()
                dlg_h = dialog.winfo_height()
                x = (screen_w - dlg_w) // 2
                y = (screen_h - dlg_h) // 2
                dialog.geometry(f"+{x}+{y}")
        except Exception:
            pass

    @staticmethod
    def _auto_close_dialog(dialog: tk.Toplevel, title: str) -> None:
        try:
            if dialog and dialog.winfo_exists():
                logger.warning(f"⚠️ Dialog Timeout: '{title}'")
                dialog.destroy()
        except tk.TclError:
            pass

    @staticmethod
    def _cancel_timeout(dialog: tk.Toplevel, timeout_id: str) -> None:
        if timeout_id:
            try:
                dialog.after_cancel(timeout_id)
            except Exception:
                pass

    @staticmethod
    def _on_closing(result: dict, close_val: bool) -> None:
        result["value"] = close_val

    @staticmethod
    def _fallback_messagebox(
        title: str, message: str, msg_type: str, buttons: bool = False
    ) -> Optional[bool]:
        try:
            import tkinter.messagebox as mb

            if buttons:
                return mb.askokcancel(title, message)
            getattr(mb, f"show{msg_type}")(title, message)
            return None
        except Exception:
            logger.info(f"💬 {title}: {message}")
            return False if buttons else None


# -----------------------------------------------------------------------------
# Kontextmenüs
# -----------------------------------------------------------------------------
class DarkContextMenu:
    def __init__(self, text_widget: tk.Text) -> None:
        self.text_widget = text_widget
        self.menu = tk.Menu(
            text_widget,
            tearoff=0,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            activebackground=CURRENT_THEME.BG_HOVER,
            activeforeground=CURRENT_THEME.TEXT_ACCENT,
            borderwidth=1,
            relief="solid",
        )
        self.menu.add_command(label="Copy", command=self.copy_text)
        self.menu.add_command(label="Paste", command=self.paste_text)  # --- Änderung: Paste hinzugefügt
        self.menu.add_command(label="Select All", command=self.select_all)
        self.menu.add_separator()
        self.menu.add_command(label="Delete", command=self.clear_text)
        text_widget.bind("<Button-3>", self.show_menu)

    def show_menu(self, event: tk.Event) -> None:
        try:
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu.grab_release()

    def copy_text(self) -> None:
        try:
            if self.text_widget.tag_ranges(tk.SEL):
                selected_text = self.text_widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.text_widget.clipboard_clear()
                self.text_widget.clipboard_append(selected_text)
        except tk.TclError:
            pass

    def paste_text(self) -> None:
        """Fügt den Inhalt der Zwischenablage an der Cursorposition ein."""
        try:
            self.text_widget.event_generate("<<Paste>>")
        except tk.TclError:
            pass

    def select_all(self) -> None:
        try:
            self.text_widget.tag_add(tk.SEL, "1.0", tk.END)
            self.text_widget.mark_set(tk.INSERT, "1.0")
            self.text_widget.see(tk.INSERT)
        except tk.TclError:
            pass

    def clear_text(self) -> None:
        try:
            self.text_widget.delete("1.0", tk.END)
        except tk.TclError:
            pass


class DarkEntryContextMenu:
    def __init__(self, entry_widget: tk.Entry) -> None:
        self.entry_widget = entry_widget
        self.menu = tk.Menu(
            entry_widget,
            tearoff=0,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            activebackground=CURRENT_THEME.BG_HOVER,
            activeforeground=CURRENT_THEME.TEXT_ACCENT,
            borderwidth=1,
            relief="solid",
        )
        self.menu.add_command(label="Cut", command=self.cut_text)
        self.menu.add_command(label="Copy", command=self.copy_text)
        self.menu.add_command(label="Paste", command=self.paste_text)
        self.menu.add_separator()
        self.menu.add_command(label="Select All", command=self.select_all)
        self.menu.add_command(label="Delete", command=self.delete_text)
        entry_widget.bind("<Button-3>", self.show_menu)

    def show_menu(self, event: tk.Event) -> None:
        try:
            self.menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.menu.grab_release()

    def cut_text(self) -> None:
        self.entry_widget.event_generate("<<Cut>>")

    def copy_text(self) -> None:
        self.entry_widget.event_generate("<<Copy>>")

    def paste_text(self) -> None:
        self.entry_widget.event_generate("<<Paste>>")

    def select_all(self) -> None:
        self.entry_widget.select_range(0, "end")
        self.entry_widget.icursor("end")

    def delete_text(self) -> None:
        self.entry_widget.delete(0, "end")


# -----------------------------------------------------------------------------
# ToolTip
# -----------------------------------------------------------------------------
class ToolTip:
    def __init__(self, widget: tk.Widget, text: str, delay: int = 500) -> None:
        self.widget = widget
        self.text = text
        self.delay = delay
        self.tip_window: Optional[tk.Toplevel] = None
        self.after_id: Optional[str] = None
        widget.bind("<Enter>", self.enter)
        widget.bind("<Leave>", self.leave)
        widget.bind("<ButtonPress>", self.leave)

    def enter(self, event: Optional[tk.Event] = None) -> None:
        self.schedule()

    def leave(self, event: Optional[tk.Event] = None) -> None:
        self.unschedule()
        self.hide_tip()

    def schedule(self) -> None:
        self.unschedule()
        self.after_id = self.widget.after(self.delay, self.show_tip)

    def unschedule(self) -> None:
        if self.after_id:
            self.widget.after_cancel(self.after_id)
            self.after_id = None

    def show_tip(self) -> None:
        if self.tip_window or not self.text:
            return
        x, y = self.widget.winfo_pointerxy()
        x += 10
        y += 10
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(
            tw,
            text=self.text,
            background="#ffffe0",
            relief="solid",
            borderwidth=1,
            font=("Segoe UI", 9),
        )
        label.pack()
        tw.bind("<Leave>", lambda e: self.hide_tip())

    def hide_tip(self) -> None:
        if self.tip_window:
            try:
                self.tip_window.destroy()
            except tk.TclError:
                pass
            self.tip_window = None


# -----------------------------------------------------------------------------
# ProgressDialog
# -----------------------------------------------------------------------------
class ProgressDialog:
    def __init__(self, parent: tk.Tk, title: str = "Processing...") -> None:
        self.parent = parent
        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        self.dialog.geometry("300x120")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.protocol("WM_DELETE_WINDOW", self.cancel)

        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.dialog.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")

        content_frame = tk.Frame(
            self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=20, pady=20
        )
        content_frame.pack(fill="both", expand=True)

        self.message_label = tk.Label(
            content_frame,
            text="Analyzing video...",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        )
        self.message_label.pack(pady=(0, 10))

        self.progress = ttk.Progressbar(content_frame, mode="indeterminate", length=250)
        self.progress.pack(pady=(0, 15))

        button_frame = tk.Frame(content_frame, bg=CURRENT_THEME.BG_PRIMARY)
        button_frame.pack()

        self.cancel_button = tk.Button(
            button_frame,
            text="Cancel",
            command=self.cancel,
            bg=CURRENT_THEME.ERROR,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            relief="flat",
            padx=15,
        )
        self.cancel_button.pack()

        self.is_cancelled = False
        self._is_running = True
        self._update_interval = 100
        self._after_id = None

        self.progress.start(10)
        self._schedule_updates()

    def _schedule_updates(self) -> None:
        if not self._is_running:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.update_idletasks()
                self._after_id = self.dialog.after(
                    self._update_interval, self._schedule_updates
                )
        except tk.TclError:
            self._is_running = False
            self._after_id = None

    def cancel(self) -> None:
        if not self._is_running:
            return
        self.is_cancelled = True
        self._is_running = False

        try:
            if self.message_label and self.message_label.winfo_exists():
                self.message_label.config(text="Cancelling...", fg=CURRENT_THEME.ERROR)
            if self.cancel_button and self.cancel_button.winfo_exists():
                self.cancel_button.config(text="Cancelling...", state="disabled")
        except tk.TclError:
            pass

        if self._after_id:
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.after_cancel(self._after_id)
            except tk.TclError:
                pass
            self._after_id = None

        self.close()

    def update_message(self, message: str) -> None:
        if not self._is_running:
            return

        def _update():
            try:
                if self.message_label and self.message_label.winfo_exists():
                    self.message_label.config(text=message)
            except tk.TclError:
                pass

        if self.dialog and self.dialog.winfo_exists():
            self.dialog.after(0, _update)

    def close(self) -> None:
        self._is_running = False
        if self._after_id:
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.after_cancel(self._after_id)
            except tk.TclError:
                pass
            self._after_id = None

        try:
            self.progress.stop()
        except Exception:
            pass

        try:
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.destroy()
        except Exception:
            pass
        finally:
            self.dialog = None
            self.message_label = None
            self.progress = None
            self.cancel_button = None


# -----------------------------------------------------------------------------
# SummarizeDialog
# -----------------------------------------------------------------------------
class SummarizeDialog:
    def __init__(self, parent: Any, text: str, gui_ref: Any) -> None:
        self.parent = parent
        self.text = text
        self.gui = gui_ref
        self.dialog: Optional[tk.Toplevel] = None
        self.summarizer = None
        self._destroyed = False
        self._request_cancel = threading.Event()

        # Summarizer erstellen und Server-Reachability prüfen
        self.summarizer = OllamaSummarizer(parent)
        if not self.summarizer.is_server_reachable():
            DarkMessageBox.showwarning(
                "Ollama nicht erreichbar",
                "Der Ollama-Server läuft nicht oder ist nicht erreichbar.\n"
                "Bitte starte 'ollama serve' und versuche es erneut.",
                parent=self.parent,
            )
            return

        # Dialog erstellen
        self.create_dialog()

        # Dialog zur Verwaltung hinzufügen (falls gewünscht)
        if hasattr(self.gui, '_open_dialogs') and self.dialog is not None:
            self.gui._open_dialogs.append(self.dialog)
            self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)

    def _on_close(self):
        if hasattr(self.gui, '_open_dialogs') and self.dialog in self.gui._open_dialogs:
            self.gui._open_dialogs.remove(self.dialog)
        self.dialog.destroy()

    def create_dialog(self) -> None:
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("🐉 Zusammenfassung mit Ollama")
        self.dialog.geometry("750x650")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(self.parent)
        self.dialog.grab_set()

        def on_close() -> None:
            self._destroyed = True
            self._request_cancel.set()
            try:
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.destroy()
            except tk.TclError:
                pass

        self.dialog.protocol("WM_DELETE_WINDOW", on_close)

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        model_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        model_frame.pack(fill="x", pady=5)
        tk.Label(
            model_frame,
            text="Modell:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left")
        self.model_var = tk.StringVar()
        available = self.summarizer.get_available_models()
        if available:
            preferred = ["qwen2.5:7b", "glm4:9b", "llama3.1:8b"]
            values = []
            for pref in preferred:
                if pref in available:
                    values.append(pref)
            for m in available:
                if m not in values:
                    values.append(m)
            self.model_combo = ttk.Combobox(
                model_frame,
                textvariable=self.model_var,
                values=values,
                width=20,
                state="readonly",
                style="Dark.TCombobox",
            )
            self.model_var.set(values[0] if values else "")
        else:
            self.model_combo = ttk.Combobox(
                model_frame,
                textvariable=self.model_var,
                values=["(keine Modelle gefunden)"],
                width=20,
                state="disabled",
                style="Dark.TCombobox",
            )
            self.model_var.set("(keine Modelle)")
        self.model_combo.pack(side="left", padx=10)
        ToolTip(self.model_combo, "Wähle das Ollama-Modell für die Zusammenfassung")

        tk.Label(
            model_frame,
            text="Temperatur:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left", padx=(20, 5))
        self.temp_var = tk.DoubleVar(value=0.0)
        temp_scale = tk.Scale(
            model_frame,
            from_=0.0,
            to=1.0,
            resolution=0.1,
            orient=tk.HORIZONTAL,
            variable=self.temp_var,
            length=150,
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            highlightbackground=CURRENT_THEME.BG_PRIMARY,
        )
        temp_scale.pack(side="left")
        tk.Label(
            model_frame,
            text="(0 = deterministisch)",
            font=("Segoe UI", 7),
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_SECONDARY,
        ).pack(side="left", padx=5)
        ToolTip(temp_scale, "Zufälligkeit der Ausgabe (höher = kreativer)")

        lang_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        lang_frame.pack(fill="x", pady=5)
        tk.Label(
            lang_frame,
            text="Zusammenfassen auf:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left")
        self.summary_lang_var = tk.StringVar()
        supported_summary_langs = ["Deutsch", "Englisch", "Spanisch", "Koreanisch"]
        current_lang_name = SUPPORTED_LANGUAGES.get(
            self.gui.current_language, "Deutsch"
        )
        if current_lang_name not in supported_summary_langs:
            current_lang_name = "Deutsch"
        self.summary_lang_var.set(current_lang_name)
        lang_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.summary_lang_var,
            values=supported_summary_langs,
            width=15,
            state="readonly",
            style="Dark.TCombobox",
        )
        lang_combo.pack(side="left", padx=10)
        lang_combo.bind("<<ComboboxSelected>>", lambda e: self._set_default_prompt())
        ToolTip(lang_combo, "Sprache der Zusammenfassung")

        tk.Label(
            main,
            text="Prompt (optional):",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(anchor="w", pady=(10, 2))
        self.prompt_text = scrolledtext.ScrolledText(
            main,
            height=4,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
        )
        self.prompt_text.pack(fill="x", pady=(0, 10))
        self._set_default_prompt()
        ToolTip(self.prompt_text, "Optionaler Prompt – wird an das Modell gesendet")

        tk.Label(
            main,
            text="Zusammenfassung:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(anchor="w")
        self.summary_text = scrolledtext.ScrolledText(
            main,
            height=10,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
        )
        self.summary_text.pack(fill="both", expand=True, pady=10)
        DarkContextMenu(self.summary_text)

        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x", pady=5)

        btn_row1 = tk.Frame(btn_frame, bg=CURRENT_THEME.BG_PRIMARY)
        btn_row1.pack(fill="x", pady=2)
        self.summarize_btn = tk.Button(
            btn_row1,
            text="🤖 Zusammenfassen",
            command=self.start_summarize,
            bg=CURRENT_THEME.DRAGON_GREEN,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
        )
        self.summarize_btn.pack(side="left", padx=5)

        self.cancel_btn = tk.Button(
            btn_row1,
            text="⏹️ Abbrechen",
            command=self.cancel_request,
            bg=CURRENT_THEME.ERROR,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
            state="disabled",
        )
        self.cancel_btn.pack(side="left", padx=5)

        self.copy_btn = tk.Button(
            btn_row1,
            text="📋 Kopieren",
            command=self.copy_summary,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
        )
        self.copy_btn.pack(side="left", padx=5)

        btn_row2 = tk.Frame(btn_frame, bg=CURRENT_THEME.BG_PRIMARY)
        btn_row2.pack(fill="x", pady=2)
        self.save_btn = tk.Button(
            btn_row2,
            text="💾 Speichern",
            command=self.save_summary,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
        )
        self.save_btn.pack(side="left", padx=5)

        self.translate_btn = tk.Button(
            btn_row2,
            text="🌐 Übersetzen",
            command=self.translate_summary,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
            state="disabled",
        )
        self.translate_btn.pack(side="left", padx=5)

        self.close_btn = tk.Button(
            btn_row2,
            text="Schließen",
            command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        )
        self.close_btn.pack(side="left", padx=5)

        self.status_label = tk.Label(
            main, text="", bg=CURRENT_THEME.BG_PRIMARY, fg=CURRENT_THEME.TEXT_SECONDARY
        )
        self.status_label.pack(pady=5)

        self.full_summary = ""
        self._chunk_results: List[str] = []
        self._chunks: List[str] = []
        self._current_chunk = 0
        self._prompt = ""
        self._temp = 0.0

    def _set_default_prompt(self) -> None:
        target = self.summary_lang_var.get()
        if target == "Deutsch":
            prompt = (
                "Fasse den folgenden Text in 3-5 Sätzen auf Deutsch zusammen. "
                "Konzentriere dich auf die Hauptaussagen, ignoriere Wiederholungen und Füllwörter. "
                "Verwende vollständige, idiomatische Sätze.\n\n"
                "WICHTIG: Die Ausgabe MUSS auf DEUTSCH erfolgen und darf KEINE Zeichen aus anderen "
                "Sprachen (wie Chinesisch, Japanisch oder Koreanisch) enthalten. "
                "Fachbegriffe wie 'Sakura' oder 'Kimchi' können als Lehnwörter verwendet werden, "
                "aber bitte in lateinischer Schrift."
            )
        elif target == "Englisch":
            prompt = (
                "Summarize the following text in 3-5 sentences in English. "
                "Focus on the main points, ignore repetitions and filler words. "
                "Use complete, idiomatic sentences.\n\n"
                "IMPORTANT: The output MUST be in ENGLISH and must NOT contain characters from other "
                "languages (like Chinese, Japanese, or Korean). Loanwords are acceptable in Latin script."
            )
        elif target == "Spanisch":
            prompt = (
                "Resume el siguiente texto en 3-5 oraciones en español. "
                "Concéntrate en las ideas principales, ignora repeticiones y palabras de relleno. "
                "Utiliza oraciones completas e idiomáticas.\n\n"
                "IMPORTANTE: La salida DEBE estar en ESPAÑOL y NO debe contener caracteres de otros "
                "idiomas (como chino, japonés o coreano). Los préstamos son aceptables en escritura latina."
            )
        elif target == "Koreanisch":
            prompt = (
                "다음 텍스트를 한국어로 3-5문장으로 요약하세요. "
                "핵심 내용에 집중하고, 반복과 불용어는 무시하세요. "
                "완전하고 관용적인 문장을 사용하세요.\n\n"
                "중요: 출력은 반드시 한국어로 작성되어야 하며 다른 언어(예: 중국어, 일본어)의 문자가 포함되어서는 안 됩니다. "
                "외래어는 라틴 문자로 표기해도 됩니다."
            )
        else:
            prompt = (
                f"Summarize the following text in 3-5 sentences in {target}. "
                f"Focus on the main points, ignore repetitions and filler words. "
                f"Use complete, idiomatic sentences.\n\n"
                f"IMPORTANT: The output MUST be in {target.upper()} and must NOT contain characters from other languages."
            )
        self.prompt_text.delete("1.0", "end")
        self.prompt_text.insert("1.0", prompt)

    def _split_text(self, text: str, max_words: int = 2000) -> List[str]:
        words = text.split()
        if len(words) <= max_words:
            return [text]
        chunks = []
        chunk = []
        word_count = 0
        for word in words:
            chunk.append(word)
            word_count += 1
            if word_count >= max_words:
                chunks.append(" ".join(chunk))
                chunk = []
                word_count = 0
        if chunk:
            chunks.append(" ".join(chunk))
        return chunks

    def start_summarize(self) -> None:
        if self._destroyed:
            return
        model = self.model_var.get().strip()
        if model == "--- Modell auswählen ---" or model.startswith("(keine"):
            self.status_label.config(text="❌ Bitte ein gültiges Modell auswählen")
            return

        if not self.summarizer.is_model_available(model):
            self.status_label.config(
                text=f"❌ Modell '{model}' nicht auf Server gefunden."
            )
            return

        self.summarizer.model = model
        if hasattr(self.gui, "advanced_settings"):
            self.summarizer.host = self.gui.advanced_settings.ollama_host
        else:
            self.summarizer.host = "http://localhost:11434"

        target_lang = self.summary_lang_var.get()
        system_prompt = f"Du bist ein Assistent, der Texte auf {target_lang} zusammenfasst. Antworte ausschließlich auf {target_lang}."
        self.summarizer.system_prompt = system_prompt

        prompt = self.prompt_text.get("1.0", "end-1c").strip()
        temp = self.temp_var.get()

        word_count = len(self.text.split())
        if word_count > 2000:
            self.status_label.config(text="⏳ Text wird in Abschnitte zerlegt...")
            self.summarize_btn.config(state="disabled", text="⏳ Warte...")
            self.cancel_btn.config(state="normal")
            self.copy_btn.config(state="disabled")
            self.save_btn.config(state="disabled")
            self.translate_btn.config(state="disabled")
            self._request_cancel.clear()
            self.full_summary = ""
            self.summary_text.delete("1.0", "end")

            chunks = self._split_text(self.text, max_words=2000)
            self.status_label.config(text=f"⏳ Verarbeite {len(chunks)} Abschnitte...")
            self._chunks = chunks
            self._prompt = prompt
            self._temp = temp
            self._chunk_results = []
            self._current_chunk = 0
            self._process_next_chunk()
        else:
            self.summarize_btn.config(state="disabled", text="⏳ Warte...")
            self.cancel_btn.config(state="normal")
            self.copy_btn.config(state="disabled")
            self.save_btn.config(state="disabled")
            self.translate_btn.config(state="disabled")
            self.status_label.config(text="Sende Anfrage an Ollama...")
            self.summary_text.delete("1.0", "end")
            self.full_summary = ""
            self._request_cancel.clear()

            def on_complete() -> None:
                if self._destroyed:
                    return
                if self.dialog and self.dialog.winfo_exists():
                    self.dialog.after(0, self._reset_ui)

            self.summarizer.summarize(
                self.text,
                prompt,
                temp,
                callback=self.on_chunk,
                error_callback=self.on_error,
                complete_callback=on_complete,
                cancel_event=self._request_cancel,
            )

    def _process_next_chunk(self, retry_count: int = 0) -> None:
        if self._destroyed or self._request_cancel.is_set():
            self._finalize_chunks()
            return

        if self._current_chunk >= len(self._chunks):
            self._create_final_summary()
            return

        chunk_text = self._chunks[self._current_chunk]
        self.status_label.config(
            text=f"⏳ Verarbeite Abschnitt {self._current_chunk+1}/{len(self._chunks)}... (Versuch {retry_count+1})"
        )

        def on_chunk_complete():
            if self.summarizer.last_result:
                self._chunk_results.append(self.summarizer.last_result)
            else:
                self._chunk_results.append("")
            self._current_chunk += 1
            self.dialog.after(100, lambda: self._process_next_chunk(0))

        def on_chunk_error(error: str):
            if retry_count < 3:
                delay_ms = 1000 * (2 ** retry_count)
                self.status_label.config(text=f"⏳ Fehler, Wiederholung in {delay_ms//1000}s... (Versuch {retry_count+2}/4)")
                self.dialog.after(delay_ms, lambda: self._process_next_chunk(retry_count + 1))
            else:
                self._chunk_results.append(f"[Fehler in Abschnitt {self._current_chunk+1} nach 3 Versuchen: {error}]")
                self._current_chunk += 1
                self.dialog.after(100, lambda: self._process_next_chunk(0))

        self.summarizer.summarize(
            chunk_text,
            self._prompt,
            self._temp,
            callback=lambda chunk: None,
            error_callback=on_chunk_error,
            complete_callback=on_chunk_complete,
            cancel_event=self._request_cancel,
        )

    def _create_final_summary(self) -> None:
        if not self._chunk_results:
            self.status_label.config(text="❌ Keine Teilzusammenfassungen vorhanden.")
            self._reset_ui()
            return

        valid_chunks = [res for res in self._chunk_results if res and not res.startswith("[Fehler")]
        if not valid_chunks:
            self.status_label.config(text="❌ Keine gültigen Teilzusammenfassungen.")
            self._reset_ui()
            return

        combined = "\n\n".join(valid_chunks)
        target_lang = self.summary_lang_var.get()
        final_prompt = (f"Fasse die folgenden Teilzusammenfassungen zu einer kohärenten Gesamtzusammenfassung in {target_lang} zusammen. "
                        f"Antworte ausschließlich auf {target_lang}.\n\n{combined}")

        self.status_label.config(text="⏳ Erstelle finale Zusammenfassung...")

        def on_final_complete():
            if self._destroyed:
                return
            if self.dialog and self.dialog.winfo_exists():
                self.dialog.after(0, self._reset_ui)

        self.summarizer.summarize(
            combined,
            final_prompt,
            self._temp,
            callback=self.on_chunk,
            error_callback=self.on_error,
            complete_callback=on_final_complete,
            cancel_event=self._request_cancel,
        )

    def _reset_ui(self) -> None:
        if self._destroyed:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                if hasattr(self, "summarize_btn") and self.summarize_btn.winfo_exists():
                    self.summarize_btn.config(state="normal", text="🤖 Zusammenfassen")
                if hasattr(self, "cancel_btn") and self.cancel_btn.winfo_exists():
                    self.cancel_btn.config(state="disabled")
                if hasattr(self, "copy_btn") and self.copy_btn.winfo_exists():
                    self.copy_btn.config(state="normal")
                if hasattr(self, "save_btn") and self.save_btn.winfo_exists():
                    self.save_btn.config(state="normal")
                if hasattr(self, "translate_btn") and self.translate_btn.winfo_exists():
                    self.translate_btn.config(
                        state="normal" if self.full_summary else "disabled"
                    )
                if hasattr(self, "status_label") and self.status_label.winfo_exists():
                    self.status_label.config(text="✅ Zusammenfassung abgeschlossen")

                if self.full_summary and hasattr(self.gui, "advanced_settings"):
                    blacklist = getattr(self.gui.advanced_settings, "blacklist", [])
                    found = []
                    for phrase in blacklist:
                        if phrase and phrase.lower() in self.full_summary.lower():
                            found.append(phrase)
                    if found:
                        logger.warning(f"Zusammenfassung enthält Blacklist-Phrasen: {found}")
                        self.status_label.config(
                            text=f"⚠️ Warnung: Blacklist-Phrasen gefunden: {', '.join(found[:2])}"
                        )

                try:
                    from langdetect import detect
                    detected = detect(self.full_summary)
                    target = self.summary_lang_var.get()
                    target_map = {"Deutsch": "de", "Englisch": "en", "Spanisch": "es", "Koreanisch": "ko"}
                    expected = target_map.get(target, "de")
                    if detected != expected:
                        logger.warning(f"Zusammenfassung in falscher Sprache erkannt: {detected}, erwartet {expected}")
                        self.status_label.config(
                            text=f"⚠️ Warnung: Sprache erkannt als {detected}, erwartet {target}"
                        )
                except ImportError:
                    pass
                except Exception as e:
                    logger.debug(f"Spracherkennung fehlgeschlagen: {e}")

        except tk.TclError:
            pass

    def cancel_request(self) -> None:
        self._request_cancel.set()
        self.status_label.config(text="⏹️ Abbruch eingeleitet...")
        self.cancel_btn.config(state="disabled")
        self.dialog.after(500, self._reset_ui)

    def on_chunk(self, chunk: str) -> None:
        if self._destroyed:
            return

        def update() -> None:
            if self._destroyed:
                return
            try:
                if self.dialog and self.dialog.winfo_exists():
                    if (
                        hasattr(self, "summary_text")
                        and self.summary_text.winfo_exists()
                    ):
                        self.summary_text.insert("end", chunk)
                        self.summary_text.see("end")
                        self.full_summary += chunk
            except tk.TclError:
                pass

        if self.dialog and not self._destroyed and self.dialog.winfo_exists():
            self.dialog.after(0, update)

    def on_error(self, error: str) -> None:
        if self._destroyed:
            return

        def update() -> None:
            if self._destroyed:
                return
            try:
                if self.dialog and self.dialog.winfo_exists():
                    if (
                        hasattr(self, "summary_text")
                        and self.summary_text.winfo_exists()
                    ):
                        self.summary_text.delete("1.0", "end")
                        self.summary_text.insert("1.0", f"Fehler: {error}")
                    self._reset_ui()
                    if (
                        hasattr(self, "status_label")
                        and self.status_label.winfo_exists()
                    ):
                        self.status_label.config(text="❌ Fehler")
            except tk.TclError:
                pass

        if self.dialog and not self._destroyed and self.dialog.winfo_exists():
            self.dialog.after(0, update)

    def copy_summary(self) -> None:
        if self._destroyed:
            return
        try:
            if self.dialog and self.dialog.winfo_exists():
                if self.full_summary:
                    self.dialog.clipboard_clear()
                    self.dialog.clipboard_append(self.full_summary)
                    if (
                        hasattr(self, "status_label")
                        and self.status_label.winfo_exists()
                    ):
                        self.status_label.config(text="✅ In Zwischenablage kopiert")
                else:
                    if (
                        hasattr(self, "status_label")
                        and self.status_label.winfo_exists()
                    ):
                        self.status_label.config(
                            text="⚠️ Keine Zusammenfassung vorhanden"
                        )
        except tk.TclError:
            pass

    def save_summary(self) -> None:
        if not self.full_summary:
            self.status_label.config(text="⚠️ Keine Zusammenfassung zum Speichern")
            return
        from tkinter import filedialog

        filename = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Textdateien", "*.txt"), ("Alle Dateien", "*.*")],
            title="Zusammenfassung speichern",
        )
        if filename:
            try:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(self.full_summary)
                self.status_label.config(
                    text=f"💾 Gespeichert: {os.path.basename(filename)}"
                )
            except Exception as e:
                self.status_label.config(text=f"❌ Fehler beim Speichern: {e}")

    def translate_summary(self) -> None:
        if not self.full_summary:
            self.status_label.config(text="⚠️ Keine Zusammenfassung zum Übersetzen")
            return
        if not hasattr(self.gui, "translation_engine"):
            self.status_label.config(text="❌ Keine Übersetzungs-Engine verfügbar")
            return
        engine = self.gui.translation_engine
        if hasattr(engine, "is_functional") and not engine.is_functional():
            self.status_label.config(
                text="⚠️ Übersetzungs-Engine derzeit nicht verfügbar"
            )
            return
        TranslationDialog(self.dialog, engine, initial_text=self.full_summary)


# -----------------------------------------------------------------------------
# TranslationDialog – mit verbesserter Engine-Auswahl und satzweiser Übersetzung
# -----------------------------------------------------------------------------
class TranslationDialog:
    def __init__(
        self,
        parent: tk.Widget,
        translation_engine: BaseTranslationEngine,
        initial_text: str = "",
    ) -> None:
        self.parent = parent
        self.engine = translation_engine
        self.gui = parent.gui if hasattr(parent, "gui") else None
        self.initial_text = initial_text
        self.dialog: Optional[tk.Toplevel] = None
        self._cancel_event = threading.Event()
        self.engine_var = tk.StringVar(value=self.engine.__class__.__name__.replace('TranslationEngine', '').lower())
        self.create_dialog()

        if self.gui and hasattr(self.gui, '_open_dialogs'):
            self.gui._open_dialogs.append(self.dialog)
            self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)

    def _on_close(self):
        if self.gui and hasattr(self.gui, '_open_dialogs') and self.dialog in self.gui._open_dialogs:
            self.gui._open_dialogs.remove(self.dialog)
        self.dialog.destroy()

    def create_dialog(self) -> None:
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("🐉 Text Translation")
        self.dialog.geometry("600x600")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(self.parent)
        self.dialog.grab_set()

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        tk.Label(
            main,
            text="Source text:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        ).pack(anchor="w")
        self.source_text = scrolledtext.ScrolledText(
            main,
            height=8,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
        )
        self.source_text.pack(fill="x", pady=(0, 10))
        if self.initial_text:
            self.source_text.insert("1.0", self.initial_text)
        DarkContextMenu(self.source_text)

        lang_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        lang_frame.pack(fill="x", pady=5)

        tk.Label(
            lang_frame,
            text="From:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left", padx=(0, 5))
        self.src_lang_var = tk.StringVar(value="auto")
        src_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.src_lang_var,
            values=["auto"] + [name for name, code in SORTED_LANGUAGES],
            width=15,
            state="readonly",
        )
        src_combo.pack(side="left", padx=(0, 20))

        tk.Label(
            lang_frame,
            text="To:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left", padx=(0, 5))

        current_target_name = SUPPORTED_LANGUAGES.get(
            self.engine.default_target_lang, "Deutsch"
        )
        self.tgt_lang_var = tk.StringVar(value=current_target_name)
        tgt_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.tgt_lang_var,
            values=[name for name, code in SORTED_LANGUAGES if name != "auto"],
            width=15,
            state="readonly",
        )
        tgt_combo.pack(side="left")

        engine_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        engine_frame.pack(fill="x", pady=5)
        tk.Label(
            engine_frame,
            text="Engine:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(side="left", padx=(0, 5))
        engine_combo = ttk.Combobox(
            engine_frame,
            textvariable=self.engine_var,
            values=["google", "ollama", "argos"],
            width=10,
            state="readonly",
        )
        engine_combo.pack(side="left")

        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x", pady=10)

        self.translate_btn = tk.Button(
            btn_frame,
            text="🌐 Translate",
            command=self.translate,
            bg=CURRENT_THEME.DRAGON_GREEN,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=20,
        )
        self.translate_btn.pack(side="left")

        self.cancel_btn = tk.Button(
            btn_frame,
            text="⏹️ Cancel",
            command=self.cancel_translation,
            bg=CURRENT_THEME.ERROR,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=20,
            state="disabled",
        )
        self.cancel_btn.pack(side="left", padx=10)

        tk.Label(
            main,
            text="Translation:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        ).pack(anchor="w", pady=(10, 0))
        self.target_text = scrolledtext.ScrolledText(
            main,
            height=8,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
            state="normal",
        )
        self.target_text.pack(fill="both", expand=True, pady=(5, 0))
        DarkContextMenu(self.target_text)

        self.progress_label = tk.Label(
            main,
            text="",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_SECONDARY,
            font=Fonts.SMALL,
        )
        self.progress_label.pack(pady=2)

        close_btn = tk.Button(
            main,
            text="Close",
            command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        )
        close_btn.pack(pady=10)

    # --- Änderung: Hilfsmethode zur Satzaufteilung (auch asiatische Satzzeichen)
    def _split_sentences(self, text: str) -> List[str]:
        """Teilt Text in Sätze auf, berücksichtigt .!? sowie asiatische Satzzeichen 。！？"""
        # Regulärer Ausdruck für Satzenden: Punkt, Ausrufezeichen, Fragezeichen,
        # sowie chinesische/japanische Zeichen 。！？, gefolgt von Leerzeichen oder Zeilenende
        sentence_endings = r'(?<=[.!?。！？])\s+'
        sentences = re.split(sentence_endings, text)
        # Leere Einträge entfernen
        return [s.strip() for s in sentences if s.strip()]

    def translate(self) -> None:
        source = self.source_text.get("1.0", "end-1c").strip()
        if not source:
            return

        src_name = self.src_lang_var.get().strip()
        tgt_name = self.tgt_lang_var.get().strip()
        engine_name = self.engine_var.get().strip()

        valid_language_names = [name for name, code in SORTED_LANGUAGES]

        if src_name not in valid_language_names and src_name != "auto":
            src_name = "auto"
            self.src_lang_var.set("auto")

        if tgt_name not in valid_language_names:
            tgt_name = "Deutsch"
            self.tgt_lang_var.set("Deutsch")

        try:
            src_code = (
                "auto"
                if src_name == "auto"
                else next(code for name, code in SORTED_LANGUAGES if name == src_name)
            )
            tgt_code = next(code for name, code in SORTED_LANGUAGES if name == tgt_name)
        except StopIteration:
            self._display_error("Ungültige Sprachauswahl")
            return

        # --- Temporäre Engine erstellen
        if self.gui and hasattr(self.gui, "_create_temporary_translation_engine"):
            engine = self.gui._create_temporary_translation_engine(engine_name, tgt_code)
            if engine is None:
                self._display_error(f"Engine '{engine_name}' nicht verfügbar (siehe Log)")
                return
        else:
            engine = self.engine
            old_target = engine.default_target_lang
            engine.set_target_language(tgt_code)

        # UI deaktivieren
        self.translate_btn.config(state="disabled", text="⏳ Translating...")
        self.cancel_btn.config(state="normal")
        self.target_text.delete("1.0", "end")
        self.progress_label.config(text="")
        self._cancel_event.clear()

        # --- Satzweise Übersetzung
        sentences = self._split_sentences(source)
        total = len(sentences)
        if total == 0:
            self._reset_ui()
            return
        if total == 1:
            # Ein Satz – normale Übersetzung
            self._translate_single(sentences[0], src_code, engine, old_target if not self.gui else None)
        else:
            # Mehrere Sätze – mit Fortschrittsanzeige
            self._translate_multi(sentences, src_code, engine, old_target)

    def _translate_single(self, text: str, src_code: str, engine: BaseTranslationEngine, old_target: Optional[str]):
        """Übersetzt einen einzelnen Satz (synchron im Thread)."""
        def worker():
            try:
                result = engine.translate_text(text, src_code)
                if result and result.translated and not self._cancel_event.is_set():
                    self.dialog.after(0, self._display_result, result.translated)
                elif not self._cancel_event.is_set():
                    self.dialog.after(0, self._display_result, "(Übersetzung fehlgeschlagen)")
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                if not self._cancel_event.is_set():
                    self.dialog.after(0, self._display_error, str(e))
            finally:
                if not self.gui and old_target is not None:
                    engine.set_target_language(old_target)
                if not self._cancel_event.is_set():
                    self.dialog.after(0, self._reset_ui)

        threading.Thread(target=worker, daemon=True).start()

    def _translate_multi(self, sentences: List[str], src_code: str, engine: BaseTranslationEngine, old_target: Optional[str]):
        """Übersetzt mehrere Sätze nacheinander und sammelt die Ergebnisse."""
        translated_sentences = []
        current = 0
        total = len(sentences)

        def translate_next():
            nonlocal current
            if current >= total or self._cancel_event.is_set():
                # Fertig oder abgebrochen
                if not self._cancel_event.is_set():
                    # Alle übersetzt
                    full_translation = " ".join(translated_sentences)
                    self.dialog.after(0, self._display_result, full_translation)
                self.dialog.after(0, self._reset_ui)
                return

            sentence = sentences[current]
            self.progress_label.config(text=f"Übersetze Satz {current+1}/{total}...")

            def worker():
                nonlocal current
                try:
                    result = engine.translate_text(sentence, src_code)
                    if result and result.translated and not self._cancel_event.is_set():
                        translated_sentences.append(result.translated)
                    else:
                        translated_sentences.append("[Fehler]")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    translated_sentences.append(f"[Fehler: {str(e)}]")
                finally:
                    current += 1
                    self.dialog.after(10, translate_next)

            threading.Thread(target=worker, daemon=True).start()

        translate_next()

    def _display_result(self, translated_text: str) -> None:
        self.target_text.delete("1.0", "end")
        self.target_text.insert("1.0", translated_text)

    def _display_error(self, error: str) -> None:
        self.target_text.delete("1.0", "end")
        self.target_text.insert("1.0", f"Error: {error}")

    def _reset_ui(self) -> None:
        self.translate_btn.config(state="normal", text="🌐 Translate")
        self.cancel_btn.config(state="disabled")
        self.progress_label.config(text="")

    def cancel_translation(self) -> None:
        self._cancel_event.set()
        self._reset_ui()
        self.target_text.insert("1.0", "(Translation cancelled)")


# -----------------------------------------------------------------------------
# ShortcutsDialog
# -----------------------------------------------------------------------------
class ShortcutsDialog:
    def __init__(self, parent):
        self.parent = parent
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("🐉 Tastenkürzel")
        self.dialog.geometry("500x400")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()

        if hasattr(parent, '_open_dialogs'):
            parent._open_dialogs.append(self.dialog)
            self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=20, pady=20)
        main.pack(fill="both", expand=True)

        tk.Label(
            main,
            text="Tastenkürzel",
            font=("Segoe UI", 14, "bold"),
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(pady=(0, 15))

        mod = "Cmd" if IS_MACOS else "Strg"
        shortcuts = [
            (f"{mod}+O", "Datei öffnen"),
            (f"{mod}+V", "URL einfügen"),
            (f"{mod}+Return", "Transkription starten"),
            (f"{mod}+Q", "Programm beenden"),
            (f"{mod}+S", "Transkript speichern"),
            (f"{mod}+L", "Layout umschalten"),
            (f"{mod}+T", "Übersetzung ein/aus"),
            (f"{mod}+E", "Untertitel exportieren"),
            (f"{mod}+Umschalt+C", "Alles löschen"),
            ("F1", "Diese Hilfe anzeigen"),
        ]

        frame = tk.Frame(main, bg=CURRENT_THEME.BG_SECONDARY)
        frame.pack(fill="both", expand=True)

        for i, (key, desc) in enumerate(shortcuts):
            row = tk.Frame(
                frame,
                bg=CURRENT_THEME.BG_TERTIARY if i % 2 else CURRENT_THEME.BG_SECONDARY,
            )
            row.pack(fill="x", pady=1)

            key_label = tk.Label(
                row,
                text=key,
                font=("Segoe UI", 10, "bold"),
                bg=row["bg"],
                fg=CURRENT_THEME.TEXT_ACCENT,
                width=15,
                anchor="w",
            )
            key_label.pack(side="left", padx=(10, 5), pady=5)

            desc_label = tk.Label(
                row,
                text=desc,
                font=("Segoe UI", 10),
                bg=row["bg"],
                fg=CURRENT_THEME.TEXT_PRIMARY,
                anchor="w",
            )
            desc_label.pack(side="left", fill="x", expand=True, padx=5, pady=5)

        close_btn = tk.Button(
            main,
            text="Schließen",
            command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        )
        close_btn.pack(pady=10)

        self.dialog.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.dialog.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.dialog.winfo_height()) // 2
        self.dialog.geometry(f"+{x}+{y}")

    def _on_close(self):
        if hasattr(self.parent, '_open_dialogs') and self.dialog in self.parent._open_dialogs:
            self.parent._open_dialogs.remove(self.dialog)
        self.dialog.destroy()


# -----------------------------------------------------------------------------
# InstallDependencyDialog
# -----------------------------------------------------------------------------
class InstallDependencyDialog:
    def __init__(self, parent, gui_ref):
        self.parent = parent
        self.gui = gui_ref
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("🐉 Fehlende Pakete installieren")
        self.dialog.geometry("600x450")
        self.dialog.configure(bg=CURRENT_THEME.BG_PRIMARY)
        self.dialog.transient(parent)
        self.dialog.grab_set()
        self.dialog.protocol("WM_DELETE_WINDOW", self._on_close)

        if hasattr(self.gui, '_open_dialogs'):
            self.gui._open_dialogs.append(self.dialog)
            

        self._install_thread = None
        self._process = None
        self._stop_requested = False

        main = tk.Frame(self.dialog, bg=CURRENT_THEME.BG_PRIMARY, padx=15, pady=15)
        main.pack(fill="both", expand=True)

        tk.Label(
            main,
            text="Optionale Pakete, die installiert werden können:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        ).pack(anchor="w", pady=(0, 10))

        self.packages = {}
        frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        frame.pack(fill="x")

        if not WHISPER_AVAILABLE:
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(
                frame,
                text="faster-whisper (benötigt für Transkription)",
                variable=var,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                selectcolor=CURRENT_THEME.BG_TERTIARY,
                activebackground=CURRENT_THEME.BG_HOVER,
            )
            cb.pack(anchor="w")
            self.packages["faster-whisper"] = var

        if not TRANSLATOR_AVAILABLE:
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(
                frame,
                text="deep-translator (für Übersetzungen)",
                variable=var,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                selectcolor=CURRENT_THEME.BG_TERTIARY,
                activebackground=CURRENT_THEME.BG_HOVER,
            )
            cb.pack(anchor="w")
            self.packages["deep-translator"] = var

        if not FastLazyLoader.is_available("psutil"):
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(
                frame,
                text="psutil (Systemmonitoring)",
                variable=var,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                selectcolor=CURRENT_THEME.BG_TERTIARY,
                activebackground=CURRENT_THEME.BG_HOVER,
            )
            cb.pack(anchor="w")
            self.packages["psutil"] = var

        if not FastLazyLoader.is_available("pynvml"):
            var = tk.BooleanVar(value=True)
            cb = tk.Checkbutton(
                frame,
                text="pynvml (genaue GPU‑Auslastung)",
                variable=var,
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.TEXT_PRIMARY,
                selectcolor=CURRENT_THEME.BG_TERTIARY,
                activebackground=CURRENT_THEME.BG_HOVER,
            )
            cb.pack(anchor="w")
            self.packages["pynvml"] = var

        if not self.packages:
            tk.Label(
                main,
                text="✅ Alle optionalen Pakete sind bereits installiert.",
                bg=CURRENT_THEME.BG_PRIMARY,
                fg=CURRENT_THEME.SUCCESS,
            ).pack(pady=20)
            self.dialog.after(2000, self.dialog.destroy)
            return

        tk.Label(
            main,
            text="Installationsausgabe:",
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        ).pack(anchor="w", pady=(10, 2))
        self.output_text = scrolledtext.ScrolledText(
            main,
            height=12,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            wrap=tk.WORD,
            state="normal",
        )
        self.output_text.pack(fill="both", expand=True, pady=5)

        self.status_var = tk.StringVar(value="Bereit")
        status_label = tk.Label(
            main,
            textvariable=self.status_var,
            bg=CURRENT_THEME.BG_PRIMARY,
            fg=CURRENT_THEME.TEXT_SECONDARY,
            font=Fonts.SMALL,
        )
        status_label.pack(fill="x", pady=(5, 0))

        btn_frame = tk.Frame(main, bg=CURRENT_THEME.BG_PRIMARY)
        btn_frame.pack(fill="x", pady=5)

        self.install_btn = tk.Button(
            btn_frame,
            text="Installieren",
            command=self.install_selected,
            bg=CURRENT_THEME.DRAGON_GREEN,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
        )
        self.install_btn.pack(side="left", padx=5)

        self.cancel_btn = tk.Button(
            btn_frame,
            text="Abbrechen",
            command=self.cancel_installation,
            bg=CURRENT_THEME.ERROR,
            fg=CURRENT_THEME.TEXT_PRIMARY,
            font=Fonts.BUTTON,
            padx=15,
            state="disabled",
        )
        self.cancel_btn.pack(side="left", padx=5)

        self.close_btn = tk.Button(
            btn_frame,
            text="Schließen",
            command=self.dialog.destroy,
            bg=CURRENT_THEME.BG_TERTIARY,
            fg=CURRENT_THEME.TEXT_PRIMARY,
        )
        self.close_btn.pack(side="right", padx=5)

    def _on_close(self):
        self.cancel_installation()
        if hasattr(self.gui, '_open_dialogs') and self.dialog in self.gui._open_dialogs:
            self.gui._open_dialogs.remove(self.dialog)
        self.dialog.destroy()

    def cancel_installation(self):
        if self._process and self._process.poll() is None:
            self._stop_requested = True
            try:
                self._process.terminate()
                self._process.wait(timeout=2)
            except subprocess.TimeoutExpired:
                try:
                    self._process.kill()
                except Exception:
                    pass
            except Exception:
                pass
        self._stop_requested = False
        self._enable_ui(True)
        self.status_var.set("Installation abgebrochen")

    def _enable_ui(self, enabled: bool):
        state = "normal" if enabled else "disabled"
        self.install_btn.config(state=state)
        self.cancel_btn.config(state="disabled" if enabled else "normal")
        for var in self.packages.values():
            cb = var._widget
            if cb:
                cb.config(state=state)

    def install_selected(self):
        packages = [pkg for pkg, var in self.packages.items() if var.get()]
        if not packages:
            return

        self.output_text.delete("1.0", "end")
        self.output_text.insert(
            "end", f"Starte Installation von: {', '.join(packages)}...\n"
        )
        self.dialog.update()

        self._enable_ui(False)
        self.status_var.set("Installation läuft...")

        self._install_thread = threading.Thread(
            target=self._install_worker, args=(packages,), daemon=True
        )
        self._install_thread.start()

    def _install_worker(self, packages):
        python_exe = sys.executable
        cmd = [python_exe, "-m", "pip", "install"] + packages

        try:
            self._process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
                encoding="utf-8",
                errors="ignore",
            )

            for line in self._process.stdout:
                if self._stop_requested:
                    break
                self.dialog.after(0, self._append_output, line)
                time.sleep(0.01)

            returncode = self._process.wait()

            if self._stop_requested:
                self.dialog.after(0, self._installation_finished, "abgebrochen")
            elif returncode == 0:
                self.dialog.after(0, self._installation_finished, "erfolgreich")
            else:
                self.dialog.after(
                    0,
                    self._installation_finished,
                    f"fehlgeschlagen (Code {returncode})",
                )

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.dialog.after(0, self._installation_finished, f"Fehler: {str(e)}")
        finally:
            self._process = None

    def _append_output(self, line: str):
        try:
            self.output_text.insert("end", line)
            self.output_text.see("end")
        except tk.TclError:
            pass

    def _installation_finished(self, status: str):
        if status == "erfolgreich":
            self.output_text.insert("end", "\n✅ Installation erfolgreich!\n")
            self.output_text.insert(
                "end",
                "Bitte starten Sie das Programm neu, um die neuen Pakete zu nutzen.\n",
            )
            self.status_var.set("✅ Erfolgreich installiert")
        elif status == "abgebrochen":
            self.output_text.insert("end", "\n⏹️ Installation abgebrochen.\n")
            self.status_var.set("⏹️ Abgebrochen")
        else:
            self.output_text.insert("end", f"\n❌ {status}\n")
            self.status_var.set("❌ Fehlgeschlagen")

        self._enable_ui(True)
        self.cancel_btn.config(state="disabled")
        self._install_thread = None


# -----------------------------------------------------------------------------
# WhisperLayoutManager
# -----------------------------------------------------------------------------
class WhisperLayoutManager:
    def __init__(self, gui_ref: "DragonWhispererGUI") -> None:
        self.gui_ref = gui_ref
        self.root = gui_ref.root
        self._batch_timer_id: Optional[str] = None
        try:
            self.gui_ref._text_update_queue = queue.Queue(maxsize=150)
            self.gui_ref.gui_queue = queue.Queue(maxsize=200)
            logger.info("✅ Queues erfolgreich erstellt")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Queue-Erstellung fehlgeschlagen: {e}")
            self.gui_ref._text_update_queue = DummyQueue(maxsize=150)
            self.gui_ref.gui_queue = DummyQueue(maxsize=200)
            logger.warning("⚠️ Verwende Dummy-Queues (eingeschränkte Funktionalität)")

    def setup_gui(self) -> None:
        self.root.configure(bg=self.gui_ref.current_theme.BG_PRIMARY)
        self.root.title("🐉 Dragon Whisperer - Plattformunabhängig")
        self.root.geometry("900x700")
        self.root.minsize(850, 600)
        self.root.grid_rowconfigure(0, weight=0)
        self.root.grid_rowconfigure(1, weight=0)
        self.root.grid_rowconfigure(2, weight=0)
        self.root.grid_rowconfigure(3, weight=10)
        self.root.grid_rowconfigure(4, weight=0)
        self.root.grid_columnconfigure(0, weight=1)
        self.setup_dark_styles()
        self.center_window()
        self.root.protocol("WM_DELETE_WINDOW", self.gui_ref._safe_exit_dialog)
        self.create_layout()
        self.root.after(100, self.start_batch_updates)

    def setup_dark_styles(self) -> None:
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(
            "Dark.TCombobox",
            fieldbackground=self.gui_ref.current_theme.COMBO_BG,
            background=self.gui_ref.current_theme.COMBO_BG,
            foreground=self.gui_ref.current_theme.COMBO_FG,
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            selectforeground=self.gui_ref.current_theme.TEXT_PRIMARY,
            insertcolor=self.gui_ref.current_theme.TEXT_PRIMARY,
            borderwidth=1,
            relief="flat",
            arrowsize=12,
            padding=5,
        )
        style.map(
            "Dark.TCombobox",
            fieldbackground=[
                ("readonly", self.gui_ref.current_theme.COMBO_BG),
                ("active", self.gui_ref.current_theme.BG_HOVER),
            ],
            background=[
                ("readonly", self.gui_ref.current_theme.COMBO_BG),
                ("active", self.gui_ref.current_theme.BG_HOVER),
            ],
            foreground=[
                ("readonly", self.gui_ref.current_theme.COMBO_FG),
                ("active", self.gui_ref.current_theme.TEXT_PRIMARY),
            ],
        )
        style.configure(
            "Dark.Horizontal.TProgressbar",
            background=self.gui_ref.current_theme.SUCCESS,
            troughcolor=self.gui_ref.current_theme.BG_TERTIARY,
            bordercolor=self.gui_ref.current_theme.BORDER,
        )
        self.root.option_add(
            "*TCombobox*Listbox.background", self.gui_ref.current_theme.COMBO_BG
        )
        self.root.option_add(
            "*TCombobox*Listbox.foreground", self.gui_ref.current_theme.COMBO_FG
        )
        self.root.option_add(
            "*TCombobox*Listbox.selectBackground",
            self.gui_ref.current_theme.COMBO_SELECTION,
        )
        self.root.option_add(
            "*TCombobox*Listbox.selectForeground",
            self.gui_ref.current_theme.TEXT_PRIMARY,
        )

    def center_window(self) -> None:
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f"+{x}+{y}")

    def create_layout(self) -> None:
        header_frame = tk.Frame(
            self.root, bg=self.gui_ref.current_theme.BG_PRIMARY, height=35
        )
        header_frame.grid(row=0, column=0, sticky="ew", padx=12, pady=8)
        header_frame.grid_propagate(False)
        title_label = tk.Label(
            header_frame,
            text="🐉 Dragon Whisperer - Livestream Transcription & Translation",
            font=Fonts.TITLE,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.DRAGON_GREEN,
        )
        title_label.pack(side="left")
        self.gui_ref.status_label = tk.Label(
            header_frame,
            text="✅ READY",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
        )
        self.gui_ref.status_label.pack(side="right")
        self.create_stream_info_display()
        self.gui_ref.stream_info_frame.grid(
            row=1, column=0, sticky="ew", padx=12, pady=3
        )
        input_frame = tk.Frame(self.root, bg=self.gui_ref.current_theme.BG_PRIMARY)
        input_frame.grid(row=2, column=0, sticky="ew", padx=12, pady=3)
        url_frame = tk.Frame(input_frame, bg=self.gui_ref.current_theme.BG_PRIMARY)
        url_frame.pack(fill="x", pady=2)
        tk.Label(
            url_frame,
            text="URL:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.url_entry = tk.Entry(
            url_frame,
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            insertbackground=self.gui_ref.current_theme.TEXT_PRIMARY,
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            width=60,
        )
        self.gui_ref.url_entry.pack(side="left", fill="x", expand=True, padx=(5, 5))
        self.gui_ref.url_entry.insert(
            0, self.gui_ref.settings.last_url if self.gui_ref.settings.last_url else ""
        )
        DarkEntryContextMenu(self.gui_ref.url_entry)
        self.gui_ref.language_info_label = tk.Label(
            url_frame,
            text="",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
        )
        self.gui_ref.language_info_label.pack(side="right", padx=(5, 0))
        self.create_compact_control_panel(input_frame)
        self.setup_status_bar()
        self.gui_ref.status_bar_frame.grid(row=4, column=0, sticky="ew", pady=(2, 0))
        self.create_text_areas()
        self.gui_ref.text_container.grid(
            row=3, column=0, sticky="nsew", padx=12, pady=8
        )
        self.gui_ref.url_entry.bind("<KeyRelease>", self.gui_ref.on_url_change)
        self.gui_ref.url_entry.bind("<FocusOut>", self.gui_ref.on_url_change)

    def create_stream_info_display(self) -> None:
        self.gui_ref.stream_info_frame = tk.Frame(
            self.root, bg=self.gui_ref.current_theme.BG_SECONDARY, height=50
        )
        self.gui_ref.stream_info_frame.grid_propagate(True)
        self.gui_ref.stream_title_label = tk.Label(
            self.gui_ref.stream_info_frame,
            text="📡 No active stream",
            font=Fonts.SUBTITLE,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            wraplength=700,
            justify="left",
        )
        self.gui_ref.stream_title_label.pack(fill="x", padx=8, pady=(6, 2))
        self.gui_ref.stream_details_label = tk.Label(
            self.gui_ref.stream_info_frame,
            text="Ready to connect...",
            font=Fonts.PRIMARY,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            justify="left",
        )
        self.gui_ref.stream_details_label.pack(fill="x", padx=8, pady=(2, 6))

    def create_compact_control_panel(self, parent: tk.Frame) -> None:
        control_frame = tk.Frame(parent, bg=self.gui_ref.current_theme.BG_PRIMARY)
        control_frame.pack(fill="x", pady=8)
        left_controls = tk.Frame(
            control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY
        )
        left_controls.pack(side="left")
        action_buttons = [
            ("📁", self.gui_ref.select_file_dark, "Datei auswählen"),
            ("📋", self.gui_ref.paste_url, "URL aus Zwischenablage einfügen"),
        ]
        for icon, command, tooltip in action_buttons:
            btn = tk.Button(
                left_controls,
                text=icon,
                command=command,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                bd=0,
                font=("Segoe UI", 9),
                cursor="hand2",
            )
            btn.pack(side="left", padx=1)
            ToolTip(btn, tooltip)

        self.gui_ref.layout_btn = tk.Button(
            left_controls,
            text="🔄",
            command=self.gui_ref.toggle_layout,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            bd=0,
            font=("Segoe UI", 9),
            cursor="hand2",
        )
        self.gui_ref.layout_btn.pack(side="left", padx=5)
        ToolTip(self.gui_ref.layout_btn, "Layout umschalten (vertikal/horizontal)")

        center_controls = tk.Frame(
            control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY
        )
        center_controls.pack(side="left", padx=15)

        src_lang_frame = tk.Frame(
            center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY
        )
        src_lang_frame.pack(side="left", padx=5)
        tk.Label(
            src_lang_frame,
            text="From:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.src_lang_var = tk.StringVar(value="Automatisch")
        self.gui_ref.src_lang_combo = ttk.Combobox(
            src_lang_frame,
            textvariable=self.gui_ref.src_lang_var,
            values=[name for name, code in SORTED_LANGUAGES],
            width=10,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.src_lang_combo.pack(side="left", padx=3)
        ToolTip(
            self.gui_ref.src_lang_combo,
            "Quellsprache (Automatisch = Whisper-Erkennung)",
        )

        model_frame = tk.Frame(
            center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY
        )
        model_frame.pack(side="left", padx=5)
        tk.Label(
            model_frame,
            text="Model:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.model_var = tk.StringVar(value=self.gui_ref.settings.default_model)
        self.gui_ref.model_combo = ttk.Combobox(
            model_frame,
            textvariable=self.gui_ref.model_var,
            values=WHISPER_MODELS,
            width=8,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.model_combo.pack(side="left", padx=3)
        ToolTip(
            self.gui_ref.model_combo,
            "Whisper-Modell auswählen (größer = genauer, aber langsamer)",
        )

        if getattr(self.gui_ref, "demo_mode", False):
            self.gui_ref.model_combo.config(state="disabled")
            self.gui_ref.model_var.set("dummy (Demo)")

        lang_frame = tk.Frame(center_controls, bg=self.gui_ref.current_theme.BG_PRIMARY)
        lang_frame.pack(side="left", padx=5)
        tk.Label(
            lang_frame,
            text="Translate:",
            bg=self.gui_ref.current_theme.BG_PRIMARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=Fonts.PRIMARY,
        ).pack(side="left")
        self.gui_ref.lang_var = tk.StringVar()
        language_groups = {
            "Common": ["German", "English", "French", "Spanish", "Italian"],
            "Asian": ["Japanese", "Chinese", "Korean", "Vietnamese", "Thai"],
            "More": [
                name
                for name, code in SORTED_LANGUAGES
                if name
                not in [
                    "German",
                    "English",
                    "French",
                    "Spanish",
                    "Italian",
                    "Japanese",
                    "Chinese",
                    "Korean",
                    "Vietnamese",
                    "Thai",
                ]
            ],
        }
        all_languages: List[str] = []
        for group_name, languages in language_groups.items():
            if languages:
                all_languages.append(f"--- {group_name} ---")
                all_languages.extend(languages)
        self.gui_ref.lang_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.gui_ref.lang_var,
            values=all_languages,
            width=12,
            style="Dark.TCombobox",
            state="readonly",
        )
        self.gui_ref.lang_combo.pack(side="left", padx=3)
        ToolTip(self.gui_ref.lang_combo, "Zielsprache für Übersetzung")
        default_lang_name = SUPPORTED_LANGUAGES.get(
            self.gui_ref.settings.default_language, "German"
        )
        self.gui_ref.lang_var.set(default_lang_name)
        self.gui_ref.lang_combo.bind(
            "<<ComboboxSelected>>", self.gui_ref.on_language_change
        )

        right_controls = tk.Frame(
            control_frame, bg=self.gui_ref.current_theme.BG_PRIMARY
        )
        right_controls.pack(side="right")

        self.gui_ref.start_button = tk.Button(
            right_controls,
            text="🚀 START",
            command=self.gui_ref._on_start_click,
            bg=self.gui_ref.current_theme.SUCCESS,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=("Segoe UI", 9, "bold"),
            relief="flat",
            padx=20,
        )
        self.gui_ref.start_button.pack(side="left", padx=2)
        ToolTip(self.gui_ref.start_button, "Transkription/Übersetzung starten")

        self.gui_ref.stop_button = tk.Button(
            right_controls,
            text="⏹️ STOP",
            command=self.gui_ref.controller.stop_processing,
            bg=self.gui_ref.current_theme.ERROR,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            state="disabled",
            font=("Segoe UI", 9, "bold"),
            relief="flat",
            padx=20,
        )
        self.gui_ref.stop_button.pack(side="left", padx=2)
        ToolTip(self.gui_ref.stop_button, "Laufende Verarbeitung stoppen")

        self.gui_ref.translate_btn = tk.Button(
            right_controls,
            text="🌐 ON" if self.gui_ref.translate_active else "🌐 OFF",
            command=self.gui_ref.toggle_translation,
            bg=(
                self.gui_ref.current_theme.SUCCESS
                if self.gui_ref.translate_active
                else self.gui_ref.current_theme.BG_TERTIARY
            ),
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            padx=6,
        )
        self.gui_ref.translate_btn.pack(side="left", padx=2)
        ToolTip(self.gui_ref.translate_btn, "Übersetzung ein/aus")

        self.gui_ref.subtitle_btn = tk.Button(
            right_controls,
            text="🎬",
            command=self.gui_ref.toggle_subtitle_mode,
            bg=self.gui_ref.current_theme.SUBTITLE_INACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            bd=0,
            font=("Segoe UI", 9),
            cursor="hand2",
        )
        self.gui_ref.subtitle_btn.pack(side="left", padx=5)
        ToolTip(self.gui_ref.subtitle_btn, "Untertitel-Modus (mit Zeitstempeln)")

        self.gui_ref.model_combo.bind(
            "<<ComboboxSelected>>", self.gui_ref.on_model_change
        )

    def create_text_areas(
        self,
    ) -> Tuple[
        Optional[scrolledtext.ScrolledText], Optional[scrolledtext.ScrolledText]
    ]:
        layout_changed = False
        current_layout = getattr(self.gui_ref, "_current_layout", None)
        if current_layout != self.gui_ref.layout_mode:
            layout_changed = True
            logger.info(
                f"🔄 Layout change detected: {current_layout} → {self.gui_ref.layout_mode}"
            )
        if hasattr(self.gui_ref, "text_container") and layout_changed:
            try:
                if self.gui_ref.text_container.winfo_exists():
                    logger.info("   🗑️ Destroying old container for layout change")
                    self.gui_ref.text_container.destroy()
                    time.sleep(0.02)
            except tk.TclError:
                pass
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"   ⚠️ Container destroy warning: {e}")
        if layout_changed or not hasattr(self.gui_ref, "text_container"):
            self.gui_ref.text_container = tk.Frame(
                self.root, bg=self.gui_ref.current_theme.BG_PRIMARY
            )
            self.gui_ref._current_layout = self.gui_ref.layout_mode
            logger.info(
                f"   ✅ New container created for {self.gui_ref.layout_mode} layout"
            )
        if self.gui_ref.layout_mode == "horizontal":
            self.create_horizontal_layout()
        else:
            self.create_vertical_layout()
        self.gui_ref.text_container.grid(
            row=3, column=0, sticky="nsew", padx=12, pady=8
        )
        self.root.grid_rowconfigure(3, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        self.root.update_idletasks()
        return (
            getattr(self.gui_ref, "transcript_text", None),
            getattr(self.gui_ref, "translation_text", None),
        )

    def create_vertical_layout(self) -> None:
        main_frame = tk.LabelFrame(
            self.gui_ref.text_container,
            text="Live Transkription & Übersetzung",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.SUBTITLE,
            padx=8,
            pady=8,
        )
        main_frame.pack(fill="both", expand=True)
        trans_frame = tk.Frame(main_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        trans_frame.pack(fill="x", pady=(0, 3))
        trans_header = tk.Frame(trans_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        trans_header.pack(fill="x")
        tk.Label(
            trans_header,
            text="🎤 Transkription:",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        ).pack(side="left")
        self.gui_ref.transcript_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            trans_header,
            variable=self.gui_ref.transcript_scroll_var,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            activebackground=self.gui_ref.current_theme.BG_SECONDARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            trans_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.transcript_text = self.create_text_widget(main_frame, height=6)
        transla_frame = tk.Frame(main_frame, bg=self.gui_ref.current_theme.BG_SECONDARY)
        transla_frame.pack(fill="x", pady=(8, 0))
        transla_header = tk.Frame(
            transla_frame, bg=self.gui_ref.current_theme.BG_SECONDARY
        )
        transla_header.pack(fill="x")
        self.gui_ref.translation_header = tk.Label(
            transla_header,
            text="🌐 Übersetzung:",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        )
        self.gui_ref.translation_header.pack(side="left")
        self.gui_ref.translation_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            transla_header,
            variable=self.gui_ref.translation_scroll_var,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            activebackground=self.gui_ref.current_theme.BG_SECONDARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            transla_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.translation_text = self.create_text_widget(main_frame, height=6)

    def create_horizontal_layout(self) -> None:
        main_frame = tk.LabelFrame(
            self.gui_ref.text_container,
            text="Live Transkription & Übersetzung",
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.SUBTITLE,
            padx=8,
            pady=8,
        )
        main_frame.pack(fill="both", expand=True)
        self.gui_ref.paned_window = tk.PanedWindow(
            main_frame,
            orient=tk.HORIZONTAL,
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            sashrelief="raised",
            sashwidth=4,
            sashpad=0,
        )
        self.gui_ref.paned_window.pack(fill="both", expand=True)
        left_frame = tk.Frame(
            self.gui_ref.paned_window, bg=self.gui_ref.current_theme.BG_TERTIARY
        )
        self.gui_ref.paned_window.add(left_frame, stretch="always", width=400)
        trans_header = tk.Frame(left_frame, bg=self.gui_ref.current_theme.BG_TERTIARY)
        trans_header.pack(fill="x", padx=5, pady=2)
        tk.Label(
            trans_header,
            text="🎤 Transkription",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        ).pack(side="left")
        self.gui_ref.transcript_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            trans_header,
            variable=self.gui_ref.transcript_scroll_var,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            activebackground=self.gui_ref.current_theme.BG_TERTIARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            trans_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.transcript_text = self.create_text_widget(left_frame)
        right_frame = tk.Frame(
            self.gui_ref.paned_window, bg=self.gui_ref.current_theme.BG_TERTIARY
        )
        self.gui_ref.paned_window.add(right_frame, stretch="always", width=400)
        transla_header = tk.Frame(
            right_frame, bg=self.gui_ref.current_theme.BG_TERTIARY
        )
        transla_header.pack(fill="x", padx=5, pady=2)
        self.gui_ref.translation_header = tk.Label(
            transla_header,
            text="🌐 Übersetzung",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_ACCENT,
            font=Fonts.SUBTITLE,
        )
        self.gui_ref.translation_header.pack(side="left")
        self.gui_ref.translation_scroll_var = tk.BooleanVar(value=True)
        scroll_cb = tk.Checkbutton(
            transla_header,
            variable=self.gui_ref.translation_scroll_var,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            activebackground=self.gui_ref.current_theme.BG_TERTIARY,
            selectcolor=self.gui_ref.current_theme.CHECKBOX_ACTIVE,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
        )
        scroll_cb.pack(side="right", padx=3)
        tk.Label(
            transla_header,
            text="Auto-Scroll",
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            font=("Segoe UI", 7),
        ).pack(side="right", padx=1)
        self.gui_ref.translation_text = self.create_text_widget(right_frame)
        self.gui_ref.paned_window.paneconfig(left_frame, minsize=250, width=400)
        self.gui_ref.paned_window.paneconfig(right_frame, minsize=250, width=400)

    def create_text_widget(
        self, parent: tk.Frame, height: Optional[int] = None
    ) -> scrolledtext.ScrolledText:
        text_widget = scrolledtext.ScrolledText(
            parent,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            font=Fonts.MONOSPACE,
            insertbackground=self.gui_ref.current_theme.TEXT_PRIMARY,
            wrap=tk.WORD,
            relief="flat",
            selectbackground=self.gui_ref.current_theme.COMBO_SELECTION,
            selectforeground=self.gui_ref.current_theme.TEXT_PRIMARY,
            maxundo=30,
            undo=True,
        )
        if height:
            text_widget.config(height=height)
        text_widget.pack(fill="both", expand=True, padx=5, pady=5)
        DarkContextMenu(text_widget)
        return text_widget

    def setup_status_bar(self) -> None:
        self.gui_ref.status_bar_frame = tk.Frame(
            self.root, bg=self.gui_ref.current_theme.BG_SECONDARY, height=50
        )
        self.gui_ref.status_bar_frame.grid_propagate(True)
        separator = tk.Frame(
            self.gui_ref.status_bar_frame,
            height=2,
            bg=self.gui_ref.current_theme.DRAGON_GREEN,
        )
        separator.pack(fill="x", side="top")
        main_container = tk.Frame(
            self.gui_ref.status_bar_frame, bg=self.gui_ref.current_theme.BG_SECONDARY
        )
        main_container.pack(fill="x", expand=True, padx=12, pady=8)

        main_container.columnconfigure(0, weight=0)
        main_container.columnconfigure(1, weight=1)
        main_container.columnconfigure(2, weight=0)

        left_panel = tk.Frame(
            main_container, bg=self.gui_ref.current_theme.BG_SECONDARY
        )
        left_panel.grid(row=0, column=0, sticky="w", padx=5)

        quick_actions = [
            ("🗑️", self.gui_ref.clear_all, "Alles löschen"),
            ("💾", self.gui_ref.save_transcript, "Transkription speichern"),
            ("📝", self.gui_ref.export_subtitles, "Untertitel exportieren"),
            ("📊", self.gui_ref.show_simple_stats, "Statistiken anzeigen"),
            ("⚙️", self.gui_ref.show_advanced_settings, "Erweiterte Einstellungen"),
            ("🌐", self.gui_ref.show_translation_dialog, "Text übersetzen"),
            ("🤖", self.gui_ref.show_summarize_dialog, "Mit Ollama zusammenfassen"),
        ]

        if getattr(self.gui_ref, "demo_mode", False) or not TRANSLATOR_AVAILABLE:
            install_btn = tk.Button(
                left_panel,
                text="📦",
                command=self.gui_ref.show_install_dialog,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4,
                pady=2,
                activebackground=self.gui_ref.current_theme.BG_HOVER,
            )
            install_btn.grid(row=0, column=len(quick_actions) + 1, padx=1, sticky="w")
            ToolTip(install_btn, "Fehlende Pakete installieren")

        for i, (icon, command, tooltip) in enumerate(quick_actions):
            btn = tk.Button(
                left_panel,
                text=icon,
                command=command,
                bg=self.gui_ref.current_theme.BG_TERTIARY,
                fg=self.gui_ref.current_theme.TEXT_PRIMARY,
                relief="flat",
                font=("Segoe UI", 9),
                cursor="hand2",
                padx=4,
                pady=2,
                activebackground=self.gui_ref.current_theme.BG_HOVER,
            )
            btn.grid(row=0, column=i, padx=1, sticky="w")
            ToolTip(btn, tooltip)

        center_panel = tk.Frame(
            main_container, bg=self.gui_ref.current_theme.BG_SECONDARY
        )
        center_panel.grid(row=0, column=1, sticky="ew", padx=5)

        self.gui_ref.progress_bar = ttk.Progressbar(
            center_panel,
            mode="determinate",
            length=150,
            style="Dark.Horizontal.TProgressbar",
        )
        self.gui_ref.progress_bar.pack(side="left", padx=(10, 10))

        self.gui_ref.progress_label = tk.Label(
            center_panel,
            text="",
            font=("Segoe UI", 8),
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
        )
        self.gui_ref.progress_label.pack(side="left", padx=(0, 10))

        if IS_WINDOWS:
            default_text = "🪟 Windows | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        elif IS_MACOS:
            if IS_ARM:
                default_text = "🍎 macOS (Apple Silicon) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
            else:
                default_text = (
                    "🍎 macOS (Intel) | CPU: --% | RAM: --MB | GPU: --% | Model: --"
                )
        elif IS_LINUX:
            default_text = "🐧 Linux | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        else:
            default_text = "🌐 Unknown OS | CPU: --% | RAM: --MB | GPU: --% | Model: --"
        self.gui_ref.system_info_label = tk.Label(
            center_panel,
            text=default_text,
            font=("Segoe UI", 8, "normal"),
            bg=self.gui_ref.current_theme.BG_SECONDARY,
            fg=self.gui_ref.current_theme.TEXT_SECONDARY,
            padx=5,
        )
        self.gui_ref.system_info_label.pack(side="left", fill="x", expand=True)

        right_panel = tk.Frame(
            main_container, bg=self.gui_ref.current_theme.BG_SECONDARY
        )
        right_panel.grid(row=0, column=2, sticky="e", padx=5)

        self.gui_ref.exit_button = tk.Button(
            right_panel,
            text=" ⏻ EXIT ",
            command=self.gui_ref.controller.safe_exit,
            bg="#dc3545",
            fg="white",
            font=("Segoe UI", 9, "bold"),
            relief="raised",
            cursor="hand2",
            padx=12,
            pady=3,
            activebackground="#c82333",
        )
        self.gui_ref.exit_button.pack(side="right")
        ToolTip(self.gui_ref.exit_button, "Programm beenden (Strg+Q / Cmd+Q)")

        self.gui_ref.correct_btn = tk.Button(
            right_panel,
            text="🔧",
            command=self.gui_ref.correct_transcript_with_ollama,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        self.gui_ref.correct_btn.pack(side="right", padx=2)
        ToolTip(self.gui_ref.correct_btn, "Transkript mit Ollama korrigieren")

        help_btn = tk.Button(
            right_panel,
            text="⌨️",
            command=self.gui_ref.show_shortcuts_help,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        help_btn.pack(side="right", padx=2)
        ToolTip(help_btn, "Tastenkürzel anzeigen (F1)")

        self.gui_ref.tts_btn = tk.Button(
            right_panel,
            text="🔊",
            command=self.gui_ref.speak_current_text,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 9),
            cursor="hand2",
            padx=4,
        )
        self.gui_ref.tts_btn.pack(side="right", padx=2)
        ToolTip(self.gui_ref.tts_btn, "Ausgewählten Text vorlesen (TTS)")

        # VAD-Fallback-Button (Status wird später aktualisiert)
        self.gui_ref.vad_fallback_btn = tk.Button(
            right_panel,
            text="🔁 VAD-Fallback ON",
            command=self.gui_ref.toggle_vad_fallback,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 8),
            padx=4,
        )
        self.gui_ref.vad_fallback_btn.pack(side="right", padx=2)
        ToolTip(self.gui_ref.vad_fallback_btn, "VAD-Fallback aktivieren/deaktivieren (wenn aus, werden leere Chunks ignoriert)")

        # Live-Mode-Umschalter
        self.gui_ref.live_mode_btn = tk.Button(
            right_panel,
            text="⏱️ 20s",
            command=self.gui_ref.toggle_live_mode,
            bg=self.gui_ref.current_theme.BG_TERTIARY,
            fg=self.gui_ref.current_theme.TEXT_PRIMARY,
            relief="flat",
            font=("Segoe UI", 8),
            padx=4,
        )
        self.gui_ref.live_mode_btn.pack(side="right", padx=2)
        ToolTip(self.gui_ref.live_mode_btn, "Chunk-Dauer umschalten (20s/10s)")

    def process_batch_text_updates(self) -> None:
        if not hasattr(self.gui_ref, "_shutting_down") or getattr(
            self.gui_ref, "_shutting_down", False
        ):
            return
        if (
            not hasattr(self, "root")
            or self.root is None
            or not self.root.winfo_exists()
        ):
            return
        if not hasattr(self.gui_ref, "_text_update_queue"):
            return
        queue_obj = self.gui_ref._text_update_queue
        if queue_obj is None:
            return
        if not hasattr(queue_obj, "empty") or not callable(queue_obj.empty):
            return
        if queue_obj.empty():
            return

        if logger.isEnabledFor(logging.DEBUG):
            logger.debug(
                f"[QUEUE] process_batch_text_updates: Queue-Größe vor Verarbeitung: {queue_obj.qsize()}"
            )

        processed = 0
        max_updates = 5
        start_time = time.time()
        while processed < max_updates and (time.time() - start_time) < 0.05:
            if not self.root.winfo_exists():
                break
            try:
                if hasattr(queue_obj, "get_nowait") and callable(queue_obj.get_nowait):
                    item = queue_obj.get_nowait()
                else:
                    item = queue_obj.get(block=False)
                if isinstance(item, tuple) and len(item) == 2:
                    update_type, text_data = item
                    self._process_update(update_type, text_data)
                processed += 1
                try:
                    self.root.update_idletasks()
                except tk.TclError:
                    break
            except Exception as e:
                if "Empty" in type(e).__name__ or "empty" in str(e).lower():
                    break
                logger.warning(f"⚠️ Queue processing error: {e}")
                break
        if logger.isEnabledFor(logging.DEBUG) and processed > 0:
            logger.debug(
                f"[QUEUE] Verarbeitet: {processed} Elemente, verbleibende Größe: {queue_obj.qsize()}"
            )
        self._schedule_next_update()

    def _process_update(self, update_type: str, text_data: str) -> None:
        if not hasattr(self, "gui_ref") or self.gui_ref is None:
            return
        try:
            if update_type == "transcript":
                widget = getattr(self.gui_ref, "transcript_text", None)
                if widget is not None and widget.winfo_exists():
                    widget.insert("end", text_data)
                    self._auto_scroll("transcript")
            elif update_type == "translation":
                widget = getattr(self.gui_ref, "translation_text", None)
                if widget is not None and widget.winfo_exists():
                    widget.insert("end", text_data)
                    self._auto_scroll("translation")
        except tk.TclError:
            pass
        except AttributeError:
            pass
        except Exception as e:
            logger.warning(f"⚠️ GUI update error: {e}")

    def _auto_scroll(self, text_type: str) -> None:
        try:
            if text_type == "transcript":
                if (
                    hasattr(self.gui_ref, "transcript_scroll_var")
                    and self.gui_ref.transcript_scroll_var is not None
                    and self.gui_ref.transcript_scroll_var.get()
                ):
                    self.gui_ref.transcript_text.see("end")
            elif text_type == "translation":
                if (
                    hasattr(self.gui_ref, "translation_scroll_var")
                    and self.gui_ref.translation_scroll_var is not None
                    and self.gui_ref.translation_scroll_var.get()
                ):
                    self.gui_ref.translation_text.see("end")
        except Exception:
            pass

    def _schedule_next_update(self) -> None:
        try:
            if (
                hasattr(self, "root")
                and self.root is not None
                and self.root.winfo_exists()
            ):
                interval = 150
                if hasattr(self.gui_ref, "_batch_update_interval"):
                    try:
                        interval = self.gui_ref._batch_update_interval
                    except Exception:
                        pass
                if hasattr(self, "_batch_timer_id") and self._batch_timer_id:
                    try:
                        self.root.after_cancel(self._batch_timer_id)
                    except Exception:
                        pass
                self._batch_timer_id = self.root.after(
                    interval, self.process_batch_text_updates
                )
            else:
                self._batch_timer_id = None
        except Exception as e:
            logger.warning(f"⚠️ Timer scheduling error: {e}")

    def start_batch_updates(self) -> None:
        try:
            if (
                hasattr(self, "root")
                and self.root is not None
                and self.root.winfo_exists()
            ):
                if (
                    not hasattr(self.gui_ref, "_text_update_queue")
                    or self.gui_ref._text_update_queue is None
                ):
                    try:
                        self.gui_ref._text_update_queue = queue.Queue(maxsize=150)
                    except Exception:
                        self.gui_ref._text_update_queue = DummyQueue(maxsize=150)
                        logger.warning("⚠️ Queue-Fallback in start_batch_updates")
                self.root.after(100, self.process_batch_text_updates)
                logger.info("✅ Batch updates gestartet")
        except Exception as e:
            logger.warning(f"⚠️ Start batch updates error: {e}")


# =============================================================================
# DragonWhispererGUI – Hauptklasse
class DragonWhispererGUI:
    class RateLimiter:
        def __init__(self, max_updates_per_second: int = 30) -> None:
            self.min_interval = 1.0 / max_updates_per_second
            self.last_calls: Dict[str, float] = {}
            self._lock = threading.RLock()

        def can_update(self, update_type: str = "default") -> bool:
            with self._lock:
                now = time.time()
                if update_type not in self.last_calls:
                    self.last_calls[update_type] = 0.0
                last = self.last_calls[update_type]
                if now - last >= self.min_interval:
                    self.last_calls[update_type] = now
                    return True
                return False

        def reset(self, update_type: Optional[str] = None) -> None:
            with self._lock:
                if update_type is None:
                    self.last_calls.clear()
                elif update_type in self.last_calls:
                    del self.last_calls[update_type]

    ASIAN_LANGUAGES = ['zh', 'ja', 'ko', 'th', 'vi']

    def __init__(self) -> None:
        self._gui_update_limiter = self.RateLimiter(max_updates_per_second=15)
        self._shutting_down = False
        self._exit_dialog_active = False
        self.is_processing = False
        self.subtitle_mode = False
        self.exit_confirmed = False
        self.current_stream_info: Optional[StreamInfo] = None
        self.current_video_language: Optional[str] = None
        self._progress_bar_started = False
        self.translate_active = True
        self.translation_executor = ThreadPoolExecutor(
            max_workers=1, thread_name_prefix="Translation"
        )
        self._history_lock = threading.RLock()
        self._duplicate_lock = threading.RLock()

        self._open_dialogs: List[tk.Toplevel] = []

        if not GUI_AVAILABLE:
            logger.error("❌ Tkinter nicht verfügbar. Versuche Fallback...")
            self._try_fallback_gui()
            return

        try:
            self.settings = AppSettings.load_from_file()
            if not self.settings.last_url:
                self.settings.last_url = ""
            self.advanced_settings = AdvancedSettings.load_from_file()
            self.advanced_settings.repair()
            validation_issues = self.advanced_settings.validate()
            if validation_issues:
                logger.warning(f"⚠️ Settings validation issues: {validation_issues}")
            logger.debug(f"Blacklist geladen: {self.advanced_settings.blacklist}") 
            logger.info(
                f"✅ Settings ready: SAMPLE_RATE={self.advanced_settings.config.SAMPLE_RATE}, "
                f"CHANNELS={self.advanced_settings.config.CHANNELS}, "
                f"CHUNK_SIZE_BYTES={self.advanced_settings.config.CHUNK_SIZE_BYTES}"
            )
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Settings load failed: {e}, using defaults")
            self.settings = AppSettings()
            self.settings.last_url = ""
            self.advanced_settings = AdvancedSettings()

        self._apply_precision_optimizations()

        if not self.settings.cookies_notice_shown and self.settings.use_browser_cookies:
            self._show_cookie_notice = True
        else:
            self._show_cookie_notice = False

        self.app_context = AppContext()
        self.current_theme = self.app_context.theme
        self.demo_mode = not WHISPER_AVAILABLE
        self.layout_mode = getattr(self.settings, "layout_mode", "vertical")
        self.current_language = getattr(self.settings, "default_language", "de")
        self._translation_reset_counter = 0
        self.progress_dialog: Optional[ProgressDialog] = None

        self._last_valid_language = SUPPORTED_LANGUAGES.get(
            self.settings.default_language, "German"
        )

        try:
            self.root = tk.Tk()
            self.root.withdraw()
        except (tk.TclError, RuntimeError) as e:
            raise RuntimeError(f"Tkinter Fehler: {e}")

        self._batch_update_interval = 150
        self._last_batch_update = 0.0
        self._last_gui_update_time = 0.0
        self._processing_lock = threading.Lock()
        self.transcript_history: Deque[TranscriptionResult] = deque(maxlen=1000)
        self.translation_history: Deque[TranslationResult] = deque(maxlen=500)
        self._last_transcription_text = ""
        self._last_translation_text = ""
        self.performance_monitor = SimplePerformanceTracker()
        self.gui_queue: queue.Queue = queue.Queue(maxsize=200)
        self._text_update_queue: queue.Queue = queue.Queue(maxsize=150)
        self._last_text_queue_size = 0

        self.stream_info_extractor = StreamInfoExtractor()
        self.stream_info_extractor.use_browser_cookies = (
            self.settings.use_browser_cookies
        )

        try:
            self.controller = WhisperController(gui_ref=self)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Controller Fehler: {e}")
            self._show_error_and_exit(f"Controller Fehler: {e}")
            return
        self.controller.set_callbacks(
            on_transcription=self.handle_transcription,
            on_translation=self.handle_translation,
            on_info=self.handle_info,
            on_error=self.handle_error,
            on_status=self._handle_status_update,
            on_finished=self._on_processing_finished,
        )

        self._init_managers()
        self._init_engines()

        try:
            self.layout = WhisperLayoutManager(gui_ref=self)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Layout Fehler: {e}")
            self._show_error_and_exit(f"Layout Fehler: {e}")
            return

        try:
            self.layout.setup_gui()
            self._setup_callbacks()
            self.queue_manager = QueueManager(self)
            self.tts_manager = TTSManager(self.advanced_settings)

            # --- Änderung: VAD-Fallback Variable mit gespeichertem Wert initialisieren
            self.vad_fallback_enabled = tk.BooleanVar(value=self.advanced_settings.vad_fallback_enabled)
            # Status des Buttons später in StatusBar setzen, aber dort wird der Wert direkt aus der Variable gelesen

            self.status_bar = StatusBar(self.root, self)
            self.status_bar.frame.grid(row=4, column=0, sticky="ew", pady=(2, 0))

            # Button-Texte aktualisieren
            self._update_vad_fallback_button()
            self._update_live_mode_button()

            self.root.after(100, self._start_gui_updaters)

            if hasattr(self, "url_entry") and self.url_entry.winfo_exists():
                self.url_entry.delete(0, "end")
                self.url_entry.insert(0, self.settings.last_url)
            else:

                def set_initial_url() -> None:
                    if hasattr(self, "url_entry") and self.url_entry.winfo_exists():
                        self.url_entry.delete(0, "end")
                        self.url_entry.insert(0, self.settings.last_url)

                self.root.after(200, set_initial_url)

            self.root.deiconify()
            self.root.title("🐉 Dragon Whisperer")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ GUI Setup Fehler: {e}")
            self._show_error_and_exit(f"GUI konnte nicht erstellt werden: {e}")
            return

        self._register_signal_handlers()
        if self._show_cookie_notice:
            self.root.after(500, self._show_cookie_notice_dialog)
        self._bind_shortcuts()
        self.root.after(1000, self._start_system_monitoring)
        self.root.after(2000, self._final_initialization_check)
        self._schedule_gui_health_check()

    def _apply_precision_optimizations(self):
        self.advanced_settings.vad_filter = False
        self.advanced_settings.duplicate_similarity_threshold = 0.98
        self.advanced_settings.adaptive_chunk = False
        self.advanced_settings.chunk_duration = 20.0

        free_vram = None
        if TORCH_AVAILABLE:
            try:
                import torch
                if torch.cuda.is_available():
                    total = torch.cuda.get_device_properties(0).total_memory
                    free_vram = total / (1024**3)
                    free_vram = max(0, free_vram - 0.5)
            except Exception:
                pass

        if free_vram is not None and free_vram > 6:
            self.settings.default_model = "large-v3"
        else:
            self.settings.default_model = "medium"

        logger.info("🔧 Präzisionsoptimierungen angewendet: VAD=aus, Duplikatschwelle=0.98, adaptive_chunk=aus, chunk_dauer=20s, modell={}".format(self.settings.default_model))

    def _init_managers(self) -> None:
        self.stream_manager = StreamManager(
            enable_debug=(DEBUG_LEVEL >= 1),
            use_browser_cookies=self.settings.use_browser_cookies,
        )
        self.ffmpeg_manager = FFmpegManager(
            self.advanced_settings.config,
            self.stream_manager,
            self.advanced_settings,
        )
        self.export_manager = ExportManager()
        self.resource_manager = ResourceManager()
        self.memory_manager = MemoryManager()
        if IS_LINUX:
            self.performance_optimizer = LinuxPerformanceOptimizer(gui_ref=self)

    def _init_engines(self) -> None:
        if WHISPER_AVAILABLE:
            self.transcription_engine = TranscriptionEngine(
                self.advanced_settings, cache_manager=self.app_context.cache_manager
            )
        else:
            self.transcription_engine = DummyTranscriptionEngine(
                self.advanced_settings, cache_manager=self.app_context.cache_manager
            )

        self.language_detector = LanguageDetector(self.transcription_engine)

        self.translation_engine = self._create_translation_engine()

        self.audio_processor = AudioProcessor(
            controller_ref=self.controller,
            ffmpeg_manager=self.ffmpeg_manager,
            settings=self.advanced_settings,
        )

        fallback_engine = None
        if self.advanced_settings.translation_engine == "google" and OLLAMA_AVAILABLE:
            fallback_engine = OllamaTranslationEngine(
                target_lang=self.current_language,
                settings=self.advanced_settings,
                model=self.advanced_settings.ollama_model,
                host=self.advanced_settings.ollama_host,
                cache_manager=self.app_context.cache_manager,
            )
            logger.info("✅ Fallback-Engine (Ollama) bereitgestellt")

        self.audio_processor.set_engines(
            transcription_engine=self.transcription_engine,
            translation_engine=self.translation_engine,
            fallback_translation_engine=fallback_engine,
            plugin_manager=None,
        )

    def _create_translation_engine(self) -> BaseTranslationEngine:
        if self.advanced_settings.translation_engine == "ollama":
            if OLLAMA_AVAILABLE:
                logger.info("✅ OllamaTranslationEngine aktiviert")
                return OllamaTranslationEngine(
                    target_lang=self.current_language,
                    settings=self.advanced_settings,
                    model=self.advanced_settings.ollama_model,
                    host=self.advanced_settings.ollama_host,
                    cache_manager=self.app_context.cache_manager,
                )
            else:
                logger.warning(
                    "⚠️ Ollama not available, falling back to Google Translate"
                )
                return self._create_google_translation_engine()

        elif self.advanced_settings.translation_engine == "argos":
            if ARGOS_AVAILABLE:
                logger.info("✅ ArgosTranslateEngine aktiviert")
                return ArgosTranslateEngine(
                    self.current_language,
                    self.advanced_settings,
                    cache_manager=self.app_context.cache_manager,
                )
            else:
                logger.warning(
                    "⚠️ argos-translate missing, falling back to Google Translate"
                )
                return self._create_google_translation_engine()

        else:
            return self._create_google_translation_engine()

    def _create_google_translation_engine(self) -> BaseTranslationEngine:
        if TRANSLATOR_AVAILABLE:
            logger.info("✅ GoogleTranslationEngine aktiviert")
            return GoogleTranslationEngine(
                self.current_language,
                self.advanced_settings,
                cache_manager=self.app_context.cache_manager,
            )
        else:
            logger.warning("⚠️ Keine Übersetzungs-Engine verfügbar, verwende Dummy")
            return DummyTranslationEngine(
                self.current_language,
                self.advanced_settings,
                cache_manager=self.app_context.cache_manager,
            )

    # --- Änderung: Verbesserte temporäre Engine-Erstellung mit Logging
    def _create_temporary_translation_engine(self, engine_name: str, target_lang: str) -> Optional[BaseTranslationEngine]:
        logger.debug(f"Temporäre Engine angefragt: {engine_name} für Ziel {target_lang}")
        if engine_name == "google":
            if TRANSLATOR_AVAILABLE:
                logger.info("Erstelle temporäre GoogleTranslationEngine")
                return GoogleTranslationEngine(
                    target_lang=target_lang,
                    settings=self.advanced_settings,
                    cache_manager=self.app_context.cache_manager,
                )
            else:
                logger.warning("Google-Engine nicht verfügbar (deep-translator fehlt)")
        elif engine_name == "ollama":
            if OLLAMA_AVAILABLE:
                logger.info("Erstelle temporäre OllamaTranslationEngine")
                return OllamaTranslationEngine(
                    target_lang=target_lang,
                    settings=self.advanced_settings,
                    model=self.advanced_settings.ollama_model,
                    host=self.advanced_settings.ollama_host,
                    cache_manager=self.app_context.cache_manager,
                )
            else:
                logger.warning("Ollama-Engine nicht verfügbar (requests fehlt)")
        elif engine_name == "argos":
            if ARGOS_AVAILABLE:
                logger.info("Erstelle temporäre ArgosTranslateEngine")
                return ArgosTranslateEngine(
                    target_lang=target_lang,
                    settings=self.advanced_settings,
                    cache_manager=self.app_context.cache_manager,
                )
            else:
                logger.warning("Argos-Engine nicht verfügbar (argostranslate fehlt)")
        return None

    def _setup_callbacks(self) -> None:
        pass

    def _apply_theme(self, theme_name: str) -> None:
        if theme_name == "dark":
            self.current_theme = DarkTheme()
        elif theme_name == "light":
            self.current_theme = LightTheme()
        elif theme_name == "highcontrast":
            self.current_theme = HighContrastTheme()
        else:
            self.current_theme = DarkTheme()
        global CURRENT_THEME
        CURRENT_THEME = self.current_theme

        self.root.configure(bg=self.current_theme.BG_PRIMARY)

        self._update_ttk_styles()

        self._update_widget_tree(self.root)

        if hasattr(self, "layout") and hasattr(self.layout, "apply_theme"):
            self.layout.apply_theme()

        logger.info(f"🎨 Theme gewechselt zu: {theme_name}")

    def _update_ttk_styles(self) -> None:
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(
            "Dark.TCombobox",
            fieldbackground=self.current_theme.COMBO_BG,
            background=self.current_theme.COMBO_BG,
            foreground=self.current_theme.COMBO_FG,
            selectbackground=self.current_theme.COMBO_SELECTION,
            selectforeground=self.current_theme.TEXT_PRIMARY,
            insertcolor=self.current_theme.TEXT_PRIMARY,
            borderwidth=1,
            relief="flat",
            arrowsize=12,
            padding=5,
        )
        style.map(
            "Dark.TCombobox",
            fieldbackground=[
                ("readonly", self.current_theme.COMBO_BG),
                ("active", self.current_theme.BG_HOVER),
            ],
            background=[
                ("readonly", self.current_theme.COMBO_BG),
                ("active", self.current_theme.BG_HOVER),
            ],
            foreground=[
                ("readonly", self.current_theme.COMBO_FG),
                ("active", self.current_theme.TEXT_PRIMARY),
            ],
        )
        style.configure(
            "Dark.Horizontal.TProgressbar",
            background=self.current_theme.SUCCESS,
            troughcolor=self.current_theme.BG_TERTIARY,
            bordercolor=self.current_theme.BORDER,
        )

    def _update_widget_tree(self, parent: tk.Widget) -> None:
        widget_updates = {
            tk.Label: {"bg": "BG_PRIMARY", "fg": "TEXT_PRIMARY"},
            tk.Button: {
                "bg": "BG_TERTIARY",
                "fg": "TEXT_PRIMARY",
                "activebackground": "BG_HOVER",
                "activeforeground": "TEXT_ACCENT",
            },
            tk.Entry: {
                "bg": "BG_TERTIARY",
                "fg": "TEXT_PRIMARY",
                "insertbackground": "TEXT_PRIMARY",
                "selectbackground": "COMBO_SELECTION",
                "selectforeground": "TEXT_PRIMARY",
            },
            tk.Frame: {"bg": "BG_PRIMARY"},
            tk.LabelFrame: {"bg": "BG_SECONDARY", "fg": "TEXT_PRIMARY"},
            tk.Text: {
                "bg": "BG_TERTIARY",
                "fg": "TEXT_PRIMARY",
                "insertbackground": "TEXT_PRIMARY",
                "selectbackground": "COMBO_SELECTION",
                "selectforeground": "TEXT_PRIMARY",
            },
            tk.Checkbutton: {
                "bg": "BG_SECONDARY",
                "fg": "TEXT_PRIMARY",
                "selectcolor": "BG_TERTIARY",
                "activebackground": "BG_SECONDARY",
                "activeforeground": "TEXT_ACCENT",
            },
            tk.Radiobutton: {
                "bg": "BG_SECONDARY",
                "fg": "TEXT_PRIMARY",
                "selectcolor": "BG_TERTIARY",
                "activebackground": "BG_SECONDARY",
                "activeforeground": "TEXT_ACCENT",
            },
            tk.Listbox: {
                "bg": "BG_TERTIARY",
                "fg": "TEXT_PRIMARY",
                "selectbackground": "COMBO_SELECTION",
                "selectforeground": "TEXT_PRIMARY",
            },
            tk.Scrollbar: {
                "bg": "SCROLLBAR",
                "activebackground": "SCROLLBAR_HOVER",
                "troughcolor": "BG_TERTIARY",
            },
            tk.PanedWindow: {
                "bg": "BG_SECONDARY",
                "sashrelief": "raised",
                "sashwidth": 4,
            },
            ttk.Combobox: {
            },
        }

        special_updates = {
            "start_button": {"bg": "SUCCESS"},
            "stop_button": {"bg": "ERROR"},
            "translate_btn": {
                "bg": "SUCCESS" if self.translate_active else "BG_TERTIARY"
            },
            "subtitle_btn": {
                "bg": "SUBTITLE_ACTIVE" if self.subtitle_mode else "SUBTITLE_INACTIVE"
            },
            "status_label": {"bg": "BG_PRIMARY", "fg": "TEXT_SECONDARY"},
            "system_info_label": {"bg": "BG_SECONDARY", "fg": "TEXT_SECONDARY"},
            "progress_label": {"bg": "BG_SECONDARY", "fg": "TEXT_SECONDARY"},
            "stream_info_frame": {"bg": "BG_SECONDARY"},
            "stream_title_label": {"bg": "BG_SECONDARY", "fg": "TEXT_ACCENT"},
            "stream_details_label": {"bg": "BG_SECONDARY", "fg": "TEXT_SECONDARY"},
            "status_bar_frame": {"bg": "BG_SECONDARY"},
            "tts_btn": {"bg": "BG_TERTIARY", "fg": "TEXT_PRIMARY"},
            "vad_fallback_btn": {"bg": "BG_TERTIARY", "fg": "TEXT_PRIMARY"},
            "live_mode_btn": {"bg": "BG_TERTIARY", "fg": "TEXT_PRIMARY"},
        }

        widget_class = parent.__class__
        if widget_class in widget_updates:
            updates = {}
            for opt, theme_attr in widget_updates[widget_class].items():
                if hasattr(self.current_theme, theme_attr):
                    updates[opt] = getattr(self.current_theme, theme_attr)
            try:
                parent.configure(**updates)
            except tk.TclError:
                pass

        for name, updates in special_updates.items():
            if hasattr(self, name) and getattr(self, name) is parent:
                conf = {}
                for opt, theme_attr in updates.items():
                    if hasattr(self.current_theme, theme_attr):
                        conf[opt] = getattr(self.current_theme, theme_attr)
                try:
                    parent.configure(**conf)
                except tk.TclError:
                    pass

        try:
            for child in parent.winfo_children():
                self._update_widget_tree(child)
        except tk.TclError:
            pass

    def _show_cookie_notice_dialog(self) -> None:
        result = DarkMessageBox.askyesno(
            "Datenschutzhinweis",
            "Dragon Whisperer kann auf gespeicherte Browser-Cookies zugreifen, "
            "um YouTube-Streams zuverlässiger abzurufen. Dies kann Ihre Privatsphäre "
            "beeinträchtigen.\n\nMöchten Sie die Nutzung von Browser-Cookies erlauben?\n\n"
            "(Sie können diese Einstellung später in den erweiterten Einstellungen ändern.)",
            parent=self.root,
        )
        self.settings.use_browser_cookies = result
        self.settings.cookies_notice_shown = True
        self.settings.save_to_file()
        if hasattr(self, "stream_manager"):
            self.stream_manager.use_browser_cookies = result
        if hasattr(self, "stream_info_extractor"):
            self.stream_info_extractor.use_browser_cookies = result

    def _schedule_gui_health_check(self) -> None:
        if (
            not self._shutting_down
            and hasattr(self, "root")
            and self.root.winfo_exists()
        ):
            self.root.after(30000, self._perform_gui_health_check)

    def _perform_gui_health_check(self) -> None:
        try:
            checks: List[str] = []
            start_time = time.time()

            try:
                self.root.update_idletasks()
                responsiveness = time.time() - start_time
                if responsiveness > 0.5:
                    msg = f"⚠️ GUI responsiveness slow: {responsiveness:.1f}s"
                    checks.append(msg)
                    if logger.isEnabledFor(logging.DEBUG):
                        log_debug("gui", msg)
            except tk.TclError:
                return

            if hasattr(self, "memory_manager"):
                try:
                    mem_stats = self.memory_manager.get_memory_stats()
                    process_usage = mem_stats.get("process_usage_percent", 0)
                    if process_usage > 80:
                        msg = f"⚠️ High process memory usage: {process_usage:.1f}%"
                        checks.append(msg)
                        if process_usage > 85:
                            self.memory_manager.aggressive_cleanup()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    if logger.isEnabledFor(logging.DEBUG):
                        log_debug("memory", f"Health check error: {e}")

            if hasattr(self, "gui_queue") and self.gui_queue.qsize() > 50:
                old_size = self.gui_queue.qsize()
                self.queue_manager._cleanup_queue(self.gui_queue, 30)
                new_size = self.gui_queue.qsize()
                checks.append(f"🧹 GUI queue cleaned: {old_size} → {new_size}")
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("queue", f"GUI queue cleaned: {old_size} → {new_size}")

            if (
                hasattr(self, "_text_update_queue")
                and self._text_update_queue.qsize() > 150
            ):
                old_size = self._text_update_queue.qsize()
                self.queue_manager._cleanup_queue(self._text_update_queue, 75)
                new_size = self._text_update_queue.qsize()
                checks.append(f"🧹 Text queue cleaned: {old_size} → {new_size}")
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("queue", f"Text queue cleaned: {old_size} → {new_size}")

            active_threads = threading.enumerate()
            if len(active_threads) > 15:
                checks.append(f"⚠️ Many active threads: {len(active_threads)}")
                if logger.isEnabledFor(logging.DEBUG):
                    thread_names = [t.name for t in active_threads]
                    log_debug("threads", f"Active threads: {thread_names}")

            try:
                cache_stats = self.app_context.cache_manager.get_stats()
                for cache_name, stats in cache_stats.items():
                    size = stats.get("size", 0)
                    maxsize = stats.get("maxsize", 100)
                    if size > maxsize * 0.9:
                        checks.append(f"⚠️ {cache_name} nearly full ({size}/{maxsize})")
                        if logger.isEnabledFor(logging.DEBUG):
                            log_debug(
                                "cache", f"{cache_name} nearly full: {size}/{maxsize}"
                            )
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                if logger.isEnabledFor(logging.DEBUG):
                    log_debug("cache", f"Stats error: {e}")

            for attr in ["transcript_text", "translation_text"]:
                widget = getattr(self, attr, None)
                if widget and widget.winfo_exists():
                    try:
                        lines = int(widget.index("end-1c").split(".")[0])
                        max_lines = (self.advanced_settings.transcript_max_lines if attr == "transcript_text" else self.advanced_settings.translation_max_lines)
                        if lines > max_lines:
                            keep_lines = max_lines - 100
                            delete_to = f"{lines - keep_lines}.0"
                            widget.delete("1.0", delete_to)
                            checks.append(f"🧹 {attr} bereinigt: {lines} → {keep_lines} Zeilen")
                    except Exception as e:
                        logger.debug(f"Fehler bei Zeilenbereinigung {attr}: {e}")

            if any("memory usage" in c for c in checks) and hasattr(
                self, "memory_manager"
            ):
                logger.info("🧹 Triggering aggressive cleanup due to high memory")
                self.memory_manager.aggressive_cleanup()

            if checks:
                log_checks = checks[:3]
                if len(checks) > 3:
                    log_checks.append(f"... und {len(checks)-3} weitere")
                logger.info(f"🔍 GUI Health Check: {', '.join(log_checks)}")

            if hasattr(self, "audio_processor") and hasattr(
                self.audio_processor, "_chunk_counter"
            ):
                if self.audio_processor._chunk_counter % 50 == 0:
                    gc.collect()
            else:
                if not hasattr(self, "_health_check_counter"):
                    self._health_check_counter = 0
                self._health_check_counter += 1
                if self._health_check_counter % 10 == 0:
                    gc.collect()

        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Health check error: {e}")
        finally:
            if (
                not self._shutting_down
                and hasattr(self, "root")
                and self.root.winfo_exists()
            ):
                self.root.after(30000, self._perform_gui_health_check)

    def _safe_audio_processor_dispose(self):
        if hasattr(self, 'audio_processor') and self.audio_processor is not None:
            self.audio_processor.dispose()

    def _safe_ffmpeg_manager_dispose(self):
        if hasattr(self, 'ffmpeg_manager') and self.ffmpeg_manager is not None:
            self.ffmpeg_manager.dispose()

    def _safe_transcription_engine_dispose(self):
        if hasattr(self, 'transcription_engine') and self.transcription_engine is not None:
            self.transcription_engine.dispose()

    def _safe_translation_engine_dispose(self):
        if hasattr(self, 'translation_engine') and self.translation_engine is not None:
            self.translation_engine.dispose()

    def _safe_memory_manager_dispose(self):
        if hasattr(self, 'memory_manager') and self.memory_manager is not None:
            self.memory_manager.dispose()

    def _safe_resource_manager_cleanup(self):
        if hasattr(self, 'resource_manager') and self.resource_manager is not None:
            self.resource_manager.cleanup()

    def _safe_stream_manager_dispose(self):
        if hasattr(self, 'stream_manager') and self.stream_manager is not None:
            self.stream_manager.dispose()

    def _register_signal_handlers(self) -> None:
        try:
            logger.info("🔧 Registering cleanup handlers with SignalHandler...")
            SignalHandler.register_cleanup(
                self._safe_stop_all_processes,
                name="StopAllProcesses",
                priority=ShutdownPriority.CRITICAL,
                timeout=3.0,
                essential=True,
            )
            SignalHandler.register_cleanup(
                self._safe_audio_processor_dispose,
                name="AudioProcessorDispose",
                priority=ShutdownPriority.HIGH,
                timeout=2.0,
            )
            if hasattr(self, "ffmpeg_manager") and self.ffmpeg_manager:
                SignalHandler.register_cleanup(
                    self._safe_ffmpeg_manager_dispose,
                    name="FFmpegManagerDispose",
                    priority=ShutdownPriority.HIGH,
                    timeout=2.0,
                )
                logger.info("   ✅ Registered FFmpegManager cleanup")
            SignalHandler.register_cleanup(
                self._safe_transcription_engine_dispose,
                name="TranscriptionEngineDispose",
                priority=ShutdownPriority.MEDIUM,
                timeout=1.0,
            )
            SignalHandler.register_cleanup(
                self._safe_translation_engine_dispose,
                name="TranslationEngineDispose",
                priority=ShutdownPriority.MEDIUM,
                timeout=1.0,
            )
            SignalHandler.register_cleanup(
                self._safe_memory_manager_dispose,
                name="MemoryManagerDispose",
                priority=ShutdownPriority.LOW,
            )
            SignalHandler.register_cleanup(
                self._safe_resource_manager_cleanup,
                name="ResourceManagerCleanup",
                priority=ShutdownPriority.LOW,
            )
            SignalHandler.register_cleanup(
                self._safe_stream_manager_dispose,
                name="StreamManagerDispose",
                priority=ShutdownPriority.LOW,
            )
            if IS_LINUX and hasattr(self, "performance_optimizer"):
                SignalHandler.register_cleanup(
                    self._safe_linux_optimizer_cleanup,
                    name="LinuxOptimizerCleanup",
                    priority=ShutdownPriority.LOW,
                    timeout=1.0,
                )
            SignalHandler.register_cleanup(
                lambda: (
                    self.app_context.cache_manager.transcription_cache.clear(),
                    self.app_context.cache_manager.translation_cache.clear(),
                    self.app_context.cache_manager.audio_cache.clear(),
                ),
                name="ClearGlobalCaches",
                priority=ShutdownPriority.LOW,
            )
            SignalHandler.register_cleanup(
                self._cleanup_queues,
                name="CleanupQueues",
                priority=ShutdownPriority.LOW,
            )
            SignalHandler.register_cleanup(
                lambda: self.translation_executor.shutdown(wait=False),
                name="TranslationExecutorShutdown",
                priority=ShutdownPriority.LOW,
            )
            SignalHandler.register_cleanup(
                lambda: _EXECUTOR.shutdown(wait=False),
                name="GlobalExecutorShutdown",
                priority=ShutdownPriority.LOW,
            )
            try:
                import torch

                if torch.cuda.is_available():
                    SignalHandler.register_cleanup(
                        torch.cuda.empty_cache,
                        name="GPUMemoryCleanup",
                        priority=ShutdownPriority.LOW,
                        timeout=1.0,
                    )
                    logger.info("   ✅ Registered GPU cleanup")
            except ImportError:
                pass
            count = sum(len(ops) for ops in SignalHandler._cleanup_operations.values())
            logger.info(f"✅ Registered {count} cleanup handlers")
            SignalHandler.setup(verbose=False, silent=True, atexit_enabled=False)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ SignalHandler registration error: {e}")

    def _safe_stop_all_processes(self) -> None:
        logger.info("🛑 Safely stopping all processes...")
        self._shutting_down = True
        self.is_processing = False
        if hasattr(self, "controller"):
            try:
                self.controller._shutdown_event.set()
                self.controller.stop_processing(wait=False)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Controller stop error: {e}")
        if hasattr(self, "audio_processor"):
            try:
                self.audio_processor._processing.clear()
                if hasattr(self.audio_processor, "_stop_event"):
                    self.audio_processor._stop_event.set()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Audio processor stop error: {e}")
        if hasattr(self, "ffmpeg_manager"):
            try:
                self.ffmpeg_manager._shutting_down = True
                if hasattr(self.ffmpeg_manager, "_processes"):
                    import os
                    import signal

                    for process_id, process_info in list(
                        self.ffmpeg_manager._processes.items()
                    ):
                        try:
                            process = process_info.get("process")
                            if process and hasattr(process, "pid"):
                                try:
                                    if IS_WINDOWS:
                                        process.terminate()
                                    else:
                                        os.kill(process.pid, signal.SIGTERM)
                                    time.sleep(0.1)
                                except Exception:
                                    try:
                                        if IS_WINDOWS:
                                            process.kill()
                                        else:
                                            os.kill(process.pid, signal.SIGKILL)
                                    except Exception:
                                        pass
                        except Exception:
                            pass
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ FFmpeg stop error: {e}")
        try:
            self.translation_executor.shutdown(wait=False)
        except Exception:
            pass
        logger.info("✅ All processes stopped")

    def _safe_linux_optimizer_cleanup(self) -> None:
        if not IS_LINUX or not hasattr(self, "performance_optimizer"):
            return
        logger.info("🐧 Safe Linux optimizer cleanup...")
        try:
            gui_exists = False
            try:
                if hasattr(self, "root") and self.root.winfo_exists():
                    gui_exists = True
            except Exception:
                gui_exists = False
            if gui_exists:
                try:
                    self.performance_optimizer.restore_normal_mode()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ restore_normal_mode failed: {e}")
                    try:
                        self.performance_optimizer.dispose()
                    except Exception:
                        pass
            else:
                try:
                    self.performance_optimizer.dispose()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"⚠️ dispose failed: {e}")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Linux optimizer cleanup error: {e}")

    def _cleanup_queues(self) -> None:
        logger.info("🗑️ Cleaning up queues...")

        def drain_queue(q: queue.Queue, name: str, max_items: int = 100) -> int:
            removed = 0
            try:
                while removed < max_items:
                    try:
                        q.get_nowait()
                        removed += 1
                    except queue.Empty:
                        break
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Fehler beim Leeren der {name}-Queue: {e}")
            return removed

        if hasattr(self, "gui_queue") and self.gui_queue is not None:
            cleared = drain_queue(self.gui_queue, "GUI")
            if cleared > 0:
                logger.info(f"  Cleared GUI queue: {cleared} items")
        if hasattr(self, "_text_update_queue") and self._text_update_queue is not None:
            cleared = drain_queue(self._text_update_queue, "Text")
            if cleared > 0:
                logger.info(f"  Cleared text queue: {cleared} items")

    def _safe_exit_dialog(self) -> None:
        if self._shutting_down or self._exit_dialog_active:
            return
        self._exit_dialog_active = True
        try:
            if not hasattr(self, "root") or not self.root.winfo_exists():
                self._direct_shutdown()
                return
            result = DarkMessageBox.askyesno(
                "🐉 Dragon Whisperer - Beenden",
                "Programm wirklich beenden?\n\n"
                "● Laufende Transkriptionen werden gestoppt\n"
                "● Nicht gespeicherte Daten gehen verloren\n\n"
                "Sicher beenden?",
                parent=self.root,
            )
            if result:
                logger.info("✅ User confirmed exit - shutting down...")
                self._direct_shutdown()
            else:
                logger.info("↩️ Exit cancelled by user")
                if hasattr(self, "exit_button") and self.exit_button.winfo_exists():
                    self.exit_button.config(state="normal", text=" ⏻ EXIT ")
                self._exit_dialog_active = False
        except tk.TclError:
            logger.warning("⚠️ GUI destroyed, performing direct shutdown...")
            self._direct_shutdown()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Exit dialog error: {e}")
            self._direct_shutdown()
        finally:
            if not self._shutting_down:
                self._exit_dialog_active = False

    def _direct_shutdown(self) -> None:
        if self._shutting_down:
            logger.warning("⚠️ Shutdown already in progress, skipping...")
            return
        logger.info("🔧 Performing confirmed shutdown...")
        self._shutting_down = True
        self._safe_stop_all_processes()
        for dlg in self._open_dialogs[:]:
            try:
                if dlg.winfo_exists():
                    dlg.destroy()
            except Exception:
                pass
        self._open_dialogs.clear()
        time.sleep(0.3)
        try:
            if hasattr(self, "root") and self.root.winfo_exists():
                self.root.title("🐉 Dragon Whisperer - Beendet...")
                self.root.update_idletasks()
                self.root.quit()
                self.root.destroy()
        except tk.TclError:
            pass
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ GUI shutdown error: {e}")

    def _start_gui_updaters(self) -> None:
        if hasattr(self, "root") and self.root.winfo_exists():
            self.queue_manager.start()
            self.root.after(5000, self._check_queue_sizes)

    def _check_queue_sizes(self) -> None:
        if self.queue_manager:
            self.queue_manager._check_queue_sizes()

    @gui_operation_decorator
    def select_file_dark(self) -> None:
        try:
            filename = filedialog.askopenfilename(
                title="🎬 Select Audio/Video File - Dragon Whisperer",
                filetypes=[
                    ("Media files", "*.mp3 *.wav *.m4a *.mp4 *.avi *.mkv *.mov *.flac"),
                    ("All files", "*.*"),
                ],
            )
            if filename:
                file_url = f"file://{filename}"
                self.url_entry.delete(0, "end")
                self.url_entry.insert(0, file_url)
                self.update_status(f"📁 File selected: {os.path.basename(filename)}")

                def async_language_detection() -> None:
                    try:
                        self.analyze_video_language(filename)
                    except Exception:
                        pass

                detection_thread = threading.Thread(
                    target=async_language_detection, daemon=True
                )
                if hasattr(self, "resource_manager"):
                    self.resource_manager.register_thread(detection_thread)
                detection_thread.start()
                info = self.stream_info_extractor.extract_stream_info(file_url)
                self.update_stream_info(info)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.update_status(f"❌ File selection failed: {e}")

    @gui_operation_decorator
    def paste_url(self) -> None:
        try:
            clipboard = self.root.clipboard_get().strip()
            if clipboard:
                try:
                    cleaned_url = self.clean_and_validate_url(clipboard)
                except ValueError as e:
                    self.update_status(f"❌ Ungültige URL: {e}")
                    return
                self.url_entry.delete(0, "end")
                self.url_entry.insert(0, cleaned_url)
                self.update_status("📋 URL eingefügt")
                if cleaned_url.startswith("file://"):
                    file_path = cleaned_url[7:]
                    if os.path.exists(file_path):

                        def async_detection() -> None:
                            try:
                                self.analyze_video_language(file_path)
                            except Exception:
                                pass

                        detection_thread = threading.Thread(
                            target=async_detection, daemon=True
                        )
                        if hasattr(self, "resource_manager"):
                            self.resource_manager.register_thread(detection_thread)
                        detection_thread.start()
            else:
                self.update_status("❌ Zwischenablage ist leer")
        except tk.TclError:
            self.update_status("❌ Konnte nicht auf Zwischenablage zugreifen")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.update_status(f"❌ Fehler beim Einfügen: {str(e)[:50]}")

    def clean_and_validate_url(self, url: str) -> str:
        if not url:
            raise ValueError("URL cannot be empty")

        url = url.strip()
        if not url:
            raise ValueError("URL is empty after stripping")

        if url.startswith(("dvb://", "dvb-s://", "file://")):
            return url

        if not url.startswith(("http://", "https://")):
            url = "https://" + url

        if len(url) < 10:
            raise ValueError("URL too short (minimum 10 characters)")

        if " " in url:
            raise ValueError("URL cannot contain spaces")

        try:
            parsed = urllib.parse.urlparse(url)
            if not parsed.netloc:
                raise ValueError("Invalid URL format (no network location)")
        except Exception as e:
            raise ValueError(f"Invalid URL: {e}")

        return url

    def analyze_video_language(self, file_path: str) -> None:
        if hasattr(self, "language_info_label"):
            self.root.after(
                0, lambda: self.language_info_label.config(text="🔍 Analyzing...")
            )

        def language_detection_worker() -> None:
            try:
                detection_result = self.language_detector.detect_video_language(
                    file_path
                )
                logger.debug(f"Language detection result: {detection_result}")

                def update_result() -> None:
                    if hasattr(self, "language_info_label"):
                        if "error" in detection_result:
                            self.language_info_label.config(
                                text=f"❌ {detection_result['error']}"
                            )
                        elif "info" in detection_result:
                            self.language_info_label.config(
                                text=f"ℹ️ {detection_result['info']}"
                            )
                        else:
                            language_name = detection_result["language_name"]
                            confidence = detection_result["confidence"]
                            self.current_video_language = detection_result[
                                "detected_language"
                            ]
                            language_icons = {
                                "zh": "㊗️",
                                "ja": "🗾",
                                "ko": "₩",
                                "th": "🇹🇭",
                                "vi": "🇻🇳",
                            }
                            icon = language_icons.get(self.current_video_language, "✅")
                            display_text = f"{icon} {language_name} ({confidence:.0%})"
                            self.language_info_label.config(text=display_text)

                if hasattr(self, "root"):
                    self.root.after(0, update_result)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.error(f"Language detection exception: {e}")

                def update_error() -> None:
                    if hasattr(self, "language_info_label"):
                        self.language_info_label.config(text="❌ Analysis failed")

                if hasattr(self, "root"):
                    self.root.after(0, update_error)

        detection_thread = threading.Thread(
            target=language_detection_worker, daemon=True
        )
        if hasattr(self, "resource_manager"):
            self.resource_manager.register_thread(detection_thread)
        detection_thread.start()

    def on_url_change(self, event: Optional[tk.Event] = None) -> None:
        if not hasattr(self, "url_entry"):
            return
        url = self.url_entry.get().strip()
        if url.startswith("file://"):
            file_path = url[7:]
            if os.path.exists(file_path):

                def async_detection() -> None:
                    try:
                        self.analyze_video_language(file_path)
                    except Exception:
                        pass

                detection_thread = threading.Thread(target=async_detection, daemon=True)
                if hasattr(self, "resource_manager"):
                    self.resource_manager.register_thread(detection_thread)
                detection_thread.start()
            else:
                if hasattr(self, "language_info_label"):
                    self.language_info_label.config(text="❌ File not found")
        else:
            if hasattr(self, "language_info_label"):
                self.language_info_label.config(text="")
            self.current_video_language = None

    def on_language_change(self, event: Optional[tk.Event] = None) -> None:
        try:
            selected_name = self.lang_var.get()
            lang_code: Optional[str] = None
            for name, code in SORTED_LANGUAGES:
                if name == selected_name:
                    lang_code = code
                    break
            if lang_code and lang_code != self.current_language:
                self.current_language = lang_code
                if hasattr(self, "translation_engine"):
                    self.translation_engine.set_target_language(lang_code)
                lang_display = LANGUAGE_SHORT_CODES.get(lang_code, lang_code)
                if hasattr(self, "translation_header"):
                    self.translation_header.config(
                        text=f"🌐 Translation ({lang_display})"
                    )
                self.update_status(f"🌍 Target language: {selected_name}")
        except Exception:
            pass

    def _on_language_selected(self, event):
        selected = self.lang_var.get()
        if selected.startswith("---"):
            self.lang_var.set(self._last_valid_language)
        else:
            self._last_valid_language = selected
            self.on_language_change(event)

    def on_model_change(self, event: Optional[tk.Event] = None) -> None:
        if not hasattr(self, "model_var"):
            return
        new_model = self.model_var.get()
        if new_model not in WHISPER_MODELS:
            logger.warning(f"⚠️ Invalid model selected: {new_model}")
            current = self.transcription_engine.get_current_model()
            self.model_var.set(current)
            return
        if not hasattr(self, "transcription_engine"):
            return
        current_model = self.transcription_engine.get_current_model()
        if new_model == current_model:
            return
        if self.transcription_engine.is_model_loading():
            self.update_status("🔄 Model already loading...")
            return
        success = self.transcription_engine.reload_model(new_model)
        if success:
            self.update_status(f"🔄 Switching to {new_model}...")
            self._check_model_loading_complete(new_model)
        else:
            self.update_status("❌ Model switch failed")
            self.model_var.set(current_model)

    def _check_model_loading_complete(self, target_model: str) -> None:
        if self.transcription_engine.is_model_loading():
            self.root.after(
                200, lambda: self._check_model_loading_complete(target_model)
            )
        else:
            current = self.transcription_engine.get_current_model()
            if current == target_model:
                self.update_status(f"✅ Model switched to {target_model}")
            else:
                self.update_status("❌ Model switch failed")
                self.model_var.set(current)

    def toggle_translation(self) -> None:
        self.translate_active = not self.translate_active
        if hasattr(self, "audio_processor"):
            if self.translate_active:
                self.audio_processor._translation_enabled.set()
            else:
                self.audio_processor._translation_enabled.clear()
        if hasattr(self, "translate_btn"):
            if self.translate_active:
                self.translate_btn.config(text="🌐 ON", bg=self.current_theme.SUCCESS)
                self.update_status("✅ Translation active")
            else:
                self.translate_btn.config(
                    text="🌐 OFF", bg=self.current_theme.BG_TERTIARY
                )
                self.update_status("❌ Translation inactive")

    def toggle_subtitle_mode(self) -> None:
        self.subtitle_mode = not self.subtitle_mode
        if hasattr(self, "audio_processor"):
            self.audio_processor.enable_subtitle_mode(self.subtitle_mode)
        if hasattr(self, "subtitle_btn"):
            if self.subtitle_mode:
                self.subtitle_btn.config(
                    bg=self.current_theme.SUBTITLE_ACTIVE,
                    fg=self.current_theme.TEXT_PRIMARY,
                )
                self.update_status("🎬 SUBTITLE MODE: Timestamps activated")
            else:
                self.subtitle_btn.config(
                    bg=self.current_theme.SUBTITLE_INACTIVE,
                    fg=self.current_theme.TEXT_PRIMARY,
                )
                self.update_status("📝 NORMAL MODE: Continuous text")

    def _on_start_click(self):
        if self.is_processing:
            self.update_status("⚠️ Bereits aktiv")
            return
        try:
            self.start_button.config(state="disabled")
        except Exception:
            pass
        self.controller.start_processing()

    def toggle_layout(self) -> None:
        try:
            logger.info(f"🔄 Starting layout toggle from {self.layout_mode}")
            old_transcript = ""
            old_translation = ""
            try:
                if hasattr(self, "transcript_text") and self.transcript_text:
                    old_transcript = self.transcript_text.get("1.0", "end-1c")
                    logger.info(f"  📝 Saved transcript: {len(old_transcript)} chars")
            except (tk.TclError, AttributeError) as e:
                logger.warning(f"  ⚠️ Could not save transcript: {e}")
            try:
                if hasattr(self, "translation_text") and self.translation_text:
                    old_translation = self.translation_text.get("1.0", "end-1c")
                    logger.info(f"  📝 Saved translation: {len(old_translation)} chars")
            except (tk.TclError, AttributeError) as e:
                logger.warning(f"  ⚠️ Could not save translation: {e}")
            if self.layout_mode == "vertical":
                self.layout_mode = "horizontal"
                new_mode_text = "Horizontal"
            else:
                self.layout_mode = "vertical"
                new_mode_text = "Vertical"
            logger.info(f"  🔄 Switching to: {self.layout_mode}")
            if hasattr(self, "settings"):
                self.settings.layout_mode = self.layout_mode
                try:
                    self.settings.save_to_file()
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(f"  ⚠️ Settings save error: {e}")
            self.update_status(f"🔄 Switching to {new_mode_text} layout...")
            if hasattr(self, "layout"):
                new_transcript, new_translation = self.layout.create_text_areas()
                if new_transcript and old_transcript:
                    try:
                        new_transcript.insert("1.0", old_transcript)
                        logger.info("  ✅ Restored transcript to new widget")
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.warning(f"  ❌ Failed to restore transcript: {e}")
                if new_translation and old_translation:
                    try:
                        new_translation.insert("1.0", old_translation)
                        logger.info("  ✅ Restored translation to new widget")
                    except Exception as e:
                        if isinstance(e, (KeyboardInterrupt, SystemExit)):
                            raise
                        logger.warning(f"  ❌ Failed to restore translation: {e}")
            self.update_status(f"✅ {new_mode_text} layout active")
            logger.info("✅ Layout toggle completed successfully")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ CRITICAL Layout toggle error: {e}")
            self.update_status("❌ Layout change failed")
            try:
                self.layout_mode = "vertical"
                if hasattr(self, "layout"):
                    self.layout.create_text_areas()
            except Exception:
                pass

    def clear_all(self) -> None:
        try:
            if hasattr(self, "transcript_text"):
                self.transcript_text.delete("1.0", "end")
            if hasattr(self, "translation_text"):
                self.translation_text.delete("1.0", "end")
        except Exception:
            pass
        with self._history_lock:
            self.transcript_history.clear()
            self.translation_history.clear()
        if hasattr(self, "memory_manager"):
            self.memory_manager.clear_component("transcript")
            self.memory_manager.clear_component("translation")
        self._last_transcription_text = ""
        self._last_translation_text = ""
        self._translation_reset_counter = 0
        self.update_status("🗑️ Cleared & optimizations reset")

    @gui_operation_decorator
    def save_transcript(self) -> None:
        try:
            if not self.transcript_history:
                DarkMessageBox.showinfo(
                    "WARNING", "No transcriptions available to save.", self.root
                )
                return

            if self.current_stream_info and self.current_stream_info.title:
                base_name = re.sub(r"[^\w\-_\. ]", "", self.current_stream_info.title)
                base_name = base_name.strip().replace(" ", "_")[:50]
            else:
                base_name = "transcript"
            suggested = f"{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

            filename = filedialog.asksaveasfilename(
                defaultextension=".txt",
                initialfile=suggested,
                filetypes=[
                    ("Text files", "*.txt"),
                    ("SRT subtitles", "*.srt"),
                    ("WebVTT", "*.vtt"),
                    ("JSON", "*.json"),
                    ("Word document", "*.docx"),
                    ("All files", "*.*"),
                ],
            )
            if not filename:
                return

            file_ext = Path(filename).suffix.lower()
            success = False

            if file_ext == ".srt":
                success = self.export_manager.export_subtitles(
                    list(self.transcript_history), None, "srt", filename
                )
            elif file_ext == ".vtt":
                success = self.export_manager.export_subtitles(
                    list(self.transcript_history), None, "vtt", filename
                )
            elif file_ext == ".json":
                success = self.export_manager.export_json(
                    list(self.transcript_history),
                    list(self.translation_history),
                    filename,
                )
            elif file_ext == ".docx":
                success = self.export_manager.export_docx(
                    list(self.transcript_history), filename
                )
            else:
                with open(filename, "w", encoding="utf-8") as f:
                    if self.current_stream_info:
                        f.write("=== STREAM INFORMATION ===\n")
                        f.write(f"Title: {self.current_stream_info.title}\n")
                        f.write(f"Uploader: {self.current_stream_info.uploader}\n")
                        f.write(f"Duration: {self.current_stream_info.duration}\n")
                        f.write(f"Platform: {self.current_stream_info.platform}\n")
                        f.write(
                            f"Saved at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
                        )
                    f.write("=== TRANSCRIPT ===\n")
                    if hasattr(self, "transcript_text"):
                        f.write(self.transcript_text.get("1.0", "end-1c"))
                    f.write("\n\n=== TRANSLATION ===\n")
                    if hasattr(self, "translation_text"):
                        f.write(self.translation_text.get("1.0", "end-1c"))
                success = True

            if success:
                self.update_status(f"💾 Saved: {os.path.basename(filename)}")
            else:
                self.update_status("❌ Export failed")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.update_status(f"❌ Save failed: {e}")

    def export_subtitles(self) -> None:
        if (
            not hasattr(self, "audio_processor")
            or not self.audio_processor._timed_transcriptions
        ):
            DarkMessageBox.showinfo(
                "WARNING",
                "No subtitle data available.\n\n"
                "Tip: First activate '🎬 Subtitle mode' "
                "and start a transcription.",
                self.root,
            )
            return
        try:
            filename = filedialog.asksaveasfilename(
                defaultextension=".srt",
                filetypes=[
                    ("SRT subtitles", "*.srt"),
                    ("VTT subtitles", "*.vtt"),
                    ("All files", "*.*"),
                ],
                title="Export subtitles",
            )
            if not filename:
                return
            file_ext = Path(filename).suffix.lower()
            format_type = "srt" if file_ext == ".srt" else "vtt"
            with self.audio_processor._subtitle_lock:
                timed_trans = list(self.audio_processor._timed_transcriptions)
                timed_transl = list(self.audio_processor._timed_translations)
            success = self.export_manager.export_subtitles(
                timed_trans,
                timed_transl,
                format=format_type,
                filename=filename,
            )
            if success:
                segment_count = len(timed_trans)
                translation_count = len(timed_transl)
                self.update_status(
                    f"📝 {format_type.upper()} exported: {os.path.basename(filename)}"
                )
                DarkMessageBox.showinfo(
                    "Success",
                    f"Subtitles successfully exported!\n\n"
                    f"• File: {os.path.basename(filename)}\n"
                    f"• Segments: {segment_count}\n"
                    f"• Translations: {translation_count}\n"
                    f"• Format: {format_type.upper()}\n\n"
                    f"Can be directly imported into video editors.",
                    self.root,
                )
            else:
                self.update_status("❌ Subtitle export failed")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.update_status(f"❌ Subtitle export failed: {e}")
            DarkMessageBox.showerror("Error", f"Export failed:\n{str(e)}", self.root)

    def show_simple_stats(self) -> None:
        try:
            stats = self.performance_monitor.get_basic_stats()
            try:
                import psutil

                cpu = psutil.cpu_percent()
                memory = psutil.virtual_memory()
                memory_used = memory.used // (1024**2)
                memory_percent = memory.percent
                health_status = (
                    "Healthy" if cpu < 90 and memory_percent < 85 else "Degraded"
                )
            except Exception:
                cpu = 0.0
                memory_used = 0
                memory_percent = 0
                health_status = "Unknown"
            stats_text = f"""📊 STATISTIKEN:

🤖 PERFORMANCE:
⏱️ Runtime: {stats["uptime_minutes"]:.1f} minutes
📝 Transcriptions: {stats["transcriptions"]}
🌐 Translations: {stats["translations"]}
🎯 Cache Hit Rate: {stats["cache_hit_rate"]}

💻 SYSTEM:
🖥️ CPU: {cpu:.1f}%
🧠 RAM: {memory_used}MB ({memory_percent:.1f}%)
⚡ Status: {health_status}

🎬 Subtitle mode: {"Active" if self.subtitle_mode else "Inactive"}
"""
            DarkMessageBox.showinfo("Performance Statistics", stats_text, self.root)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            self.update_status(f"❌ Statistics error: {e}")

    def show_advanced_settings(self) -> None:
        AdvancedSettingsDialog(self.root, self)

    def _handle_status_update(self, state_info: Dict[str, Any]) -> None:
        def update_task() -> None:
            try:
                if "status" in state_info and hasattr(self, "status_label"):
                    if self.status_label.winfo_exists():
                        self.status_label.config(text=state_info["status"][:100])
                if "buttons" in state_info:
                    buttons = state_info["buttons"]
                    if (
                        hasattr(self, "start_button")
                        and self.start_button.winfo_exists()
                    ):
                        self.start_button.config(state=buttons.get("start", "normal"))
                    if hasattr(self, "stop_button") and self.stop_button.winfo_exists():
                        self.stop_button.config(state=buttons.get("stop", "disabled"))
                elif "processing_state" in state_info:
                    processing = state_info["processing_state"]
                    if (
                        hasattr(self, "start_button")
                        and self.start_button.winfo_exists()
                    ):
                        self.start_button.config(
                            state="disabled" if processing else "normal"
                        )
                    if hasattr(self, "stop_button") and self.stop_button.winfo_exists():
                        self.stop_button.config(
                            state="normal" if processing else "disabled"
                        )
                elif "controller_state" in state_info:
                    state_name = state_info["controller_state"]
                    if state_name == "PROCESSING":
                        if (
                            hasattr(self, "start_button")
                            and self.start_button.winfo_exists()
                        ):
                            self.start_button.config(state="disabled")
                        if (
                            hasattr(self, "stop_button")
                            and self.stop_button.winfo_exists()
                        ):
                            self.stop_button.config(state="normal")
                    elif state_name in ("IDLE", "ERROR"):
                        if (
                            hasattr(self, "start_button")
                            and self.start_button.winfo_exists()
                        ):
                            self.start_button.config(state="normal")
                        if (
                            hasattr(self, "stop_button")
                            and self.stop_button.winfo_exists()
                        ):
                            self.stop_button.config(state="disabled")
                if "stream_info" in state_info:
                    stream_info = state_info["stream_info"]
                    self.current_stream_info = stream_info
                    if (
                        hasattr(self, "stream_title_label")
                        and self.stream_title_label.winfo_exists()
                    ):
                        title = (
                            stream_info.title[:80] + "..."
                            if len(stream_info.title) > 80
                            else stream_info.title
                        )
                        self.stream_title_label.config(text=f"📡 {title}")
                    if (
                        hasattr(self, "stream_details_label")
                        and self.stream_details_label.winfo_exists()
                    ):
                        details = f"👤 {stream_info.uploader}"
                        if stream_info.duration and stream_info.duration != "Live":
                            details += f" | ⏱️ {stream_info.duration}"
                        self.stream_details_label.config(text=details)
                if state_info.get("file_finished"):
                    logger.info("📂 Dateiende erkannt – öffne Speicherdialog")
                    if self.settings.auto_save_on_completion:
                        self.save_transcript()
                    else:
                        self.update_status(
                            "✅ Dateiende – zum Speichern auf 💾 klicken"
                        )
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Status update error: {e}")

        if hasattr(self, "gui_queue") and hasattr(self, "queue_manager"):
            self.queue_manager.safe_put("gui", ("status_update", update_task))
        else:
            if hasattr(self, "root") and self.root.winfo_exists():
                self.root.after(0, update_task)

    def handle_transcription(self, result: TranscriptionResult) -> None:
        if not result or not result.text or not result.text.strip():
            return
        current_text = result.text.strip()
        current_text = re.sub(r'[^\w\s\.\,\!\?\:\;\-]', '', current_text)

        blacklist = getattr(self.advanced_settings, "blacklist", [])
        if blacklist:
            mode = getattr(self.advanced_settings, "blacklist_mode", "word")
            detected_lang = getattr(result, "language", "unknown")
            use_substring = (detected_lang in self.ASIAN_LANGUAGES) or (mode == "substring")
            current_lower = current_text.lower()
            for phrase in blacklist:
                if not phrase:
                    continue
                if use_substring:
                    if phrase.lower() in current_lower:
                        logger.debug(f"Blacklist-Treffer (substring): '{phrase}' in '{current_text}'")
                        return
                else:
                    pattern = r'\b' + re.escape(phrase.lower()) + r'\b'
                    if re.search(pattern, current_lower):
                        logger.debug(f"Blacklist-Treffer (word): '{phrase}' in '{current_text}'")
                        return

        with self._duplicate_lock:
            if current_text == self._last_transcription_text:
                return
            self._last_transcription_text = current_text

        self.performance_monitor.log_transcription()
        with self._history_lock:
            self.transcript_history.append(result)    

        try:
            if self.subtitle_mode and result.start is not None:
                timestamp = self.export_manager._format_timestamp_srt(result.start)
            else:
                timestamp = datetime.now().strftime("%H:%M:%S")
            detected_lang = getattr(result, "language", "unknown")
            lang_code = LANGUAGE_SHORT_CODES.get(detected_lang, "??")
            text = f"[{timestamp}] [{lang_code}] {current_text}\n"
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Error preparing transcription text: {e}")
            return

        if hasattr(self, "_text_update_queue") and hasattr(self, "queue_manager"):
            self.queue_manager.safe_put("text", ("transcript", text))
        else:
            logger.warning("⚠️ _text_update_queue missing, update wird verworfen")

    def handle_translation(self, result: TranslationResult) -> None:
        if not result or not result.translated or not result.translated.strip():
            return
        current_text = result.translated.strip()

        blacklist = getattr(self.advanced_settings, "blacklist", [])
        mode = getattr(self.advanced_settings, "blacklist_mode", "word")
        use_substring = (mode == "substring")
        for phrase in blacklist:
            if not phrase:
                continue
            if use_substring:
                if phrase.lower() in current_text.lower():
                    logger.debug(f"Blacklist-Treffer (substring, translation): '{phrase}' in '{current_text}' – Segment ignoriert")
                    return
            else:
                pattern = r'\b' + re.escape(phrase) + r'\b'
                if re.search(pattern, current_text, re.IGNORECASE):
                    logger.debug(f"Blacklist-Treffer (word, translation): '{phrase}' in '{current_text}' – Segment ignoriert")
                    return

        with self._duplicate_lock:
            if current_text == self._last_translation_text:
                return
            self._last_translation_text = current_text
        self.performance_monitor.log_translation()
        with self._history_lock:
            self.translation_history.append(result)
        try:
            if self.subtitle_mode and result.start is not None:
                timestamp = self.export_manager._format_timestamp_srt(result.start)
            else:
                timestamp = datetime.now().strftime("%H:%M:%S")
            text = f"[{timestamp}] {current_text}\n"
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Error preparing translation text: {e}")
            return
        if hasattr(self, "_text_update_queue") and hasattr(self, "queue_manager"):
            self.queue_manager.safe_put("text", ("translation", text))
        else:
            logger.warning("⚠️ _text_update_queue missing, update wird verworfen")

    def handle_info(self, info_msg: str) -> None:
        def update() -> None:
            if not self._shutting_down:
                self.update_status(f"ℹ️ {info_msg}")

        if (
            not self._shutting_down
            and hasattr(self, "root")
            and self.root.winfo_exists()
        ):
            self.root.after(0, update)

    def handle_error(self, error_msg: str) -> None:
        def update() -> None:
            try:
                self.update_status(f"❌ {error_msg}")
                logger.debug("handle_error: setzte Controller zurück")
                self.controller.stop_processing()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Fehler im error‑Handler: {e}")

        if (
            not self._shutting_down
            and hasattr(self, "root")
            and self.root.winfo_exists()
        ):
            self.root.after(0, update)
        else:
            logger.debug(
                f"handle_error übersprungen (Shutdown oder kein Root): {error_msg}"
            )

    def _on_processing_finished(self) -> None:
        logger.info("Processing finished – GUI kann reagieren")

    def update_status(self, message: str) -> None:
        if self._shutting_down:
            return
        if not hasattr(self, "root") or not self.root.winfo_exists():
            return
        short_msg = message[:100]

        def _update():
            try:
                if hasattr(self, "status_label") and self.status_label.winfo_exists():
                    self.status_label.config(text=short_msg)
            except tk.TclError:
                pass
            except Exception:
                pass

        try:
            self.root.after(0, _update)
        except Exception:
            pass

    def update_stream_info(self, info: StreamInfo) -> None:
        def update_gui() -> None:
            try:
                self.current_stream_info = info
                if (
                    hasattr(self, "stream_title_label")
                    and self.stream_title_label.winfo_exists()
                ):
                    title = (
                        info.title[:80] + "..." if len(info.title) > 80 else info.title
                    )
                    self.stream_title_label.config(text=f"📡 {title}")
                if (
                    hasattr(self, "stream_details_label")
                    and self.stream_details_label.winfo_exists()
                ):
                    details = f"👤 {info.uploader}"
                    if info.duration and info.duration != "Live":
                        details += f" | ⏱️ {info.duration}"
                    self.stream_details_label.config(text=details)
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Stream info update error: {e}")

        if (
            not self._shutting_down
            and hasattr(self, "root")
            and self.root.winfo_exists()
        ):
            self.root.after(0, update_gui)

    def _try_fallback_gui(self) -> None:
        logger.info("ℹ️ Starte im Kommandozeilen-Modus...")
        raise RuntimeError("Bitte installieren Sie Tkinter: pip install tk")

    def _show_error_and_exit(self, message: str) -> None:
        logger.error(f"💥 KRITISCHER FEHLER: {message}")
        try:
            import tkinter.messagebox as mb

            mb.showerror("Dragon Whisperer - Fehler", message)
        except Exception:
            pass
        self._emergency_cleanup()
        sys.exit(1)

    def _show_warning(self, message: str) -> None:
        logger.warning(f"⚠️ WARNUNG: {message}")

    def show_translation_dialog(self) -> None:
        if hasattr(self, "translation_engine"):
            TranslationDialog(self.root, self.translation_engine)
        else:
            DarkMessageBox.showerror(
                "Error", "Translation engine not available", self.root
            )

    def show_summarize_dialog(self) -> None:
        if not OLLAMA_AVAILABLE:
            DarkMessageBox.showerror(
                "Fehler",
                "Ollama nicht verfügbar (requests nicht installiert)",
                self.root,
            )
            return
        if hasattr(self, "transcript_text") and self.transcript_text.winfo_exists():
            text = self.transcript_text.get("1.0", "end-1c").strip()
        else:
            text = ""
        if not text:
            DarkMessageBox.showwarning(
                "Kein Text",
                "Kein Transkriptions-Text zum Zusammenfassen vorhanden.",
                self.root,
            )
            return
        SummarizeDialog(self.root, text, self)

    def show_install_dialog(self) -> None:
        InstallDependencyDialog(self.root, self)

    def _start_system_monitoring(self) -> None:
        def monitor() -> None:
            try:
                import psutil

                psutil_available = True
            except ImportError:
                psutil_available = False

            pynvml_available = False
            try:
                import pynvml

                pynvml_available = True
            except ImportError:
                pass

            try:
                if psutil_available:
                    cpu = psutil.cpu_percent(interval=None)
                    memory = psutil.virtual_memory()
                    ram_used = memory.used // (1024**2)
                    ram_total = memory.total // (1024**2)
                else:
                    cpu = ram_used = ram_total = 0

                gpu_text = ""
                pynvml_warning = ""
                if pynvml_available:
                    try:
                        pynvml.nvmlInit()
                        handle = pynvml.nvmlDeviceGetHandleByIndex(0)
                        memory_info = pynvml.nvmlDeviceGetMemoryInfo(handle)
                        total_gb = memory_info.total / (1024**3)
                        used_gb = 0.0
                        pid = os.getpid()
                        try:
                            processes = pynvml.nvmlDeviceGetComputeRunningProcesses(
                                handle
                            )
                            for proc in processes:
                                if proc.pid == pid:
                                    used_gb = proc.usedGpuMemory / (1024**3)
                                    break
                        except pynvml.NVMLError:
                            pass
                        if used_gb == 0.0:
                            try:
                                processes = (
                                    pynvml.nvmlDeviceGetGraphicsRunningProcesses(handle)
                                )
                                for proc in processes:
                                    if proc.pid == pid:
                                        used_gb = proc.usedGpuMemory / (1024**3)
                                        break
                            except pynvml.NVMLError:
                                pass
                        gpu_text = f" | 🎮 VRAM: {used_gb:.1f}/{total_gb:.1f}GB"
                    except pynvml.NVMLError:
                        gpu_text = " | 🎮 GPU: Fehler"
                else:
                    if TORCH_AVAILABLE:
                        try:
                            torch = FastLazyLoader.load("torch")
                            if torch.cuda.is_available():
                                used_gb = torch.cuda.memory_allocated() / (1024**3)
                                total_gb = torch.cuda.get_device_properties(
                                    0
                                ).total_memory / (1024**3)
                                gpu_text = f" | 🎮 VRAM: {used_gb:.1f}/{total_gb:.1f}GB (via torch)"
                                pynvml_warning = " ⚠️ pynvml fehlt – Werte ungenau (pip install pynvml)"
                            else:
                                gpu_text = " | 🎮 GPU: ❌"
                        except Exception:
                            gpu_text = " | 🎮 GPU: Fehler"
                    else:
                        gpu_text = " | 🎮 GPU: N/A"

                current_model = "None"
                if hasattr(self, "transcription_engine"):
                    current_model = self.transcription_engine.get_current_model()
                demo_hint = " | ⚠️ Demo" if getattr(self, "demo_mode", False) else ""

                if IS_WINDOWS:
                    info = f"🪟 Windows | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text}{pynvml_warning} | 🤖 Model: {current_model}{demo_hint}"
                elif IS_MACOS:
                    if IS_ARM:
                        info = f"🍎 macOS ARM | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text}{pynvml_warning} | 🤖 Model: {current_model}{demo_hint}"
                    else:
                        info = f"🍎 macOS Intel | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text}{pynvml_warning} | 🤖 Model: {current_model}{demo_hint}"
                elif IS_LINUX:
                    info = f"🐧 Linux | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text}{pynvml_warning} | 🤖 Model: {current_model}{demo_hint}"
                else:
                    info = f"🌐 System | 🔧 CPU: {cpu:.0f}% | 💾 RAM: {ram_used}/{ram_total}MB{gpu_text}{pynvml_warning} | 🤖 Model: {current_model}{demo_hint}"

                if (
                    hasattr(self, "system_info_label")
                    and self.system_info_label.winfo_exists()
                ):
                    self.system_info_label.config(text=info)

            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ System monitoring error: {e}")

            if hasattr(self, "root") and self.root.winfo_exists():
                self.root.after(3000, monitor)

        if hasattr(self, "root") and self.root.winfo_exists():
            self.root.after(1000, monitor)

    def _update_ollama_button_state(self):
        if not hasattr(self, "correct_btn"):
            return
        if not OLLAMA_AVAILABLE:
            self.correct_btn.config(state="disabled", text="🔧 (kein Ollama)")
            return
        summarizer = OllamaSummarizer(
            self,
            model=self.advanced_settings.ollama_model,
            host=self.advanced_settings.ollama_host,
        )
        if not summarizer.is_server_reachable():
            self.correct_btn.config(state="disabled", text="🔧 (Server aus)")
            ToolTip(
                self.correct_btn, "Ollama-Server läuft nicht – starte 'ollama serve'"
            )
        else:
            self.correct_btn.config(state="normal", text="🔧")
            ToolTip(self.correct_btn, "Transkript mit Ollama korrigieren")

    @gui_operation_decorator
    def correct_transcript_with_ollama(self):
        if (
            not hasattr(self, "transcript_text")
            or not self.transcript_text.winfo_exists()
        ):
            return
        text = self.transcript_text.get("1.0", "end-1c").strip()
        if not text:
            DarkMessageBox.showinfo(
                "Hinweis", "Kein Text zum Korrigieren vorhanden.", self.root
            )
            return

        if not OLLAMA_AVAILABLE:
            DarkMessageBox.showerror(
                "Fehler", "Ollama nicht verfügbar (requests fehlt).", self.root
            )
            return

        summarizer = OllamaSummarizer(
            self,
            model=self.advanced_settings.ollama_model,
            host=self.advanced_settings.ollama_host,
        )
        if not summarizer.is_server_reachable():
            DarkMessageBox.showerror(
                "Fehler",
                "Ollama-Server läuft nicht.\nBitte starte 'ollama serve' und versuche es erneut.",
                self.root,
            )
            return

        progress = DarkMessageBox.show_progress(
            "Korrektur läuft",
            "Sende Text an Ollama zur Korrektur...",
            parent=self.root,
            indeterminate=True,
        )

        corrected_parts = []
        error_occurred = False

        def on_chunk(chunk: str):
            corrected_parts.append(chunk)
            self.root.after(
                0,
                lambda: progress.update_message(
                    f"Empfange Daten... ({len(corrected_parts)} Teile)"
                ),
            )

        def on_error(error: str):
            nonlocal error_occurred
            error_occurred = True
            self.root.after(0, progress.close)
            self.root.after(
                0,
                lambda: DarkMessageBox.showerror(
                    "Fehler", f"Korrektur fehlgeschlagen:\n{error}", self.root
                ),
            )

        def on_complete():
            self.root.after(0, progress.close)
            if not error_occurred:
                corrected_text = "".join(corrected_parts).strip()
                if corrected_text:
                    self.root.after(
                        0,
                        lambda: self._update_transcript_with_correction(corrected_text),
                    )
                else:
                    self.root.after(
                        0, lambda: self.update_status("⚠️ Korrektur ergab leeren Text")
                    )

        summarizer.correct_transcript(
            text,
            callback=on_chunk,
            error_callback=on_error,
            complete_callback=on_complete,
        )

    @gui_operation_decorator
    def _update_transcript_with_correction(self, corrected_text: str):
        try:
            if self.transcript_text and self.transcript_text.winfo_exists():
                self.transcript_text.delete("1.0", "end")
                self.transcript_text.insert("1.0", corrected_text)
                self.update_status("✅ Transkription korrigiert")
        except tk.TclError:
            pass

    def speak_current_text(self):
        focused = self.root.focus_get()
        text = None
        source = ""

        if focused == self.transcript_text:
            try:
                if self.transcript_text.tag_ranges(tk.SEL):
                    text = self.transcript_text.get(tk.SEL_FIRST, tk.SEL_LAST)
                    source = "Auswahl (Transkript)"
                else:
                    text = self.transcript_text.get("1.0", "end-1c")
                    source = "gesamtes Transkript"
            except tk.TclError:
                pass
        elif focused == self.translation_text:
            try:
                if self.translation_text.tag_ranges(tk.SEL):
                    text = self.translation_text.get(tk.SEL_FIRST, tk.SEL_LAST)
                    source = "Auswahl (Übersetzung)"
                else:
                    text = self.translation_text.get("1.0", "end-1c")
                    source = "gesamte Übersetzung"
            except tk.TclError:
                pass

        if not text or not text.strip():
            self.update_status("❌ Kein Text zum Vorlesen ausgewählt.")
            return

        if not self.tts_manager.is_available():
            self.update_status("❌ Keine TTS-Engine verfügbar (piper oder pyttsx3 fehlen).")
            return

        self.update_status(f"🔊 Lese {source} vor...")

        def tts_callback(success: bool, message: str):
            if success:
                self.update_status("✅ Sprachausgabe beendet.")
            else:
                self.update_status(f"❌ Fehler bei Sprachausgabe: {message}")

        self.tts_manager.speak(text, callback=tts_callback)

    # --- Änderung: VAD-Fallback umschalten und speichern
    def toggle_vad_fallback(self):
        current = self.vad_fallback_enabled.get()
        self.vad_fallback_enabled.set(not current)
        self._update_vad_fallback_button()
        self.advanced_settings.vad_fallback_enabled = not current
        self.advanced_settings.save_to_file()
        if hasattr(self, 'audio_processor'):
            self.audio_processor.set_vad_fallback_enabled(not current)
        self.update_status(f"VAD-Fallback {'aktiviert' if not current else 'deaktiviert'}")

    def _update_vad_fallback_button(self):
        if hasattr(self, 'vad_fallback_btn') and self.vad_fallback_btn.winfo_exists():
            new_text = "🔁 VAD-Fallback ON" if self.vad_fallback_enabled.get() else "🔁 VAD-Fallback OFF"
            self.vad_fallback_btn.config(text=new_text)

    # --- Änderung: Live-Mode umschalten (20s <-> 10s)
    def toggle_live_mode(self):
        if self.is_processing:
            DarkMessageBox.showwarning(
                "Achtung",
                "Live-Modus kann während einer laufenden Transkription nicht geändert werden.\nBitte stoppen Sie zuerst die Verarbeitung.",
                self.root
            )
            return
        current_duration = self.advanced_settings.chunk_duration
        if current_duration == 20.0:
            new_duration = 10.0
        else:
            new_duration = 20.0
        self.advanced_settings.chunk_duration = new_duration
        self.advanced_settings.save_to_file()
        if hasattr(self, 'audio_processor'):
            self.audio_processor.config.CHUNK_DURATION = new_duration
            self.audio_processor._update_chunk_size()
        self._update_live_mode_button()
        self.update_status(f"Chunk-Dauer auf {new_duration:.0f}s umgestellt")

    def _update_live_mode_button(self):
        if hasattr(self, 'live_mode_btn') and self.live_mode_btn.winfo_exists():
            current = self.advanced_settings.chunk_duration
            self.live_mode_btn.config(text=f"⏱️ {current:.0f}s")

    def _final_initialization_check(self) -> None:
        logger.info("✅ Dragon Whisperer initialisiert")
        if getattr(self, "demo_mode", False):
            self.update_status(
                "⚠️ Demo-Modus: Whisper nicht verfügbar – verwende Dummy-Transkriptionen"
            )
        self.root.after(500, self._update_ollama_button_state)
        # VAD-Fallback-Button initial aktualisieren
        self._update_vad_fallback_button()
        self._update_live_mode_button()

    def run(self):
        self.root.mainloop()

    def _emergency_cleanup(self) -> None:
        logger.info("🆘 Emergency cleanup...")
        self._minimal_emergency_cleanup()

    @gui_operation_decorator
    def update_progress(
        self, processed: int, total: Optional[int], chunks: int
    ) -> None:
        if not hasattr(self, "progress_bar") or not self.progress_bar.winfo_exists():
            return
        try:
            if total is not None and total > 0:
                if not self.progress_bar.winfo_ismapped():
                    self.progress_bar.pack(side="left", padx=(10, 10))
                percent = (processed / total) * 100
                self.progress_bar.config(mode="determinate", value=percent)
                mb = processed // (1024 * 1024)
                tb = total // (1024 * 1024)
                remaining_text = ""
                if self.audio_processor and self.audio_processor._expected_duration:
                    expected = self.audio_processor._expected_duration
                    processed_secs = self.audio_processor._processed_seconds
                    remaining = max(0, expected - processed_secs)
                    remaining_text = f" | ⏳ {int(remaining//60):02d}:{int(remaining%60):02d} verbleibend"
                self.progress_label.config(text=f"{mb}MB/{tb}MB{remaining_text}")
            else:
                if self.progress_bar.winfo_ismapped():
                    self.progress_bar.pack_forget()
                self.progress_label.config(
                    text=f"Chunks: {chunks}  |  Daten: {processed // 1024} KB"
                )
        except tk.TclError:
            pass

    @gui_operation_decorator
    def _reset_progress(self):
        if hasattr(self, "progress_bar") and self.progress_bar.winfo_exists():
            self.progress_bar.stop()
            self.progress_bar.config(mode="determinate", value=0)
        self._progress_bar_started = False
        if hasattr(self, "progress_label"):
            self.progress_label.config(text="")

    def _bind_shortcuts(self):
        mod = "Command" if IS_MACOS else "Control"
        self.root.bind(f"<{mod}-o>", lambda e: self.select_file_dark())
        self.root.bind(f"<{mod}-v>", lambda e: self.paste_url())
        self.root.bind(f"<{mod}-Return>", lambda e: self._on_start_click())
        self.root.bind(f"<{mod}-q>", lambda e: self._safe_exit_dialog())
        self.root.bind(f"<{mod}-s>", lambda e: self.save_transcript())
        self.root.bind(f"<{mod}-l>", lambda e: self.toggle_layout())
        self.root.bind(f"<{mod}-t>", lambda e: self.toggle_translation())
        self.root.bind(f"<{mod}-e>", lambda e: self.export_subtitles())
        self.root.bind(f"<{mod}-u>", lambda e: self.toggle_subtitle_mode())
        self.root.bind(f"<{mod}-Shift-c>", lambda e: self.clear_all())
        self.root.bind(f"<{mod}-h>", lambda e: self.show_shortcuts_help())
        self.root.bind("<F1>", lambda e: self.show_shortcuts_help())
        self.url_entry.bind(f"<{mod}-v>", lambda e: "break")

    def show_shortcuts_help(self):
        ShortcutsDialog(self.root)

    def update_translation_engine(self):
        old_engine = self.translation_engine
        self.translation_engine = self._create_translation_engine()
        if hasattr(self, 'audio_processor'):
            self.audio_processor.translation_engine = self.translation_engine
        if hasattr(old_engine, 'dispose'):
            old_engine.dispose()
        logger.info("Übersetzungs-Engine aktualisiert")


# =============================================================================
# 12. LINUX PERFORMANCE OPTIMIZER
# =============================================================================

PSUTIL_AVAILABLE = importlib.util.find_spec("psutil") is not None
if not PSUTIL_AVAILABLE:
    logger.warning(
        "⚠️ psutil nicht verfügbar – Linux Performance Optimizer läuft im Dummy-Modus"
    )

if IS_LINUX and PSUTIL_AVAILABLE:

    class LinuxPerformanceOptimizer:
        def __init__(self, gui_ref: "DragonWhispererGUI") -> None:
            self.gui = gui_ref
            self.is_processing = False
            self._original_settings: Dict[str, Any] = {}
            self._optimization_active = False
            self._monitoring_thread: Optional[threading.Thread] = None
            self._shutdown_event = threading.Event()
            self._monitoring_lock = threading.RLock()
            self._last_gui_access_time = 0.0
            self._gui_access_warning_printed = False

        def _is_gui_available_safe(self) -> bool:
            try:
                if self.gui is None:
                    return False
                return (
                    hasattr(self.gui, "root")
                    and self.gui.root is not None
                    and self.gui.root.winfo_exists()
                    and not getattr(self.gui, "_shutting_down", False)
                )
            except Exception:
                return False

        def optimize_for_processing(self) -> None:
            if not IS_LINUX or self._optimization_active:
                return
            with self._monitoring_lock:
                self._shutdown_event.clear()
                if self._monitoring_thread and self._monitoring_thread.is_alive():
                    logger.warning(
                        "⚠️ Optimize: Monitoring-Thread läuft bereits – überspringe"
                    )
                    return
                logger.info("🔧 Aktiviere Linux-Performance-Optimierungen...")
                self._optimize_text_widget("transcript_text")
                self._optimize_text_widget("translation_text")
                self._schedule_batch_interval_increase()
                self._clean_queue_safe(getattr(self.gui, "gui_queue", None), 15)
                self._apply_linux_specific_optimizations()
                self._optimization_active = True
                self.is_processing = True
                self._start_performance_monitoring()
                logger.info("✅ Linux-Performance-Optimierungen aktiviert")

        def _schedule_batch_interval_increase(self) -> None:
            if not self._is_gui_available_safe():
                return
            if hasattr(self.gui, "_batch_update_interval"):
                self._original_settings["batch_update_interval"] = (
                    self.gui._batch_update_interval
                )

            def task() -> None:
                if self._is_gui_available_safe() and hasattr(self.gui, "_batch_update_interval"):
                    self.gui._batch_update_interval = 250

            if self._is_gui_available_safe():
                self.gui.root.after(0, task)

        def _optimize_text_widget(self, attr_name: str) -> None:
            if not self._is_gui_available_safe():
                return
            widget = getattr(self.gui, attr_name, None)
            if widget and widget.winfo_exists():
                self._original_settings[attr_name] = {
                    "maxundo": widget.cget("maxundo"),
                    "undo": widget.cget("undo"),
                    "autoseparators": widget.cget("autoseparators"),
                }
                widget.configure(maxundo=5, undo=True, autoseparators=True, height=12)

        def _apply_linux_specific_optimizations(self) -> None:
            if not self._is_gui_available_safe():
                return
            self._detect_compositor()
            self._increase_resource_limits()

        def _detect_compositor(self) -> None:
            try:
                import psutil

                for proc in psutil.process_iter(["name"]):
                    try:
                        name = proc.info["name"].lower()
                        if any(
                            c in name for c in ["compton", "picom", "compiz", "kwin"]
                        ):
                            logger.info(f"  ↪ Compositor erkannt: {name}")
                            break
                    except (psutil.NoSuchProcess, psutil.AccessDenied):
                        continue
            except (ImportError, AttributeError):
                pass

        def _increase_resource_limits(self) -> None:
            try:
                import resource

                try:
                    soft, hard = resource.getrlimit(resource.RLIMIT_DATA)
                    new_soft = min(hard, 1024 * 1024 * 1024)
                    if new_soft > soft:
                        resource.setrlimit(resource.RLIMIT_DATA, (new_soft, hard))
                        logger.info(f"  ↪ Daten-Limit erhöht: {soft} → {new_soft}")
                except (resource.error, ValueError, PermissionError) as e:
                    logger.debug(f"  ⚠️ Daten-Limit konnte nicht erhöht werden: {e}")

                try:
                    soft_fd, hard_fd = resource.getrlimit(resource.RLIMIT_NOFILE)
                    new_soft_fd = min(hard_fd, 8192)
                    if new_soft_fd > soft_fd:
                        resource.setrlimit(
                            resource.RLIMIT_NOFILE, (new_soft_fd, hard_fd)
                        )
                        logger.info(
                            f"  ↪ Dateideskriptoren-Limit erhöht: {soft_fd} → {new_soft_fd}"
                        )
                except (resource.error, ValueError, PermissionError) as e:
                    logger.debug(
                        f"  ⚠️ Dateideskriptoren-Limit konnte nicht erhöht werden: {e}"
                    )

                try:
                    os.nice(-5)
                    logger.info("  ↪ CPU-Priorität erhöht (nice -5)")
                except PermissionError:
                    logger.debug("  ⚠️ CPU-Priorität konnte nicht erhöht werden (keine root-Rechte)")
                except Exception as e:
                    logger.warning(
                        f"  ⚠️ CPU-Priorität konnte nicht angepasst werden: {e}"
                    )

            except ImportError:
                logger.debug(
                    "  resource-Modul nicht verfügbar – überspringe Limit-Erhöhung"
                )
            except Exception as e:
                logger.warning(
                    f"  ⚠️ Unerwarteter Fehler in _increase_resource_limits: {e}"
                )

        def _start_performance_monitoring(self) -> None:
            with self._monitoring_lock:
                if self._monitoring_thread and self._monitoring_thread.is_alive():
                    return

                def monitor_worker() -> None:
                    if self._shutdown_event.is_set():
                        return
                    logger.info("🔍 Linux-Performance-Monitoring gestartet")
                    check_count = 0
                    max_checks = 240

                    while (
                        not self._shutdown_event.is_set() and self._optimization_active
                    ):
                        try:
                            time.sleep(30)
                            if self._shutdown_event.is_set():
                                break
                            if not self._is_gui_available_safe():
                                logger.warning(
                                    "⚠️ GUI nicht mehr verfügbar – stoppe Monitoring"
                                )
                                break

                            check_count += 1
                            if check_count > max_checks:
                                logger.info("⏰ Monitoring-Zeitlimit erreicht – beende")
                                break

                            system_load = self._get_system_load()
                            memory_usage = self._get_memory_usage()

                            if system_load > 0.8 or memory_usage > 0.85:
                                self._schedule_adjustments(system_load, memory_usage)

                            if check_count % 12 == 0:
                                self._print_performance_report(
                                    system_load, memory_usage
                                )

                        except (OSError, ValueError) as e:
                            logger.warning(f"⚠️ Fehler im Monitoring-Thread: {e}")

                    logger.info("✅ Linux-Performance-Monitoring beendet")
                    with self._monitoring_lock:
                        self._monitoring_thread = None

                self._monitoring_thread = threading.Thread(
                    target=monitor_worker, daemon=True, name="LinuxPerfMon"
                )
                self._monitoring_thread.start()

        def _get_system_load(self) -> float:
            try:
                load_avg = os.getloadavg()
                cpu_count = os.cpu_count() or 1
                return load_avg[0] / cpu_count
            except (OSError, ValueError):
                return 0.0

        def _get_memory_usage(self) -> float:
            try:
                import psutil

                memory = psutil.virtual_memory()
                return memory.percent / 100.0
            except Exception:
                return 0.0

        def _schedule_adjustments(self, load: float, memory: float) -> None:
            if not self._is_gui_available_safe():
                return

            def task() -> None:
                if not self._optimization_active or self._shutdown_event.is_set():
                    return
                if not self._is_gui_available_safe():
                    return
                adjustments: List[str] = []

                if memory > 0.85 and hasattr(self.gui, "gui_queue"):
                    cleared = self._clean_queue_safe(self.gui.gui_queue, 5)
                    if cleared > 0:
                        adjustments.append(f"Queue: -{cleared}")

                if load > 0.8 and hasattr(self.gui, "_batch_update_interval"):
                    current = self.gui._batch_update_interval
                    if current < 500:
                        self.gui._batch_update_interval = min(500, current + 50)
                        adjustments.append(
                            f"Update: {current}→{self.gui._batch_update_interval}ms"
                        )

                if adjustments:
                    logger.info(f"🔧 Anpassungen: {', '.join(adjustments)}")

            if self._is_gui_available_safe():
                self.gui.root.after(0, task)

        def _clean_queue_safe(
            self, queue_obj: Optional[queue.Queue], target_size: int
        ) -> int:
            if not queue_obj:
                return 0
            try:
                if queue_obj.qsize() <= target_size:
                    return 0
            except Exception:
                return 0

            cleared = 0
            try:
                while queue_obj.qsize() > target_size and cleared < 50:
                    queue_obj.get_nowait()
                    cleared += 1
            except queue.Empty:
                pass
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ Queue-Cleanup-Fehler: {e}")
            return cleared

        def _print_performance_report(self, load: float, memory: float) -> None:
            try:
                stats = {
                    "System-Last": f"{load:.1%}",
                    "RAM-Auslastung": f"{memory:.1%}",
                    "Status": "Aktiv" if self.is_processing else "Inaktiv",
                }
                logger.info(
                    "🐧 Linux-Performance-Report: "
                    + " | ".join(f"{k}: {v}" for k, v in stats.items())
                )
            except Exception:
                pass

        def restore_normal_mode(self) -> None:
            if not IS_LINUX or not self._optimization_active:
                return
            logger.info("🔧 Linux-Optimierer: Fahre herunter...")
            self._shutdown_event.set()
            self._optimization_active = False
            self.is_processing = False

            if self._monitoring_thread and self._monitoring_thread.is_alive():
                self._monitoring_thread.join(timeout=2.0)
                self._monitoring_thread = None

            if self._is_gui_available_safe():
                logger.info("  ↪ Stelle GUI-Einstellungen wieder her...")
                self._restore_text_widget("transcript_text")
                self._restore_text_widget("translation_text")

                if "batch_update_interval" in self._original_settings:
                    saved_interval = self._original_settings["batch_update_interval"]

                    def restore_batch_interval(saved=saved_interval):
                        if self._is_gui_available_safe() and hasattr(self.gui, "_batch_update_interval"):
                            self.gui._batch_update_interval = saved

                    if self._is_gui_available_safe():
                        self.gui.root.after(0, restore_batch_interval)

            self._original_settings.clear()
            logger.info("✅ Linux-Optimierer heruntergefahren")

        def _restore_text_widget(self, attr_name: str) -> None:
            if not self._is_gui_available_safe():
                return
            widget = getattr(self.gui, attr_name, None)
            if (
                widget
                and widget.winfo_exists()
                and attr_name in self._original_settings
            ):
                try:
                    widget.configure(**self._original_settings[attr_name])
                    logger.info(f"    ✅ {attr_name} wiederhergestellt")
                except Exception as e:
                    if isinstance(e, (KeyboardInterrupt, SystemExit)):
                        raise
                    logger.warning(
                        f"    ⚠️ {attr_name} konnte nicht wiederhergestellt werden: {e}"
                    )

        def emergency_optimize(self) -> None:
            if self._shutdown_event.is_set():
                return
            logger.info("🚨 Führe Notfall-Optimierungen durch...")
            if not self._is_gui_available_safe():
                return

            def task() -> None:
                if not self._is_gui_available_safe():
                    return
                self._clean_queue_safe(getattr(self.gui, "gui_queue", None), 3)
                self._clean_queue_safe(getattr(self.gui, "_text_update_queue", None), 2)

                for attr in ["transcript_text", "translation_text"]:
                    widget = getattr(self.gui, attr, None)
                    if widget and widget.winfo_exists():
                        try:
                            widget.configure(height=6, maxundo=1)
                        except Exception:
                            pass

                if hasattr(self.gui, "_batch_update_interval"):
                    self.gui._batch_update_interval = 500

                gc.collect()
                logger.info("✅ Notfall-Optimierungen abgeschlossen")

            if self._is_gui_available_safe():
                self.gui.root.after(0, task)

        def get_optimization_status(self) -> Dict[str, Any]:
            return {
                "platform": SYSTEM,
                "optimization_active": self._optimization_active,
                "processing_active": self.is_processing,
                "monitoring_active": self._monitoring_thread
                and self._monitoring_thread.is_alive(),
                "shutdown_event_set": self._shutdown_event.is_set(),
                "original_settings_count": len(self._original_settings),
                "linux_specific": IS_LINUX,
                "gui_available": self._is_gui_available_safe(),
            }

        def dispose(self) -> None:
            logger.info("🧹 Linux-Performance-Optimierer wird entsorgt...")
            self._shutdown_event.set()
            try:
                self.restore_normal_mode()
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ restore_normal_mode fehlgeschlagen: {e}")
            self._original_settings.clear()
            gc.collect()
            logger.info("✅ Linux-Performance-Optimierer entsorgt")

elif IS_LINUX:

    class LinuxPerformanceOptimizer:
        def __init__(self, gui_ref: "DragonWhispererGUI") -> None:
            self.gui = gui_ref
            self.is_processing = False
            logger.info("🐧 LinuxPerformanceOptimizer: Dummy-Modus (psutil fehlt)")

        def optimize_for_processing(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): optimize_for_processing")

        def restore_normal_mode(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): restore_normal_mode")

        def emergency_optimize(self) -> None:
            logger.debug("LinuxPerformanceOptimizer (Dummy): emergency_optimize")

        def get_optimization_status(self) -> Dict[str, Any]:
            return {"platform": SYSTEM, "dummy_mode": True, "psutil_missing": True}

        def dispose(self) -> None:
            pass


# =============================================================================
# 13. HAUPTSCHLEIFE & START
# =============================================================================

BASE_LANGUAGES = {
    "auto": "Automatisch",
    "de": "Deutsch",
    "en": "Englisch",
    "fr": "Französisch",
    "es": "Spanisch",
    "it": "Italienisch",
    "pt": "Portugiesisch",
    "nl": "Niederländisch",
    "pl": "Polnisch",
    "ru": "Russisch",
    "ja": "Japanisch",
    "zh": "Chinesisch",
    "ko": "Koreanisch",
    "ar": "Arabisch",
    "hi": "Hindi",
    "tr": "Türkisch",
    "vi": "Vietnamesisch",
    "th": "Thailändisch",
    "id": "Indonesisch",
    "ms": "Malaysisch",
    "fa": "Persisch",
    "he": "Hebräisch",
    "bn": "Bengalisch",
    "ta": "Tamil",
    "te": "Telugu",
    "ml": "Malayalam",
    "kn": "Kannada",
    "mr": "Marathi",
    "gu": "Gujarati",
    "pa": "Punjabi",
    "ur": "Urdu",
    "sv": "Schwedisch",
    "da": "Dänisch",
    "no": "Norwegisch",
    "fi": "Finnisch",
    "cs": "Tschechisch",
    "hu": "Ungarisch",
    "ro": "Rumänisch",
    "bg": "Bulgarisch",
    "el": "Griechisch",
    "sk": "Slowakisch",
    "hr": "Kroatisch",
    "sr": "Serbisch",
    "uk": "Ukrainisch",
    "ca": "Katalanisch",
    "eu": "Baskisch",
    "gl": "Galizisch",
}

SUPPORTED_LANGUAGES = BASE_LANGUAGES.copy()
SORTED_LANGUAGES = sorted(
    [(name, code) for code, name in BASE_LANGUAGES.items()], key=lambda x: x[0]
)
LANGUAGE_SHORT_CODES = {code: name[:3] for code, name in BASE_LANGUAGES.items()}

WHISPER_MODELS: List[str] = [
    "tiny",
    "tiny.en",
    "base",
    "base.en",
    "small",
    "small.en",
    "medium",
    "medium.en",
    "large-v2",
    "large-v3",
]


# -----------------------------------------------------------------------------
# AdvancedSettings
# -----------------------------------------------------------------------------
@dataclass
class AdvancedSettings:
    # --- Modell-Parameter ---
    beam_size: int = 10
    temperature: float = Constants.DEFAULT_TEMPERATURE
    vad_filter: bool = False
    gpu_acceleration: bool = True
    max_cache_size: int = 200
    auto_save_interval: int = 300
    enable_sentiment_analysis: bool = False
    enable_speaker_diarization: bool = False
    max_memory_mb: int = 1024
    optimize_translations: bool = False

    config_type: str = "high_accuracy"

    transcript_max_lines: int = 400
    translation_max_lines: int = 300

    translation_engine: str = "google"
    ollama_model: str = "llama3.1:8b"
    ollama_host: str = "http://localhost:11434"

    asian_mode: bool = False
    precision_mode: bool = False

    audio_profile: str = "transcription"

    adaptive_chunk: bool = False
    duplicate_similarity_threshold: float = 0.98
    adaptive_chunk_low_words: int = 3
    adaptive_chunk_high_words: int = 10

    min_confidence: float = 0.1

    vad_threshold: float = 0.2
    vad_min_speech_duration_ms: int = 150
    vad_min_silence_duration_ms: int = 50

    max_empty_reads: int = 30

    blacklist: List[str] = field(default_factory=list)

    blacklist_mode: str = "word"

    enable_noise_reduction: bool = False

    enable_duplicate_check: bool = False

    hotwords: str = ""

    transcription_workers: int = 2
    translation_workers: int = 1

    tts_engine: str = "piper"

    best_of: int = 5
    patience: float = 1.0
    log_prob_threshold: float = -1.2
    compression_ratio_threshold: float = 2.8
    condition_on_previous_text: bool = True
    no_speech_threshold: float = 0.6
    suppress_tokens: str = "-1"

    # --- Änderung: VAD-Fallback-Einstellung
    vad_fallback_enabled: bool = True

    config: Config = field(init=False, repr=False, compare=False)

    _chunk_duration: float = field(default=Constants.BASE_CHUNK_DURATION, init=False, repr=False)

    @property
    def chunk_duration(self) -> float:
        return self._chunk_duration

    @chunk_duration.setter
    def chunk_duration(self, value: float):
        if hasattr(self, 'config') and self.config:
            min_dur = self.config.MIN_CHUNK_DURATION
            max_dur = self.config.MAX_CHUNK_DURATION
            if not (min_dur <= value <= max_dur):
                logger.warning(f"Chunk duration {value}s out of range, clamping to [{min_dur}, {max_dur}]")
                value = max(min_dur, min(value, max_dur))
        self._chunk_duration = value
        if hasattr(self, 'config') and self.config:
            self.config.CHUNK_DURATION = value

    def __post_init__(self) -> None:
        self._recreate_config()
        self.chunk_duration = self.config.CHUNK_DURATION
        self._apply_mode_overrides()
        logger.info("🔊 Settings initialized (dataclass)")

    def _recreate_config(self) -> None:
        if self.config_type == "realtime":
            self.config = RealtimeConfig()
        elif self.config_type == "high_accuracy":
            self.config = HighAccuracyConfig()
        elif self.config_type == "youtube":
            self.config = YouTubeOptimizedConfig()
        else:
            self.config = Config()
        self.config.SAMPLE_RATE = Constants.SAMPLE_RATE
        self.config.CHANNELS = Constants.CHANNELS
        self.config.AUDIO_FORMAT = Constants.AUDIO_FORMAT
        self._chunk_duration = self.config.CHUNK_DURATION

    def _apply_mode_overrides(self) -> None:
        if self.asian_mode:
            self.chunk_duration = 10
        if self.precision_mode:
            self.chunk_duration = 7
            self.beam_size = 15
            self.temperature = 0.0
            self.vad_threshold = 0.25
            self.vad_min_speech_duration_ms = 260
            self.vad_min_silence_duration_ms = 110
            if self.asian_mode:
                self.chunk_duration = 10

    @classmethod
    def load_from_file(cls, filename: str = "dragon_advanced_settings.json") -> "AdvancedSettings":
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            if not file_path.exists():
                logger.info("📝 Keine gespeicherten Einstellungen, verwende Standard")
                return cls()

            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            valid_fields = {f.name for f in fields(cls) if f.init}
            filtered = {k: v for k, v in data.items() if k in valid_fields}

            if "beam_size" in filtered:
                filtered["beam_size"] = int(filtered["beam_size"])
            if "temperature" in filtered:
                filtered["temperature"] = float(filtered["temperature"])
            if "chunk_duration" in filtered:
                filtered["chunk_duration"] = float(filtered["chunk_duration"])
            if "duplicate_similarity_threshold" in filtered:
                filtered["duplicate_similarity_threshold"] = float(filtered["duplicate_similarity_threshold"])
            if "min_confidence" in filtered:
                filtered["min_confidence"] = float(filtered["min_confidence"])
            if "vad_threshold" in filtered:
                filtered["vad_threshold"] = float(filtered["vad_threshold"])
            if "vad_min_speech_duration_ms" in filtered:
                filtered["vad_min_speech_duration_ms"] = int(filtered["vad_min_speech_duration_ms"])
            if "vad_min_silence_duration_ms" in filtered:
                filtered["vad_min_silence_duration_ms"] = int(filtered["vad_min_silence_duration_ms"])
            if "blacklist_mode" in filtered:
                filtered["blacklist_mode"] = filtered["blacklist_mode"]
            if "hotwords" in filtered:
                filtered["hotwords"] = filtered["hotwords"]
            if "transcription_workers" in filtered:
                filtered["transcription_workers"] = int(filtered["transcription_workers"])
            if "translation_workers" in filtered:
                filtered["translation_workers"] = int(filtered["translation_workers"])
            if "tts_engine" in filtered:
                filtered["tts_engine"] = filtered["tts_engine"]
            if "best_of" in filtered:
                filtered["best_of"] = int(filtered["best_of"])
            if "patience" in filtered:
                filtered["patience"] = float(filtered["patience"])
            if "log_prob_threshold" in filtered:
                filtered["log_prob_threshold"] = float(filtered["log_prob_threshold"])
            if "compression_ratio_threshold" in filtered:
                filtered["compression_ratio_threshold"] = float(filtered["compression_ratio_threshold"])
            if "condition_on_previous_text" in filtered:
                filtered["condition_on_previous_text"] = bool(filtered["condition_on_previous_text"])
            if "no_speech_threshold" in filtered:
                filtered["no_speech_threshold"] = float(filtered["no_speech_threshold"])
            if "suppress_tokens" in filtered:
                filtered["suppress_tokens"] = filtered["suppress_tokens"]
            if "vad_fallback_enabled" in filtered:
                filtered["vad_fallback_enabled"] = bool(filtered["vad_fallback_enabled"])

            instance = cls(**filtered)
            if "chunk_duration" in data:
                instance.chunk_duration = data["chunk_duration"]
            logger.info(f"✅ Settings loaded successfully (Config Type: {instance.config_type})")
            return instance

        except (json.JSONDecodeError, OSError, PermissionError) as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler beim Laden der Einstellungen: {e}")
            return cls()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Unerwarteter Fehler beim Laden: {e}")
            return cls()

    def save_to_file(self, filename: str = "dragon_advanced_settings.json") -> None:
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            config_dir.mkdir(parents=True, exist_ok=True)
            file_path = config_dir / filename
            temp_path = file_path.with_suffix(".tmp")

            data = asdict(self)
            data.pop("config", None)
            data.pop("_chunk_duration", None)
            data["chunk_duration"] = self.chunk_duration

            with open(temp_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            try:
                os.replace(temp_path, file_path)
            except AttributeError:
                temp_path.replace(file_path)
            except FileNotFoundError:
                config_dir.mkdir(parents=True, exist_ok=True)
                os.replace(temp_path, file_path)

            logger.info(f"💾 Settings saved to {file_path} (Config Type: {self.config_type})")
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler beim Speichern der Einstellungen: {e}")

    def repair(self) -> List[str]:
        repairs = []
        default_instance = AdvancedSettings()

        for field_name, field_value in asdict(default_instance).items():
            if not hasattr(self, field_name) or getattr(self, field_name) is None:
                setattr(self, field_name, field_value)
                repairs.append(f"Added missing {field_name}")

        if hasattr(self, "ollama_model") and self.ollama_model == "llama3":
            self.ollama_model = "llama3.1:8b"
            repairs.append("Updated ollama_model from llama3 to llama3.1:8b")

        if not hasattr(self, "config") or self.config is None:
            self._recreate_config()
            repairs.append("Recreated config")

        self._apply_mode_overrides()
        self.chunk_duration = self._chunk_duration

        if repairs:
            logger.info(f"✅ Repairs made: {', '.join(repairs)}")
            self.save_to_file()
        else:
            logger.info("✅ No repairs needed")
        return repairs

    def validate(self) -> List[str]:
        issues = []

        if not (1 <= self.beam_size <= 20):
            issues.append(f"beam_size {self.beam_size} außerhalb 1-20")
        if not (0.0 <= self.temperature <= 2.0):
            issues.append(f"temperature {self.temperature} außerhalb 0.0-2.0")
        if not (100 <= self.max_memory_mb <= 16384):
            issues.append(f"max_memory_mb {self.max_memory_mb} außerhalb 100-16384")

        valid_config_types = ["default", "realtime", "high_accuracy", "youtube"]
        if self.config_type not in valid_config_types:
            issues.append(f"Ungültiger config_type '{self.config_type}'")

        if hasattr(self, 'config') and self.config:
            min_dur = self.config.MIN_CHUNK_DURATION
            max_dur = self.config.MAX_CHUNK_DURATION
            if not (min_dur <= self.chunk_duration <= max_dur):
                issues.append(
                    f"chunk_duration {self.chunk_duration}s außerhalb [{min_dur}, {max_dur}]"
                )

        if not (100 <= self.transcript_max_lines <= 5000):
            issues.append(
                f"transcript_max_lines {self.transcript_max_lines} außerhalb 100-5000"
            )
        if not (100 <= self.translation_max_lines <= 5000):
            issues.append(
                f"translation_max_lines {self.translation_max_lines} außerhalb 100-5000"
            )

        if self.translation_engine not in ("google", "ollama", "argos"):
            issues.append(f"translation_engine '{self.translation_engine}' ungültig")

        if not (0.5 <= self.duplicate_similarity_threshold <= 1.0):
            issues.append(
                f"duplicate_similarity_threshold {self.duplicate_similarity_threshold} außerhalb 0.5-1.0"
            )
        if not (0.0 <= self.min_confidence <= 1.0):
            issues.append(f"min_confidence {self.min_confidence} außerhalb 0.0-1.0")

        if not (0.0 <= self.vad_threshold <= 1.0):
            issues.append(f"vad_threshold {self.vad_threshold} außerhalb 0.0-1.0")
        if self.vad_min_speech_duration_ms < 0:
            issues.append("vad_min_speech_duration_ms negativ")
        if self.vad_min_silence_duration_ms < 0:
            issues.append("vad_min_silence_duration_ms negativ")

        if self.blacklist_mode not in ("word", "substring"):
            issues.append(f"blacklist_mode '{self.blacklist_mode}' ungültig")

        if self.tts_engine not in ("piper", "pyttsx3"):
            issues.append(f"tts_engine '{self.tts_engine}' ungültig")

        if not (1 <= self.best_of <= 20):
            issues.append(f"best_of {self.best_of} außerhalb 1-20")
        if not (0.0 <= self.patience <= 2.0):
            issues.append(f"patience {self.patience} außerhalb 0.0-2.0")
        if not (0.0 <= self.no_speech_threshold <= 1.0):
            issues.append(f"no_speech_threshold {self.no_speech_threshold} außerhalb 0.0-1.0")
        if not (-5.0 <= self.log_prob_threshold <= 0.0):
            issues.append(f"log_prob_threshold {self.log_prob_threshold} außerhalb -5.0-0.0")
        if not (1.0 <= self.compression_ratio_threshold <= 5.0):
            issues.append(f"compression_ratio_threshold {self.compression_ratio_threshold} außerhalb 1.0-5.0")

        return issues

    def set_config_type(self, config_type: str) -> bool:
        if config_type == self.config_type:
            return True
        valid_types = ["default", "realtime", "high_accuracy", "youtube"]
        if config_type not in valid_types:
            logger.warning(f"⚠️ Ungültiger config_type '{config_type}'")
            return False

        old_type = self.config_type
        self.config_type = config_type
        old_chunk = self.chunk_duration 
        self._recreate_config()
        self.chunk_duration = old_chunk
        self._apply_mode_overrides()
        logger.info(
            f"🔄 Config type geändert: {old_type} → {config_type}, neue CHUNK_DURATION: {self.config.CHUNK_DURATION}s"
        )
        return True

    def get_audio_filter(
        self, language: Optional[str] = None, profile: Optional[str] = None
    ) -> str:
        return self.config.get_audio_filter(language, profile or self.audio_profile)

    def get_youtube_headers(self, is_manifest: bool = False) -> Dict[str, str]:
        return self.config.get_youtube_headers(is_manifest)

    def get_platform_config(self, platform: Optional[str] = None) -> Dict[str, Any]:
        return self.config.get_platform_config(platform)

    def print_config_summary(self) -> None:
        lines = [
            "\n" + "=" * 60,
            "⚙️ SETTINGS CONFIGURATION",
            "=" * 60,
            "\n🤖 AI Model Parameters:",
            f"  • Beam Size: {self.beam_size}",
            f"  • Temperature: {self.temperature}",
            f"  • VAD Filter: {self.vad_filter}",
            f"  • VAD Threshold: {self.vad_threshold}",
            f"  • VAD Min Speech (ms): {self.vad_min_speech_duration_ms}",
            f"  • VAD Min Silence (ms): {self.vad_min_silence_duration_ms}",
            f"  • GPU Acceleration: {self.gpu_acceleration}",
            "\n🎵 Audio Configuration (from Config):",
            f"  • Sample Rate: {self.config.SAMPLE_RATE} Hz",
            f"  • Channels: {self.config.CHANNELS} ({'Mono' if self.config.CHANNELS == 1 else 'Stereo'})",
            f"  • Chunk Duration: {self.config.CHUNK_DURATION}s",
            f"  • Chunk Size: {self.config.CHUNK_SIZE_BYTES:,} bytes",
            f"  • Bytes/sec: {self.config.BYTES_PER_SECOND:,}",
            f"  • Audio Filter Profiles: {len(self.config.FILTER_PROFILES)}",
            f"  • Language Filters: {len(self.config.LANGUAGE_FILTERS)} languages",
            "\n⚡ Performance Settings:",
            f"  • Max Cache Size: {self.max_cache_size}",
            f"  • Max Memory: {self.max_memory_mb} MB",
            f"  • Auto Save Interval: {self.auto_save_interval}s",
            "\n🔧 Features:",
            f"  • Sentiment Analysis: {self.enable_sentiment_analysis}",
            f"  • Speaker Diarization: {self.enable_speaker_diarization}",
            f"  • Optimize Translations: {self.optimize_translations}",
            f"  • Noise Reduction: {self.enable_noise_reduction}",
            f"  • Duplicate Check: {self.enable_duplicate_check}",
            "\n🖥️ GUI Display:",
            f"  • Transcript Max Lines: {self.transcript_max_lines}",
            f"  • Translation Max Lines: {self.translation_max_lines}",
            "\n🌐 Translation Engine:",
            f"  • Engine: {self.translation_engine}",
        ]
        if self.translation_engine == "ollama":
            lines.append(f"  • Ollama Model: {self.ollama_model}")
            lines.append(f"  • Ollama Host: {self.ollama_host}")
        lines.extend(
            [
                "\n🗾 Asian Language Mode:",
                f"  • Active: {self.asian_mode}",
                "\n🎯 Precision Mode:",
                f"  • Active: {self.precision_mode}",
                "\n🎛️ Audio Profile:",
                f"  • Profile: {self.audio_profile}",
                "\n⚙️ Advanced Options:",
                f"  • Adaptive Chunk: {self.adaptive_chunk}",
                f"  • Duplicate Similarity Threshold: {self.duplicate_similarity_threshold:.2f}",
                f"  • Adaptive Chunk Low Words: {self.adaptive_chunk_low_words}",
                f"  • Adaptive Chunk High Words: {self.adaptive_chunk_high_words}",
                f"\n🚫 Blacklist Entries: {len(self.blacklist)}",
                f"\n🔤 Blacklist Mode: {self.blacklist_mode}",
                f"\n🔑 Hotwords: {self.hotwords}",
                f"\n⚙️ Executor Workers: Transcribe={self.transcription_workers}, Translate={self.translation_workers}",
                f"\n🎯 Config Type: {self.config_type.upper()}",
                f"\n🔊 TTS Engine: {self.tts_engine}",
                "\n🔧 Erweiterte Whisper-Parameter:",
                f"  • Best of: {self.best_of}",
                f"  • Patience: {self.patience}",
                f"  • No Speech Threshold: {self.no_speech_threshold}",
                f"  • Log Prob Threshold: {self.log_prob_threshold}",
                f"  • Compression Ratio: {self.compression_ratio_threshold}",
                f"  • Condition on previous text: {self.condition_on_previous_text}",
                f"  • Suppress Tokens: {self.suppress_tokens}",
                f"\n🔁 VAD Fallback Enabled: {self.vad_fallback_enabled}",
            ]
        )
        issues = self.validate()
        if issues:
            lines.append("\n⚠️ Validation Issues:")
            for issue in issues:
                lines.append(f"  • {issue}")
        else:
            lines.append("\n✅ All settings valid")
        lines.append("=" * 60)
        logger.info("\n".join(lines))

    def __repr__(self) -> str:
        return (
            f"AdvancedSettings(type={self.config_type}, "
            f"beam_size={self.beam_size}, "
            f"chunk={self.config.CHUNK_DURATION}s/{self.config.CHUNK_SIZE_BYTES:,}B, "
            f"gpu={self.gpu_acceleration}, "
            f"trans_engine={self.translation_engine}, "
            f"asian_mode={self.asian_mode}, "
            f"precision_mode={self.precision_mode}, "
            f"tts_engine={self.tts_engine})"
        )


# -----------------------------------------------------------------------------
# AppSettings
# -----------------------------------------------------------------------------
@dataclass
class AppSettings:
    last_url: str = ""
    default_model: str = "large-v3"
    default_language: str = "auto"
    layout_mode: str = "horizontal"
    recent_urls: List[str] = field(default_factory=list)
    enable_plugins: bool = True
    export_format: str = "srt"
    auto_save_on_completion: bool = False
    theme: str = "dark"
    use_browser_cookies: bool = True
    cookies_notice_shown: bool = False

    @classmethod
    def load_from_file(cls, filename: str = "dragon_settings.json") -> "AppSettings":
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            file_path = config_dir / filename
            if not file_path.exists():
                return cls()

            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            valid_fields = {f.name for f in fields(cls)}
            filtered = {k: v for k, v in data.items() if k in valid_fields}

            if "recent_urls" in filtered and not isinstance(
                filtered["recent_urls"], list
            ):
                filtered["recent_urls"] = []

            if "default_model" in filtered:
                model = filtered["default_model"]
                if model not in WHISPER_MODELS and model != "large-v3":
                    filtered["default_model"] = "large-v3"

            return cls(**filtered)

        except (json.JSONDecodeError, OSError, PermissionError) as e:
            logger.warning(f"⚠️ Konnte Einstellungen nicht laden: {e}")
            return cls()
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Unerwarteter Fehler beim Laden der Einstellungen: {e}")
            return cls()

    def save_to_file(self, filename: str = "dragon_settings.json") -> bool:
        try:
            config_dir = PlatformUtils.get_platform_config_dir()
            config_dir.mkdir(parents=True, exist_ok=True)

            file_path = config_dir / filename
            temp_path = file_path.with_suffix(".tmp")

            data = {
                "last_url": self.last_url,
                "default_model": self.default_model,
                "default_language": self.default_language,
                "layout_mode": self.layout_mode,
                "recent_urls": self.recent_urls[:10],
                "enable_plugins": self.enable_plugins,
                "export_format": self.export_format,
                "auto_save_on_completion": self.auto_save_on_completion,
                "theme": self.theme,
                "use_browser_cookies": self.use_browser_cookies,
                "cookies_notice_shown": self.cookies_notice_shown,
            }

            with open(temp_path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            temp_path.replace(file_path)
            logger.debug(f"💾 Einstellungen gespeichert: {file_path}")
            return True

        except (OSError, PermissionError) as e:
            logger.warning(f"⚠️ Konnte Einstellungen nicht speichern: {e}")
            return False
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.error(f"❌ Fehler beim Speichern der Einstellungen: {e}")
            return False

    def add_recent_url(self, url: str) -> None:
        if not url:
            return
        if url in self.recent_urls:
            self.recent_urls.remove(url)
        self.recent_urls.insert(0, url)
        self.recent_urls = self.recent_urls[:10]
        self.save_to_file()

    def update_last_url(self, url: str) -> None:
        if url != self.last_url:
            self.last_url = url
            self.save_to_file()

    def validate(self) -> List[str]:
        issues = []

        if self.default_model not in WHISPER_MODELS and self.default_model != "large-v3":
            issues.append(f"Ungültiges Modell: {self.default_model}")

        valid_languages = [code for code, name in SUPPORTED_LANGUAGES.items()]
        if self.default_language not in valid_languages:
            issues.append(f"Ungültige Sprache: {self.default_language}")

        if self.layout_mode not in ("vertical", "horizontal"):
            issues.append(f"Ungültiges Layout: {self.layout_mode}")

        if self.export_format not in ("txt", "srt", "vtt", "json", "docx"):
            issues.append(f"Ungültiges Exportformat: {self.export_format}")

        if self.theme not in ("dark", "light", "pastel", "system", "highcontrast"):
            issues.append(f"Ungültiges Theme: {self.theme}")

        return issues

    def repair(self) -> List[str]:
        repairs = []
        default = AppSettings()

        if self.default_model not in WHISPER_MODELS and self.default_model != "large-v3":
            old = self.default_model
            self.default_model = default.default_model
            repairs.append(f"Modell {old} -> {self.default_model}")

        valid_languages = [code for code, name in SUPPORTED_LANGUAGES.items()]
        if self.default_language not in valid_languages:
            old = self.default_language
            self.default_language = default.default_language
            repairs.append(f"Sprache {old} -> {self.default_language}")

        if self.layout_mode not in ("vertical", "horizontal"):
            old = self.layout_mode
            self.layout_mode = default.layout_mode
            repairs.append(f"Layout {old} -> {self.layout_mode}")

        if self.export_format not in ("txt", "srt", "vtt", "json", "docx"):
            old = self.export_format
            self.export_format = default.export_format
            repairs.append(f"Exportformat {old} -> {self.export_format}")

        if self.theme not in ("dark", "light", "pastel", "system", "highcontrast"):
            old = self.theme
            self.theme = default.theme
            repairs.append(f"Theme {old} -> {self.theme}")

        if repairs:
            logger.info(f"✅ AppSettings repariert: {', '.join(repairs)}")
            self.save_to_file()
        return repairs

    def __repr__(self) -> str:
        return (
            f"AppSettings(model={self.default_model}, "
            f"lang={self.default_language}, layout={self.layout_mode}, "
            f"theme={self.theme}, cookies={self.use_browser_cookies})"
        )


# =============================================================================
# 14. HILFSFUNKTIONEN FÜR MAIN
# =============================================================================

_original_console_mode: Optional[int] = None
_original_codepage: Optional[int] = None


def _save_console_state() -> None:
    global _original_console_mode, _original_codepage
    if not IS_WINDOWS:
        return
    try:
        import ctypes

        kernel32 = ctypes.windll.kernel32
        handle = kernel32.GetStdHandle(-11)
        mode = ctypes.c_ulong()
        if kernel32.GetConsoleMode(handle, ctypes.byref(mode)):
            _original_console_mode = mode.value
        _original_codepage = kernel32.GetConsoleOutputCP()
    except Exception:
        pass


def _restore_console_state() -> None:
    if not IS_WINDOWS:
        return
    try:
        import ctypes

        kernel32 = ctypes.windll.kernel32
        if _original_console_mode is not None:
            handle = kernel32.GetStdHandle(-11)
            kernel32.SetConsoleMode(handle, _original_console_mode)
        if _original_codepage is not None:
            kernel32.SetConsoleOutputCP(_original_codepage)
            kernel32.SetConsoleCP(_original_codepage)
    except Exception:
        pass


_save_console_state()
atexit.register(_restore_console_state)


def _setup_windows_console() -> None:
    if not IS_WINDOWS:
        return
    try:
        import codecs

        sys.stdout = codecs.getwriter("utf-8")(sys.stdout.buffer)
        import ctypes

        kernel32 = ctypes.windll.kernel32
        kernel32.SetConsoleOutputCP(65001)
        kernel32.SetConsoleCP(65001)
        handle = kernel32.GetStdHandle(-11)
        mode = ctypes.c_ulong()
        if kernel32.GetConsoleMode(handle, ctypes.byref(mode)):
            ENABLE_VIRTUAL_TERMINAL_PROCESSING = 0x0004
            new_mode = mode.value | ENABLE_VIRTUAL_TERMINAL_PROCESSING
            kernel32.SetConsoleMode(handle, new_mode)
    except Exception:
        pass


def check_platform_compatibility() -> List[str]:
    issues: List[str] = []
    if sys.version_info < (3, 8):
        issues.append(
            f"Python 3.8+ required (you have {sys.version_info.major}.{sys.version_info.minor})"
        )
    if IS_WINDOWS:
        try:
            import platform

            win_ver = platform.version()
            major_ver = int(win_ver.split(".")[0]) if "." in win_ver else 0
            if major_ver < 10:
                issues.append("Windows 10+ recommended for best experience")
        except Exception:
            pass
    return issues


def _print_help() -> None:
    print("🐉 Dragon Whisperer - Ultimate Stream Transcription & Translation")
    print("=" * 60)
    print("\nUsage:")
    print("  python dragon_whisperer.py [options]")
    print("\nOptions:")
    print("  --quiet, -q    Quiet mode (minimal output)")
    print("  --debug[=N]    Debug mode with level (1-3). Examples: --debug, --debug=2")
    print("  --debug=COMP   Component-specific debug (e.g., --debug=network,vad)")
    print("  --check        System compatibility check")
    print("  --help, -h     Show this help")
    print("  --version, -v  Show version")
    print("\nExamples:")
    print("  Normal use:   python dragon_whisperer.py")
    print("  Quiet mode:   python dragon_whisperer.py --quiet")
    print("  Debug level2: python dragon_whisperer.py --debug=2")
    print("  VAD debug:    python dragon_whisperer.py --debug=vad")
    print("  System check: python dragon_whisperer.py --check")
    print("\nFeatures:")
    print("  • Live stream transcription (YouTube, Twitch, etc.)")
    print("  • Real-time translation to 50+ languages")
    print("  • Subtitle export (SRT, VTT)")
    print("  • Batch processing")
    print("  • Dark mode GUI")


def _run_system_check() -> int:
    issues = check_platform_compatibility()
    print("🔍 Dragon Whisperer - System Compatibility Check")
    print("=" * 50)
    if issues:
        print("⚠️  Issues found:")
        for issue in issues:
            print(f"  • {issue}")
    else:
        print("✅ No compatibility issues")
    print("\n📦 Dependency Check:")
    print(f"  FFmpeg:        {'✅' if shutil.which('ffmpeg') else '❌'}")
    print(f"  yt-dlp:        {'✅' if shutil.which('yt-dlp') else '❌'}")
    print(f"  Tkinter:       {'✅' if GUI_AVAILABLE else '❌'}")
    print(f"  faster-whisper:{'✅' if WHISPER_AVAILABLE else '❌'}")
    print(f"  NumPy:         {'✅' if NUMPY_AVAILABLE else '❌'}")
    print(f"  PyTorch:       {'✅' if TORCH_AVAILABLE else '❌'}")
    print(f"  deep-translator:{'✅' if TRANSLATOR_AVAILABLE else '❌'}")
    print(f"  SciPy:         {'✅' if SCIPY_AVAILABLE else '❌'}")
    print(f"  psutil:        {'✅' if 'psutil' in sys.modules else '❌'}")
    print("\n💻 System Info:")
    print(f"  Platform: {SYSTEM}")
    print(f"  Architecture: {'ARM' if IS_ARM else 'x86'}")
    print(f"  Python: {sys.version.split()[0]}")
    if IS_WINDOWS:
        try:
            import platform

            print(f"  Windows: {platform.version()}")
        except Exception:
            pass
    return 0 if not issues else 1


def _show_user_error(message: str) -> None:
    if IS_WINDOWS and not sys.stdin.isatty():
        try:
            import ctypes

            ctypes.windll.user32.MessageBoxW(
                0,
                f"Dragon Whisperer - Setup Required\n\n{message}\n\n"
                "Please install missing components and try again.",
                "Setup Error",
                0x10,
            )
        except Exception:
            pass
    else:
        if "Tkinter" in message:
            print("\n💡 INSTALLATION HELP:")
            print("  pip install tk")
        elif "FFmpeg" in message:
            print("\n💡 INSTALLATION HELP:")
            if IS_WINDOWS:
                print("  Download from: https://ffmpeg.org/download.html")
            elif IS_MACOS:
                print("  brew install ffmpeg")
            else:
                print("  sudo apt install ffmpeg")
        else:
            print(f"\n❌ {message}")


def _show_critical_error(message: str) -> None:
    if IS_WINDOWS and not sys.stdin.isatty():
        try:
            import ctypes

            short_msg = message[:200] + "..." if len(message) > 200 else message
            ctypes.windll.user32.MessageBoxW(
                0,
                f"Dragon Whisperer - Critical Error\n\n{short_msg}\n\n"
                "Please check the console for details.\n"
                "Try running with --debug flag for more information.",
                "Critical Error",
                0x10,
            )
        except Exception:
            pass
    else:
        print(f"\n💥 {message}")


def print_system_info_debug3():
    import platform

    print("\n" + "=" * 60)
    print("🐉 DEBUG LEVEL 3 - SYSTEM INFORMATION")
    print("=" * 60)
    print(f"Platform: {platform.platform()}")
    print(f"Python: {sys.version}")
    print(f"CPU count: {os.cpu_count()}")
    try:
        soft, hard = resource.getrlimit(resource.RLIMIT_NOFILE)
        print(f"RLIMIT_NOFILE: soft={soft}, hard={hard}")
    except Exception:
        pass
    print("\n📦 Environment variables set by script:")
    for key in os.environ:
        if key.startswith(("FFMPEG_", "TORCH_", "PYTORCH_", "AV_", "OPENCV_")):
            print(f"  {key}={os.environ[key]}")
    print("\n📚 Library versions:")
    libs = [
        ("torch", TORCH_AVAILABLE),
        ("faster_whisper", FASTER_WHISPER_AVAILABLE),
        ("whisper", OPENAI_WHISPER_AVAILABLE),
        ("numpy", NUMPY_AVAILABLE),
        ("scipy", SCIPY_AVAILABLE),
        ("psutil", importlib.util.find_spec("psutil") is not None),
        ("requests", OLLAMA_AVAILABLE),
        ("deep_translator", TRANSLATOR_AVAILABLE),
        ("pynvml", importlib.util.find_spec("pynvml") is not None),
    ]
    for name, avail in libs:
        if avail:
            try:
                mod = __import__(name)
                ver = getattr(mod, "__version__", "unknown")
                print(f"  {name}: {ver}")
            except Exception:
                print(f"  {name}: available (version unknown)")
        else:
            print(f"  {name}: not available")
    print("\n🎮 GPU info:")
    if TORCH_AVAILABLE:
        torch = FastLazyLoader.load("torch")
        if torch.cuda.is_available():
            print(f"  CUDA available: {torch.version.cuda}")
            print(f"  Device: {torch.cuda.get_device_name(0)}")
            print(
                f"  Total VRAM: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.2f} GB"
            )
        else:
            print("  CUDA not available")
    print("=" * 60 + "\n")


# =============================================================================
# 15. MAIN
# =============================================================================

def global_exception_handler(exc_type, exc_value, exc_traceback):
    if issubclass(exc_type, KeyboardInterrupt):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logger.critical("Unhandled exception", exc_info=(exc_type, exc_value, exc_traceback))
    try:
        import tkinter.messagebox as mb
        mb.showerror("Dragon Whisperer - Unerwarteter Fehler",
                     f"Ein unerwarteter Fehler ist aufgetreten:\n\n{exc_value}\n\n"
                     "Details in der Konsole.")
    except Exception:
        pass
    sys.__excepthook__(exc_type, exc_value, exc_traceback)

sys.excepthook = global_exception_handler

def main() -> int:
    warnings.filterwarnings("ignore")
    if IS_WINDOWS:
        _setup_windows_console()
    cli_args = {
        "debug": "--debug" in sys.argv,
        "quiet": "--quiet" in sys.argv or "-q" in sys.argv,
        "check": "--check" in sys.argv,
        "help": "--help" in sys.argv or "-h" in sys.argv,
        "version": "--version" in sys.argv or "-v" in sys.argv,
    }
    if cli_args["help"]:
        _print_help()
        return 0
    if cli_args["version"]:
        print("🐉 Dragon Whisperer v2.1 - überarbeitet")
        print(f"Platform: {SYSTEM} {'ARM' if IS_ARM else 'x86'}")
        return 0
    debug_level = DEBUG_LEVEL
    if cli_args["check"]:
        return _run_system_check()
    if DEBUG_LEVEL >= 3:
        print_system_info_debug3()
    app: Optional[DragonWhispererGUI] = None
    exit_code = 0
    try:
        logger.info("🐉 Dragon Whisperer starting...")
        logger.debug("🔍 Checking dependencies...")
        try:
            PlatformUtils.check_platform_dependencies()
        except RuntimeError as e:
            _show_user_error(str(e))
            return 1
        logger.debug("✅ Dependencies OK")
        if not GUI_AVAILABLE:
            raise RuntimeError(
                "Tkinter/GUI not available. Install with: pip install tk"
            )
        logger.debug("⚡ Setting up signal handlers...")
        SignalHandler.setup(verbose=False, silent=True, max_cleanup_time=10.0, atexit_enabled=False)
        SignalHandler.register_cleanup(
            lambda: _EXECUTOR.shutdown(wait=False),
            name="GlobalExecutorShutdown",
            priority=ShutdownPriority.LOW,
        )
        logger.debug("🖥️ Initializing GUI...")
        app = DragonWhispererGUI()
        if debug_level >= 1 and not cli_args["quiet"]:
            print("\n" + "=" * 50)
            print("🐉 DRAGON WHISPERER READY")
            print("=" * 50)
            print(f"Platform: {SYSTEM} {'ARM' if IS_ARM else 'x86'}")
            print(f"Python: {sys.version.split()[0]}")
            print(f"Working Dir: {os.getcwd()}")
            if hasattr(app, "transcription_engine"):
                current_model = app.transcription_engine.get_current_model()
                print(f"Model: {current_model if current_model else 'Not loaded'}")
            print(f"Layout: {getattr(app, 'layout_mode', 'vertical')}")
            print("=" * 50 + "\n")
        logger.info("🚀 Starting main loop...")
        app.run()
        logger.info("✅ Application closed normally")
    except KeyboardInterrupt:
        logger.info("\n🛑 Interrupted by user")
        exit_code = 0
    except RuntimeError as e:
        error_msg = str(e)
        logger.error(f"❌ {error_msg}")
        _show_user_error(error_msg)
        exit_code = 1
    except Exception as e:
        if isinstance(e, (KeyboardInterrupt, SystemExit)):
            raise
        error_msg = str(e)
        logger.error(f"💥 Unexpected error: {error_msg}")
        if debug_level >= 2:
            import traceback

            traceback.print_exc()
        _show_critical_error(error_msg)
        exit_code = 2
    finally:
        logger.debug("🧹 Final minimal cleanup...")
        if app is not None:
            try:
                if hasattr(app, "_safe_stop_all_processes"):
                    app._safe_stop_all_processes()
                if app.app_context:
                    app.app_context.cache_manager.transcription_cache.clear()
                    app.app_context.cache_manager.translation_cache.clear()
                    app.app_context.cache_manager.audio_cache.clear()
                app = None
            except Exception as e:
                if isinstance(e, (KeyboardInterrupt, SystemExit)):
                    raise
                logger.warning(f"⚠️ App cleanup error: {e}")
        try:
            _EXECUTOR.shutdown(wait=True)
        except Exception as e:
            if isinstance(e, (KeyboardInterrupt, SystemExit)):
                raise
            logger.warning(f"⚠️ Executor shutdown error: {e}")
        if IS_WINDOWS:
            _restore_console_state()
        gc.collect()
        logger.debug("✅ Shutdown complete")
    return exit_code


if __name__ == "__main__":
    try:
        if "__file__" in globals():
            script_dir = os.path.dirname(os.path.abspath(__file__))
            if os.getcwd() != script_dir:
                os.chdir(script_dir)
                if DEBUG_LEVEL >= 1:
                    print(f"📁 Working directory set to: {script_dir}")
        exit_code = main()
        sys.exit(exit_code)
    except KeyboardInterrupt:
        print("\n🛑 Program interrupted by user (KeyboardInterrupt)")
        sys.exit(0)
    except SystemExit as e:
        raise e
    except Exception as e:
        if isinstance(e, (KeyboardInterrupt, SystemExit)):
            raise
        error_type = type(e).__name__
        error_msg = str(e)
        print(f"\n💥 FATAL ERROR in main guard: {error_type}: {error_msg}")
        if DEBUG_LEVEL >= 2:
            import traceback

            traceback.print_exc()
            print("\n🔧 Debug Info:")
            print(f"  Python: {sys.version}")
            print(f"  Platform: {sys.platform}")
            print(f"  Executable: {sys.executable}")
            print(f"  CWD: {os.getcwd()}")
            print(f"  Script: {__file__ if '__file__' in globals() else 'Unknown'}")
        sys.exit(99)
